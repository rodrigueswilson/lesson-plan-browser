
import sys
import os
from docx import Document
from pathlib import Path

def analyze_template(path):
    doc = Document(path)
    print(f"Analysis of: {Path(path).name}\n")
    
    # 1. Look for Metadata Table
    print("--- PAGE STRUCTURE ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns")
        for r_idx, row in enumerate(table.rows):
            # Just print the first few character of each cell to understand the grid
            cells = [cell.text.strip().replace('\n', ' ')[:30] for cell in row.cells]
            print(f"Row {r_idx}: | {' | '.join(cells)} |")

    # 2. Look for explicit headers
    print("\n--- STANDALONE HEADERS ---")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') or para.text.isupper():
            if para.text.strip():
                print(f"Header: {para.text.strip()}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python analyze_template.py <path_to_docx>")
        sys.exit(1)
    analyze_template(sys.argv[1])
