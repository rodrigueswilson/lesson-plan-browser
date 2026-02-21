"""
Compare primary teacher lesson plans with combined_originals.docx.

This script identifies differences between:
1. Primary teacher DOCX files (source)
2. Database stored original lesson plans
3. Combined originals DOCX output

Purpose: Understand why combined_originals.docx contains different slots/content
than what's in the primary teacher files.

Usage:
    python scripts/compare_primary_vs_combined_originals.py --week-folder "F:/rodri/Documents/OneDrive/AS/Lesson Plan/25 W51" --week-of "12/16-12/20" --user-id "04fe8898-cb89-4a73-affb-64a97a98f820"

    Or use defaults (W51 example):
    python scripts/compare_primary_vs_combined_originals.py

Arguments:
    --week-folder: Path to week folder (default: F:/rodri/Documents/OneDrive/AS/Lesson Plan/25 W51)
    --week-of: Week date range in MM/DD-MM/DD format (default: 12/16-12/20)
    --user-id: User ID to query database (default: auto-detect Wilson Rodrigues)
    --output: Optional path for JSON report (default: week_folder/comparison_report_*.json)

Output:
    - Console report showing differences
    - JSON report with detailed comparison data
"""

import asyncio
import json
import sys
from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.config import settings
from backend.database import get_db
from backend.file_manager import get_file_manager
from tools.docx_parser import DOCXParser


class ComparisonReport:
    """Stores comparison results and generates reports."""

    def __init__(self, week_folder: Path, week_of: str, user_id: str):
        self.week_folder = week_folder
        self.week_of = week_of
        self.user_id = user_id
        self.primary_files: Dict[str, Dict[str, Any]] = {}
        self.db_plans: List[Any] = []
        self.combined_originals_path: Optional[Path] = None
        self.combined_originals_slots: Dict[int, Dict[str, Any]] = {}
        self.differences: List[Dict[str, Any]] = []

    def add_difference(
        self,
        category: str,
        slot_number: Optional[int],
        subject: Optional[str],
        message: str,
        details: Optional[Dict[str, Any]] = None,
    ):
        """Add a difference to the report."""
        self.differences.append(
            {
                "category": category,
                "slot_number": slot_number,
                "subject": subject,
                "message": message,
                "details": details or {},
            }
        )

    def print_report(self):
        """Print a formatted comparison report."""
        print("\n" + "=" * 80)
        print(f"COMPARISON REPORT: Week {self.week_of}")
        print(f"Week Folder: {self.week_folder}")
        print("=" * 80)

        # 1. Primary Files Summary
        print("\n[1] PRIMARY TEACHER FILES FOUND:")
        print("-" * 80)
        if not self.primary_files:
            print("  No primary teacher files found in week folder")
        else:
            for file_path, file_info in self.primary_files.items():
                print(f"  File: {file_info['name']}")
                print(f"    Path: {file_path}")
                print(f"    Slots detected: {file_info.get('slots', [])}")
                print(f"    Subjects: {file_info.get('subjects', [])}")
                print()

        # 2. Database Plans Summary
        print("\n[2] DATABASE ORIGINAL LESSON PLANS:")
        print("-" * 80)
        if not self.db_plans:
            print("  No original lesson plans found in database")
        else:
            print(f"  Total plans in database: {len(self.db_plans)}")
            for plan in sorted(self.db_plans, key=lambda x: x.slot_number):
                print(f"  Slot {plan.slot_number}: {plan.subject} (Grade {plan.grade})")
                print(f"    Source file: {plan.source_file_name}")
                print(f"    Primary teacher: {plan.primary_teacher_name or 'N/A'}")
                print(f"    Available days: {plan.available_days or []}")
                print(f"    Status: {plan.status}")
                print()

        # 3. Combined Originals Summary
        print("\n[3] COMBINED ORIGINALS.DOCX:")
        print("-" * 80)
        if not self.combined_originals_path:
            print("  No combined_originals.docx file found")
        else:
            print(f"  File: {self.combined_originals_path.name}")
            print(f"  Path: {self.combined_originals_path}")
            print(f"  Slots found in document: {sorted(self.combined_originals_slots.keys())}")
            for slot_num, slot_info in sorted(self.combined_originals_slots.items()):
                print(f"    Slot {slot_num}: {slot_info.get('subject', 'Unknown')}")
            print()

        # 4. Differences Analysis
        print("\n[4] DIFFERENCES ANALYSIS:")
        print("-" * 80)
        if not self.differences:
            print("  No differences found - everything matches!")
        else:
            # Group by category
            by_category = defaultdict(list)
            for diff in self.differences:
                by_category[diff["category"]].append(diff)

            for category, diffs in by_category.items():
                print(f"\n  {category}:")
                for diff in diffs:
                    slot_info = f"Slot {diff['slot_number']}" if diff['slot_number'] else "General"
                    subject_info = f" ({diff['subject']})" if diff['subject'] else ""
                    print(f"    - {slot_info}{subject_info}: {diff['message']}")
                    if diff['details']:
                        for key, value in diff['details'].items():
                            print(f"        {key}: {value}")

        # 5. Slot Mapping
        print("\n[5] SLOT MAPPING:")
        print("-" * 80)
        self._print_slot_mapping()

        print("\n" + "=" * 80)

    def _print_slot_mapping(self):
        """Print a mapping table showing where each slot appears."""
        # Collect all slot numbers
        primary_slots = set()
        for file_info in self.primary_files.values():
            primary_slots.update(file_info.get("slots", []))

        db_slots = {plan.slot_number for plan in self.db_plans}
        combined_slots = set(self.combined_originals_slots.keys())

        all_slots = primary_slots | db_slots | combined_slots

        if not all_slots:
            print("  No slots found")
            return

        print(f"{'Slot':<6} {'Primary Files':<20} {'Database':<15} {'Combined DOCX':<15}")
        print("-" * 60)

        for slot in sorted(all_slots):
            in_primary = "YES" if slot in primary_slots else "NO"
            in_db = "YES" if slot in db_slots else "NO"
            in_combined = "YES" if slot in combined_slots else "NO"

            # Get subject info
            db_subject = next(
                (p.subject for p in self.db_plans if p.slot_number == slot), "N/A"
            )
            combined_subject = self.combined_originals_slots.get(slot, {}).get(
                "subject", "N/A"
            )

            print(f"{slot:<6} {in_primary:<20} {in_db:<15} {in_combined:<15}")
            if db_subject != "N/A" or combined_subject != "N/A":
                print(f"       DB: {db_subject:<18} Combined: {combined_subject}")

    def save_json_report(self, output_path: Path):
        """Save detailed JSON report."""
        report_data = {
            "week_of": self.week_of,
            "week_folder": str(self.week_folder),
            "user_id": self.user_id,
            "primary_files": {
                path: {
                    "name": info["name"],
                    "slots": info.get("slots", []),
                    "subjects": info.get("subjects", []),
                }
                for path, info in self.primary_files.items()
            },
            "database_plans": [
                {
                    "id": plan.id,
                    "slot_number": plan.slot_number,
                    "subject": plan.subject,
                    "grade": plan.grade,
                    "source_file_name": plan.source_file_name,
                    "source_file_path": plan.source_file_path,
                    "primary_teacher_name": plan.primary_teacher_name,
                    "available_days": plan.available_days,
                    "status": plan.status,
                }
                for plan in self.db_plans
            ],
            "combined_originals": {
                "path": str(self.combined_originals_path) if self.combined_originals_path else None,
                "slots": {
                    slot_num: {
                        "subject": info.get("subject", "Unknown"),
                        "grade": info.get("grade", "Unknown"),
                        "teacher_name": info.get("primary_teacher_name", "Unknown"),
                        "has_content": info.get("has_content", False),
                        "db_source_file": info.get("db_source_file"),
                        "db_primary_teacher": info.get("db_primary_teacher"),
                        "sample_content": info.get("sample_content", []),
                    }
                    for slot_num, info in self.combined_originals_slots.items()
                },
            },
            "differences": self.differences,
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(report_data, f, indent=2, default=str)

        print(f"\nDetailed JSON report saved to: {output_path}")


def find_primary_teacher_files(week_folder: Path) -> Dict[str, Dict[str, Any]]:
    """Find and analyze primary teacher DOCX files in week folder."""
    primary_files = {}
    file_manager = get_file_manager()

    # List all DOCX files in week folder
    docx_files = list(week_folder.glob("*.docx"))

    # Exclude combined_originals and output files
    excluded_patterns = [
        "combined_originals",
        "Lesson_plan_",
        "_output.docx",
    ]

    for docx_file in docx_files:
        # Skip excluded files
        if any(pattern in docx_file.name for pattern in excluded_patterns):
            continue

        try:
            parser = DOCXParser(str(docx_file))
            subjects = parser.list_available_subjects()

            # Try to detect slots by analyzing tables
            # Each slot typically has 2 tables (metadata + daily plan)
            table_count = len(parser.doc.tables)
            # Check if last table is signature
            has_signature = False
            if parser.doc.tables:
                last_table = parser.doc.tables[-1]
                if last_table.rows and last_table.rows[0].cells:
                    first_cell = last_table.rows[0].cells[0].text.strip().lower()
                    if "signature" in first_cell:
                        has_signature = True

            if has_signature:
                estimated_slots = (table_count - 1) // 2
            else:
                estimated_slots = table_count // 2

            primary_files[str(docx_file)] = {
                "name": docx_file.name,
                "path": str(docx_file),
                "subjects": subjects,
                "slots": list(range(1, estimated_slots + 1)) if estimated_slots > 0 else [],
                "table_count": table_count,
                "has_signature": has_signature,
            }
        except Exception as e:
            print(f"Warning: Could not parse {docx_file.name}: {e}")
            primary_files[str(docx_file)] = {
                "name": docx_file.name,
                "path": str(docx_file),
                "error": str(e),
            }

    return primary_files


def find_combined_originals(week_folder: Path) -> Optional[Path]:
    """Find the combined_originals.docx file in the week folder."""
    originals_dir = week_folder / "originals"
    if not originals_dir.exists():
        return None

    # Look for combined_originals files (most recent)
    combined_files = list(originals_dir.glob("combined_originals_*.docx"))
    if not combined_files:
        return None

    # Return most recently modified
    return max(combined_files, key=lambda p: p.stat().st_mtime)


def extract_slots_from_combined_docx(docx_path: Path) -> Dict[int, Dict[str, Any]]:
    """Extract slot information and content from combined_originals.docx.
    
    The combined docx is created by merging individual slot documents.
    Each slot should have 2 tables: metadata table + daily plans table.
    We can detect slots by counting table pairs and extracting subject from metadata.
    """
    slots = {}
    try:
        doc = Document(str(docx_path))
        import re

        # Method 1: Analyze table structure
        # Each slot has 2 tables (metadata + daily plans)
        # Count table pairs to estimate slots
        table_count = len(doc.tables)
        estimated_slots = table_count // 2

        # Method 2: Extract from metadata tables (first table of each pair)
        # Metadata table typically has: Name, Grade, Subject, Week, Homeroom
        for slot_num in range(1, estimated_slots + 1):
            metadata_table_idx = (slot_num - 1) * 2
            if metadata_table_idx < table_count:
                metadata_table = doc.tables[metadata_table_idx]
                daily_plans_table = doc.tables[metadata_table_idx + 1] if (metadata_table_idx + 1) < table_count else None
                
                # Extract subject from metadata table
                subject = "Unknown"
                grade = "Unknown"
                teacher_name = "Unknown"
                primary_teacher_name = "Unknown"
                
                # Look through table cells for subject, grade, and teacher
                for row in metadata_table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # Look for subject patterns (including ELA/SS)
                        subject_match = re.search(r'\b(ELA/SS|ELA|Math|Science|Social Studies|SS)\b', cell_text, re.IGNORECASE)
                        if subject_match:
                            subject = subject_match.group(1)
                        # Look for grade patterns
                        grade_match = re.search(r'\bGrade\s*(\d+[A-Z]?)\b', cell_text, re.IGNORECASE)
                        if grade_match:
                            grade = grade_match.group(1)
                        # Look for teacher/name patterns
                        if re.search(r'\b(Name|Teacher):\s*(.+)', cell_text, re.IGNORECASE):
                            name_match = re.search(r'\b(Name|Teacher):\s*(.+)', cell_text, re.IGNORECASE)
                            if name_match:
                                teacher_name = name_match.group(2).strip()
                                primary_teacher_name = teacher_name

                # Extract sample content from daily plans table (first few rows)
                sample_content = []
                if daily_plans_table:
                    for row_idx, row in enumerate(daily_plans_table.rows[:5]):  # First 5 rows
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            sample_content.append(row_text[:100])  # First 100 chars

                slots[slot_num] = {
                    "subject": subject,
                    "grade": grade,
                    "teacher_name": teacher_name,
                    "primary_teacher_name": primary_teacher_name,
                    "has_content": True,
                    "table_index": metadata_table_idx,
                    "sample_content": sample_content[:3],  # First 3 rows of sample
                }

        # Method 3: Also check paragraphs for slot indicators (backup)
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            slot_match = re.search(r"Slot\s*(\d+)", text, re.IGNORECASE)
            if slot_match:
                slot_num = int(slot_match.group(1))
                if slot_num not in slots:
                    slots[slot_num] = {
                        "subject": "Unknown",
                        "has_content": True,
                    }

    except Exception as e:
        print(f"Warning: Error extracting slots from combined_originals.docx: {e}")
        import traceback
        traceback.print_exc()

    return slots


def compare_slots(
    report: ComparisonReport,
    primary_files: Dict[str, Dict[str, Any]],
    db_plans: List[Any],
    combined_slots: Dict[int, Dict[str, Any]],
):
    """Compare slots across all three sources and identify differences."""

    # Collect all slot numbers
    primary_slots_by_file: Dict[str, Set[int]] = {}
    for file_path, file_info in primary_files.items():
        primary_slots_by_file[file_path] = set(file_info.get("slots", []))

    all_primary_slots = set()
    for slots in primary_slots_by_file.values():
        all_primary_slots.update(slots)

    db_slots = {plan.slot_number for plan in db_plans}
    combined_slots_set = set(combined_slots.keys())

    # 1. Slots in primary files but not in database
    missing_in_db = all_primary_slots - db_slots
    for slot in missing_in_db:
        # Find which file has this slot
        source_files = [
            path for path, slots in primary_slots_by_file.items() if slot in slots
        ]
        report.add_difference(
            "Missing in Database",
            slot,
            None,
            f"Slot {slot} found in primary file(s) but not in database",
            {"source_files": source_files},
        )

    # 2. Slots in database but not in primary files
    missing_in_primary = db_slots - all_primary_slots
    for slot in missing_in_primary:
        plan = next(p for p in db_plans if p.slot_number == slot)
        report.add_difference(
            "Missing in Primary Files",
            slot,
            plan.subject,
            f"Slot {slot} ({plan.subject}) in database but not found in any primary file",
            {
                "source_file_name": plan.source_file_name,
                "source_file_path": plan.source_file_path,
            },
        )

    # 3. Slots in database but not in combined_originals
    missing_in_combined = db_slots - combined_slots_set
    for slot in missing_in_combined:
        plan = next(p for p in db_plans if p.slot_number == slot)
        report.add_difference(
            "Missing in Combined DOCX",
            slot,
            plan.subject,
            f"Slot {slot} ({plan.subject}) in database but not in combined_originals.docx",
        )

    # 4. Slots in combined_originals but not in database
    extra_in_combined = combined_slots_set - db_slots
    for slot in extra_in_combined:
        report.add_difference(
            "Extra in Combined DOCX",
            slot,
            combined_slots.get(slot, {}).get("subject"),
            f"Slot {slot} in combined_originals.docx but not in database",
        )

    # 5. Subject mismatches
    for plan in db_plans:
        combined_info = combined_slots.get(plan.slot_number)
        if combined_info:
            combined_subject = combined_info.get("subject", "Unknown")
            if combined_subject != "Unknown" and combined_subject != plan.subject:
                report.add_difference(
                    "Subject Mismatch",
                    plan.slot_number,
                    plan.subject,
                    f"Subject mismatch: DB={plan.subject}, Combined={combined_subject}",
                    {
                        "db_grade": plan.grade,
                        "combined_grade": combined_info.get("grade", "Unknown"),
                        "db_teacher": plan.primary_teacher_name,
                        "combined_teacher": combined_info.get("primary_teacher_name", "Unknown"),
                    },
                )
        
    # 6. Content comparison - check if source file matches
    for plan in db_plans:
        combined_info = combined_slots.get(plan.slot_number)
        if combined_info:
            # Check if teacher name matches (indicates correct source file)
            combined_teacher = combined_info.get("primary_teacher_name", "").lower()
            db_teacher = (plan.primary_teacher_name or "").lower()
            if db_teacher and combined_teacher and db_teacher not in combined_teacher and combined_teacher not in db_teacher:
                report.add_difference(
                    "Teacher Name Mismatch",
                    plan.slot_number,
                    plan.subject,
                    f"Teacher mismatch suggests wrong source file: DB={plan.primary_teacher_name}, Combined={combined_info.get('primary_teacher_name')}",
                    {
                        "db_source_file": plan.source_file_name,
                    },
                )

    # 7. Source file analysis
    # Check if database plans reference files that exist
    for plan in db_plans:
        source_path = Path(plan.source_file_path)
        if not source_path.exists():
            # Try relative to week folder
            relative_path = report.week_folder / source_path.name
            if not relative_path.exists():
                report.add_difference(
                    "Source File Missing",
                    plan.slot_number,
                    plan.subject,
                    f"Source file not found: {plan.source_file_path}",
                    {"source_file_name": plan.source_file_name},
                )
    
    # 8. Detailed slot mapping with source file information
    for plan in db_plans:
        combined_info = combined_slots.get(plan.slot_number)
        if combined_info:
            # Add source file tracking
            combined_info["db_source_file"] = plan.source_file_name
            combined_info["db_primary_teacher"] = plan.primary_teacher_name


async def main():
    """Main comparison function."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Compare primary teacher files with combined_originals.docx"
    )
    parser.add_argument(
        "--week-folder",
        type=str,
        help="Path to week folder (e.g., 'F:/rodri/Documents/OneDrive/AS/Lesson Plan/25 W51')",
    )
    parser.add_argument(
        "--week-of",
        type=str,
        help="Week date range (e.g., '12/16-12/20')",
    )
    parser.add_argument(
        "--user-id",
        type=str,
        help="User ID to query database",
    )
    parser.add_argument(
        "--output",
        type=str,
        help="Output path for JSON report (optional)",
    )

    args = parser.parse_args()

    # Determine week folder
    if args.week_folder:
        week_folder = Path(args.week_folder)
    else:
        # Default to W51 example
        week_folder = Path(
            r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W51"
        )

    if not week_folder.exists():
        print(f"Error: Week folder not found: {week_folder}")
        return

    # Determine week_of
    if args.week_of:
        week_of = args.week_of
    else:
        # Try to extract from folder name or use default
        folder_name = week_folder.name
        # Try to match pattern like "25 W51" or extract dates
        # For now, use a default that matches W51
        week_of = "12/16-12/20"  # Default for W51, user should override

    # Determine user_id
    if args.user_id:
        user_id = args.user_id
    else:
        # Try to find user from database or use default
        # For W51, it's likely Wilson Rodrigues
        user_id = None
        db = get_db()
        # Try to find user by name
        users = db.list_users()
        for user in users:
            if "wilson" in user.name.lower() or "rodrigues" in user.name.lower():
                user_id = user.id
                break

        if not user_id:
            print("Error: Could not determine user_id. Please provide --user-id")
            print("Available users:")
            for user in users:
                print(f"  {user.id}: {user.name}")
            return

    print(f"Using week folder: {week_folder}")
    print(f"Using week_of: {week_of}")
    print(f"Using user_id: {user_id}")

    # Initialize report
    report = ComparisonReport(week_folder, week_of, user_id)

    # 1. Find primary teacher files
    print("\n[Step 1] Finding primary teacher files...")
    report.primary_files = find_primary_teacher_files(week_folder)

    # 2. Query database for original lesson plans
    print("\n[Step 2] Querying database for original lesson plans...")
    db = get_db(user_id=user_id)
    report.db_plans = db.get_original_lesson_plans_for_week(user_id, week_of)

    # 3. Find combined_originals.docx
    print("\n[Step 3] Finding combined_originals.docx...")
    report.combined_originals_path = find_combined_originals(week_folder)
    if report.combined_originals_path:
        report.combined_originals_slots = extract_slots_from_combined_docx(
            report.combined_originals_path
        )

    # 4. Compare and identify differences
    print("\n[Step 4] Comparing sources...")
    compare_slots(
        report,
        report.primary_files,
        report.db_plans,
        report.combined_originals_slots,
    )

    # 5. Print report
    report.print_report()

    # 6. Save JSON report if requested
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = week_folder / f"comparison_report_{week_of.replace('/', '-')}.json"

    report.save_json_report(output_path)


if __name__ == "__main__":
    asyncio.run(main())
