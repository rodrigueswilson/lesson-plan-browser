"""Detailed check of all hyperlinks and Referenced Links sections."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

print(f"Detailed Analysis: {output_file.name}\n")
print("="*80)

doc = Document(output_file)

# Check all tables
print(f"\nTotal tables in document: {len(doc.tables)}\n")

for idx, table in enumerate(doc.tables):
    print(f"\n{'='*80}")
    print(f"TABLE {idx}")
    print(f"{'='*80}")
    
    # Check first cell of first row to identify table type
    if table.rows and table.rows[0].cells:
        first_cell_text = table.rows[0].cells[0].text[:100]
        print(f"First cell: {first_cell_text}")
    
    # Count hyperlinks in this table
    link_count = 0
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                if hyperlinks:
                    link_count += len(hyperlinks)
                    # Show first few links
                    if link_count <= 5:
                        for hyperlink in hyperlinks:
                            text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                            print(f"  Row {row_idx}, Cell {cell_idx}: {text[:60]}")
    
    print(f"\nTotal links in this table: {link_count}")
    
    # Check if this is a "Referenced Links" section
    is_referenced = False
    for row in table.rows:
        for cell in row.cells:
            if 'Referenced Links' in cell.text or 'Referenced Media' in cell.text:
                is_referenced = True
                break
    
    if is_referenced:
        print("⚠️  THIS IS A 'REFERENCED LINKS' SECTION")

# Summary
print(f"\n\n{'='*80}")
print("SUMMARY")
print(f"{'='*80}\n")

# Count all Referenced Links sections
referenced_sections = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if 'Referenced Links' in cell.text or 'Referenced Media' in cell.text:
                referenced_sections += 1
                break

print(f"Number of 'Referenced Links' sections: {referenced_sections}")

# Count total links
total_links = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                total_links += len(hyperlinks)

print(f"Total hyperlinks in document: {total_links}")
