"""
Batch processor for handling multiple class slots and generating combined lesson plans.
Processes all user's class slots and combines them into a single DOCX output.
"""

import asyncio
import re
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

from backend.config import settings
from backend.database import get_db
from backend.file_manager import get_file_manager
from backend.llm_service import LLMService
from backend.original_lesson_schema_models import (
    OriginalAnticipatorySet,
    OriginalAssessment,
    OriginalDayPlanSingleSlot,
    OriginalHomework,
    OriginalHyperlink,
    OriginalHyperlinks,
    OriginalInstruction,
    OriginalMaterials,
    OriginalMisconceptions,
    OriginalObjective,
    OriginalTailoredInstruction,
)
from backend.performance_tracker import get_tracker
from backend.progress import progress_tracker
from backend.schema import OriginalLessonPlan
from backend.services.objectives_utils import normalize_objectives_in_lesson
from backend.telemetry import logger
from backend.utils.date_formatter import format_week_dates
from tools.docx_parser import DOCXParser
from tools.docx_renderer import DOCXRenderer

# DEBUG FLAG: Set to True to skip actual LLM calls and return mock data
# This helps isolate where the ModelPrivateAttr error is occurring
MOCK_LLM_CALL = False  # Changed to False to enable real LLM calls


@dataclass
class SlotProcessingContext:
    """Context for processing a single slot across phases.

    This dataclass tracks the state of a slot through the two-phase processing:
    - Phase 1: Sequential file operations (extract content)
    - Phase 2: Parallel LLM processing (transform content)
    """

    slot: Dict[str, Any]
    slot_index: int
    total_slots: int
    primary_file: Optional[str] = None
    extracted_content: Optional[str] = None
    available_days: Optional[List[str]] = None
    no_school_days: Optional[List[str]] = None
    lesson_json: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    cache_hit: bool = False

    # Parallel processing metrics (for analytics)
    is_parallel: bool = False
    parallel_slot_count: Optional[int] = None
    sequential_time_ms: Optional[float] = None
    rate_limit_errors: int = 0
    concurrency_level: Optional[int] = None
    tpm_usage: Optional[int] = None
    rpm_usage: Optional[int] = None
    link_map: Dict[str, Dict[str, Any]] = field(default_factory=dict)


class BatchProcessor:
    """Batch processor for multi-slot lesson plan generation."""

    def __init__(self, llm_service: LLMService):
        """Initialize batch processor.

        Args:
            llm_service: LLM service instance for transformations
        """
        self.llm_service = llm_service
        self.db = get_db()
        self.tracker = get_tracker()
        self._file_locks = {}  # Cache of asyncio.Lock per file path

    def _build_teacher_name(self, user: Dict[str, Any], slot: Dict[str, Any]) -> str:
        """
        Build teacher name as "Primary First Last / Bilingual First Last".

        Fallback strategy:
        1. Try structured first/last names
        2. Fall back to legacy 'name' and 'primary_teacher_name' fields
        3. Return "Unknown" if all fail

        Args:
            user: User dictionary from database
            slot: Slot dictionary from database

        Returns:
            Formatted teacher name string
        """
        # Primary teacher name - normalize None to empty string before strip
        primary_first = (slot.get("primary_teacher_first_name") or "").strip()
        primary_last = (slot.get("primary_teacher_last_name") or "").strip()

        if primary_first and primary_last:
            primary_name = f"{primary_first} {primary_last}"
        elif primary_first or primary_last:
            primary_name = primary_first or primary_last
        else:
            # Fallback to legacy field
            primary_name = (slot.get("primary_teacher_name") or "").strip()

        # Bilingual teacher name - normalize None to empty string before strip
        bilingual_first = (user.get("first_name") or "").strip()
        bilingual_last = (user.get("last_name") or "").strip()

        if bilingual_first and bilingual_last:
            bilingual_name = f"{bilingual_first} {bilingual_last}"
        elif bilingual_first or bilingual_last:
            bilingual_name = bilingual_first or bilingual_last
        else:
            # Fallback to legacy field
            bilingual_name = (user.get("name") or "").strip()

        # Combine
        if primary_name and bilingual_name:
            return f"{primary_name} / {bilingual_name}"
        elif primary_name:
            return primary_name
        elif bilingual_name:
            return bilingual_name
        else:
            return "Unknown"

    def _no_school_day_stub(self) -> Dict[str, Any]:
        """Return the minimal day structure for a No School day (shared by sequential and parallel paths)."""
        return {
            "unit_lesson": "No School",
            "objective": {
                "content_objective": "No School",
                "student_goal": "No School",
                "wida_objective": "No School",
            },
            "anticipatory_set": {
                "original_content": "",
                "bilingual_bridge": "",
            },
            "tailored_instruction": {
                "original_content": "",
                "co_teaching_model": {},
                "ell_support": [],
                "special_needs_support": [],
                "materials": [],
            },
            "misconceptions": {
                "original_content": "",
                "linguistic_note": {},
            },
            "assessment": {
                "primary_assessment": "",
                "bilingual_overlay": {},
            },
            "homework": {
                "original_content": "",
                "family_connection": "",
            },
        }

    def _scrub_hyperlinks(self, context: SlotProcessingContext):
        """Pre-processing: Replace links with [[LINK_n]] placeholders, tracking which day they belong to."""
        if not context.extracted_content:
            return

        link_count = 0
        link_map = {}

        def replace_with_token(match, day):
            nonlocal link_count
            link_count += 1
            token = f"[[LINK_{link_count}]]"
            original = match.group(0)
            # Store original link and the day it was found in for post-LLM validation
            link_map[token] = {"original": original, "day": day}
            return token

        markdown_pattern = r"\[([^\]]+)\]\((https?://[^\s\)]+)\)"
        raw_url_pattern = r"https?://[^\s<>{}\"|\\^`\[\]]+"

        # Split content into days to preserve day-link association
        days_of_week = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY"]
        day_regex = r"\n(" + "|".join(days_of_week) + r")"
        
        parts = re.split(day_regex, context.extracted_content)
        
        # parts[0] is content before any day header (metadata)
        current_day = "metadata"
        processed_parts = []
        
        # Scrub pre-day content
        pre_content = parts[0]
        pre_content = re.sub(markdown_pattern, lambda m: replace_with_token(m, current_day), pre_content)
        pre_content = re.sub(raw_url_pattern, lambda m: replace_with_token(m, current_day), pre_content)
        processed_parts.append(pre_content)
        
        # Scrub each day's content
        for i in range(1, len(parts), 2):
            day_header = parts[i]
            current_day = day_header.lower()
            day_content = parts[i+1] if i+1 < len(parts) else ""
            
            day_content = re.sub(markdown_pattern, lambda m: replace_with_token(m, current_day), day_content)
            day_content = re.sub(raw_url_pattern, lambda m: replace_with_token(m, current_day), day_content)
            
            processed_parts.append(f"\n{day_header}{day_content}")
            
        context.extracted_content = "".join(processed_parts)
        context.link_map = link_map

        if link_count > 0:
            logger.info(
                "links_scrubbed_with_days",
                extra={
                    "slot": context.slot.get("slot_number"),
                    "link_count": link_count,
                    "days_tracked": list(set(entry["day"] for entry in link_map.values())),
                },
            )

    def _restore_hyperlinks(
        self, data: Any, link_map: Dict[str, Dict[str, Any]]
    ) -> Tuple[Any, set]:
        """Post-processing: Recursively swap placeholders back for original links with day-matching validation.
        Returns a tuple of (restored_data, set_of_original_links_restored).
        """
        restored_links = set()

        if not link_map:
            return data, restored_links

        def recurse(item, current_day=None):
            if isinstance(item, dict):
                new_dict = {}
                for k, v in item.items():
                    # If this is a day key in the "days" dictionary, track it
                    day_context = current_day
                    if k.lower() in ["monday", "tuesday", "wednesday", "thursday", "friday"]:
                        day_context = k.lower()
                    new_dict[k] = recurse(v, day_context)
                return new_dict
            elif isinstance(item, list):
                return [recurse(i, current_day) for i in item]
            elif isinstance(item, str):
                text = item
                # Sort tokens by length descending to avoid partial matching issues
                sorted_tokens = sorted(link_map.keys(), key=len, reverse=True)
                for token in sorted_tokens:
                    pattern = re.compile(re.escape(token), re.IGNORECASE)
                    if pattern.search(text):
                        link_info = link_map[token]
                        original = link_info["original"]
                        source_day = link_info["day"]
                        
                        # DAY VALIDATION: Reject if token is placed in the wrong day
                        if current_day and source_day and source_day != "metadata" and source_day != current_day:
                            logger.warning(
                                "cross_day_link_restoration_rejected",
                                extra={
                                    "token": token,
                                    "source_day": source_day,
                                    "target_day": current_day,
                                    "link": original
                                }
                            )
                            # Replace token with original text (without link) so renderer can place it properly later
                            text_only = original
                            if original.startswith("[") and "]" in original:
                                text_only = original[1:original.find("]")]
                            text = pattern.sub(text_only, text)
                            continue
                            
                        restored_links.add(original)
                        # Robustness: If the original was a raw URL, ensure it's wrapped in markdown
                        replacement = original
                        if replacement.startswith("http") and not replacement.startswith("["):
                            replacement = f"[{replacement}]({replacement})"
                        text = pattern.sub(replacement, text)
                return text
            return item

        return recurse(data), restored_links

    async def process_user_week(
        self,
        user_id: str,
        week_of: str,
        provider: str = "openai",
        week_folder_path: Optional[str] = None,
        slot_ids: Optional[list] = None,
        plan_id: Optional[str] = None,
        partial: bool = False,
        missing_only: bool = False,
        force_slots: Optional[List[int]] = None,
    ) -> Dict[str, Any]:
        """Process all class slots for a user's week.

        Args:
            user_id: User ID
            week_of: Week date range (MM/DD-MM/DD)
            provider: LLM provider to use
            week_folder_path: Optional path to week folder (overrides auto-detection)
            slot_ids: Optional list of specific slot IDs to process
            plan_id: Optional plan ID for progress tracking (if already created)
            partial: Whether to merge with an existing plan for the week
            missing_only: Whether to automatically identify and process only missing slots

        Returns:
            Dict with processing results
        """
        originals_docx = None
        try:
            logger.info(
                "process_user_week_START",
                extra={"user_id": user_id, "week_of": week_of},
            )
            print(
                f"\n{'=' * 60}\nPROCESS_USER_WEEK STARTED\nUser: {user_id}\nWeek: {week_of}\n{'=' * 60}\n"
            )
            start_time = datetime.now()
        except Exception as e:
            print(f"\n{'=' * 60}\nERROR IN PROCESS_USER_WEEK INIT: {e}\n{'=' * 60}\n")
            raise

        # Get user-specific database instance to ensure we're using the correct project
        db = get_db(user_id=user_id)
        print(f"DEBUG: Using database instance for user_id: {user_id}")

        # Get user and slots (wrapped in asyncio.to_thread to avoid blocking event loop)
        print("DEBUG: About to call db.get_user()")
        user_raw = await asyncio.to_thread(db.get_user, user_id)
        if not user_raw:
            return {"success": False, "error": f"User not found: {user_id}"}

        # Convert User object to dictionary
        if hasattr(user_raw, "model_dump"):
            user = user_raw.model_dump()
        elif hasattr(user_raw, "dict"):
            user = user_raw.dict()
        else:
            user = dict(user_raw)
        print(f"DEBUG: Got user: {user.get('name')}")

        print("DEBUG: About to call db.get_user_slots()")
        slots_raw = await asyncio.to_thread(db.get_user_slots, user_id)
        print(f"DEBUG: Got {len(slots_raw)} slots")
        if not slots_raw:
            return {
                "success": False,
                "error": f"No class slots configured for user: {user['name']}",
            }

        # Convert ClassSlot objects to dictionaries so we can enrich them
        slots = []
        for slot_obj in slots_raw:
            slot_dict = None

            # Try model_dump first (Pydantic v2 / SQLModel)
            if hasattr(slot_obj, "model_dump"):
                try:
                    slot_dict = slot_obj.model_dump(mode="python")
                except Exception:
                    try:
                        slot_dict = slot_obj.model_dump()
                    except Exception:
                        slot_dict = None

            # Try dict() method (Pydantic v1)
            if slot_dict is None and hasattr(slot_obj, "dict"):
                try:
                    slot_dict = slot_obj.dict()
                except Exception:
                    slot_dict = None

            # Already a dict - ensure all values are properly extracted
            if slot_dict is None and isinstance(slot_obj, dict):
                slot_dict = {}
                for key, value in slot_obj.items():
                    # Skip ModelPrivateAttr and similar objects
                    if hasattr(value, "__class__") and "ModelPrivateAttr" in str(
                        type(value)
                    ):
                        try:
                            # Try to get the actual value
                            slot_dict[key] = getattr(slot_obj, key, None)
                        except Exception:
                            slot_dict[key] = None
                    else:
                        slot_dict[key] = value

            # Fallback: Extract as dictionary from object attributes
            if slot_dict is None:
                slot_dict = {
                    "id": str(getattr(slot_obj, "id", ""))
                    if hasattr(slot_obj, "id")
                    else "",
                    "user_id": str(getattr(slot_obj, "user_id", ""))
                    if hasattr(slot_obj, "user_id")
                    else "",
                    "slot_number": int(getattr(slot_obj, "slot_number", 0))
                    if hasattr(slot_obj, "slot_number")
                    else 0,
                    "subject": str(getattr(slot_obj, "subject", ""))
                    if hasattr(slot_obj, "subject")
                    else "",
                    "grade": str(getattr(slot_obj, "grade", ""))
                    if hasattr(slot_obj, "grade")
                    else "",
                    "homeroom": str(getattr(slot_obj, "homeroom", ""))
                    if hasattr(slot_obj, "homeroom")
                    else None,
                    "plan_group_label": str(getattr(slot_obj, "plan_group_label", ""))
                    if hasattr(slot_obj, "plan_group_label")
                    else None,
                    "proficiency_levels": str(
                        getattr(slot_obj, "proficiency_levels", "")
                    )
                    if hasattr(slot_obj, "proficiency_levels")
                    else None,
                    "primary_teacher_file": str(
                        getattr(slot_obj, "primary_teacher_file", "")
                    )
                    if hasattr(slot_obj, "primary_teacher_file")
                    else None,
                    "primary_teacher_name": str(
                        getattr(slot_obj, "primary_teacher_name", "")
                    )
                    if hasattr(slot_obj, "primary_teacher_name")
                    else None,
                    "primary_teacher_first_name": str(
                        getattr(slot_obj, "primary_teacher_first_name", "")
                    )
                    if hasattr(slot_obj, "primary_teacher_first_name")
                    else None,
                    "primary_teacher_last_name": str(
                        getattr(slot_obj, "primary_teacher_last_name", "")
                    )
                    if hasattr(slot_obj, "primary_teacher_last_name")
                    else None,
                    "primary_teacher_file_pattern": str(
                        getattr(slot_obj, "primary_teacher_file_pattern", "")
                    )
                    if hasattr(slot_obj, "primary_teacher_file_pattern")
                    else None,
                    "display_order": int(getattr(slot_obj, "display_order", 0))
                    if hasattr(slot_obj, "display_order")
                    else 0,
                }

            # Ensure all values are safe (no ModelPrivateAttr objects)
            for key in list(slot_dict.keys()):
                value = slot_dict[key]
                if hasattr(value, "__class__") and "ModelPrivateAttr" in str(
                    type(value)
                ):
                    slot_dict[key] = None
                elif not isinstance(
                    value, (str, int, float, bool, type(None), list, dict)
                ):
                    # Convert other non-serializable types to None
                    try:
                        slot_dict[key] = str(value)
                    except Exception:
                        slot_dict[key] = None

            slots.append(slot_dict)

        # Enrich slots with schedule entry data (start_time, end_time)
        print("DEBUG: Enriching slots with schedule data")
        schedule_entries = await asyncio.to_thread(db.get_user_schedule, user_id)
        # Convert schedule entries to dictionaries to avoid ModelPrivateAttr issues
        schedule_entries_dict = []
        for entry in schedule_entries:
            if hasattr(entry, "model_dump"):
                entry_dict = entry.model_dump()
            elif hasattr(entry, "dict"):
                entry_dict = entry.dict()
            else:
                entry_dict = {
                    "slot_number": getattr(entry, "slot_number", None),
                    "is_active": getattr(entry, "is_active", True),
                    "start_time": getattr(entry, "start_time", None),
                    "end_time": getattr(entry, "end_time", None),
                }
            schedule_entries_dict.append(entry_dict)

        def normalize_subj(s):
            if not s:
                return ""
            # Lowercase and keep only alphanumeric characters for comparison
            return re.sub(r"[^a-z0-9]", "", s.lower())

        for slot in slots:
            slot_subj_norm = normalize_subj(slot.get("subject", ""))
            slot_num = slot.get("slot_number")

            # Find matching schedule entry using a multi-stage strategy
            # Stage 1: Match by both subject and slot number (tightest)
            matching_entries = [
                e
                for e in schedule_entries_dict
                if normalize_subj(e.get("subject", "")) == slot_subj_norm
                and e.get("slot_number") == slot_num
                and e.get("is_active", True)
            ]

            # Stage 2: Match by subject only (useful if slot numbers shifted)
            if not matching_entries:
                matching_entries = [
                    e
                    for e in schedule_entries_dict
                    if normalize_subj(e.get("subject", "")) == slot_subj_norm
                    and e.get("is_active", True)
                ]

            # Stage 3: Match by slot number only (legacy fallback)
            if not matching_entries:
                matching_entries = [
                    e
                    for e in schedule_entries_dict
                    if e.get("slot_number") == slot_num
                    and e.get("is_active", True)
                ]

            if matching_entries:
                # Store all day-specific times for accurate chronological ordering in PDFs
                day_times = {}
                for e in matching_entries:
                    day = e.get("day_of_week", "").lower()
                    if day:
                        day_times[day] = {
                            "start_time": e.get("start_time"),
                            "end_time": e.get("end_time"),
                        }
                slot["day_times"] = day_times

                # Keep the "most common" time as the default/fallback
                from collections import Counter

                times = [
                    (e.get("start_time"), e.get("end_time")) for e in matching_entries
                ]
                most_common_time = Counter(times).most_common(1)[0][0]
                slot["start_time"] = most_common_time[0]
                slot["end_time"] = most_common_time[1]
                print(
                    f"DEBUG: Enriched slot {slot.get('slot_number')} ({slot.get('subject')}) with day_times for {list(day_times.keys())}"
                )
            else:
                print(
                    f"DEBUG: No schedule entry found for slot {slot.get('slot_number')} ({slot.get('subject')})"
                )

        # Filter slots if specific slot_ids provided
        # Convert slot_ids to a set of strings for efficient lookup
        if slot_ids:
            print(f"DEBUG: Filtering slots with provided slot_ids: {slot_ids}")
            print(f"DEBUG: Total slots before filtering: {len(slots)}")
            print(f"DEBUG: Slot IDs in database: {[slot.get('id') for slot in slots]}")

            slot_ids_set = set()
            for sid in slot_ids:
                if isinstance(sid, str):
                    slot_ids_set.add(sid)
                elif hasattr(sid, "id"):
                    slot_ids_set.add(str(sid.id))
                else:
                    slot_ids_set.add(str(sid))

            print(f"DEBUG: slot_ids_set after conversion: {slot_ids_set}")

            slots = [slot for slot in slots if slot.get("id") in slot_ids_set]
            print(f"DEBUG: Total slots after filtering: {len(slots)}")
            print(f"DEBUG: Filtered slot IDs: {[slot.get('id') for slot in slots]}")

            if not slots:
                error_msg = f"No matching slots found for provided slot_ids. Requested: {slot_ids_set}, Available: {[slot.get('id') for slot in slots_raw if hasattr(slot, 'id') or isinstance(slot, dict)]}"
                print(f"ERROR: {error_msg}")
                return {
                    "success": False,
                    "error": error_msg,
                }

        # Always check for existing plan for this week (Optimization cache)
        existing_lesson_json = None
        print("DEBUG: Checking for existing plan for week process optimization")
        # Try to get existing plan for this week
        existing_plans = await asyncio.to_thread(db.get_user_plans, user_id, limit=5)
        # Find the first one with the same week_of
        existing_plan = next(
            (p for p in existing_plans if p.week_of == week_of), None
        )

        if existing_plan:
            print(f"DEBUG: Found existing plan {existing_plan.id} for week {week_of}")
            # Use current plan_id if not provided
            if not plan_id:
                plan_id = existing_plan.id

            # Load existing lesson_json
            existing_lesson_json = existing_plan.lesson_json
            if isinstance(existing_lesson_json, str):
                import json

                try:
                    existing_lesson_json = json.loads(existing_lesson_json)
                except Exception:
                    existing_lesson_json = None

            if existing_lesson_json:
                # Enrich from steps if needed (api.py logic)
                from backend.utils.lesson_json_enricher import (
                    enrich_lesson_json_from_steps,
                )

                existing_lesson_json = enrich_lesson_json_from_steps(
                    existing_lesson_json, existing_plan.id, db
                )
                print(
                    f"DEBUG: Loaded and enriched existing lesson_json for plan {existing_plan.id}"
                )
                logger.info(
                    "skip_logic_plan_found",
                    extra={
                        "plan_id": existing_plan.id,
                        "status": "valid_json_loaded",
                        "week_of": week_of
                    }
                )

                if missing_only:
                    print("DEBUG: Identifying missing slots...")
                    # Extract existing slot IDs or slot numbers from existing_lesson_json
                    existing_slots_data = []
                    for day in existing_lesson_json.get("days", {}).values():
                        existing_slots_data.extend(day.get("slots", []))

                    existing_slot_numbers = {
                        s.get("slot_number")
                        for s in existing_slots_data
                        if s.get("slot_number") is not None
                    }
                    print(f"DEBUG: Existing slot numbers: {existing_slot_numbers}")

                    # Filter current slots to only those NOT in existing_slots_data
                    # We use slot_number as the primary identifier for "missing"
                    original_count = len(slots)
                    slots = [
                        s
                        for s in slots
                        if s.get("slot_number") not in existing_slot_numbers
                    ]
                    print(
                        f"DEBUG: Filtered from {original_count} to {len(slots)} missing slots"
                    )
            else:
                logger.info(
                    "skip_logic_plan_invalid",
                    extra={
                        "plan_id": existing_plan.id,
                        "status": "json_empty_or_invalid",
                        "week_of": week_of
                    }
                )
        else:
            print(
                f"DEBUG: No existing plan found for week {week_of}, proceeding with generation"
            )
            logger.info(
                "skip_logic_no_plan",
                extra={
                    "status": "plan_not_found_in_db",
                    "week_of": week_of
                }
            )
            # If no existing plan, we can't do "partial" merge at the end
            partial = False
            missing_only = False

        # Store user's metadata for file resolution and rendering
        self._current_user_id = user_id
        # Prefer base_path_override, but fall back to base_path if available in the user object
        self._user_base_path = user.get("base_path_override") or user.get("base_path")
        self._user_first_name = user.get("first_name", "")
        self._user_last_name = user.get("last_name", "")
        self._user_name = user.get("name", "")
        logger.info(
            "batch_user_base_path",
            extra={
                "user_id": user_id,
                "user_name": user["name"],
                "base_path": self._user_base_path or "default",
            },
        )

        # Determine if this is a consolidated plan
        is_consolidated = len(slots) > 1

        # Create or use existing weekly plan record
        if not plan_id:
            plan_id = await asyncio.to_thread(
                db.create_weekly_plan,
                user_id,
                week_of,
                output_file="",  # Will be updated after generation
                week_folder_path=week_folder_path,
                consolidated=is_consolidated,
                total_slots=len(slots),
            )
            await asyncio.to_thread(db.update_weekly_plan, plan_id, status="processing")

        # Store plan_id for tracking
        self._current_plan_id = plan_id

        # Start tracking total batch duration
        batch_op_id = ""
        try:
            batch_op_id = self.tracker.start_operation(
                plan_id,
                "batch_process",
                metadata={"total_slots": len(slots), "consolidated": is_consolidated},
            )
        except Exception as e:
            print(f"WARNING: Failed to start batch tracking: {e}")

        processing_weight = settings.PROGRESS_PROCESSING_WEIGHT
        rendering_weight = settings.PROGRESS_RENDERING_WEIGHT

        # Ensure weights sum to 1 for progress tracking
        if abs((processing_weight + rendering_weight) - 1.0) > 1e-6:
            total = processing_weight + rendering_weight
            processing_weight = processing_weight / total
            rendering_weight = rendering_weight / total

        # Process each slot
        lessons = []
        errors = []

        print(f"DEBUG: About to start processing {len(slots)} slots")
        # Update progress: starting
        progress_tracker.update(
            plan_id, "processing", 0, f"Starting to process {len(slots)} slots..."
        )
        print("DEBUG: Progress tracker updated - starting")

        # Check if parallel processing is enabled
        use_parallel = settings.PARALLEL_LLM_PROCESSING and len(slots) > 1

        if use_parallel:
            # Two-phase processing: Extract sequentially, transform in parallel
            print(f"DEBUG: Using parallel processing for {len(slots)} slots")

            # Phase 1: Parallel Extraction with DB Caching and Grouping
            contexts = await self._extract_slots_parallel_db(
                slots,
                week_of,
                week_folder_path,
                self._user_base_path,
                plan_id,
                progress_tracker,
            )

            # AUTO-GENERATE ORIGINALS AUDIT DOCX (Parallel Path)
            try:
                logger.info("batch_auto_generating_originals_docx_parallel")
                originals_docx = await self._generate_combined_original_docx(
                    user_id, week_of, plan_id, week_folder_path
                )
            except Exception as e:
                logger.error(
                    "batch_auto_originals_generation_failed_parallel",
                    extra={"error": str(e)},
                )

            # Re-collect errors from contexts
            for ctx in contexts:
                if ctx.error:
                    errors.append(
                        f"Failed slot {ctx.slot_index}/{ctx.total_slots}: {ctx.error}"
                    )

            # NEW: Optimize Phase 2 - Avoid LLM calls for unchanged slots if we already have transformed versions
            if existing_lesson_json:
                print(
                    "DEBUG: Checking for slots to skip LLM transformation (already in DB and unchanged source)..."
                )
                existing_slot_plans = self._reconstruct_slots_from_json(
                    existing_lesson_json
                )

                for ctx in contexts:
                    if ctx.cache_hit and not ctx.error:
                        slot_num = ctx.slot.get("slot_number")
                        if slot_num in existing_slot_plans and slot_num not in (
                            force_slots or []
                        ):
                            print(
                                f"[DEBUG] REUSING TRANSFORMED PLAN for slot {slot_num} ({ctx.slot['subject']}) - Source file unchanged."
                            )
                            # CRITICAL: Create a deep copy to prevent shared state in parallel processing
                            # If multiple slots share the same lesson_json reference, mutations in one
                            # parallel task could affect others, causing all slots to show the same teacher
                            import copy
                            ctx.lesson_json = copy.deepcopy(existing_slot_plans[slot_num][
                                "lesson_json"
                            ])
                        elif slot_num in existing_slot_plans and slot_num in (
                            force_slots or []
                        ):
                            print(
                                f"[DEBUG] FORCING AI transformation for slot {slot_num} ({ctx.slot['subject']}) as requested."
                            )

            # Phase 2: Transform with LLM in parallel
            print(
                f"DEBUG: Starting Phase 2: Parallel LLM transformation for {len(contexts)} slots"
            )
            transform_count = len(
                [c for c in contexts if not c.error and not c.lesson_json]
            )
            progress_tracker.update(
                plan_id,
                "processing",
                20,
                f"Transforming {transform_count} slots with AI in parallel...",
            )

            contexts = await self._process_slots_parallel_llm(
                contexts,
                week_of,
                provider,
                plan_id,
            )

            # Collect results
            for context in contexts:
                if context.error:
                    errors.append(
                        f"Failed slot {context.slot_index}/{context.total_slots}: {context.error}"
                    )
                elif context.lesson_json:
                    # LOGGING: Verify hyperlinks are present before collecting
                    hyperlinks_in_json = context.lesson_json.get("_hyperlinks", [])
                    print(
                        f"[DEBUG] Collecting parallel result: Slot {context.slot.get('slot_number')}, "
                        f"Subject {context.slot.get('subject')}, "
                        f"Hyperlinks in lesson_json: {len(hyperlinks_in_json)}"
                    )
                    logger.info(
                        "parallel_result_collection",
                        extra={
                            "slot": context.slot.get("slot_number"),
                            "subject": context.slot.get("subject"),
                            "has_lesson_json": context.lesson_json is not None,
                            "hyperlinks_count": len(hyperlinks_in_json),
                            "has_hyperlinks_key": "_hyperlinks" in context.lesson_json,
                        },
                    )
                    # Include original slot data for primary teacher fields
                    # context.slot is already a dict with primary teacher fields from database
                    slot_data = context.slot.copy() if isinstance(context.slot, dict) else {
                        "slot_number": getattr(context.slot, "slot_number", context.slot_index),
                        "subject": getattr(context.slot, "subject", "Unknown"),
                        "primary_teacher_name": getattr(context.slot, "primary_teacher_name", None),
                        "primary_teacher_first_name": getattr(context.slot, "primary_teacher_first_name", None),
                        "primary_teacher_last_name": getattr(context.slot, "primary_teacher_last_name", None),
                    }
                    # Debug: Verify primary teacher fields are present
                    if isinstance(slot_data, dict):
                        print(f"[DEBUG] BATCH_PROCESSOR: Slot {slot_data.get('slot_number')} slot_data has primary_teacher_name: {slot_data.get('primary_teacher_name')}")
                    lessons.append(
                        {
                            "slot_number": slot_data.get("slot_number") if isinstance(slot_data, dict) else getattr(context.slot, "slot_number", context.slot_index),
                            "subject": slot_data.get("subject") if isinstance(slot_data, dict) else getattr(context.slot, "subject", "Unknown"),
                            "lesson_json": context.lesson_json,
                            "slot_data": slot_data,  # Include original slot data for teacher fields
                        }
                    )
                    # Track parallel processing metrics
                    if plan_id and context.is_parallel:
                        # Store metrics in performance tracker result
                        # This will be picked up by end_operation
                        pass  # Metrics are already in context

        else:
            # Sequential processing (fallback or single slot)
            print(
                "DEBUG: Using sequential processing (parallel disabled or single slot)"
            )

            for i, slot in enumerate(slots, 1):
                # Ensure slot is a dictionary with safe access
                slot_num = (
                    slot.get("slot_number")
                    if isinstance(slot, dict)
                    else getattr(slot, "slot_number", i)
                )
                slot_subject = (
                    slot.get("subject")
                    if isinstance(slot, dict)
                    else getattr(slot, "subject", "Unknown")
                )

                print(
                    f"\nDEBUG: === Starting slot {i}/{len(slots)} (slot_number={slot_num}): {slot_subject} ==="
                )
                try:
                    # CRITICAL: Sanitize slot immediately before ANY access
                    # This handles Pydantic/SQLModel PrivateAttr issues at the source
                    try:
                        slot = self._sanitize_slot(slot)
                    except Exception as sanitize_err:
                        print(f"ERROR: _sanitize_slot failed: {sanitize_err}")
                        # Fallback to dict conversion
                        if hasattr(slot, "model_dump"):
                            slot = slot.model_dump()
                        elif hasattr(slot, "dict"):
                            slot = slot.dict()
                        else:
                            slot = dict(slot)

                    # Ensure slot is fully a dictionary for safe access
                    if not isinstance(slot, dict):
                        if hasattr(slot, "model_dump"):
                            slot = slot.model_dump()
                        elif hasattr(slot, "dict"):
                            slot = slot.dict()
                        else:
                            # Convert to dict manually
                            slot = {
                                "id": getattr(slot, "id", None),
                                "slot_number": getattr(slot, "slot_number", i),
                                "subject": getattr(slot, "subject", "Unknown"),
                                "grade": getattr(slot, "grade", None),
                                "homeroom": getattr(slot, "homeroom", None),
                                "primary_teacher_name": getattr(
                                    slot, "primary_teacher_name", None
                                ),
                                "primary_teacher_first_name": getattr(
                                    slot, "primary_teacher_first_name", None
                                ),
                                "primary_teacher_last_name": getattr(
                                    slot, "primary_teacher_last_name", None
                                ),
                                "primary_teacher_file": getattr(
                                    slot, "primary_teacher_file", None
                                ),
                                "primary_teacher_file_pattern": getattr(
                                    slot, "primary_teacher_file_pattern", None
                                ),
                            }

                    # RE-SANITIZE just in case conversion re-introduced issues
                    slot = self._sanitize_slot(slot)

                    # Update slot in list to use dictionary version
                    slots[i - 1] = slot

                    # Update progress: processing slot
                    progress_pct = int((i - 1) / len(slots) * processing_weight * 100)
                    print(
                        f"DEBUG: Updating progress tracker for slot {i} (slot_number={slot.get('slot_number')})"
                    )
                    progress_tracker.update(
                        plan_id,
                        "processing",
                        progress_pct,
                        f"Processing slot {i}/{len(slots)}: {slot.get('subject', 'Unknown')} ({slot.get('primary_teacher_name', 'No teacher')})",
                    )
                    print(f"DEBUG: Progress tracker updated for slot {i}")

                    logger.info(
                        "batch_slot_processing",
                        extra={
                            "plan_id": plan_id,
                            "slot_index": i,
                            "total_slots": len(slots),
                            "subject": slot.get("subject", "Unknown"),
                            "teacher": slot.get("primary_teacher_name"),
                        },
                    )
                    print(
                        f"DEBUG: About to call _process_slot for slot {i} (slot_number={slot.get('slot_number')}, subject={slot.get('subject')})"
                    )
                    lesson_json = await self._process_slot(
                        slot,
                        week_of,
                        provider,
                        week_folder_path,
                        self._user_base_path,
                        plan_id,
                        i,
                        len(slots),
                        processing_weight,
                        existing_lesson_json=existing_lesson_json,
                        force_ai=slot.get("slot_number") in (force_slots or []),
                    )
                    print(f"[DEBUG] _process_slot (SEQUENTIAL) completed for slot {i}")
                    # LOGGING: Verify hyperlinks are present before collecting
                    hyperlinks_in_json = lesson_json.get("_hyperlinks", [])
                    print(
                        f"[DEBUG] Collecting sequential result: Slot {slot.get('slot_number')}, "
                        f"Subject {slot.get('subject')}, "
                        f"Hyperlinks in lesson_json: {len(hyperlinks_in_json)}"
                    )
                    logger.info(
                        "sequential_result_collection",
                        extra={
                            "slot": slot.get("slot_number"),
                            "subject": slot.get("subject"),
                            "hyperlinks_count": len(hyperlinks_in_json),
                            "has_hyperlinks_key": "_hyperlinks" in lesson_json,
                        },
                    )
                    # Include original slot data for primary teacher fields
                    # slot is already a dict with primary teacher fields from database
                    slot_data = slot.copy() if isinstance(slot, dict) else {
                        "slot_number": getattr(slot, "slot_number", i),
                        "subject": getattr(slot, "subject", "Unknown"),
                        "primary_teacher_name": getattr(slot, "primary_teacher_name", None),
                        "primary_teacher_first_name": getattr(slot, "primary_teacher_first_name", None),
                        "primary_teacher_last_name": getattr(slot, "primary_teacher_last_name", None),
                    }
                    # Debug: Verify primary teacher fields are present
                    if isinstance(slot_data, dict):
                        print(f"[DEBUG] BATCH_PROCESSOR: Slot {slot_data.get('slot_number')} slot_data has primary_teacher_name: {slot_data.get('primary_teacher_name')}")
                    lessons.append(
                        {
                            "slot_number": slot_data.get("slot_number") if isinstance(slot_data, dict) else getattr(slot, "slot_number", i),
                            "subject": slot_data.get("subject") if isinstance(slot_data, dict) else getattr(slot, "subject", "Unknown"),
                            "lesson_json": lesson_json,
                            "slot_data": slot_data,  # Include original slot data for teacher fields
                        }
                    )
                    print(f"[DEBUG] Lesson appended for slot {i}")
                    logger.info(
                        "batch_slot_completed",
                        extra={
                            "plan_id": plan_id,
                            "slot_index": i,
                            "total_slots": len(slots),
                            "subject": slot.get("subject", "Unknown"),
                        },
                    )
                    print(f"DEBUG: === Completed slot {i}/{len(slots)} ===")

                    # Update progress: slot completed
                    progress_pct = int(i / len(slots) * processing_weight * 100)
                    progress_tracker.update(
                        plan_id,
                        "processing",
                        progress_pct,
                        f"Completed slot {i}/{len(slots)}: {slot.get('subject', 'Unknown')}",
                    )
                except Exception as e:
                    print(
                        f"ERROR: Exception in process_weekly_plan loop for slot {i}: {e}"
                    )
                    traceback.print_exc()
                    # Use safe access for error message
                    slot_num = (
                        slot.get("slot_number")
                        if isinstance(slot, dict)
                        else getattr(slot, "slot_number", i)
                    )
                    slot_subject = (
                        slot.get("subject")
                        if isinstance(slot, dict)
                        else getattr(slot, "subject", "Unknown")
                    )
                    # Safely format error message, handling encoding errors
                    try:
                        error_str = str(e)
                    except UnicodeEncodeError:
                        error_str = repr(e).encode("ascii", "replace").decode("ascii")
                    error_msg = f"Slot {slot_num} ({slot_subject}): {error_str}"
                    errors.append(error_msg)
                    logger.error(
                        "batch_slot_error",
                        extra={
                            "plan_id": plan_id,
                            "slot_index": i,
                            "total_slots": len(slots),
                            "subject": slot.get("subject", "Unknown"),
                            "error": error_msg,
                        },
                    )

                    # Update progress: slot failed
                    progress_tracker.update(
                        plan_id,
                        "error",
                        int(i / len(slots) * processing_weight * 100),
                        f"Failed slot {i}/{len(slots)}: {error_msg}",
                    )

        # Sequential extraction/transformation loop ends here

        # AUTO-GENERATE ORIGINALS AUDIT DOCX (Sequential/Fallback Path)
        # ONLY if not already generated in parallel path
        if not originals_docx:
            try:
                logger.info("batch_auto_generating_originals_docx_sequential")
                originals_docx = await self._generate_combined_original_docx(
                    user_id, week_of, plan_id, week_folder_path
                )
            except Exception as e:
                logger.error(
                    "batch_auto_originals_generation_failed_sequential",
                    extra={"error": str(e)},
                )

        # Handle merging with existing lessons if partial is requested
        all_lessons_for_rendering = lessons.copy()
        if (partial or missing_only) and existing_lesson_json:
            print("DEBUG: Reconstructing existing slots for merge...")
            existing_slots_by_number = self._reconstruct_slots_from_json(
                existing_lesson_json
            )

            # Combine: new lessons override existing ones
            combined_lessons_by_number = existing_slots_by_number
            for lesson in lessons:
                combined_lessons_by_number[lesson["slot_number"]] = lesson

            all_lessons_for_rendering = list(combined_lessons_by_number.values())
            # Sort by slot number for consistent rendering
            all_lessons_for_rendering.sort(key=lambda x: x.get("slot_number", 0))
            print(
                f"DEBUG: Combined to {len(all_lessons_for_rendering)} total slots for rendering"
            )

        # Generate combined DOCX if we have any successful lessons
        print(
            f"\nDEBUG: Finished processing all slots, {len(all_lessons_for_rendering)} total for rendering"
        )
        output_file = None
        if all_lessons_for_rendering:
            try:
                print(
                    f"DEBUG: About to render {len(all_lessons_for_rendering)} lessons"
                )
                # Update progress: rendering
                rendering_progress = int(
                    processing_weight * 100 + rendering_weight * 50
                )
                print("DEBUG: Updating progress tracker for rendering")
                progress_tracker.update(
                    plan_id,
                    "rendering",
                    rendering_progress,
                    f"Rendering {len(all_lessons_for_rendering)} lessons to DOCX...",
                )
                print("DEBUG: Progress tracker updated for rendering")

                print("DEBUG: Calling _combine_lessons (wrapped in asyncio.to_thread)")
                output_file = await asyncio.to_thread(
                    self._combine_lessons,
                    user,
                    all_lessons_for_rendering,
                    week_of,
                    start_time,
                    plan_id,
                )
                print(f"DEBUG: _combine_lessons returned: {output_file}")

                # Note: Objectives and Sentence Frames PDFs are already generated in _combine_lessons_impl
                # with correct enrichment and proper filenames. Removing duplicate generation that was
                # creating incorrect files with _Objectives_ and _SentenceFrames_ patterns.

                # Update progress: complete
                print(
                    f"DEBUG: Updating progress tracker to complete (plan_id={plan_id})"
                )
                progress_tracker.update(
                    plan_id,
                    "complete",
                    100,
                    f"Successfully created lesson plan with {len(all_lessons_for_rendering)} slots",
                )
                print("DEBUG: Progress tracker updated to complete")

                # Also mark as complete using the tracker's complete method
                progress_tracker.complete(plan_id)
                print("DEBUG: Progress tracker marked complete via complete() method")
            except Exception as e:
                error_msg = f"Failed to combine lessons: {str(e)}"
                print(f"ERROR: Exception in _combine_lessons: {error_msg}")
                traceback.print_exc()
                errors.append(error_msg)
                progress_tracker.update(
                    plan_id, "error", rendering_progress, f"Failed to render: {str(e)}"
                )

        # Update plan status
        total_time = (datetime.now() - start_time).total_seconds() * 1000

        if output_file:
            # Get merged lesson_json for database storage
            merged_lesson_json = None
            if all_lessons_for_rendering:
                if len(all_lessons_for_rendering) == 1:
                    # Single-slot: convert to slots format for consistency
                    single_lesson_json = all_lessons_for_rendering[0]["lesson_json"]
                    merged_lesson_json = self._convert_single_slot_to_slots_format(
                        single_lesson_json,
                        all_lessons_for_rendering[0]["slot_number"],
                        all_lessons_for_rendering[0]["subject"],
                    )
                else:
                    # Multi-slot: merge lesson JSONs (already creates slots format)
                    from tools.json_merger import merge_lesson_jsons

                    merged_lesson_json = merge_lesson_jsons(all_lessons_for_rendering)

            await asyncio.to_thread(
                db.update_weekly_plan,
                plan_id,
                status="completed",
                output_file=output_file,
                lesson_json=merged_lesson_json,
                total_slots=len(all_lessons_for_rendering),
            )
            # Update performance summary
            await asyncio.to_thread(self.tracker.update_plan_summary, plan_id)
        else:
            await asyncio.to_thread(
                db.update_weekly_plan,
                plan_id,
                status="failed",
                error_message="; ".join(errors),
            )

        # If we have no output file and no errors were collected, something went wrong silently
        if not output_file and not errors:
            error_msg = "Processing failed: No output file generated and no errors were collected. This may indicate an unhandled exception."
            errors = [error_msg]
            logger.error(
                "batch_process_silent_failure",
                extra={
                    "plan_id": plan_id,
                    "lessons_count": len(lessons),
                    "slots_count": len(slots),
                    "output_file": output_file,
                },
            )
            print(f"ERROR: {error_msg}")
            print(
                f"DEBUG: lessons={len(lessons)}, slots={len(slots)}, output_file={output_file}"
            )

        # End tracking total batch duration
        try:
            if batch_op_id:
                self.tracker.end_operation(
                    batch_op_id,
                    result={
                        "success": bool(output_file),
                        "processed_slots": len(lessons),
                        "failed_slots": len(errors),
                        "error": "; ".join(errors) if errors else None,
                    },
                )
        except Exception as e:
            print(f"WARNING: Failed to end batch tracking: {e}")

        return {
            "success": bool(output_file),
            "plan_id": plan_id,
            "output_file": output_file or "",
            "originals_docx": originals_docx,
            "processed_slots": len(lessons),
            "failed_slots": len(errors),
            "total_time_ms": total_time,
            "consolidated": is_consolidated,
            "total_slots": len(slots),
            "errors": errors if errors else None,
        }

    def _sanitize_value(self, value: Any) -> Any:
        """Recursively sanitize a value to remove ModelPrivateAttr objects."""
        # Check for ModelPrivateAttr
        if hasattr(value, "__class__") and "ModelPrivateAttr" in str(type(value)):
            return None

        # Handle lists
        if isinstance(value, list):
            return [self._sanitize_value(item) for item in value]

        # Handle dicts
        if isinstance(value, dict):
            return {k: self._sanitize_value(v) for k, v in value.items()}

        # Handle Pydantic models (v1 and v2)
        if hasattr(value, "model_dump") and callable(value.model_dump):
            try:
                return self._sanitize_value(value.model_dump())
            except Exception:
                pass

        if hasattr(value, "dict") and callable(value.dict):
            try:
                return self._sanitize_value(value.dict())
            except Exception:
                pass

        return value

    async def _open_docx_with_retry(
        self,
        file_path: str,
        plan_id: Optional[str] = None,
        slot_number: Optional[int] = None,
        subject: Optional[str] = None,
        max_retries: int = 3,
        initial_delay: float = 1.0,
    ) -> DOCXParser:
        """
        Open DOCX file with retry logic for OneDrive sync and file locking issues.

        Args:
            file_path: Path to DOCX file
            plan_id: Optional plan ID for tracking
            slot_number: Optional slot number for error messages
            subject: Optional subject for error messages
            max_retries: Maximum number of retry attempts
            initial_delay: Initial delay in seconds (exponential backoff)

        Returns:
            DOCXParser instance

        Raises:
            FileNotFoundError: If file doesn't exist
            PermissionError: If file is locked after all retries
            Exception: If file cannot be opened after all retries
        """
        from docx.opc.exceptions import PackageNotFoundError

        file_path_obj = Path(file_path)

        # Check if file exists
        if not file_path_obj.exists():
            error_msg = f"File not found: {file_path}" + (
                f" (Slot {slot_number}, {subject})" if slot_number and subject else ""
            )
            logger.error(
                "docx_file_not_found",
                extra={
                    "file": file_path,
                    "slot_number": slot_number,
                    "subject": subject,
                },
            )
            raise FileNotFoundError(error_msg)

        # Try to close file in Word if it's open (Windows only)
        def _try_close_word_file(file_path: Path) -> Tuple[bool, Optional[str]]:
            """Attempt to close file in Microsoft Word if it's open (Windows only)."""
            try:
                import sys

                if sys.platform != "win32":
                    return False, "Word automation only available on Windows"

                try:
                    import win32com.client
                except ImportError:
                    return False, "pywin32 not installed (required for Word automation)"

                try:
                    word_app = win32com.client.GetActiveObject("Word.Application")
                except Exception:
                    # Word is not running
                    return True, None  # Not an error - Word just isn't running

                # Check if file is open in Word
                file_path_abs = str(file_path.resolve())
                document_closed = False

                for doc in word_app.Documents:
                    try:
                        doc_path = str(Path(doc.FullName).resolve())
                        if doc_path.lower() == file_path_abs.lower():
                            # File is open - try to close it
                            doc.Close(SaveChanges=False)  # Don't save changes
                            document_closed = True
                            logger.info(
                                "word_file_closed",
                                extra={
                                    "file": file_path,
                                    "slot_number": slot_number,
                                    "subject": subject,
                                },
                            )
                            break
                    except Exception:
                        # Document might be in an invalid state, skip it
                        continue

                return True, "File closed in Word" if document_closed else None

            except Exception as e:
                # Don't fail if Word automation fails - just log and continue
                logger.debug(
                    "word_automation_failed",
                    extra={
                        "file": file_path,
                        "error": str(e),
                        "error_type": type(e).__name__,
                    },
                )
                return False, f"Word automation failed: {str(e)}"

        # Check file accessibility before attempting to open
        def _check_file_accessible(file_path: Path) -> Tuple[bool, Optional[str]]:
            """Check if file is accessible (not locked by another process)."""
            try:
                # Try to open file in append mode to check if it's locked
                with open(file_path, "a+b"):
                    pass
                return True, None
            except PermissionError as e:
                # File is locked - try to close it in Word first
                close_result, close_msg = _try_close_word_file(file_path)
                if close_result and close_msg:
                    # File was closed in Word, wait a moment and retry
                    import time

                    time.sleep(0.5)
                    try:
                        with open(file_path, "a+b"):
                            pass
                        return True, None  # File is now accessible
                    except Exception:
                        pass  # Still locked, continue with error

                return (
                    False,
                    f"File is locked by another process (e.g., Microsoft Word): {str(e)}",
                )
            except OSError as e:
                return False, f"File access error (possibly OneDrive syncing): {str(e)}"
            except Exception as e:
                return False, f"Unexpected error checking file: {str(e)}"

        # Retry loop
        last_error = None
        for attempt in range(max_retries):
            # Check file accessibility first
            is_accessible, access_error = await asyncio.to_thread(
                _check_file_accessible, file_path_obj
            )

            if not is_accessible:
                if attempt < max_retries - 1:
                    delay = initial_delay * (2**attempt)  # Exponential backoff
                    error_msg = (
                        f"File not accessible (attempt {attempt + 1}/{max_retries}): {access_error}. "
                        f"Retrying in {delay:.1f} seconds..."
                    )
                    logger.warning(
                        "docx_file_not_accessible_retrying",
                        extra={
                            "file": file_path,
                            "attempt": attempt + 1,
                            "max_retries": max_retries,
                            "delay": delay,
                            "slot_number": slot_number,
                            "subject": subject,
                            "error": access_error,
                        },
                    )
                    await asyncio.sleep(delay)
                    continue
                else:
                    # Final attempt failed
                    error_msg = (
                        f"Cannot access file after {max_retries} attempts: {access_error}. "
                        f"Possible causes:\n"
                        f"  1. File is open in Microsoft Word or another program - please close it\n"
                        f"  2. OneDrive is syncing the file - wait for sync to complete\n"
                        f"  3. File permissions issue - check file permissions\n"
                        f"  4. File is corrupted - try opening it manually"
                    )
                    logger.error(
                        "docx_file_not_accessible_final",
                        extra={
                            "file": file_path,
                            "slot_number": slot_number,
                            "subject": subject,
                            "error": access_error,
                        },
                    )
                    raise PermissionError(error_msg)

            # Try to open the DOCX file
            try:
                if plan_id:

                    def _create_parser():
                        with self.tracker.track_operation(plan_id, "parse_open_docx"):
                            return DOCXParser(file_path)

                    parser = await asyncio.to_thread(_create_parser)
                else:
                    parser = await asyncio.to_thread(DOCXParser, file_path)

                # Success!
                logger.info(
                    "docx_opened_successfully",
                    extra={
                        "file": file_path,
                        "attempt": attempt + 1,
                        "slot_number": slot_number,
                        "subject": subject,
                    },
                )
                return parser

            except PackageNotFoundError as e:
                last_error = e
                if attempt < max_retries - 1:
                    # Try to close file in Word before retrying
                    close_result, close_msg = await asyncio.to_thread(
                        _try_close_word_file, file_path_obj
                    )
                    if close_result and close_msg:
                        logger.info(
                            "word_file_closed_on_retry",
                            extra={
                                "file": file_path,
                                "attempt": attempt + 1,
                                "slot_number": slot_number,
                                "subject": subject,
                            },
                        )
                        # Wait a moment for file to be released
                        await asyncio.sleep(0.5)

                    delay = initial_delay * (2**attempt)
                    error_msg = (
                        f"Package not found (attempt {attempt + 1}/{max_retries}): {str(e)}. "
                        f"This usually means the file is locked or OneDrive is syncing. "
                        f"{'File closed in Word. ' if close_result and close_msg else ''}"
                        f"Retrying in {delay:.1f} seconds..."
                    )
                    logger.warning(
                        "docx_package_not_found_retrying",
                        extra={
                            "file": file_path,
                            "attempt": attempt + 1,
                            "max_retries": max_retries,
                            "delay": delay,
                            "slot_number": slot_number,
                            "subject": subject,
                            "error": str(e),
                            "word_file_closed": close_result and close_msg,
                        },
                    )
                    await asyncio.sleep(delay)
                    continue
                else:
                    # Final attempt failed
                    error_msg = (
                        f"Cannot open DOCX file after {max_retries} attempts: {str(e)}\n\n"
                        f"Possible causes:\n"
                        f"  1. File is open in Microsoft Word - please close the file\n"
                        f"  2. OneDrive is syncing - wait for sync to complete (check OneDrive icon in system tray)\n"
                        f"  3. File is corrupted - try opening it manually in Word\n"
                        f"  4. File permissions issue - check that you have read access\n\n"
                        f"File: {file_path}"
                        + (
                            f"\nSlot: {slot_number} ({subject})"
                            if slot_number and subject
                            else ""
                        )
                    )
                    logger.error(
                        "docx_package_not_found_final",
                        extra={
                            "file": file_path,
                            "slot_number": slot_number,
                            "subject": subject,
                            "error": str(e),
                        },
                    )
                    raise PackageNotFoundError(error_msg)

            except Exception as e:
                last_error = e
                if attempt < max_retries - 1:
                    delay = initial_delay * (2**attempt)
                    logger.warning(
                        "docx_open_error_retrying",
                        extra={
                            "file": file_path,
                            "attempt": attempt + 1,
                            "max_retries": max_retries,
                            "delay": delay,
                            "slot_number": slot_number,
                            "subject": subject,
                            "error": str(e),
                            "error_type": type(e).__name__,
                        },
                    )
                    await asyncio.sleep(delay)
                    continue
                else:
                    error_msg = (
                        f"Failed to open DOCX file after {max_retries} attempts: {str(e)}\n\n"
                        f"File: {file_path}"
                        + (
                            f"\nSlot: {slot_number} ({subject})"
                            if slot_number and subject
                            else ""
                        )
                    )
                    logger.error(
                        "docx_open_error_final",
                        extra={
                            "file": file_path,
                            "slot_number": slot_number,
                            "subject": subject,
                            "error": str(e),
                            "error_type": type(e).__name__,
                        },
                    )
                    raise

        # Should never reach here, but just in case
        raise Exception(
            f"Failed to open file after {max_retries} attempts: {last_error}"
        )

    def _sanitize_slot(self, slot: Dict[str, Any]) -> Dict[str, Any]:
        """Ensure slot dictionary contains no ModelPrivateAttr objects."""
        print(f"DEBUG: _sanitize_slot called for slot type {type(slot)}")
        if not isinstance(slot, dict):
            # Try to convert to dict if it's a model
            if hasattr(slot, "model_dump"):
                try:
                    slot = slot.model_dump()
                except Exception:
                    pass
            elif hasattr(slot, "dict"):
                try:
                    slot = slot.dict()
                except Exception:
                    pass
            else:
                try:
                    slot = dict(slot)
                except Exception:
                    pass

        if not isinstance(slot, dict):
            return slot  # Can't do much else

        sanitized = {}
        for k, v in slot.items():
            sanitized[k] = self._sanitize_value(v)

        return sanitized

    def _convert_single_slot_to_slots_format(
        self, lesson_json: Dict[str, Any], slot_number: int, subject: str
    ) -> Dict[str, Any]:
        """
        Convert a single-slot lesson_json (old format) to the new slots format.

        Old format: days = { "monday": { "unit_lesson": ..., "objective": ..., ... } }
        New format: days = { "monday": { "slots": [{ "slot_number": 1, "subject": "...", ... }] } }

        Args:
            lesson_json: Lesson JSON in old single-slot format
            slot_number: Slot number to assign to the slot
            subject: Subject name to assign to the slot

        Returns:
            Lesson JSON in new slots format
        """
        if not lesson_json or not isinstance(lesson_json, dict):
            return lesson_json

        # If already in slots format, return as-is
        days = lesson_json.get("days", {})
        if not days:
            return lesson_json

        # Check if already in slots format (has slots array in at least one day)
        is_already_in_slots_format = any(
            isinstance(day_data, dict)
            and "slots" in day_data
            and isinstance(day_data["slots"], list)
            for day_data in days.values()
        )

        if is_already_in_slots_format:
            return lesson_json

        # Convert to slots format
        converted_days = {}
        for day_name, day_data in days.items():
            if not isinstance(day_data, dict):
                converted_days[day_name] = day_data
                continue

            # Create a slot from the day data
            slot = {
                "slot_number": slot_number,
                "subject": subject,
            }

            # Copy all day-level fields to the slot
            for key, value in day_data.items():
                slot[key] = value

            # Add metadata if available in the day data
            metadata = lesson_json.get("metadata", {})
            if "grade" in metadata and "grade" not in slot:
                slot["grade"] = metadata.get("grade")
            if "homeroom" in metadata and "homeroom" not in slot:
                slot["homeroom"] = metadata.get("homeroom")
            if "teacher_name" in metadata and "teacher_name" not in slot:
                slot["teacher_name"] = metadata.get("teacher_name")

            # Create new format with slots array
            converted_days[day_name] = {"slots": [slot]}

        # Create converted lesson_json
        converted_lesson_json = {**lesson_json, "days": converted_days}

        return converted_lesson_json

    def _map_day_content_to_schema(
        self,
        day_content: Dict[str, str],
        slot_info: Dict[str, Any],
        day_hyperlinks: List[Dict[str, str]] = None,
    ) -> Dict[str, Any]:
        """
        Map raw extracted string content to OriginalDayPlanSingleSlot schema.

        Args:
            day_content: Dictionary of {row_label: cell_text}
            slot_info: Slot metadata (subject, grade, etc.)

        Returns:
            Dictionary representation of OriginalDayPlanSingleSlot
        """
        # Normalize keys for fuzzy matching
        normalized_content = {k.lower().strip(): v for k, v in day_content.items()}

        # Helper to find text by fuzzy key
        def find_text(keywords: List[str]) -> Optional[str]:
            for k, v in normalized_content.items():
                if any(kw in k for kw in keywords):
                    return v.strip()
            return None

        # 1. Unit/Lesson
        unit_lesson = find_text(["unit", "lesson", "module"]) or "N/A"

        # 2. Objective
        objective_text = find_text(["objective", "goal", "target", "swbat"])
        objective = (
            OriginalObjective(content_objective=objective_text)
            if objective_text
            else None
        )

        # 3. Anticipatory Set
        ant_set_text = find_text(["anticipatory", "do now", "warm up", "entry"])
        anticipatory_set = (
            OriginalAnticipatorySet(original_content=ant_set_text)
            if ant_set_text
            else None
        )

        # 4. Instruction/Activities
        instruction_text = find_text(["activity", "procedure", "instruction", "lesson"])

        # 4b. Tailored Instruction
        tailored_text = find_text(["tailored", "differentiation", "scaffold"])
        tailored_instruction = (
            OriginalTailoredInstruction(content=tailored_text)
            if tailored_text
            else None
        )

        # Fallback: If no general instruction but we have tailored,
        # use tailored as the activities context for the AI
        if not instruction_text and tailored_text:
            instruction_text = f"[Tailored Instruction provided]: {tailored_text}"

        instruction = (
            OriginalInstruction(activities=instruction_text)
            if instruction_text
            else None
        )

        # 4c. Misconceptions
        misconception_text = find_text(["misconception", "pitfall", "error"])
        misconceptions = (
            OriginalMisconceptions(content=misconception_text)
            if misconception_text
            else None
        )

        # 5. Materials
        materials_text = find_text(["material", "resource", "supplies"])
        materials = (
            OriginalMaterials(
                root=[m.strip() for m in materials_text.split("\n") if m.strip()]
            )
            if materials_text
            else None
        )

        # 6. Assessment
        assessment_text = find_text(["assessment", "evaluate", "exit ticket", "check"])
        assessment = (
            OriginalAssessment(primary_assessment=assessment_text)
            if assessment_text
            else None
        )

        # 7. Homework
        homework_text = find_text(["homework", "assignment"])
        homework = (
            OriginalHomework(original_content=homework_text) if homework_text else None
        )

        # Create model and dump to dict
        try:
            plan = OriginalDayPlanSingleSlot(
                unit_lesson=unit_lesson,
                objective=objective,
                anticipatory_set=anticipatory_set,
                instruction=instruction,
                tailored_instruction=tailored_instruction,
                misconceptions=misconceptions,
                materials=materials,
                assessment=assessment,
                homework=homework,
                hyperlinks=OriginalHyperlinks(
                    root=[
                        OriginalHyperlink(text=h["text"], url=h["url"])
                        for h in day_hyperlinks
                        if h.get("text") and h.get("url")
                    ]
                )
                if day_hyperlinks
                else None,
            )
            return plan.model_dump()
        except Exception as e:
            logger.warning(
                "original_schema_mapping_failed",
                extra={"error": str(e), "slot": slot_info.get("slot_number")},
            )
            # Fallback to dictionary version of what we have
            return {
                "unit_lesson": unit_lesson,
                "objective": objective.model_dump() if objective else None,
                "anticipatory_set": anticipatory_set.model_dump()
                if anticipatory_set
                else None,
                "instruction": instruction.model_dump() if instruction else None,
                "tailored_instruction": tailored_instruction.model_dump()
                if tailored_instruction
                else None,
                "misconceptions": misconceptions.model_dump()
                if misconceptions
                else None,
                "materials": materials.model_dump() if materials else None,
                "assessment": assessment.model_dump() if assessment else None,
                "homework": homework.model_dump() if homework else None,
                "hyperlinks": {"root": day_hyperlinks} if day_hyperlinks else None,
            }

    async def _persist_original_lesson_plan(
        self,
        slot: Dict[str, Any],
        week_of: str,
        primary_file: str,
        teacher_name: str,
        content: Dict[str, Any],
        hyperlinks: List[Dict[str, Any]],
        available_days: Optional[List[str]],
        plan_id: Optional[str] = None,
    ) -> str:
        """Helper to persist original lesson plan data to the database."""
        try:
            # 1. Map days to structured schema
            structured_days = {}
            if "table_content" in content:
                for day, day_data in content["table_content"].items():
                    day_lower = day.lower().strip()
                    if day_lower in [
                        "monday",
                        "tuesday",
                        "wednesday",
                        "thursday",
                        "friday",
                    ]:
                        # Filter hyperlinks for this day
                        day_links = [
                            h
                            for h in hyperlinks
                            if h.get("day_hint", "").lower().strip() == day_lower
                        ]
                        # Map this day's content to the schema
                        structured_days[f"{day_lower}_content"] = (
                            self._map_day_content_to_schema(
                                day_data, slot, day_hyperlinks=day_links
                            )
                        )

            # 2. Prepare the DB record
            import hashlib

            unique_str = f"{self._current_user_id}_{week_of}_{slot['slot_number']}_{slot['subject']}"
            stable_id = f"orig_{hashlib.md5(unique_str.encode()).hexdigest()}"

            original_plan_data = {
                "id": stable_id,
                "user_id": self._current_user_id,
                "week_of": week_of,
                "slot_number": slot["slot_number"],
                "subject": slot["subject"],
                "grade": slot.get("grade") or "N/A",
                "homeroom": slot.get("homeroom"),
                "source_file_path": str(primary_file),
                "source_file_name": Path(primary_file).name,
                "primary_teacher_name": teacher_name,
                "content_json": content,  # Full raw extraction
                "full_text": content.get("full_text", ""),
                "available_days": available_days,
                "status": "extracted",
                "has_no_school": not available_days or len(available_days) == 0,
                **structured_days,
            }

            # 3. Save to DB
            print(
                f"DEBUG: Saving OriginalLessonPlan for slot {slot['slot_number']} to DB..."
            )

            def _save_original():
                return self.db.create_original_lesson_plan(original_plan_data)

            res_id = await asyncio.to_thread(_save_original)
            print(f"DEBUG: OriginalLessonPlan saved with ID: {res_id}")

            return res_id

        except Exception as e:
            logger.error(
                "original_lesson_plan_persistence_failed",
                extra={"error": str(e), "slot": slot["slot_number"]},
            )
            print(f"ERROR: Persistence failed for slot {slot['slot_number']}: {e}")
            traceback.print_exc()
            return None

    async def _extract_slot_content(
        self,
        context: SlotProcessingContext,
        week_of: str,
        week_folder_path: Optional[str] = None,
        user_base_path: Optional[str] = None,
        plan_id: Optional[str] = None,
    ) -> SlotProcessingContext:
        """Phase 1: Extract content from DOCX file (sequential, file I/O).

        This method handles all file operations sequentially to prevent locking:
        - Resolve primary file
        - Open DOCX with retry logic
        - Extract content, images, hyperlinks
        - Detect available days

        Args:
            context: SlotProcessingContext with slot information
            week_of: Week date range
            week_folder_path: Optional week folder override
            user_base_path: User's base path override
            plan_id: Plan ID for progress tracking

        Returns:
            Updated SlotProcessingContext with extracted_content and available_days

        Raises:
            ValueError: If primary file not found or extraction fails
        """
        slot = context.slot

        from backend.progress import progress_tracker

        # Helper function to update progress during extraction
        def update_extraction_progress(progress: int, message: str):
            """Update progress for extraction phase."""
            if plan_id:
                # Phase 1: 0-20% of total processing
                overall_progress = int(progress * 0.20)
                progress_tracker.update(
                    plan_id, "processing", overall_progress, message
                )

        # Sanitize slot
        slot = self._sanitize_slot(slot)
        context.slot = slot

        update_extraction_progress(
            5, f"Finding lesson plan file for {slot['subject']}..."
        )

        # Resolve primary file
        if plan_id:

            def _resolve_with_tracking():
                try:
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_resolve_file",
                        metadata={
                            "slot_number": slot.get("slot_number"),
                            "subject": slot["subject"],
                            "week_of": week_of,
                        },
                    ):
                        return self._resolve_primary_file(
                            slot, week_of, week_folder_path, user_base_path
                        )
                except Exception as e:
                    logger.error(f"Error in _resolve_with_tracking: {e}")
                    raise

            primary_file = await asyncio.to_thread(_resolve_with_tracking)
        else:
            primary_file = self._resolve_primary_file(
                slot, week_of, week_folder_path, user_base_path
            )

        if not primary_file:
            # Provide helpful error message
            if week_folder_path:
                week_folder = Path(week_folder_path)
            else:
                file_mgr = get_file_manager(base_path=user_base_path)
                week_folder = file_mgr.get_week_folder(week_of)

            teacher_pattern = (
                slot.get("primary_teacher_file_pattern")
                or slot.get("primary_teacher_name")
                or (
                    f"{slot.get('primary_teacher_first_name', '')} {slot.get('primary_teacher_last_name', '')}".strip()
                    if slot.get("primary_teacher_first_name")
                    or slot.get("primary_teacher_last_name")
                    else None
                )
            )

            error_msg = (
                f"No primary teacher file found for slot {slot['slot_number']} (subject: '{slot['subject']}').\n"
                f"Week folder: {week_folder} (exists: {week_folder.exists()})\n"
            )

            if teacher_pattern:
                error_msg += f"\nSearched for pattern: '{teacher_pattern}'\n"

            context.error = error_msg
            raise ValueError(error_msg)

        context.primary_file = primary_file
        update_extraction_progress(10, "Reading lesson plan document...")

        # Open DOCX with retry logic
        try:
            parser = await self._open_docx_with_retry(
                primary_file,
                plan_id=plan_id,
                slot_number=slot.get("slot_number"),
                subject=slot.get("subject"),
            )
        except Exception as e:
            logger.error(
                "docx_parser_init_failed",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "file": primary_file,
                    "error": str(e),
                },
            )
            context.error = f"Failed to open DOCX file: {str(e)}"
            raise

        # Check for "No School" week (entire document)
        if parser.is_no_school_day():
            logger.info(
                "no_school_week_detected",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "file": primary_file,
                },
            )
            # Return special marker in context
            context.extracted_content = "__NO_SCHOOL_WEEK__"
            return context

        # Find actual slot number by subject
        primary_first = slot.get("primary_teacher_first_name", "").strip()
        primary_last = slot.get("primary_teacher_last_name", "").strip()
        teacher_name = (
            f"{primary_first} {primary_last}".strip()
            or slot.get("primary_teacher_name", "").strip()
        )

        update_extraction_progress(15, "Locating slot in document...")

        try:
            if plan_id:
                _parser = parser

                def _find_slot_with_tracking():
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_find_slot",
                        metadata={
                            "requested_slot": slot["slot_number"],
                            "subject": slot["subject"],
                        },
                    ):
                        return _parser.find_slot_by_subject(
                            slot["subject"],
                            teacher_name,
                            slot.get("homeroom"),
                            slot.get("grade"),
                        )

                actual_slot_num = await asyncio.to_thread(_find_slot_with_tracking)
            else:
                actual_slot_num = await asyncio.to_thread(
                    parser.find_slot_by_subject,
                    slot["subject"],
                    teacher_name,
                    slot.get("homeroom"),
                    slot.get("grade"),
                )
            slot_num = actual_slot_num
        except ValueError:
            # Subject not found - use requested slot as fallback
            slot_num = slot["slot_number"]

        # Extract images and hyperlinks
        update_extraction_progress(18, "Extracting images and hyperlinks...")

        try:
            if plan_id:
                _parser = parser

                def _extract_media_with_tracking():
                    images_result = None
                    hyperlinks_result = None
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_extract_images",
                        metadata={"slot_number": slot_num, "subject": slot["subject"]},
                    ):
                        images_result = _parser.extract_images_for_slot(slot_num)
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_extract_hyperlinks",
                        metadata={"slot_number": slot_num, "subject": slot["subject"]},
                    ):
                        hyperlinks_result = _parser.extract_hyperlinks_for_slot(
                            slot_num
                        )
                    return images_result, hyperlinks_result

                images, hyperlinks = await asyncio.to_thread(
                    _extract_media_with_tracking
                )
            else:
                images = await asyncio.to_thread(
                    parser.extract_images_for_slot, slot_num
                )
                hyperlinks = await asyncio.to_thread(
                    parser.extract_hyperlinks_for_slot, slot_num
                )
        except Exception as e:
            logger.error("media_extraction_failed", extra={"error": str(e)})
            images = []
            hyperlinks = []

        # Extract subject content
        update_extraction_progress(20, "Extracting lesson content...")

        if plan_id:
            _parser = parser

            def _extract_content_with_tracking():
                with self.tracker.track_operation(
                    plan_id,
                    "parse_extract_content",
                    metadata={"slot_number": slot_num, "subject": slot["subject"]},
                ):
                    # CRITICAL: strip_urls=False so we can see markdown links/raw URLs
                    # and replace them with tokens before sending to LLM.
                    return _parser.extract_subject_content_for_slot(
                        slot_num, slot["subject"], teacher_name, strip_urls=False
                    )

            content = await asyncio.to_thread(_extract_content_with_tracking)
        else:
            content = await asyncio.to_thread(
                parser.extract_subject_content_for_slot,
                slot_num,
                slot["subject"],
                teacher_name,
                strip_urls=False,
            )

        # Detect available days
        available_days = []
        if "table_content" in content:
            for day, day_content in content["table_content"].items():
                day_lower = day.lower().strip()
                if day_lower == "all days":
                    available_days = ["monday"]
                    break
                elif day_lower in [
                    "monday",
                    "tuesday",
                    "wednesday",
                    "thursday",
                    "friday",
                ]:
                    day_text = (
                        " ".join(day_content.values())
                        if isinstance(day_content, dict)
                        else str(day_content)
                    )
                    if day_text and day_text.strip().lower() not in [
                        "no school",
                        "n/a",
                        "",
                    ]:
                        available_days.append(day_lower)

        if not available_days:
            available_days = None  # None means generate all days

        # Get full text content for LLM
        primary_content = content.get("full_text", "")

        # Store metadata for later use (images, hyperlinks, etc.)
        # We'll attach these to the lesson_json after LLM transformation
        context.slot["_extracted_images"] = images
        context.slot["_extracted_hyperlinks"] = hyperlinks
        context.slot["_extracted_content_dict"] = content  # Store full content dict

        # --- NEW: Map to Schema and Persist to DB ---
        # Persistence (extraction cache)
        await self._persist_original_lesson_plan(
            slot,
            week_of,
            primary_file,
            teacher_name,
            content,
            hyperlinks,
            available_days,
            plan_id,
        )

        # Store in context
        context.extracted_content = primary_content
        context.available_days = available_days

        print(
            f"[DEBUG] _extract_slot_content (PARALLEL): Stored {len(hyperlinks)} hyperlinks "
            f"in context.slot['_extracted_hyperlinks'] for slot {slot['slot_number']}, subject {slot['subject']}"
        )
        logger.info(
            "parallel_extract_slot_content_hyperlinks_stored",
            extra={
                "slot": slot["slot_number"],
                "subject": slot["subject"],
                "hyperlinks_count": len(hyperlinks),
            },
        )

        return context

    async def _transform_slot_with_llm(
        self,
        context: SlotProcessingContext,
        week_of: str,
        provider: str,
        plan_id: Optional[str] = None,
    ) -> SlotProcessingContext:
        """Phase 2: Transform content with LLM (can run in parallel).

        This method calls the LLM service to transform extracted content.
        Multiple instances can run in parallel.

        Args:
            context: SlotProcessingContext with extracted_content
            week_of: Week date range
            provider: LLM provider name
            plan_id: Plan ID for progress tracking

        Returns:
            Updated SlotProcessingContext with lesson_json

        Raises:
            ValueError: If LLM transformation fails
        """
        if context.lesson_json:
            print(
                f"[DEBUG] _transform_slot_with_llm: Reusing existing lesson_json for slot {context.slot.get('slot_number')}"
            )
            # CRITICAL: Update metadata with slot-specific values even when reusing existing lesson_json
            # This ensures grade/homeroom are correct for each slot, not just the first one
            slot = context.slot
            if "metadata" not in context.lesson_json:
                context.lesson_json["metadata"] = {}
            
            # Force update metadata from current slot configuration
            context.lesson_json["metadata"]["slot_number"] = slot.get("slot_number")
            context.lesson_json["metadata"]["grade"] = slot.get("grade")
            context.lesson_json["metadata"]["homeroom"] = slot.get("homeroom")
            context.lesson_json["metadata"]["subject"] = slot.get("subject")
            if slot.get("start_time"):
                context.lesson_json["metadata"]["start_time"] = slot.get("start_time")
            if slot.get("end_time"):
                context.lesson_json["metadata"]["end_time"] = slot.get("end_time")
            # CRITICAL: Update primary teacher fields from current slot data
            # This ensures cached JSON gets the correct slot-specific teacher, not the old combined name
            context.lesson_json["metadata"]["primary_teacher_name"] = slot.get("primary_teacher_name")
            context.lesson_json["metadata"]["primary_teacher_first_name"] = slot.get("primary_teacher_first_name")
            context.lesson_json["metadata"]["primary_teacher_last_name"] = slot.get("primary_teacher_last_name")
            
            # CRITICAL: Also update teacher_name with combined format "Primary / Bilingual"
            # This ensures the metadata table shows "Kelsey Lang / Wilson Rodrigues" format
            try:
                combined_teacher_name = self._build_teacher_name(
                    {
                        "first_name": getattr(self, "_user_first_name", ""),
                        "last_name": getattr(self, "_user_last_name", ""),
                        "name": getattr(self, "_user_name", ""),
                    },
                    slot,
                )
                context.lesson_json["metadata"]["teacher_name"] = combined_teacher_name
            except Exception as e:
                print(f"[DEBUG] Error building combined teacher name in _transform_slot_with_llm: {e}")
                # Fallback to just primary teacher name
                context.lesson_json["metadata"]["teacher_name"] = slot.get("primary_teacher_name") or "Unknown"
            
            print(f"[DEBUG] Updated metadata for slot {slot.get('slot_number')}: grade={slot.get('grade')}, homeroom={slot.get('homeroom')}, teacher={context.lesson_json['metadata'].get('teacher_name')}")
            return context

        slot = context.slot
        total_slots = context.total_slots

        if context.error:
            # Skip if extraction failed
            return context

        if context.extracted_content == "__NO_SCHOOL_WEEK__":
            # Handle No School week
            # Get user info from instance attributes (set in process_user_week)
            user_dict = {
                "first_name": getattr(self, "_user_first_name", ""),
                "last_name": getattr(self, "_user_last_name", ""),
                "name": getattr(self, "_user_name", ""),
            }
            context.lesson_json = {
                "metadata": {
                    "teacher_name": self._build_teacher_name(user_dict, slot),
                    "grade": slot.get("grade", ""),
                    "subject": slot["subject"],
                    "week_of": week_of,
                    "homeroom": slot.get("homeroom", ""),
                    "slot_number": slot["slot_number"],
                },
                "days": {
                    day: {"unit_lesson": "No School"}
                    for day in ["monday", "tuesday", "wednesday", "thursday", "friday"]
                },
                "_images": [],
                "_hyperlinks": [],
            }
            return context

        # Track parallel processing metrics
        import time

        start_time = time.time()
        is_parallel = settings.PARALLEL_LLM_PROCESSING and total_slots > 1

        # --- PHASE 1: Link Placeholder Strategy (Pre-LLM Scrubbing) ---
        # Scrub hyperlinks from content before sending to LLM.
        # This works for both fresh extraction and cache hits.
        self._scrub_hyperlinks(context)

        # Add preservation instruction to content (immutable tokens)
        if context.link_map:
            preserve_msg = (
                f"\n\nIMPORTANT: Your input contains placeholders like {', '.join(list(context.link_map.keys())[:5])}. "
                f"These represent hyperlinks. You MUST preserve these tokens exactly in your output."
            )
            # We modify a local copy for the LLM call to keep context.extracted_content clean
            llm_primary_content = context.extracted_content + preserve_msg
        else:
            llm_primary_content = context.extracted_content

        # Create progress callback
        from backend.progress import progress_tracker

        def llm_progress_callback(stage: str, llm_progress: int, message: str):
            """Callback to update progress during LLM transformation."""
            if plan_id:
                # Map LLM progress (10-90%) to Phase 2 range (20-80% of total)
                phase2_min = 20
                phase2_max = 80
                phase2_progress = phase2_min + int(
                    (llm_progress - 10) / 80 * (phase2_max - phase2_min)
                )
                progress_tracker.update(plan_id, stage, phase2_progress, message)

        # Call LLM service
        try:
            success, lesson_json, error = await asyncio.to_thread(
                self.llm_service.transform_lesson,
                primary_content=llm_primary_content,
                grade=slot["grade"],
                subject=slot["subject"],
                week_of=week_of,
                teacher_name=None,
                homeroom=slot.get("homeroom"),
                plan_id=plan_id,
                available_days=context.available_days,
                progress_callback=llm_progress_callback,
            )

            if not success:
                context.error = f"LLM transformation failed: {error}"
                raise ValueError(context.error)

            # Ensure lesson_json is a dictionary
            if not isinstance(lesson_json, dict):
                if hasattr(lesson_json, "model_dump"):
                    lesson_json = lesson_json.model_dump(mode="python")
                elif hasattr(lesson_json, "dict"):
                    lesson_json = lesson_json.dict()
                else:
                    lesson_json = dict(lesson_json) if lesson_json else {}

            # Attach images and hyperlinks
            images = context.slot.get("_extracted_images", [])
            hyperlinks = context.slot.get("_extracted_hyperlinks", [])

            # --- PHASE 2: Link Placeholder Strategy (Post-LLM Restoration) ---
            # Restore links from placeholders in the generated JSON
            if context.link_map:
                logger.info(
                    "restoring_links_from_placeholders",
                    extra={
                        "slot": slot.get("slot_number"),
                        "link_count": len(context.link_map),
                    },
                )
                lesson_json, restored_originals = self._restore_hyperlinks(
                    lesson_json, context.link_map
                )

                # Filter out the hyperlinks that were already restored inline
                # so the renderer doesn't try to place them again (prevent duplication)
                if restored_originals:
                    initial_count = len(hyperlinks)
                    # We match against the original string (either markdown link or raw URL)
                    hyperlinks = [
                        h
                        for h in hyperlinks
                        if f"[{h['text']}]({h['url']})" not in restored_originals
                        and h["url"] not in restored_originals
                    ]
                    removed_count = initial_count - len(hyperlinks)
                    if removed_count > 0:
                        logger.info(
                            "filtering_redundant_hyperlinks",
                            extra={
                                "slot": slot.get("slot_number"),
                                "removed": removed_count,
                                "remaining": len(hyperlinks),
                            },
                        )

            # Remove usage information (tracked separately)
            lesson_json.pop("_usage", None)
            lesson_json.pop("_model", None)
            lesson_json.pop("_provider", None)

            # LOGGING: Trace hyperlink flow in parallel path
            print(
                f"[DEBUG] _transform_slot_with_llm (PARALLEL): Slot {slot.get('slot_number')}, "
                f"Retrieved {len(hyperlinks)} hyperlinks from context.slot['_extracted_hyperlinks']"
            )
            logger.info(
                "parallel_hyperlink_retrieval",
                extra={
                    "slot": slot.get("slot_number"),
                    "subject": slot.get("subject"),
                    "hyperlinks_count": len(hyperlinks),
                    "has_hyperlinks_in_context": "_extracted_hyperlinks"
                    in context.slot,
                },
            )

            # CRITICAL: Add slot and subject metadata to hyperlinks for proper filtering in multi-slot scenarios
            slot_number = slot.get("slot_number")
            subject = slot.get("subject")
            if hyperlinks:
                print(
                    f"[DEBUG] _transform_slot_with_llm (PARALLEL): Processing {len(hyperlinks)} hyperlinks "
                    f"for slot {slot_number}, subject {subject}"
                )
                for hyperlink in hyperlinks:
                    # Ensure hyperlink has slot/subject metadata for renderer filtering
                    if "_source_slot" not in hyperlink:
                        hyperlink["_source_slot"] = slot_number
                    if "_source_subject" not in hyperlink:
                        hyperlink["_source_subject"] = subject

            if images:
                lesson_json["_images"] = images
                print(
                    f"[DEBUG] _transform_slot_with_llm (PARALLEL): Added {len(images)} images to lesson_json"
                )
            if hyperlinks:
                print(
                    f"[DEBUG] _transform_slot_with_llm (PARALLEL): Adding {len(hyperlinks)} hyperlinks to lesson_json"
                )
                lesson_json["_hyperlinks"] = hyperlinks
                logger.info(
                    "parallel_hyperlinks_attached",
                    extra={
                        "slot": slot_number,
                        "subject": subject,
                        "hyperlinks_count": len(hyperlinks),
                    },
                )
            else:
                print(
                    f"[WARN] _transform_slot_with_llm (PARALLEL): No hyperlinks to attach for slot {slot_number}, subject {subject}"
                )
                logger.warning(
                    "parallel_no_hyperlinks",
                    extra={
                        "slot": slot_number,
                        "subject": subject,
                        "context_slot_keys": list(context.slot.keys()),
                    },
                )

            # CRITICAL: Set schema version for coordinate-based placement (if media present)
            # This must match the sequential path to ensure consistent rendering
            if images or hyperlinks:
                lesson_json["_media_schema_version"] = (
                    "2.0"  # Use v2.0 for coordinate placement
                )

            # CRITICAL: Update metadata with structured teacher name and formatted week
            # This must match the sequential path to ensure consistency
            if "metadata" not in lesson_json:
                lesson_json["metadata"] = {}

            # Build teacher name using structured fields with fallback
            try:
                teacher_name = self._build_teacher_name(
                    {
                        "first_name": getattr(self, "_user_first_name", ""),
                        "last_name": getattr(self, "_user_last_name", ""),
                        "name": getattr(self, "_user_name", ""),
                    },
                    slot,
                )
                lesson_json["metadata"]["teacher_name"] = teacher_name
            except Exception as e:
                print(f"DEBUG: Error in _build_teacher_name (parallel path): {e}")
                traceback.print_exc()
                # Continue with default
                lesson_json["metadata"]["teacher_name"] = "Unknown"

            # Format week dates consistently
            from backend.utils.date_formatter import format_week_dates

            lesson_json["metadata"]["week_of"] = format_week_dates(week_of)

            # Preserve slot metadata (homeroom, grade, subject, time) from slot data
            # This ensures correct values even if LLM returns incorrect/missing data
            try:
                # Unconditionally preserve metadata
                lesson_json["metadata"]["slot_number"] = slot.get("slot_number")
                lesson_json["metadata"]["homeroom"] = slot.get("homeroom")
                lesson_json["metadata"]["grade"] = slot.get("grade")
                lesson_json["metadata"]["subject"] = slot.get("subject")
                lesson_json["metadata"]["start_time"] = slot.get("start_time")
                lesson_json["metadata"]["end_time"] = slot.get("end_time")
                lesson_json["metadata"]["day_times"] = slot.get("day_times")
                
                # DIAGNOSTIC LOGGING
                if slot.get("slot_number") == 2:
                    print(f"DEBUG: Preserving Slot 2 Metadata -> Grade: '{lesson_json['metadata']['grade']}', Homeroom: '{lesson_json['metadata']['homeroom']}'")
            except Exception as e:
                print(f"DEBUG: Error copying slot metadata (parallel path): {e}")
                traceback.print_exc()

            # Replace No School days with minimal content (parity with sequential path)
            if context.no_school_days:
                for day in context.no_school_days:
                    day_lower = day.lower().strip()
                    if day_lower in lesson_json.get("days", {}):
                        lesson_json["days"][day_lower] = self._no_school_day_stub()

            # Store parallel processing metrics
            elapsed_time = (time.time() - start_time) * 1000  # Convert to ms
            context.lesson_json = lesson_json
            context.is_parallel = is_parallel
            context.parallel_slot_count = total_slots if is_parallel else None
            # Sequential time would be estimated as: elapsed_time * total_slots (if parallel)
            if is_parallel:
                context.sequential_time_ms = elapsed_time * total_slots

            # Extract rate limit errors and usage from lesson_json if available
            usage = lesson_json.get("_usage", {})
            if usage:
                # These would be set by LLM service if rate limit tracking is implemented
                context.tpm_usage = usage.get("tpm_usage")
                context.rpm_usage = usage.get("rpm_usage")

        except Exception as e:
            context.error = str(e)
            logger.error(
                "llm_transformation_failed",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "error": str(e),
                },
            )
            raise

        return context

    async def _process_slots_parallel_llm(
        self,
        contexts: List[SlotProcessingContext],
        week_of: str,
        provider: str,
        plan_id: Optional[str] = None,
    ) -> List[SlotProcessingContext]:
        """Process all slots' LLM calls in parallel with concurrency limit.

        Args:
            contexts: List of SlotProcessingContext with extracted_content
            week_of: Week date range
            provider: LLM provider name
            plan_id: Plan ID for progress tracking

        Returns:
            List of updated SlotProcessingContext with lesson_json
        """
        # Get concurrency limit from settings
        limit = settings.MAX_CONCURRENT_LLM_REQUESTS
        semaphore = asyncio.Semaphore(limit)

        async def limited_transform(
            ctx: SlotProcessingContext,
        ) -> SlotProcessingContext:
            """Execute transformation within semaphore to limit concurrency."""
            async with semaphore:
                return await self._transform_slot_with_llm(
                    ctx, week_of, provider, plan_id
                )

        # Create tasks for limited parallel execution
        tasks = [limited_transform(ctx) for ctx in contexts]

        # Execute in parallel with error handling
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Process results
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                contexts[i].error = str(result)
                logger.error(
                    "parallel_llm_slot_failed",
                    extra={
                        "slot_index": contexts[i].slot_index,
                        "slot_number": contexts[i].slot.get("slot_number"),
                        "error": str(result),
                    },
                )
            else:
                contexts[i] = result

        return contexts

    async def _process_slot(
        self,
        slot: Dict[str, Any],
        week_of: str,
        provider: str,
        week_folder_path: Optional[str] = None,
        user_base_path: Optional[str] = None,
        plan_id: Optional[str] = None,
        slot_index: int = 1,
        total_slots: int = 1,
        processing_weight: float = 0.8,
        existing_lesson_json: Optional[Dict[str, Any]] = None,
        force_ai: bool = False,
    ) -> Dict[str, Any]:
        """Process a single class slot.

        Args:
            slot: Class slot data
            week_of: Week date range
            provider: LLM provider
            week_folder_path: Optional week folder override
            user_base_path: User's base path override
            plan_id: Plan ID for progress tracking
            slot_index: Current slot index (1-based)
            total_slots: Total number of slots being processed
            processing_weight: Weight for processing phase (0-1)

        Returns:
            Lesson plan JSON
        """
        # CRITICAL: Sanitize slot input to remove any ModelPrivateAttr objects
        # This prevents "argument of type 'ModelPrivateAttr' is not iterable" errors
        slot = self._sanitize_slot(slot)

        from backend.progress import progress_tracker

        # Helper function to update progress during slot processing
        def update_slot_progress(stage: str, slot_progress: int, message: str):
            """Update progress for current slot processing."""
            if plan_id:
                # Calculate overall progress: base progress for previous slots + current slot progress
                base_progress = int(
                    (slot_index - 1) / total_slots * processing_weight * 100
                )
                current_slot_progress = int(
                    slot_progress / 100 * (1 / total_slots) * processing_weight * 100
                )
                overall_progress = base_progress + current_slot_progress
                progress_tracker.update(plan_id, stage, overall_progress, message)

        print(
            f"DEBUG: _process_slot - Resolving primary file for slot {slot['slot_number']}"
        )
        update_slot_progress(
            "processing", 10, f"Finding lesson plan file for {slot['subject']}..."
        )

        # Track file resolution
        print("DEBUG: _process_slot - calling _resolve_primary_file")
        if plan_id:

            def _resolve_with_tracking():
                try:
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_resolve_file",
                        metadata={
                            "slot_number": slot.get("slot_number"),
                            "subject": slot["subject"],
                            "week_of": week_of,
                        },
                    ):
                        return self._resolve_primary_file(
                            slot, week_of, week_folder_path, user_base_path
                        )
                except Exception as e:
                    print(f"DEBUG: Error in _resolve_with_tracking: {e}")
                    traceback.print_exc()
                    raise

            primary_file = await asyncio.to_thread(_resolve_with_tracking)
        else:
            primary_file = self._resolve_primary_file(
                slot, week_of, week_folder_path, user_base_path
            )
        print(f"DEBUG: _process_slot - Primary file resolved: {primary_file}")

        if not primary_file:
            # Provide helpful error message with guidance
            # Get week folder path for error message
            if week_folder_path:
                week_folder = Path(week_folder_path)
            else:
                file_mgr = get_file_manager(base_path=user_base_path)
                week_folder = file_mgr.get_week_folder(week_of)

            teacher_pattern = (
                slot.get("primary_teacher_file_pattern")
                or slot.get("primary_teacher_name")
                or (
                    f"{slot.get('primary_teacher_first_name', '')} {slot.get('primary_teacher_last_name', '')}".strip()
                    if slot.get("primary_teacher_first_name")
                    or slot.get("primary_teacher_last_name")
                    else None
                )
            )

            error_msg = (
                f"No primary teacher file found for slot {slot['slot_number']} (subject: '{slot['subject']}').\n"
                f"Week folder: {week_folder} (exists: {week_folder.exists()})\n"
            )

            if not teacher_pattern:
                error_msg += (
                    f"\nCONFIGURATION MISSING: Slot {slot['slot_number']} needs primary teacher information.\n"
                    f"Please configure one of the following:\n"
                    f"  - primary_teacher_file_pattern (e.g., 'Davies', 'Savoca')\n"
                    f"  - primary_teacher_name\n"
                    f"  - primary_teacher_first_name + primary_teacher_last_name\n"
                )
            else:
                error_msg += f"\nSearched for pattern: '{teacher_pattern}'\n"
                if week_folder.exists():
                    docx_files = list(week_folder.glob("*.docx"))
                    error_msg += f"Available files ({len(docx_files)} total):\n"
                    for f in docx_files[:10]:
                        error_msg += f"  - {f.name}\n"
                    if len(docx_files) > 10:
                        error_msg += f"  ... and {len(docx_files) - 10} more files\n"
                    error_msg += (
                        "\nTROUBLESHOOTING:\n"
                        "1. Check if the teacher's name appears in any filename\n"
                        "2. Verify the filename format matches the pattern\n"
                        "3. Ensure files are in the correct week folder\n"
                    )
                else:
                    error_msg += "\nWeek folder does not exist. Please create it or check the path.\n"

            raise ValueError(error_msg)

        # Parse the found file (sync I/O - run in thread)
        print("DEBUG: _process_slot - Creating DOCXParser")
        update_slot_progress("processing", 15, "Reading lesson plan document...")

        try:
            # Check file accessibility and retry if needed
            parser = await self._open_docx_with_retry(
                primary_file,
                plan_id=plan_id,
                slot_number=slot.get("slot_number"),
                subject=slot.get("subject"),
            )
            print("DEBUG: _process_slot - DOCXParser created successfully")

            # NEW: Check DB cache for already processed slot (Optimization)
            user_id = getattr(self, "_current_user_id", "unknown")
            db_record = await asyncio.to_thread(
                self.db.get_original_lesson_plan,
                user_id,
                week_of,
                slot["slot_number"],
            )

            # Cache validation: check if source file hasn't changed
            cache_hit = False
            if db_record and db_record.source_file_path == primary_file:
                path_obj = Path(primary_file)
                if path_obj.exists():
                    current_mtime = path_obj.stat().st_mtime
                    if db_record.extracted_at.timestamp() > (current_mtime + 2):
                        print(
                            f"[DEBUG] DB Cache hit for slot {slot['slot_number']} ({slot['subject']})"
                        )
                        cache_hit = True

            # If cache hit and we have existing transformed results, reuse them
            # UNLESS force_ai is requested
            if cache_hit and existing_lesson_json and not force_ai:
                existing_slot_plans = self._reconstruct_slots_from_json(
                    existing_lesson_json
                )
                slot_num = slot.get("slot_number")
                if slot_num in existing_slot_plans:
                    print(
                        f"[DEBUG] REUSING TRANSFORMED PLAN for slot {slot_num} ({slot['subject']}) - Source file unchanged."
                    )
                    update_slot_progress(
                        "processing",
                        100,
                        f"Reusing existing plan for {slot['subject']}",
                    )
                    # CRITICAL: Clean up parser before returning
                    del parser
                    import gc

                    gc.collect()

                    # Ensure metadata is up-to-date with current slot configuration even when using cache
                    # CRITICAL: Create a deep copy to prevent shared state in parallel processing
                    # If multiple slots share the same lesson_json reference, mutations in one
                    # parallel task could affect others, causing all slots to show the same teacher
                    import copy
                    cached_json = copy.deepcopy(existing_slot_plans[slot_num]["lesson_json"])
                    
                    if "metadata" not in cached_json:
                        cached_json["metadata"] = {}
                    
                    # CRITICAL: Update primary teacher fields from current slot data
                    # This ensures cached JSON gets the correct slot-specific teacher, not the old combined name
                    cached_json["metadata"]["primary_teacher_name"] = slot.get("primary_teacher_name")
                    cached_json["metadata"]["primary_teacher_first_name"] = slot.get("primary_teacher_first_name")
                    cached_json["metadata"]["primary_teacher_last_name"] = slot.get("primary_teacher_last_name")
                    
                    # CRITICAL: Also update teacher_name with combined format "Primary / Bilingual"
                    # This ensures the metadata table shows "Kelsey Lang / Wilson Rodrigues" format
                    try:
                        user_dict = {
                            "first_name": getattr(self, "_user_first_name", ""),
                            "last_name": getattr(self, "_user_last_name", ""),
                            "name": getattr(self, "_user_name", ""),
                        }
                        combined_teacher_name = self._build_teacher_name(user_dict, slot)
                        cached_json["metadata"]["teacher_name"] = combined_teacher_name
                    except Exception as e:
                        print(f"[DEBUG] Error building combined teacher name in cache reuse: {e}")
                        # Fallback to just primary teacher name
                        cached_json["metadata"]["teacher_name"] = slot.get("primary_teacher_name") or "Unknown"
                    
                    # Force update key fields from current slot config (Unconditionally)
                    cached_json["metadata"]["grade"] = slot.get("grade")
                    cached_json["metadata"]["homeroom"] = slot.get("homeroom")
                    cached_json["metadata"]["subject"] = slot.get("subject")
                    
                    if slot.get("slot_number") == 2:
                        print(f"DEBUG: Cache Inject Slot 2 -> Grade: '{cached_json['metadata']['grade']}', Homeroom: '{cached_json['metadata']['homeroom']}', Teacher: '{cached_json['metadata'].get('teacher_name')}'")
                        
                    return cached_json
            elif cache_hit and existing_lesson_json and force_ai:
                print(
                    f"[DEBUG] FORCING AI transformation for slot {slot.get('slot_number')} ({slot['subject']}) as requested."
                )

            # Check for "No School" week
            # If the entire document indicates "No School", we generate a standardized
            # "No School" lesson plan instead of trying to parse content.
            if parser.is_no_school_day():
                logger.info(
                    "no_school_week_detected",
                    extra={
                        "slot": slot["slot_number"],
                        "subject": slot["subject"],
                        "file": primary_file,
                    },
                )

                # Reconstruct user dict for teacher name builder
                user_dict = {
                    "first_name": self._user_first_name,
                    "last_name": self._user_last_name,
                    "name": self._user_name,
                }

                # Create standardized "No School" lesson JSON
                return {
                    "metadata": {
                        "teacher_name": self._build_teacher_name(user_dict, slot),
                        "grade": slot.get("grade", ""),
                        "subject": slot["subject"],
                        "week_of": week_of,
                        "homeroom": slot.get("homeroom", ""),
                        "slot_number": slot["slot_number"],
                    },
                    "days": {
                        day: {"unit_lesson": "No School"}
                        for day in [
                            "monday",
                            "tuesday",
                            "wednesday",
                            "thursday",
                            "friday",
                        ]
                    },
                    "_images": [],
                    "_hyperlinks": [],
                }

            # Early validation: Check if requested slot exists in document
            print("DEBUG: _process_slot - Validating slot availability")
            total_tables = len(parser.doc.tables)

            # Check for signature table
            has_signature = False
            if total_tables > 0:
                last_table = parser.doc.tables[-1]
                if last_table.rows and last_table.rows[0].cells:
                    first_cell = last_table.rows[0].cells[0].text.strip().lower()
                    if "signature" in first_cell or "required signatures" in first_cell:
                        has_signature = True

            # Calculate available slots
            if has_signature:
                available_slots = (total_tables - 1) // 2
            else:
                available_slots = total_tables // 2

            # Store available_slots for later validation
            # Don't validate immediately - let subject detection try first
            # If subject matches slot 1, we can auto-map slot 5 -> slot 1
            requested_slot = slot["slot_number"]
            print(
                f"DEBUG: _process_slot - Document has {available_slots} available slots, requested: {requested_slot}"
            )

            # Don't validate here - let subject detection and auto-mapping handle it
            # This allows slot 5 to be auto-mapped to slot 1 when document only has 1 slot

            # INSTRUMENTATION: Log table structure for analysis
            print("DEBUG: _process_slot - Logging table structure")
            logger.info(
                "document_table_structure",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "file": Path(primary_file).name,
                    "total_tables": total_tables,
                },
            )

            # Log each table's structure
            for table_idx, table in enumerate(parser.doc.tables):
                first_cell = ""
                if table.rows and table.rows[0].cells:
                    first_cell = table.rows[0].cells[0].text.strip()

                # Get first row text for additional context
                first_row_text = ""
                if table.rows:
                    first_row_text = " | ".join(
                        cell.text.strip()[:30] for cell in table.rows[0].cells[:5]
                    )

                logger.info(
                    "table_structure_detail",
                    extra={
                        "slot": slot["slot_number"],
                        "file": Path(primary_file).name,
                        "table_idx": table_idx,
                        "first_cell": first_cell[:100],
                        "first_row": first_row_text[:200],
                        "row_count": len(table.rows),
                        "col_count": len(table.columns) if table.rows else 0,
                    },
                )

        except Exception as e:
            logger.error(
                "docx_parser_init_failed",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "file": primary_file,
                    "error": str(e),
                    "error_type": type(e).__name__,
                },
            )
            raise

        # CRITICAL: Extract primary teacher name for disambiguation (normalize once)
        primary_first = slot.get("primary_teacher_first_name", "").strip()
        primary_last = slot.get("primary_teacher_last_name", "").strip()
        teacher_name = (
            f"{primary_first} {primary_last}".strip()
            or slot.get("primary_teacher_name", "").strip()
        )

        # CRITICAL: Find actual slot number by subject in source file
        # This handles cases where slot numbers don't align (e.g., Savoca ELA/SS is Slot 1, Wilson config is Slot 2)
        print(
            f"DEBUG: _process_slot - Finding actual slot for subject '{slot['subject']}'"
        )
        try:
            # Track slot finding operation
            if plan_id:
                # Capture parser in closure to avoid linter warnings
                _parser = parser

                def _find_slot_with_tracking():
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_find_slot",
                        metadata={
                            "requested_slot": slot["slot_number"],
                            "subject": slot["subject"],
                        },
                    ):
                        return _parser.find_slot_by_subject(
                            slot["subject"],
                            teacher_name,
                            slot.get("homeroom"),
                            slot.get("grade"),
                        )

                actual_slot_num = await asyncio.to_thread(_find_slot_with_tracking)
            else:
                actual_slot_num = await asyncio.to_thread(
                    parser.find_slot_by_subject,
                    slot["subject"],
                    teacher_name,
                    slot.get("homeroom"),
                    slot.get("grade"),
                )
            if actual_slot_num != slot["slot_number"]:
                print(
                    f"DEBUG: _process_slot - Slot mismatch! Requested slot {slot['slot_number']}, found subject in slot {actual_slot_num}"
                )

                # Determine if this is expected behavior (single-slot document)
                is_single_slot = available_slots == 1
                is_expected = is_single_slot

                # Build enhanced warning message
                warning_type = "slot_subject_mismatch"
                if is_single_slot:
                    warning_type = "slot_subject_mismatch_single_slot"
                    message = (
                        f"Slot {slot['slot_number']} requested, but document '{Path(primary_file).name}' "
                        f"has only 1 slot. Content correctly extracted from slot 1. "
                        f"This is expected behavior for single-slot documents."
                    )
                else:
                    warning_type = "slot_subject_mismatch_multi_slot"
                    message = (
                        f"Slot {slot['slot_number']} requested for '{slot['subject']}', but document "
                        f"'{Path(primary_file).name}' has '{slot['subject']}' in slot {actual_slot_num}. "
                        f"Content correctly extracted via subject-based detection. "
                        f"Consider updating slot configuration to match document structure (slot {actual_slot_num})."
                    )

                logger.warning(
                    warning_type,
                    extra={
                        "requested_slot": slot["slot_number"],
                        "actual_slot": actual_slot_num,
                        "subject": slot["subject"],
                        "file": Path(primary_file).name,
                        "available_slots": available_slots,
                        "is_single_slot": is_single_slot,
                        "is_expected": is_expected,
                        "message": message,
                        "teacher": teacher_name,
                        "grade": slot.get("grade"),
                        "homeroom": slot.get("homeroom"),
                    },
                )
            slot_num = actual_slot_num
        except ValueError as e:
            # Subject not found - check if requested slot is valid
            print(f"DEBUG: _process_slot - Subject detection failed: {e}")

            # If requested slot exceeds available slots, auto-map to slot 1
            # This handles cases where config says slot 5 but document only has 1 slot
            if slot["slot_number"] > available_slots:
                print(
                    f"DEBUG: _process_slot - Requested slot {slot['slot_number']} > available slots ({available_slots}). "
                    f"Auto-mapping to slot 1."
                )
                logger.warning(
                    "slot_auto_mapped",
                    extra={
                        "requested_slot": slot["slot_number"],
                        "available_slots": available_slots,
                        "mapped_to": 1,
                        "subject": slot["subject"],
                        "file": Path(primary_file).name,
                        "message": (
                            f"Slot {slot['slot_number']} requested, but document '{Path(primary_file).name}' "
                            f"has only {available_slots} slot(s). Auto-mapped to slot 1. "
                            f"This is expected behavior for single-slot documents."
                        ),
                        "teacher": teacher_name,
                        "grade": slot.get("grade"),
                        "homeroom": slot.get("homeroom"),
                    },
                )
                slot_num = 1  # Auto-map to first slot
            else:
                # Requested slot is valid, use it
                print(
                    f"DEBUG: _process_slot - Using requested slot {slot['slot_number']} as fallback (validated: {available_slots} slots available)"
                )
                slot_num = slot["slot_number"]

        # Extract images and hyperlinks (metadata, won't be sent to LLM) - sync I/O
        # Use slot-aware extraction (table-only, no paragraphs) to prevent cross-contamination
        print(
            f"DEBUG: _process_slot - Extracting images and hyperlinks for slot {slot_num} (subject: {slot['subject']})"
        )
        update_slot_progress("processing", 20, "Extracting content from lesson plan...")

        try:
            # Track image and hyperlink extraction
            if plan_id:
                # Capture parser in closure to avoid linter warnings
                _parser = parser

                def _extract_media_with_tracking():
                    images_result = None
                    hyperlinks_result = None
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_extract_images",
                        metadata={
                            "slot_number": slot_num,
                            "subject": slot["subject"],
                        },
                    ):
                        images_result = _parser.extract_images_for_slot(slot_num)
                    with self.tracker.track_operation(
                        plan_id,
                        "parse_extract_hyperlinks",
                        metadata={
                            "slot_number": slot_num,
                            "subject": slot["subject"],
                        },
                    ):
                        hyperlinks_result = _parser.extract_hyperlinks_for_slot(
                            slot_num
                        )
                    return images_result, hyperlinks_result

                images, hyperlinks = await asyncio.to_thread(
                    _extract_media_with_tracking
                )
            else:
                images = await asyncio.to_thread(
                    parser.extract_images_for_slot, slot_num
                )
                hyperlinks = await asyncio.to_thread(
                    parser.extract_hyperlinks_for_slot, slot_num
                )
            print(
                f"[DEBUG] _process_slot (SEQUENTIAL): Extracted {len(images)} images, {len(hyperlinks)} hyperlinks "
                f"for slot {slot['slot_number']}, subject {slot['subject']}"
            )

            # DIAGNOSTIC: Log extracted hyperlinks
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_hyperlinks_extracted(
                slot["slot_number"], slot["subject"], hyperlinks, primary_file
            )

            logger.info(
                "slot_media_extracted",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "images_count": len(images),
                    "hyperlinks_count": len(hyperlinks),
                    "extraction_mode": "slot_aware",
                    "processing_path": "sequential",
                },
            )

            # LOGGING: Log hyperlink details for debugging
            if hyperlinks:
                print(
                    f"[DEBUG] _process_slot (SEQUENTIAL): Hyperlink details: "
                    f"{[(h.get('text', '')[:30], h.get('url', '')[:50]) for h in hyperlinks[:3]]}"
                )
        except ValueError as e:
            # Structure validation failed - log and re-raise
            logger.error(
                "slot_structure_invalid",
                extra={
                    "slot": slot_num,
                    "subject": slot["subject"],
                    "file": primary_file,
                    "error": str(e),
                },
            )
            raise
        except Exception as e:
            logger.error(
                "media_extraction_failed",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "file": primary_file,
                    "error": str(e),
                    "error_type": type(e).__name__,
                },
            )
            # Continue without media - don't crash the whole process
            images = []
            hyperlinks = []

        # Check for "No School" day - skip LLM processing if detected
        print("DEBUG: _process_slot - Checking for No School day")
        if parser.is_no_school_day():
            print(
                "DEBUG: _process_slot - No School day detected, returning minimal JSON"
            )
            logger.info(
                "no_school_day_skipped",
                extra={
                    "slot": slot["slot_number"],
                    "subject": slot["subject"],
                    "file": primary_file,
                },
            )

            # Filter out all hyperlinks since entire document is No School
            if hyperlinks:
                filtered_count = len(hyperlinks)
                hyperlinks = []
                print(
                    f"DEBUG: _process_slot - Filtered {filtered_count} hyperlinks (entire document is No School)"
                )
                logger.info(
                    "hyperlinks_filtered_no_school",
                    extra={
                        "slot": slot["slot_number"],
                        "subject": slot["subject"],
                        "reason": "entire_document_no_school",
                        "filtered_count": filtered_count,
                    },
                )

            # Return minimal JSON structure indicating "No School"
            # This will be handled specially during rendering
            no_school_json = {
                "metadata": {
                    "week_of": week_of,
                    "grade": slot["grade"],
                    "subject": slot["subject"],
                    "homeroom": slot.get("homeroom"),
                    "no_school": True,
                },
                "days": {
                    day: {
                        "unit_lesson": "No School",
                        "objective": {
                            "content_objective": "No School",
                            "student_goal": "No School",
                            "wida_objective": "No School",
                        },
                        "anticipatory_set": {
                            "original_content": "No School",
                            "bilingual_bridge": "",
                        },
                        "tailored_instruction": {
                            "original_content": "No School",
                            "co_teaching_model": {},
                            "ell_support": [],
                            "special_needs_support": [],
                            "materials": [],
                        },
                        "misconceptions": {
                            "original_content": "No School",
                            "linguistic_note": {},
                        },
                        "assessment": {
                            "primary_assessment": "No School",
                            "bilingual_overlay": {},
                        },
                        "homework": {
                            "original_content": "No School",
                            "family_connection": "",
                        },
                    }
                    for day in ["monday", "tuesday", "wednesday", "thursday", "friday"]
                },
            }

            # Only include hyperlinks if there are any (should be empty after filtering)
            if hyperlinks:
                no_school_json["_hyperlinks"] = hyperlinks
                no_school_json["_media_schema_version"] = "2.0"

            # CRITICAL: Clean up parser before returning
            del parser
            import gc

            gc.collect()

            return no_school_json

        print("DEBUG: _process_slot - Extracting subject content (SLOT-AWARE)")
        update_slot_progress("processing", 25, "Parsing lesson content...")
        # Track content extraction
        if plan_id:
            # Capture parser in closure to avoid linter warnings
            _parser = parser

            def _extract_content_with_tracking():
                with self.tracker.track_operation(
                    plan_id,
                    "parse_extract_content",
                    metadata={
                        "slot_number": slot_num,
                        "subject": slot["subject"],
                    },
                ):
                    # CRITICAL: strip_urls=False so we can see markdown links/raw URLs
                    # and replace them with tokens before sending to LLM.
                    return _parser.extract_subject_content_for_slot(
                        slot_num, slot["subject"], teacher_name, strip_urls=False
                    )

            content = await asyncio.to_thread(_extract_content_with_tracking)
        else:
            content = await asyncio.to_thread(
                parser.extract_subject_content_for_slot,
                slot_num,
                slot["subject"],
                teacher_name,
                strip_urls=False,
            )
        print(
            f"DEBUG: _process_slot - Content extracted, length: {len(content.get('full_text', ''))}, slot: {slot_num}"
        )

        # Detect which days actually have content
        available_days = []
        if "table_content" in content:
            for day, day_content in content["table_content"].items():
                day_lower = day.lower().strip()
                # Check if it's "All Days" (single-lesson format)
                if day_lower == "all days":
                    # For single-lesson documents, assume it's Monday
                    available_days = ["monday"]
                    print(
                        "DEBUG: _process_slot - Single-lesson format detected, generating only Monday"
                    )
                    break
                # Check if this day has actual content (not empty, not "No School")
                elif day_lower in [
                    "monday",
                    "tuesday",
                    "wednesday",
                    "thursday",
                    "friday",
                ]:
                    # Check if day has content
                    day_text = (
                        " ".join(day_content.values())
                        if isinstance(day_content, dict)
                        else str(day_content)
                    )
                    if day_text and day_text.strip().lower() not in [
                        "no school",
                        "n/a",
                        "",
                    ]:
                        available_days.append(day_lower)
                        print(f"DEBUG: _process_slot - Day {day_lower} has content")

        # If no days detected, default to all 5 days (backward compatibility)
        if not available_days:
            available_days = None  # None means generate all days
            print(
                "DEBUG: _process_slot - No specific days detected, will generate all 5 days"
            )
        else:
            print(
                f"DEBUG: _process_slot - Generating content for days: {available_days}"
            )

        # Persistence (extraction cache)
        await self._persist_original_lesson_plan(
            slot,
            week_of,
            primary_file,
            teacher_name,
            content,
            hyperlinks,
            available_days,
            plan_id,
        )

        # CRITICAL: Extract original unit/lesson and objective content for preservation
        original_unit_lessons = {}  # day -> original unit/lesson text
        original_objectives = {}  # day -> original objective text

        if "table_content" in content:
            for day, day_content in content["table_content"].items():
                # day_content is {label: text} dict
                # Find unit/lesson row (stop after first match per day)
                day_lower = day.lower()

                for label, text in day_content.items():
                    label_lower = label.lower().strip()

                    # Match Unit/Lesson row (be specific to avoid false matches)
                    # Look for labels like "Unit, Lesson #, Module:", "Unit/Lesson:", etc.
                    if day_lower not in original_unit_lessons:
                        if (
                            (
                                "unit" in label_lower
                                and ("lesson" in label_lower or "module" in label_lower)
                            )
                            or label_lower.startswith("unit")
                            or label_lower.startswith("lesson")
                        ):
                            original_unit_lessons[day_lower] = text
                            print(
                                f"DEBUG: Extracted unit/lesson for {day}: '{text[:50]}...'"
                            )

                    # Match Objective row (be specific)
                    if day_lower not in original_objectives:
                        if label_lower.startswith("objective"):
                            original_objectives[day_lower] = text
                            print(
                                f"DEBUG: Extracted objective for {day}: '{text[:50]}...'"
                            )

            print(
                f"DEBUG: Extracted {len(original_unit_lessons)} unit/lessons, {len(original_objectives)} objectives"
            )

        # Transform with LLM - track performance
        primary_content = content["full_text"]

        # Check for No School days - we'll handle them post-LLM
        no_school_days = content.get("no_school_days", [])
        if no_school_days:
            print(f"DEBUG: _process_slot - No School days detected: {no_school_days}")
            # Don't add instruction - we'll replace the LLM output for these days later

            # Filter out hyperlinks from No School days
            # Normalize day names for comparison (day_hint might be lowercase, no_school_days might be uppercase)
            no_school_days_normalized = {day.lower().strip() for day in no_school_days}
            initial_count = len(hyperlinks)

            # Track links without day_hint for diagnostics
            links_without_day_hint = [h for h in hyperlinks if not h.get("day_hint")]

            hyperlinks = [
                h
                for h in hyperlinks
                if not h.get("day_hint")
                or h.get("day_hint", "").lower().strip()
                not in no_school_days_normalized
            ]

            filtered_count = initial_count - len(hyperlinks)

            # Log diagnostics for links preserved without day_hint
            if links_without_day_hint:
                logger.info(
                    "hyperlinks_preserved_no_day_hint",
                    extra={
                        "slot": slot["slot_number"],
                        "subject": slot["subject"],
                        "no_school_days": no_school_days,
                        "links_without_day_hint_count": len(links_without_day_hint),
                        "preserved_links": [
                            {
                                "text": h.get("text", "")[:50],
                                "section_hint": h.get("section_hint"),
                            }
                            for h in links_without_day_hint[:5]  # Limit to first 5
                        ],
                    },
                )

            if filtered_count > 0:
                logger.info(
                    "hyperlinks_filtered_no_school",
                    extra={
                        "slot": slot["slot_number"],
                        "subject": slot["subject"],
                        "no_school_days": no_school_days,
                        "filtered_count": filtered_count,
                        "remaining_count": len(hyperlinks),
                    },
                )

        # Get full text content for LLM
        primary_content = content.get("full_text", "")

        # --- PHASE 1: Link Placeholder Strategy (Pre-LLM Scrubbing) ---
        # Initialize a temporary context for scrubbing
        temp_context = SlotProcessingContext(
            slot=slot,
            slot_index=slot_index,
            total_slots=total_slots,
            extracted_content=primary_content,
        )
        self._scrub_hyperlinks(temp_context)
        # Use scrubbed content for LLM
        scrubbed_primary_content = temp_context.extracted_content

        # Add preservation instruction if placeholders were found
        if temp_context.link_map:
            preserve_msg = (
                f"\n\nIMPORTANT: Your input contains placeholders like {', '.join(list(temp_context.link_map.keys())[:5])}. "
                f"These represent hyperlinks. You MUST preserve these exact tokens in your output."
            )
            scrubbed_primary_content += preserve_msg

        print("DEBUG: _process_slot - Starting LLM transformation")
        update_slot_progress("processing", 30, "Preparing for transformation...")

        # Note: Detailed tracking is now handled by sub-operations (LLM service tracks internally)
        # No need for top-level process_slot wrapper - let sub-operations be tracked individually

        try:
            print("DEBUG: _process_slot - Calling LLM service transform_lesson")
            update_slot_progress(
                "processing", 40, f"Transforming {slot['subject']} with AI..."
            )

            # LLM call is sync HTTP request - run in thread to avoid blocking event loop
            # Create a progress callback that updates slot progress
            def llm_progress_callback(stage: str, llm_progress: int, message: str):
                """Callback to update progress during LLM transformation."""
                # Map LLM progress (0-100) to slot progress range (8-70%)
                # LLM starts at 10% and goes to 90%, map to slot range 8-70%
                slot_progress_min = 8
                slot_progress_max = 70
                slot_progress = slot_progress_min + int(
                    (llm_progress - 10)
                    / 80
                    * (slot_progress_max - slot_progress_min)
                )
                update_slot_progress(stage, slot_progress, message)

            success, lesson_json, error = await asyncio.to_thread(
                self.llm_service.transform_lesson,
                primary_content=scrubbed_primary_content,
                grade=slot["grade"],
                subject=slot["subject"],
                week_of=week_of,
                teacher_name=None,  # Will be added later
                homeroom=slot.get("homeroom"),
                plan_id=plan_id,  # Pass plan_id for detailed tracking
                available_days=available_days,  # Pass detected days
                progress_callback=llm_progress_callback,
            )

            print(
                f"DEBUG: _process_slot - LLM transform_lesson returned, success: {success}"
            )
            update_slot_progress(
                "processing", 70, "Processing transformation results..."
            )

            if not success:
                print(f"DEBUG: _process_slot - LLM transformation failed: {error}")
                raise ValueError(f"LLM transformation failed: {error}")

            # Ensure lesson_json is a dictionary immediately after LLM call
            # LLM service may return SQLModel/Pydantic objects that cause ModelPrivateAttr errors
            if not isinstance(lesson_json, dict):
                if hasattr(lesson_json, "model_dump"):
                    lesson_json = lesson_json.model_dump(mode="python")
                elif hasattr(lesson_json, "dict"):
                    lesson_json = lesson_json.dict()
                else:
                    lesson_json = dict(lesson_json) if lesson_json else {}

            # --- PHASE 2: Link Placeholder Strategy (Post-LLM Restoration) ---
            # Restore links from placeholders in the generated JSON
            if temp_context.link_map:
                print(
                    f"DEBUG: _process_slot - Restoring {len(temp_context.link_map)} links from placeholders"
                )
                lesson_json, restored_originals = self._restore_hyperlinks(
                    lesson_json, temp_context.link_map
                )

                # Filter out the hyperlinks that were already restored inline
                # so the renderer doesn't try to place them again (prevent duplication)
                if restored_originals:
                    initial_count = len(hyperlinks)
                    hyperlinks = [
                        h
                        for h in hyperlinks
                        if f"[{h['text']}]({h['url']})" not in restored_originals
                        and h["url"] not in restored_originals
                    ]
                    removed_count = initial_count - len(hyperlinks)
                    if removed_count > 0:
                        print(
                            f"DEBUG: _process_slot - Filtered {removed_count} redundant hyperlinks"
                        )

            # Remove usage information from lesson_json (already tracked by LLM service)
            lesson_json.pop("_usage", {})
            lesson_json.pop("_model", "")
            lesson_json.pop("_provider", "")

            # CRITICAL: Restore original unit/lesson and objective content
            print(
                "DEBUG: _process_slot - Restoring original unit/lesson and objective content"
            )
            for day_lower, day_data in lesson_json.get("days", {}).items():
                # Restore original unit/lesson (exact copy from input)
                if day_lower in original_unit_lessons:
                    day_data["unit_lesson"] = original_unit_lessons[day_lower]
                    print(f"DEBUG: Restored unit/lesson for {day_lower}")

                # Restore original objective content (preserve teacher's wording)
                if day_lower in original_objectives and "objective" in day_data:
                    day_data["objective"]["content_objective"] = original_objectives[
                        day_lower
                    ]
                    print(f"DEBUG: Restored objective content for {day_lower}")

            # Replace No School days with minimal content
            update_slot_progress("processing", 85, "Finalizing lesson plan...")
            if no_school_days:
                print(
                    f"DEBUG: _process_slot - Replacing {len(no_school_days)} No School days in output"
                )
                for day in no_school_days:
                    day_lower = day.lower()
                    if day_lower in lesson_json.get("days", {}):
                        lesson_json["days"][day_lower] = self._no_school_day_stub()

            # Store images and hyperlinks as metadata (underscore prefix = won't be sent to LLM)
            # CRITICAL: Add slot and subject metadata to hyperlinks for proper filtering in multi-slot scenarios
            slot_number = slot.get("slot_number")
            subject = slot.get("subject")
            if hyperlinks:
                for hyperlink in hyperlinks:
                    # Ensure hyperlink has slot/subject metadata for renderer filtering
                    if "_source_slot" not in hyperlink:
                        hyperlink["_source_slot"] = slot_number
                    if "_source_subject" not in hyperlink:
                        hyperlink["_source_subject"] = subject

            if images:
                lesson_json["_images"] = images
            if hyperlinks:
                print(
                    f"[DEBUG] _process_slot (SEQUENTIAL): Adding {len(hyperlinks)} hyperlinks to lesson_json "
                    f"for slot {slot.get('slot_number')}, subject {slot.get('subject')}"
                )
                lesson_json["_hyperlinks"] = hyperlinks
                logger.info(
                    "sequential_hyperlinks_attached",
                    extra={
                        "slot": slot.get("slot_number"),
                        "subject": slot.get("subject"),
                        "hyperlinks_count": len(hyperlinks),
                    },
                )

                # DIAGNOSTIC: Log lesson JSON with hyperlinks
                from tools.diagnostic_logger import get_diagnostic_logger

                diag = get_diagnostic_logger()
                diag.log_lesson_json_created(
                    slot["slot_number"], slot["subject"], lesson_json
                )
            else:
                print(
                    f"[WARN] _process_slot (SEQUENTIAL): No hyperlinks extracted for slot {slot.get('slot_number')}, "
                    f"subject {slot.get('subject')}!"
                )
                logger.warning(
                    "sequential_no_hyperlinks",
                    extra={
                        "slot": slot.get("slot_number"),
                        "subject": slot.get("subject"),
                    },
                )

            # Set schema version for coordinate-based placement (if media present)
            if images or hyperlinks:
                lesson_json["_media_schema_version"] = (
                    "2.0"  # Use v2.0 for coordinate placement
                )

            # Update metadata with structured teacher name and formatted week
            # Ensure lesson_json is a dictionary before using 'in' operator
            if not isinstance(lesson_json, dict):
                if hasattr(lesson_json, "model_dump"):
                    lesson_json = lesson_json.model_dump()
                elif hasattr(lesson_json, "dict"):
                    lesson_json = lesson_json.dict()
                else:
                    lesson_json = dict(lesson_json) if lesson_json else {}

            if "metadata" not in lesson_json:
                lesson_json["metadata"] = {}

            # Build teacher name using structured fields with fallback
            try:
                teacher_name = self._build_teacher_name(
                    {
                        "first_name": getattr(self, "_user_first_name", ""),
                        "last_name": getattr(self, "_user_last_name", ""),
                        "name": getattr(self, "_user_name", ""),
                    },
                    slot,
                )
                lesson_json["metadata"]["teacher_name"] = teacher_name
            except Exception as e:
                print(f"DEBUG: Error in _build_teacher_name: {e}")
                traceback.print_exc()
                # Continue with default
                lesson_json["metadata"]["teacher_name"] = "Unknown"

            # Format week dates consistently
            lesson_json["metadata"]["week_of"] = format_week_dates(week_of)

            # Preserve slot metadata (homeroom, grade, subject, time, primary teacher) from slot data
            # This ensures correct values even if LLM returns incorrect/missing data
            # Use .get() instead of 'in' to avoid ModelPrivateAttr issues
            try:
                # Unconditionally preserve metadata
                lesson_json["metadata"]["slot_number"] = slot.get("slot_number")
                lesson_json["metadata"]["homeroom"] = slot.get("homeroom")
                lesson_json["metadata"]["grade"] = slot.get("grade")
                lesson_json["metadata"]["subject"] = slot.get("subject")
                lesson_json["metadata"]["start_time"] = slot.get("start_time")
                lesson_json["metadata"]["end_time"] = slot.get("end_time")
                lesson_json["metadata"]["day_times"] = slot.get("day_times")
                # Preserve primary teacher fields for slot-specific teacher display
                lesson_json["metadata"]["primary_teacher_name"] = slot.get("primary_teacher_name")
                lesson_json["metadata"]["primary_teacher_first_name"] = slot.get("primary_teacher_first_name")
                lesson_json["metadata"]["primary_teacher_last_name"] = slot.get("primary_teacher_last_name")

                # DIAGNOSTIC LOGGING
                if slot.get("slot_number") == 2:
                    print(f"DEBUG: Preserving Slot 2 Metadata -> Grade: '{lesson_json['metadata']['grade']}', Homeroom: '{lesson_json['metadata']['homeroom']}'")
            except Exception as e:
                print(f"DEBUG: Error copying slot metadata: {e}")
                traceback.print_exc()

            try:
                normalize_objectives_in_lesson(lesson_json)
            except Exception as e:
                print(f"DEBUG: Error in normalize_objectives_in_lesson: {e}")
                traceback.print_exc()

            # CRITICAL: Clean up parser to release file handle
            # This prevents "Package not found" errors when multiple slots use the same file
            del parser
            import gc

            gc.collect()

            return lesson_json

        except Exception:
            # Error tracking is handled by sub-operations
            raise

    def _resolve_primary_file(
        self,
        slot: Dict[str, Any],
        week_of: str,
        week_folder_path: Optional[str] = None,
        user_base_path: Optional[str] = None,
    ) -> Optional[str]:
        """
        Resolve primary teacher file using hybrid approach.

        Priority:
        1. Direct path if absolute and exists
        2. Direct path relative to week folder if exists
        3. Pattern-based search in week folder
        4. Fallback to input/ directory for testing

        Args:
            slot: Class slot data
            week_of: Week date range
            week_folder_path: Optional week folder override
            user_base_path: User's base path override

        Returns:
            Path to primary teacher file or None
        """
        from pathlib import Path

        # Get week folder
        if week_folder_path:
            week_folder = Path(week_folder_path)
        else:
            # Use user's base path override if available
            file_mgr = get_file_manager(base_path=user_base_path)
            week_folder = file_mgr.get_week_folder(week_of)

        logger.debug(
            "primary_file_resolve_start",
            extra={
                "slot_number": slot.get("slot_number"),
                "week_folder": str(week_folder),
                "week_folder_exists": week_folder.exists(),
                "primary_teacher_file": slot.get("primary_teacher_file"),
                "primary_teacher_name": slot.get("primary_teacher_name"),
                "primary_teacher_file_pattern": slot.get(
                    "primary_teacher_file_pattern"
                ),
            },
        )

        # Priority 1: Absolute path
        if slot.get("primary_teacher_file"):
            file_path = Path(slot["primary_teacher_file"])
            if file_path.is_absolute() and file_path.exists():
                logger.debug(
                    "primary_file_resolve_absolute",
                    extra={
                        "slot_number": slot.get("slot_number"),
                        "path": str(file_path),
                    },
                )
                return str(file_path)

            # Priority 2: Relative to week folder
            if week_folder.exists():
                relative_path = week_folder / file_path.name
                if relative_path.exists():
                    logger.debug(
                        "primary_file_resolve_relative",
                        extra={
                            "slot_number": slot.get("slot_number"),
                            "path": str(relative_path),
                        },
                    )
                    return str(relative_path)

        # Priority 3: Pattern-based search in week folder
        # Try pattern first, but if it doesn't match, fall back to last name

        def _safe_get(d, key, default=None):
            val = d.get(key, default)
            # Check if value is ModelPrivateAttr (which is not iterable and breaks boolean checks)
            if hasattr(val, "__class__") and "ModelPrivateAttr" in str(type(val)):
                return default
            return val

        teacher_pattern = _safe_get(slot, "primary_teacher_file_pattern")
        primary_name = _safe_get(slot, "primary_teacher_name")
        first_name = _safe_get(slot, "primary_teacher_first_name", "")
        last_name = _safe_get(slot, "primary_teacher_last_name", "")

        # If pattern is set, try it first
        if teacher_pattern:
            print(f"DEBUG: Using primary_teacher_file_pattern: '{teacher_pattern}'")
        # If no pattern, try primary_teacher_name
        elif primary_name:
            teacher_pattern = primary_name
            print(f"DEBUG: Using primary_teacher_name: '{teacher_pattern}'")
        # Fallback to first_name + last_name
        elif first_name or last_name:
            teacher_pattern = f"{first_name} {last_name}".strip()
            print(f"DEBUG: Using first_name + last_name: '{teacher_pattern}'")
        else:
            teacher_pattern = None
            print("DEBUG: No teacher pattern available")
        if teacher_pattern and week_folder.exists():
            print(f"DEBUG: Searching for pattern '{teacher_pattern}' in week folder")
            logger.debug(
                "primary_file_resolve_pattern",
                extra={
                    "slot_number": slot.get("slot_number"),
                    "pattern": teacher_pattern,
                },
            )

            # List available files for debugging
            if week_folder.exists():
                docx_files = list(week_folder.glob("*.docx"))
                file_names = [f.name for f in docx_files[:10]]
                print(
                    f"DEBUG: Found {len(docx_files)} DOCX files in week folder: {file_names}"
                )
                logger.debug(
                    "primary_file_candidates",
                    extra={
                        "slot_number": slot.get("slot_number"),
                        "candidate_count": len(docx_files),
                        "candidates": file_names,
                    },
                )

            file_mgr = get_file_manager(base_path=user_base_path)

            # Try pattern first
            patterns_to_try = []
            if teacher_pattern:
                patterns_to_try.append(teacher_pattern)

            # If pattern doesn't match and we have last_name, try that
            if last_name and last_name not in (patterns_to_try or []):
                patterns_to_try.append(last_name)

            # Also try full name if available
            if primary_name and primary_name not in patterns_to_try:
                patterns_to_try.append(primary_name)
            elif (
                first_name and last_name
            ) and f"{first_name} {last_name}".strip() not in patterns_to_try:
                patterns_to_try.append(f"{first_name} {last_name}".strip())

            found = None
            for pattern in patterns_to_try:
                print(
                    f"DEBUG: Trying pattern '{pattern}' for subject '{slot['subject']}'"
                )
                found = file_mgr.find_primary_teacher_file(
                    week_folder, pattern, slot["subject"]
                )
                if found:
                    print(
                        f"DEBUG: Found primary file using pattern '{pattern}': {found}"
                    )
                    logger.debug(
                        "primary_file_resolve_pattern_success",
                        extra={
                            "slot_number": slot.get("slot_number"),
                            "pattern": pattern,
                            "path": found,
                        },
                    )
                    return found

            # None of the patterns matched
            print(f"DEBUG: None of the patterns matched: {patterns_to_try}")
            # List all files that weren't skipped for debugging
            all_files = list(week_folder.glob("*.docx"))
            skipped_files = [
                f.name for f in all_files if file_mgr._should_skip_file(f.name)
            ]
            available_files = [
                f.name for f in all_files if not file_mgr._should_skip_file(f.name)
            ]
            print(f"DEBUG: Available files (not skipped): {available_files}")
            if skipped_files:
                print(f"DEBUG: Skipped files: {skipped_files}")
            logger.debug(
                "primary_file_resolve_pattern_failed",
                extra={
                    "slot_number": slot.get("slot_number"),
                    "patterns_tried": patterns_to_try,
                    "available_files": available_files,
                    "skipped_files": skipped_files,
                },
            )
        else:
            if not teacher_pattern:
                print(
                    f"DEBUG: Slot {slot.get('slot_number')} has no teacher pattern or name configured"
                )
                print(
                    f"DEBUG: primary_teacher_file_pattern: {slot.get('primary_teacher_file_pattern')}"
                )
                print(
                    f"DEBUG: primary_teacher_name: {slot.get('primary_teacher_name')}"
                )
                logger.debug(
                    "primary_file_no_pattern",
                    extra={
                        "slot_number": slot.get("slot_number"),
                        "primary_teacher_file_pattern": slot.get(
                            "primary_teacher_file_pattern"
                        ),
                        "primary_teacher_name": slot.get("primary_teacher_name"),
                    },
                )
            if not week_folder.exists():
                print(f"DEBUG: Week folder does not exist: {week_folder}")
                logger.debug(
                    "primary_file_week_folder_missing",
                    extra={
                        "slot_number": slot.get("slot_number"),
                        "week_folder": str(week_folder),
                    },
                )

        # Priority 4: Fallback to input/ for testing
        if slot.get("primary_teacher_file"):
            fallback = Path("input") / Path(slot["primary_teacher_file"]).name
            if fallback.exists():
                logger.debug(
                    "primary_file_resolve_fallback",
                    extra={
                        "slot_number": slot.get("slot_number"),
                        "path": str(fallback),
                    },
                )
                return str(fallback)

        # Enhanced error logging with all diagnostic info
        print(f"DEBUG: PRIMARY FILE NOT FOUND for slot {slot.get('slot_number')}")
        print(f"DEBUG: Week folder: {week_folder} (exists: {week_folder.exists()})")
        print("DEBUG: Slot config:")
        print(f"  - primary_teacher_file: {slot.get('primary_teacher_file')}")
        print(
            f"  - primary_teacher_file_pattern: {slot.get('primary_teacher_file_pattern')}"
        )
        print(f"  - primary_teacher_name: {slot.get('primary_teacher_name')}")
        print(
            f"  - primary_teacher_first_name: {slot.get('primary_teacher_first_name')}"
        )
        print(f"  - primary_teacher_last_name: {slot.get('primary_teacher_last_name')}")
        print(f"  - subject: {slot.get('subject')}")

        if week_folder.exists():
            docx_files = list(week_folder.glob("*.docx"))
            print(f"DEBUG: Available files in week folder ({len(docx_files)} total):")
            for f in docx_files[:20]:
                print(f"  - {f.name}")

        logger.warning(
            "primary_file_not_found",
            extra={
                "slot_number": slot.get("slot_number"),
                "week_folder": str(week_folder),
                "week_folder_exists": week_folder.exists(),
                "primary_teacher_file": slot.get("primary_teacher_file"),
                "primary_teacher_file_pattern": slot.get(
                    "primary_teacher_file_pattern"
                ),
                "primary_teacher_name": slot.get("primary_teacher_name"),
                "available_files": [
                    f.name for f in list(week_folder.glob("*.docx"))[:20]
                ]
                if week_folder.exists()
                else [],
            },
        )
        return None

    def _combine_lessons(
        self,
        user: Dict[str, Any],
        lessons: List[Dict[str, Any]],
        week_of: str,
        generated_at: datetime,
        plan_id: Optional[str] = None,
    ) -> str:
        """
        Combine multiple lessons into a single DOCX using JSON merging.

        Args:
            user: User data
            lessons: List of lesson dicts with slot_number, subject, lesson_json
            week_of: Week date range
            generated_at: Generation timestamp

        Returns:
            Path to output DOCX file
        """
        # Sort lessons by slot number
        lessons.sort(key=lambda x: x["slot_number"])

        logger.info(
            "batch_combining_lessons",
            extra={"lesson_count": len(lessons), "user_id": user["id"]},
        )

        # Track combine lessons operation - wrap the entire method logic
        if plan_id:
            with self.tracker.track_operation(
                plan_id,
                "render_combine_lessons",
                metadata={
                    "slot_count": len(lessons),
                    "consolidated": len(lessons) > 1,
                },
            ):
                return self._combine_lessons_impl(
                    user, lessons, week_of, generated_at, plan_id
                )
        else:
            return self._combine_lessons_impl(
                user, lessons, week_of, generated_at, plan_id
            )

    def _combine_lessons_impl(
        self,
        user: Dict[str, Any],
        lessons: List[Dict[str, Any]],
        week_of: str,
        generated_at: datetime,
        plan_id: Optional[str] = None,
    ) -> str:
        """Internal implementation of _combine_lessons (without tracking wrapper)."""
        from tools.diagnostic_logger import finalize_diagnostics
        from tools.json_merger import (
            get_merge_summary,
            merge_lesson_jsons,
            validate_merged_json,
        )

        def _safe_finalize():
            """Safely finalize diagnostics, never crash."""
            try:
                finalize_diagnostics()
            except Exception:
                logger.warning("diagnostic_finalize_failed", exc_info=True)

        # Merge all lesson JSONs into a single multi-slot structure
        merged_json = merge_lesson_jsons(lessons)

        # Validate merged structure
        is_valid, error_msg = validate_merged_json(merged_json)
        if not is_valid:
            raise ValueError(f"Merged JSON validation failed: {error_msg}")

        # Log merge summary
        logger.debug(
            "batch_merge_summary", extra={"summary": get_merge_summary(merged_json)}
        )

        # CRITICAL: Enrich merged_json with day-specific start_time/end_time from schedule
        # This ensures slots are ordered correctly for each day based on actual schedule times
        # (subjects/grades/rooms vary by day, so we need day-specific times for correct ordering)
        from backend.api import enrich_lesson_json_with_times
        user_id = user.get("id") or user.get("user_id")
        if user_id:
            enrich_lesson_json_with_times(merged_json, user_id)

        # Add user info to metadata
        merged_json["metadata"]["user_name"] = user["name"]

        # Get output path using user's base path
        file_mgr = get_file_manager(base_path=user.get("base_path_override"))
        week_folder = file_mgr.get_week_folder(week_of)

        # Determine output filename based on number of slots
        if len(lessons) > 1:
            # Multi-slot: Use consolidated "Weekly" filename with timestamp
            import re
            from datetime import datetime

            # Extract week number from folder name if it matches pattern "YY W##" (e.g., "26 W02")
            # This ensures filename matches the folder name
            week_num = None
            folder_name = week_folder.name
            week_match = re.search(
                r"(\d{2})\s*W\s*(\d{1,2})\b", folder_name, re.IGNORECASE
            )
            if week_match:
                week_num = int(week_match.group(2))
                logger.debug(
                    "week_number_extracted_from_folder",
                    extra={"folder_name": folder_name, "week_num": week_num},
                )

            # Fallback to calculating from week_of if folder name doesn't match pattern
            if week_num is None:
                week_num = self._get_week_num(week_of)
                logger.debug(
                    "week_number_calculated_from_date",
                    extra={"week_of": week_of, "week_num": week_num},
                )

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{user['name'].replace(' ', '_')}_Weekly_W{week_num:02d}_{week_of.replace('/', '-')}_{timestamp}.docx"
            output_path = str(week_folder / filename)

            # Add consolidated metadata
            merged_json["metadata"]["total_slots"] = len(lessons)

            # Extract unique teachers and subjects
            teachers = set()
            subjects = set()
            for lesson in lessons:
                if (
                    lesson.get("lesson_json", {})
                    .get("metadata", {})
                    .get("teacher_name")
                ):
                    teachers.add(lesson["lesson_json"]["metadata"]["teacher_name"])
                if lesson.get("subject"):
                    subjects.add(lesson["subject"])

            # Update metadata for multi-slot display - join names instead of "Multiple..."
            if len(teachers) > 1:
                merged_json["metadata"]["teacher_name"] = " / ".join(sorted(teachers))
            if len(subjects) > 1:
                merged_json["metadata"]["subject"] = " / ".join(sorted(subjects))
        else:
            # Single slot: Use timestamped filename format to prevent overwrites
            output_path = file_mgr.get_output_path_with_timestamp(
                week_folder, user["name"], week_of
            )

        # Render to DOCX - use user-specific template if available, otherwise default
        template_path = (
            user.get("template_path") or "input/Lesson Plan Template SY'25-26.docx"
        )
        if not Path(template_path).exists():
            logger.warning(
                "user_template_not_found",
                extra={"user_id": user.get("id"), "template_path": template_path},
            )
            # Fallback to default template
            template_path = "input/Lesson Plan Template SY'25-26.docx"

        # Verify template exists before creating renderer
        if not Path(template_path).exists():
            raise FileNotFoundError(
                f"Template file not found: {template_path}. "
                f"Please ensure the template exists at this path."
            )

        renderer = DOCXRenderer(template_path)

        if len(lessons) == 1:
            # Single slot: render directly
            # CRITICAL: Add slot metadata to the lesson JSON before rendering
            lesson = lessons[0]
            slot_num = lesson["slot_number"]
            subject = lesson["subject"]
            lesson_json = lesson["lesson_json"]

            # Ensure lesson_json is a dictionary before using 'in' operator
            if not isinstance(lesson_json, dict):
                if hasattr(lesson_json, "model_dump"):
                    lesson_json = lesson_json.model_dump()
                elif hasattr(lesson_json, "dict"):
                    lesson_json = lesson_json.dict()
                else:
                    lesson_json = dict(lesson_json) if lesson_json else {}

            # Add slot metadata to lesson JSON metadata
            if "metadata" not in lesson_json:
                lesson_json["metadata"] = {}
            lesson_json["metadata"]["slot_number"] = slot_num
            lesson_json["metadata"]["subject"] = subject

            # Add slot metadata to hyperlinks and images

            if "_hyperlinks" in lesson_json:
                for link in lesson_json.get("_hyperlinks", []):
                    if isinstance(link, dict):
                        link["_source_slot"] = slot_num
                        link["_source_subject"] = subject

            if "_images" in lesson_json:
                for image in lesson_json.get("_images", []):
                    if isinstance(image, dict):
                        image["_source_slot"] = slot_num
                        image["_source_subject"] = subject

            logger.info(
                "batch_render_single_slot",
                extra={"output_file": Path(output_path).name, "week_of": week_of},
            )

            # DIAGNOSTIC: Log before rendering
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_before_render(slot_num, subject, lesson_json, "single_slot")

            # Verify output directory exists and is writable before rendering
            output_dir = Path(output_path).parent
            try:
                output_dir.mkdir(parents=True, exist_ok=True)
                # Test write permissions by creating a temp file
                test_file = output_dir / ".write_test"
                try:
                    test_file.touch()
                    test_file.unlink()
                except Exception as perm_error:
                    raise PermissionError(
                        f"Cannot write to output directory '{output_dir}': {perm_error}"
                    )
            except Exception as dir_error:
                error_msg = f"Failed to create or access output directory '{output_dir}': {str(dir_error)}"
                print(f"ERROR: {error_msg}")
                traceback.print_exc()
                raise RuntimeError(error_msg) from dir_error

            # Pass plan_id for detailed tracking
            try:
                render_success = renderer.render(
                    lesson_json, output_path, plan_id=plan_id
                )
                if not render_success:
                    raise RuntimeError(
                        "Renderer returned False - rendering failed silently. Check logs for details."
                    )
            except Exception as e:
                error_msg = (
                    f"Renderer failed to generate file '{output_path}': {str(e)}"
                )
                print(f"ERROR: {error_msg}")
                traceback.print_exc()
                raise RuntimeError(error_msg) from e

            if not Path(output_path).exists():
                error_msg = f"Renderer failed to create output file: {output_path}. File does not exist after render() call."
                print(f"ERROR: {error_msg}")
                # Check if parent directory exists
                output_dir = Path(output_path).parent
                if not output_dir.exists():
                    error_msg += f" Parent directory does not exist: {output_dir}"
                else:
                    error_msg += f" Parent directory exists: {output_dir}"
                raise FileNotFoundError(error_msg)

            # Modify the existing signature table in the rendered document
            # (add date and signature image/name)
            try:
                doc = Document(output_path)
            except Exception as e:
                error_msg = f"Failed to open rendered file '{output_path}': {str(e)}"
                print(f"ERROR: {error_msg}")
                traceback.print_exc()
                raise FileNotFoundError(error_msg)

            signature_image_path = user.get("signature_image_path")
            if signature_image_path and not signature_image_path.strip():
                signature_image_path = None

            # Special case: Wilson Rodrigues should use PNG signature
            # Check if database path exists, if not (or not set), try default path
            user_name = user.get("name", "")
            if user_name and "Wilson Rodrigues" in user_name:
                # Verify database path exists, or try default if not set/invalid
                if signature_image_path and not Path(signature_image_path).exists():
                    logger.warning(
                        "wilson_signature_db_path_not_found",
                        extra={
                            "user": user_name,
                            "db_path": signature_image_path,
                            "falling_back_to_default": True,
                        },
                    )
                    signature_image_path = None  # Clear invalid path to try default

                # If no valid path from database, try default paths
                if not signature_image_path:
                    # Try multiple possible paths (case-insensitive matching)
                    possible_paths = [
                        r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.PNG",
                        r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.png",
                        r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.Png",
                    ]
                    for sig_path in possible_paths:
                        if Path(sig_path).exists():
                            signature_image_path = sig_path
                            logger.info(
                                "using_default_wilson_signature",
                                extra={"path": sig_path, "user": user_name},
                            )
                            break
                    else:
                        logger.warning(
                            "wilson_signature_not_found",
                            extra={
                                "user": user_name,
                                "attempted_paths": possible_paths,
                            },
                        )

            try:
                self._modify_existing_signature_table(
                    doc, generated_at, signature_image_path, user.get("name")
                )
                doc.save(output_path)
            except Exception as e:
                error_msg = f"Failed to modify/save document '{output_path}': {str(e)}"
                print(f"ERROR: {error_msg}")
                traceback.print_exc()
                raise

            # JSON file writing removed - using database lesson_json column as single source of truth
            # Ensure lesson_json is saved to database via update_weekly_plan() call

            # Generate objectives DOCX file
            try:
                from backend.services.objectives_printer import ObjectivesPrinter

                objectives_printer = ObjectivesPrinter()
                objectives_output_path = Path(output_path).with_name(
                    Path(output_path).stem + "_objectives.docx"
                )
                objectives_printer.generate_docx(
                    lesson_json,
                    str(objectives_output_path),
                    user_name=user.get("name"),
                    week_of=week_of,
                )
                logger.info(
                    "batch_objectives_docx_generated",
                    extra={"objectives_path": str(objectives_output_path)},
                )
            except Exception as e:
                logger.warning(
                    "batch_objectives_docx_failed",
                    extra={"error": str(e), "error_type": type(e).__name__},
                    exc_info=True,
                )
                # Don't fail the entire process if objectives generation fails

            # Generate objectives PDF and HTML files
            # Write debug marker to confirm this code is reached
            debug_marker = (
                Path(output_path).parent
                / f"{Path(output_path).stem}_objectives_debug.txt"
            )
            try:
                with open(debug_marker, "w", encoding="utf-8") as f:
                    f.write(f"Objectives generation started at {datetime.now()}\n")
                    f.write(f"Output path: {output_path}\n")
            except Exception:
                pass  # Don't fail if we can't write debug file

            try:
                from backend.services.objectives_pdf_generator import (
                    generate_objectives_pdf,
                )

                objectives_pdf_path = Path(output_path).with_name(
                    Path(output_path).stem + "_objectives.pdf"
                )
                print(f"DEBUG: Generating objectives PDF at: {objectives_pdf_path}")
                try:
                    with open(debug_marker, "a", encoding="utf-8") as f:
                        f.write(f"PDF path: {objectives_pdf_path}\n")
                        f.write("About to call generate_objectives_pdf\n")
                        f.write("Sanitizing lesson_json before PDF generation\n")
                except Exception:
                    pass

                # CRITICAL: Enrich single-slot lesson_json with day-specific start_time/end_time from schedule
                # This ensures slots are ordered correctly for each day based on actual schedule times
                # (subjects/grades/rooms vary by day, so we need day-specific times for correct ordering)
                from backend.api import enrich_lesson_json_with_times
                user_id = user.get("id") or user.get("user_id")
                if user_id:
                    enrich_lesson_json_with_times(lesson_json, user_id)

                # CRITICAL: Sanitize lesson_json to remove ModelPrivateAttr objects before PDF generation
                lesson_json_for_pdf = self._sanitize_value(lesson_json)

                logger.info(
                    "batch_objectives_pdf_html_starting",
                    extra={"objectives_pdf_path": str(objectives_pdf_path)},
                )
                # Generate PDF with keep_html=True to also create the HTML file
                generate_objectives_pdf(
                    lesson_json_for_pdf,
                    str(objectives_pdf_path),
                    user_name=user.get("name"),
                    keep_html=True,  # Keep HTML file (it will be created with .html extension)
                )
                # The HTML file will be created with the same name but .html extension
                objectives_html_path = objectives_pdf_path.with_suffix(".html")

                # Verify files were created
                if not objectives_pdf_path.exists():
                    logger.error(
                        "batch_objectives_pdf_not_created",
                        extra={"expected_path": str(objectives_pdf_path)},
                    )
                if not objectives_html_path.exists():
                    logger.error(
                        "batch_objectives_html_not_created",
                        extra={"expected_path": str(objectives_html_path)},
                    )

                print(
                    f"DEBUG: Objectives PDF generated: {objectives_pdf_path.exists()}"
                )
                print(
                    f"DEBUG: Objectives HTML generated: {objectives_html_path.exists()}"
                )
                try:
                    with open(debug_marker, "a", encoding="utf-8") as f:
                        f.write(f"PDF exists: {objectives_pdf_path.exists()}\n")
                        f.write(f"HTML exists: {objectives_html_path.exists()}\n")
                except Exception:
                    pass
                logger.info(
                    "batch_objectives_pdf_html_generated",
                    extra={
                        "objectives_pdf_path": str(objectives_pdf_path),
                        "objectives_html_path": str(objectives_html_path),
                        "pdf_exists": objectives_pdf_path.exists(),
                        "html_exists": objectives_html_path.exists(),
                    },
                )

                # Generate sentence frames PDF/HTML for the same lesson JSON
                try:
                    print("DEBUG: Attempting to import sentence_frames_pdf_generator")
                    from backend.services.sentence_frames_pdf_generator import (
                        generate_sentence_frames_pdf,
                    )

                    print("DEBUG: Successfully imported sentence_frames_pdf_generator")

                    # Use the same directory as the output_path to ensure files are in the correct folder
                    output_dir = Path(output_path).parent
                    output_stem = Path(output_path).stem
                    sentence_frames_pdf_path = (
                        output_dir / f"{output_stem}_sentence_frames.pdf"
                    )

                    logger.info(
                        "batch_sentence_frames_pdf_html_starting",
                        extra={
                            "sentence_frames_pdf_path": str(sentence_frames_pdf_path),
                            "output_path": str(output_path),
                            "output_dir": str(output_dir),
                        },
                    )

                    print(f"DEBUG: Output path: {output_path}")
                    print(f"DEBUG: Output directory: {output_dir}")
                    print(
                        f"DEBUG: Sentence frames PDF path: {sentence_frames_pdf_path}"
                    )

                    # Check if sentence_frames exist in the data
                    # Use the original lesson_json, not the sanitized one, to ensure we don't miss frames
                    has_frames = False
                    frames_count = 0
                    if "days" in lesson_json:
                        for day_name, day in lesson_json["days"].items():
                            if not isinstance(day, dict):
                                continue
                            # Check day-level frames
                            day_frames = day.get("sentence_frames")
                            if (
                                day_frames
                                and isinstance(day_frames, list)
                                and len(day_frames) > 0
                            ):
                                has_frames = True
                                frames_count += len(day_frames)
                            # Check slot-level frames
                            slots = day.get("slots", [])
                            if isinstance(slots, list):
                                for slot in slots:
                                    if not isinstance(slot, dict):
                                        continue
                                    slot_frames = slot.get("sentence_frames")
                                    if (
                                        slot_frames
                                        and isinstance(slot_frames, list)
                                        and len(slot_frames) > 0
                                    ):
                                        has_frames = True
                                        frames_count += len(slot_frames)

                    print(
                        f"DEBUG: lesson_json has sentence_frames: {has_frames} (count: {frames_count})"
                    )

                    # Also check the sanitized version as a fallback
                    if not has_frames and "days" in lesson_json_for_pdf:
                        for day in lesson_json_for_pdf["days"].values():
                            if not isinstance(day, dict):
                                continue
                            if "sentence_frames" in day and day["sentence_frames"]:
                                has_frames = True
                                frames_count += len(day.get("sentence_frames", []))
                                break
                            if "slots" in day:
                                for s in day["slots"]:
                                    if not isinstance(s, dict):
                                        continue
                                    if "sentence_frames" in s and s["sentence_frames"]:
                                        has_frames = True
                                        frames_count += len(
                                            s.get("sentence_frames", [])
                                        )
                                        break
                        print(
                            f"DEBUG: After checking sanitized version, has_frames: {has_frames} (count: {frames_count})"
                        )

                    if has_frames:
                        # Ensure output directory exists
                        output_dir.mkdir(parents=True, exist_ok=True)

                        # Use the sanitized version for generation to avoid ModelPrivateAttr issues
                        generate_sentence_frames_pdf(
                            lesson_json_for_pdf,
                            str(sentence_frames_pdf_path),
                            user_name=user.get("name"),
                            keep_html=True,
                        )

                        # Generate DOCX file
                        sentence_frames_docx_path = (
                            sentence_frames_pdf_path.with_suffix(".docx")
                        )
                        from backend.services.sentence_frames_pdf_generator import (
                            generate_sentence_frames_docx,
                        )

                        generate_sentence_frames_docx(
                            lesson_json_for_pdf,
                            str(sentence_frames_docx_path),
                            user_name=user.get("name"),
                        )

                        # Verify files were created
                        sentence_frames_html_path = (
                            sentence_frames_pdf_path.with_suffix(".html")
                        )
                        if not sentence_frames_pdf_path.exists():
                            logger.error(
                                "batch_sentence_frames_pdf_not_created",
                                extra={"expected_path": str(sentence_frames_pdf_path)},
                            )
                        if not sentence_frames_html_path.exists():
                            logger.error(
                                "batch_sentence_frames_html_not_created",
                                extra={"expected_path": str(sentence_frames_html_path)},
                            )
                        if not sentence_frames_docx_path.exists():
                            logger.error(
                                "batch_sentence_frames_docx_not_created",
                                extra={"expected_path": str(sentence_frames_docx_path)},
                            )

                        print(
                            f"DEBUG: Sentence frames PDF generated: {sentence_frames_pdf_path.exists()}"
                        )
                        print(
                            f"DEBUG: Sentence frames HTML generated: {sentence_frames_html_path.exists()}"
                        )
                        print(
                            f"DEBUG: Sentence frames DOCX generated: {sentence_frames_docx_path.exists()}"
                        )
                        print(
                            "DEBUG: generate_sentence_frames_pdf completed successfully"
                        )
                        logger.info(
                            "batch_sentence_frames_pdf_generated",
                            extra={
                                "path": str(sentence_frames_pdf_path),
                                "pdf_exists": sentence_frames_pdf_path.exists(),
                                "html_exists": sentence_frames_html_path.exists(),
                                "docx_exists": sentence_frames_docx_path.exists(),
                                "frames_count": frames_count,
                            },
                        )
                    else:
                        print(
                            "DEBUG: Skipping sentence frames PDF generation (no frames found)"
                        )
                        logger.warning(
                            "batch_sentence_frames_skipped_no_data",
                            extra={
                                "plan_id": plan_id,
                                "output_path": str(output_path),
                                "has_days": "days" in lesson_json,
                            },
                        )
                except ImportError as ie:
                    print(
                        f"ERROR: Could not import sentence_frames_pdf_generator: {ie}"
                    )
                    logger.error(
                        "batch_sentence_frames_import_failed", extra={"error": str(ie)}
                    )
                except Exception as e:
                    print(f"ERROR: Failed to generate sentence frames PDF: {e}")
                    traceback.print_exc()
                    logger.warning(
                        "batch_sentence_frames_pdf_html_failed",
                        extra={"error": str(e), "error_type": type(e).__name__},
                        exc_info=True,
                    )
            except Exception as e:
                error_msg = f"Failed to generate objectives PDF/HTML: {str(e)}"
                print(f"ERROR: {error_msg}")
                try:
                    with open(debug_marker, "a", encoding="utf-8") as f:
                        f.write(f"ERROR: {error_msg}\n")
                        f.write(f"Error type: {type(e).__name__}\n")
                        try:
                            f.write(f"Traceback:\n{traceback.format_exc()}\n")
                        except Exception:
                            f.write("Traceback: (unavailable)\n")
                except Exception as debug_err:
                    # If we can't write to debug file, just log it
                    print(f"WARNING: Could not write to debug file: {debug_err}")
                logger.warning(
                    "batch_objectives_pdf_html_failed",
                    extra={"error": str(e), "error_type": type(e).__name__},
                    exc_info=True,
                )
                try:
                    traceback.print_exc()
                except Exception:
                    print(f"Could not print traceback: {e}")
                # Don't fail the entire process if PDF/HTML generation fails

            logger.info(
                "batch_render_single_slot_success", extra={"output_path": output_path}
            )
            _safe_finalize()
            return output_path
        else:
            # Multi-slot: render each slot separately to preserve 2-table structure
            # Then merge DOCX files
            if not lessons:
                error_msg = "No lessons provided to combine"
                logger.error("batch_combine_no_lessons")
                raise ValueError(error_msg)
            
            logger.info(
                "batch_render_multi_slot_start",
                extra={
                    "lesson_count": len(lessons),
                    "output_file": Path(output_path).name,
                },
            )

            temp_files = []
            all_unplaced_hyperlinks = []
            all_unplaced_images = []
            week_folder = file_mgr.get_week_folder(week_of)

            # Render each slot to temporary file
            # Add slot metadata to hyperlinks/images BEFORE rendering for proper filtering
            for lesson in lessons:
                slot_num = lesson["slot_number"]
                subject = lesson["subject"]
                lesson_json = lesson["lesson_json"]
                slot_data = lesson.get("slot_data", {})  # Get original slot data with primary teacher fields
                print(f"[DEBUG] BATCH_PROCESSOR: Rendering slot {slot_num} ({subject}), slot_data present: {bool(slot_data)}")
                if slot_data:
                    if isinstance(slot_data, dict):
                        print(f"[DEBUG] BATCH_PROCESSOR: slot_data dict keys: {list(slot_data.keys())[:10]}")
                        print(f"[DEBUG] BATCH_PROCESSOR: slot_data.primary_teacher_name: {slot_data.get('primary_teacher_name')}")
                    else:
                        print(f"[DEBUG] BATCH_PROCESSOR: slot_data is object, primary_teacher_name: {getattr(slot_data, 'primary_teacher_name', None)}")

                # CRITICAL: Add slot metadata to lesson JSON metadata
                # The renderer will extract this to enable strict filtering
                # Ensure lesson_json is a dictionary before using 'in' operator
                if not isinstance(lesson_json, dict):
                    if hasattr(lesson_json, "model_dump"):
                        lesson_json = lesson_json.model_dump()
                    elif hasattr(lesson_json, "dict"):
                        lesson_json = lesson_json.dict()
                    else:
                        lesson_json = dict(lesson_json) if lesson_json else {}

                if "metadata" not in lesson_json:
                    lesson_json["metadata"] = {}
                lesson_json["metadata"]["slot_number"] = slot_num
                lesson_json["metadata"]["subject"] = subject
                
                # CRITICAL: Update primary teacher fields from slot_data
                # This ensures each slot shows its own primary teacher, not the combined name
                if slot_data:
                    if isinstance(slot_data, dict):
                        primary_teacher_name = slot_data.get("primary_teacher_name")
                        primary_teacher_first_name = slot_data.get("primary_teacher_first_name")
                        primary_teacher_last_name = slot_data.get("primary_teacher_last_name")
                        lesson_json["metadata"]["primary_teacher_name"] = primary_teacher_name
                        lesson_json["metadata"]["primary_teacher_first_name"] = primary_teacher_first_name
                        lesson_json["metadata"]["primary_teacher_last_name"] = primary_teacher_last_name
                    else:
                        # Handle database object
                        primary_teacher_name = getattr(slot_data, "primary_teacher_name", None)
                        primary_teacher_first_name = getattr(slot_data, "primary_teacher_first_name", None)
                        primary_teacher_last_name = getattr(slot_data, "primary_teacher_last_name", None)
                        lesson_json["metadata"]["primary_teacher_name"] = primary_teacher_name
                        lesson_json["metadata"]["primary_teacher_first_name"] = primary_teacher_first_name
                        lesson_json["metadata"]["primary_teacher_last_name"] = primary_teacher_last_name
                    
                    # CRITICAL: Also update teacher_name with combined format "Primary / Bilingual"
                    # This ensures the metadata table shows "Kelsey Lang / Wilson Rodrigues" format
                    try:
                        combined_teacher_name = self._build_teacher_name(
                            {
                                "first_name": getattr(self, "_user_first_name", ""),
                                "last_name": getattr(self, "_user_last_name", ""),
                                "name": getattr(self, "_user_name", ""),
                            },
                            slot_data if isinstance(slot_data, dict) else {
                                "primary_teacher_name": primary_teacher_name,
                                "primary_teacher_first_name": primary_teacher_first_name,
                                "primary_teacher_last_name": primary_teacher_last_name,
                            },
                        )
                        lesson_json["metadata"]["teacher_name"] = combined_teacher_name
                        print(f"[DEBUG] BATCH_PROCESSOR: Updated slot {slot_num} with combined teacher_name={combined_teacher_name}")
                    except Exception as e:
                        print(f"[DEBUG] BATCH_PROCESSOR: Error building combined teacher name for slot {slot_num}: {e}")
                        # Fallback to just primary teacher name
                        lesson_json["metadata"]["teacher_name"] = primary_teacher_name or "Unknown"
                else:
                    print(f"[DEBUG] BATCH_PROCESSOR: WARNING - No slot_data for slot {slot_num}, cannot update primary teacher fields")

                # CRITICAL: Ensure media schema version is set to 2.0
                # Without this, the renderer defaults to 1.0 and ignores all hyperlinks/images
                lesson_json["_media_schema_version"] = "2.0"

                # Also add slot metadata to hyperlinks and images
                # This enables the renderer's strict filtering to work correctly
                if "_hyperlinks" in lesson_json:
                    for link in lesson_json.get("_hyperlinks", []):
                        if isinstance(link, dict):
                            link["_source_slot"] = slot_num
                            link["_source_subject"] = subject

                if "_images" in lesson_json:
                    for image in lesson_json.get("_images", []):
                        if isinstance(image, dict):
                            image["_source_slot"] = slot_num
                            image["_source_subject"] = subject

                # Create temp filename
                temp_filename = f"_temp_slot{slot_num}_{subject.replace('/', '_')}.docx"
                temp_path = str(week_folder / temp_filename)

                logger.debug(
                    "batch_render_slot",
                    extra={
                        "slot_number": slot_num,
                        "subject": subject,
                        "hyperlinks": len(lesson_json.get("_hyperlinks", [])),
                        "images": len(lesson_json.get("_images", [])),
                    },
                )

                # DIAGNOSTIC: Log before rendering
                from tools.diagnostic_logger import get_diagnostic_logger

                diag = get_diagnostic_logger()
                diag.log_before_render(slot_num, subject, lesson_json, "multi_slot")

                # Render this slot (will have 2 tables: original + bilingual)
                # Pass plan_id for detailed tracking
                # CRITICAL: Set skip_fallback_sections=True to prevent duplicate "Referenced Links"
                # after every table. These will be consolidated at the end.
                # #region agent log
                try:
                    import json
                    import time
                    with open(r'd:\LP\hyperlink_debug.log', 'a', encoding='utf-8') as f:
                        f.write(json.dumps({
                            "sessionId": "debug-session",
                            "runId": "run1",
                            "hypothesisId": "RENDER_CALL",
                            "location": "batch_processor.py:_combine_lessons_impl:before_render",
                            "message": "About to call renderer.render",
                            "data": {
                                "slot_number": slot_num,
                                "subject": subject,
                                "temp_path": temp_path,
                                "has_lesson_json": lesson_json is not None,
                            },
                            "timestamp": int(time.time() * 1000)
                        }) + '\n')
                except Exception as e:
                    logger.warning(f"Failed to log RENDER_CALL: {e}")
                # #endregion
                
                success, unplaced_hl, unplaced_img = renderer.render(
                    lesson_json, temp_path, plan_id=plan_id, skip_fallback_sections=True
                )

                # #region agent log
                try:
                    import json
                    import time
                    # Path is already imported at module level (line 12)
                    import json
                    import time
                    # Path is already imported at module level (line 12)
                    with open(r'd:\LP\hyperlink_debug.log', 'a', encoding='utf-8') as f:
                        f.write(json.dumps({
                            "sessionId": "debug-session",
                            "runId": "run1",
                            "hypothesisId": "RENDER_RESULT",
                            "location": "batch_processor.py:_combine_lessons_impl:after_render",
                            "message": "Renderer.render returned",
                            "data": {
                                "slot_number": slot_num,
                                "subject": subject,
                                "success": success,
                                "temp_path": temp_path,
                                "file_exists": Path(temp_path).exists() if success else False,
                            },
                            "timestamp": int(time.time() * 1000)
                        }) + '\n')
                except Exception as e:
                    logger.warning(f"Failed to log RENDER_RESULT: {e}")
                # #endregion

                if success:
                    temp_files.append(temp_path)
                    all_unplaced_hyperlinks.extend(unplaced_hl)
                    all_unplaced_images.extend(unplaced_img)
                else:
                    # #region agent log
                    import json
                    with open(r'd:\LP\.cursor\debug.log', 'a', encoding='utf-8') as f:
                        f.write(json.dumps({
                            "sessionId": "debug-session",
                            "runId": "run1",
                            "hypothesisId": "RENDER_FAIL",
                            "location": "batch_processor.py:_combine_lessons_impl:render_failed",
                            "message": "Slot rendering failed",
                            "data": {
                                "slot_number": slot_num,
                                "subject": subject,
                                "temp_path": temp_path,
                                "has_lesson_json": lesson_json is not None,
                                "lesson_json_keys": list(lesson_json.keys()) if isinstance(lesson_json, dict) else None,
                            },
                            "timestamp": int(__import__('time').time() * 1000)
                        }) + '\n')
                    # #endregion
                    logger.error(
                        "batch_render_slot_failed",
                        extra={
                            "slot_number": slot_num,
                            "subject": subject,
                            "temp_path": temp_path,
                        },
                    )

            # Check if we have any files to merge before attempting merge
            if not temp_files:
                error_msg = (
                    f"No files to merge: All {len(lessons)} slot(s) failed to render. "
                    f"Check logs for rendering errors."
                )
                logger.error("batch_merge_no_files", extra={"lesson_count": len(lessons)})
                raise ValueError(error_msg)

            # Merge all temp files into one consolidated DOCX
            logger.info(
                "batch_merge_slots",
                extra={"slot_count": len(temp_files), "output_file": output_path},
            )
            try:
                # Track file merging operation
                if plan_id:
                    with self.tracker.track_operation(
                        plan_id,
                        "render_merge_files",
                        metadata={"file_count": len(temp_files)},
                    ):
                        self._merge_docx_files(temp_files, output_path)
                else:
                    self._merge_docx_files(temp_files, output_path)
                logger.info(
                    "batch_merge_slots_success", extra={"slot_count": len(temp_files)}
                )
            except Exception as e:
                logger.exception("batch_merge_slots_error", extra={"error": str(e)})
                raise

            # Remove all signature boxes from merged document, then add one at the end
            logger.debug("batch_clean_signature_boxes")
            doc = Document(output_path)
            # Track signature operations
            # Use user-specific signature image if available (optional - only if path is provided)
            signature_image_path = user.get("signature_image_path")
            # Only use signature image if path is provided and not empty
            if signature_image_path and not signature_image_path.strip():
                signature_image_path = None

            # Special case: Wilson Rodrigues should use PNG signature if not set in database
            user_name = user.get("name", "")
            if (
                not signature_image_path
                and user_name
                and "Wilson Rodrigues" in user_name
            ):
                # Try multiple possible paths (case-insensitive matching)
                possible_paths = [
                    r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.PNG",
                    r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.png",
                    r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.Png",
                ]
                for sig_path in possible_paths:
                    if Path(sig_path).exists():
                        signature_image_path = sig_path
                        logger.info(
                            "using_default_wilson_signature",
                            extra={"path": sig_path, "user": user_name},
                        )
                        break
                else:
                    logger.warning(
                        "wilson_signature_not_found",
                        extra={
                            "user": user_name,
                            "attempted_paths": possible_paths,
                        },
                    )
            if plan_id:
                with self.tracker.track_operation(plan_id, "render_remove_signatures"):
                    self._remove_signature_boxes(doc)
                with self.tracker.track_operation(plan_id, "render_add_signature"):
                    self._add_signature_box(
                        doc,
                        generated_at,
                        template_path,
                        signature_image_path,
                        user.get("name"),
                    )
                doc.save(output_path)
            else:
                self._remove_signature_boxes(doc)
                self._add_signature_box(
                    doc,
                    generated_at,
                    template_path,
                    signature_image_path,
                    user.get("name"),
                )
                doc.save(output_path)

            # CONSOLIDATED FALLBACK MEDIA: Add a single "Referenced Links" section at the end
            # if any hyperlinks or images were not placed inline in their respective slots.
            if all_unplaced_hyperlinks or all_unplaced_images:
                logger.info(
                    "consolidating_fallback_media",
                    extra={
                        "hyperlinks": len(all_unplaced_hyperlinks),
                        "images": len(all_unplaced_images),
                    },
                )

                # De-duplicate hyperlinks by (text, url)
                unique_hl = []
                seen_hl = set()
                for hl in all_unplaced_hyperlinks:
                    key = (hl.get("text", "").strip(), hl.get("url", "").strip())
                    if key not in seen_hl:
                        seen_hl.add(key)
                        unique_hl.append(hl)

                # De-duplicate images by data hash
                unique_img = []
                seen_img = set()
                import hashlib

                for img in all_unplaced_images:
                    data_hash = hashlib.md5(img["data"].encode()).hexdigest()
                    if data_hash not in seen_img:
                        seen_img.add(data_hash)
                        unique_img.append(img)

                # Append to the final document
                # Re-open doc to ensure it's fresh after signature saving
                doc = Document(output_path)
                renderer._append_unmatched_media(doc, unique_hl, unique_img)
                doc.save(output_path)

            # JSON file writing removed - using database lesson_json column as single source of truth
            # Ensure lesson_json is saved to database via update_weekly_plan() call

            # Generate objectives DOCX file
            try:
                from backend.services.objectives_printer import ObjectivesPrinter

                objectives_printer = ObjectivesPrinter()
                objectives_output_path = Path(output_path).with_name(
                    Path(output_path).stem + "_objectives.docx"
                )
                objectives_printer.generate_docx(
                    merged_json,
                    str(objectives_output_path),
                    user_name=user.get("name"),
                    week_of=week_of,
                )
                logger.info(
                    "batch_objectives_docx_generated",
                    extra={"objectives_path": str(objectives_output_path)},
                )
            except Exception as e:
                logger.warning(
                    "batch_objectives_docx_failed",
                    extra={"error": str(e), "error_type": type(e).__name__},
                    exc_info=True,
                )
                # Don't fail the entire process if objectives generation fails

            # Generate objectives PDF and HTML files
            try:
                from backend.services.objectives_pdf_generator import (
                    generate_objectives_pdf,
                )

                objectives_pdf_path = Path(output_path).with_name(
                    Path(output_path).stem + "_objectives.pdf"
                )
                print(f"DEBUG: Generating objectives PDF at: {objectives_pdf_path}")

                # CRITICAL: Sanitize merged_json to remove ModelPrivateAttr objects before PDF generation
                merged_json_for_pdf = self._sanitize_value(merged_json)

                logger.info(
                    "batch_objectives_pdf_html_starting",
                    extra={"objectives_pdf_path": str(objectives_pdf_path)},
                )
                # Generate PDF with keep_html=True to also create the HTML file
                generate_objectives_pdf(
                    merged_json_for_pdf,
                    str(objectives_pdf_path),
                    user_name=user.get("name"),
                    keep_html=True,  # Keep HTML file (it will be created with .html extension)
                )
                # The HTML file will be created with the same name but .html extension
                objectives_html_path = objectives_pdf_path.with_suffix(".html")

                # Verify files were created
                if not objectives_pdf_path.exists():
                    logger.error(
                        "batch_objectives_pdf_not_created",
                        extra={"expected_path": str(objectives_pdf_path)},
                    )
                if not objectives_html_path.exists():
                    logger.error(
                        "batch_objectives_html_not_created",
                        extra={"expected_path": str(objectives_html_path)},
                    )

                print(
                    f"DEBUG: Objectives PDF generated: {objectives_pdf_path.exists()}"
                )
                print(
                    f"DEBUG: Objectives HTML generated: {objectives_html_path.exists()}"
                )
                logger.info(
                    "batch_objectives_pdf_html_generated",
                    extra={
                        "objectives_pdf_path": str(objectives_pdf_path),
                        "objectives_html_path": str(objectives_html_path),
                        "pdf_exists": objectives_pdf_path.exists(),
                        "html_exists": objectives_html_path.exists(),
                    },
                )
            except Exception as e:
                error_msg = f"Failed to generate objectives PDF/HTML: {str(e)}"
                print(f"ERROR: {error_msg}")
                logger.warning(
                    "batch_objectives_pdf_html_failed",
                    extra={"error": str(e), "error_type": type(e).__name__},
                    exc_info=True,
                )
                try:
                    traceback.print_exc()
                except Exception:
                    print(f"Could not print traceback: {e}")
                # Don't fail the entire process if PDF/HTML generation fails

            # Generate sentence frames PDF/HTML for multi-slot plans
            try:
                print(
                    "DEBUG: Attempting to import sentence_frames_pdf_generator for multi-slot"
                )
                from backend.services.sentence_frames_pdf_generator import (
                    generate_sentence_frames_pdf,
                )

                print("DEBUG: Successfully imported sentence_frames_pdf_generator")

                # Use the same directory as the output_path to ensure files are in the correct folder
                output_dir = Path(output_path).parent
                output_stem = Path(output_path).stem
                sentence_frames_pdf_path = (
                    output_dir / f"{output_stem}_sentence_frames.pdf"
                )

                logger.info(
                    "batch_sentence_frames_pdf_html_starting",
                    extra={
                        "sentence_frames_pdf_path": str(sentence_frames_pdf_path),
                        "output_path": str(output_path),
                        "output_dir": str(output_dir),
                        "multi_slot": True,
                    },
                )

                print(f"DEBUG: Output path: {output_path}")
                print(f"DEBUG: Output directory: {output_dir}")
                print(f"DEBUG: Sentence frames PDF path: {sentence_frames_pdf_path}")

                # Check if sentence_frames exist in the merged JSON
                has_frames = False
                frames_count = 0
                if "days" in merged_json:
                    for day_name, day in merged_json["days"].items():
                        if not isinstance(day, dict):
                            continue
                        # Check day-level frames
                        day_frames = day.get("sentence_frames")
                        if (
                            day_frames
                            and isinstance(day_frames, list)
                            and len(day_frames) > 0
                        ):
                            has_frames = True
                            frames_count += len(day_frames)
                        # Check slot-level frames
                        slots = day.get("slots", [])
                        if isinstance(slots, list):
                            for slot in slots:
                                if not isinstance(slot, dict):
                                    continue
                                slot_frames = slot.get("sentence_frames")
                                if (
                                    slot_frames
                                    and isinstance(slot_frames, list)
                                    and len(slot_frames) > 0
                                ):
                                    has_frames = True
                                    frames_count += len(slot_frames)

                print(
                    f"DEBUG: merged_json has sentence_frames: {has_frames} (count: {frames_count})"
                )

                # Also check the sanitized version as a fallback
                if not has_frames and "days" in merged_json_for_pdf:
                    for day in merged_json_for_pdf["days"].values():
                        if not isinstance(day, dict):
                            continue
                        if "sentence_frames" in day and day["sentence_frames"]:
                            has_frames = True
                            frames_count += len(day.get("sentence_frames", []))
                            break
                        if "slots" in day:
                            for s in day["slots"]:
                                if not isinstance(s, dict):
                                    continue
                                if "sentence_frames" in s and s["sentence_frames"]:
                                    has_frames = True
                                    frames_count += len(s.get("sentence_frames", []))
                                    break
                    print(
                        f"DEBUG: After checking sanitized version, has_frames: {has_frames} (count: {frames_count})"
                    )

                if has_frames:
                    # Ensure output directory exists
                    output_dir.mkdir(parents=True, exist_ok=True)

                    # Use the sanitized version for generation to avoid ModelPrivateAttr issues
                    generate_sentence_frames_pdf(
                        merged_json_for_pdf,
                        str(sentence_frames_pdf_path),
                        user_name=user.get("name"),
                        keep_html=True,
                    )

                    # Generate DOCX file
                    sentence_frames_docx_path = sentence_frames_pdf_path.with_suffix(
                        ".docx"
                    )
                    from backend.services.sentence_frames_pdf_generator import (
                        generate_sentence_frames_docx,
                    )

                    generate_sentence_frames_docx(
                        merged_json_for_pdf,
                        str(sentence_frames_docx_path),
                        user_name=user.get("name"),
                    )

                    # Verify files were created
                    sentence_frames_html_path = sentence_frames_pdf_path.with_suffix(
                        ".html"
                    )
                    if not sentence_frames_pdf_path.exists():
                        logger.error(
                            "batch_sentence_frames_pdf_not_created",
                            extra={"expected_path": str(sentence_frames_pdf_path)},
                        )
                    if not sentence_frames_html_path.exists():
                        logger.error(
                            "batch_sentence_frames_html_not_created",
                            extra={"expected_path": str(sentence_frames_html_path)},
                        )
                    if not sentence_frames_docx_path.exists():
                        logger.error(
                            "batch_sentence_frames_docx_not_created",
                            extra={"expected_path": str(sentence_frames_docx_path)},
                        )

                    print(
                        f"DEBUG: Sentence frames PDF generated: {sentence_frames_pdf_path.exists()}"
                    )
                    print(
                        f"DEBUG: Sentence frames HTML generated: {sentence_frames_html_path.exists()}"
                    )
                    print(
                        f"DEBUG: Sentence frames DOCX generated: {sentence_frames_docx_path.exists()}"
                    )
                    print(
                        "DEBUG: generate_sentence_frames_pdf completed successfully for multi-slot"
                    )
                    logger.info(
                        "batch_sentence_frames_pdf_generated",
                        extra={
                            "path": str(sentence_frames_pdf_path),
                            "pdf_exists": sentence_frames_pdf_path.exists(),
                            "html_exists": sentence_frames_html_path.exists(),
                            "docx_exists": sentence_frames_docx_path.exists(),
                            "frames_count": frames_count,
                            "multi_slot": True,
                        },
                    )
                else:
                    print(
                        "DEBUG: Skipping sentence frames PDF generation (no frames found in merged JSON)"
                    )
                    logger.warning(
                        "batch_sentence_frames_skipped_no_data",
                        extra={
                            "plan_id": plan_id,
                            "output_path": str(output_path),
                            "has_days": "days" in merged_json,
                            "multi_slot": True,
                        },
                    )
            except ImportError as ie:
                print(f"ERROR: Could not import sentence_frames_pdf_generator: {ie}")
                logger.error(
                    "batch_sentence_frames_import_failed",
                    extra={"error": str(ie), "multi_slot": True},
                )
            except Exception as e:
                print(
                    f"ERROR: Failed to generate sentence frames PDF for multi-slot: {e}"
                )
                traceback.print_exc()
                logger.warning(
                    "batch_sentence_frames_pdf_html_failed",
                    extra={
                        "error": str(e),
                        "error_type": type(e).__name__,
                        "multi_slot": True,
                    },
                    exc_info=True,
                )

            # Clean up temp files - ensure they're deleted even if there's an error
            for temp_file in temp_files:
                try:
                    temp_path = Path(temp_file)
                    if temp_path.exists():
                        temp_path.unlink()
                        logger.debug(
                            "batch_temp_file_deleted", extra={"temp_file": temp_file}
                        )
                except Exception as e:
                    logger.warning(
                        "batch_temp_cleanup_failed",
                        extra={"temp_file": temp_file, "error": str(e)},
                    )

            logger.info(
                "batch_render_multi_slot_success",
                extra={"lesson_count": len(lessons), "output_path": output_path},
            )
            _safe_finalize()
            return output_path

    def _merge_docx_files(
        self,
        file_paths: List[str],
        output_path: str,
        master_template_path: Optional[str] = None,
    ):
        """
        Merge multiple DOCX files into one using docxcompose.
        Each document is separated by a page break.

        Args:
            file_paths: List of DOCX file paths to merge
            output_path: Path for merged output file
            master_template_path: Optional path to a template to use as the base
        """
        from docx import Document
        from docxcompose.composer import Composer

        if not file_paths:
            raise ValueError("No files to merge")

        logger.debug("batch_merge_docx_start", extra={"file_count": len(file_paths)})

        # Load the base document
        if master_template_path:
            master = Document(master_template_path)
            # Clear all content from the master template to start fresh
            for p in list(master.paragraphs):
                p._element.getparent().remove(p._element)
            for t in list(master.tables):
                t._element.getparent().remove(t._element)
            composer = Composer(master)
            start_idx = 0
        else:
            master = Document(file_paths[0])
            composer = Composer(master)
            start_idx = 1

        # Append documents with explicit page breaks
        for i, file_path in enumerate(file_paths[start_idx:], start_idx + 1):
            logger.debug(
                "batch_merge_docx_append",
                extra={"index": i, "file_name": Path(file_path).name},
            )
            doc = Document(file_path)

            # Add page break at the beginning of each document being appended
            # (except for the very first document if we're not using a template)
            if i > 1:
                page_break_para = doc.add_page_break()
                para_element = page_break_para._element
                para_element.getparent().remove(para_element)
                doc._element.body.insert(0, para_element)

            composer.append(doc)

        # Save the merged document
        logger.debug(
            "batch_merge_docx_save", extra={"output_file": Path(output_path).name}
        )
        composer.save(output_path)

    def _get_week_num(self, week_of: str) -> int:
        """Extract week number from week_of string.

        Supports formats:
        - MM/DD-MM/DD (e.g., 10/13-10/17)
        - MM-DD-MM-DD (e.g., 10-13-10-17)
        """
        try:
            from datetime import datetime

            # Normalize format: convert all hyphens to slashes first, then split
            # Handle both "10/13-10/17" and "10-13-10-17" formats
            if "/" in week_of:
                # Format: MM/DD-MM/DD
                first_date = week_of.split("-")[0].strip()
                month, day = map(int, first_date.split("/"))
            else:
                # Format: MM-DD-MM-DD (all hyphens)
                parts = week_of.split("-")
                if len(parts) >= 2:
                    month, day = int(parts[0]), int(parts[1])
                else:
                    raise ValueError(f"Invalid week_of format: {week_of}")

            year = datetime.now().year
            date_obj = datetime(year, month, day)
            week_num = date_obj.isocalendar()[1]
            logger.debug(
                "week_number_calculated",
                extra={
                    "week_of": week_of,
                    "month": month,
                    "day": day,
                    "year": year,
                    "week_number": week_num,
                },
            )
            return week_num
        except Exception as e:
            logger.warning(
                "week_number_calculation_failed",
                extra={"week_of": week_of, "error": str(e)},
            )
            return 1

    def _remove_signature_boxes(self, doc: Document):
        """Remove signature boxes/tables from document.

        Args:
            doc: Document to remove signatures from
        """
        # Find and remove tables containing signature text
        tables_to_remove = []
        for table in doc.tables:
            # Check if table contains signature-related text
            table_text = "\n".join(
                [cell.text for row in table.rows for cell in row.cells]
            )
            if any(
                keyword in table_text
                for keyword in [
                    "Required Signatures",
                    "Teacher Signature:",
                    "Administrator Signature:",
                    "Administrative Feedback:",
                    "I certify that these lessons",
                ]
            ):
                tables_to_remove.append(table)

        # Remove the signature tables
        for table in tables_to_remove:
            tbl_element = table._element
            tbl_element.getparent().remove(tbl_element)

        # Also remove paragraphs containing signature text
        paragraphs_to_remove = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("Signature:") or text.startswith("Date:"):
                paragraphs_to_remove.append(para)

        # Remove the paragraphs
        for para in paragraphs_to_remove:
            p_element = para._element
            p_element.getparent().remove(p_element)

    def _modify_existing_signature_table(
        self,
        doc: Document,
        date: datetime,
        signature_image_path: str = None,
        user_name: str = None,
    ):
        """Modify existing signature table in document (for single-slot documents).

        Args:
            doc: Document containing the signature table
            date: Date to display
            signature_image_path: Optional path to signature image file
            user_name: User's name to use if no signature image is provided
        """
        # Find the signature table in the document
        signature_table = None
        for table in doc.tables:
            table_text = "\n".join(
                [cell.text for row in table.rows for cell in row.cells]
            )
            if "Required Signatures" in table_text:
                signature_table = table
                break

        if not signature_table:
            logger.warning("signature_table_not_found_in_doc")
            return

        # Fill date and add signature/name using the same logic as _add_signature_box
        date_formatted = date.strftime("%m/%d/%Y")

        # Find the cell containing "Teacher Signature:" and update the Date that appears AFTER it
        for row in signature_table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Look for cell containing "Teacher Signature:"
                if "Teacher Signature:" in cell_text:
                    # Find the SPECIFIC paragraph that contains "Teacher Signature:" followed by "Date:"
                    # (not the certification paragraph which might have other text)
                    target_para = None
                    for para in cell.paragraphs:
                        para_text = para.text.strip()
                        # Must start with or contain "Teacher Signature:" and have "Date:" after it
                        if "Teacher Signature:" in para_text:
                            teacher_pos = para_text.find("Teacher Signature:")
                            date_pos = (
                                para_text.find("Date:", teacher_pos)
                                if teacher_pos >= 0
                                else -1
                            )
                            # Only select if Date appears AFTER Teacher Signature
                            if date_pos > teacher_pos:
                                target_para = para
                                break

                    if target_para:
                        para = target_para
                        para_text = para.text
                        # Find the "Date:" that appears AFTER "Teacher Signature:" in the paragraph
                        teacher_pos = para_text.find("Teacher Signature:")
                        date_pos = para_text.find(
                            "Date:", teacher_pos
                        )  # Find Date after Teacher Signature

                        if (
                            date_pos > teacher_pos
                        ):  # Date appears after Teacher Signature
                            # Find the run containing this Date (the one after Teacher Signature)
                            date_run = None
                            original_font_size = None
                            original_font_name = None

                            # Track position in paragraph to find the right Date run
                            current_pos = 0
                            for run in para.runs:
                                run_text = run.text
                                run_start = current_pos
                                run_end = current_pos + len(run_text)

                                # Check if this run contains the Date we're looking for (after Teacher Signature)
                                if "Date:" in run_text and run_start >= teacher_pos:
                                    # This is the Date run after Teacher Signature
                                    date_run = run
                                    if run.font.size:
                                        original_font_size = run.font.size
                                    if run.font.name:
                                        original_font_name = run.font.name
                                    break

                                current_pos = run_end

                            if date_run:
                                # Update this Date field using TAB STOPS for precise alignment
                                run_text = date_run.text
                                if "Date:" in run_text:
                                    # Clear and rebuild with tab stop
                                    para.clear()

                                    # Add tab stop at 5.5 inches (matches Administrator row)
                                    # This ensures perfect alignment regardless of font metrics
                                    DATE_TAB_POSITION = Inches(5.5)
                                    para.paragraph_format.tab_stops.add_tab_stop(
                                        DATE_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT
                                    )

                                    # Add "Teacher Signature:" label
                                    teacher_run = para.add_run("Teacher Signature: ")
                                    teacher_run.font.bold = True
                                    if original_font_size:
                                        teacher_run.font.size = original_font_size
                                    if original_font_name:
                                        teacher_run.font.name = original_font_name

                                    # Placeholder for signature/name (will be replaced by _add_signature_image_to_table or _add_user_name_to_table)
                                    placeholder_run = para.add_run("_" * 49)
                                    if original_font_size:
                                        placeholder_run.font.size = original_font_size
                                    if original_font_name:
                                        placeholder_run.font.name = original_font_name

                                    # Add TAB character to jump to tab stop position
                                    para.add_run("\t")

                                    # Add Date label and value
                                    date_label_run = para.add_run("Date: ")
                                    date_label_run.font.bold = True
                                    if original_font_size:
                                        date_label_run.font.size = original_font_size
                                    if original_font_name:
                                        date_label_run.font.name = original_font_name

                                    date_value_run = para.add_run(date_formatted)
                                    date_value_run.font.underline = True
                                    if original_font_size:
                                        date_value_run.font.size = original_font_size
                                    if original_font_name:
                                        date_value_run.font.name = original_font_name

                                    logger.info("signature_date_updated_with_tab_stop")
                                break
                    break

        # Add signature image or user name
        if (
            signature_image_path
            and signature_image_path.strip()
            and Path(signature_image_path).exists()
        ):
            self._add_signature_image_to_table(signature_table, signature_image_path)
        elif user_name:
            self._add_user_name_to_table(signature_table, user_name)

    def _add_signature_image_to_table(self, table, signature_image_path: str):
        """Add signature image to the signature table."""
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if "Teacher Signature:" in cell_text:
                    for para in cell.paragraphs:
                        # Re-read paragraph text to get current state (may have been modified by date insertion)
                        para_text = para.text
                        if "Teacher Signature:" in para_text:
                            # Get font size for image sizing
                            text_font_size = Pt(11)
                            for run in para.runs:
                                if run.font.size:
                                    text_font_size = run.font.size
                                    break

                            text_height_inches = text_font_size.pt / 72.0
                            image_height_inches = text_height_inches * 1.4

                            # Find the position of "Teacher Signature:" and "Date:" in the CURRENT paragraph
                            teacher_pos = para_text.find("Teacher Signature:")
                            date_pos = (
                                para_text.find("Date:", teacher_pos)
                                if teacher_pos >= 0
                                else -1
                            )

                            if teacher_pos >= 0 and date_pos > teacher_pos:
                                try:
                                    # Get text before "Teacher Signature:"
                                    before_teacher = para_text[:teacher_pos]

                                    # Get text between "Teacher Signature:" and "Date:"
                                    # This is where we'll insert the image
                                    between_text = para_text[
                                        teacher_pos
                                        + len("Teacher Signature:") : date_pos
                                    ].strip()
                                    # Remove underscores that are placeholders
                                    between_text = between_text.lstrip("_").strip()

                                    # Check if date was already updated (look for formatted date in paragraph)
                                    date_updated = False
                                    date_value = None
                                    # Look for date pattern MM/DD/YYYY in the paragraph after "Date:"
                                    import re

                                    date_pattern = r"Date:\s*(\d{1,2}/\d{1,2}/\d{4})"
                                    date_match = re.search(date_pattern, para_text)
                                    if date_match:
                                        date_updated = True
                                        date_value = date_match.group(1)

                                    # Get text after "Date:" (not needed - we rebuild the paragraph)
                                    # Removed unused after_date variable

                                    # Save font properties
                                    original_font_size = None
                                    original_font_name = None
                                    for run in para.runs:
                                        if run.font.size:
                                            original_font_size = run.font.size
                                        if run.font.name:
                                            original_font_name = run.font.name
                                        if original_font_size and original_font_name:
                                            break

                                    # Rebuild the paragraph
                                    para.clear()

                                    # Add text before "Teacher Signature:"
                                    if before_teacher.strip():
                                        before_run = para.add_run(before_teacher)
                                        if original_font_size:
                                            before_run.font.size = original_font_size
                                        if original_font_name:
                                            before_run.font.name = original_font_name

                                    # Add "Teacher Signature:" label (bold)
                                    teacher_run = para.add_run("Teacher Signature: ")
                                    teacher_run.font.bold = True
                                    if original_font_size:
                                        teacher_run.font.size = original_font_size
                                    if original_font_name:
                                        teacher_run.font.name = original_font_name

                                    # Add the signature image
                                    image_run = para.add_run()
                                    try:
                                        if not Path(signature_image_path).exists():
                                            logger.error(
                                                "signature_image_file_not_found",
                                                extra={"path": signature_image_path},
                                            )
                                            raise FileNotFoundError(
                                                f"Signature image not found: {signature_image_path}"
                                            )

                                        with open(
                                            signature_image_path, "rb"
                                        ) as img_file:
                                            img_stream = BytesIO(img_file.read())
                                            image_run.add_picture(
                                                img_stream,
                                                height=Inches(image_height_inches),
                                            )

                                        # Verify image was added
                                        if not image_run._element.xpath(".//a:blip"):
                                            logger.error(
                                                "signature_image_not_inserted",
                                                extra={"path": signature_image_path},
                                            )
                                    except Exception as img_error:
                                        logger.error(
                                            "signature_image_insertion_error",
                                            extra={
                                                "path": signature_image_path,
                                                "error": str(img_error),
                                            },
                                        )
                                        import traceback

                                        logger.debug(
                                            "signature_image_error_traceback",
                                            extra={"traceback": traceback.format_exc()},
                                        )
                                        # Don't re-raise - continue without image

                                    # Add any text between signature and date (should be minimal/empty)
                                    if between_text:
                                        between_run = para.add_run(between_text)
                                        if original_font_size:
                                            between_run.font.size = original_font_size
                                        if original_font_name:
                                            between_run.font.name = original_font_name

                                    # Use TAB STOP for precise alignment (same as Administrator row)
                                    # Add tab stop at 5.5 inches
                                    DATE_TAB_POSITION = Inches(5.5)
                                    para.paragraph_format.tab_stops.add_tab_stop(
                                        DATE_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT
                                    )

                                    # Add TAB character to jump to tab stop position
                                    para.add_run("\t")

                                    # Add Date label and value (preserve if already updated, otherwise add placeholder)
                                    # Date label should be bold
                                    date_label_run = para.add_run("Date: ")
                                    date_label_run.font.bold = True
                                    if original_font_size:
                                        date_label_run.font.size = original_font_size
                                    if original_font_name:
                                        date_label_run.font.name = original_font_name

                                    if date_updated and date_value:
                                        # Date was already updated, preserve it with underline
                                        date_value_run = para.add_run(date_value)
                                        date_value_run.font.underline = True
                                        if original_font_size:
                                            date_value_run.font.size = (
                                                original_font_size
                                            )
                                        if original_font_name:
                                            date_value_run.font.name = (
                                                original_font_name
                                            )
                                    else:
                                        # Date not updated yet, add placeholder
                                        date_placeholder_run = para.add_run(
                                            "__________________"
                                        )
                                        if original_font_size:
                                            date_placeholder_run.font.size = (
                                                original_font_size
                                            )
                                        if original_font_name:
                                            date_placeholder_run.font.name = (
                                                original_font_name
                                            )

                                    logger.info(
                                        "signature_image_added",
                                        extra={"path": signature_image_path},
                                    )
                                    # Verify the result
                                    final_text = para.text
                                    if "Teacher Signature:" not in final_text:
                                        logger.error(
                                            "signature_text_missing_after_insertion",
                                            extra={
                                                "path": signature_image_path,
                                                "final_text": final_text,
                                            },
                                        )
                                except Exception as e:
                                    logger.warning(
                                        "signature_image_failed",
                                        extra={
                                            "path": signature_image_path,
                                            "error": str(e),
                                        },
                                    )
                                    import traceback

                                    logger.debug(
                                        "signature_image_traceback",
                                        extra={"traceback": traceback.format_exc()},
                                    )
                                break  # Break from paragraph loop after processing
                    break  # Break from cell loop after finding Teacher Signature

    def _add_user_name_to_table(self, table, user_name: str):
        """Add user name (underlined) to the signature table."""
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if "Teacher Signature:" in cell_text:
                    for para in cell.paragraphs:
                        # Re-read paragraph text to get current state (may have been modified by date insertion)
                        para_text = para.text
                        if "Teacher Signature:" in para_text:
                            # Find the position of "Teacher Signature:" and "Date:" in the CURRENT paragraph
                            teacher_pos = para_text.find("Teacher Signature:")
                            date_pos = (
                                para_text.find("Date:", teacher_pos)
                                if teacher_pos >= 0
                                else -1
                            )

                            if teacher_pos >= 0 and date_pos > teacher_pos:
                                # Get text before "Teacher Signature:"
                                before_teacher = para_text[:teacher_pos]

                                # Get text between "Teacher Signature:" and "Date:"
                                between_text = para_text[
                                    teacher_pos + len("Teacher Signature:") : date_pos
                                ].strip()
                                between_text = between_text.lstrip("_").strip()

                                # Check if date was already updated (look for formatted date in paragraph)
                                date_updated = False
                                date_value = None
                                import re

                                date_pattern = r"Date:\s*(\d{1,2}/\d{1,2}/\d{4})"
                                date_match = re.search(date_pattern, para_text)
                                if date_match:
                                    date_updated = True
                                    date_value = date_match.group(1)

                                # Get text after "Date:" (not needed - we rebuild the paragraph)
                                # Removed unused after_date variable

                                # Save font properties
                                original_font_size = None
                                original_font_name = None
                                for run in para.runs:
                                    if run.font.size:
                                        original_font_size = run.font.size
                                    if run.font.name:
                                        original_font_name = run.font.name
                                    if original_font_size and original_font_name:
                                        break

                                # Rebuild the paragraph
                                para.clear()

                                # Add text before "Teacher Signature:"
                                if before_teacher.strip():
                                    before_run = para.add_run(before_teacher)
                                    if original_font_size:
                                        before_run.font.size = original_font_size
                                    if original_font_name:
                                        before_run.font.name = original_font_name

                                # Add "Teacher Signature:" label (bold)
                                teacher_run = para.add_run("Teacher Signature: ")
                                teacher_run.font.bold = True
                                if original_font_size:
                                    teacher_run.font.size = original_font_size
                                if original_font_name:
                                    teacher_run.font.name = original_font_name

                                # Add user's name underlined
                                name_run = para.add_run(user_name)
                                name_run.font.underline = True
                                if original_font_size:
                                    name_run.font.size = original_font_size
                                if original_font_name:
                                    name_run.font.name = original_font_name

                                # Add any text between signature and date
                                if between_text:
                                    between_run = para.add_run(between_text)
                                    if original_font_size:
                                        between_run.font.size = original_font_size
                                    if original_font_name:
                                        between_run.font.name = original_font_name

                                # Use TAB STOP for precise alignment (same as Administrator row)
                                # Add tab stop at 5.5 inches
                                DATE_TAB_POSITION = Inches(5.5)
                                para.paragraph_format.tab_stops.add_tab_stop(
                                    DATE_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT
                                )

                                # Add TAB character to jump to tab stop position
                                para.add_run("\t")

                                # Add Date label and value (preserve if already updated, otherwise add placeholder)
                                # Date label should be bold
                                date_label_run = para.add_run("Date: ")
                                date_label_run.font.bold = True
                                if original_font_size:
                                    date_label_run.font.size = original_font_size
                                if original_font_name:
                                    date_label_run.font.name = original_font_name

                                if date_updated and date_value:
                                    # Date was already updated, preserve it with underline
                                    date_value_run = para.add_run(date_value)
                                    date_value_run.font.underline = True
                                    if original_font_size:
                                        date_value_run.font.size = original_font_size
                                    if original_font_name:
                                        date_value_run.font.name = original_font_name
                                else:
                                    # Date not updated yet, add placeholder
                                    date_placeholder_run = para.add_run(
                                        "__________________"
                                    )
                                    if original_font_size:
                                        date_placeholder_run.font.size = (
                                            original_font_size
                                        )
                                    if original_font_name:
                                        date_placeholder_run.font.name = (
                                            original_font_name
                                        )

                                # Verify the result
                                final_text = para.text
                                if "Teacher Signature:" not in final_text:
                                    logger.error(
                                        "signature_text_missing_after_name_insertion",
                                        extra={
                                            "user_name": user_name,
                                            "final_text": final_text,
                                        },
                                    )
                            break  # Break from paragraph loop after processing
                    break  # Break from cell loop after finding Teacher Signature

    def _add_signature_box(
        self,
        doc: Document,
        date: datetime,
        template_path: str,
        signature_image_path: str = None,
        user_name: str = None,
    ):
        """Add signature box from template to the end of document.

        Args:
            doc: Document to add signature to
            date: Date to display
            template_path: Path to the template document (user-specific or default)
            signature_image_path: Optional path to signature image file (user-specific)
            user_name: User's name to use if no signature image is provided
        """
        # Load the template to get the signature table
        template_doc = Document(template_path)

        # Find the signature table in the template (look for "Required Signatures")
        signature_table = None
        for table in template_doc.tables:
            table_text = "\n".join(
                [cell.text for row in table.rows for cell in row.cells]
            )
            if "Required Signatures" in table_text:
                signature_table = table
                break

        if signature_table:
            # Add some space before signature
            doc.add_paragraph()

            # Copy the signature table from template
            from copy import deepcopy

            new_table = deepcopy(signature_table._element)
            doc._element.body.append(new_table)

            # Get the newly added table (last table in doc)
            added_table = doc.tables[-1]

            # Find and fill the date field in the signature table (underlined)
            # Find the cell containing "Teacher Signature:" and update the Date that appears AFTER it
            date_formatted = date.strftime("%m/%d/%Y")

            for row in added_table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Look for cell containing "Teacher Signature:"
                    if "Teacher Signature:" in cell_text:
                        # Find the paragraph containing both "Teacher Signature:" and "Date:"
                        for para in cell.paragraphs:
                            para_text = para.text
                            if (
                                "Teacher Signature:" in para_text
                                and "Date:" in para_text
                            ):
                                # Find the "Date:" that appears AFTER "Teacher Signature:" in the paragraph
                                teacher_pos = para_text.find("Teacher Signature:")
                                date_pos = para_text.find(
                                    "Date:", teacher_pos
                                )  # Find Date after Teacher Signature

                                if (
                                    date_pos > teacher_pos
                                ):  # Date appears after Teacher Signature
                                    # Find the run containing this Date (the one after Teacher Signature)
                                    date_run = None
                                    original_font_size = None
                                    original_font_name = None

                                    # Track position in paragraph to find the right Date run
                                    current_pos = 0
                                    for run in para.runs:
                                        run_text = run.text
                                        run_start = current_pos
                                        run_end = current_pos + len(run_text)

                                        # Check if this run contains the Date we're looking for (after Teacher Signature)
                                        if (
                                            "Date:" in run_text
                                            and run_start >= teacher_pos
                                        ):
                                            # This is the Date run after Teacher Signature
                                            date_run = run
                                            if run.font.size:
                                                original_font_size = run.font.size
                                            if run.font.name:
                                                original_font_name = run.font.name
                                            break

                                        current_pos = run_end

                                    if date_run:
                                        # Update this Date field
                                        run_text = date_run.text
                                        if "Date:" in run_text:
                                            para.clear()
                                            # Rebuild paragraph: keep everything before Date, then add Date with value
                                            before_date_text = para_text[:date_pos]
                                            if before_date_text.strip():
                                                before_run = para.add_run(
                                                    before_date_text
                                                )
                                                if original_font_size:
                                                    before_run.font.size = (
                                                        original_font_size
                                                    )
                                                if original_font_name:
                                                    before_run.font.name = (
                                                        original_font_name
                                                    )

                                            # Add Date label and value (Date label should be bold)
                                            # NOTE: This spacing will be overwritten by _add_signature_image_to_table
                                            # or _add_user_name_to_table which rebuild the paragraph.
                                            # Actual spacing is set in those functions (91 spaces with image, 45 without)
                                            spacing_text = (
                                                " " * 90
                                            )  # This value is overwritten - see _add_signature_image_to_table/_add_user_name_to_table
                                            date_label_run = para.add_run(
                                                spacing_text + "Date: "
                                            )
                                            date_label_run.font.bold = True
                                            date_value_run = para.add_run(
                                                date_formatted
                                            )
                                            date_value_run.font.underline = True
                                            if original_font_size:
                                                date_label_run.font.size = (
                                                    original_font_size
                                                )
                                                date_value_run.font.size = (
                                                    original_font_size
                                                )
                                            if original_font_name:
                                                date_label_run.font.name = (
                                                    original_font_name
                                                )
                                                date_value_run.font.name = (
                                                    original_font_name
                                                )

                                            # Add any text after Date
                                            after_date_text = para_text[
                                                date_pos + len("Date: ") :
                                            ]
                                            after_date_text = after_date_text.lstrip(
                                                "_"
                                            ).strip()
                                            if after_date_text:
                                                after_run = para.add_run(
                                                    after_date_text
                                                )
                                                if original_font_size:
                                                    after_run.font.size = (
                                                        original_font_size
                                                    )
                                                if original_font_name:
                                                    after_run.font.name = (
                                                        original_font_name
                                                    )
                                        break
                        break

            # Now find the "Teacher Signature: " cell and add image/name inline
            # Use the helper methods for consistency
            if (
                signature_image_path
                and signature_image_path.strip()
                and Path(signature_image_path).exists()
            ):
                self._add_signature_image_to_table(added_table, signature_image_path)
            elif user_name:
                self._add_user_name_to_table(added_table, user_name)
        else:
            # Fallback to simple signature if template table not found
            doc.add_paragraph()
            doc.add_paragraph()
            sig_para = doc.add_paragraph()
            sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sig_run = sig_para.add_run("Signature: _" + "_" * 50)
            sig_run.font.size = Pt(11)
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Date: {date.strftime('%m/%d/%Y')}")
            date_run.font.size = Pt(11)

    def _calculate_week_number(self, week_of: str) -> int:
        """Calculate week number from date range.

        Args:
            week_of: Week date range (MM/DD-MM/DD)

        Returns:
            Week number (1-52)
        """
        try:
            # Parse first date from range
            first_date = week_of.split("-")[0].strip()
            month, day = map(int, first_date.split("/"))

            # Simple week calculation (can be improved)
            # Assuming school year starts in September
            if month >= 9:  # Sept-Dec
                week_num = (month - 9) * 4 + (day // 7) + 1
            else:  # Jan-June
                week_num = (month + 3) * 4 + (day // 7) + 1

            return min(week_num, 52)
        except Exception:
            return 1

    async def _extract_slots_parallel_db(
        self,
        slots: List[Dict[str, Any]],
        week_of: str,
        week_folder_path: Optional[str],
        user_base_path: Optional[str],
        plan_id: Optional[str],
        progress_tracker: Any,
    ) -> List[SlotProcessingContext]:
        """Parallel extraction using DB cache and file grouping."""
        # 1. Resolve all files first (synchronous bit)
        file_to_slots = {}
        for i, slot in enumerate(slots, 1):
            slot = self._sanitize_slot(slot)
            primary_file = self._resolve_primary_file(
                slot, week_of, week_folder_path, user_base_path
            )
            # Use "None" string as key for slots without files
            key = primary_file if primary_file else "None"
            if key not in file_to_slots:
                file_to_slots[key] = []
            file_to_slots[key].append((i, slot))

        # 2. Process each file group in parallel
        semaphore = asyncio.Semaphore(3)  # Limit concurrent DOCX parsing
        tasks = []
        for file_path_key, group in file_to_slots.items():
            file_path = None if file_path_key == "None" else file_path_key
            tasks.append(
                self._process_file_group(
                    file_path,
                    group,
                    week_of,
                    week_folder_path,
                    user_base_path,
                    plan_id,
                    semaphore,
                )
            )

        # 3. Gather all contexts
        results = await asyncio.gather(*tasks)

        # Flatten results (list of lists of contexts)
        flattened_contexts = []
        for group_contexts in results:
            flattened_contexts.extend(group_contexts)

        # Sort by slot_index to maintain original order
        return sorted(flattened_contexts, key=lambda x: x.slot_index)

    async def _process_file_group(
        self,
        file_path: Optional[str],
        group: List[Tuple[int, Dict[str, Any]]],
        week_of: str,
        week_folder_path: Optional[str],
        user_base_path: Optional[str],
        plan_id: Optional[str],
        semaphore: asyncio.Semaphore,
    ) -> List[SlotProcessingContext]:
        """Process a group of slots that share the same source file."""
        contexts = []
        user_id = getattr(self, "_current_user_id", "unknown")

        if not file_path:
            # Handle error for all slots in group
            for i, slot in group:
                ctx = SlotProcessingContext(
                    slot=slot, slot_index=i, total_slots=len(group)
                )
                ctx.error = "No primary teacher file found."
                contexts.append(ctx)
            return contexts

        # Ensure we have a lock for this file path to prevent concurrent parsing of same file
        if file_path not in self._file_locks:
            self._file_locks[file_path] = asyncio.Lock()

        async with self._file_locks[file_path]:
            # Check DB cache for each slot in group
            remaining_group = []

            # Get mtime/size for validation
            path_obj = Path(file_path)
            current_mtime = 0
            if path_obj.exists():
                stat = path_obj.stat()
                current_mtime = stat.st_mtime

            for i, slot in group:
                context = SlotProcessingContext(
                    slot=slot,
                    slot_index=i,
                    total_slots=len(group),  # Total in this specific call
                    primary_file=file_path,
                )

                # Check DB
                db_record = await asyncio.to_thread(
                    self.db.get_original_lesson_plan,
                    user_id,
                    week_of,
                    slot["slot_number"],
                )

                # Cache validation: check if extracted_at is after mtime (approximate)
                if db_record and db_record.source_file_path == file_path:
                    if path_obj.exists():
                        # We use 2 seconds buffer for FS precision
                        if db_record.extracted_at.timestamp() > (current_mtime + 2):
                            logger.info(
                                f"DB Cache hit for slot {slot['slot_number']} ({slot['subject']})"
                            )
                            context.extracted_content = db_record.full_text
                            context.available_days = db_record.available_days
                            context.cache_hit = True

                            # Restore media from cache metadata
                            cached_hyperlinks = []
                            if db_record.content_json:
                                context.slot["_extracted_images"] = (
                                    db_record.content_json.get("_extracted_images", [])
                                )
                                cached_hyperlinks = db_record.content_json.get(
                                    "_extracted_hyperlinks", []
                                )
                                context.slot["_extracted_hyperlinks"] = (
                                    cached_hyperlinks
                                )
                                context.no_school_days = (
                                    db_record.content_json.get("no_school_days", [])
                                )

                            # CRITICAL: If cache doesn't have hyperlinks, we still need to extract them
                            # This handles cases where cache was created before hyperlink extraction was fixed
                            if not cached_hyperlinks:
                                print(
                                    f"[DEBUG] DB Cache hit for slot {slot['slot_number']} but no hyperlinks in cache. "
                                    f"Will extract hyperlinks from file."
                                )
                                # Don't continue - add to remaining_group to extract hyperlinks
                                remaining_group.append((i, slot, context))
                            else:
                                print(
                                    f"[DEBUG] DB Cache hit for slot {slot['slot_number']} ({slot['subject']}) "
                                    f"with {len(cached_hyperlinks)} hyperlinks"
                                )
                                contexts.append(context)
                            continue

                remaining_group.append((i, slot, context))

            if not remaining_group:
                return contexts

            # Parse file if any slots remain
            try:
                async with semaphore:
                    logger.info(f"Opening DOCX for parsing group: {file_path}")
                    # Open file (thread-safe/blocking I/O)
                    parser = await self._open_docx_with_retry(
                        file_path, plan_id=plan_id
                    )

                    # Process each remaining slot
                    for i, slot, context in remaining_group:
                        try:
                            # 1. Resolve teacher name for finding slot
                            primary_first = slot.get(
                                "primary_teacher_first_name", ""
                            ).strip()
                            primary_last = slot.get(
                                "primary_teacher_last_name", ""
                            ).strip()
                            teacher_name = (
                                f"{primary_first} {primary_last}".strip()
                                or slot.get("primary_teacher_name", "").strip()
                            )

                            # 2. Find actual slot number
                            actual_slot_num = await asyncio.to_thread(
                                parser.find_slot_by_subject,
                                slot["subject"],
                                teacher_name,
                                slot.get("homeroom"),
                                slot.get("grade"),
                            )

                            # 3. Extract content
                            content_data = await asyncio.to_thread(
                                parser.extract_subject_content_for_slot,
                                actual_slot_num,
                                slot["subject"],
                                None,  # teacher_name
                                False,  # strip_urls=False
                            )

                            # 4. Extract images and hyperlinks
                            # CRITICAL: If context already has hyperlinks from cache, use those; otherwise extract
                            if (
                                "_extracted_hyperlinks" in context.slot
                                and context.slot["_extracted_hyperlinks"]
                            ):
                                # Hyperlinks already in context from cache, use them
                                hyperlinks = context.slot["_extracted_hyperlinks"]
                                print(
                                    f"[DEBUG] _process_file_group (PARALLEL): Using {len(hyperlinks)} hyperlinks "
                                    f"from cache for slot {slot['slot_number']}, subject {slot['subject']}"
                                )
                            else:
                                # Extract hyperlinks from file
                                hyperlinks = await asyncio.to_thread(
                                    parser.extract_hyperlinks_for_slot, actual_slot_num
                                )
                                print(
                                    f"[DEBUG] _process_file_group (PARALLEL): Extracted {len(hyperlinks)} hyperlinks "
                                    f"from file for slot {slot['slot_number']}, subject {slot['subject']}"
                                )

                            # Always extract images (they're smaller and less likely to be cached)
                            images = await asyncio.to_thread(
                                parser.extract_images_for_slot, actual_slot_num
                            )

                            print(
                                f"[DEBUG] _process_file_group (PARALLEL): Final counts - {len(images)} images, "
                                f"{len(hyperlinks)} hyperlinks for slot {slot['slot_number']}, subject {slot['subject']}"
                            )
                            logger.info(
                                "parallel_file_group_media_extracted",
                                extra={
                                    "slot": slot["slot_number"],
                                    "subject": slot["subject"],
                                    "images_count": len(images),
                                    "hyperlinks_count": len(hyperlinks),
                                },
                            )

                            # 5. Handle No School day/week filtering (parity with sequential path)
                            initial_hyperlink_count = len(hyperlinks)
                            if parser.is_no_school_day():
                                print(
                                    f"[DEBUG] _process_file_group (PARALLEL): No School day detected, "
                                    f"filtering out all {len(hyperlinks)} hyperlinks"
                                )
                                hyperlinks = []
                            elif "no_school_days" in content_data:
                                no_school_days_normalized = {
                                    day.lower().strip()
                                    for day in content_data["no_school_days"]
                                }
                                hyperlinks = [
                                    h
                                    for h in hyperlinks
                                    if not h.get("day_hint")
                                    or h.get("day_hint", "").lower().strip()
                                    not in no_school_days_normalized
                                ]
                                filtered_count = initial_hyperlink_count - len(
                                    hyperlinks
                                )
                                if filtered_count > 0:
                                    print(
                                        f"[DEBUG] _process_file_group (PARALLEL): Filtered out {filtered_count} hyperlinks "
                                        f"for No School days, {len(hyperlinks)} remaining"
                                    )
                                    logger.info(
                                        "parallel_hyperlinks_filtered_no_school",
                                        extra={
                                            "slot": slot["slot_number"],
                                            "subject": slot["subject"],
                                            "filtered_count": filtered_count,
                                            "remaining_count": len(hyperlinks),
                                        },
                                    )

                            context.extracted_content = content_data.get(
                                "full_text", ""
                            )
                            context.available_days = content_data.get(
                                "available_days", []
                            )
                            context.no_school_days = content_data.get(
                                "no_school_days", []
                            )

                            # Attach media to slot for transformation phase
                            context.slot["_extracted_images"] = images
                            context.slot["_extracted_hyperlinks"] = hyperlinks

                            print(
                                f"[DEBUG] _process_file_group (PARALLEL): Stored {len(hyperlinks)} hyperlinks "
                                f"in context.slot['_extracted_hyperlinks'] for slot {slot['slot_number']}"
                            )

                            # 6. Include media in content_json for DB caching
                            content_data["_extracted_images"] = images
                            content_data["_extracted_hyperlinks"] = hyperlinks
                            # 7. Add hyperlink preservation instruction to extracted_content
                            hyperlink_texts = [
                                h["text"] for h in hyperlinks if h.get("text")
                            ]
                            if hyperlink_texts:
                                preserve_instruction = (
                                    f"\n\nIMPORTANT: Preserve these exact phrases in your output "
                                    f"(they are hyperlinks): {', '.join(hyperlink_texts[:20])}"
                                )
                                context.extracted_content += preserve_instruction

                            # 8. Map to Schema and Update DB cache
                            metadata = content_data.get("metadata", {})
                            primary_teacher_name = (
                                metadata.get("primary_teacher_name")
                                or slot.get("primary_teacher_name")
                                or slot.get("teacher_name")
                            )

                            structured_days = {}
                            if "table_content" in content_data:
                                for day, day_data in content_data[
                                    "table_content"
                                ].items():
                                    day_lower = day.lower().strip()
                                    if day_lower in [
                                        "monday",
                                        "tuesday",
                                        "wednesday",
                                        "thursday",
                                        "friday",
                                    ]:
                                        # Filter hyperlinks for this day
                                        day_links = [
                                            h
                                            for h in hyperlinks
                                            if h.get("day", "").lower().strip()
                                            == day_lower
                                        ]
                                        structured_days[f"{day_lower}_content"] = (
                                            self._map_day_content_to_schema(
                                                day_data, slot, day_hyperlinks=day_links
                                            )
                                        )

                            import uuid

                            plan_data = {
                                "id": f"orig_{uuid.uuid4()}",
                                "user_id": user_id,
                                "week_of": week_of,
                                "slot_number": slot["slot_number"],
                                "subject": slot["subject"],
                                "grade": metadata.get("grade")
                                or slot.get("grade", "Unknown"),
                                "homeroom": metadata.get("homeroom")
                                or slot.get("homeroom"),
                                "source_file_path": file_path,
                                "source_file_name": path_obj.name,
                                "primary_teacher_name": primary_teacher_name,
                                "extracted_at": datetime.utcnow(),
                                "content_json": content_data,
                                "full_text": context.extracted_content,
                                "available_days": context.available_days,
                                "has_no_school": parser.is_no_school_day(),
                                "status": "extracted",
                                **structured_days,
                            }
                            await asyncio.to_thread(
                                self.db.create_original_lesson_plan, plan_data
                            )
                            contexts.append(context)

                        except Exception as e:
                            logger.error(
                                f"Error extracting slot {slot['slot_number']}: {e}"
                            )
                            context.error = str(e)
                            contexts.append(context)

            except Exception as e:
                logger.error(f"Failed to parse group file {file_path}: {e}")
                for i, slot, context in remaining_group:
                    context.error = f"Failed to parse DOCX: {str(e)}"
                    contexts.append(context)

        return contexts

    def _convert_originals_to_json(
        self, plans: List[OriginalLessonPlan]
    ) -> Dict[str, Any]:
        """
        Convert a list of OriginalLessonPlan objects to a multi-slot lesson JSON.
        This allows reusing the standard DOCXRenderer for original plans.
        """
        if not plans:
            return {}

        # 1. Metadata from first plan
        # When rendering a single plan, use that plan's specific metadata
        # When rendering multiple plans, use first plan's metadata but mark as combined
        first = plans[0]
        is_single_plan = len(plans) == 1
        
        metadata = {
            "primary_teacher_name": first.primary_teacher_name,
            "week_of": first.week_of,
            "grade": first.grade or "Unknown",
            # Always use the actual subject, not "Combined Originals" even for single plans
            # This ensures subject matches between database and rendered output
            "subject": first.subject,
            # Always preserve slot_number for single plans to maintain consistency
            "slot_number": first.slot_number if is_single_plan else None,
            "homeroom": first.homeroom,
            "source_type": "originals",
        }

        # 2. Days & Hyperlinks
        days_data = {}
        all_hyperlinks = []

        for day in ["monday", "tuesday", "wednesday", "thursday", "friday"]:
            slots = []
            for plan in plans:
                content = getattr(plan, f"{day}_content", None)
                if content:
                    # Enrich content with slot/subject for filtering in renderer
                    # content is already a dict from SQLModel JSON column
                    enriched = content.copy()

                    # 3. SCHEMA ALIGNMENT FOR RENDERER
                    # The renderer expects specific keys and nesting

                    # A. Hyperlinks: Collect for root elevation
                    # Renderer expects them in root _hyperlinks key with slot/subject metadata
                    if "hyperlinks" in enriched and enriched["hyperlinks"]:
                        links = enriched["hyperlinks"]
                        # Handle RootModel wrapping
                        if isinstance(links, dict) and "root" in links:
                            links = links["root"]

                        if isinstance(links, list):
                            for link in links:
                                if isinstance(link, dict):
                                    link_copy = link.copy()
                                    link_copy["_source_slot"] = plan.slot_number
                                    link_copy["_source_subject"] = plan.subject
                                    # All daily plan links belong in Table 1
                                    link_copy["table_idx"] = 1
                                    all_hyperlinks.append(link_copy)

                    # B. Instruction merge: Move instruction and tailored_instruction to renderer's expected keys
                    orig_instruction = enriched.get("instruction") or {}
                    orig_tailored = enriched.get("tailored_instruction") or {}

                    # Combine activities and tailored content for the main Instruction row
                    instr_parts = []
                    activities = (
                        orig_instruction.get("activities")
                        if isinstance(orig_instruction, dict)
                        else None
                    )
                    if activities:
                        instr_parts.append(activities)

                    tailored = (
                        orig_tailored.get("content")
                        if isinstance(orig_tailored, dict)
                        else None
                    )
                    if tailored:
                        # Deduplicate if activities already contains this text (common in fallbacks)
                        is_duplicate = any(
                            tailored in part or part in tailored for part in instr_parts
                        )
                        if not is_duplicate:
                            instr_parts.append(tailored)

                    # Materials must be nested inside tailored_instruction for the renderer's format function
                    materials_list = []
                    materials_data = enriched.get("materials")
                    if materials_data:
                        if isinstance(materials_data, list):
                            materials_list = materials_data
                        elif (
                            isinstance(materials_data, dict)
                            and "root" in materials_data
                        ):
                            materials_list = materials_data["root"]

                    # renderer uses 'tailored_instruction' field for the 'Instruction' row
                    enriched["tailored_instruction"] = {
                        "original_content": "\n\n".join(instr_parts),
                        "materials": materials_list,
                    }

                    # C. Misconceptions alignment
                    misconceptions = enriched.get("misconceptions")
                    if (
                        misconceptions
                        and isinstance(misconceptions, dict)
                        and "content" in misconceptions
                    ):
                        # Renderer expects 'original_content'
                        misconceptions["original_content"] = misconceptions["content"]

                    enriched["slot_number"] = plan.slot_number
                    enriched["subject"] = plan.subject
                    slots.append(enriched)

            if slots:
                if len(plans) == 1:
                    # For a single plan, return flat structure for _fill_single_slot_day
                    days_data[day] = slots[0]
                else:
                    # For multiple plans, use the multi-slot structure
                    days_data[day] = {"slots": slots}

        return {
            "metadata": metadata,
            "days": days_data,
            "_hyperlinks": all_hyperlinks,
            "_media_schema_version": "2.0",  # Enable semantic anchoring support
        }

    def _reconstruct_slots_from_json(
        self, lesson_json: Dict[str, Any]
    ) -> Dict[int, Dict[str, Any]]:
        """
        Reconstruct individual slot lesson plans from a multi-slot lesson JSON.
        Returns a mapping of slot_number -> lesson_entry.
        
        CRITICAL: Uses deep copies to prevent shared state in parallel processing.
        """
        import copy
        if not lesson_json or "days" not in lesson_json:
            return {}

        existing_slots_by_number = {}

        # 1. Extract slots from days
        for day_name, day_data in lesson_json.get("days", {}).items():
            day_slots = day_data.get("slots", [])
            if not day_slots and any(
                k in day_data for k in ["unit_lesson", "objective"]
            ):
                # Handle single-slot flattened format by temporarily wrapping it
                day_slots = [{**day_data}]

            for slot_data in day_slots:
                slot_num = slot_data.get("slot_number")
                if slot_num is not None:
                    if slot_num not in existing_slots_by_number:
                        existing_slots_by_number[slot_num] = {
                            "slot_number": slot_num,
                            "subject": slot_data.get("subject", "Unknown"),
                            "lesson_json": {
                                # CRITICAL: Use deep copy to prevent shared state in parallel processing
                                # Shallow copy (.copy()) would share nested dicts, causing race conditions
                                "metadata": copy.deepcopy(lesson_json.get("metadata", {})),
                                "days": {
                                    d: {}
                                    for d in [
                                        "monday",
                                        "tuesday",
                                        "wednesday",
                                        "thursday",
                                        "friday",
                                    ]
                                },
                            },
                        }
                        # Inject slot metadata for the renderer
                        existing_slots_by_number[slot_num]["lesson_json"]["metadata"][
                            "slot_number"
                        ] = slot_num
                        existing_slots_by_number[slot_num]["lesson_json"]["metadata"][
                            "subject"
                        ] = slot_data.get("subject", "Unknown")

                    # Add this day's content to the reconstructed lesson_json
                    day_content = slot_data.copy()
                    # Remove slot metadata fields
                    for field in [
                        "slot_number",
                        "subject",
                        "teacher_name",
                        "grade",
                        "homeroom",
                        "start_time",
                        "end_time",
                    ]:
                        day_content.pop(field, None)

                    existing_slots_by_number[slot_num]["lesson_json"]["days"][
                        day_name
                    ] = day_content

        # 2. Reconstruct _images and _hyperlinks per slot
        all_existing_images = lesson_json.get("_images", [])
        all_existing_links = lesson_json.get("_hyperlinks", [])

        # For robustness: if we have links but none have _source_slot,
        # and it's a single-slot plan (flattened), assign them all to the first slot.
        has_source_metadata = any(
            link.get("_source_slot") is not None for link in all_existing_links
        )
        is_effectively_single_slot = len(existing_slots_by_number) == 1

        for slot_num, lesson_entry in existing_slots_by_number.items():
            # Carry over media schema version (v2.0 required for links)
            lesson_entry["lesson_json"]["_media_schema_version"] = "2.0"

            if not has_source_metadata and is_effectively_single_slot:
                # Legacy or single-slot: assign all media to the only slot
                # and inject metadata for consistent filtering
                images_with_meta = []
                for img in all_existing_images:
                    img_copy = img.copy()
                    img_copy["_source_slot"] = slot_num
                    images_with_meta.append(img_copy)
                lesson_entry["lesson_json"]["_images"] = images_with_meta

                links_with_meta = []
                for link in all_existing_links:
                    link_copy = link.copy()
                    link_copy["_source_slot"] = slot_num
                    links_with_meta.append(link_copy)
                lesson_entry["lesson_json"]["_hyperlinks"] = links_with_meta
            else:
                # Standard multi-slot: filter by source slot
                lesson_entry["lesson_json"]["_images"] = [
                    img
                    for img in all_existing_images
                    if img.get("_source_slot") == slot_num
                ]
                lesson_entry["lesson_json"]["_hyperlinks"] = [
                    link
                    for link in all_existing_links
                    if link.get("_source_slot") == slot_num
                ]

        return existing_slots_by_number

    async def _generate_combined_original_docx(
        self,
        user_id: str,
        week_of: str,
        plan_id: str,
        week_folder_path: Optional[str] = None,
    ) -> Optional[str]:
        """
        Generates a combined DOCX of all original plans for the week using structured schema data.
        Reuses the standard DOCXRenderer for visual consistency.
        """
        logger.info(f"Generating combined original DOCX for {user_id}, week {week_of}")

        # 1. Fetch
        plans = await asyncio.to_thread(
            self.db.get_original_lesson_plans_for_week, user_id, week_of
        )
        if not plans:
            logger.warning(
                "combined_originals_skipped_no_plans",
                extra={"user_id": user_id, "week": week_of},
            )
            return None

        plans.sort(key=lambda x: x.slot_number)

        # Log all plans fetched from database
        logger.info(
            "combined_originals_plans_fetched",
            extra={
                "total_plans": len(plans),
                "plans": [
                    {
                        "slot_number": p.slot_number,
                        "subject": p.subject,
                        "grade": p.grade,
                        "source_file_name": p.source_file_name,
                        "primary_teacher_name": p.primary_teacher_name,
                    }
                    for p in plans
                ],
            },
        )

        # 2. Render each slot individually and merge
        def _render_originals():
            temp_files = []
            rendering_results = {}  # Track success/failure for each plan
            try:
                from docx import Document

                from tools.docx_renderer import DOCXRenderer
                from tools.docx_utils import normalize_styles_from_master

                # Prepare output path
                file_mgr = get_file_manager(
                    base_path=getattr(self, "_user_base_path", None)
                )
                week_folder = (
                    Path(week_folder_path)
                    if week_folder_path
                    else file_mgr.get_week_folder(week_of)
                )
                originals_dir = week_folder / "originals"
                originals_dir.mkdir(parents=True, exist_ok=True)

                safe_week = (
                    week_of.replace("/", "-").replace("\\", "-").replace(" ", "_")
                )
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"combined_originals_{safe_week}_{timestamp}.docx"
                output_path = originals_dir / filename

                logger.info(
                    "combined_originals_processing_started",
                    extra={
                        "total_plans": len(plans),
                    },
                )

                # Instantiate renderer
                renderer = DOCXRenderer(settings.DOCX_TEMPLATE_PATH)
                renderer.is_originals = True

                # Master doc for style truth
                style_master = Document(settings.DOCX_TEMPLATE_PATH)

                # Deduplicate plans by (Source File, Slot Number, Subject) to handle cases
                # where the same slot number from different files with different subjects should both be included.
                # 1. Same File, Different Slot -> KEEP (Different subjects).
                # 2. Same File, Same Slot -> MERGE (Duplicate).
                # 3. Different File, Same Slot, Different Subject -> KEEP (Both needed).
                unique_slot_files = set()
                deduplicated_plans = []
                removed_plans = []
                
                for plan in plans:
                    source_key = getattr(plan, "source_file_path", None)
                    slot_key = plan.slot_number
                    subject_key = plan.subject
                    
                    # Key includes file path, slot number, AND subject to handle edge cases
                    dedup_key = (source_key, slot_key, subject_key)
                    
                    if dedup_key not in unique_slot_files:
                        unique_slot_files.add(dedup_key)
                        deduplicated_plans.append(plan)
                    else:
                        removed_plans.append({
                            "slot_number": plan.slot_number,
                            "subject": plan.subject,
                            "source_file_name": plan.source_file_name,
                        })

                logger.info(
                    "combined_originals_deduplication",
                    extra={
                        "original_count": len(plans),
                        "deduplicated_count": len(deduplicated_plans),
                        "unique_keys_count": len(unique_slot_files),
                        "removed_plans": removed_plans,
                        "deduplicated_plans": [
                            {
                                "slot_number": p.slot_number,
                                "subject": p.subject,
                                "source_file_name": p.source_file_name,
                            }
                            for p in deduplicated_plans
                        ],
                    }
                )
                
                # Validation: Verify all database plans are in deduplicated_plans
                original_slot_subjects = {(p.slot_number, p.subject) for p in plans}
                dedup_slot_subjects = {(p.slot_number, p.subject) for p in deduplicated_plans}
                missing_combinations = original_slot_subjects - dedup_slot_subjects
                if missing_combinations:
                    logger.warning(
                        "combined_originals_missing_after_dedup",
                        extra={
                            "missing_combinations": [
                                {"slot_number": s, "subject": subj}
                                for s, subj in missing_combinations
                            ]
                        },
                    )

                for plan in deduplicated_plans:
                    # Create temp filename for this slot
                    temp_filename = f"_temp_orig_slot{plan.slot_number}_{plan.subject.replace('/', '_')}.docx"
                    temp_path = str(originals_dir / temp_filename)

                    try:
                        logger.info(
                            "combined_originals_rendering_plan",
                            extra={
                                "slot_number": plan.slot_number,
                                "subject": plan.subject,
                                "source_file_name": plan.source_file_name,
                                "temp_path": temp_path,
                            },
                        )

                        # Convert single plan to JSON
                        slot_json = self._convert_originals_to_json([plan])
                        
                        # Validate JSON contains correct slot_number and subject
                        if slot_json.get("metadata", {}).get("slot_number") != plan.slot_number:
                            logger.warning(
                                "combined_originals_slot_number_mismatch",
                                extra={
                                    "expected": plan.slot_number,
                                    "actual": slot_json.get("metadata", {}).get("slot_number"),
                                },
                            )
                        if slot_json.get("metadata", {}).get("subject") != plan.subject:
                            logger.warning(
                                "combined_originals_subject_mismatch",
                                extra={
                                    "expected": plan.subject,
                                    "actual": slot_json.get("metadata", {}).get("subject"),
                                },
                            )

                        # Render this slot to temporary file
                        # CRITICAL: Set skip_fallback_sections=True to prevent duplicate sections
                        # Use unpacking to handle tuple return when skip_fallback_sections=True
                        render_result = renderer.render(
                            slot_json,
                            temp_path,
                            plan_id=plan_id,
                            skip_fallback_sections=True,
                        )
                        success, unplaced_hl, unplaced_img = (
                            render_result
                            if isinstance(render_result, tuple)
                            else (render_result, [], [])
                        )

                        if success:
                            # Re-open and apply aggressive cleanups
                            sub_doc = Document(temp_path)
                            from tools.docx_utils import (
                                remove_headers_footers,
                                strip_custom_styles,
                                strip_problematic_elements,
                                strip_sections,
                            )

                            strip_problematic_elements(sub_doc)
                            remove_headers_footers(sub_doc)
                            strip_sections(sub_doc)
                            strip_custom_styles(sub_doc, style_master)

                            # Normalize styles
                            from tools.docx_utils import normalize_styles_from_master

                            normalize_styles_from_master(style_master, sub_doc)
                            if hasattr(sub_doc, "_normalized_stream"):
                                sub_doc._normalized_stream.seek(0)
                                sub_doc = Document(sub_doc._normalized_stream)

                            sub_doc.save(temp_path)
                            temp_files.append(temp_path)
                            rendering_results[plan.slot_number] = {
                                "success": True,
                                "subject": plan.subject,
                                "temp_path": temp_path,
                            }
                            logger.info(
                                "combined_originals_plan_rendered_successfully",
                                extra={
                                    "slot_number": plan.slot_number,
                                    "subject": plan.subject,
                                },
                            )
                        else:
                            rendering_results[plan.slot_number] = {
                                "success": False,
                                "subject": plan.subject,
                                "error": "Render returned False",
                            }
                            logger.error(
                                "combined_originals_plan_render_failed",
                                extra={
                                    "slot_number": plan.slot_number,
                                    "subject": plan.subject,
                                    "unplaced_hyperlinks": len(unplaced_hl),
                                    "unplaced_images": len(unplaced_img),
                                },
                            )
                    except Exception as e:
                        rendering_results[plan.slot_number] = {
                            "success": False,
                            "subject": plan.subject,
                            "error": str(e),
                        }
                        logger.error(
                            "combined_originals_plan_render_exception",
                            extra={
                                "slot_number": plan.slot_number,
                                "subject": plan.subject,
                                "error": str(e),
                            },
                            exc_info=True,
                        )

                # Log rendering results summary
                successful_slots = [
                    slot_num
                    for slot_num, result in rendering_results.items()
                    if result.get("success")
                ]
                failed_slots = [
                    slot_num
                    for slot_num, result in rendering_results.items()
                    if not result.get("success")
                ]
                
                logger.info(
                    "combined_originals_rendering_summary",
                    extra={
                        "total_plans": len(deduplicated_plans),
                        "successful_renders": len(successful_slots),
                        "failed_renders": len(failed_slots),
                        "successful_slots": successful_slots,
                        "failed_slots": [
                            {
                                "slot_number": slot_num,
                                "subject": rendering_results[slot_num].get("subject"),
                                "error": rendering_results[slot_num].get("error"),
                            }
                            for slot_num in failed_slots
                        ],
                    },
                )

                if not temp_files:
                    logger.warning(
                        "combined_originals_no_files_generated",
                        extra={
                            "total_plans": len(deduplicated_plans),
                            "rendering_results": rendering_results,
                        },
                    )
                    return None

                # Validation: Check if expected slots are missing
                expected_slots = {p.slot_number for p in deduplicated_plans}
                rendered_slots = {p.slot_number for p in deduplicated_plans if rendering_results.get(p.slot_number, {}).get("success")}
                missing_slots = expected_slots - rendered_slots
                if missing_slots:
                    logger.warning(
                        "combined_originals_missing_slots",
                        extra={
                            "missing_slots": [
                                {
                                    "slot_number": slot_num,
                                    "subject": next(
                                        (p.subject for p in deduplicated_plans if p.slot_number == slot_num),
                                        "Unknown",
                                    ),
                                }
                                for slot_num in missing_slots
                            ],
                        },
                    )

                # Use the existing _merge_docx_files method to combine files from disk
                # Use the template as the master to ensure 100% style consistency
                logger.info(
                    "combined_originals_merging_files",
                    extra={
                        "file_count": len(temp_files),
                        "file_paths": [Path(f).name for f in temp_files],
                    },
                )
                self._merge_docx_files(
                    temp_files,
                    str(output_path),
                    master_template_path=settings.DOCX_TEMPLATE_PATH,
                )

                # Log final merged document structure
                try:
                    merged_doc = Document(str(output_path))
                    table_count = len(merged_doc.tables)
                    logger.info(
                        "combined_originals_merged_structure",
                        extra={
                            "table_count": table_count,
                            "estimated_slots": table_count // 2,
                        },
                    )
                except Exception as e:
                    logger.warning(
                        "combined_originals_structure_check_failed",
                        extra={"error": str(e)},
                    )

                logger.info(
                    "combined_originals_generated_successfully",
                    extra={
                        "user_id": user_id,
                        "week": week_of,
                        "output_path": str(output_path.absolute()),
                        "size_bytes": output_path.stat().st_size if output_path.exists() else 0,
                        "successful_slots": successful_slots,
                        "failed_slots": failed_slots,
                    },
                )
                print(f"DEBUG: Combined originals DOCX created: {output_path}")
                return str(output_path)

            except Exception as e:
                logger.error(f"Error rendering originals doc: {e}")
                traceback.print_exc()
                return None
            finally:
                # Cleanup temp files
                import os

                for tf in temp_files:
                    try:
                        if os.path.exists(tf):
                            os.remove(tf)
                    except Exception:
                        pass

        return await asyncio.to_thread(_render_originals)


async def process_batch(
    user_id: str, week_of: str, provider: str = "openai"
) -> Dict[str, Any]:
    """Convenience function to process a batch.

    Args:
        user_id: User ID
        week_of: Week date range
        provider: LLM provider

    Returns:
        Processing results
    """
    from backend.llm_service import get_llm_service

    llm_service = get_llm_service()
    processor = BatchProcessor(llm_service)

    return await processor.process_user_week(user_id, week_of, provider)
