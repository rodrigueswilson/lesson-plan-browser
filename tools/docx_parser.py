"""
DOCX Parser for extracting content from primary teacher lesson plan files.
Handles various DOCX formats and structures to extract lesson content.
"""

import base64
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from backend.performance_tracker import get_tracker


def validate_slot_structure(doc, slot_number: int) -> Tuple[int, int]:
    """
    Validate document structure for slot extraction.

    Ensures document has expected table layout (metadata + daily plans per slot + signature).
    Fails loudly with detailed error messages for unexpected structures.

    Args:
        doc: Document object (from python-docx)
        slot_number: Requested slot number (1-indexed)

    Raises:
        ValueError: If structure is invalid with detailed message

    Returns:
        tuple: (table_start, table_end) indices for the slot (0-indexed)
    """
    from backend.telemetry import logger

    table_count = len(doc.tables)

    if table_count == 0:
        raise ValueError("Document has no tables")

    # Check for signature table (last table) - optional validation
    last_table = doc.tables[-1]
    has_signature = False

    if last_table.rows and last_table.rows[0].cells:
        first_cell = last_table.rows[0].cells[0].text.strip().lower()
        if "signature" in first_cell or "required signatures" in first_cell:
            has_signature = True

    # Calculate available slots (excluding signature if present)
    if has_signature:
        available_slots = (table_count - 1) // 2
    else:
        # No signature table - all tables are slot tables
        available_slots = table_count // 2
        logger.warning(
            "no_signature_table",
            extra={
                "table_count": table_count,
                "available_slots": available_slots,
                "last_table_first_cell": last_table.rows[0].cells[0].text.strip() if last_table.rows and last_table.rows[0].cells else "empty",
            },
        )

    # Validate table count matches expected pattern
    if has_signature:
        expected_table_count = (available_slots * 2) + 1
        if table_count != expected_table_count:
            # Generate valid counts dynamically (1-10 slots)
            valid_counts = [
                2 * n + 1 for n in range(1, 11)
            ]  # 3, 5, 7, 9, 11, 13, 15, 17, 19, 21
            valid_examples = ", ".join(
                f"{c} ({(c - 1) // 2} slot{'s' if (c - 1) // 2 > 1 else ''})"
                for c in valid_counts[:5]
            )
            raise ValueError(
                f"Unexpected table count: {table_count}. "
                f"Expected {expected_table_count} for {available_slots} slots. "
                f"Examples of valid counts: {valid_examples}."
            )
    else:
        # No signature table - must have even number of tables (pairs of metadata + daily plan)
        expected_table_count = available_slots * 2
        if table_count != expected_table_count:
            raise ValueError(
                f"Unexpected table count: {table_count}. "
                f"Expected {expected_table_count} (even number) for {available_slots} slots "
                f"without signature table. Each slot requires 2 tables (metadata + daily plan)."
            )

    # Validate requested slot exists
    if slot_number < 1:
        raise ValueError(f"Slot number must be >= 1, got {slot_number}")

    if slot_number > available_slots:
        raise ValueError(
            f"Slot {slot_number} requested but only {available_slots} slots available. "
            f"Document has {table_count} tables."
        )

    # Calculate table indices for this slot
    table_start = (slot_number - 1) * 2
    table_end = table_start + 1

    # Validate metadata table (relaxed: case-insensitive, whitespace-tolerant)
    meta_table = doc.tables[table_start]
    if not meta_table.rows or not meta_table.rows[0].cells:
        raise ValueError(
            f"Slot {slot_number} metadata table (index {table_start}) is empty"
        )

    first_cell = meta_table.rows[0].cells[0].text.strip().lower()
    # Allow variations: "name:", "name(s):", "teacher:", "teacher name:", etc.
    metadata_indicators = ["name", "teacher"]
    has_metadata_indicator = any(
        indicator in first_cell for indicator in metadata_indicators
    )

    if not has_metadata_indicator:
        raise ValueError(
            f"Slot {slot_number} table {table_start} doesn't look like metadata table. "
            f"Expected 'Name:', 'Teacher:', or similar but got: '{meta_table.rows[0].cells[0].text.strip()[:50]}'"
        )

    # Validate daily plans table (relaxed: check all cells, strip punctuation)
    daily_table = doc.tables[table_end]
    if not daily_table.rows or not daily_table.rows[0].cells:
        raise ValueError(
            f"Slot {slot_number} daily plans table (index {table_end}) is empty"
        )

    # Check all cells in first row, not just first 5
    first_row_text = " ".join(cell.text.strip() for cell in daily_table.rows[0].cells)
    # Remove punctuation and normalize
    import string

    first_row_clean = first_row_text.translate(
        str.maketrans("", "", string.punctuation)
    ).upper()

    weekdays = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY"]
    has_weekday = any(day in first_row_clean for day in weekdays)

    # Also check for metadata indicators - if found, this might be a metadata table instead
    metadata_indicators = ["UNIT", "LESSON", "MODULE", "SUBJECT", "OBJECTIVE"]
    has_metadata_indicator = any(
        indicator in first_row_clean for indicator in metadata_indicators
    )

    if not has_weekday:
        # If it looks like metadata instead of daily plans, log a warning but proceed
        # Some documents may have different structures
        if has_metadata_indicator:
            logger.warning(
                "daily_plans_table_looks_like_metadata",
                extra={
                    "slot_number": slot_number,
                    "table_index": table_end,
                    "first_row": first_row_text[:200],
                    "note": "Table appears to contain metadata instead of daily plans. Proceeding anyway."
                }
            )
            # Don't raise - allow processing to continue
        else:
            raise ValueError(
                f"Slot {slot_number} table {table_end} doesn't look like daily plans table. "
                f"Expected weekday headers but got: '{first_row_text[:100]}'"
            )

    logger.info(
        "slot_structure_validated",
        extra={
            "slot_number": slot_number,
            "table_start": table_start,
            "table_end": table_end,
            "total_tables": table_count,
            "available_slots": available_slots,
        },
    )

    return table_start, table_end


class DOCXParser:
    """Parser for primary teacher DOCX lesson plan files."""

    def __init__(self, file_path: str):
        """Initialize parser with DOCX file.

        Args:
            file_path: Path to DOCX file
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        self.doc = Document(str(self.file_path))
        self.paragraphs = [
            p.text.strip() for p in self.doc.paragraphs if p.text.strip()
        ]
        self.tables = self._extract_tables()

    def _extract_tables(self) -> List[List[List[str]]]:
        """Extract all tables from document.

        Returns:
            List of tables, where each table is a list of rows,
            and each row is a list of cell values
        """
        tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def _is_signature_content(self, text: str) -> bool:
        """Check if text is part of signature/certification section.

        Args:
            text: Text to check

        Returns:
            True if text is signature content to exclude
        """
        signature_patterns = [
            r"i certify",
            r"certification",
            r"teacher signature",
            r"principal signature",
            r"administrator signature",
            r"aligned with.*standards",
            r"new jersey student learning",
            r"njsls",
            r"signature:",
            r"date:.*signature",
        ]

        text_lower = text.lower()
        return any(re.search(pattern, text_lower) for pattern in signature_patterns)

    def get_full_text(
        self, exclude_signatures: bool = True, strip_urls: bool = False
    ) -> str:
        """Get all text content from document.

        Args:
            exclude_signatures: If True, filter out signature/certification sections
            strip_urls: If True, remove URLs to save tokens (hyperlinks preserved separately)

        Returns:
            Full text content
        """
        paragraphs = self.paragraphs

        # Filter out signature content
        if exclude_signatures:
            paragraphs = [p for p in paragraphs if not self._is_signature_content(p)]

        # Strip URLs if requested (saves LLM tokens)
        if strip_urls:
            url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
            paragraphs = [re.sub(url_pattern, "", p) for p in paragraphs]

        return "\n".join(paragraphs)

    def _get_no_school_patterns(self) -> list:
        """Get regex patterns for detecting No School days."""
        return [
            # Core "No School" patterns
            r"no\s+school",  # "No School"
            r"no\s*-\s*school",  # "No-School" or "No - School"
            r"school\s+closed",  # "School Closed"
            # Holiday patterns
            r"holiday",  # "Holiday"
            r"vacation\s+day",  # "Vacation Day"
            # Development/Training patterns
            r"professional\s+development",  # "Professional Development"
            r"staff\s+development",  # "Staff Development"
            r"teacher\s+development",  # "Teacher Development"
            r"pd\s+day",  # "PD Day"
            r"in[-\s]?service",  # "In-Service" or "Inservice"
            # Workday/Planning patterns
            r"teacher\s+workday",  # "Teacher Workday"
            r"planning\s+day",  # "Planning Day"
            r"prep\s+day",  # "Prep Day"
            # Conference patterns
            r"conference\s+day",  # "Conference Day"
            r"parent[-\s]teacher\s+conference",  # "Parent-Teacher Conference"
            # Dismissal patterns (often no regular instruction)
            r"early\s+dismissal",  # "Early Dismissal"
            r"half\s+day",  # "Half Day"
            r"early\s+release",  # "Early Release"
        ]

        # Get full text and convert to lowercase for case-insensitive matching
        full_text = self.get_full_text().lower()

        # Check each pattern
        for pattern in patterns:
            if re.search(pattern, full_text):
                # Log detection for debugging
                from backend.telemetry import logger

                logger.info(
                    "no_school_detected",
                    extra={"pattern": pattern, "file": str(self.file_path.name)},
                )
                return True

        return False

    def is_no_school_day(self) -> bool:
        """
        Check if document indicates 'No School' day.

        Detects various patterns indicating school is closed.

        Returns:
            True if "No School" indicators found, False otherwise
        """
        patterns = self._get_no_school_patterns()
        full_text = self.get_full_text(exclude_signatures=True, strip_urls=False)

        for pattern in patterns:
            if re.search(pattern, full_text, re.IGNORECASE):
                return True

        return False

    def is_day_no_school(self, day_text: str) -> bool:
        """
        Check if a specific day's content indicates 'No School'.

        Args:
            day_text: Text content for a specific day

        Returns:
            True if day is "No School", False otherwise
        """
        if not day_text or len(day_text.strip()) < 5:
            return False

        # Normalize day text for checking
        day_text_lower = day_text.lower().strip()
        
        # Require more specific patterns to avoid false positives
        # Only match if the pattern appears as a standalone phrase or at the start
        # This prevents matching "no school" when it appears in context like "no school supplies needed"
        specific_patterns = [
            r"^no\s+school\s*$",  # "No School" as standalone
            r"^no\s+school\s+[,\-\.]",  # "No School," or "No School -"
            r"^no\s*-\s*school\s*$",  # "No-School" as standalone
            r"^holiday\s*$",  # "Holiday" as standalone
            r"^school\s+closed\s*$",  # "School Closed" as standalone
        ]
        
        # Check specific patterns first (more restrictive)
        for pattern in specific_patterns:
            if re.search(pattern, day_text_lower, re.IGNORECASE):
                return True
        
        # If day text is very short and matches a pattern, it's likely "No School"
        # Include development/workday/planning/conference patterns (same as document-level)
        # but only for short text to avoid false positives in long lesson content
        if len(day_text.strip()) < 30:  # Very short content
            core_patterns = [
                r"no\s+school",
                r"no\s*-\s*school",
                r"school\s+closed",
                r"^holiday$",
                r"staff\s+development",
                r"professional\s+development",
                r"pd\s+day",
                r"planning\s+day",
                r"teacher\s+workday",
                r"prep\s+day",
                r"conference\s+day",
                r"in[-\s]?service",
            ]
            for pattern in core_patterns:
                if re.search(pattern, day_text_lower, re.IGNORECASE):
                    return True
        
        # For longer content, be more conservative - only match if pattern is clearly indicating no school
        # and not part of lesson content
        if len(day_text.strip()) >= 30:
            # Only match if "no school" appears at the beginning or as a clear standalone statement
            strict_patterns = [
                r"^(no\s+school|no\s*-\s*school|school\s+closed|holiday)[\s,\-\.]",
                r"[\s,\-\.](no\s+school|no\s*-\s*school|school\s+closed|holiday)\s*$",
            ]
            for pattern in strict_patterns:
                if re.search(pattern, day_text_lower, re.IGNORECASE):
                    # Additional check: make sure it's not part of a longer sentence
                    # If there are multiple sentences or complex content, it's probably not "No School"
                    if day_text_lower.count('.') <= 1 and day_text_lower.count(',') <= 2:
                        return True

        return False

    def find_subject_sections(self) -> Dict[str, Any]:
        """
        Find different subject sections in the document.
        Handles both text-based and table-based formats.

        Returns:
            Dictionary mapping subject names to their table indices or text positions
        """
        sections = {}

        # Method 1: Check tables for subject headers (Davies' format)
        for i, table in enumerate(self.doc.tables):
            # Check if this is a header table (1 row with subject info)
            if len(table.rows) == 1:
                for cell in table.rows[0].cells:
                    text = cell.text.strip()
                    if "Subject:" in text:
                        subject = text.replace("Subject:", "").strip()
                        sections[subject] = {
                            "type": "table",
                            "header_table": i,
                            "lesson_table": i + 1
                            if i + 1 < len(self.doc.tables)
                            else None,
                        }
                        break

        # Method 2: Check text for subject headings (fallback)
        if not sections:
            full_text = self.get_full_text()
            subjects = [
                "math",
                "science",
                "ela",
                "english",
                "social studies",
                "history",
            ]

            for subject in subjects:
                pattern = rf"\b{subject}\b"
                if re.search(pattern, full_text, re.IGNORECASE):
                    sections[subject.title()] = {"type": "text", "found": True}

        return sections

    def extract_by_heading(
        self, heading_text: str, case_sensitive: bool = False
    ) -> List[str]:
        """Extract content following a specific heading.

        Args:
            heading_text: Text to search for in headings
            case_sensitive: Whether to match case exactly

        Returns:
            List of paragraphs following the heading
        """
        content = []
        found_heading = False

        for para in self.paragraphs:
            if not case_sensitive:
                match = heading_text.lower() in para.lower()
            else:
                match = heading_text in para

            if match and len(para) < 100:  # Likely a heading
                found_heading = True
                continue

            if found_heading:
                # Stop at next heading (short paragraph)
                if len(para) < 50 and para.isupper():
                    break
                content.append(para)

        return content

    def extract_table_by_header(self, header_text: str) -> Optional[List[List[str]]]:
        """Extract table that contains specific header text.

        Args:
            header_text: Text to search for in table headers

        Returns:
            Table data or None if not found
        """
        for table in self.tables:
            if not table:
                continue

            # Check first row for header
            first_row = table[0]
            for cell in first_row:
                if header_text.lower() in cell.lower():
                    return table

        return None

    def extract_table_lesson(self, table_index: int) -> Dict[str, Any]:
        """
        Extract lesson content from a table (Davies' format).

        Args:
            table_index: Index of the lesson table

        Returns:
            Dictionary with extracted lesson content
        """
        if table_index >= len(self.doc.tables):
            return {}

        table = self.doc.tables[table_index]
        content = {}

        # Get column headers (days of week)
        days = []
        for cell in table.rows[0].cells[1:]:  # Skip first column
            day = cell.text.strip()
            if day:
                days.append(day)

        # Get row labels and content
        for row in table.rows[1:]:
            row_label = row.cells[0].text.strip()
            if not row_label:
                continue

            # Extract content for each day
            for i, day in enumerate(days):
                if day not in content:
                    content[day] = {}

                cell = row.cells[i + 1]
                cell_text = self._get_cell_text_with_hyperlinks(cell)
                if cell_text:
                    content[day][row_label] = cell_text

        return content

    def _get_cell_text_with_hyperlinks(self, cell) -> str:
        """Extract text from a cell, injecting markdown hyperlinks where they exist."""
        from docx.oxml.ns import qn
        
        full_text = []
        for paragraph in cell.paragraphs:
            para_parts = []
            
            # Use XPath to find all runs and hyperlinks in order
            for child in paragraph._element:
                if child.tag == qn("w:r"):
                    # Regular text run
                    text = "".join(node.text for node in child.xpath(".//w:t") if node.text)
                    if text:
                        para_parts.append(text)
                elif child.tag == qn("w:hyperlink"):
                    # Hyperlink - Convert to Markdown [text](url)
                    try:
                        r_id = child.get(qn("r:id"))
                        if r_id and r_id in paragraph.part.rels:
                            url = paragraph.part.rels[r_id].target_ref
                            text = "".join(node.text for node in child.xpath(".//w:t") if node.text)
                            if text and url:
                                # Inject as markdown right in the text stream
                                para_parts.append(f"[{text}]({url})")
                            elif text:
                                para_parts.append(text)
                    except Exception:
                        # Fallback to plain text if link resolution fails
                        text = "".join(node.text for node in child.xpath(".//w:t") if node.text)
                        if text:
                            para_parts.append(text)
            
            para_text = "".join(para_parts).strip()
            if para_text:
                full_text.append(para_text)
                
        return "\n".join(full_text).strip()

    def extract_subject_content(
        self, subject: str, strip_urls: bool = True
    ) -> Dict[str, Any]:
        """
        Extract content for a specific subject.
        Handles both table-based (Davies) and text-based formats.

        Args:
            subject: Subject name to extract
            strip_urls: If True, remove URLs from text (saves LLM tokens)

        Returns:
            Dictionary with subject content
        """
        # Normalize subject names for matching
        subject_mappings = {
            "language arts": ["ela", "english", "reading", "literacy"],
            "math": ["mathematics", "math"],
            "science": ["science", "sci", "science/health"],
            "social studies": ["ss", "social studies", "history"],
            "ela": ["ela", "english", "language arts", "reading"],
            "ela/ss": ["ela/ss", "language arts/social studies"],
        }

        sections = self.find_subject_sections()

        # Find matching subject (case-insensitive with aliases)
        subject_lower = subject.lower()

        # Get possible aliases for this subject
        possible_names = [subject_lower]
        for canonical, aliases in subject_mappings.items():
            if subject_lower in aliases or subject_lower == canonical:
                possible_names.extend(aliases)
                break

        for key, info in sections.items():
            key_lower = key.lower()
            # Check if any possible name matches
            if any(name in key_lower for name in possible_names):
                # Handle table-based format
                if info.get("type") == "table" and info.get("lesson_table") is not None:
                    table_content = self.extract_table_lesson(info["lesson_table"])

                    # Convert to text format, filtering No School days
                    full_text_parts = []
                    no_school_days = []

                    for day, day_content in table_content.items():
                        # Check if this day is "No School" - but preserve all content regardless
                        day_text = " ".join(day_content.values())
                        is_no_school = self.is_day_no_school(day_text)
                        
                        if is_no_school:
                            no_school_days.append(day)
                            # Mark as "No School" but still include all content from DOCX
                            # This allows LLM to preserve tailored_instruction and other fields
                            full_text_parts.append(f"\n{day.upper()}")
                            full_text_parts.append("Unit/Lesson: No School")
                            # Include all other content from DOCX
                            for label, text in day_content.items():
                                # Skip unit/lesson field if it's "No School" to avoid duplication
                                label_lower = label.lower()
                                if "unit" not in label_lower and "lesson" not in label_lower:
                                    full_text_parts.append(f"{label} {text}")
                                elif text.strip() and text.strip().lower() != "no school":
                                    # If unit/lesson has actual content (not just "No School"), include it
                                    full_text_parts.append(f"{label} {text}")
                        else:
                            # Include full content for regular days
                            full_text_parts.append(f"\n{day.upper()}")
                            for label, text in day_content.items():
                                full_text_parts.append(f"{label} {text}")

                    return {
                        "subject": key,
                        "full_text": "\n".join(full_text_parts),
                        "table_content": table_content,
                        "no_school_days": no_school_days,
                        "found": True,
                        "format": "table",
                    }

                # Handle text-based format
                elif isinstance(info, list):
                    return {
                        "subject": key,
                        "full_text": "\n".join(info),
                        "found": True,
                        "format": "text",
                    }

        # If not found, return all text
        return {
            "subject": subject,
            "full_text": self.get_full_text(strip_urls=strip_urls),
            "found": False,
            "format": "unknown",
        }

    def _extract_objectives(self, content: List[str]) -> List[str]:
        """Extract learning objectives from content.

        Args:
            content: List of content paragraphs

        Returns:
            List of objectives
        """
        objectives = []
        in_objectives = False

        for para in content:
            para_lower = para.lower()

            # Check for objectives section
            if any(
                keyword in para_lower
                for keyword in [
                    "objective",
                    "goal",
                    "learning target",
                    "swbat",
                    "students will",
                ]
            ):
                in_objectives = True

                # If objective is in same line
                if ":" in para:
                    obj_text = para.split(":", 1)[1].strip()
                    if obj_text:
                        objectives.append(obj_text)
                continue

            if in_objectives:
                # Stop at next section
                if any(
                    keyword in para_lower
                    for keyword in ["activity", "procedure", "assessment", "material"]
                ):
                    in_objectives = False
                else:
                    objectives.append(para)

        return objectives

    def _extract_activities(self, content: List[str]) -> List[str]:
        """Extract learning activities from content.

        Args:
            content: List of content paragraphs

        Returns:
            List of activities
        """
        activities = []
        in_activities = False

        for para in content:
            para_lower = para.lower()

            # Check for activities section
            if any(
                keyword in para_lower
                for keyword in [
                    "activity",
                    "activities",
                    "procedure",
                    "lesson",
                    "instruction",
                ]
            ):
                in_activities = True

                # If activity is in same line
                if ":" in para:
                    act_text = para.split(":", 1)[1].strip()
                    if act_text:
                        activities.append(act_text)
                continue

            if in_activities:
                # Stop at next section
                if any(
                    keyword in para_lower
                    for keyword in ["assessment", "material", "homework", "closure"]
                ):
                    in_activities = False
                else:
                    activities.append(para)

        return activities

    def _extract_assessments(self, content: List[str]) -> List[str]:
        """Extract assessments from content.

        Args:
            content: List of content paragraphs

        Returns:
            List of assessments
        """
        assessments = []
        in_assessments = False

        for para in content:
            para_lower = para.lower()

            # Check for assessment section
            if any(
                keyword in para_lower
                for keyword in [
                    "assessment",
                    "evaluate",
                    "check for understanding",
                    "exit ticket",
                ]
            ):
                in_assessments = True

                # If assessment is in same line
                if ":" in para:
                    assess_text = para.split(":", 1)[1].strip()
                    if assess_text:
                        assessments.append(assess_text)
                continue

            if in_assessments:
                # Stop at next section
                if any(
                    keyword in para_lower
                    for keyword in ["material", "homework", "closure", "extension"]
                ):
                    in_assessments = False
                else:
                    assessments.append(para)

        return assessments

    def _extract_materials(self, content: List[str]) -> List[str]:
        """Extract materials/resources from content.

        Args:
            content: List of content paragraphs

        Returns:
            List of materials
        """
        materials = []
        in_materials = False

        for para in content:
            para_lower = para.lower()

            # Check for materials section
            if any(
                keyword in para_lower
                for keyword in ["material", "resource", "supplies", "equipment"]
            ):
                in_materials = True

                # If materials are in same line
                if ":" in para:
                    mat_text = para.split(":", 1)[1].strip()
                    if mat_text:
                        materials.append(mat_text)
                continue

            if in_materials:
                # Stop at next section
                if any(
                    keyword in para_lower
                    for keyword in ["objective", "activity", "procedure", "assessment"]
                ):
                    in_materials = False
                else:
                    materials.append(para)

        return materials

    def list_available_subjects(self) -> List[str]:
        """List all potential subjects found in the document.

        Returns:
            List of subject names found
        """
        sections = self.find_subject_sections()
        return list(sections.keys())

    def extract_week_info(self) -> Optional[str]:
        """Try to extract week information from document.

        Returns:
            Week string if found, None otherwise
        """
        # Look for week patterns
        week_patterns = [
            r"week\s+of\s+(\d{1,2}/\d{1,2})",
            r"week\s+(\d{1,2})",
            r"(\d{1,2}/\d{1,2}\s*-\s*\d{1,2}/\d{1,2})",
        ]

        full_text = self.get_full_text()

        for pattern in week_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                return match.group(1)

        return None

    def extract_images(self, exclude_signatures: bool = True) -> List[Dict[str, Any]]:
        """Extract all images from document with context for semantic anchoring.

        Args:
            exclude_signatures: If True, filter out signature/certification images

        Returns:
            List of image dictionaries with data, content_type, filename, and context
        """
        images = []
        try:
            for rel_id, rel in self.doc.part.rels.items():
                if "image" in rel.target_ref.lower():
                    try:
                        # Find context for this image
                        context_info = self._find_image_context(rel_id)

                        # Skip signature images if requested
                        if exclude_signatures:
                            context = context_info.get("context", "")
                            if self._is_signature_content(context):
                                continue

                        image_data = {
                            "data": base64.b64encode(rel.target_part.blob).decode(
                                "utf-8"
                            ),
                            "content_type": rel.target_part.content_type,
                            "filename": rel.target_ref.split("/")[-1],
                            "rel_id": rel_id,
                            "context_snippet": context_info.get("context", ""),
                            "context_type": context_info.get("type", "unknown"),
                            "section_hint": context_info.get("section"),
                            "day_hint": context_info.get("day"),
                            "row_label": context_info.get("row_label"),
                            "row_idx": context_info.get("row_idx"),
                            "cell_index": context_info.get("cell_index"),
                            "table_idx": context_info.get(
                                "table_idx"
                            ),  # CRITICAL: needed for slot filtering
                        }
                        images.append(image_data)
                    except Exception as e:
                        from backend.telemetry import logger

                        logger.warning(
                            "image_extraction_failed",
                            extra={"rel_id": rel_id, "error": str(e)},
                        )
        except Exception as e:
            from backend.telemetry import logger

            logger.warning("images_extraction_error", extra={"error": str(e)})

        return images

    def extract_hyperlinks(self) -> List[Dict[str, str]]:
        """Extract all hyperlinks from document with context for semantic anchoring.

        Schema v2.0: Now includes coordinate information for table-based placement.

        Returns:
            List of hyperlink dictionaries with:
            - v1.1 fields: text, url, context_snippet, context_type, section_hint, day_hint
            - v2.0 fields: schema_version, table_idx, row_idx, cell_idx, row_label, col_header
        """
        hyperlinks = []
        try:
            # Extract from paragraphs (non-table links)
            for paragraph in self.doc.paragraphs:
                for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                    try:
                        r_id = hyperlink.get(qn("r:id"))
                        if r_id and r_id in paragraph.part.rels:
                            url = paragraph.part.rels[r_id].target_ref
                            text = "".join(
                                node.text
                                for node in hyperlink.xpath(".//w:t")
                                if node.text
                            )
                            if text and url:
                                hyperlinks.append(
                                    {
                                        # Schema version
                                        "schema_version": "2.0",
                                        # v1.1 fields
                                        "text": text,
                                        "url": url,
                                        "context_snippet": self._get_context_snippet(
                                            paragraph, text
                                        ),
                                        "context_type": "paragraph",
                                        "section_hint": self._infer_section(
                                            paragraph.text
                                        ),
                                        "day_hint": None,
                                        # v2.0 fields (None for paragraph links)
                                        "table_idx": None,
                                        "row_idx": None,
                                        "cell_idx": None,
                                        "row_label": None,
                                        "col_header": None,
                                    }
                                )
                    except Exception as e:
                        from backend.telemetry import logger

                        logger.warning(
                            "hyperlink_extraction_failed", extra={"error": str(e)}
                        )

            # Extract from tables with coordinate capture (v2.0)
            for table_idx, table in enumerate(self.doc.tables):
                # Get column headers from first row
                col_headers = []
                if table.rows:
                    col_headers = [cell.text.strip() for cell in table.rows[0].cells]

                for row_idx, row in enumerate(table.rows):
                    # Get row label (first cell)
                    row_label = row.cells[0].text.strip() if row.cells else ""

                    for cell_idx, cell in enumerate(row.cells):
                        # Get column header for this cell
                        col_header = (
                            col_headers[cell_idx] if cell_idx < len(col_headers) else ""
                        )

                        # Extract day hint from column header
                        day_hint = self._extract_day_from_header(col_header)

                        for paragraph in cell.paragraphs:
                            for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                                try:
                                    r_id = hyperlink.get(qn("r:id"))
                                    if r_id and r_id in paragraph.part.rels:
                                        url = paragraph.part.rels[r_id].target_ref
                                        text = "".join(
                                            node.text
                                            for node in hyperlink.xpath(".//w:t")
                                            if node.text
                                        )
                                        if text and url:
                                            hyperlinks.append(
                                                {
                                                    # Schema version
                                                    "schema_version": "2.0",
                                                    # v1.1 fields
                                                    "text": text,
                                                    "url": url,
                                                    "context_snippet": self._get_context_snippet(
                                                        paragraph, text
                                                    ),
                                                    "context_type": "table_cell",
                                                    "section_hint": self._infer_section(
                                                        cell.text
                                                    ),
                                                    "day_hint": day_hint,
                                                    # v2.0 fields (coordinates)
                                                    "table_idx": table_idx,
                                                    "row_idx": row_idx,
                                                    "cell_idx": cell_idx,
                                                    "row_label": row_label,
                                                    "col_header": col_header,
                                                }
                                            )
                                except Exception as e:
                                    from backend.telemetry import logger

                                    logger.warning(
                                        "table_hyperlink_extraction_failed",
                                        extra={"error": str(e)},
                                    )
        except Exception as e:
            from backend.telemetry import logger

            logger.warning("hyperlinks_extraction_error", extra={"error": str(e)})

        # Validate and clean extracted hyperlinks
        validated_hyperlinks = []
        for link in hyperlinks:
            # Ensure row_label is clean
            if link.get("row_label"):
                link["row_label"] = link["row_label"].strip()

            # Ensure day_hint is valid
            if link.get("day_hint"):
                original_day_hint = link["day_hint"]
                link["day_hint"] = link["day_hint"].lower().strip()

                # Normalize compound formats before strict validation
                normalized_day = self._normalize_day_hint(link["day_hint"])

                if normalized_day:
                    # Use normalized day if found
                    link["day_hint"] = normalized_day
                else:
                    # If normalization didn't find a valid day, try strict validation
                    valid_days = [
                        "monday",
                        "tuesday",
                        "wednesday",
                        "thursday",
                        "friday",
                    ]
                    if link["day_hint"] not in valid_days:
                        from backend.telemetry import logger

                        logger.warning(
                            "invalid_day_hint",
                            extra={
                                "day_hint": original_day_hint,
                                "normalized_attempt": link["day_hint"],
                                "link_text": link.get("text", "")[:50],
                            },
                        )
                        link["day_hint"] = None

            # Ensure section_hint is valid
            valid_sections = [
                "unit_lesson",
                "objective",
                "anticipatory_set",
                "instruction",
                "misconceptions",
                "assessment",
                "homework",
            ]
            if link.get("section_hint") and link["section_hint"] not in valid_sections:
                from backend.telemetry import logger

                logger.debug(
                    "invalid_section_hint",
                    extra={
                        "section_hint": link["section_hint"],
                        "link_text": link.get("text", "")[:50],
                    },
                )
                # Don't null it out - might still be useful for fuzzy matching

            validated_hyperlinks.append(link)

        return validated_hyperlinks

    def extract_hyperlinks_for_slot(
        self, slot_number: int = None, subject: str = None, teacher_name: str = None
    ) -> List[Dict[str, str]]:
        """
        Extract hyperlinks for specific slot only.

        This prevents cross-slot contamination by only extracting from the slot's tables,
        excluding paragraph-level hyperlinks.

        Supports both slot-number-based and subject-based extraction to handle
        misaligned slot numbers between source files and weekly plan configs.

        Args:
            slot_number: Slot number (1-indexed) - optional if subject provided
            subject: Subject name to find correct slot - optional if slot_number provided
            teacher_name: Teacher name for disambiguating duplicate subjects

        Returns:
            List of hyperlink dictionaries with slot-specific metadata

        Raises:
            ValueError: If slot structure is invalid or neither slot_number nor subject provided
        """
        from backend.telemetry import logger

        # If subject provided, find actual slot number
        if subject:
            slot_number = self.find_slot_by_subject(subject, teacher_name)
            logger.debug(
                "hyperlinks_extraction_via_subject",
                extra={
                    "subject": subject,
                    "teacher": teacher_name,
                    "resolved_slot": slot_number,
                },
            )
        elif slot_number is None:
            raise ValueError(
                "Must provide either slot_number or subject for slot-aware extraction"
            )

        # Validate and get table indices
        table_start, table_end = validate_slot_structure(self.doc, slot_number)

        logger.info(
            "extracting_slot_hyperlinks",
            extra={
                "slot_number": slot_number,
                "table_start": table_start,
                "table_end": table_end,
            },
        )

        hyperlinks = []
        seen_links = set()  # To de-duplicate hyperlinks in this slot

        # Extract ONLY from slot's tables (no paragraphs!)
        for table_idx in range(table_start, table_end + 1):
            table = self.doc.tables[table_idx]

            # Get column headers
            col_headers = []
            if table.rows:
                col_headers = [cell.text.strip() for cell in table.rows[0].cells]

            for row_idx, row in enumerate(table.rows):
                row_label = row.cells[0].text.strip() if row.cells else ""

                for cell_idx, cell in enumerate(row.cells):
                    col_header = (
                        col_headers[cell_idx] if cell_idx < len(col_headers) else ""
                    )
                    day_hint = self._extract_day_from_header(col_header)

                    # Normalize day hint
                    if day_hint:
                        day_hint = self._normalize_day_hint(day_hint)
                    elif cell_idx > 0:
                        # Check if this is a metadata column (not expected to have day hints)
                        # Common metadata headers: Grade, Homeroom, Subject, Week of
                        col_header_lower = col_header.lower() if col_header else ""
                        is_metadata_column = any(
                            pattern in col_header_lower
                            for pattern in [
                                "grade:",
                                "homeroom:",
                                "subject:",
                                "week of:",
                                "week:",
                            ]
                        )
                        
                        # Only log warning if this isn't a metadata column
                        # Metadata columns are expected to not have day hints
                        if not is_metadata_column:
                            # Log warning if day hint is missing for non-label, non-metadata column
                            # This helps identify why some links become "orphans"
                            logger.warning(
                                "hyperlink_extraction_missing_day_hint",
                                extra={
                                    "slot": slot_number,
                                    "table_idx": table_idx,
                                    "row": row_idx,
                                    "col": cell_idx,
                                    "header": col_header[:50] if col_header else "empty"
                                }
                            )

                    for paragraph in cell.paragraphs:
                        for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                            try:
                                r_id = hyperlink.get(qn("r:id"))
                                if r_id and r_id in paragraph.part.rels:
                                    url = paragraph.part.rels[r_id].target_ref
                                    text = "".join(
                                        node.text
                                        for node in hyperlink.xpath(".//w:t")
                                        if node.text
                                    )

                                    if text and url:
                                        # Normalize text/url for de-duplication
                                        text_clean = text.strip()
                                        url_clean = url.strip()
                                        
                                        # De-duplicate by text, url, and location
                                        # This prevents identical links in the same cell from being extracted multiple times
                                        link_key = (text_clean, url_clean, table_idx, row_idx, cell_idx)
                                        if link_key in seen_links:
                                            continue
                                        seen_links.add(link_key)

                                        hyperlinks.append(
                                            {
                                                "schema_version": "2.0",
                                                "text": text_clean,
                                                "url": url_clean,
                                                "context_snippet": self._get_context_snippet(
                                                    paragraph, text
                                                ),
                                                "context_type": "table",
                                                "section_hint": self._infer_section(
                                                    row_label
                                                ),
                                                "day_hint": day_hint,
                                                "table_idx": table_idx,
                                                "row_idx": row_idx,
                                                "cell_idx": cell_idx,
                                                "row_label": row_label,
                                                "col_header": col_header,
                                            }
                                        )
                            except Exception as e:
                                logger.warning(
                                    "hyperlink_extraction_failed",
                                    extra={"error": str(e), "table_idx": table_idx},
                                )

        # Final de-duplication: if multiple identical links exist in the same slot but 
        # in different cells, we keep them for placement attempts. 
        # But if they end up in Referenced Links, they will be de-duplicated there later or here.
        # For now, we keep them if coordinates differ.

        logger.info(
            "slot_hyperlinks_extracted",
            extra={"slot_number": slot_number, "hyperlink_count": len(hyperlinks)},
        )

        return hyperlinks

    def extract_images_for_slot(
        self, slot_number: int = None, subject: str = None, teacher_name: str = None
    ) -> List[Dict]:
        """
        Extract images for specific slot only.

        Supports both slot-number-based and subject-based extraction to handle
        misaligned slot numbers between source files and weekly plan configs.

        Args:
            slot_number: Slot number (1-indexed) - optional if subject provided
            subject: Subject name to find correct slot - optional if slot_number provided
            teacher_name: Teacher name for disambiguating duplicate subjects

        Returns:
            List of image dictionaries from slot's tables only

        Raises:
            ValueError: If slot structure is invalid or neither slot_number nor subject provided
        """
        from backend.telemetry import logger

        # If subject provided, find actual slot number
        if subject:
            slot_number = self.find_slot_by_subject(subject, teacher_name)
            logger.debug(
                "images_extraction_via_subject",
                extra={
                    "subject": subject,
                    "teacher": teacher_name,
                    "resolved_slot": slot_number,
                },
            )
        elif slot_number is None:
            raise ValueError(
                "Must provide either slot_number or subject for slot-aware extraction"
            )

        # Validate and get table indices
        table_start, table_end = validate_slot_structure(self.doc, slot_number)

        # Extract all images (with table_idx now populated)
        all_images = self.extract_images()

        # Filter to slot's tables only and de-duplicate
        slot_images = []
        seen_images = set()  # To de-duplicate images in this slot

        for img in all_images:
            if img.get("table_idx") is not None and table_start <= img["table_idx"] <= table_end:
                # De-duplicate by data (hashed) and location
                # This prevents identical images in the same cell from being extracted multiple times
                import hashlib
                data_hash = hashlib.md5(img["data"].encode()).hexdigest()
                # Use .get() so images from older code paths or paragraph context don't raise KeyError
                img_key = (data_hash, img["table_idx"], img.get("row_idx", -1), img.get("cell_index", -1))
                
                if img_key in seen_images:
                    continue
                seen_images.add(img_key)
                slot_images.append(img)

        logger.info(
            "slot_images_extracted",
            extra={
                "slot_number": slot_number,
                "total_images": len(all_images),
                "slot_images": len(slot_images),
                "table_start": table_start,
                "table_end": table_end,
            },
        )

        return slot_images

    def find_slot_by_subject(self, subject: str, teacher_name: str = None, homeroom: str = None, grade: str = None) -> int:
        """
        Find which slot contains the given subject by scanning metadata tables.

        This handles cases where slot numbers don't align between source file
        and weekly plan configuration (e.g., Savoca has ELA/SS as Slot 1, but
        Wilson config has it as Slot 2).

        Args:
            subject: Subject name to find (e.g., "ELA/SS", "Math", "Science")
            teacher_name: Optional teacher name to disambiguate duplicate subjects
            homeroom: Optional homeroom to disambiguate duplicate subjects (e.g., "T5")
            grade: Optional grade to disambiguate duplicate subjects (e.g., "3")

        Returns:
            Slot number (1-indexed) containing the subject

        Raises:
            ValueError: If subject not found in any slot
        """
        import string

        from backend.telemetry import logger

        def normalize_text(text):
            """Normalize text: lowercase, strip punctuation and whitespace."""
            # Remove punctuation
            text = text.translate(str.maketrans("", "", string.punctuation))
            # Lowercase and strip
            return text.lower().strip()

        def tokenize_subject(text):
            """Split combined subjects like 'ela/ss' into tokens."""
            # Split on common separators
            tokens = []
            for sep in ["/", "-", "&", "and"]:
                if sep in text:
                    tokens.extend(text.split(sep))
                    break
            else:
                tokens = [text]
            return [normalize_text(t) for t in tokens if t.strip()]

        # Normalize requested subject
        subject_normalized = normalize_text(subject)
        subject_tokens = tokenize_subject(subject)

        # Subject aliases - NON-OVERLAPPING to prevent cross-matching
        #
        # IMPORTANT: These aliases are intentionally non-overlapping to prevent
        # incorrect matches. For example:
        # - 'ela' does NOT include 'ela/ss' (would cause SS to match ELA)
        # - 'social studies' does NOT include 'ela/ss' (would cause ELA to match SS)
        # - Combined subjects like 'ela/ss' are handled separately with token matching
        #
        # Adding new aliases:
        # 1. Ensure they don't overlap with other categories
        # 2. Use normalize_text() format (lowercase, no punctuation)
        # 3. Add to appropriate category or create new one
        # 4. Update tests in test_subject_slot_detection.py
        #
        subject_mappings = {
            "ela": ["ela", "english", "language arts", "reading", "literacy"],
            "math": ["math", "mathematics"],
            "science": ["science", "sci"],
            "social studies": ["ss", "social studies", "history"],
            # Combined subjects handled separately with token-based matching
            "ela/ss": [
                "ela/ss",
                "elass",
                "language arts/social studies",
                "language artssocial studies",
            ],
        }

        # Get possible aliases for this subject
        possible_aliases = [subject_normalized]
        for canonical, aliases in subject_mappings.items():
            if subject_normalized in aliases or subject_normalized == normalize_text(
                canonical
            ):
                possible_aliases.extend([normalize_text(a) for a in aliases])
                break

        # Normalize teacher name if provided
        teacher_normalized = normalize_text(teacher_name) if teacher_name else None
        
        # Normalize homeroom and grade if provided
        homeroom_normalized = normalize_text(homeroom) if homeroom else None
        grade_normalized = normalize_text(grade) if grade else None

        # Scan each slot's metadata table
        total_tables = len(self.doc.tables)
        
        # Check for signature table (last table) - optional
        has_signature = False
        if total_tables > 0:
            last_table = self.doc.tables[-1]
            if last_table.rows and last_table.rows[0].cells:
                first_cell = last_table.rows[0].cells[0].text.strip().lower()
                if "signature" in first_cell or "required signatures" in first_cell:
                    has_signature = True
        
        # Calculate available slots (excluding signature if present)
        if has_signature:
            available_slots = (total_tables - 1) // 2  # Exclude signature table
        else:
            available_slots = total_tables // 2  # No signature table

        matches = []  # Store all matches for disambiguation

        for slot_num in range(1, available_slots + 1):
            try:
                # Get metadata table for this slot
                table_start, table_end = validate_slot_structure(self.doc, slot_num)
                meta_table = self.doc.tables[table_start]

                # Extract metadata from table
                slot_subject = None
                slot_teacher = None
                slot_homeroom = None
                slot_grade = None

                for row in meta_table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_lower = cell_text.lower()

                        # Extract subject
                        if "subject:" in cell_lower:
                            slot_subject = cell_text.split(":", 1)[-1].strip()

                        # Extract teacher name
                        if "name:" in cell_lower or "teacher:" in cell_lower:
                            slot_teacher = cell_text.split(":", 1)[-1].strip()

                        # Extract homeroom
                        if "homeroom:" in cell_lower:
                            slot_homeroom = cell_text.split(":", 1)[-1].strip()
                        
                        # Extract grade
                        if "grade:" in cell_lower:
                            slot_grade = cell_text.split(":", 1)[-1].strip()

                if not slot_subject:
                    continue

                # Normalize metadata subject
                subject_value_normalized = normalize_text(slot_subject)
                subject_value_tokens = tokenize_subject(slot_subject)

                # Match logic: Check for exact match or token overlap
                matched = False
                matched_alias = None

                # 1. Try exact match with aliases
                for alias in possible_aliases:
                    if subject_value_normalized == alias:
                        matched = True
                        matched_alias = alias
                        logger.debug(
                            "subject_exact_match",
                            extra={
                                "slot": slot_num,
                                "requested": subject,
                                "metadata": slot_subject,
                                "matched_alias": alias,
                            },
                        )
                        break

                # 2. For combined subjects (e.g., ELA/SS), check token overlap
                if not matched and len(subject_tokens) > 1:
                    # Both tokens must match
                    if all(
                        any(st == vt for vt in subject_value_tokens)
                        for st in subject_tokens
                    ):
                        matched = True
                        matched_alias = "token_match"
                        logger.debug(
                            "subject_token_match",
                            extra={
                                "slot": slot_num,
                                "requested_tokens": subject_tokens,
                                "metadata_tokens": subject_value_tokens,
                            },
                        )

                if matched:
                    matches.append(
                        {
                            "slot_num": slot_num,
                            "subject": slot_subject,
                            "teacher": slot_teacher,
                            "homeroom": slot_homeroom,
                            "grade": slot_grade,
                            "matched_alias": matched_alias,
                        }
                    )

            except ValueError:
                # Invalid slot structure, skip
                continue

        # No matches found - try flexible matching for combined subjects
        if not matches and len(subject_tokens) > 1:
            # For combined subjects like "ELA/SS", try matching individual tokens
            # This handles cases where config says "ELA/SS" but file only has "ELA"
            logger.info(
                "trying_flexible_subject_match",
                extra={
                    "requested_subject": subject,
                    "tokens": subject_tokens,
                    "message": "No exact match found, trying individual tokens",
                },
            )

            for slot_num in range(1, available_slots + 1):
                try:
                    table_start, table_end = validate_slot_structure(self.doc, slot_num)
                    meta_table = self.doc.tables[table_start]

                    slot_subject = None
                    slot_teacher = None
                    slot_homeroom = None
                    slot_grade = None
                    for row in meta_table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            cell_lower = cell_text.lower()

                            if "subject:" in cell_lower:
                                slot_subject = cell_text.split(":", 1)[-1].strip()
                            if "name:" in cell_lower or "teacher:" in cell_lower:
                                slot_teacher = cell_text.split(":", 1)[-1].strip()
                            if "homeroom:" in cell_lower:
                                slot_homeroom = cell_text.split(":", 1)[-1].strip()
                            if "grade:" in cell_lower:
                                slot_grade = cell_text.split(":", 1)[-1].strip()

                    if slot_subject:
                        subject_value_normalized = normalize_text(slot_subject)
                        # Try matching any individual token
                        for token in subject_tokens:
                            if token == subject_value_normalized:
                                logger.info(
                                    "flexible_subject_match_found",
                                    extra={
                                        "requested_subject": subject,
                                        "matched_token": token,
                                        "metadata_subject": slot_subject,
                                        "slot": slot_num,
                                    },
                                )
                                matches.append(
                                    {
                                        "slot_num": slot_num,
                                        "subject": slot_subject,
                                        "teacher": slot_teacher,
                                        "homeroom": slot_homeroom,
                                        "grade": slot_grade,
                                        "matched_alias": f"flexible_token_{token}",
                                    }
                                )
                                break
                except ValueError:
                    continue

        # Still no matches found
        if not matches:
            raise ValueError(
                f"Subject '{subject}' not found in any slot. "
                f"Available slots: {available_slots}. "
                f"Scanned metadata tables but found no matching subject."
            )

        # Single match - return it
        if len(matches) == 1:
            match = matches[0]
            logger.info(
                "subject_slot_found",
                extra={
                    "requested_subject": subject,
                    "found_in_slot": match["slot_num"],
                    "metadata_subject": match["subject"],
                    "metadata_teacher": match["teacher"],
                    "matched_via": match["matched_alias"],
                },
            )
            return match["slot_num"]

        # Multiple matches - disambiguate by teacher, then homeroom, then grade
        if teacher_normalized:
            for match in matches:
                if match["teacher"] and teacher_normalized in normalize_text(
                    match["teacher"]
                ):
                    logger.info(
                        "subject_slot_found_via_teacher",
                        extra={
                            "requested_subject": subject,
                            "requested_teacher": teacher_name,
                            "found_in_slot": match["slot_num"],
                            "metadata_subject": match["subject"],
                            "metadata_teacher": match["teacher"],
                            "total_matches": len(matches),
                        },
                    )
                    return match["slot_num"]
        
        # If teacher disambiguation failed or not provided, try homeroom
        if homeroom_normalized:
            homeroom_matches = []
            for match in matches:
                if match["homeroom"] and homeroom_normalized == normalize_text(
                    match["homeroom"]
                ):
                    homeroom_matches.append(match)
            
            if len(homeroom_matches) == 1:
                match = homeroom_matches[0]
                logger.info(
                    "subject_slot_found_via_homeroom",
                    extra={
                        "requested_subject": subject,
                        "requested_homeroom": homeroom,
                        "found_in_slot": match["slot_num"],
                        "metadata_subject": match["subject"],
                        "metadata_homeroom": match["homeroom"],
                        "total_matches": len(matches),
                    },
                )
                return match["slot_num"]
            elif len(homeroom_matches) > 1:
                # Multiple matches with same homeroom - try grade if available
                if grade_normalized:
                    grade_matches = []
                    for match in homeroom_matches:
                        if match["grade"] and grade_normalized == normalize_text(
                            match["grade"]
                        ):
                            grade_matches.append(match)
                    
                    if len(grade_matches) == 1:
                        match = grade_matches[0]
                        logger.info(
                            "subject_slot_found_via_homeroom_grade",
                            extra={
                                "requested_subject": subject,
                                "requested_homeroom": homeroom,
                                "requested_grade": grade,
                                "found_in_slot": match["slot_num"],
                                "metadata_subject": match["subject"],
                                "metadata_homeroom": match["homeroom"],
                                "metadata_grade": match["grade"],
                                "total_matches": len(matches),
                            },
                        )
                        return match["slot_num"]
                    elif len(grade_matches) > 1:
                        # Still multiple matches - return first one and warn
                        match = grade_matches[0]
                        logger.warning(
                            "multiple_matches_after_homeroom_grade",
                            extra={
                                "requested_subject": subject,
                                "requested_homeroom": homeroom,
                                "requested_grade": grade,
                                "total_matches": len(grade_matches),
                                "selected_slot": match["slot_num"],
                            },
                        )
                        return match["slot_num"]
        
        # Multiple matches, can't disambiguate - return first and warn
        match = matches[0]
        logger.warning(
            "multiple_subject_matches",
            extra={
                "requested_subject": subject,
                "requested_teacher": teacher_name,
                "requested_homeroom": homeroom,
                "requested_grade": grade,
                "total_matches": len(matches),
                "matches": [
                    {
                        "slot": m["slot_num"], 
                        "teacher": m["teacher"],
                        "homeroom": m["homeroom"],
                        "grade": m.get("grade")
                    } for m in matches
                ],
                "selected_slot": match["slot_num"],
                "message": "Multiple slots match subject, returning first match",
            },
        )
        return match["slot_num"]

    def extract_subject_content_for_slot(
        self,
        slot_number: int,
        subject: str,
        teacher_name: str = None,
        strip_urls: bool = True,
    ) -> Dict[str, Any]:
        """
        Extract content for a specific slot only (slot-aware).

        This prevents cross-slot contamination by only extracting from the slot's tables.
        Now supports subject-based slot detection to handle misaligned slot numbers.

        Args:
            slot_number: Slot number from weekly plan config (used as fallback)
            subject: Subject name - will be used to find correct slot in source file
            teacher_name: Optional teacher name for disambiguation when multiple slots match subject
            strip_urls: If True, remove URLs from text

        Returns:
            Dictionary with slot content
        """
        import string

        from backend.telemetry import logger

        def normalize_text(text):
            """Normalize text: lowercase, strip punctuation and whitespace."""
            if not text:
                return ""
            # Remove punctuation
            text = text.translate(str.maketrans("", "", string.punctuation))
            # Lowercase and strip
            return text.lower().strip()

        # CRITICAL: Find the actual slot number by subject, not by position
        # This handles cases where Savoca has ELA/SS as Slot 1 but Wilson config has it as Slot 2
        try:
            actual_slot_number = self.find_slot_by_subject(
                subject, teacher_name=teacher_name
            )

            # If find_slot_by_subject returned a different slot, check if we should trust it or use requested slot_number
            # Priority: If teacher_name successfully disambiguated, trust it. Otherwise, prefer requested slot_number if it's valid.
            if actual_slot_number != slot_number:
                # Check if the requested slot_number also contains this subject and teacher
                try:
                    table_start, table_end = validate_slot_structure(
                        self.doc, slot_number
                    )
                    meta_table = self.doc.tables[table_start]

                    # Extract subject, teacher, and homeroom from metadata table
                    slot_subject = None
                    slot_teacher = None
                    slot_homeroom = None
                    for row in meta_table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            cell_lower = cell_text.lower()
                            if "subject:" in cell_lower:
                                slot_subject = cell_text.split(":", 1)[-1].strip()
                            if "name:" in cell_lower or "teacher:" in cell_lower:
                                slot_teacher = cell_text.split(":", 1)[-1].strip()
                            if "homeroom:" in cell_lower:
                                slot_homeroom = cell_text.split(":", 1)[-1].strip()

                    # Normalize for comparison
                    subject_normalized = normalize_text(subject)
                    slot_subject_normalized = (
                        normalize_text(slot_subject) if slot_subject else ""
                    )

                    # Check if requested slot has the subject
                    subject_matches = (
                        subject_normalized in slot_subject_normalized
                        or slot_subject_normalized in subject_normalized
                    )

                    # Check if teacher_name matches requested slot's teacher (if teacher_name provided)
                    teacher_matches = False
                    if teacher_name and slot_teacher:
                        teacher_normalized = normalize_text(teacher_name)
                        slot_teacher_normalized = normalize_text(slot_teacher)
                        teacher_matches = (
                            teacher_normalized in slot_teacher_normalized
                            or slot_teacher_normalized in teacher_normalized
                        )

                    # Use requested slot_number if:
                    # 1. It has the subject AND
                    # 2. Either no teacher_name provided OR teacher_name matches requested slot's teacher
                    # This means: if teacher_name was provided and successfully disambiguated to a different slot, trust that.
                    if subject_matches and (not teacher_name or teacher_matches):
                        # Requested slot_number is valid - use it
                        logger.info(
                            "using_requested_slot_number",
                            extra={
                                "requested_slot": slot_number,
                                "found_slot": actual_slot_number,
                                "subject": subject,
                                "teacher_name": teacher_name,
                                "reason": "Requested slot contains subject and matches teacher (if provided)",
                            },
                        )
                        slot_number = slot_number  # Keep requested slot_number
                    else:
                        # Teacher_name disambiguated to a different slot, or requested slot doesn't match - trust found slot
                        logger.warning(
                            "slot_number_mismatch",
                            extra={
                                "requested_slot": slot_number,
                                "actual_slot": actual_slot_number,
                                "subject": subject,
                                "teacher_name": teacher_name,
                                "message": f"Slot {slot_number} requested but subject '{subject}' found in slot {actual_slot_number}. Teacher disambiguation overrides.",
                            },
                        )
                        slot_number = actual_slot_number
                except ValueError:
                    # Requested slot_number is invalid - use found slot
                    logger.warning(
                        "slot_number_mismatch",
                        extra={
                            "requested_slot": slot_number,
                            "actual_slot": actual_slot_number,
                            "subject": subject,
                            "teacher_name": teacher_name,
                            "message": f"Slot {slot_number} requested but subject '{subject}' found in slot {actual_slot_number}",
                        },
                    )
                    slot_number = actual_slot_number
            # If actual_slot_number == slot_number, no change needed
            else:
                print(f"DEBUG: DOCXParser - actual_slot_number matches requested slot_number {slot_number}")
                pass
        except ValueError as e:
            # Subject not found, fall back to original slot number
            logger.warning(
                "subject_slot_detection_failed",
                extra={
                    "slot_number": slot_number,
                    "subject": subject,
                    "teacher_name": teacher_name,
                    "error": str(e),
                    "fallback": "using requested slot number",
                },
            )

        # Validate slot structure and get table indices
        try:
            table_start, table_end = validate_slot_structure(self.doc, slot_number)
        except ValueError as e:
            total_tables = len(self.doc.tables)
            if total_tables >= 2:
                table_start = (slot_number - 1) * 2
                table_end = table_start + 1
            else:
                raise

        # Extract metadata from table_start
        meta_table = self.doc.tables[table_start]
        meta_data = {}
        for row in meta_table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                lower = text.lower()
                if "subject:" in lower:
                    meta_data["subject"] = text.split(":", 1)[-1].strip()
                elif "name:" in lower or "teacher:" in lower:
                    meta_data["primary_teacher_name"] = text.split(":", 1)[-1].strip()
                elif "homeroom:" in lower:
                    meta_data["homeroom"] = text.split(":", 1)[-1].strip()
                elif "grade:" in lower:
                    meta_data["grade"] = text.split(":", 1)[-1].strip()
                elif "week of:" in lower:
                    meta_data["week_of"] = text.split(":", 1)[-1].strip()

        # The daily plans table is at table_end
        daily_table_idx = table_end

        logger.info(
            "extracting_slot_content",
            extra={
                "slot_number": slot_number,
                "subject": subject,
                "table_idx": daily_table_idx,
            },
        )

        # Check if this is a single-lesson format (table_end contains metadata, not daily plans)
        is_single_lesson = False
        if daily_table_idx < len(self.doc.tables):
            daily_table = self.doc.tables[daily_table_idx]
            if daily_table.rows:
                first_row_text = " ".join(
                    cell.text.strip() for cell in daily_table.rows[0].cells
                )
                first_row_clean = first_row_text.translate(
                    str.maketrans("", "", string.punctuation)
                ).upper()
                
                metadata_indicators = ["UNIT", "LESSON", "MODULE", "SUBJECT", "OBJECTIVE"]
                weekdays = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY"]
                has_metadata_indicator = any(
                    indicator in first_row_clean for indicator in metadata_indicators
                )
                has_weekday = any(day in first_row_clean for day in weekdays)
                
                if has_metadata_indicator and not has_weekday:
                    is_single_lesson = True

        if is_single_lesson:
            # Single-lesson format: extract content from both metadata and lesson tables
            full_text_parts = []
            
            # Extract metadata table content
            meta_table = self.doc.tables[table_start]
            for row in meta_table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_parts.append(row_text)
            
            # Extract lesson metadata table content (table_end)
            lesson_table = self.doc.tables[table_end]
            for row in lesson_table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_parts.append(row_text)
            
            # For single-lesson documents, create a generic daily structure
            table_content = {"All Days": {"Lesson Content": "\n".join(full_text_parts)}}
            no_school_days = []
            
            logger.info(
                "single_lesson_content_extracted",
                extra={
                    "slot_number": slot_number,
                    "subject": subject,
                    "content_length": len("\n".join(full_text_parts)),
                },
            )
        else:
            # Standard format: extract from daily plans table
            table_content = self.extract_table_lesson(daily_table_idx)

            # Convert to text format, filtering No School days
            full_text_parts = []
            no_school_days = []

            for day, day_content in table_content.items():
                # Check if this day is "No School" - but preserve all content regardless
                day_text = " ".join(day_content.values())
                is_no_school = self.is_day_no_school(day_text)
                
                if is_no_school:
                    no_school_days.append(day)
                    # Mark as "No School" but still include all content from DOCX
                    # This allows LLM to preserve tailored_instruction and other fields
                    full_text_parts.append(f"\n{day.upper()}")
                    full_text_parts.append("Unit/Lesson: No School")
                    # Include all other content from DOCX
                    for label, text in day_content.items():
                        # Skip unit/lesson field if it's "No School" to avoid duplication
                        label_lower = label.lower()
                        if "unit" not in label_lower and "lesson" not in label_lower:
                            full_text_parts.append(f"{label} {text}")
                        elif text.strip() and text.strip().lower() != "no school":
                            # If unit/lesson has actual content (not just "No School"), include it
                            full_text_parts.append(f"{label} {text}")
                else:
                    # Include full content for regular days
                    full_text_parts.append(f"\n{day.upper()}")
                    for label, text in day_content.items():
                        full_text_parts.append(f"{label} {text}")

            logger.info(
                "slot_content_extracted",
                extra={
                    "slot_number": slot_number,
                    "subject": subject,
                    "content_length": len("\n".join(full_text_parts)),
                    "no_school_days": len(no_school_days),
                },
            )

        return {
            "subject": subject,
            "full_text": "\n".join(full_text_parts),
            "table_content": table_content,
            "no_school_days": no_school_days,
            "found": True,
            "format": "table",
            "slot_number": slot_number,
            "table_idx": daily_table_idx,
            "metadata": meta_data if 'meta_data' in locals() else {}
        }

    def _get_context_snippet(
        self, paragraph, link_text: str, window: int = None
    ) -> str:
        """Extract text context around element for semantic matching.

        Args:
            paragraph: Paragraph object containing the element
            link_text: Text of the hyperlink/element
            window: Number of characters to capture (defaults to config setting)

        Returns:
            Context snippet string
        """
        if window is None:
            try:
                from backend.config import settings

                window = settings.MEDIA_CONTEXT_WINDOW_CHARS
            except:
                window = 100

        full_text = paragraph.text
        if link_text in full_text:
            pos = full_text.index(link_text)
            start = max(0, pos - window // 2)
            end = min(len(full_text), pos + len(link_text) + window // 2)
            return full_text[start:end].strip()
        return full_text[:window].strip()

    def _infer_section(self, text: str) -> Optional[str]:
        """Infer lesson plan section from text content.

        Args:
            text: Text content to analyze

        Returns:
            Section name or None
        """
        import re

        text_lower = text.lower()

        # Check unit/lesson first (more specific) - use word boundaries to avoid false matches
        if (
            re.search(r"\bunit\b", text_lower)
            or "unit/lesson" in text_lower
            or "lesson #" in text_lower
            or re.search(r"\bmodule\b", text_lower)
        ):
            return "unit_lesson"
        elif any(
            kw in text_lower
            for kw in ["objective", "goal", "swbat", "students will be able"]
        ):
            return "objective"
        elif any(
            kw in text_lower for kw in ["warm-up", "hook", "anticipatory", "do now"]
        ):
            return "anticipatory_set"
        elif any(
            kw in text_lower
            for kw in ["instruction", "activity", "procedure", "tailored"]
        ):
            return "instruction"
        elif any(kw in text_lower for kw in ["misconception", "common error"]):
            return "misconceptions"
        elif any(
            kw in text_lower
            for kw in ["assessment", "check", "evaluate", "exit ticket"]
        ):
            return "assessment"
        elif any(kw in text_lower for kw in ["homework", "assignment", "practice"]):
            return "homework"

        return None

    def _detect_day_from_table(self, table) -> Optional[str]:
        """Try to detect day from table header row.

        Args:
            table: Table object

        Returns:
            Day name (lowercase) or None
        """
        if not table.rows:
            return None

        header_text = " ".join(cell.text.lower() for cell in table.rows[0].cells)
        days = ["monday", "tuesday", "wednesday", "thursday", "friday"]

        for day in days:
            if day in header_text:
                return day

        return None

    def _extract_day_from_header(self, col_header: str) -> Optional[str]:
        """Extract day name from column header with enhanced format support.

        Handles various formats:
        - "MONDAY" → "monday"
        - "MONDAY 9/22" → "monday"
        - "Monday, Sept 22" → "monday"
        - "Mon" → "monday"
        - "M 10/21" → "monday"

        Args:
            col_header: Column header text

        Returns:
            Day name (lowercase) or None
        """
        if not col_header:
            return None

        header_lower = col_header.lower().strip()

        # Full day names (check first for accuracy)
        days_full = ["monday", "tuesday", "wednesday", "thursday", "friday"]
        for day in days_full:
            if day in header_lower:
                return day

        # Common abbreviations (3-letter)
        day_abbrev_3 = {
            "mon": "monday",
            "tue": "tuesday",
            "wed": "wednesday",
            "thu": "thursday",
            "fri": "friday",
        }

        for abbrev, full in day_abbrev_3.items():
            if abbrev in header_lower:
                return full

        # Single letter abbreviations (less reliable, check with boundaries)
        day_abbrev_1 = {
            "m": "monday",
            "t": "tuesday",
            "w": "wednesday",
            "r": "thursday",  # Common in academic schedules
            "f": "friday",
        }

        # Only match single letters if they're at word boundaries
        words = header_lower.split()
        if words:
            first_word = words[0]
            if len(first_word) == 1 and first_word in day_abbrev_1:
                return day_abbrev_1[first_word]

        return None

    def _normalize_day_hint(self, day_hint: str) -> Optional[str]:
        """Normalize day hint from compound formats to a single valid day name.

        Handles various compound formats:
        - "Monday/Wednesday" → "monday" (takes first valid day)
        - "Mon-Wed" → "monday"
        - "Monday, Wednesday" → "monday"
        - "Monday and Wednesday" → "monday"
        - "Monday/Wednesday/Friday" → "monday"
        - "monday" → "monday" (already valid)
        - "Flex Day" → None (no valid day found)
        - "Day 1" → None (no valid day found, unless lookup table added)

        Args:
            day_hint: Day hint string (may be compound format)

        Returns:
            Normalized day name (lowercase) or None if no valid day found
        """
        if not day_hint:
            return None

        day_hint_lower = day_hint.lower().strip()

        # Valid day names (full names)
        valid_days = ["monday", "tuesday", "wednesday", "thursday", "friday"]

        # Check if already a valid day (exact match)
        if day_hint_lower in valid_days:
            return day_hint_lower

        # Check for full day names in the hint (handles compound formats)
        for day in valid_days:
            if day in day_hint_lower:
                return day

        # Common abbreviations (3-letter)
        day_abbrev_3 = {
            "mon": "monday",
            "tue": "tuesday",
            "wed": "wednesday",
            "thu": "thursday",
            "fri": "friday",
        }

        # Split on common separators: "/", "-", ",", "and"
        import re

        # Split on separators and whitespace
        parts = re.split(r"[/\-,]|\band\b", day_hint_lower)

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # Check full day name
            for day in valid_days:
                if day in part:
                    return day

            # Check abbreviations
            for abbrev, full in day_abbrev_3.items():
                if abbrev in part:
                    return full

        # No valid day found
        return None

    def _find_image_context(self, rel_id: str) -> Dict[str, Any]:
        """Find context for an image by locating it in document structure.

        Uses structure-based detection (row label + cell index) for exact placement.
        Now includes table_idx for slot-aware filtering.

        Args:
            rel_id: Relationship ID of the image

        Returns:
            Dict with context, type, section, day, row_label, cell_index, and table_idx
        """
        # Search paragraphs
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                try:
                    # Use string matching instead of XPath
                    if rel_id in str(run._element.xml):
                        return {
                            "context": paragraph.text[:100],
                            "type": "paragraph",
                            "section": self._infer_section(paragraph.text),
                            "day": None,
                            "row_label": None,
                            "cell_index": None,
                            "table_idx": None,  # Not in a table
                        }
                except:
                    pass

        # Search tables with structure detection
        for table_idx, table in enumerate(self.doc.tables):
            day_hint = self._detect_day_from_table(table)

            for row_idx, row in enumerate(table.rows):
                # Get row label from first cell
                row_label = row.cells[0].text.strip() if row.cells else ""

                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                # Use string matching instead of XPath
                                if rel_id in str(run._element.xml):
                                    # Infer section from row label
                                    section = (
                                        self._infer_section(row_label)
                                        if row_label
                                        else self._infer_section(cell.text)
                                    )

                                    # Map cell index to day
                                    day_map = {
                                        1: "monday",
                                        2: "tuesday",
                                        3: "wednesday",
                                        4: "thursday",
                                        5: "friday",
                                    }
                                    day_from_cell = day_map.get(cell_idx, day_hint)

                                    return {
                                        "context": cell.text[:100],
                                        "type": "table_cell",
                                        "section": section,
                                        "day": day_from_cell,
                                        "row_label": row_label,
                                        "row_idx": row_idx,
                                        "cell_index": cell_idx,
                                        "table_idx": table_idx,
                                    }
                            except:
                                pass

        return {
            "context": "",
            "type": "unknown",
            "section": None,
            "day": None,
            "row_label": None,
            "cell_index": None,
            "table_idx": None,  # Not found in any table
        }

    def get_metadata(self) -> Dict[str, Any]:
        """Extract metadata from document.

        Returns:
            Dict with document metadata
        """
        core_props = self.doc.core_properties

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat()
            if core_props.modified
            else None,
            "paragraph_count": len(self.paragraphs),
            "table_count": len(self.tables),
            "available_subjects": self.list_available_subjects(),
            "week_info": self.extract_week_info(),
        }


def parse_docx(
    file_path: str, subject: Optional[str] = None, plan_id: Optional[str] = None
) -> Dict[str, Any]:
    """Convenience function to parse DOCX file.

    Args:
        file_path: Path to DOCX file
        subject: Optional subject to extract (if None, returns all content)
        plan_id: Optional plan ID for performance tracking

    Returns:
        Parsed content dict
    """
    tracker = get_tracker()

    # Track file location/validation
    if plan_id:
        with tracker.track_operation(plan_id, "parse_locate_file"):
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")

    # Track DOCX opening and initial parsing
    if plan_id:
        with tracker.track_operation(plan_id, "parse_open_docx"):
            parser = DOCXParser(file_path)
    else:
        parser = DOCXParser(file_path)

    # Track content extraction
    if subject:
        if plan_id:
            with tracker.track_operation(plan_id, "parse_extract_subject"):
                result = parser.extract_subject_content(subject)
        else:
            result = parser.extract_subject_content(subject)
        return result
    else:
        result = {}

        if plan_id:
            with tracker.track_operation(plan_id, "parse_extract_text"):
                result["full_text"] = parser.get_full_text()

            with tracker.track_operation(plan_id, "parse_extract_metadata"):
                result["metadata"] = parser.get_metadata()

            with tracker.track_operation(plan_id, "parse_list_subjects"):
                result["available_subjects"] = parser.list_available_subjects()

            # Tables already extracted in __init__, just reference them
            result["tables"] = parser.tables
        else:
            result = {
                "full_text": parser.get_full_text(),
                "metadata": parser.get_metadata(),
                "available_subjects": parser.list_available_subjects(),
                "tables": parser.tables,
            }

        return result
