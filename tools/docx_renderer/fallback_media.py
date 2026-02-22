"""
Append unmatched hyperlinks and images to the end of a DOCX (fallback sections).

Used when semantic anchoring does not place media inline. All functions take
the renderer instance as first argument.
"""

import base64
from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.shared import Inches, Pt

from . import logger
from . import hyperlink_placement as _hyperlink_module


def append_unmatched_media(
    renderer,
    doc: Document,
    hyperlinks: List[Dict],
    images: List[Dict],
) -> None:
    """Append unmatched media to end with context for traceability.

    Args:
        renderer: DOCXRenderer instance (for add_hyperlink)
        doc: Document object
        hyperlinks: List of unmatched hyperlinks
        images: List of unmatched images
    """
    if hyperlinks:
        doc.add_page_break()
        heading = doc.add_paragraph("Referenced Links")
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(12)

        for link in hyperlinks:
            para = doc.add_paragraph()
            try:
                para.style = "List Bullet"
            except KeyError:
                para.add_run("• ")

            _hyperlink_module.add_hyperlink(
                renderer, para, link["text"], link["url"]
            )

            if link.get("context_snippet"):
                context_para = doc.add_paragraph()
                context_run = context_para.add_run(
                    f'  Context: "{link["context_snippet"][:80]}..."'
                )
                context_run.font.size = Pt(9)
                context_run.italic = True

    if images:
        if not hyperlinks:
            doc.add_page_break()

        heading = doc.add_paragraph("Attached Images")
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(12)

        for i, image in enumerate(images, 1):
            try:
                image_data = base64.b64decode(image["data"])
                image_stream = BytesIO(image_data)

                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(image_stream, width=Inches(4.0))

                caption = doc.add_paragraph()
                caption_run = caption.add_run(
                    f"Image {i}: {image.get('filename', 'Untitled')}"
                )
                caption_run.italic = True
                caption_run.font.size = Pt(10)

                if image.get("context_snippet"):
                    context_para = doc.add_paragraph()
                    context_run = context_para.add_run(
                        f'Context: "{image["context_snippet"][:80]}..."'
                    )
                    context_run.font.size = Pt(9)
                    context_run.italic = True

            except Exception as e:
                logger.warning(
                    "fallback_image_failed",
                    extra={"filename": image.get("filename"), "error": str(e)},
                )


def insert_images(
    renderer,
    doc: Document,
    images: List[Dict],
) -> None:
    """Insert images at the end of the document (legacy v1.0 behavior).

    Args:
        renderer: DOCXRenderer instance (unused; kept for API consistency)
        doc: Document object
        images: List of image dictionaries with base64-encoded data
    """
    if not images:
        return

    try:
        doc.add_page_break()

        heading = doc.add_paragraph()
        heading_run = heading.add_run("Attached Images")
        heading_run.bold = True
        heading_run.font.size = Pt(14)

        for i, image in enumerate(images, 1):
            try:
                image_data = base64.b64decode(image["data"])
                image_stream = BytesIO(image_data)

                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(4.0))

                caption = doc.add_paragraph()
                caption_run = caption.add_run(
                    f"Image {i}: {image.get('filename', 'Untitled')}"
                )
                caption_run.italic = True
                caption_run.font.size = Pt(10)

            except Exception as e:
                logger.warning(
                    "image_insertion_failed",
                    extra={
                        "image_index": i,
                        "filename": image.get("filename", "unknown"),
                        "error": str(e),
                    },
                )
    except Exception as e:
        logger.warning("images_insertion_error", extra={"error": str(e)})
