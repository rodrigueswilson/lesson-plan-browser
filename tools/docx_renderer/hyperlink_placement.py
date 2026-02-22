"""
Hyperlink placement and restoration for DOCX rendering.

All functions take the renderer instance as first argument and use
style.force_font_tnr8 for link formatting.
"""

import re
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from . import logger
from tools.table_structure import StructureMetadata

from .style import force_font_tnr8, FUZZY_MATCH_THRESHOLD


def inject_hyperlink_inline(renderer, cell, hyperlink: Dict, row_idx: int = None) -> None:
    """Inject hyperlink into cell on its own line."""
    link_text = hyperlink.get("text", "")
    link_url = hyperlink.get("url", "")
    cell_text = cell.text or ""

    markdown_pattern = rf"\[{re.escape(link_text)}\]\({re.escape(link_url)}\)"
    if re.search(markdown_pattern, cell_text, re.IGNORECASE):
        logger.debug(
            f"Skipping duplicate hyperlink injection: '{link_text}' already in markdown format in cell"
        )
        return

    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    is_bold = row_idx == renderer.UNIT_LESSON_ROW
    add_hyperlink(renderer, para, hyperlink["text"], hyperlink["url"], bold=is_bold)


def place_hyperlink_hybrid(
    renderer, link: Dict, table, structure: StructureMetadata
) -> str:
    """Place hyperlink using hybrid strategy. Returns strategy used."""
    diagnostic = {
        "link_text": link.get("text", "")[:50],
        "url": link.get("url", "")[:50],
        "input_coords": f"table={link.get('table_idx')}, row={link.get('row_idx')}, cell={link.get('cell_idx')}",
        "row_label": link.get("row_label", ""),
        "col_header": link.get("col_header", ""),
        "day_hint": link.get("day_hint", ""),
        "section_hint": link.get("section_hint", ""),
        "structure_type": structure.structure_type,
        "strategy_attempted": [],
    }

    if link.get("table_idx") is None:
        diagnostic["strategy_attempted"].append("skipped_non_table")
        diagnostic["result"] = "fallback"
        logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
        return "fallback"

    if structure.structure_type == "standard_8x6":
        diagnostic["strategy_attempted"].append("coordinate")
        if try_coordinate_placement(renderer, link, table, structure):
            diagnostic["result"] = "success_coordinate"
            logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
            return "coordinate"
        diagnostic["coordinate_failure"] = "mismatch_or_bounds"

    diagnostic["strategy_attempted"].append("label_day")
    diagnostic["label_lookup"] = structure.get_row_index(link.get("row_label", ""))
    diagnostic["day_lookup"] = structure.get_col_index(link.get("day_hint", ""))

    if try_label_day_placement(renderer, link, table, structure):
        diagnostic["result"] = "success_label_day"
        logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
        return "label_day"
    diagnostic["label_day_failure"] = (
        f"row={diagnostic['label_lookup']}, col={diagnostic['day_lookup']}"
    )

    diagnostic["strategy_attempted"].append("fuzzy")
    if try_fuzzy_placement(renderer, link, table, threshold=FUZZY_MATCH_THRESHOLD):
        diagnostic["result"] = "success_fuzzy"
        logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
        return "fuzzy"
    diagnostic["fuzzy_failure"] = "no_match_above_threshold"

    diagnostic["result"] = "fallback"
    logger.warning("hyperlink_placement_fallback", extra=diagnostic)
    return "fallback"


def try_coordinate_placement(
    renderer, link: Dict, table, structure: StructureMetadata
) -> bool:
    """Try to place link at exact coordinates."""
    row_idx = link.get("row_idx")
    cell_idx = link.get("cell_idx")

    if row_idx is None or cell_idx is None:
        return False

    target_row = row_idx + structure.row_offset

    try:
        if target_row >= len(table.rows):
            logger.warning(
                f"Row {target_row} out of bounds (table has {len(table.rows)} rows)"
            )
            return False

        row = table.rows[target_row]
        if cell_idx >= len(row.cells):
            logger.warning(
                f"Cell {cell_idx} out of bounds (row has {len(row.cells)} cells)"
            )
            return False

        cell = row.cells[cell_idx]
        inject_hyperlink_inline(renderer, cell, link)
        logger.debug(f"Placed '{link['text']}' at coordinates ({target_row}, {cell_idx})")
        return True

    except (IndexError, AttributeError) as e:
        logger.warning(f"Coordinate placement failed for '{link['text']}': {e}")
        return False


def try_label_day_placement(
    renderer, link: Dict, table, structure: StructureMetadata
) -> bool:
    """Try to place link using row label + day matching."""
    row_label = link.get("row_label", "").strip().lower().rstrip(":")
    day_hint = link.get("day_hint", "").strip().lower()

    if not row_label or not day_hint:
        return False

    target_row = structure.get_row_index(row_label)
    target_col = structure.get_col_index(day_hint)

    if target_row is None or target_col is None:
        logger.debug(
            f"Could not find row/col for '{link['text']}': row={row_label}, day={day_hint}"
        )
        return False

    try:
        if target_row >= len(table.rows) or target_col >= len(
            table.rows[target_row].cells
        ):
            logger.warning(
                f"Label/day placement out of bounds: ({target_row}, {target_col})"
            )
            return False

        cell = table.rows[target_row].cells[target_col]
        inject_hyperlink_inline(renderer, cell, link)
        logger.debug(
            f"Placed '{link['text']}' via label/day at ({target_row}, {target_col})"
        )
        return True

    except (IndexError, AttributeError) as e:
        logger.warning(f"Label/day placement failed for '{link['text']}': {e}")
        return False


def try_fuzzy_placement(renderer, link: Dict, table, threshold: float = 0.65) -> bool:
    """Try to place link using fuzzy text matching."""
    best_match = {
        "confidence": 0.0,
        "location": None,
        "cell_text_preview": "",
        "match_type": "none",
    }

    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            day_name = None
            section_name = None

            if table.rows and cell_idx < len(table.rows[0].cells):
                col_header = table.rows[0].cells[cell_idx].text.strip().lower()
                days = ["monday", "tuesday", "wednesday", "thursday", "friday"]
                for day in days:
                    if day in col_header:
                        day_name = day
                        break

            if row.cells:
                row_label = row.cells[0].text.strip().lower()
                if "objective" in row_label:
                    section_name = "objective"
                elif "instruction" in row_label or "tailored" in row_label:
                    section_name = "instruction"
                elif "assessment" in row_label:
                    section_name = "assessment"
                elif "homework" in row_label:
                    section_name = "homework"

            confidence, match_type = renderer._calculate_match_confidence(
                cell.text, link, day_name=day_name, section_name=section_name
            )

            if confidence > best_match["confidence"]:
                best_match["confidence"] = confidence
                best_match["location"] = f"({row_idx}, {cell_idx})"
                best_match["cell_text_preview"] = cell.text[:80]
                best_match["match_type"] = match_type

            if confidence >= threshold:
                inject_hyperlink_inline(renderer, cell, link)
                logger.debug(
                    f"Placed '{link['text']}' via fuzzy matching "
                    f"at ({row_idx}, {cell_idx}), confidence={confidence:.2f}, "
                    f"match_type={match_type}"
                )
                return True

    logger.debug(
        "fuzzy_match_best_attempt",
        extra={
            "link_text": link.get("text", "")[:50],
            "best_confidence": best_match["confidence"],
            "best_location": best_match["location"],
            "threshold": threshold,
            "cell_preview": best_match["cell_text_preview"],
            "match_type": best_match["match_type"],
        },
    )
    return False


def restore_hyperlinks(renderer, doc: Document, hyperlinks: List[Dict]) -> None:
    """Restore hyperlinks by adding them at the end of the document."""
    if not hyperlinks:
        return

    logger.info(
        "hyperlink_fallback_placement",
        extra={
            "total_fallback_links": len(hyperlinks),
            "links": [
                {
                    "text": link.get("text", "")[:50],
                    "url": link.get("url", "")[:50],
                    "day_hint": link.get("day_hint"),
                    "section_hint": link.get("section_hint"),
                    "slot": link.get("_source_slot"),
                    "subject": link.get("_source_subject"),
                }
                for link in hyperlinks
            ],
        },
    )

    try:
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Referenced Links")
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        for link in hyperlinks:
            try:
                paragraph = doc.add_paragraph()
                try:
                    paragraph.style = "List Bullet"
                except KeyError:
                    paragraph.style = "Normal"
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                    paragraph.add_run("• ")

                add_hyperlink(renderer, paragraph, link["text"], link["url"])

            except Exception as e:
                logger.warning(
                    "hyperlink_restoration_failed",
                    extra={
                        "text": link.get("text", "unknown"),
                        "url": link.get("url", "unknown"),
                        "error": str(e),
                    },
                )
    except Exception as e:
        logger.warning("hyperlinks_restoration_error", extra={"error": str(e)})


def add_hyperlink(
    renderer,
    paragraph,
    text: str,
    url: str,
    bold: bool = False,
    insert_at: int = None,
) -> None:
    """Add a hyperlink to a paragraph using direct OxmlElement construction."""
    part = paragraph.part

    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r_elem = OxmlElement("w:r")
    from docx.text.run import Run

    temp_run = Run(r_elem, paragraph)
    force_font_tnr8(temp_run, is_bold=bold, is_hyperlink=True)

    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    r_elem.append(t_elem)

    hyperlink.append(r_elem)

    if insert_at is not None and insert_at < len(paragraph._p):
        paragraph._p.insert(insert_at, hyperlink)
    else:
        paragraph._p.append(hyperlink)
