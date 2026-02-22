"""
Render pipeline: load template, fill metadata and days, normalize, save.

Runs the main steps of DOCX rendering. Called by DOCXRenderer.render().
"""

from pathlib import Path
from typing import Any, Dict, List

from docx import Document

from backend.performance_tracker import get_tracker
from backend.utils.metadata_utils import get_teacher_name

from . import logger
from . import get_indices as _get_indices_module
from . import fallback_media as _fallback_media_module
from . import style as _style_module
from . import table_cell


def run_render_pipeline(
    renderer,
    json_data: Dict,
    output_path: Any,
    plan_id: str,
    skip_fallback_sections: bool,
):
    """Run the full render pipeline. Returns same as render()."""
    tracker = get_tracker()
    renderer._reset_state()

    is_stream = hasattr(output_path, "write")
    renderer.current_file = "stream.docx" if is_stream else Path(output_path).name

    metadata = json_data.get("metadata", {})
    renderer.current_teacher = get_teacher_name(metadata)
    renderer._current_metadata = metadata
    renderer.is_originals = metadata.get("source_type") == "originals"

    try:
        if renderer.template_buffer:
            renderer.template_buffer.seek(0)
        doc = Document(
            renderer.template_buffer if renderer.template_buffer else renderer.template_path
        )

        _style_module.ensure_hyperlink_style(doc)

        section = doc.sections[0]
        available_width_emus = (
            section.page_width - section.left_margin - section.right_margin
        )
        available_width_inches = available_width_emus / 914400

        if plan_id:
            with tracker.track_operation(plan_id, "render_fill_metadata"):
                table_cell.fill_metadata(renderer, doc, json_data)
        else:
            table_cell.fill_metadata(renderer, doc, json_data)

        from tools.docx_utils import normalize_table_column_widths

        normalize_table_column_widths(
            doc.tables[renderer.METADATA_TABLE_IDX],
            total_width_inches=available_width_inches,
        )

        schema_version = json_data.get("_media_schema_version", "1.0")
        pending_hyperlinks = (
            json_data.get("_hyperlinks", []).copy()
            if schema_version in ["1.1", "2.0"]
            else []
        )
        pending_images = (
            json_data.get("_images", []).copy()
            if schema_version in ["1.1", "2.0"]
            else []
        )

        slot_number = json_data.get("metadata", {}).get("slot_number")
        subject = json_data.get("metadata", {}).get("subject")

        logger.info(
            "renderer_slot_metadata_extracted",
            extra={
                "slot_number": slot_number,
                "subject": subject,
                "has_hyperlinks": len(json_data.get("_hyperlinks", [])),
                "teacher": get_teacher_name(json_data.get("metadata", {})),
            },
        )

        from tools.diagnostic_logger import get_diagnostic_logger

        diag = get_diagnostic_logger()
        diag.log_renderer_extracted_metadata(
            slot_number,
            subject,
            len(json_data.get("_hyperlinks", [])),
            get_teacher_name(json_data.get("metadata", {})),
        )

        table = doc.tables[renderer.DAILY_PLANS_TABLE_IDX]

        def _fill_all_days():
            for day_name, day_data in json_data["days"].items():
                if not day_data:
                    continue
                col_idx = _get_indices_module.get_col_index(renderer, day_name)
                if col_idx == -1:
                    continue
                if "slots" in day_data:
                    table_cell.fill_multi_slot_day(
                        renderer,
                        table,
                        col_idx,
                        day_data["slots"],
                        metadata=json_data.get("metadata", {}),
                        day_name=day_name,
                        pending_hyperlinks=pending_hyperlinks,
                        pending_images=pending_images,
                    )
                else:
                    table_cell.fill_single_slot_day(
                        renderer,
                        table,
                        col_idx,
                        day_data,
                        day_name=day_name,
                        pending_hyperlinks=pending_hyperlinks,
                        pending_images=pending_images,
                        slot_number=slot_number,
                        subject=subject,
                    )

        if plan_id:
            with tracker.track_operation(plan_id, "render_fill_days"):
                _fill_all_days()
        else:
            _fill_all_days()

        if not skip_fallback_sections:
            if pending_hyperlinks or pending_images:
                _fallback_media_module.append_unmatched_media(
                    renderer, doc, pending_hyperlinks, pending_images
                )
                logger.info(
                    "unmatched_media_appended",
                    extra={
                        "hyperlinks": len(pending_hyperlinks),
                        "images": len(pending_images),
                    },
                )

            if schema_version == "1.0":
                if "_images" in json_data and json_data["_images"]:
                    if plan_id:
                        with tracker.track_operation(
                            plan_id, "render_insert_images"
                        ):
                            _fallback_media_module.insert_images(
                                renderer, doc, json_data["_images"]
                            )
                            logger.info(
                                "images_inserted",
                                extra={"count": len(json_data["_images"])},
                            )
                    else:
                        _fallback_media_module.insert_images(
                            renderer, doc, json_data["_images"]
                        )
                        logger.info(
                            "images_inserted",
                            extra={"count": len(json_data["_images"])},
                        )

                if "_hyperlinks" in json_data and json_data["_hyperlinks"]:
                    if plan_id:
                        with tracker.track_operation(
                            plan_id, "render_restore_hyperlinks"
                        ):
                            renderer._restore_hyperlinks(
                                doc, json_data["_hyperlinks"]
                            )
                            logger.info(
                                "hyperlinks_restored",
                                extra={"count": len(json_data["_hyperlinks"])},
                            )
                    else:
                        renderer._restore_hyperlinks(
                            doc, json_data["_hyperlinks"]
                        )
                        logger.info(
                            "hyperlinks_restored",
                            extra={"count": len(json_data["_hyperlinks"])},
                        )
        else:
            logger.info(
                "unmatched_media_skipped",
                extra={
                    "hyperlinks": len(pending_hyperlinks),
                    "images": len(pending_images),
                    "note": "consolidation_active",
                },
            )

        if renderer.is_originals:
            _style_module.apply_originals_cleanup(
                doc,
                renderer.DAILY_PLANS_TABLE_IDX,
                _style_module.is_signature_table,
            )

        if plan_id:
            with tracker.track_operation(plan_id, "render_normalize_tables"):
                from tools.docx_utils import normalize_all_tables

                table_count = normalize_all_tables(
                    doc, total_width_inches=available_width_inches
                )
                logger.info(
                    "tables_normalized",
                    extra={
                        "count": table_count,
                        "width_inches": available_width_inches,
                    },
                )
        else:
            from tools.docx_utils import normalize_all_tables

            table_count = normalize_all_tables(
                doc, total_width_inches=available_width_inches
            )
            logger.info(
                "tables_normalized",
                extra={
                    "count": table_count,
                    "width_inches": available_width_inches,
                },
            )

        if plan_id:
            with tracker.track_operation(plan_id, "render_save_docx"):
                if not is_stream:
                    output_path_obj = Path(output_path)
                    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
                doc.save(output_path)
        else:
            if not is_stream:
                output_path_obj = Path(output_path)
                output_path_obj.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)

        logger.info("docx_render_success", extra={"output_path": str(output_path)})

        if skip_fallback_sections:
            return True, pending_hyperlinks, pending_images
        return True

    except Exception as e:
        logger.exception(
            "docx_render_error",
            extra={
                "output_path": str(output_path),
                "error": str(e),
                "error_type": type(e).__name__,
            },
        )
        if skip_fallback_sections:
            return False, [], []
        raise RuntimeError(
            f"Renderer failed to create DOCX file '{output_path}': {str(e)}"
        ) from e
