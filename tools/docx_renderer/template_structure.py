"""
Initialize table structure metadata from the DOCX template.

Finds the daily-plans table (first non-signature table after metadata) and
sets renderer.DAILY_PLANS_TABLE_IDX and renderer.structure_metadata.
"""

from docx import Document

from .style import is_signature_table


def initialize_structure(renderer) -> None:
    """Initialize table structure metadata from template.

    Sets renderer.DAILY_PLANS_TABLE_IDX and renderer.structure_metadata.
    """
    try:
        source = (
            renderer.template_buffer
            if renderer.template_buffer
            else renderer.template_path
        )
        if not source:
            return
        if hasattr(source, "seek"):
            source.seek(0)
        doc = Document(source)
        if len(doc.tables) <= 1:
            renderer.logger.warning(
                "Template has fewer than 2 tables, cannot detect structure"
            )
            return

        for i in range(1, len(doc.tables)):
            table = doc.tables[i]
            if not is_signature_table(table):
                renderer.DAILY_PLANS_TABLE_IDX = i
                renderer.structure_metadata = (
                    renderer.structure_detector.detect_structure(table)
                )
                renderer.logger.info(
                    "template_structure_detected",
                    extra={
                        "type": renderer.structure_metadata.structure_type,
                        "rows": renderer.structure_metadata.num_rows,
                        "cols": renderer.structure_metadata.num_cols,
                        "has_day_row": renderer.structure_metadata.has_day_row,
                        "daily_plans_table_idx": renderer.DAILY_PLANS_TABLE_IDX,
                    },
                )
                return

        renderer.logger.error("Could not locate daily plans table in template")
    except Exception as e:
        renderer.logger.error(f"Failed to initialize structure: {e}")
