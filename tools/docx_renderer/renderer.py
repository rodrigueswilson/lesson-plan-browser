"""
DOCX Renderer - Convert validated JSON to DOCX using district template.

This renderer:
1. Loads the district template (input/Lesson Plan Template SY'25-26.docx)
2. Fills in metadata (teacher name, grade, subject, week, homeroom)
3. Fills in daily plans for Monday-Friday
4. Preserves all original formatting
"""

import base64
import json
import sys
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.run import Run
import re
import time
import json

from backend.performance_tracker import get_tracker
from . import logger
from backend.utils.metadata_utils import get_teacher_name
from tools.table_structure import TableStructureDetector

from . import style as _style_module
from . import hyperlink_placement as _hyperlink_module
from . import table_cell
from .style import FUZZY_MATCH_THRESHOLD

# Handle imports for both CLI and module usage
try:
    from tools.markdown_to_docx import MarkdownToDocx
except ImportError:
    from markdown_to_docx import MarkdownToDocx


sanitize_xml_text = _style_module.sanitize_xml_text
is_signature_table = _style_module.is_signature_table


class DOCXRenderer:
    """Render validated JSON lesson plans to DOCX format."""

    # Table indices in the template document
    METADATA_TABLE_IDX = 0  # First table contains metadata (Name, Grade, etc.)
    DAILY_PLANS_TABLE_IDX = 1  # Second table contains daily lesson plans

    # Row indices in the daily plans table (typical layout; dynamic via _get_row_index when structure is detected)
    UNIT_LESSON_ROW = 1  # Unit/Lesson row (first data row after headers)
    OBJECTIVE_ROW = 2  # Objective row
    ANTICIPATORY_SET_ROW = 3  # Anticipatory set row
    INSTRUCTION_ROW = 4  # Tailored instruction row
    ASSESSMENT_ROW = 6  # Assessment row

    def __init__(self, template_path: str):
        """Initialize renderer with template path.

        Args:
            template_path: Path to the DOCX template file
        """
        self.template_path = template_path
        self.is_originals = False # Initialized to False, set by renderer logic
        self.logger = logger.bind(component="docx_renderer")
        self.structure_detector = TableStructureDetector()
        
        # Cache template in memory to avoid repeated disk I/O when reused
        self.template_buffer = None
        try:
            with open(template_path, "rb") as f:
                self.template_buffer = BytesIO(f.read())
        except Exception as e:
            self.logger.error(f"Failed to cache template buffer: {e}")

        # Initialize structure metadata
        self.structure_metadata = None
        self._initialize_structure()

        # Reset per-render state
        self._reset_state()

    def _reset_state(self):
        """Reset internal state before each render call."""
        self.placement_stats = {
            "coordinate": 0,
            "label_day": 0,
            "fuzzy": 0,
            "fallback": 0,
        }
        self.current_file = None
        self.current_teacher = None
        self._current_metadata = {}
        self.is_originals = False

    def _initialize_structure(self):
        """Initialize table structure metadata from template."""
        try:
            # Use buffer if available for structure detection too
            source = self.template_buffer if self.template_buffer else self.template_path
            if source:
                if hasattr(source, "seek"):
                    source.seek(0)
                doc = Document(source)
                if len(doc.tables) > 1:
                
                    # Find the daily plans table dynamically
                    # It's usually the first table after the metadata table (index 0)
                    # that is not a signature table.
                    daily_plans_table_found = False
                    for i in range(1, len(doc.tables)): # Start from index 1, assuming metadata is at 0
                        table = doc.tables[i]
                        if not is_signature_table(table):
                            self.DAILY_PLANS_TABLE_IDX = i
                            self.structure_metadata = self.structure_detector.detect_structure(table)
                            daily_plans_table_found = True
                            break
                    
                    if daily_plans_table_found:
                        self.logger.info(
                            "template_structure_detected",
                            extra={
                                "type": self.structure_metadata.structure_type,
                                "rows": self.structure_metadata.num_rows,
                                "cols": self.structure_metadata.num_cols,
                                "has_day_row": self.structure_metadata.has_day_row,
                                "daily_plans_table_idx": self.DAILY_PLANS_TABLE_IDX,
                            },
                        )
                    else:
                        self.logger.error("Could not locate daily plans table in template")
                else:
                    self.logger.warning(
                        "Template has fewer than 2 tables, cannot detect structure"
                    )
        except Exception as e:
            self.logger.error(f"Failed to initialize structure: {e}")

    def _get_row_index(self, label: str) -> int:
        """Get row index for a given label using structure metadata."""
        if self.structure_metadata:
            idx = self.structure_metadata.get_row_index(label)
            if idx is not None:
                # Apply offset if needed (e.g. for day row)
                return idx + self.structure_metadata.row_offset
        return -1

    def _get_col_index(self, day: str) -> int:
        """Get column index for a given day using structure metadata."""
        if self.structure_metadata:
            idx = self.structure_metadata.get_col_index(day)
            return idx if idx is not None else -1
        return -1

    def render(
        self,
        json_data: Dict,
        output_path: Any,  # Can be str, Path, or BytesIO
        plan_id: str = None,
        skip_fallback_sections: bool = False,
    ) -> Any:
        """
        Render JSON data to DOCX with semantic anchoring for media.

        Args:
            json_data: Validated lesson plan JSON (supports both single-slot and multi-slot)
            output_path: Path to save DOCX file or BytesIO stream
            plan_id: Optional plan ID for performance tracking
            skip_fallback_sections: If True, don't append "Referenced Links" or "Attached Images" sections
                                   (used when rendering slots for later merging)

        Returns:
            If skip_fallback_sections=True: Tuple (success: bool, pending_hyperlinks: List, pending_images: List)
            Otherwise: True if successful, False otherwise
        """
        tracker = get_tracker()
        
        # Reset per-render state
        self._reset_state()

        # Set context for logging
        is_stream = hasattr(output_path, "write")
        self.current_file = "stream.docx" if is_stream else Path(output_path).name
        
        metadata = json_data.get("metadata", {})
        self.current_teacher = get_teacher_name(metadata)
        # Store metadata for use in helper methods
        self._current_metadata = metadata
        self.is_originals = (metadata.get("source_type") == "originals")

        try:
            # Load template
            if self.template_buffer:
                self.template_buffer.seek(0)
            doc = Document(self.template_buffer if self.template_buffer else self.template_path)
            
            # NOTE: We no longer modify global styles or docDefaults here to avoid 
            # corruption during docxcompose merging. All formatting is applied 
            # per-run in the daily grid cells.

            _style_module.ensure_hyperlink_style(doc)

            # Prepare structure metadata
            # Calculate table width once for use throughout rendering
            section = doc.sections[0]
            available_width_emus = (
                section.page_width - section.left_margin - section.right_margin
            )
            available_width_inches = available_width_emus / 914400

            # Track metadata filling
            if plan_id:
                with tracker.track_operation(plan_id, "render_fill_metadata"):
                    table_cell.fill_metadata(self, doc, json_data)
            else:
                table_cell.fill_metadata(self, doc, json_data)

            # Normalize metadata table immediately after filling to prevent Word auto-resizing
            from tools.docx_utils import normalize_table_column_widths

            normalize_table_column_widths(
                doc.tables[self.METADATA_TABLE_IDX],
                total_width_inches=available_width_inches,
            )

            # Prepare media for semantic anchoring
            schema_version = json_data.get("_media_schema_version", "1.0")
            # v1.1 and v2.0 both use semantic anchoring (v2.0 adds coordinates)
            pending_hyperlinks = (
                json_data.get("_hyperlinks", []).copy()
                if schema_version in ["1.1", "2.0"]
                else []
            )
            pending_images = (
                json_data.get("_images", []).copy()
                if schema_version in ["1.1", "2.0"]
                else []
            )

            # v2.0: Hyperlink placement is now handled during cell filling to ensure 
            # that inline replacement is always attempted before top-of-cell placement.
            # Coordination and fuzzy matching have been unified within _fill_cell.
            is_multi_slot = any(
                "slots" in day_data for day_data in json_data.get("days", {}).values()
            )

            # Extract slot metadata for filtering (if present)
            # In multi-slot batch processing, each lesson JSON has slot metadata
            metadata = json_data.get("metadata", {})
            slot_number = metadata.get("slot_number")
            subject = metadata.get("subject")

            # DEBUG: Log slot metadata extraction
            logger.info(
                "renderer_slot_metadata_extracted",
                extra={
                    "slot_number": slot_number,
                    "subject": subject,
                    "has_hyperlinks": len(json_data.get("_hyperlinks", [])),
                    "teacher": get_teacher_name(json_data.get("metadata", {})),
                },
            )

            # DIAGNOSTIC: Log renderer metadata extraction
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_renderer_extracted_metadata(
                slot_number,
                subject,
                len(json_data.get("_hyperlinks", [])),
                get_teacher_name(json_data.get("metadata", {})),
            )

            # Fill daily plans
            table = doc.tables[self.DAILY_PLANS_TABLE_IDX]
            
            def _fill_all_days():
                for day_name, day_data in json_data["days"].items():
                    # Skip if day_data is None
                    if not day_data:
                        continue
                        
                    col_idx = self._get_col_index(day_name)
                    # print(f"  [DEBUG RENDERER] Day: {day_name}, col_idx: {col_idx}")
                    if col_idx == -1:
                        continue

                    # Check if this day has multiple slots
                    if "slots" in day_data:
                        table_cell.fill_multi_slot_day(
                            self, table, col_idx, day_data["slots"],
                            metadata=json_data.get("metadata", {}),
                            day_name=day_name,
                            pending_hyperlinks=pending_hyperlinks,
                            pending_images=pending_images,
                        )
                    else:
                        table_cell.fill_single_slot_day(
                            self, table, col_idx, day_data,
                            day_name=day_name,
                            pending_hyperlinks=pending_hyperlinks,
                            pending_images=pending_images,
                            slot_number=slot_number,
                            subject=subject,
                        )  # Append unmatched media to fallback sections

            if plan_id:
                with tracker.track_operation(plan_id, "render_fill_days"):
                    _fill_all_days()
            else:
                _fill_all_days()

            if not skip_fallback_sections:
                if pending_hyperlinks or pending_images:
                    self._append_unmatched_media(doc, pending_hyperlinks, pending_images)
                    logger.info(
                        "unmatched_media_appended",
                        extra={
                            "hyperlinks": len(pending_hyperlinks),
                            "images": len(pending_images),
                        },
                    )

                # Legacy behavior for schema v1.0
                if schema_version == "1.0":
                    if "_images" in json_data and json_data["_images"]:
                        if plan_id:
                            with tracker.track_operation(plan_id, "render_insert_images"):
                                self._insert_images(doc, json_data["_images"])
                                logger.info(
                                    "images_inserted",
                                    extra={"count": len(json_data["_images"])},
                                )
                        else:
                            self._insert_images(doc, json_data["_images"])
                            logger.info(
                                "images_inserted",
                                extra={"count": len(json_data["_images"])},
                            )

                    if "_hyperlinks" in json_data and json_data["_hyperlinks"]:
                        if plan_id:
                            with tracker.track_operation(
                                plan_id, "render_restore_hyperlinks"
                            ):
                                self._restore_hyperlinks(doc, json_data["_hyperlinks"])
                                logger.info(
                                    "hyperlinks_restored",
                                    extra={"count": len(json_data["_hyperlinks"])},
                                )
                        else:
                            self._restore_hyperlinks(doc, json_data["_hyperlinks"])
                            logger.info(
                                    "hyperlinks_restored",
                                extra={"count": len(json_data["_hyperlinks"])},
                            )
            else:
                # Still log remaining count even if skipped
                logger.info(
                    "unmatched_media_skipped",
                    extra={
                        "hyperlinks": len(pending_hyperlinks),
                        "images": len(pending_images),
                        "note": "consolidation_active",
                    },
                )

            if self.is_originals:
                _style_module.apply_originals_cleanup(
                    doc, self.DAILY_PLANS_TABLE_IDX, _style_module.is_signature_table
                )

            # Track table normalization
            # Use the same width calculated earlier for consistency

            if plan_id:
                with tracker.track_operation(plan_id, "render_normalize_tables"):
                    from tools.docx_utils import normalize_all_tables

                    table_count = normalize_all_tables(
                        doc, total_width_inches=available_width_inches
                    )
                    logger.info(
                        "tables_normalized",
                        extra={
                            "count": table_count,
                            "width_inches": available_width_inches,
                        },
                    )
            else:
                from tools.docx_utils import normalize_all_tables

                table_count = normalize_all_tables(
                    doc, total_width_inches=available_width_inches
                )
                logger.info(
                    "tables_normalized",
                    extra={
                        "count": table_count,
                        "width_inches": available_width_inches,
                    },
                )

            # Track file saving
            if plan_id:
                with tracker.track_operation(plan_id, "render_save_docx"):
                    if not is_stream:
                        output_path_obj = Path(output_path)
                        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
                    doc.save(output_path)
            else:
                if not is_stream:
                    output_path_obj = Path(output_path)
                    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
                doc.save(output_path)

            logger.info("docx_render_success", extra={"output_path": str(output_path)})
            
            # #region agent log
            import json
            with open(r'd:\LP\.cursor\debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "run1",
                    "hypothesisId": "RENDER_SUCCESS",
                    "location": "docx_renderer.py:render:success",
                    "message": "Render completed successfully",
                    "data": {
                        "output_path": str(output_path),
                        "skip_fallback_sections": skip_fallback_sections,
                        "is_stream": hasattr(output_path, "write"),
                    },
                    "timestamp": int(__import__('time').time() * 1000)
                }) + '\n')
            # #endregion
            
            # RETURN VALUE CHANGE: Return unplaced media if skip_fallback_sections is True
            if skip_fallback_sections:
                return True, pending_hyperlinks, pending_images
            
            return True

        except Exception as e:
            # #region agent log
            import json
            import traceback
            with open(r'd:\LP\.cursor\debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "run1",
                    "hypothesisId": "RENDER_EXCEPTION",
                    "location": "docx_renderer.py:render:exception",
                    "message": "Renderer exception caught",
                    "data": {
                        "error": str(e),
                        "error_type": type(e).__name__,
                        "output_path": str(output_path) if "output_path" in locals() else None,
                        "traceback": traceback.format_exc()[:500],
                    },
                    "timestamp": int(__import__('time').time() * 1000)
                }) + '\n')
            # #endregion
            logger.exception(
                "docx_render_error",
                extra={
                    "output_path": str(output_path) if "output_path" in locals() else None,
                    "error": str(e),
                    "error_type": type(e).__name__,
                },
            )
            
            if skip_fallback_sections:
                return False, [], []
            
            # Re-raise the exception instead of returning False
            # This allows callers to see the actual error
            raise RuntimeError(
                f"Renderer failed to create DOCX file '{output_path}': {str(e)}"
            ) from e

    def _abbreviate_content(
        self, content: str, num_slots: int, max_length: int = None
    ) -> str:
        return table_cell.abbreviate_content(self, content, num_slots, max_length)

    def _fill_day(
        self,
        doc,
        day_name: str,
        day_data: Dict,
        pending_hyperlinks: List[Dict] = None,
        pending_images: List[Dict] = None,
        slot_number: int = None,
        subject: str = None,
    ):
        table_cell.fill_day(
            self, doc, day_name, day_data,
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            slot_number=slot_number,
            subject=subject,
        )

    def _fill_cell(
        self,
        table,
        row_idx: int,
        col_idx: int,
        text: str,
        day_name: str = None,
        section_name: str = None,
        pending_hyperlinks: List[Dict] = None,
        pending_images: List[Dict] = None,
        current_slot_number: int = None,
        current_subject: str = None,
        append_mode: bool = False,
    ):
        table_cell.fill_cell(
            self, table, row_idx, col_idx, text,
            day_name=day_name,
            section_name=section_name,
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=current_slot_number,
            current_subject=current_subject,
            append_mode=append_mode,
        )

    def _calculate_match_confidence(
        self,
        cell_text: str,
        media: Dict,
        day_name: str = None,
        section_name: str = None,
    ) -> tuple:
        return table_cell.calculate_match_confidence(
            self, cell_text, media, day_name, section_name
        )

    def _inject_hyperlink_inline(self, cell, hyperlink: Dict, row_idx: int = None):
        """Inject hyperlink into cell on its own line."""
        _hyperlink_module.inject_hyperlink_inline(self, cell, hyperlink, row_idx=row_idx)

    def _inject_image_inline(self, cell, image: Dict, max_width: float):
        """Inject image into cell with width constraints.

        Args:
            cell: Cell object
            image: Image dictionary with base64 data
            max_width: Maximum width in inches
        """
        try:
            image_data = base64.b64decode(image["data"])
            image_stream = BytesIO(image_data)

            # Add to new paragraph
            para = cell.add_paragraph()
            run = para.add_run()
            run.add_picture(
                image_stream, width=Inches(max_width * 0.9)
            )  # 90% of column width

            # Optional: tiny caption
            caption = cell.add_paragraph()
            caption_run = caption.add_run(f"[{image.get('filename', 'image')}]")
            caption_run.font.size = Pt(7)
            caption_run.italic = True

        except Exception as e:
            logger.warning(
                "inline_image_failed",
                extra={"filename": image.get("filename"), "error": str(e)},
            )

    def _append_unmatched_media(
        self, doc: Document, hyperlinks: List[Dict], images: List[Dict]
    ):
        """Append unmatched media to end with context for traceability.

        Args:
            doc: Document object
            hyperlinks: List of unmatched hyperlinks
            images: List of unmatched images
        """
        if hyperlinks:
            doc.add_page_break()
            heading = doc.add_paragraph("Referenced Links")
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(12)

            for link in hyperlinks:
                para = doc.add_paragraph()
                try:
                    para.style = "List Bullet"
                except KeyError:
                    para.add_run("• ")

                self._add_hyperlink(para, link["text"], link["url"])

                # Add context hint
                if link.get("context_snippet"):
                    context_para = doc.add_paragraph()
                    context_run = context_para.add_run(
                        f'  Context: "{link["context_snippet"][:80]}..."'
                    )
                    context_run.font.size = Pt(9)
                    context_run.italic = True

        if images:
            if not hyperlinks:
                doc.add_page_break()

            heading = doc.add_paragraph("Attached Images")
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(12)

            for i, image in enumerate(images, 1):
                try:
                    image_data = base64.b64decode(image["data"])
                    image_stream = BytesIO(image_data)

                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(4.0))

                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(
                        f"Image {i}: {image.get('filename', 'Untitled')}"
                    )
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)

                    # Add context hint
                    if image.get("context_snippet"):
                        context_para = doc.add_paragraph()
                        context_run = context_para.add_run(
                            f'Context: "{image["context_snippet"][:80]}..."'
                        )
                        context_run.font.size = Pt(9)
                        context_run.italic = True

                except Exception as e:
                    logger.warning(
                        "fallback_image_failed",
                        extra={"filename": image.get("filename"), "error": str(e)},
                    )

    def _insert_images(self, doc: Document, images: List[Dict]):
        """Insert images at the end of the document.

        Args:
            doc: Document object
            images: List of image dictionaries with base64-encoded data
        """
        if not images:
            return

        try:
            # Add a page break before images section
            doc.add_page_break()

            # Add section header
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Attached Images")
            heading_run.bold = True
            heading_run.font.size = Pt(14)

            # Insert each image
            for i, image in enumerate(images, 1):
                try:
                    # Decode base64 image data
                    image_data = base64.b64decode(image["data"])
                    image_stream = BytesIO(image_data)

                    # Add image with caption
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()

                    # Insert image with reasonable width (4 inches)
                    run.add_picture(image_stream, width=Inches(4.0))

                    # Add caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(
                        f"Image {i}: {image.get('filename', 'Untitled')}"
                    )
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)

                except Exception as e:
                    logger.warning(
                        "image_insertion_failed",
                        extra={
                            "image_index": i,
                            "filename": image.get("filename", "unknown"),
                            "error": str(e),
                        },
                    )
        except Exception as e:
            logger.warning("images_insertion_error", extra={"error": str(e)})

    def _restore_hyperlinks(self, doc: Document, hyperlinks: List[Dict]):
        """Restore hyperlinks by adding them at the end of the document."""
        _hyperlink_module.restore_hyperlinks(self, doc, hyperlinks)

    def _force_font_tnr8(self, run, is_bold: bool = False, is_hyperlink: bool = False):
        """Force Times New Roman 8pt on a Run object."""
        _style_module.force_font_tnr8(run, is_bold=is_bold, is_hyperlink=is_hyperlink)

    def _force_font_arial10(self, run, is_bold: bool = False):
        """Force Arial 10pt on a Run object. Used for metadata table formatting."""
        _style_module.force_font_arial10(run, is_bold=is_bold)

    def _add_hyperlink(self, paragraph, text: str, url: str, bold: bool = False, insert_at: int = None):
        """Add a hyperlink to a paragraph."""
        _hyperlink_module.add_hyperlink(self, paragraph, text, url, bold=bold, insert_at=insert_at)

def main():
    """CLI entry point."""
    if len(sys.argv) < 3:
        print(
            "Usage: python docx_renderer.py <input.json> <output.docx> [template.docx]"
        )
        print("\nExample:")
        print(
            "  python docx_renderer.py tests/fixtures/valid_lesson_minimal.json output/lesson.docx"
        )
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    template_path = (
        sys.argv[3] if len(sys.argv) > 3 else "input/Lesson Plan Template SY'25-26.docx"
    )

    # Load JSON
    with open(input_path, "r", encoding="utf-8") as f:
        json_data = json.load(f)

    # Render
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
