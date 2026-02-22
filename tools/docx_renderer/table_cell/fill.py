"""
Fill logic for DOCX table/cell rendering.

Provides: fill_metadata, fill_day, fill_single_slot_day, fill_multi_slot_day, fill_cell.
"""

import re
import time
from typing import Any, Dict, List, Optional

from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run

from .. import logger
from backend.services.sorting_utils import sort_slots
from backend.utils.metadata_utils import get_homeroom, get_subject, get_teacher_name

try:
    from tools.markdown_to_docx import MarkdownToDocx
except ImportError:
    from markdown_to_docx import MarkdownToDocx

from .. import style as _style_module
from . import format as format_module
from . import placement as placement_module

sanitize_xml_text = _style_module.sanitize_xml_text


def fill_metadata(renderer, doc, json_data: Dict) -> None:
    """
    Fill metadata table (Table 0).

    Template structure:
    | Name: | Grade: | Homeroom: | Subject: | Week of: | Room: |

    For multi-slot lessons, extracts slot-specific metadata from the first slot
    found across all days, using standardized metadata utilities.

    Args:
        renderer: DOCXRenderer instance
        doc: Document object
        json_data: Full lesson plan JSON (supports both single-slot and multi-slot)
    """
    metadata = json_data.get("metadata", {})

    is_multi_slot = any(
        day_data and "slots" in day_data
        for day_data in json_data.get("days", {}).values()
    )

    representative_slot = None
    if "days" in json_data and isinstance(json_data["days"], dict):
        all_slots = []
        for day_name, day_data in json_data["days"].items():
            if day_data and "slots" in day_data and isinstance(day_data["slots"], list):
                all_slots.extend(day_data["slots"])
        if all_slots:
            sorted_slots = sorted(all_slots, key=lambda x: x.get("slot_number", 0))
            representative_slot = sorted_slots[0]

    table = doc.tables[renderer.METADATA_TABLE_IDX]
    row = table.rows[0]

    teacher_name = get_teacher_name(metadata, slot=representative_slot)
    cell = row.cells[0]
    cell.text = ""
    para = cell.paragraphs[0]
    run1 = para.add_run(sanitize_xml_text("Name: "))
    renderer._force_font_arial10(run1, is_bold=True)
    run2 = para.add_run(
        sanitize_xml_text(
            teacher_name if teacher_name and teacher_name != "Unknown" else "Unknown"
        )
    )
    renderer._force_font_arial10(run2, is_bold=False)

    grade = None
    if representative_slot:
        grade = representative_slot.get("grade")
    if not grade or grade == "N/A":
        grade = metadata.get("grade")
    if grade and grade != "Unknown" and grade != "N/A":
        cell = row.cells[1]
        cell.text = ""
        para = cell.paragraphs[0]
        run1 = para.add_run(sanitize_xml_text("Grade: "))
        renderer._force_font_arial10(run1, is_bold=True)
        run2 = para.add_run(sanitize_xml_text(grade))
        renderer._force_font_arial10(run2, is_bold=False)

    homeroom = get_homeroom(metadata, slot=representative_slot)
    if homeroom and homeroom != "Unknown":
        cell = row.cells[2]
        cell.text = ""
        para = cell.paragraphs[0]
        run1 = para.add_run(sanitize_xml_text("Homeroom: "))
        renderer._force_font_arial10(run1, is_bold=True)
        run2 = para.add_run(sanitize_xml_text(homeroom))
        renderer._force_font_arial10(run2, is_bold=False)

    subject = get_subject(metadata, slot=representative_slot)
    if subject and subject != "Unknown":
        cell = row.cells[3]
        cell.text = ""
        para = cell.paragraphs[0]
        run1 = para.add_run(sanitize_xml_text("Subject: "))
        renderer._force_font_arial10(run1, is_bold=True)
        run2 = para.add_run(sanitize_xml_text(subject))
        renderer._force_font_arial10(run2, is_bold=False)

    week_of = metadata.get("week_of", "Unknown")
    if week_of and week_of != "Unknown":
        cell = row.cells[4]
        cell.text = ""
        para = cell.paragraphs[0]
        run1 = para.add_run(sanitize_xml_text("Week of: "))
        renderer._force_font_arial10(run1, is_bold=True)
        run2 = para.add_run(sanitize_xml_text(week_of))
        renderer._force_font_arial10(run2, is_bold=False)

    if len(row.cells) > 5:
        room = None
        if representative_slot:
            room = representative_slot.get("room")
        if (
            not room
            or room == "N/A"
            or (isinstance(room, str) and not room.strip())
        ):
            room = metadata.get("room", "")
        if room and room.strip() and room != "N/A" and room != "Unknown":
            cell = row.cells[5]
            cell.text = ""
            para = cell.paragraphs[0]
            run1 = para.add_run(sanitize_xml_text("Room: "))
            renderer._force_font_arial10(run1, is_bold=True)
            run2 = para.add_run(sanitize_xml_text(room))
            renderer._force_font_arial10(run2, is_bold=False)


def fill_day(
    renderer,
    doc,
    day_name: str,
    day_data: Dict,
    pending_hyperlinks: List[Dict] = None,
    pending_images: List[Dict] = None,
    slot_number: int = None,
    subject: str = None,
) -> None:
    """
    Fill a single day's data with media injection.

    Args:
        renderer: DOCXRenderer instance
        doc: Document object
        day_name: Day name (monday, tuesday, etc.)
        day_data: Day's lesson plan data (can be single-slot or multi-slot)
        pending_hyperlinks: List of hyperlinks awaiting placement
        pending_images: List of images awaiting placement
        slot_number: Slot number for filtering (if rendering single slot)
        subject: Subject name for filtering (if rendering single slot)
    """
    table = doc.tables[renderer.DAILY_PLANS_TABLE_IDX]
    col_idx = renderer._get_col_index(day_name)
    if col_idx == -1:
        return

    hyperlinks_before = len(pending_hyperlinks) if pending_hyperlinks else 0

    if "slots" in day_data and isinstance(day_data["slots"], list):
        metadata = getattr(renderer, "_current_metadata", {})
        fill_multi_slot_day(
            renderer,
            table,
            col_idx,
            day_data["slots"],
            metadata=metadata,
            day_name=day_name,
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
        )
    else:
        fill_single_slot_day(
            renderer,
            table,
            col_idx,
            day_data,
            day_name=day_name,
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            slot_number=slot_number,
            subject=subject,
        )

    hyperlinks_after = len(pending_hyperlinks) if pending_hyperlinks else 0
    placed_count = hyperlinks_before - hyperlinks_after
    if hyperlinks_before > 0:
        logger.info(
            "hyperlink_placement_outcome",
            extra={
                "day": day_name,
                "slot": slot_number,
                "subject": subject,
                "total_links": hyperlinks_before,
                "placed_count": placed_count,
                "remaining_count": hyperlinks_after,
                "placement_rate": placed_count / hyperlinks_before
                if hyperlinks_before > 0
                else 0.0,
            },
        )


def fill_single_slot_day(
    renderer,
    table,
    col_idx: int,
    day_data: Dict,
    day_name: str = None,
    pending_hyperlinks: List[Dict] = None,
    pending_images: List[Dict] = None,
    slot_number: int = None,
    subject: str = None,
) -> None:
    """
    Fill a single slot's data for one day with media injection.
    """
    unit_lesson = day_data.get("unit_lesson", "")
    if unit_lesson and "no school" in unit_lesson.lower():
        fill_cell(
            renderer,
            table,
            renderer.UNIT_LESSON_ROW,
            col_idx,
            unit_lesson,
            day_name=day_name,
            section_name="unit_lesson",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )
        for row_label in [
            "objective",
            "anticipatory",
            "instruction",
            "misconception",
            "assessment",
            "homework",
        ]:
            row_idx = renderer._get_row_index(row_label)
            if row_idx != -1:
                table.rows[row_idx].cells[col_idx].text = ""
        return

    fill_cell(
        renderer,
        table,
        renderer._get_row_index("unit"),
        col_idx,
        unit_lesson,
        day_name=day_name,
        section_name="unit_lesson",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )

    objective_text = format_module.format_objective(renderer, day_data.get("objective", {}))
    fill_cell(
        renderer,
        table,
        renderer._get_row_index("objective"),
        col_idx,
        objective_text,
        day_name=day_name,
        section_name="objective",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )

    anticipatory_text = format_module.format_anticipatory_set(
        renderer, day_data.get("anticipatory_set", {})
    )
    fill_cell(
        renderer,
        table,
        renderer._get_row_index("anticipatory"),
        col_idx,
        anticipatory_text,
        day_name=day_name,
        section_name="anticipatory_set",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )

    instruction_text = format_module.format_tailored_instruction(
        renderer,
        day_data.get("tailored_instruction", {}),
        day_data.get("vocabulary_cognates"),
        day_data.get("sentence_frames"),
    )
    fill_cell(
        renderer,
        table,
        renderer._get_row_index("instruction"),
        col_idx,
        instruction_text,
        day_name=day_name,
        section_name="instruction",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )

    misconceptions_text = format_module.format_misconceptions(
        renderer, day_data.get("misconceptions", {})
    )
    fill_cell(
        renderer,
        table,
        renderer._get_row_index("misconception"),
        col_idx,
        misconceptions_text,
        day_name=day_name,
        section_name="misconceptions",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )

    assessment_text = format_module.format_assessment(
        renderer, day_data.get("assessment", {})
    )
    fill_cell(
        renderer,
        table,
        renderer._get_row_index("assessment"),
        col_idx,
        assessment_text,
        day_name=day_name,
        section_name="assessment",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )

    homework_text = format_module.format_homework(
        renderer, day_data.get("homework", "")
    )
    fill_cell(
        renderer,
        table,
        renderer._get_row_index("homework"),
        col_idx,
        homework_text,
        day_name=day_name,
        section_name="homework",
        pending_hyperlinks=pending_hyperlinks,
        pending_images=pending_images,
        current_slot_number=slot_number,
        current_subject=subject,
    )


def _format_slot_field(
    renderer, field_name: str, slot_content: Any
) -> Optional[str]:
    """Format a slot field using the appropriate format function."""
    if field_name == "unit_lesson":
        return slot_content if isinstance(slot_content, str) else None
    if field_name == "objective":
        return format_module.format_objective(renderer, slot_content)
    if field_name == "anticipatory_set":
        return format_module.format_anticipatory_set(renderer, slot_content)
    if field_name == "tailored_instruction":
        return format_module.format_tailored_instruction(
            renderer,
            slot_content,
            slot_content.get("vocabulary_cognates") if isinstance(slot_content, dict) else None,
            slot_content.get("sentence_frames") if isinstance(slot_content, dict) else None,
        )
    if field_name == "misconceptions":
        return format_module.format_misconceptions(renderer, slot_content)
    if field_name == "assessment":
        return format_module.format_assessment(renderer, slot_content)
    if field_name == "homework":
        return format_module.format_homework(renderer, slot_content)
    return None


def fill_multi_slot_day(
    renderer,
    table,
    col_idx: int,
    slots: List[Dict],
    metadata: Dict = None,
    day_name: str = None,
    pending_hyperlinks: List[Dict] = None,
    pending_images: List[Dict] = None,
) -> None:
    """
    Fill multiple slots' data for one day, with per-slot hyperlink placement.
    """
    if not slots:
        return

    sorted_slots = sort_slots(slots)
    num_slots = len(sorted_slots)

    slots_have_content = []
    for slot in sorted_slots:
        has_content = any(
            [
                slot.get("unit_lesson"),
                slot.get("objective"),
                slot.get("anticipatory_set"),
                slot.get("tailored_instruction"),
                slot.get("misconceptions"),
                slot.get("assessment"),
                slot.get("homework"),
            ]
        )
        slots_have_content.append(has_content)

    rows_config = [
        ("unit_lesson", renderer._get_row_index("unit"), "[No unit/lesson specified]", 100),
        ("objective", renderer._get_row_index("objective"), "[No objective specified]", None),
        ("anticipatory_set", renderer._get_row_index("anticipatory"), None, None),
        ("tailored_instruction", renderer._get_row_index("instruction"), None, None),
        ("misconceptions", renderer._get_row_index("misconception"), None, None),
        ("assessment", renderer._get_row_index("assessment"), None, None),
        ("homework", renderer._get_row_index("homework"), None, 100),
    ]

    if metadata is None:
        metadata = {}

    for (field_name, row_idx, placeholder_text, max_length) in rows_config:
        written_any = False

        for i, slot in enumerate(sorted_slots):
            slot_num = slot.get("slot_number", "?")
            subject = slot.get("subject", "Unknown")
            teacher = get_teacher_name(metadata, slot=slot)
            has_content = slots_have_content[i]

            is_no_school = "no school" in (slot.get("unit_lesson") or "").lower()
            if is_no_school and field_name != "unit_lesson":
                continue

            slot_header = f"**Slot {slot_num}: {subject}**"
            if teacher:
                slot_header += f" ({teacher})"

            slot_content = slot.get(field_name)
            if slot_content:
                slot_text = _format_slot_field(renderer, field_name, slot_content)
                if slot_text is None and field_name == "unit_lesson":
                    slot_text = slot_content
                if slot_text:
                    slot_text = placement_module.abbreviate_content(
                        renderer, slot_text, num_slots, max_length=max_length
                    )
                    slot_text = f"{slot_header}\n{slot_text}"
                else:
                    slot_text = None
            elif has_content and placeholder_text:
                slot_text = f"{slot_header}\n{placeholder_text}"
            else:
                slot_text = None

            if not slot_text:
                continue

            has_next_slot_with_content = False
            for j in range(i + 1, len(sorted_slots)):
                next_slot = sorted_slots[j]
                next_slot_content = next_slot.get(field_name)
                next_has_content = slots_have_content[j]
                if next_slot_content:
                    next_text = _format_slot_field(renderer, field_name, next_slot_content)
                    if next_text is None and field_name == "unit_lesson":
                        next_text = next_slot_content
                    if next_text:
                        has_next_slot_with_content = True
                        break
                elif next_has_content and placeholder_text:
                    has_next_slot_with_content = True
                    break

            if has_next_slot_with_content:
                slot_text += "\n\n---"

            fill_cell(
                renderer,
                table,
                row_idx,
                col_idx,
                slot_text,
                day_name=day_name,
                section_name=field_name,
                pending_hyperlinks=pending_hyperlinks,
                pending_images=pending_images,
                current_slot_number=slot_num,
                current_subject=subject,
                append_mode=written_any,
            )
            written_any = True


def fill_cell(
    renderer,
    table,
    row_idx: int,
    col_idx: int,
    text: str,
    day_name: str = None,
    section_name: str = None,
    pending_hyperlinks: List[Dict] = None,
    pending_images: List[Dict] = None,
    current_slot_number: int = None,
    current_subject: str = None,
    append_mode: bool = False,
) -> None:
    """
    Fill a cell with formatted text and inject matched media inline.
    """
    from backend.config import settings

    _ = (re, qn)

    cell = table.rows[row_idx].cells[col_idx]

    existing_hyperlinks = cell._element.xpath(".//w:hyperlink")
    has_coordinate_hyperlinks = len(existing_hyperlinks) > 0

    if has_coordinate_hyperlinks and text:
        hyperlinks_to_remove = []
        for hl_elem in existing_hyperlinks:
            try:
                r_id = hl_elem.get(qn("r:id"))
                if r_id and cell.paragraphs:
                    para = cell.paragraphs[0]
                    if r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        link_text = "".join(
                            node.text for node in hl_elem.xpath(".//w:t") if node.text
                        )
                        if link_text and url:
                            markdown_pattern = rf"\[{re.escape(link_text)}\]\({re.escape(url)}\)"
                            if re.search(markdown_pattern, text, re.IGNORECASE):
                                hyperlinks_to_remove.append(hl_elem)
            except Exception:
                pass

        for hl_elem in hyperlinks_to_remove:
            try:
                parent = hl_elem.getparent()
                if parent is not None:
                    parent.remove(hl_elem)
                    try:
                        para_elem = parent.getparent()
                        if para_elem is not None and para_elem.tag == qn("w:p"):
                            runs = para_elem.xpath(".//w:r")
                            has_text = any(run.xpath(".//w:t") for run in runs)
                            if not has_text:
                                para_parent = para_elem.getparent()
                                if para_parent is not None:
                                    para_parent.remove(para_elem)
                    except Exception:
                        pass
            except Exception as e:
                logger.warning(
                    "failed_to_remove_duplicate_hyperlink",
                    extra={
                        "error": str(e),
                        "cell": f"{day_name}_{section_name}"
                        if day_name and section_name
                        else "unknown",
                    },
                )

        existing_hyperlinks = cell._element.xpath(".//w:hyperlink")
        has_coordinate_hyperlinks = len(existing_hyperlinks) > 0
        if hyperlinks_to_remove:
            logger.info(
                "removed_duplicate_coordinate_hyperlinks",
                extra={
                    "removed_count": len(hyperlinks_to_remove),
                    "cell": f"{day_name}_{section_name}"
                    if day_name and section_name
                    else "unknown",
                    "slot": current_slot_number,
                },
            )

    if not has_coordinate_hyperlinks and not append_mode:
        paras_to_remove = list(cell.paragraphs)
        for para in paras_to_remove:
            p = para._element
            p.getparent().remove(p)
        new_para = cell.add_paragraph()
        new_para.paragraph_format.space_after = Pt(0)
        new_para.paragraph_format.space_before = Pt(0)
        new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    matching_hyperlinks = []
    if pending_hyperlinks and text:
        logger.info(
            "hyperlink_filtering_context",
            extra={
                "current_slot_number": current_slot_number,
                "current_subject": current_subject,
                "pending_hyperlinks_count": len(pending_hyperlinks),
                "cell": f"{day_name}_{section_name}"
                if day_name and section_name
                else "unknown",
            },
        )
        try:
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_filtering_context(
                current_slot_number,
                current_subject,
                len(pending_hyperlinks),
                f"{day_name}_{section_name}"
                if day_name and section_name
                else "unknown",
            )
        except Exception:
            pass

        for hyperlink in pending_hyperlinks[:]:
            if current_slot_number is not None:
                link_slot = hyperlink.get("_source_slot")
                if link_slot is None:
                    logger.warning(
                        "hyperlink_missing_slot_metadata",
                        extra={
                            "text": hyperlink["text"][:50],
                            "url": hyperlink["url"],
                            "current_slot": current_slot_number,
                        },
                    )
                    continue
                if link_slot != current_slot_number:
                    try:
                        from tools.diagnostic_logger import get_diagnostic_logger

                        get_diagnostic_logger().log_hyperlink_filtered(
                            hyperlink["text"],
                            link_slot,
                            current_slot_number,
                            "slot_mismatch",
                        )
                    except Exception:
                        pass
                    continue

            if current_subject is not None:
                link_subject = hyperlink.get("_source_subject")
                if link_subject is None:
                    logger.warning(
                        "hyperlink_missing_subject_metadata",
                        extra={
                            "text": hyperlink["text"][:50],
                            "url": hyperlink["url"],
                            "current_subject": current_subject,
                        },
                    )
                    continue
                if link_subject != current_subject:
                    continue

            is_structural_match = False
            match_type = "none"
            confidence = 0.0

            if (
                hyperlink.get("row_idx") is not None
                and hyperlink.get("cell_idx") is not None
            ):
                row_offset = (
                    renderer.structure_metadata.row_offset
                    if renderer.structure_metadata
                    else 0
                )
                target_row = hyperlink["row_idx"] + row_offset
                if target_row == row_idx and hyperlink["cell_idx"] == col_idx:
                    is_structural_match = True
                    match_type = "coordinate"
                    confidence = 1.0

            if (
                not is_structural_match
                and hyperlink.get("row_label")
                and hyperlink.get("day_hint")
            ):
                if renderer.structure_metadata:
                    target_row_idx = renderer.structure_metadata.get_row_index(
                        hyperlink["row_label"]
                    )
                    target_col_idx = renderer.structure_metadata.get_col_index(
                        hyperlink["day_hint"]
                    )
                    if target_row_idx is not None:
                        target_row_idx += renderer.structure_metadata.row_offset
                    if target_row_idx == row_idx and target_col_idx == col_idx:
                        is_structural_match = True
                        match_type = "label_day"
                        confidence = 1.0

            if not is_structural_match:
                confidence, match_type = placement_module.calculate_match_confidence(
                    renderer, text, hyperlink, day_name, section_name
                )

            if confidence >= settings.MEDIA_MATCH_CONFIDENCE_THRESHOLD:
                matching_hyperlinks.append((hyperlink, confidence, match_type))

    if matching_hyperlinks and text:
        sorted_matching = sorted(
            matching_hyperlinks,
            key=lambda m: len(m[0].get("text", "")),
            reverse=True,
        )
        for hyperlink, confidence, match_type in sorted_matching:
            link_text = hyperlink.get("text", "")
            link_url = hyperlink.get("url", "")
            if not link_text or not link_url:
                continue

            hl_day = hyperlink.get("day_hint") or hyperlink.get("day")
            if hl_day and day_name:
                if hl_day.lower().strip() != day_name.lower().strip():
                    continue

            search_pattern_raw = re.escape(link_text)
            search_pattern_flexible = re.sub(r"(\\ )+", r"\\s+", search_pattern_raw)
            search_pattern_final = f"{search_pattern_flexible}[.,;:]?"
            match = re.search(search_pattern_final, text, re.IGNORECASE)

            if match:
                raw_match_text = match.group(0)
                stripped_match_text = raw_match_text.strip()
                total_pattern = rf"(\[\s*{re.escape(stripped_match_text)}\s*\]\([^)]*\))|({search_pattern_final})"
                found_and_replaced = [False]

                def replacement_logic(m):
                    if m.group(1):
                        return m.group(1)
                    if not found_and_replaced[0]:
                        found_and_replaced[0] = True
                        plain_text = m.group(2)
                        leading_ws = plain_text[
                            : plain_text.find(stripped_match_text)
                        ]
                        trailing_ws = plain_text[
                            plain_text.find(stripped_match_text)
                            + len(stripped_match_text) :
                        ]
                        return f"{leading_ws}[{stripped_match_text}]({link_url}){trailing_ws}"
                    return m.group(0)

                new_text = re.sub(
                    total_pattern, replacement_logic, text, flags=re.IGNORECASE
                )
                if found_and_replaced[0]:
                    text = new_text
                    logger.info(
                        "hyperlink_placed_smart_inline",
                        extra={
                            "text": stripped_match_text,
                            "url": link_url,
                            "cell": f"{day_name}_{section_name}",
                            "slot": current_slot_number,
                        },
                    )
                else:
                    logger.info(
                        "hyperlink_already_inline_or_unmatched",
                        extra={
                            "text": stripped_match_text,
                            "cell": f"{day_name}_{section_name}",
                            "slot": current_slot_number,
                        },
                    )
                matching_hyperlinks = [
                    m for m in matching_hyperlinks if m[0] != hyperlink
                ]
                if hyperlink in pending_hyperlinks:
                    pending_hyperlinks.remove(hyperlink)

    if not has_coordinate_hyperlinks and not append_mode:
        if text:
            MarkdownToDocx.add_multiline_text(cell, text)
    else:
        if text:
            lines = text.split("\n")
            for i, line in enumerate(lines):
                if line.strip():
                    if i == 0:
                        new_para = cell.add_paragraph()
                        MarkdownToDocx.add_formatted_text(new_para, line)
                    else:
                        MarkdownToDocx.add_paragraph(cell, line)

    if cell.paragraphs:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in para.runs:
                renderer._force_font_tnr8(
                    run, is_bold=(row_idx == renderer.UNIT_LESSON_ROW)
                )
            for hl in para._p.findall(qn("w:hyperlink")):
                for r_elem in hl.findall(qn("w:r")):
                    hl_run = Run(r_elem, para)
                    is_hl_bold = (
                        row_idx == renderer.UNIT_LESSON_ROW or hl_run.font.bold
                    )
                    renderer._force_font_tnr8(
                        hl_run, is_bold=is_hl_bold, is_hyperlink=True
                    )

    links = pending_hyperlinks or []
    imgs = pending_images or []
    if section_name and (links or imgs):
        for hyperlink in links[:]:
            if current_slot_number is not None:
                link_slot = hyperlink.get("_source_slot")
                if link_slot is not None and link_slot != current_slot_number:
                    continue
            hl_day = hyperlink.get("day_hint")
            if hl_day and day_name:
                if hl_day.lower().strip() != day_name.lower().strip():
                    continue
            hint = (hyperlink.get("section_hint") or "").lower()
            is_section_match = False
            section_mappings = {
                "unit_lesson": ["unit", "lesson", "module"],
                "objective": ["objective", "goal", "swbat"],
                "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                "tailored_instruction": [
                    "instruction",
                    "activity",
                    "procedure",
                    "lesson",
                    "tailored",
                    "differentiation",
                ],
                "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                "homework": ["homework", "assignment", "practice"],
            }
            if hint == section_name:
                is_section_match = True
            elif section_name in section_mappings:
                if hint in section_mappings[section_name]:
                    is_section_match = True
                elif any(
                    kw in hint
                    for kw in section_mappings[section_name]
                    if len(kw) > 3
                ):
                    is_section_match = True
            if is_section_match:
                p = cell.add_paragraph()
                MarkdownToDocx.add_formatted_text(
                    p, f"\u2022 [{hyperlink['text']}]({hyperlink['url']})"
                )
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)
                if pending_hyperlinks is not None:
                    pending_hyperlinks.remove(hyperlink)
                logger.info(
                    "hyperlink_placed_fallback",
                    extra={"text": hyperlink["text"], "section": section_name},
                )

        for image in imgs[:]:
            if current_slot_number is not None:
                img_slot = image.get("_source_slot")
                if img_slot is not None and img_slot != current_slot_number:
                    continue
            hint = (image.get("section_hint") or "").lower()
            is_section_match = False
            section_mappings = {
                "unit_lesson": ["unit", "lesson", "module"],
                "objective": ["objective", "goal", "swbat"],
                "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                "tailored_instruction": [
                    "instruction",
                    "activity",
                    "procedure",
                    "lesson",
                    "tailored",
                    "differentiation",
                ],
                "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                "homework": ["homework", "assignment", "practice"],
            }
            if hint == section_name:
                is_section_match = True
            elif section_name in section_mappings:
                if hint in section_mappings[section_name]:
                    is_section_match = True
                elif any(
                    kw in hint
                    for kw in section_mappings[section_name]
                    if len(kw) > 3
                ):
                    is_section_match = True
            if is_section_match:
                renderer._inject_image_inline(cell, image, max_width=1.3)
                if pending_images is not None:
                    pending_images.remove(image)
                logger.info(
                    "image_placed_fallback",
                    extra={"filename": image["filename"], "section": section_name},
                )

    has_content = False
    if cell.paragraphs:
        for para in cell.paragraphs:
            if para.text.strip() or para.runs:
                for run in para.runs:
                    if run.text and run.text.strip():
                        has_content = True
                        break
                if has_content:
                    break
                if para._element.xpath(".//w:hyperlink"):
                    has_content = True
                    break

    if not has_content and not has_coordinate_hyperlinks and not append_mode:
        empty_paras = []
        for para in cell.paragraphs:
            if not para.text.strip() and not para.runs:
                empty_paras.append(para)
        for para in empty_paras:
            p = para._element
            p.getparent().remove(p)

    if pending_images:
        estimated_col_width = 6.5 / 5
        if estimated_col_width >= settings.IMAGE_INLINE_MIN_COLUMN_WIDTH_INCHES:
            for image in pending_images[:]:
                if current_slot_number is not None:
                    image_slot = image.get("_source_slot")
                    if image_slot is None:
                        logger.warning(
                            "image_missing_slot_metadata",
                            extra={
                                "filename": image.get("filename", "unknown"),
                                "current_slot": current_slot_number,
                            },
                        )
                        continue
                    if image_slot != current_slot_number:
                        continue
                if current_subject is not None:
                    image_subject = image.get("_source_subject")
                    if image_subject is None:
                        logger.warning(
                            "image_missing_subject_metadata",
                            extra={
                                "filename": image.get("filename", "unknown"),
                                "current_subject": current_subject,
                            },
                        )
                        continue
                    if image_subject != current_subject:
                        continue

                if placement_module.try_structure_based_placement(
                    renderer, image, day_name, section_name, col_idx
                ):
                    renderer._inject_image_inline(
                        cell, image, max_width=estimated_col_width
                    )
                    pending_images.remove(image)
                    logger.info(
                        "image_placed_inline",
                        extra={
                            "filename": image["filename"],
                            "cell": f"{day_name}_{section_name}",
                            "confidence": 1.0,
                            "match_type": "structure_based",
                            "slot": current_slot_number,
                            "subject": current_subject,
                        },
                    )
                elif text:
                    confidence, match_type = placement_module.calculate_match_confidence(
                        renderer, text, image, day_name, section_name
                    )
                    if confidence >= settings.MEDIA_MATCH_CONFIDENCE_THRESHOLD:
                        renderer._inject_image_inline(
                            cell, image, max_width=estimated_col_width
                        )
                        pending_images.remove(image)
                        logger.info(
                            "image_placed_inline",
                            extra={
                                "filename": image["filename"],
                                "cell": f"{day_name}_{section_name}",
                                "confidence": confidence,
                                "match_type": match_type,
                                "slot": current_slot_number,
                                "subject": current_subject,
                            },
                        )
                    else:
                        logger.debug(
                            "image_match_rejected",
                            extra={
                                "filename": image["filename"],
                                "cell": f"{day_name}_{section_name}",
                                "confidence": confidence,
                            },
                        )
