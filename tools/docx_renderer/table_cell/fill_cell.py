"""
Fill a single cell with formatted text and inject matched media inline.
"""

import re
from typing import Dict, List

from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run

from .. import logger
from .. import style as style_module
from . import placement as placement_module

try:
    from tools.markdown_to_docx import MarkdownToDocx
except ImportError:
    from markdown_to_docx import MarkdownToDocx


def fill_cell(
    renderer,
    table,
    row_idx: int,
    col_idx: int,
    text: str,
    day_name: str = None,
    section_name: str = None,
    pending_hyperlinks: List[Dict] = None,
    pending_images: List[Dict] = None,
    current_slot_number: int = None,
    current_subject: str = None,
    append_mode: bool = False,
) -> None:
    """
    Fill a cell with formatted text and inject matched media inline.
    """
    from backend.config import settings

    _ = (re, qn)

    cell = table.rows[row_idx].cells[col_idx]

    existing_hyperlinks = cell._element.xpath(".//w:hyperlink")
    has_coordinate_hyperlinks = len(existing_hyperlinks) > 0

    if has_coordinate_hyperlinks and text:
        hyperlinks_to_remove = []
        for hl_elem in existing_hyperlinks:
            try:
                r_id = hl_elem.get(qn("r:id"))
                if r_id and cell.paragraphs:
                    para = cell.paragraphs[0]
                    if r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        link_text = "".join(
                            node.text for node in hl_elem.xpath(".//w:t") if node.text
                        )
                        if link_text and url:
                            markdown_pattern = rf"\[{re.escape(link_text)}\]\({re.escape(url)}\)"
                            if re.search(markdown_pattern, text, re.IGNORECASE):
                                hyperlinks_to_remove.append(hl_elem)
            except Exception:
                pass

        for hl_elem in hyperlinks_to_remove:
            try:
                parent = hl_elem.getparent()
                if parent is not None:
                    parent.remove(hl_elem)
                    try:
                        para_elem = parent.getparent()
                        if para_elem is not None and para_elem.tag == qn("w:p"):
                            runs = para_elem.xpath(".//w:r")
                            has_text = any(run.xpath(".//w:t") for run in runs)
                            if not has_text:
                                para_parent = para_elem.getparent()
                                if para_parent is not None:
                                    para_parent.remove(para_elem)
                    except Exception:
                        pass
            except Exception as e:
                logger.warning(
                    "failed_to_remove_duplicate_hyperlink",
                    extra={
                        "error": str(e),
                        "cell": f"{day_name}_{section_name}"
                        if day_name and section_name
                        else "unknown",
                    },
                )

        existing_hyperlinks = cell._element.xpath(".//w:hyperlink")
        has_coordinate_hyperlinks = len(existing_hyperlinks) > 0
        if hyperlinks_to_remove:
            logger.info(
                "removed_duplicate_coordinate_hyperlinks",
                extra={
                    "removed_count": len(hyperlinks_to_remove),
                    "cell": f"{day_name}_{section_name}"
                    if day_name and section_name
                    else "unknown",
                    "slot": current_slot_number,
                },
            )

    if not has_coordinate_hyperlinks and not append_mode:
        paras_to_remove = list(cell.paragraphs)
        for para in paras_to_remove:
            p = para._element
            p.getparent().remove(p)
        new_para = cell.add_paragraph()
        new_para.paragraph_format.space_after = Pt(0)
        new_para.paragraph_format.space_before = Pt(0)
        new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    matching_hyperlinks = []
    if pending_hyperlinks and text:
        logger.info(
            "hyperlink_filtering_context",
            extra={
                "current_slot_number": current_slot_number,
                "current_subject": current_subject,
                "pending_hyperlinks_count": len(pending_hyperlinks),
                "cell": f"{day_name}_{section_name}"
                if day_name and section_name
                else "unknown",
            },
        )
        try:
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_filtering_context(
                current_slot_number,
                current_subject,
                len(pending_hyperlinks),
                f"{day_name}_{section_name}"
                if day_name and section_name
                else "unknown",
            )
        except Exception:
            pass

        for hyperlink in pending_hyperlinks[:]:
            if current_slot_number is not None:
                link_slot = hyperlink.get("_source_slot")
                if link_slot is None:
                    logger.warning(
                        "hyperlink_missing_slot_metadata",
                        extra={
                            "text": hyperlink["text"][:50],
                            "url": hyperlink["url"],
                            "current_slot": current_slot_number,
                        },
                    )
                    continue
                if link_slot != current_slot_number:
                    try:
                        from tools.diagnostic_logger import get_diagnostic_logger

                        get_diagnostic_logger().log_hyperlink_filtered(
                            hyperlink["text"],
                            link_slot,
                            current_slot_number,
                            "slot_mismatch",
                        )
                    except Exception:
                        pass
                    continue

            if current_subject is not None:
                link_subject = hyperlink.get("_source_subject")
                if link_subject is None:
                    logger.warning(
                        "hyperlink_missing_subject_metadata",
                        extra={
                            "text": hyperlink["text"][:50],
                            "url": hyperlink["url"],
                            "current_subject": current_subject,
                        },
                    )
                    continue
                if link_subject != current_subject:
                    continue

            is_structural_match = False
            match_type = "none"
            confidence = 0.0

            if (
                hyperlink.get("row_idx") is not None
                and hyperlink.get("cell_idx") is not None
            ):
                row_offset = (
                    renderer.structure_metadata.row_offset
                    if renderer.structure_metadata
                    else 0
                )
                target_row = hyperlink["row_idx"] + row_offset
                if target_row == row_idx and hyperlink["cell_idx"] == col_idx:
                    is_structural_match = True
                    match_type = "coordinate"
                    confidence = 1.0

            if (
                not is_structural_match
                and hyperlink.get("row_label")
                and hyperlink.get("day_hint")
            ):
                if renderer.structure_metadata:
                    target_row_idx = renderer.structure_metadata.get_row_index(
                        hyperlink["row_label"]
                    )
                    target_col_idx = renderer.structure_metadata.get_col_index(
                        hyperlink["day_hint"]
                    )
                    if target_row_idx is not None:
                        target_row_idx += renderer.structure_metadata.row_offset
                    if target_row_idx == row_idx and target_col_idx == col_idx:
                        is_structural_match = True
                        match_type = "label_day"
                        confidence = 1.0

            if not is_structural_match:
                confidence, match_type = placement_module.calculate_match_confidence(
                    renderer, text, hyperlink, day_name, section_name
                )

            if confidence >= settings.MEDIA_MATCH_CONFIDENCE_THRESHOLD:
                matching_hyperlinks.append((hyperlink, confidence, match_type))

    if matching_hyperlinks and text:
        sorted_matching = sorted(
            matching_hyperlinks,
            key=lambda m: len(m[0].get("text", "")),
            reverse=True,
        )
        for hyperlink, confidence, match_type in sorted_matching:
            link_text = hyperlink.get("text", "")
            link_url = hyperlink.get("url", "")
            if not link_text or not link_url:
                continue

            hl_day = hyperlink.get("day_hint") or hyperlink.get("day")
            if hl_day and day_name:
                if hl_day.lower().strip() != day_name.lower().strip():
                    continue

            search_pattern_raw = re.escape(link_text)
            search_pattern_flexible = re.sub(r"(\\ )+", r"\\s+", search_pattern_raw)
            search_pattern_final = f"{search_pattern_flexible}[.,;:]?"
            match = re.search(search_pattern_final, text, re.IGNORECASE)

            if match:
                raw_match_text = match.group(0)
                stripped_match_text = raw_match_text.strip()
                total_pattern = rf"(\[\s*{re.escape(stripped_match_text)}\s*\]\([^)]*\))|({search_pattern_final})"
                found_and_replaced = [False]

                def replacement_logic(m):
                    if m.group(1):
                        return m.group(1)
                    if not found_and_replaced[0]:
                        found_and_replaced[0] = True
                        plain_text = m.group(2)
                        leading_ws = plain_text[
                            : plain_text.find(stripped_match_text)
                        ]
                        trailing_ws = plain_text[
                            plain_text.find(stripped_match_text)
                            + len(stripped_match_text) :
                        ]
                        return f"{leading_ws}[{stripped_match_text}]({link_url}){trailing_ws}"
                    return m.group(0)

                new_text = re.sub(
                    total_pattern, replacement_logic, text, flags=re.IGNORECASE
                )
                if found_and_replaced[0]:
                    text = new_text
                    logger.info(
                        "hyperlink_placed_smart_inline",
                        extra={
                            "text": stripped_match_text,
                            "url": link_url,
                            "cell": f"{day_name}_{section_name}",
                            "slot": current_slot_number,
                        },
                    )
                else:
                    logger.info(
                        "hyperlink_already_inline_or_unmatched",
                        extra={
                            "text": stripped_match_text,
                            "cell": f"{day_name}_{section_name}",
                            "slot": current_slot_number,
                        },
                    )
                matching_hyperlinks = [
                    m for m in matching_hyperlinks if m[0] != hyperlink
                ]
                if hyperlink in pending_hyperlinks:
                    pending_hyperlinks.remove(hyperlink)

    if not has_coordinate_hyperlinks and not append_mode:
        if text:
            MarkdownToDocx.add_multiline_text(cell, text)
    else:
        if text:
            lines = text.split("\n")
            for i, line in enumerate(lines):
                if line.strip():
                    if i == 0:
                        new_para = cell.add_paragraph()
                        MarkdownToDocx.add_formatted_text(new_para, line)
                    else:
                        MarkdownToDocx.add_paragraph(cell, line)

    if cell.paragraphs:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in para.runs:
                style_module.force_font_tnr8(
                    run, is_bold=(row_idx == renderer.UNIT_LESSON_ROW)
                )
            for hl in para._p.findall(qn("w:hyperlink")):
                for r_elem in hl.findall(qn("w:r")):
                    hl_run = Run(r_elem, para)
                    is_hl_bold = (
                        row_idx == renderer.UNIT_LESSON_ROW or hl_run.font.bold
                    )
                    style_module.force_font_tnr8(
                        hl_run, is_bold=is_hl_bold, is_hyperlink=True
                    )

    links = pending_hyperlinks or []
    imgs = pending_images or []
    if section_name and (links or imgs):
        for hyperlink in links[:]:
            if current_slot_number is not None:
                link_slot = hyperlink.get("_source_slot")
                if link_slot is not None and link_slot != current_slot_number:
                    continue
            hl_day = hyperlink.get("day_hint")
            if hl_day and day_name:
                if hl_day.lower().strip() != day_name.lower().strip():
                    continue
            hint = (hyperlink.get("section_hint") or "").lower()
            is_section_match = False
            section_mappings = {
                "unit_lesson": ["unit", "lesson", "module"],
                "objective": ["objective", "goal", "swbat"],
                "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                "tailored_instruction": [
                    "instruction",
                    "activity",
                    "procedure",
                    "lesson",
                    "tailored",
                    "differentiation",
                ],
                "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                "homework": ["homework", "assignment", "practice"],
            }
            if hint == section_name:
                is_section_match = True
            elif section_name in section_mappings:
                if hint in section_mappings[section_name]:
                    is_section_match = True
                elif any(
                    kw in hint
                    for kw in section_mappings[section_name]
                    if len(kw) > 3
                ):
                    is_section_match = True
            if is_section_match:
                p = cell.add_paragraph()
                MarkdownToDocx.add_formatted_text(
                    p, f"\u2022 [{hyperlink['text']}]({hyperlink['url']})"
                )
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)
                if pending_hyperlinks is not None:
                    pending_hyperlinks.remove(hyperlink)
                logger.info(
                    "hyperlink_placed_fallback",
                    extra={"text": hyperlink["text"], "section": section_name},
                )

        for image in imgs[:]:
            if current_slot_number is not None:
                img_slot = image.get("_source_slot")
                if img_slot is not None and img_slot != current_slot_number:
                    continue
            hint = (image.get("section_hint") or "").lower()
            is_section_match = False
            section_mappings = {
                "unit_lesson": ["unit", "lesson", "module"],
                "objective": ["objective", "goal", "swbat"],
                "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                "tailored_instruction": [
                    "instruction",
                    "activity",
                    "procedure",
                    "lesson",
                    "tailored",
                    "differentiation",
                ],
                "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                "homework": ["homework", "assignment", "practice"],
            }
            if hint == section_name:
                is_section_match = True
            elif section_name in section_mappings:
                if hint in section_mappings[section_name]:
                    is_section_match = True
                elif any(
                    kw in hint
                    for kw in section_mappings[section_name]
                    if len(kw) > 3
                ):
                    is_section_match = True
            if is_section_match:
                renderer._inject_image_inline(cell, image, max_width=1.3)
                if pending_images is not None:
                    pending_images.remove(image)
                logger.info(
                    "image_placed_fallback",
                    extra={"filename": image["filename"], "section": section_name},
                )

    has_content = False
    if cell.paragraphs:
        for para in cell.paragraphs:
            if para.text.strip() or para.runs:
                for run in para.runs:
                    if run.text and run.text.strip():
                        has_content = True
                        break
                if has_content:
                    break
                if para._element.xpath(".//w:hyperlink"):
                    has_content = True
                    break

    if not has_content and not has_coordinate_hyperlinks and not append_mode:
        empty_paras = []
        for para in cell.paragraphs:
            if not para.text.strip() and not para.runs:
                empty_paras.append(para)
        for para in empty_paras:
            p = para._element
            p.getparent().remove(p)

    if pending_images:
        estimated_col_width = 6.5 / 5
        if estimated_col_width >= settings.IMAGE_INLINE_MIN_COLUMN_WIDTH_INCHES:
            for image in pending_images[:]:
                if current_slot_number is not None:
                    image_slot = image.get("_source_slot")
                    if image_slot is None:
                        logger.warning(
                            "image_missing_slot_metadata",
                            extra={
                                "filename": image.get("filename", "unknown"),
                                "current_slot": current_slot_number,
                            },
                        )
                        continue
                    if image_slot != current_slot_number:
                        continue
                if current_subject is not None:
                    image_subject = image.get("_source_subject")
                    if image_subject is None:
                        logger.warning(
                            "image_missing_subject_metadata",
                            extra={
                                "filename": image.get("filename", "unknown"),
                                "current_subject": current_subject,
                            },
                        )
                        continue
                    if image_subject != current_subject:
                        continue

                if placement_module.try_structure_based_placement(
                    renderer, image, day_name, section_name, col_idx
                ):
                    renderer._inject_image_inline(
                        cell, image, max_width=estimated_col_width
                    )
                    pending_images.remove(image)
                    logger.info(
                        "image_placed_inline",
                        extra={
                            "filename": image["filename"],
                            "cell": f"{day_name}_{section_name}",
                            "confidence": 1.0,
                            "match_type": "structure_based",
                            "slot": current_slot_number,
                            "subject": current_subject,
                        },
                    )
                elif text:
                    confidence, match_type = placement_module.calculate_match_confidence(
                        renderer, text, image, day_name, section_name
                    )
                    if confidence >= settings.MEDIA_MATCH_CONFIDENCE_THRESHOLD:
                        renderer._inject_image_inline(
                            cell, image, max_width=estimated_col_width
                        )
                        pending_images.remove(image)
                        logger.info(
                            "image_placed_inline",
                            extra={
                                "filename": image["filename"],
                                "cell": f"{day_name}_{section_name}",
                                "confidence": confidence,
                                "match_type": match_type,
                                "slot": current_slot_number,
                                "subject": current_subject,
                            },
                        )
                    else:
                        logger.debug(
                            "image_match_rejected",
                            extra={
                                "filename": image["filename"],
                                "cell": f"{day_name}_{section_name}",
                                "confidence": confidence,
                            },
                        )
