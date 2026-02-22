"""Inline injection of hyperlinks and images into table cells."""

import base64
from io import BytesIO
from typing import Dict

from docx.shared import Inches, Pt

from . import logger
from . import hyperlink_placement as _hyperlink_module


def inject_hyperlink_inline(renderer, cell, hyperlink: Dict, row_idx: int = None) -> None:
    """Inject hyperlink into cell on its own line."""
    _hyperlink_module.inject_hyperlink_inline(renderer, cell, hyperlink, row_idx=row_idx)


def inject_image_inline(renderer, cell, image: Dict, max_width: float) -> None:
    """Inject image into cell with width constraints.

    Args:
        renderer: DOCXRenderer instance (unused, for API consistency)
        cell: Cell object
        image: Image dictionary with base64 data
        max_width: Maximum width in inches
    """
    try:
        image_data = base64.b64decode(image["data"])
        image_stream = BytesIO(image_data)

        para = cell.add_paragraph()
        run = para.add_run()
        run.add_picture(
            image_stream, width=Inches(max_width * 0.9)
        )

        caption = cell.add_paragraph()
        caption_run = caption.add_run(f"[{image.get('filename', 'image')}]")
        caption_run.font.size = Pt(7)
        caption_run.italic = True

    except Exception as e:
        logger.warning(
            "inline_image_failed",
            extra={"filename": image.get("filename"), "error": str(e)},
        )
