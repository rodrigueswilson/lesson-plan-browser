"""
Comprehensive analysis of ALL lesson plan table structures.

Scans all .docx files in:
- F:\rodri\Documents\OneDrive\AS\Daniela LP
- F:\rodri\Documents\OneDrive\AS\Lesson Plan

Analyzes:
1. Table dimensions (rows x columns)
2. Row labels and their variations
3. Column headers and their variations
4. Patterns and anomalies
5. Consistency across files

Output: Comprehensive knowledge base for table structure handling
"""

from docx import Document
from pathlib import Path
from collections import defaultdict, Counter
import json

FOLDERS = [
    r'F:\rodri\Documents\OneDrive\AS\Daniela LP',
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan'
]


def find_daily_plans_table(doc):
    """Find the daily plans table (has MONDAY/Monday in it)."""
    for idx, table in enumerate(doc.tables):
        if not table.rows or not table.rows[0].cells:
            continue
        
        # Check if any cell in first row contains day names
        for cell in table.rows[0].cells:
            text = cell.text.upper()
            if any(day in text for day in ['MONDAY', 'TUESDAY', 'WEDNESDAY', 'THURSDAY', 'FRIDAY']):
                return idx, table
    
    return None, None


def extract_table_structure(table):
    """Extract detailed structure information from a table."""
    
    if not table or not table.rows:
        return None
    
    num_rows = len(table.rows)
    num_cols = len(table.rows[0].cells) if table.rows else 0
    
    # Extract row labels (first cell of each row)
    row_labels = []
    for row in table.rows:
        if row.cells:
            label = row.cells[0].text.strip()
            row_labels.append(label)
        else:
            row_labels.append("")
    
    # Extract column headers (first row)
    col_headers = []
    if table.rows:
        for cell in table.rows[0].cells:
            header = cell.text.strip()
            col_headers.append(header)
    
    return {
        'dimensions': f"{num_rows}x{num_cols}",
        'num_rows': num_rows,
        'num_cols': num_cols,
        'row_labels': row_labels,
        'col_headers': col_headers
    }


def analyze_file(file_path):
    """Analyze a single lesson plan file."""
    
    try:
        doc = Document(file_path)
        table_idx, table = find_daily_plans_table(doc)
        
        if table is None:
            return {
                'status': 'no_daily_table',
                'total_tables': len(doc.tables)
            }
        
        structure = extract_table_structure(table)
        structure['status'] = 'success'
        structure['table_index'] = table_idx
        structure['total_tables'] = len(doc.tables)
        
        return structure
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }


def scan_all_files():
    """Scan all lesson plan files."""
    
    print("="*100)
    print("SCANNING ALL LESSON PLAN FILES")
    print("="*100)
    print()
    
    all_results = []
    
    for folder in FOLDERS:
        folder_path = Path(folder)
        
        if not folder_path.exists():
            print(f"⚠️  Folder not found: {folder}")
            continue
        
        print(f"\nScanning: {folder}")
        print("-" * 80)
        
        # Find all .docx files recursively
        docx_files = list(folder_path.rglob('*.docx'))
        
        # Filter out temp files
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]
        
        print(f"Found {len(docx_files)} .docx files\n")
        
        for file_path in docx_files:
            rel_path = file_path.relative_to(folder_path)
            print(f"  Analyzing: {rel_path}")
            
            result = analyze_file(file_path)
            result['file_path'] = str(file_path)
            result['rel_path'] = str(rel_path)
            result['folder'] = folder
            
            all_results.append(result)
            
            # Show brief status
            if result['status'] == 'success':
                print(f"    ✓ {result['dimensions']} - {len(result['row_labels'])} row labels")
            elif result['status'] == 'no_daily_table':
                print(f"    ⚠️  No daily plans table found ({result['total_tables']} tables total)")
            else:
                print(f"    ✗ Error: {result.get('error', 'Unknown')}")
    
    return all_results


def analyze_patterns(results):
    """Analyze patterns across all files."""
    
    print("\n\n" + "="*100)
    print("PATTERN ANALYSIS")
    print("="*100)
    print()
    
    # Filter successful results
    successful = [r for r in results if r['status'] == 'success']
    
    print(f"Successfully analyzed: {len(successful)}/{len(results)} files")
    print(f"Failed/No table: {len(results) - len(successful)} files\n")
    
    if not successful:
        print("No successful analyses to compare")
        return
    
    # Analyze dimensions
    print("\n" + "-"*80)
    print("TABLE DIMENSIONS")
    print("-"*80)
    
    dimension_counts = Counter(r['dimensions'] for r in successful)
    for dim, count in dimension_counts.most_common():
        pct = count / len(successful) * 100
        print(f"  {dim}: {count} files ({pct:.1f}%)")
    
    # Analyze row labels
    print("\n" + "-"*80)
    print("ROW LABELS (Unique variations)")
    print("-"*80)
    
    # Collect all unique row label sequences
    row_sequences = defaultdict(list)
    for r in successful:
        sequence = tuple(r['row_labels'])
        row_sequences[sequence].append(r['rel_path'])
    
    print(f"\nFound {len(row_sequences)} unique row label sequences:\n")
    
    for idx, (sequence, files) in enumerate(sorted(row_sequences.items(), key=lambda x: len(x[1]), reverse=True), 1):
        print(f"Pattern {idx}: ({len(files)} files)")
        for i, label in enumerate(sequence):
            print(f"  Row {i}: {label[:60]}")
        print(f"  Example files:")
        for f in files[:3]:
            print(f"    - {f}")
        if len(files) > 3:
            print(f"    ... and {len(files) - 3} more")
        print()
    
    # Analyze column headers
    print("\n" + "-"*80)
    print("COLUMN HEADERS (Unique variations)")
    print("-"*80)
    
    col_sequences = defaultdict(list)
    for r in successful:
        sequence = tuple(r['col_headers'])
        col_sequences[sequence].append(r['rel_path'])
    
    print(f"\nFound {len(col_sequences)} unique column header sequences:\n")
    
    for idx, (sequence, files) in enumerate(sorted(col_sequences.items(), key=lambda x: len(x[1]), reverse=True), 1):
        print(f"Pattern {idx}: ({len(files)} files)")
        print(f"  Columns: {' | '.join(sequence)}")
        print(f"  Example files:")
        for f in files[:3]:
            print(f"    - {f}")
        if len(files) > 3:
            print(f"    ... and {len(files) - 3} more")
        print()
    
    # Identify anomalies
    print("\n" + "-"*80)
    print("ANOMALIES & VARIATIONS")
    print("-"*80)
    print()
    
    # Most common dimension
    most_common_dim = dimension_counts.most_common(1)[0][0]
    anomalies = [r for r in successful if r['dimensions'] != most_common_dim]
    
    if anomalies:
        print(f"Files with non-standard dimensions (not {most_common_dim}):")
        for r in anomalies:
            print(f"  - {r['rel_path']}")
            print(f"    Dimensions: {r['dimensions']}")
            print(f"    Row labels: {r['row_labels'][:3]}...")
            print()
    else:
        print(f"✓ All files have standard dimensions: {most_common_dim}")
    
    # Check for missing row labels
    print("\n" + "-"*80)
    print("MISSING OR EMPTY ROW LABELS")
    print("-"*80)
    print()
    
    files_with_empty = []
    for r in successful:
        empty_rows = [i for i, label in enumerate(r['row_labels']) if not label]
        if empty_rows:
            files_with_empty.append((r['rel_path'], empty_rows))
    
    if files_with_empty:
        print(f"Found {len(files_with_empty)} files with empty row labels:")
        for file_path, empty_rows in files_with_empty[:10]:
            print(f"  - {file_path}")
            print(f"    Empty rows: {empty_rows}")
    else:
        print("✓ All files have complete row labels")


def generate_knowledge_base(results):
    """Generate a structured knowledge base."""
    
    print("\n\n" + "="*100)
    print("KNOWLEDGE BASE GENERATION")
    print("="*100)
    print()
    
    successful = [r for r in results if r['status'] == 'success']
    
    if not successful:
        print("No data to generate knowledge base")
        return None
    
    # Find the most common structure
    dimension_counts = Counter(r['dimensions'] for r in successful)
    most_common_dim = dimension_counts.most_common(1)[0][0]
    
    # Find the most common row sequence
    row_sequences = Counter(tuple(r['row_labels']) for r in successful)
    most_common_rows = row_sequences.most_common(1)[0][0]
    
    # Find the most common column sequence
    col_sequences = Counter(tuple(r['col_headers']) for r in successful)
    most_common_cols = col_sequences.most_common(1)[0][0]
    
    knowledge_base = {
        'summary': {
            'total_files_analyzed': len(results),
            'successful_analyses': len(successful),
            'failed_analyses': len(results) - len(successful)
        },
        'standard_structure': {
            'dimensions': most_common_dim,
            'row_labels': list(most_common_rows),
            'column_headers': list(most_common_cols),
            'prevalence': f"{dimension_counts[most_common_dim]}/{len(successful)} files"
        },
        'variations': {
            'dimension_variations': len(dimension_counts),
            'row_sequence_variations': len(row_sequences),
            'column_sequence_variations': len(col_sequences)
        },
        'all_unique_row_labels': sorted(set(label for r in successful for label in r['row_labels'] if label)),
        'all_unique_column_headers': sorted(set(header for r in successful for header in r['col_headers'] if header)),
        'dimension_distribution': dict(dimension_counts),
        'recommendations': []
    }
    
    # Generate recommendations
    if len(dimension_counts) == 1:
        knowledge_base['recommendations'].append({
            'type': 'structure_consistency',
            'level': 'high',
            'message': 'All files have identical dimensions - coordinate-based placement is highly reliable'
        })
    else:
        knowledge_base['recommendations'].append({
            'type': 'structure_consistency',
            'level': 'medium',
            'message': f'Found {len(dimension_counts)} different dimension patterns - need dimension-aware placement logic'
        })
    
    if len(row_sequences) == 1:
        knowledge_base['recommendations'].append({
            'type': 'row_consistency',
            'level': 'high',
            'message': 'All files have identical row labels - row_index is perfectly reliable'
        })
    else:
        knowledge_base['recommendations'].append({
            'type': 'row_consistency',
            'level': 'medium',
            'message': f'Found {len(row_sequences)} different row patterns - use row_label + row_index for matching'
        })
    
    if len(col_sequences) <= 2:
        knowledge_base['recommendations'].append({
            'type': 'column_consistency',
            'level': 'high',
            'message': 'Column headers are very consistent - column_index is reliable'
        })
    
    # Save to JSON
    output_path = Path('d:/LP/docs/knowledge/table_structure_knowledge_base.json')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(knowledge_base, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Knowledge base saved to: {output_path}")
    
    # Print summary
    print("\n" + "-"*80)
    print("KNOWLEDGE BASE SUMMARY")
    print("-"*80)
    print()
    
    print(f"Standard Structure:")
    print(f"  Dimensions: {knowledge_base['standard_structure']['dimensions']}")
    print(f"  Prevalence: {knowledge_base['standard_structure']['prevalence']}")
    print(f"\n  Standard row labels:")
    for i, label in enumerate(knowledge_base['standard_structure']['row_labels']):
        print(f"    Row {i}: {label}")
    print(f"\n  Standard column headers:")
    for i, header in enumerate(knowledge_base['standard_structure']['column_headers']):
        print(f"    Col {i}: {header}")
    
    print(f"\n\nVariations:")
    print(f"  Dimension patterns: {knowledge_base['variations']['dimension_variations']}")
    print(f"  Row sequence patterns: {knowledge_base['variations']['row_sequence_variations']}")
    print(f"  Column sequence patterns: {knowledge_base['variations']['column_sequence_variations']}")
    
    print(f"\n\nRecommendations:")
    for rec in knowledge_base['recommendations']:
        level_icon = "✓" if rec['level'] == 'high' else "⚠️"
        print(f"  {level_icon} [{rec['type']}] {rec['message']}")
    
    return knowledge_base


def main():
    print("="*100)
    print("COMPREHENSIVE TABLE STRUCTURE ANALYSIS")
    print("="*100)
    print("\nAnalyzing ALL lesson plan files to build complete knowledge base...")
    print()
    
    # Scan all files
    results = scan_all_files()
    
    # Analyze patterns
    analyze_patterns(results)
    
    # Generate knowledge base
    knowledge_base = generate_knowledge_base(results)
    
    # Save detailed results
    results_path = Path('d:/LP/docs/knowledge/table_structure_analysis_results.json')
    results_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(results_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"\n\n✓ Detailed results saved to: {results_path}")
    
    print("\n\n" + "="*100)
    print("ANALYSIS COMPLETE")
    print("="*100)
    print("\nKnowledge base files created:")
    print("  1. table_structure_knowledge_base.json - Structured knowledge for code")
    print("  2. table_structure_analysis_results.json - Detailed analysis data")
    print("\nUse these files to improve:")
    print("  - Parser logic (row/column recognition)")
    print("  - Renderer placement (coordinate-based)")
    print("  - Validation (structure checking)")
    print("  - Error handling (anomaly detection)")


if __name__ == '__main__':
    main()
