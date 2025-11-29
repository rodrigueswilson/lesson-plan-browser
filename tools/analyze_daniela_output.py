"""
Analyze Daniela's W43 lesson plan output to validate coordinate-based placement.

Compares:
1. Input files vs output file
2. Hyperlink locations (input vs output)
3. Placement accuracy
4. Inline vs fallback rates
"""

import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_parser import DOCXParser

# Files
INPUT_FILES = [
    r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Morais 10_20_25 - 10_24_25.docx',
    r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Mrs. Grande Science 10_20- 10_24.docx'
]

OUTPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Daniela_Silva_Weekly_W43_10-20-10-24_20251019_150909.docx'


def analyze_hyperlinks_in_file(file_path, file_label):
    """Analyze hyperlinks in a file."""
    
    print(f"\n{'='*80}")
    print(f"ANALYZING: {file_label}")
    print(f"{'='*80}\n")
    
    if not Path(file_path).exists():
        print(f"⚠️  File not found: {file_path}")
        return None
    
    try:
        parser = DOCXParser(file_path)
        hyperlinks = parser.extract_hyperlinks()
        
        print(f"Total hyperlinks: {len(hyperlinks)}")
        
        # Schema version
        v2_links = [h for h in hyperlinks if h.get('schema_version') == '2.0']
        print(f"Schema v2.0: {len(v2_links)}/{len(hyperlinks)}")
        
        # Table vs paragraph
        table_links = [h for h in hyperlinks if h.get('table_idx') is not None]
        para_links = [h for h in hyperlinks if h.get('table_idx') is None]
        
        print(f"Table links: {len(table_links)}")
        print(f"Paragraph links: {len(para_links)}")
        print()
        
        # Show sample locations
        print("Sample hyperlink locations:")
        print("-" * 80)
        
        for i, link in enumerate(hyperlinks[:5], 1):
            print(f"{i}. '{link['text'][:40]}'")
            if link.get('table_idx') is not None:
                print(f"   Location: Table {link['table_idx']}, Row {link['row_idx']}, Cell {link['cell_idx']}")
                print(f"   Row label: '{link.get('row_label', 'N/A')[:40]}'")
                print(f"   Col header: '{link.get('col_header', 'N/A')}'")
            else:
                print(f"   Location: Paragraph (non-table)")
            print()
        
        return {
            'file': file_label,
            'total': len(hyperlinks),
            'table_links': len(table_links),
            'para_links': len(para_links),
            'hyperlinks': hyperlinks
        }
        
    except Exception as e:
        print(f"❌ Error analyzing file: {e}")
        import traceback
        traceback.print_exc()
        return None


def check_for_referenced_links_section(file_path):
    """Check if output has a 'Referenced Links' section (fallback)."""
    
    print(f"\n{'='*80}")
    print(f"CHECKING FOR FALLBACK SECTION")
    print(f"{'='*80}\n")
    
    try:
        doc = Document(file_path)
        
        # Look for "Referenced Links" heading
        has_fallback = False
        fallback_count = 0
        
        for para in doc.paragraphs:
            if "Referenced Links" in para.text or "referenced links" in para.text.lower():
                has_fallback = True
                print(f"⚠️  Found 'Referenced Links' section")
                
                # Count links in fallback section
                # (They would be in subsequent paragraphs)
                break
        
        if not has_fallback:
            print(f"✅ No 'Referenced Links' section found")
            print(f"   All hyperlinks were placed inline!")
        
        return has_fallback
        
    except Exception as e:
        print(f"❌ Error checking fallback: {e}")
        return None


def analyze_output_table_structure(file_path):
    """Analyze the output file's table structure."""
    
    print(f"\n{'='*80}")
    print(f"OUTPUT TABLE STRUCTURE")
    print(f"{'='*80}\n")
    
    try:
        doc = Document(file_path)
        
        print(f"Total tables: {len(doc.tables)}")
        print()
        
        # Analyze main lesson plan table (usually table 1)
        if len(doc.tables) >= 2:
            table = doc.tables[1]  # Daily plans table
            
            print(f"Daily Plans Table:")
            print(f"  Dimensions: {len(table.rows)} rows x {len(table.rows[0].cells) if table.rows else 0} cols")
            print()
            
            # Show row labels
            print(f"  Row labels:")
            for i, row in enumerate(table.rows[:8]):  # First 8 rows
                label = row.cells[0].text.strip()[:50] if row.cells else ""
                print(f"    Row {i}: '{label}'")
            print()
            
            # Count hyperlinks in table
            link_count = 0
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for hyperlink in para._element.xpath('.//w:hyperlink'):
                            link_count += 1
            
            print(f"  Hyperlinks in table: {link_count}")
        
    except Exception as e:
        print(f"❌ Error analyzing structure: {e}")


def compare_input_output():
    """Compare input files with output file."""
    
    print("="*80)
    print("DANIELA W43 LESSON PLAN ANALYSIS")
    print("Coordinate-Based Placement Validation")
    print("="*80)
    
    # Analyze input files
    input_results = []
    total_input_links = 0
    
    for i, input_file in enumerate(INPUT_FILES, 1):
        result = analyze_hyperlinks_in_file(input_file, f"INPUT {i}: {Path(input_file).name}")
        if result:
            input_results.append(result)
            total_input_links += result['total']
    
    # Analyze output file
    output_result = analyze_hyperlinks_in_file(OUTPUT_FILE, f"OUTPUT: {Path(OUTPUT_FILE).name}")
    
    # Check for fallback section
    has_fallback = check_for_referenced_links_section(OUTPUT_FILE)
    
    # Analyze output structure
    analyze_output_table_structure(OUTPUT_FILE)
    
    # Summary comparison
    print(f"\n{'='*80}")
    print(f"SUMMARY COMPARISON")
    print(f"{'='*80}\n")
    
    print(f"Input files: {len(INPUT_FILES)}")
    print(f"Total input hyperlinks: {total_input_links}")
    print()
    
    if output_result:
        print(f"Output hyperlinks: {output_result['total']}")
        print(f"  Table links: {output_result['table_links']}")
        print(f"  Paragraph links: {output_result['para_links']}")
        print()
        
        # Placement assessment
        if has_fallback:
            print(f"⚠️  Placement: Some links in fallback section")
            print(f"   Inline rate: <100%")
        else:
            print(f"✅ Placement: All links inline (no fallback section)")
            print(f"   Inline rate: 100%")
        print()
        
        # Expected vs actual
        print(f"Expected hyperlinks: {total_input_links}")
        print(f"Actual hyperlinks: {output_result['total']}")
        
        if output_result['total'] == total_input_links:
            print(f"✅ All hyperlinks preserved")
        else:
            diff = output_result['total'] - total_input_links
            print(f"⚠️  Difference: {diff:+d} hyperlinks")
    
    print()
    print("="*80)
    print("ANALYSIS COMPLETE")
    print("="*80)
    print()
    print("Next steps:")
    print("  1. Open output file in Word")
    print("  2. Verify hyperlinks are in correct cells")
    print("  3. Check that no 'Referenced Links' section exists")
    print("  4. Compare with input files for accuracy")


if __name__ == '__main__':
    compare_input_output()
