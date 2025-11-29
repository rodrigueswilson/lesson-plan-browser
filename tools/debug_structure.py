"""Debug script to see actual document structure."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)

print(f"Total tables: {len(doc.tables)}\n")

for idx, table in enumerate(doc.tables):
    print(f"\n{'='*80}")
    print(f"TABLE {idx}")
    print(f"{'='*80}")
    
    # Show first 3 rows
    for row_idx in range(min(3, len(table.rows))):
        print(f"\nRow {row_idx}:")
        for cell_idx in range(min(5, len(table.rows[row_idx].cells))):
            cell_text = table.rows[row_idx].cells[cell_idx].text.strip()
            if cell_text:
                print(f"  Cell {cell_idx}: {cell_text[:80]}")
    
    # Count links
    link_count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                link_count += len(hyperlinks)
    
    print(f"\nTotal links in this table: {link_count}")
    
    # Check for "Referenced Links"
    has_ref = False
    for row in table.rows:
        for cell in row.cells:
            if 'Referenced Links' in cell.text or 'Referenced Media' in cell.text:
                has_ref = True
                break
    
    if has_ref:
        print("⚠️  THIS IS A REFERENCED LINKS SECTION")
