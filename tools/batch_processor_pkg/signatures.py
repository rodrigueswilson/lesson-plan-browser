"""
Signature box and table helpers for batch processor DOCX output.
Remove, modify, or add signature tables and images. Extracted from orchestrator.
"""

import re
import traceback
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

from backend.telemetry import logger

from tools.batch_processor_pkg.signature_paragraph_helpers import (
    add_date_section_to_paragraph,
    get_paragraph_font_info,
)


def remove_signature_boxes(doc: Document) -> None:
    """Remove signature boxes/tables from document.

    Args:
        doc: Document to remove signatures from
    """
    tables_to_remove = []
    for table in doc.tables:
        table_text = "\n".join(
            [cell.text for row in table.rows for cell in row.cells]
        )
        if any(
            keyword in table_text
            for keyword in [
                "Required Signatures",
                "Teacher Signature:",
                "Administrator Signature:",
                "Administrative Feedback:",
                "I certify that these lessons",
            ]
        ):
            tables_to_remove.append(table)

    for table in tables_to_remove:
        tbl_element = table._element
        tbl_element.getparent().remove(tbl_element)

    paragraphs_to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Signature:") or text.startswith("Date:"):
            paragraphs_to_remove.append(para)

    for para in paragraphs_to_remove:
        p_element = para._element
        p_element.getparent().remove(p_element)


def modify_existing_signature_table(
    doc: Document,
    date: datetime,
    signature_image_path: Optional[str] = None,
    user_name: Optional[str] = None,
) -> None:
    """Modify existing signature table in document (for single-slot documents)."""
    signature_table = None
    for table in doc.tables:
        table_text = "\n".join(
            [cell.text for row in table.rows for cell in row.cells]
        )
        if "Required Signatures" in table_text:
            signature_table = table
            break

    if not signature_table:
        logger.warning("signature_table_not_found_in_doc")
        return

    date_formatted = date.strftime("%m/%d/%Y")

    for row in signature_table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if "Teacher Signature:" in cell_text:
                target_para = None
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if "Teacher Signature:" in para_text:
                        teacher_pos = para_text.find("Teacher Signature:")
                        date_pos = (
                            para_text.find("Date:", teacher_pos)
                            if teacher_pos >= 0
                            else -1
                        )
                        if date_pos > teacher_pos:
                            target_para = para
                            break

                if target_para:
                    para = target_para
                    para_text = para.text
                    teacher_pos = para_text.find("Teacher Signature:")
                    date_pos = para_text.find("Date:", teacher_pos)

                    if date_pos > teacher_pos:
                        date_run = None
                        original_font_size = None
                        original_font_name = None

                        current_pos = 0
                        for run in para.runs:
                            run_text = run.text
                            run_start = current_pos
                            run_end = current_pos + len(run_text)

                            if "Date:" in run_text and run_start >= teacher_pos:
                                date_run = run
                                if run.font.size:
                                    original_font_size = run.font.size
                                if run.font.name:
                                    original_font_name = run.font.name
                                break

                            current_pos = run_end

                        if date_run:
                            run_text = date_run.text
                            if "Date:" in run_text:
                                para.clear()

                                DATE_TAB_POSITION = Inches(5.5)
                                para.paragraph_format.tab_stops.add_tab_stop(
                                    DATE_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT
                                )

                                teacher_run = para.add_run("Teacher Signature: ")
                                teacher_run.font.bold = True
                                if original_font_size:
                                    teacher_run.font.size = original_font_size
                                if original_font_name:
                                    teacher_run.font.name = original_font_name

                                placeholder_run = para.add_run("_" * 49)
                                if original_font_size:
                                    placeholder_run.font.size = original_font_size
                                if original_font_name:
                                    placeholder_run.font.name = original_font_name

                                para.add_run("\t")

                                date_label_run = para.add_run("Date: ")
                                date_label_run.font.bold = True
                                if original_font_size:
                                    date_label_run.font.size = original_font_size
                                if original_font_name:
                                    date_label_run.font.name = original_font_name

                                date_value_run = para.add_run(date_formatted)
                                date_value_run.font.underline = True
                                if original_font_size:
                                    date_value_run.font.size = original_font_size
                                if original_font_name:
                                    date_value_run.font.name = original_font_name

                                logger.info("signature_date_updated_with_tab_stop")
                                break
                    break

    if (
        signature_image_path
        and signature_image_path.strip()
        and Path(signature_image_path).exists()
    ):
        add_signature_image_to_table(signature_table, signature_image_path)
    elif user_name:
        add_user_name_to_table(signature_table, user_name)


def add_signature_image_to_table(
    table, signature_image_path: str  # noqa: ANN001
) -> None:
    """Add signature image to the signature table."""
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if "Teacher Signature:" in cell_text:
                for para in cell.paragraphs:
                    para_text = para.text
                    if "Teacher Signature:" in para_text:
                        text_font_size = Pt(11)
                        for run in para.runs:
                            if run.font.size:
                                text_font_size = run.font.size
                                break

                        text_height_inches = text_font_size.pt / 72.0
                        image_height_inches = text_height_inches * 1.4

                        teacher_pos = para_text.find("Teacher Signature:")
                        date_pos = (
                            para_text.find("Date:", teacher_pos)
                            if teacher_pos >= 0
                            else -1
                        )

                        if teacher_pos >= 0 and date_pos > teacher_pos:
                            try:
                                before_teacher = para_text[:teacher_pos]

                                between_text = para_text[
                                    teacher_pos
                                    + len("Teacher Signature:") : date_pos
                                ].strip()
                                between_text = between_text.lstrip("_").strip()

                                date_updated = False
                                date_value = None
                                date_pattern = r"Date:\s*(\d{1,2}/\d{1,2}/\d{4})"
                                date_match = re.search(date_pattern, para_text)
                                if date_match:
                                    date_updated = True
                                    date_value = date_match.group(1)

                                original_font_size, original_font_name = (
                                    get_paragraph_font_info(para)
                                )

                                para.clear()

                                if before_teacher.strip():
                                    before_run = para.add_run(before_teacher)
                                    if original_font_size:
                                        before_run.font.size = original_font_size
                                    if original_font_name:
                                        before_run.font.name = original_font_name

                                teacher_run = para.add_run("Teacher Signature: ")
                                teacher_run.font.bold = True
                                if original_font_size:
                                    teacher_run.font.size = original_font_size
                                if original_font_name:
                                    teacher_run.font.name = original_font_name

                                image_run = para.add_run()
                                try:
                                    if not Path(signature_image_path).exists():
                                        logger.error(
                                            "signature_image_file_not_found",
                                            extra={"path": signature_image_path},
                                        )
                                        raise FileNotFoundError(
                                            f"Signature image not found: {signature_image_path}"
                                        )

                                    with open(
                                        signature_image_path, "rb"
                                    ) as img_file:
                                        img_stream = BytesIO(img_file.read())
                                        image_run.add_picture(
                                            img_stream,
                                            height=Inches(image_height_inches),
                                        )

                                    if not image_run._element.xpath(".//a:blip"):
                                        logger.error(
                                            "signature_image_not_inserted",
                                            extra={"path": signature_image_path},
                                        )
                                except Exception as img_error:
                                    logger.error(
                                        "signature_image_insertion_error",
                                        extra={
                                            "path": signature_image_path,
                                            "error": str(img_error),
                                        },
                                    )
                                    logger.debug(
                                        "signature_image_error_traceback",
                                        extra={"traceback": traceback.format_exc()},
                                    )

                                if between_text:
                                    between_run = para.add_run(between_text)
                                    if original_font_size:
                                        between_run.font.size = original_font_size
                                    if original_font_name:
                                        between_run.font.name = original_font_name

                                add_date_section_to_paragraph(
                                    para,
                                    date_value if date_updated else None,
                                    original_font_size,
                                    original_font_name,
                                )

                                logger.info(
                                    "signature_image_added",
                                    extra={"path": signature_image_path},
                                )
                                final_text = para.text
                                if "Teacher Signature:" not in final_text:
                                    logger.error(
                                        "signature_text_missing_after_insertion",
                                        extra={
                                            "path": signature_image_path,
                                            "final_text": final_text,
                                        },
                                    )
                            except Exception as e:
                                logger.warning(
                                    "signature_image_failed",
                                    extra={
                                        "path": signature_image_path,
                                        "error": str(e),
                                    },
                                )
                                logger.debug(
                                    "signature_image_traceback",
                                    extra={"traceback": traceback.format_exc()},
                                )
                            break
                break


def add_user_name_to_table(
    table, user_name: str  # noqa: ANN001
) -> None:
    """Add user name (underlined) to the signature table."""
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if "Teacher Signature:" in cell_text:
                for para in cell.paragraphs:
                    para_text = para.text
                    if "Teacher Signature:" in para_text:
                        teacher_pos = para_text.find("Teacher Signature:")
                        date_pos = (
                            para_text.find("Date:", teacher_pos)
                            if teacher_pos >= 0
                            else -1
                        )

                        if teacher_pos >= 0 and date_pos > teacher_pos:
                            before_teacher = para_text[:teacher_pos]

                            between_text = para_text[
                                teacher_pos + len("Teacher Signature:") : date_pos
                            ].strip()
                            between_text = between_text.lstrip("_").strip()

                            date_updated = False
                            date_value = None
                            date_pattern = r"Date:\s*(\d{1,2}/\d{1,2}/\d{4})"
                            date_match = re.search(date_pattern, para_text)
                            if date_match:
                                date_updated = True
                                date_value = date_match.group(1)

                            original_font_size, original_font_name = (
                                get_paragraph_font_info(para)
                            )

                            para.clear()

                            if before_teacher.strip():
                                before_run = para.add_run(before_teacher)
                                if original_font_size:
                                    before_run.font.size = original_font_size
                                if original_font_name:
                                    before_run.font.name = original_font_name

                            teacher_run = para.add_run("Teacher Signature: ")
                            teacher_run.font.bold = True
                            if original_font_size:
                                teacher_run.font.size = original_font_size
                            if original_font_name:
                                teacher_run.font.name = original_font_name

                            name_run = para.add_run(user_name)
                            name_run.font.underline = True
                            if original_font_size:
                                name_run.font.size = original_font_size
                            if original_font_name:
                                name_run.font.name = original_font_name

                            if between_text:
                                between_run = para.add_run(between_text)
                                if original_font_size:
                                    between_run.font.size = original_font_size
                                if original_font_name:
                                    between_run.font.name = original_font_name

                            add_date_section_to_paragraph(
                                para,
                                date_value if date_updated else None,
                                original_font_size,
                                original_font_name,
                            )

                            final_text = para.text
                            if "Teacher Signature:" not in final_text:
                                logger.error(
                                    "signature_text_missing_after_name_insertion",
                                    extra={
                                        "user_name": user_name,
                                        "final_text": final_text,
                                    },
                                )
                            break
                break


def add_signature_box(
    doc: Document,
    date: datetime,
    template_path: str,
    signature_image_path: Optional[str] = None,
    user_name: Optional[str] = None,
) -> None:
    """Add signature box from template to the end of document."""
    template_doc = Document(template_path)

    signature_table = None
    for table in template_doc.tables:
        table_text = "\n".join(
            [cell.text for row in table.rows for cell in row.cells]
        )
        if "Required Signatures" in table_text:
            signature_table = table
            break

    if signature_table:
        doc.add_paragraph()

        new_table = deepcopy(signature_table._element)
        doc._element.body.append(new_table)

        added_table = doc.tables[-1]

        date_formatted = date.strftime("%m/%d/%Y")

        for row in added_table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if "Teacher Signature:" in cell_text:
                    for para in cell.paragraphs:
                        para_text = para.text
                        if (
                            "Teacher Signature:" in para_text
                            and "Date:" in para_text
                        ):
                            teacher_pos = para_text.find("Teacher Signature:")
                            date_pos = para_text.find(
                                "Date:", teacher_pos
                            )

                            if date_pos > teacher_pos:
                                date_run = None
                                original_font_size = None
                                original_font_name = None

                                current_pos = 0
                                for run in para.runs:
                                    run_text = run.text
                                    run_start = current_pos
                                    run_end = current_pos + len(run_text)

                                    if (
                                        "Date:" in run_text
                                        and run_start >= teacher_pos
                                    ):
                                        date_run = run
                                        if run.font.size:
                                            original_font_size = run.font.size
                                        if run.font.name:
                                            original_font_name = run.font.name
                                        break

                                    current_pos = run_end

                                if date_run:
                                    run_text = date_run.text
                                    if "Date:" in run_text:
                                        para.clear()
                                        before_date_text = para_text[:date_pos]
                                        if before_date_text.strip():
                                            before_run = para.add_run(
                                                before_date_text
                                            )
                                            if original_font_size:
                                                before_run.font.size = (
                                                    original_font_size
                                                )
                                            if original_font_name:
                                                before_run.font.name = (
                                                    original_font_name
                                                )

                                        spacing_text = " " * 90
                                        date_label_run = para.add_run(
                                            spacing_text + "Date: "
                                        )
                                        date_label_run.font.bold = True
                                        date_value_run = para.add_run(
                                            date_formatted
                                        )
                                        date_value_run.font.underline = True
                                        if original_font_size:
                                            date_label_run.font.size = (
                                                original_font_size
                                            )
                                            date_value_run.font.size = (
                                                original_font_size
                                            )
                                        if original_font_name:
                                            date_label_run.font.name = (
                                                original_font_name
                                            )
                                            date_value_run.font.name = (
                                                original_font_name
                                            )

                                        after_date_text = para_text[
                                            date_pos + len("Date: ") :
                                        ]
                                        after_date_text = after_date_text.lstrip(
                                            "_"
                                        ).strip()
                                        if after_date_text:
                                            after_run = para.add_run(
                                                after_date_text
                                            )
                                            if original_font_size:
                                                after_run.font.size = (
                                                    original_font_size
                                                )
                                            if original_font_name:
                                                after_run.font.name = (
                                                    original_font_name
                                                )
                                    break
                    break

        if (
            signature_image_path
            and signature_image_path.strip()
            and Path(signature_image_path).exists()
        ):
            add_signature_image_to_table(
                added_table, signature_image_path
            )
        elif user_name:
            add_user_name_to_table(added_table, user_name)
    else:
        doc.add_paragraph()
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sig_run = sig_para.add_run("Signature: _" + "_" * 50)
        sig_run.font.size = Pt(11)
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {date.strftime('%m/%d/%Y')}")
        date_run.font.size = Pt(11)
