"""
Combine lessons into a single DOCX, merge DOCX files, and convert/reconstruct lesson JSON.
Extracted from orchestrator for single responsibility and testability.
"""

import copy
import re
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from docx import Document

from backend.schema import OriginalLessonPlan
from backend.telemetry import logger


def merge_docx_files(
    file_paths: List[str],
    output_path: str,
    master_template_path: Optional[str] = None,
) -> None:
    """
    Merge multiple DOCX files into one using docxcompose.
    Each document is separated by a page break.

    Args:
        file_paths: List of DOCX file paths to merge
        output_path: Path for merged output file
        master_template_path: Optional path to a template to use as the base
    """
    from docxcompose.composer import Composer

    if not file_paths:
        raise ValueError("No files to merge")

    logger.debug("batch_merge_docx_start", extra={"file_count": len(file_paths)})

    if master_template_path:
        master = Document(master_template_path)
        for p in list(master.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(master.tables):
            t._element.getparent().remove(t._element)
        composer = Composer(master)
        start_idx = 0
    else:
        master = Document(file_paths[0])
        composer = Composer(master)
        start_idx = 1

    for i, file_path in enumerate(file_paths[start_idx:], start_idx + 1):
        logger.debug(
            "batch_merge_docx_append",
            extra={"index": i, "file_name": Path(file_path).name},
        )
        doc = Document(file_path)

        if i > 1:
            page_break_para = doc.add_page_break()
            para_element = page_break_para._element
            para_element.getparent().remove(para_element)
            doc._element.body.insert(0, para_element)

        composer.append(doc)

    logger.debug(
        "batch_merge_docx_save", extra={"output_file": Path(output_path).name}
    )
    composer.save(output_path)


def convert_originals_to_json(
    plans: List[OriginalLessonPlan],
) -> Dict[str, Any]:
    """
    Convert a list of OriginalLessonPlan objects to a multi-slot lesson JSON.
    This allows reusing the standard DOCXRenderer for original plans.
    """
    if not plans:
        return {}

    first = plans[0]
    is_single_plan = len(plans) == 1

    metadata = {
        "primary_teacher_name": first.primary_teacher_name,
        "week_of": first.week_of,
        "grade": first.grade or "Unknown",
        "subject": first.subject,
        "slot_number": first.slot_number if is_single_plan else None,
        "homeroom": first.homeroom,
        "source_type": "originals",
    }

    days_data = {}
    all_hyperlinks = []

    for day in ["monday", "tuesday", "wednesday", "thursday", "friday"]:
        slots = []
        for plan in plans:
            content = getattr(plan, f"{day}_content", None)
            if content:
                enriched = content.copy()

                if "hyperlinks" in enriched and enriched["hyperlinks"]:
                    links = enriched["hyperlinks"]
                    if isinstance(links, dict) and "root" in links:
                        links = links["root"]

                    if isinstance(links, list):
                        for link in links:
                            if isinstance(link, dict):
                                link_copy = link.copy()
                                link_copy["_source_slot"] = plan.slot_number
                                link_copy["_source_subject"] = plan.subject
                                link_copy["table_idx"] = 1
                                all_hyperlinks.append(link_copy)

                orig_instruction = enriched.get("instruction") or {}
                orig_tailored = enriched.get("tailored_instruction") or {}

                instr_parts = []
                activities = (
                    orig_instruction.get("activities")
                    if isinstance(orig_instruction, dict)
                    else None
                )
                if activities:
                    instr_parts.append(activities)

                tailored = (
                    orig_tailored.get("content")
                    if isinstance(orig_tailored, dict)
                    else None
                )
                if tailored:
                    is_duplicate = any(
                        tailored in part or part in tailored for part in instr_parts
                    )
                    if not is_duplicate:
                        instr_parts.append(tailored)

                materials_list = []
                materials_data = enriched.get("materials")
                if materials_data:
                    if isinstance(materials_data, list):
                        materials_list = materials_data
                    elif (
                        isinstance(materials_data, dict)
                        and "root" in materials_data
                    ):
                        materials_list = materials_data["root"]

                enriched["tailored_instruction"] = {
                    "original_content": "\n\n".join(instr_parts),
                    "materials": materials_list,
                }

                misconceptions = enriched.get("misconceptions")
                if (
                    misconceptions
                    and isinstance(misconceptions, dict)
                    and "content" in misconceptions
                ):
                    misconceptions["original_content"] = misconceptions["content"]

                enriched["slot_number"] = plan.slot_number
                enriched["subject"] = plan.subject
                slots.append(enriched)

        if slots:
            if len(plans) == 1:
                days_data[day] = slots[0]
            else:
                days_data[day] = {"slots": slots}

    return {
        "metadata": metadata,
        "days": days_data,
        "_hyperlinks": all_hyperlinks,
        "_media_schema_version": "2.0",
    }


def reconstruct_slots_from_json(
    lesson_json: Dict[str, Any],
) -> Dict[int, Dict[str, Any]]:
    """
    Reconstruct individual slot lesson plans from a multi-slot lesson JSON.
    Returns a mapping of slot_number -> lesson_entry.

    CRITICAL: Uses deep copies to prevent shared state in parallel processing.
    """
    if not lesson_json or "days" not in lesson_json:
        return {}

    existing_slots_by_number = {}

    for day_name, day_data in lesson_json.get("days", {}).items():
        day_slots = day_data.get("slots", [])
        if not day_slots and any(
            k in day_data for k in ["unit_lesson", "objective"]
        ):
            day_slots = [{**day_data}]

        for slot_data in day_slots:
            slot_num = slot_data.get("slot_number")
            if slot_num is not None:
                if slot_num not in existing_slots_by_number:
                    existing_slots_by_number[slot_num] = {
                        "slot_number": slot_num,
                        "subject": slot_data.get("subject", "Unknown"),
                        "lesson_json": {
                            "metadata": copy.deepcopy(
                                lesson_json.get("metadata", {})
                            ),
                            "days": {
                                d: {}
                                for d in [
                                    "monday",
                                    "tuesday",
                                    "wednesday",
                                    "thursday",
                                    "friday",
                                ]
                            },
                        },
                    }
                    existing_slots_by_number[slot_num]["lesson_json"][
                        "metadata"
                    ]["slot_number"] = slot_num
                    existing_slots_by_number[slot_num]["lesson_json"][
                        "metadata"
                    ]["subject"] = slot_data.get("subject", "Unknown")

                day_content = slot_data.copy()
                for field in [
                    "slot_number",
                    "subject",
                    "teacher_name",
                    "grade",
                    "homeroom",
                    "start_time",
                    "end_time",
                ]:
                    day_content.pop(field, None)

                existing_slots_by_number[slot_num]["lesson_json"]["days"][
                    day_name
                ] = day_content

    all_existing_images = lesson_json.get("_images", [])
    all_existing_links = lesson_json.get("_hyperlinks", [])

    has_source_metadata = any(
        link.get("_source_slot") is not None for link in all_existing_links
    )
    is_effectively_single_slot = len(existing_slots_by_number) == 1

    for slot_num, lesson_entry in existing_slots_by_number.items():
        lesson_entry["lesson_json"]["_media_schema_version"] = "2.0"

        if not has_source_metadata and is_effectively_single_slot:
            images_with_meta = []
            for img in all_existing_images:
                img_copy = img.copy()
                img_copy["_source_slot"] = slot_num
                images_with_meta.append(img_copy)
            lesson_entry["lesson_json"]["_images"] = images_with_meta

            links_with_meta = []
            for link in all_existing_links:
                link_copy = link.copy()
                link_copy["_source_slot"] = slot_num
                links_with_meta.append(link_copy)
            lesson_entry["lesson_json"]["_hyperlinks"] = links_with_meta
        else:
            lesson_entry["lesson_json"]["_images"] = [
                img
                for img in all_existing_images
                if img.get("_source_slot") == slot_num
            ]
            lesson_entry["lesson_json"]["_hyperlinks"] = [
                link
                for link in all_existing_links
                if link.get("_source_slot") == slot_num
            ]

    return existing_slots_by_number


def combine_lessons(
    processor: Any,
    user: Dict[str, Any],
    lessons: List[Dict[str, Any]],
    week_of: str,
    generated_at: datetime,
    plan_id: Optional[str] = None,
    get_file_manager_fn: Optional[Callable[..., Any]] = None,
) -> str:
    """
    Combine multiple lessons into a single DOCX using JSON merging.
    Delegates to combine_lessons_impl; wraps with tracking when plan_id is set.
    """
    lessons.sort(key=lambda x: x["slot_number"])

    logger.info(
        "batch_combining_lessons",
        extra={"lesson_count": len(lessons), "user_id": user["id"]},
    )

    if plan_id:
        with processor.tracker.track_operation(
            plan_id,
            "render_combine_lessons",
            metadata={
                "slot_count": len(lessons),
                "consolidated": len(lessons) > 1,
            },
        ):
            return combine_lessons_impl(
                processor,
                user,
                lessons,
                week_of,
                generated_at,
                plan_id,
                get_file_manager_fn,
            )
    return combine_lessons_impl(
        processor,
        user,
        lessons,
        week_of,
        generated_at,
        plan_id,
        get_file_manager_fn,
    )


def combine_lessons_impl(
    processor: Any,
    user: Dict[str, Any],
    lessons: List[Dict[str, Any]],
    week_of: str,
    generated_at: datetime,
    plan_id: Optional[str] = None,
    get_file_manager_fn: Optional[Callable[..., Any]] = None,
) -> str:
    """
    Internal implementation of combine_lessons (without tracking wrapper).
    Merges lesson JSONs, validates, enriches with times, renders to DOCX (single or multi-slot).
    """
    from tools.diagnostic_logger import finalize_diagnostics
    from tools.json_merger import (
        get_merge_summary,
        merge_lesson_jsons,
        validate_merged_json,
    )

    if get_file_manager_fn is None:
        from tools.batch_processor import get_file_manager as _get_fm
        get_file_manager_fn = _get_fm

    def _safe_finalize():
        try:
            finalize_diagnostics()
        except Exception:
            logger.warning("diagnostic_finalize_failed", exc_info=True)

    merged_json = merge_lesson_jsons(lessons)

    is_valid, error_msg = validate_merged_json(merged_json)
    if not is_valid:
        raise ValueError(f"Merged JSON validation failed: {error_msg}")

    logger.debug(
        "batch_merge_summary", extra={"summary": get_merge_summary(merged_json)}
    )

    from backend.utils.lesson_times import enrich_lesson_json_with_times

    user_id = user.get("id") or user.get("user_id")
    if user_id:
        enrich_lesson_json_with_times(merged_json, user_id)

    merged_json["metadata"]["user_name"] = user["name"]

    file_mgr = get_file_manager_fn(base_path=user.get("base_path_override"))
    week_folder = file_mgr.get_week_folder(week_of)

    if len(lessons) > 1:
        week_num = None
        folder_name = week_folder.name
        week_match = re.search(
            r"(\d{2})\s*W\s*(\d{1,2})\b", folder_name, re.IGNORECASE
        )
        if week_match:
            week_num = int(week_match.group(2))
            logger.debug(
                "week_number_extracted_from_folder",
                extra={"folder_name": folder_name, "week_num": week_num},
            )

        if week_num is None:
            week_num = processor._get_week_num(week_of)
            logger.debug(
                "week_number_calculated_from_date",
                extra={"week_of": week_of, "week_num": week_num},
            )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{user['name'].replace(' ', '_')}_Weekly_W{week_num:02d}_{week_of.replace('/', '-')}_{timestamp}.docx"
        output_path = str(week_folder / filename)

        merged_json["metadata"]["total_slots"] = len(lessons)

        teachers = set()
        subjects = set()
        for lesson in lessons:
            if (
                lesson.get("lesson_json", {})
                .get("metadata", {})
                .get("teacher_name")
            ):
                teachers.add(lesson["lesson_json"]["metadata"]["teacher_name"])
            if lesson.get("subject"):
                subjects.add(lesson["subject"])

        if len(teachers) > 1:
            merged_json["metadata"]["teacher_name"] = " / ".join(sorted(teachers))
        if len(subjects) > 1:
            merged_json["metadata"]["subject"] = " / ".join(sorted(subjects))
    else:
        output_path = file_mgr.get_output_path_with_timestamp(
            week_folder, user["name"], week_of
        )

    template_path = (
        user.get("template_path") or "input/Lesson Plan Template SY'25-26.docx"
    )
    if not Path(template_path).exists():
        logger.warning(
            "user_template_not_found",
            extra={"user_id": user.get("id"), "template_path": template_path},
        )
        template_path = "input/Lesson Plan Template SY'25-26.docx"

    if not Path(template_path).exists():
        raise FileNotFoundError(
            f"Template file not found: {template_path}. "
            f"Please ensure the template exists at this path."
        )

    from tools.docx_renderer import DOCXRenderer

    renderer = DOCXRenderer(template_path)

    if len(lessons) == 1:
        return _render_single_slot(
            processor,
            user,
            lessons,
            week_of,
            output_path,
            template_path,
            generated_at,
            plan_id,
            renderer,
            _safe_finalize,
        )
    return _render_multi_slot(
        processor,
        user,
        lessons,
        week_of,
        output_path,
        template_path,
        generated_at,
        plan_id,
        merged_json,
        renderer,
        file_mgr,
        _safe_finalize,
    )


def _render_single_slot(
    processor: Any,
    user: Dict[str, Any],
    lessons: List[Dict[str, Any]],
    week_of: str,
    output_path: str,
    template_path: str,
    generated_at: datetime,
    plan_id: Optional[str],
    renderer: Any,
    safe_finalize: Callable[[], None],
) -> str:
    """Render a single-slot combined lesson to DOCX and add signature/objectives."""
    from tools.diagnostic_logger import get_diagnostic_logger

    lesson = lessons[0]
    slot_num = lesson["slot_number"]
    subject = lesson["subject"]
    lesson_json = lesson["lesson_json"]

    if not isinstance(lesson_json, dict):
        if hasattr(lesson_json, "model_dump"):
            lesson_json = lesson_json.model_dump()
        elif hasattr(lesson_json, "dict"):
            lesson_json = lesson_json.dict()
        else:
            lesson_json = dict(lesson_json) if lesson_json else {}

    if "metadata" not in lesson_json:
        lesson_json["metadata"] = {}
    lesson_json["metadata"]["slot_number"] = slot_num
    lesson_json["metadata"]["subject"] = subject

    if "_hyperlinks" in lesson_json:
        for link in lesson_json.get("_hyperlinks", []):
            if isinstance(link, dict):
                link["_source_slot"] = slot_num
                link["_source_subject"] = subject

    if "_images" in lesson_json:
        for image in lesson_json.get("_images", []):
            if isinstance(image, dict):
                image["_source_slot"] = slot_num
                image["_source_subject"] = subject

    logger.info(
        "batch_render_single_slot",
        extra={"output_file": Path(output_path).name, "week_of": week_of},
    )

    diag = get_diagnostic_logger()
    diag.log_before_render(slot_num, subject, lesson_json, "single_slot")

    output_dir = Path(output_path).parent
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        test_file = output_dir / ".write_test"
        try:
            test_file.touch()
            test_file.unlink()
        except Exception as perm_error:
            raise PermissionError(
                f"Cannot write to output directory '{output_dir}': {perm_error}"
            )
    except Exception as dir_error:
        error_msg = f"Failed to create or access output directory '{output_dir}': {str(dir_error)}"
        print(f"ERROR: {error_msg}")
        traceback.print_exc()
        raise RuntimeError(error_msg) from dir_error

    try:
        render_success = renderer.render(
            lesson_json, output_path, plan_id=plan_id
        )
        if not render_success:
            raise RuntimeError(
                "Renderer returned False - rendering failed silently. Check logs for details."
            )
    except Exception as e:
        error_msg = (
            f"Renderer failed to generate file '{output_path}': {str(e)}"
        )
        print(f"ERROR: {error_msg}")
        traceback.print_exc()
        raise RuntimeError(error_msg) from e

    if not Path(output_path).exists():
        error_msg = f"Renderer failed to create output file: {output_path}. File does not exist after render() call."
        print(f"ERROR: {error_msg}")
        output_dir = Path(output_path).parent
        if not output_dir.exists():
            error_msg += f" Parent directory does not exist: {output_dir}"
        else:
            error_msg += f" Parent directory exists: {output_dir}"
        raise FileNotFoundError(error_msg)

    try:
        doc = Document(output_path)
    except Exception as e:
        error_msg = f"Failed to open rendered file '{output_path}': {str(e)}"
        print(f"ERROR: {error_msg}")
        traceback.print_exc()
        raise FileNotFoundError(error_msg)

    signature_image_path = user.get("signature_image_path")
    if signature_image_path and not signature_image_path.strip():
        signature_image_path = None

    user_name = user.get("name", "")
    if user_name and "Wilson Rodrigues" in user_name:
        if signature_image_path and not Path(signature_image_path).exists():
            logger.warning(
                "wilson_signature_db_path_not_found",
                extra={
                    "user": user_name,
                    "db_path": signature_image_path,
                    "falling_back_to_default": True,
                },
            )
            signature_image_path = None

        if not signature_image_path:
            possible_paths = [
                r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.PNG",
                r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.png",
                r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.Png",
            ]
            for sig_path in possible_paths:
                if Path(sig_path).exists():
                    signature_image_path = sig_path
                    logger.info(
                        "using_default_wilson_signature",
                        extra={"path": sig_path, "user": user_name},
                    )
                    break
            else:
                logger.warning(
                    "wilson_signature_not_found",
                    extra={
                        "user": user_name,
                        "attempted_paths": possible_paths,
                    },
                )

    try:
        processor._modify_existing_signature_table(
            doc, generated_at, signature_image_path, user.get("name")
        )
        doc.save(output_path)
    except Exception as e:
        error_msg = f"Failed to modify/save document '{output_path}': {str(e)}"
        print(f"ERROR: {error_msg}")
        traceback.print_exc()
        raise

    try:
        from backend.services.objectives_printer import ObjectivesPrinter

        objectives_printer = ObjectivesPrinter()
        objectives_output_path = Path(output_path).with_name(
            Path(output_path).stem + "_objectives.docx"
        )
        objectives_printer.generate_docx(
            lesson_json,
            str(objectives_output_path),
            user_name=user.get("name"),
            week_of=week_of,
        )
        logger.info(
            "batch_objectives_docx_generated",
            extra={"objectives_path": str(objectives_output_path)},
        )
    except Exception as e:
        logger.warning(
            "batch_objectives_docx_failed",
            extra={"error": str(e), "error_type": type(e).__name__},
            exc_info=True,
        )

    debug_marker = (
        Path(output_path).parent
        / f"{Path(output_path).stem}_objectives_debug.txt"
    )
    try:
        with open(debug_marker, "w", encoding="utf-8") as f:
            f.write(f"Objectives generation started at {datetime.now()}\n")
            f.write(f"Output path: {output_path}\n")
    except Exception:
        pass

    try:
        from backend.services.objectives_pdf_generator import (
            generate_objectives_pdf,
        )

        objectives_pdf_path = Path(output_path).with_name(
            Path(output_path).stem + "_objectives.pdf"
        )
        from backend.utils.lesson_times import enrich_lesson_json_with_times

        user_id = user.get("id") or user.get("user_id")
        if user_id:
            enrich_lesson_json_with_times(lesson_json, user_id)

        lesson_json_for_pdf = processor._sanitize_value(lesson_json)

        logger.info(
            "batch_objectives_pdf_html_starting",
            extra={"objectives_pdf_path": str(objectives_pdf_path)},
        )
        generate_objectives_pdf(
            lesson_json_for_pdf,
            str(objectives_pdf_path),
            user_name=user.get("name"),
            keep_html=True,
        )
        objectives_html_path = objectives_pdf_path.with_suffix(".html")

        if not objectives_pdf_path.exists():
            logger.error(
                "batch_objectives_pdf_not_created",
                extra={"expected_path": str(objectives_pdf_path)},
            )
        if not objectives_html_path.exists():
            logger.error(
                "batch_objectives_html_not_created",
                extra={"expected_path": str(objectives_html_path)},
            )

        logger.info(
            "batch_objectives_pdf_html_generated",
            extra={
                "objectives_pdf_path": str(objectives_pdf_path),
                "objectives_html_path": str(objectives_html_path),
                "pdf_exists": objectives_pdf_path.exists(),
                "html_exists": objectives_html_path.exists(),
            },
        )

        try:
            from backend.services.sentence_frames_pdf_generator import (
                generate_sentence_frames_pdf,
                generate_sentence_frames_docx,
            )

            output_dir = Path(output_path).parent
            output_stem = Path(output_path).stem
            sentence_frames_pdf_path = (
                output_dir / f"{output_stem}_sentence_frames.pdf"
            )

            has_frames = False
            frames_count = 0
            if "days" in lesson_json:
                for day_name, day in lesson_json["days"].items():
                    if not isinstance(day, dict):
                        continue
                    day_frames = day.get("sentence_frames")
                    if (
                        day_frames
                        and isinstance(day_frames, list)
                        and len(day_frames) > 0
                    ):
                        has_frames = True
                        frames_count += len(day_frames)
                    slots = day.get("slots", [])
                    if isinstance(slots, list):
                        for slot in slots:
                            if not isinstance(slot, dict):
                                continue
                            slot_frames = slot.get("sentence_frames")
                            if (
                                slot_frames
                                and isinstance(slot_frames, list)
                                and len(slot_frames) > 0
                            ):
                                has_frames = True
                                frames_count += len(slot_frames)

            if not has_frames and "days" in lesson_json_for_pdf:
                for day in lesson_json_for_pdf["days"].values():
                    if not isinstance(day, dict):
                        continue
                    if "sentence_frames" in day and day["sentence_frames"]:
                        has_frames = True
                        frames_count += len(day.get("sentence_frames", []))
                        break
                    if "slots" in day:
                        for s in day["slots"]:
                            if not isinstance(s, dict):
                                continue
                            if "sentence_frames" in s and s["sentence_frames"]:
                                has_frames = True
                                frames_count += len(
                                    s.get("sentence_frames", [])
                                )
                                break

            if has_frames:
                output_dir.mkdir(parents=True, exist_ok=True)
                generate_sentence_frames_pdf(
                    lesson_json_for_pdf,
                    str(sentence_frames_pdf_path),
                    user_name=user.get("name"),
                    keep_html=True,
                )
                sentence_frames_docx_path = (
                    sentence_frames_pdf_path.with_suffix(".docx")
                )
                generate_sentence_frames_docx(
                    lesson_json_for_pdf,
                    str(sentence_frames_docx_path),
                    user_name=user.get("name"),
                )
                logger.info(
                    "batch_sentence_frames_pdf_generated",
                    extra={
                        "path": str(sentence_frames_pdf_path),
                        "pdf_exists": sentence_frames_pdf_path.exists(),
                        "frames_count": frames_count,
                    },
                )
            else:
                logger.warning(
                    "batch_sentence_frames_skipped_no_data",
                    extra={
                        "plan_id": plan_id,
                        "output_path": str(output_path),
                        "has_days": "days" in lesson_json,
                    },
                )
        except ImportError:
            pass
        except Exception as e:
            logger.warning(
                "batch_sentence_frames_pdf_html_failed",
                extra={"error": str(e), "error_type": type(e).__name__},
                exc_info=True,
            )
    except Exception as e:
        logger.warning(
            "batch_objectives_pdf_html_failed",
            extra={"error": str(e), "error_type": type(e).__name__},
            exc_info=True,
        )

    logger.info(
        "batch_render_single_slot_success", extra={"output_path": output_path}
    )
    safe_finalize()
    return output_path


def _render_multi_slot(
    processor: Any,
    user: Dict[str, Any],
    lessons: List[Dict[str, Any]],
    week_of: str,
    output_path: str,
    template_path: str,
    generated_at: datetime,
    plan_id: Optional[str],
    merged_json: Dict[str, Any],
    renderer: Any,
    file_mgr: Any,
    safe_finalize: Callable[[], None],
) -> str:
    """Render multi-slot: each slot to temp DOCX, merge, signatures, objectives."""
    if not lessons:
        raise ValueError("No lessons provided to combine")

    logger.info(
        "batch_render_multi_slot_start",
        extra={
            "lesson_count": len(lessons),
            "output_file": Path(output_path).name,
        },
    )

    temp_files = []
    all_unplaced_hyperlinks = []
    all_unplaced_images = []
    week_folder = file_mgr.get_week_folder(week_of)

    for lesson in lessons:
        slot_num = lesson["slot_number"]
        subject = lesson["subject"]
        lesson_json = lesson["lesson_json"]
        slot_data = lesson.get("slot_data", {})

        if not isinstance(lesson_json, dict):
            if hasattr(lesson_json, "model_dump"):
                lesson_json = lesson_json.model_dump()
            elif hasattr(lesson_json, "dict"):
                lesson_json = lesson_json.dict()
            else:
                lesson_json = dict(lesson_json) if lesson_json else {}

        if "metadata" not in lesson_json:
            lesson_json["metadata"] = {}
        lesson_json["metadata"]["slot_number"] = slot_num
        lesson_json["metadata"]["subject"] = subject

        if slot_data:
            if isinstance(slot_data, dict):
                primary_teacher_name = slot_data.get("primary_teacher_name")
                primary_teacher_first_name = slot_data.get(
                    "primary_teacher_first_name"
                )
                primary_teacher_last_name = slot_data.get(
                    "primary_teacher_last_name"
                )
                lesson_json["metadata"]["primary_teacher_name"] = (
                    primary_teacher_name
                )
                lesson_json["metadata"]["primary_teacher_first_name"] = (
                    primary_teacher_first_name
                )
                lesson_json["metadata"]["primary_teacher_last_name"] = (
                    primary_teacher_last_name
                )
            else:
                primary_teacher_name = getattr(
                    slot_data, "primary_teacher_name", None
                )
                primary_teacher_first_name = getattr(
                    slot_data, "primary_teacher_first_name", None
                )
                primary_teacher_last_name = getattr(
                    slot_data, "primary_teacher_last_name", None
                )
                lesson_json["metadata"]["primary_teacher_name"] = (
                    primary_teacher_name
                )
                lesson_json["metadata"]["primary_teacher_first_name"] = (
                    primary_teacher_first_name
                )
                lesson_json["metadata"]["primary_teacher_last_name"] = (
                    primary_teacher_last_name
                )

            try:
                combined_teacher_name = processor._build_teacher_name(
                    {
                        "first_name": getattr(processor, "_user_first_name", ""),
                        "last_name": getattr(processor, "_user_last_name", ""),
                        "name": getattr(processor, "_user_name", ""),
                    },
                    slot_data
                    if isinstance(slot_data, dict)
                    else {
                        "primary_teacher_name": primary_teacher_name,
                        "primary_teacher_first_name": primary_teacher_first_name,
                        "primary_teacher_last_name": primary_teacher_last_name,
                    },
                )
                lesson_json["metadata"]["teacher_name"] = combined_teacher_name
            except Exception:
                primary_teacher_name = (
                    slot_data.get("primary_teacher_name")
                    if isinstance(slot_data, dict)
                    else getattr(slot_data, "primary_teacher_name", None)
                )
                lesson_json["metadata"]["teacher_name"] = (
                    primary_teacher_name or "Unknown"
                )

        lesson_json["_media_schema_version"] = "2.0"

        if "_hyperlinks" in lesson_json:
            for link in lesson_json.get("_hyperlinks", []):
                if isinstance(link, dict):
                    link["_source_slot"] = slot_num
                    link["_source_subject"] = subject

        if "_images" in lesson_json:
            for image in lesson_json.get("_images", []):
                if isinstance(image, dict):
                    image["_source_slot"] = slot_num
                    image["_source_subject"] = subject

        temp_filename = f"_temp_slot{slot_num}_{subject.replace('/', '_')}.docx"
        temp_path = str(week_folder / temp_filename)

        logger.debug(
            "batch_render_slot",
            extra={
                "slot_number": slot_num,
                "subject": subject,
                "hyperlinks": len(lesson_json.get("_hyperlinks", [])),
                "images": len(lesson_json.get("_images", [])),
            },
        )

        from tools.diagnostic_logger import get_diagnostic_logger

        diag = get_diagnostic_logger()
        diag.log_before_render(slot_num, subject, lesson_json, "multi_slot")

        try:
            success, unplaced_hl, unplaced_img = renderer.render(
                lesson_json,
                temp_path,
                plan_id=plan_id,
                skip_fallback_sections=True,
            )
            if success:
                temp_files.append(temp_path)
                all_unplaced_hyperlinks.extend(unplaced_hl)
                all_unplaced_images.extend(unplaced_img)
            else:
                logger.error(
                    "batch_render_slot_failed",
                    extra={
                        "slot_number": slot_num,
                        "subject": subject,
                        "temp_path": temp_path,
                    },
                )
        except Exception:
            logger.error(
                "batch_render_slot_failed",
                extra={
                    "slot_number": slot_num,
                    "subject": subject,
                    "temp_path": temp_path,
                },
            )

    if not temp_files:
        error_msg = (
            f"No files to merge: All {len(lessons)} slot(s) failed to render. "
            f"Check logs for rendering errors."
        )
        logger.error("batch_merge_no_files", extra={"lesson_count": len(lessons)})
        raise ValueError(error_msg)

    logger.info(
        "batch_merge_slots",
        extra={"slot_count": len(temp_files), "output_file": output_path},
    )
    try:
        if plan_id:
            with processor.tracker.track_operation(
                plan_id,
                "render_merge_files",
                metadata={"file_count": len(temp_files)},
            ):
                merge_docx_files(temp_files, output_path)
        else:
            merge_docx_files(temp_files, output_path)
        logger.info(
            "batch_merge_slots_success",
            extra={"slot_count": len(temp_files)},
        )
    except Exception as e:
        logger.exception("batch_merge_slots_error", extra={"error": str(e)})
        raise

    logger.debug("batch_clean_signature_boxes")
    doc = Document(output_path)

    signature_image_path = user.get("signature_image_path")
    if signature_image_path and not signature_image_path.strip():
        signature_image_path = None

    user_name = user.get("name", "")
    if (
        not signature_image_path
        and user_name
        and "Wilson Rodrigues" in user_name
    ):
        possible_paths = [
            r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.PNG",
            r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.png",
            r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\wilson_rodrigues_signature.Png",
        ]
        for sig_path in possible_paths:
            if Path(sig_path).exists():
                signature_image_path = sig_path
                logger.info(
                    "using_default_wilson_signature",
                    extra={"path": sig_path, "user": user_name},
                )
                break
        else:
            logger.warning(
                "wilson_signature_not_found",
                extra={
                    "user": user_name,
                    "attempted_paths": possible_paths,
                },
            )

    if plan_id:
        with processor.tracker.track_operation(plan_id, "render_remove_signatures"):
            processor._remove_signature_boxes(doc)
        with processor.tracker.track_operation(plan_id, "render_add_signature"):
            processor._add_signature_box(
                doc,
                generated_at,
                template_path,
                signature_image_path,
                user.get("name"),
            )
        doc.save(output_path)
    else:
        processor._remove_signature_boxes(doc)
        processor._add_signature_box(
            doc,
            generated_at,
            template_path,
            signature_image_path,
            user.get("name"),
        )
        doc.save(output_path)

    if all_unplaced_hyperlinks or all_unplaced_images:
        logger.info(
            "consolidating_fallback_media",
            extra={
                "hyperlinks": len(all_unplaced_hyperlinks),
                "images": len(all_unplaced_images),
            },
        )

        unique_hl = []
        seen_hl = set()
        for hl in all_unplaced_hyperlinks:
            key = (hl.get("text", "").strip(), hl.get("url", "").strip())
            if key not in seen_hl:
                seen_hl.add(key)
                unique_hl.append(hl)

        unique_img = []
        seen_img = set()
        import hashlib

        for img in all_unplaced_images:
            data_hash = hashlib.md5(img["data"].encode()).hexdigest()
            if data_hash not in seen_img:
                seen_img.add(data_hash)
                unique_img.append(img)

        doc = Document(output_path)
        renderer._append_unmatched_media(doc, unique_hl, unique_img)
        doc.save(output_path)

    try:
        from backend.services.objectives_printer import ObjectivesPrinter

        objectives_printer = ObjectivesPrinter()
        objectives_output_path = Path(output_path).with_name(
            Path(output_path).stem + "_objectives.docx"
        )
        objectives_printer.generate_docx(
            merged_json,
            str(objectives_output_path),
            user_name=user.get("name"),
            week_of=week_of,
        )
        logger.info(
            "batch_objectives_docx_generated",
            extra={"objectives_path": str(objectives_output_path)},
        )
    except Exception as e:
        logger.warning(
            "batch_objectives_docx_failed",
            extra={"error": str(e), "error_type": type(e).__name__},
            exc_info=True,
        )

    merged_json_for_pdf = processor._sanitize_value(merged_json)

    try:
        from backend.services.objectives_pdf_generator import (
            generate_objectives_pdf,
        )

        objectives_pdf_path = Path(output_path).with_name(
            Path(output_path).stem + "_objectives.pdf"
        )
        logger.info(
            "batch_objectives_pdf_html_starting",
            extra={"objectives_pdf_path": str(objectives_pdf_path)},
        )
        generate_objectives_pdf(
            merged_json_for_pdf,
            str(objectives_pdf_path),
            user_name=user.get("name"),
            keep_html=True,
        )
        objectives_html_path = objectives_pdf_path.with_suffix(".html")

        if not objectives_pdf_path.exists():
            logger.error(
                "batch_objectives_pdf_not_created",
                extra={"expected_path": str(objectives_pdf_path)},
            )
        if not objectives_html_path.exists():
            logger.error(
                "batch_objectives_html_not_created",
                extra={"expected_path": str(objectives_html_path)},
            )
        logger.info(
            "batch_objectives_pdf_html_generated",
            extra={
                "objectives_pdf_path": str(objectives_pdf_path),
                "objectives_html_path": str(objectives_html_path),
                "pdf_exists": objectives_pdf_path.exists(),
                "html_exists": objectives_html_path.exists(),
            },
        )
    except Exception as e:
        logger.warning(
            "batch_objectives_pdf_html_failed",
            extra={"error": str(e), "error_type": type(e).__name__},
            exc_info=True,
        )

    try:
        from backend.services.sentence_frames_pdf_generator import (
            generate_sentence_frames_pdf,
            generate_sentence_frames_docx,
        )

        output_dir = Path(output_path).parent
        output_stem = Path(output_path).stem
        sentence_frames_pdf_path = (
            output_dir / f"{output_stem}_sentence_frames.pdf"
        )

        has_frames = False
        frames_count = 0
        if "days" in merged_json:
            for day_name, day in merged_json["days"].items():
                if not isinstance(day, dict):
                    continue
                day_frames = day.get("sentence_frames")
                if (
                    day_frames
                    and isinstance(day_frames, list)
                    and len(day_frames) > 0
                ):
                    has_frames = True
                    frames_count += len(day_frames)
                slots = day.get("slots", [])
                if isinstance(slots, list):
                    for slot in slots:
                        if not isinstance(slot, dict):
                            continue
                        slot_frames = slot.get("sentence_frames")
                        if (
                            slot_frames
                            and isinstance(slot_frames, list)
                            and len(slot_frames) > 0
                        ):
                            has_frames = True
                            frames_count += len(slot_frames)

        if not has_frames and "days" in merged_json_for_pdf:
            for day in merged_json_for_pdf["days"].values():
                if not isinstance(day, dict):
                    continue
                if "sentence_frames" in day and day["sentence_frames"]:
                    has_frames = True
                    frames_count += len(day.get("sentence_frames", []))
                    break
                if "slots" in day:
                    for s in day["slots"]:
                        if not isinstance(s, dict):
                            continue
                        if "sentence_frames" in s and s["sentence_frames"]:
                            has_frames = True
                            frames_count += len(s.get("sentence_frames", []))
                            break

        if has_frames:
            output_dir.mkdir(parents=True, exist_ok=True)
            generate_sentence_frames_pdf(
                merged_json_for_pdf,
                str(sentence_frames_pdf_path),
                user_name=user.get("name"),
                keep_html=True,
            )
            sentence_frames_docx_path = sentence_frames_pdf_path.with_suffix(
                ".docx"
            )
            generate_sentence_frames_docx(
                merged_json_for_pdf,
                str(sentence_frames_docx_path),
                user_name=user.get("name"),
            )
            logger.info(
                "batch_sentence_frames_pdf_generated",
                extra={
                    "path": str(sentence_frames_pdf_path),
                    "pdf_exists": sentence_frames_pdf_path.exists(),
                    "frames_count": frames_count,
                    "multi_slot": True,
                },
            )
        else:
            logger.warning(
                "batch_sentence_frames_skipped_no_data",
                extra={
                    "plan_id": plan_id,
                    "output_path": str(output_path),
                    "has_days": "days" in merged_json,
                    "multi_slot": True,
                },
            )
    except ImportError:
        pass
    except Exception as e:
        logger.warning(
            "batch_sentence_frames_pdf_html_failed",
            extra={
                "error": str(e),
                "error_type": type(e).__name__,
                "multi_slot": True,
            },
            exc_info=True,
        )

    for temp_file in temp_files:
        try:
            temp_path = Path(temp_file)
            if temp_path.exists():
                temp_path.unlink()
                logger.debug(
                    "batch_temp_file_deleted", extra={"temp_file": temp_file}
                )
        except Exception as e:
            logger.warning(
                "batch_temp_cleanup_failed",
                extra={"temp_file": temp_file, "error": str(e)},
            )

    logger.info(
        "batch_render_multi_slot_success",
        extra={"lesson_count": len(lessons), "output_path": output_path},
    )
    safe_finalize()
    return output_path
