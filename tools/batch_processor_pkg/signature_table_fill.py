"""
Add signature image or user name to the teacher signature cell in a table.
Used by signatures.py (modify_existing_signature_table, add_signature_box).
"""

import re
import traceback
from io import BytesIO
from pathlib import Path

from docx.shared import Inches, Pt

from backend.telemetry import logger
from tools.docx_renderer.style import sanitize_xml_text

from tools.batch_processor_pkg.signature_paragraph_helpers import (
    add_date_section_to_paragraph,
    get_paragraph_font_info,
)


def add_signature_image_to_table(
    table, signature_image_path: str  # noqa: ANN001
) -> None:
    """Add signature image to the signature table."""
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if "Teacher Signature:" in cell_text:
                for para in cell.paragraphs:
                    para_text = para.text
                    if "Teacher Signature:" in para_text:
                        text_font_size = Pt(11)
                        for run in para.runs:
                            if run.font.size:
                                text_font_size = run.font.size
                                break

                        text_height_inches = text_font_size.pt / 72.0
                        image_height_inches = text_height_inches * 1.4

                        teacher_pos = para_text.find("Teacher Signature:")
                        date_pos = (
                            para_text.find("Date:", teacher_pos)
                            if teacher_pos >= 0
                            else -1
                        )

                        if teacher_pos >= 0 and date_pos > teacher_pos:
                            try:
                                before_teacher = para_text[:teacher_pos]

                                between_text = para_text[
                                    teacher_pos
                                    + len("Teacher Signature:") : date_pos
                                ].strip()
                                between_text = between_text.lstrip("_").strip()

                                date_updated = False
                                date_value = None
                                date_pattern = r"Date:\s*(\d{1,2}/\d{1,2}/\d{4})"
                                date_match = re.search(date_pattern, para_text)
                                if date_match:
                                    date_updated = True
                                    date_value = date_match.group(1)

                                original_font_size, original_font_name = (
                                    get_paragraph_font_info(para)
                                )

                                para.clear()

                                if before_teacher.strip():
                                    before_run = para.add_run(
                                        sanitize_xml_text(before_teacher)
                                    )
                                    if original_font_size:
                                        before_run.font.size = original_font_size
                                    if original_font_name:
                                        before_run.font.name = original_font_name

                                teacher_run = para.add_run("Teacher Signature: ")
                                teacher_run.font.bold = True
                                if original_font_size:
                                    teacher_run.font.size = original_font_size
                                if original_font_name:
                                    teacher_run.font.name = original_font_name

                                image_run = para.add_run()
                                try:
                                    if not Path(signature_image_path).exists():
                                        logger.error(
                                            "signature_image_file_not_found",
                                            extra={"path": signature_image_path},
                                        )
                                        raise FileNotFoundError(
                                            f"Signature image not found: {signature_image_path}"
                                        )

                                    with open(
                                        signature_image_path, "rb"
                                    ) as img_file:
                                        img_stream = BytesIO(img_file.read())
                                        image_run.add_picture(
                                            img_stream,
                                            height=Inches(image_height_inches),
                                        )

                                    if not image_run._element.xpath(".//a:blip"):
                                        logger.error(
                                            "signature_image_not_inserted",
                                            extra={"path": signature_image_path},
                                        )
                                except Exception as img_error:
                                    logger.error(
                                        "signature_image_insertion_error",
                                        extra={
                                            "path": signature_image_path,
                                            "error": str(img_error),
                                        },
                                    )
                                    logger.debug(
                                        "signature_image_error_traceback",
                                        extra={"traceback": traceback.format_exc()},
                                    )

                                if between_text:
                                    between_run = para.add_run(
                                        sanitize_xml_text(between_text)
                                    )
                                    if original_font_size:
                                        between_run.font.size = original_font_size
                                    if original_font_name:
                                        between_run.font.name = original_font_name

                                add_date_section_to_paragraph(
                                    para,
                                    date_value if date_updated else None,
                                    original_font_size,
                                    original_font_name,
                                )

                                logger.info(
                                    "signature_image_added",
                                    extra={"path": signature_image_path},
                                )
                                final_text = para.text
                                if "Teacher Signature:" not in final_text:
                                    logger.error(
                                        "signature_text_missing_after_insertion",
                                        extra={
                                            "path": signature_image_path,
                                            "final_text": final_text,
                                        },
                                    )
                            except Exception as e:
                                logger.warning(
                                    "signature_image_failed",
                                    extra={
                                        "path": signature_image_path,
                                        "error": str(e),
                                    },
                                )
                                logger.debug(
                                    "signature_image_traceback",
                                    extra={"traceback": traceback.format_exc()},
                                )
                            break
                break


def add_user_name_to_table(
    table, user_name: str  # noqa: ANN001
) -> None:
    """Add user name (underlined) to the signature table."""
    user_name = sanitize_xml_text(user_name or "")
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if "Teacher Signature:" in cell_text:
                for para in cell.paragraphs:
                    para_text = para.text
                    if "Teacher Signature:" in para_text:
                        teacher_pos = para_text.find("Teacher Signature:")
                        date_pos = (
                            para_text.find("Date:", teacher_pos)
                            if teacher_pos >= 0
                            else -1
                        )

                        if teacher_pos >= 0 and date_pos > teacher_pos:
                            before_teacher = para_text[:teacher_pos]

                            between_text = para_text[
                                teacher_pos + len("Teacher Signature:") : date_pos
                            ].strip()
                            between_text = between_text.lstrip("_").strip()

                            date_updated = False
                            date_value = None
                            date_pattern = r"Date:\s*(\d{1,2}/\d{1,2}/\d{4})"
                            date_match = re.search(date_pattern, para_text)
                            if date_match:
                                date_updated = True
                                date_value = date_match.group(1)

                            original_font_size, original_font_name = (
                                get_paragraph_font_info(para)
                            )

                            para.clear()

                            if before_teacher.strip():
                                before_run = para.add_run(
                                    sanitize_xml_text(before_teacher)
                                )
                                if original_font_size:
                                    before_run.font.size = original_font_size
                                if original_font_name:
                                    before_run.font.name = original_font_name

                            teacher_run = para.add_run("Teacher Signature: ")
                            teacher_run.font.bold = True
                            if original_font_size:
                                teacher_run.font.size = original_font_size
                            if original_font_name:
                                teacher_run.font.name = original_font_name

                            name_run = para.add_run(user_name)
                            name_run.font.underline = True
                            if original_font_size:
                                name_run.font.size = original_font_size
                            if original_font_name:
                                name_run.font.name = original_font_name

                            if between_text:
                                between_run = para.add_run(
                                    sanitize_xml_text(between_text)
                                )
                                if original_font_size:
                                    between_run.font.size = original_font_size
                                if original_font_name:
                                    between_run.font.name = original_font_name

                            add_date_section_to_paragraph(
                                para,
                                date_value if date_updated else None,
                                original_font_size,
                                original_font_name,
                            )

                            final_text = para.text
                            if "Teacher Signature:" not in final_text:
                                logger.error(
                                    "signature_text_missing_after_name_insertion",
                                    extra={
                                        "user_name": user_name,
                                        "final_text": final_text,
                                    },
                                )
                            break
                break
