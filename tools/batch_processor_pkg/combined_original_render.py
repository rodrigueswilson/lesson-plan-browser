"""
Synchronous rendering of combined original DOCX: deduplication, per-plan render, merge, style normalization.
Used by combined_original.generate_combined_original_docx.
"""

import os
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, List, Optional

from backend.config import settings
from backend.telemetry import logger


def render_combined_originals_sync(
    processor: Any,
    plans: List[Any],
    week_of: str,
    plan_id: str,
    week_folder_path: Optional[str],
    get_file_manager_fn: Callable[..., Any],
) -> Optional[str]:
    """
    Renders a combined DOCX of the given original plans. Creates temp files per plan,
    merges them, normalizes styles. Returns output path or None.
    """
    temp_files = []
    rendering_results = {}
    try:
        from docx import Document

        from tools.docx_renderer import DOCXRenderer
        from tools.docx_utils import (
            diagnose_style_conflicts,
            normalize_styles_from_master,
            normalize_styles_via_file,
        )

        file_mgr = get_file_manager_fn(
            base_path=getattr(processor, "_user_base_path", None)
        )
        week_folder = (
            Path(week_folder_path)
            if week_folder_path
            else file_mgr.get_week_folder(week_of)
        )
        originals_dir = week_folder / "originals"
        originals_dir.mkdir(parents=True, exist_ok=True)

        safe_week = (
            week_of.replace("/", "-")
            .replace("\\", "-")
            .replace(" ", "_")
        )
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"combined_originals_{safe_week}_{timestamp}.docx"
        output_path = originals_dir / filename

        logger.info(
            "combined_originals_processing_started",
            extra={"total_plans": len(plans)},
        )

        renderer = DOCXRenderer(settings.DOCX_TEMPLATE_PATH)
        renderer.is_originals = True

        style_master = Document(settings.DOCX_TEMPLATE_PATH)

        unique_slot_files = set()
        deduplicated_plans = []
        removed_plans = []

        for plan in plans:
            source_key = getattr(plan, "source_file_path", None)
            slot_key = plan.slot_number
            subject_key = plan.subject
            dedup_key = (source_key, slot_key, subject_key)

            if dedup_key not in unique_slot_files:
                unique_slot_files.add(dedup_key)
                deduplicated_plans.append(plan)
            else:
                removed_plans.append({
                    "slot_number": plan.slot_number,
                    "subject": plan.subject,
                    "source_file_name": plan.source_file_name,
                })

        logger.info(
            "combined_originals_deduplication",
            extra={
                "original_count": len(plans),
                "deduplicated_count": len(deduplicated_plans),
                "unique_keys_count": len(unique_slot_files),
                "removed_plans": removed_plans,
                "deduplicated_plans": [
                    {
                        "slot_number": p.slot_number,
                        "subject": p.subject,
                        "source_file_name": p.source_file_name,
                    }
                    for p in deduplicated_plans
                ],
            },
        )

        original_slot_subjects = {
            (p.slot_number, p.subject) for p in plans
        }
        dedup_slot_subjects = {
            (p.slot_number, p.subject) for p in deduplicated_plans
        }
        missing_combinations = (
            original_slot_subjects - dedup_slot_subjects
        )
        if missing_combinations:
            logger.warning(
                "combined_originals_missing_after_dedup",
                extra={
                    "missing_combinations": [
                        {"slot_number": s, "subject": subj}
                        for s, subj in missing_combinations
                    ]
                },
            )

        diagnosed_style_conflicts_once = False
        for plan in deduplicated_plans:
            temp_filename = f"_temp_orig_slot{plan.slot_number}_{plan.subject.replace('/', '_')}.docx"
            temp_path = str(originals_dir / temp_filename)

            try:
                logger.info(
                    "combined_originals_rendering_plan",
                    extra={
                        "slot_number": plan.slot_number,
                        "subject": plan.subject,
                        "source_file_name": plan.source_file_name,
                        "temp_path": temp_path,
                    },
                )

                slot_json = processor._convert_originals_to_json([plan])

                if slot_json.get("metadata", {}).get(
                    "slot_number"
                ) != plan.slot_number:
                    logger.warning(
                        "combined_originals_slot_number_mismatch",
                        extra={
                            "expected": plan.slot_number,
                            "actual": slot_json.get("metadata", {}).get(
                                "slot_number"
                            ),
                        },
                    )
                if slot_json.get("metadata", {}).get(
                    "subject"
                ) != plan.subject:
                    logger.warning(
                        "combined_originals_subject_mismatch",
                        extra={
                            "expected": plan.subject,
                            "actual": slot_json.get("metadata", {}).get(
                                "subject"
                            ),
                        },
                    )

                render_result = renderer.render(
                    slot_json,
                    temp_path,
                    plan_id=plan_id,
                    skip_fallback_sections=True,
                )
                success, unplaced_hl, unplaced_img = (
                    render_result
                    if isinstance(render_result, tuple)
                    else (render_result, [], [])
                )

                if success:
                    sub_doc = Document(temp_path)
                    from tools.docx_utils import (
                        remove_headers_footers,
                        strip_custom_styles,
                        strip_problematic_elements,
                        strip_sections,
                    )

                    strip_problematic_elements(sub_doc)
                    remove_headers_footers(sub_doc)
                    strip_sections(sub_doc)
                    strip_custom_styles(sub_doc, style_master)

                    if not diagnosed_style_conflicts_once:
                        diagnosis = diagnose_style_conflicts(
                            style_master, sub_doc
                        )
                        logger.debug(
                            "combined_originals_style_diagnosis",
                            extra=diagnosis,
                        )
                        diagnosed_style_conflicts_once = True

                    normalize_styles_from_master(style_master, sub_doc)
                    if hasattr(sub_doc, "_normalized_stream"):
                        sub_doc._normalized_stream.seek(0)
                        sub_doc = Document(sub_doc._normalized_stream)

                    sub_doc.save(temp_path)
                    temp_files.append(temp_path)
                    rendering_results[plan.slot_number] = {
                        "success": True,
                        "subject": plan.subject,
                        "temp_path": temp_path,
                    }
                    logger.info(
                        "combined_originals_plan_rendered_successfully",
                        extra={
                            "slot_number": plan.slot_number,
                            "subject": plan.subject,
                        },
                    )
                else:
                    rendering_results[plan.slot_number] = {
                        "success": False,
                        "subject": plan.subject,
                        "error": "Render returned False",
                    }
                    logger.error(
                        "combined_originals_plan_render_failed",
                        extra={
                            "slot_number": plan.slot_number,
                            "subject": plan.subject,
                            "unplaced_hyperlinks": len(unplaced_hl),
                            "unplaced_images": len(unplaced_img),
                        },
                    )
            except Exception as e:
                rendering_results[plan.slot_number] = {
                    "success": False,
                    "subject": plan.subject,
                    "error": str(e),
                }
                logger.error(
                    "combined_originals_plan_render_exception",
                    extra={
                        "slot_number": plan.slot_number,
                        "subject": plan.subject,
                        "error": str(e),
                    },
                    exc_info=True,
                )

        successful_slots = [
            slot_num
            for slot_num, result in rendering_results.items()
            if result.get("success")
        ]
        failed_slots = [
            slot_num
            for slot_num, result in rendering_results.items()
            if not result.get("success")
        ]

        logger.info(
            "combined_originals_rendering_summary",
            extra={
                "total_plans": len(deduplicated_plans),
                "successful_renders": len(successful_slots),
                "failed_renders": len(failed_slots),
                "successful_slots": successful_slots,
                "failed_slots": [
                    {
                        "slot_number": slot_num,
                        "subject": rendering_results[slot_num].get(
                            "subject"
                        ),
                        "error": rendering_results[slot_num].get(
                            "error"
                        ),
                    }
                    for slot_num in failed_slots
                ],
            },
        )

        if not temp_files:
            logger.warning(
                "combined_originals_no_files_generated",
                extra={
                    "total_plans": len(deduplicated_plans),
                    "rendering_results": rendering_results,
                },
            )
            return None

        expected_slots = {
            p.slot_number for p in deduplicated_plans
        }
        rendered_slots = {
            p.slot_number
            for p in deduplicated_plans
            if rendering_results.get(p.slot_number, {}).get("success")
        }
        missing_slots = expected_slots - rendered_slots
        if missing_slots:
            logger.warning(
                "combined_originals_missing_slots",
                extra={
                    "missing_slots": [
                        {
                            "slot_number": slot_num,
                            "subject": next(
                                (
                                    p.subject
                                    for p in deduplicated_plans
                                    if p.slot_number == slot_num
                                ),
                                "Unknown",
                            ),
                        }
                        for slot_num in missing_slots
                    ],
                },
            )

        logger.info(
            "combined_originals_merging_files",
            extra={
                "file_count": len(temp_files),
                "file_paths": [Path(f).name for f in temp_files],
            },
        )
        processor._merge_docx_files(
            temp_files,
            str(output_path),
            master_template_path=settings.DOCX_TEMPLATE_PATH,
        )

        template_doc = Document(settings.DOCX_TEMPLATE_PATH)
        merged_doc = Document(str(output_path))
        normalized_stream = normalize_styles_via_file(template_doc, merged_doc)
        if normalized_stream:
            output_path.write_bytes(normalized_stream.getvalue())
            logger.info(
                "combined_originals_post_merge_styles_applied",
                extra={"output_path": str(output_path.absolute())},
            )
        else:
            logger.warning(
                "combined_originals_post_merge_styles_failed",
                extra={
                    "output_path": str(output_path.absolute()),
                    "note": "merged file unchanged; Word may show Styles 1 error",
                },
            )

        try:
            merged_doc = Document(str(output_path))
            table_count = len(merged_doc.tables)
            logger.info(
                "combined_originals_merged_structure",
                extra={
                    "table_count": table_count,
                    "estimated_slots": table_count // 2,
                },
            )
        except Exception as e:
            logger.warning(
                "combined_originals_structure_check_failed",
                extra={"error": str(e)},
            )

        logger.info(
            "combined_originals_generated_successfully",
            extra={
                "user_id": getattr(processor, "_current_user_id", "unknown"),
                "week": week_of,
                "output_path": str(output_path.absolute()),
                "size_bytes": (
                    output_path.stat().st_size
                    if output_path.exists()
                    else 0
                ),
                "successful_slots": successful_slots,
                "failed_slots": failed_slots,
            },
        )
        return str(output_path)

    except Exception as e:
        logger.error(f"Error rendering originals doc: {e}")
        traceback.print_exc()
        return None
    finally:
        for tf in temp_files:
            try:
                if os.path.exists(tf):
                    os.remove(tf)
            except Exception:
                pass
