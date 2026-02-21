"""
Combined original DOCX generation and file-group processing for batch processor.
Fetches original plans from DB, renders per-slot DOCX, merges; processes slot groups by file.
Extracted from orchestrator.
"""

import asyncio
import os
import traceback
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from backend.config import settings
from backend.telemetry import logger
from tools.batch_processor_pkg.context import SlotProcessingContext


async def process_file_group(
    processor: Any,
    file_path: Optional[str],
    group: List[Tuple[int, Dict[str, Any]]],
    week_of: str,
    week_folder_path: Optional[str],
    user_base_path: Optional[str],
    plan_id: Optional[str],
    semaphore: asyncio.Semaphore,
) -> List[SlotProcessingContext]:
    """Process a group of slots that share the same source file."""
    contexts = []
    user_id = getattr(processor, "_current_user_id", "unknown")

    if not file_path:
        for i, slot in group:
            ctx = SlotProcessingContext(
                slot=slot, slot_index=i, total_slots=len(group)
            )
            ctx.error = "No primary teacher file found."
            contexts.append(ctx)
        return contexts

    if file_path not in processor._file_locks:
        processor._file_locks[file_path] = asyncio.Lock()

    async with processor._file_locks[file_path]:
        remaining_group = []

        path_obj = Path(file_path)
        current_mtime = 0
        if path_obj.exists():
            stat = path_obj.stat()
            current_mtime = stat.st_mtime

        for i, slot in group:
            context = SlotProcessingContext(
                slot=slot,
                slot_index=i,
                total_slots=len(group),
                primary_file=file_path,
            )

            db_record = await asyncio.to_thread(
                processor.db.get_original_lesson_plan,
                user_id,
                week_of,
                slot["slot_number"],
            )

            if db_record and db_record.source_file_path == file_path:
                if path_obj.exists():
                    if db_record.extracted_at.timestamp() > (current_mtime + 2):
                        logger.info(
                            f"DB Cache hit for slot {slot['slot_number']} ({slot['subject']})"
                        )
                        context.extracted_content = db_record.full_text
                        context.available_days = db_record.available_days
                        context.cache_hit = True

                        cached_hyperlinks = []
                        if db_record.content_json:
                            context.slot["_extracted_images"] = (
                                db_record.content_json.get(
                                    "_extracted_images", []
                                )
                            )
                            cached_hyperlinks = db_record.content_json.get(
                                "_extracted_hyperlinks", []
                            )
                            context.slot["_extracted_hyperlinks"] = (
                                cached_hyperlinks
                            )
                            context.no_school_days = (
                                db_record.content_json.get(
                                    "no_school_days", []
                                )
                            )

                        if not cached_hyperlinks:
                            remaining_group.append((i, slot, context))
                        else:
                            contexts.append(context)
                        continue

            remaining_group.append((i, slot, context))

        if not remaining_group:
            return contexts

        try:
            async with semaphore:
                logger.info(f"Opening DOCX for parsing group: {file_path}")
                parser = await processor._open_docx_with_retry(
                    file_path, plan_id=plan_id
                )

                for i, slot, context in remaining_group:
                    try:
                        primary_first = slot.get(
                            "primary_teacher_first_name", ""
                        ).strip()
                        primary_last = slot.get(
                            "primary_teacher_last_name", ""
                        ).strip()
                        teacher_name = (
                            f"{primary_first} {primary_last}".strip()
                            or slot.get("primary_teacher_name", "").strip()
                        )

                        actual_slot_num = await asyncio.to_thread(
                            parser.find_slot_by_subject,
                            slot["subject"],
                            teacher_name,
                            slot.get("homeroom"),
                            slot.get("grade"),
                        )

                        content_data = await asyncio.to_thread(
                            parser.extract_subject_content_for_slot,
                            actual_slot_num,
                            slot["subject"],
                            None,
                            False,
                        )

                        if (
                            "_extracted_hyperlinks" in context.slot
                            and context.slot["_extracted_hyperlinks"]
                        ):
                            hyperlinks = context.slot["_extracted_hyperlinks"]
                        else:
                            hyperlinks = await asyncio.to_thread(
                                parser.extract_hyperlinks_for_slot,
                                actual_slot_num,
                            )

                        images = await asyncio.to_thread(
                            parser.extract_images_for_slot, actual_slot_num
                        )

                        logger.info(
                            "parallel_file_group_media_extracted",
                            extra={
                                "slot": slot["slot_number"],
                                "subject": slot["subject"],
                                "images_count": len(images),
                                "hyperlinks_count": len(hyperlinks),
                            },
                        )

                        initial_hyperlink_count = len(hyperlinks)
                        if parser.is_no_school_day():
                            hyperlinks = []
                        elif "no_school_days" in content_data:
                            no_school_days_normalized = {
                                day.lower().strip()
                                for day in content_data["no_school_days"]
                            }
                            hyperlinks = [
                                h
                                for h in hyperlinks
                                if not h.get("day_hint")
                                or h.get("day_hint", "").lower().strip()
                                not in no_school_days_normalized
                            ]
                            filtered_count = (
                                initial_hyperlink_count - len(hyperlinks)
                            )
                            if filtered_count > 0:
                                logger.info(
                                    "parallel_hyperlinks_filtered_no_school",
                                    extra={
                                        "slot": slot["slot_number"],
                                        "subject": slot["subject"],
                                        "filtered_count": filtered_count,
                                        "remaining_count": len(hyperlinks),
                                    },
                                )

                        context.extracted_content = content_data.get(
                            "full_text", ""
                        )
                        context.available_days = content_data.get(
                            "available_days", []
                        )
                        context.no_school_days = content_data.get(
                            "no_school_days", []
                        )

                        context.slot["_extracted_images"] = images
                        context.slot["_extracted_hyperlinks"] = hyperlinks

                        content_data["_extracted_images"] = images
                        content_data["_extracted_hyperlinks"] = hyperlinks
                        hyperlink_texts = [
                            h["text"] for h in hyperlinks if h.get("text")
                        ]
                        if hyperlink_texts:
                            preserve_instruction = (
                                "\n\nIMPORTANT: Preserve these exact phrases in your output "
                                f"(they are hyperlinks): {', '.join(hyperlink_texts[:20])}"
                            )
                            context.extracted_content += preserve_instruction

                        metadata = content_data.get("metadata", {})
                        primary_teacher_name = (
                            metadata.get("primary_teacher_name")
                            or slot.get("primary_teacher_name")
                            or slot.get("teacher_name")
                        )

                        structured_days = {}
                        if "table_content" in content_data:
                            for day, day_data in content_data[
                                "table_content"
                            ].items():
                                day_lower = day.lower().strip()
                                if day_lower in [
                                    "monday",
                                    "tuesday",
                                    "wednesday",
                                    "thursday",
                                    "friday",
                                ]:
                                    day_links = [
                                        h
                                        for h in hyperlinks
                                        if h.get("day", "").lower().strip()
                                        == day_lower
                                    ]
                                    structured_days[f"{day_lower}_content"] = (
                                        processor._map_day_content_to_schema(
                                            day_data,
                                            slot,
                                            day_hyperlinks=day_links,
                                        )
                                    )

                        plan_data = {
                            "id": f"orig_{uuid.uuid4()}",
                            "user_id": user_id,
                            "week_of": week_of,
                            "slot_number": slot["slot_number"],
                            "subject": slot["subject"],
                            "grade": metadata.get("grade")
                            or slot.get("grade", "Unknown"),
                            "homeroom": metadata.get("homeroom")
                            or slot.get("homeroom"),
                            "source_file_path": file_path,
                            "source_file_name": path_obj.name,
                            "primary_teacher_name": primary_teacher_name,
                            "extracted_at": datetime.utcnow(),
                            "content_json": content_data,
                            "full_text": context.extracted_content,
                            "available_days": context.available_days,
                            "has_no_school": parser.is_no_school_day(),
                            "status": "extracted",
                            **structured_days,
                        }
                        await asyncio.to_thread(
                            processor.db.create_original_lesson_plan,
                            plan_data,
                        )
                        contexts.append(context)

                    except Exception as e:
                        logger.error(
                            f"Error extracting slot {slot['slot_number']}: {e}"
                        )
                        context.error = str(e)
                        contexts.append(context)

        except Exception as e:
            logger.error(f"Failed to parse group file {file_path}: {e}")
            for i, slot, context in remaining_group:
                context.error = f"Failed to parse DOCX: {str(e)}"
                contexts.append(context)

    return contexts


async def generate_combined_original_docx(
    processor: Any,
    user_id: str,
    week_of: str,
    plan_id: str,
    week_folder_path: Optional[str] = None,
    get_file_manager_fn: Optional[Callable[..., Any]] = None,
) -> Optional[str]:
    """
    Generates a combined DOCX of all original plans for the week using
    structured schema data. Reuses the standard DOCXRenderer for visual consistency.
    """
    logger.info(
        f"Generating combined original DOCX for {user_id}, week {week_of}"
    )

    plans = await asyncio.to_thread(
        processor.db.get_original_lesson_plans_for_week, user_id, week_of
    )
    if not plans:
        logger.warning(
            "combined_originals_skipped_no_plans",
            extra={"user_id": user_id, "week": week_of},
        )
        return None

    plans.sort(key=lambda x: x.slot_number)

    logger.info(
        "combined_originals_plans_fetched",
        extra={
            "total_plans": len(plans),
            "plans": [
                {
                    "slot_number": p.slot_number,
                    "subject": p.subject,
                    "grade": p.grade,
                    "source_file_name": p.source_file_name,
                    "primary_teacher_name": p.primary_teacher_name,
                }
                for p in plans
            ],
        },
    )

    if get_file_manager_fn is None:
        from tools.batch_processor import get_file_manager as _get_fm
        get_file_manager_fn = _get_fm

    def _render_originals():
        temp_files = []
        rendering_results = {}
        try:
            from docx import Document

            from tools.docx_renderer import DOCXRenderer
            from tools.docx_utils import normalize_styles_from_master

            file_mgr = get_file_manager_fn(
                base_path=getattr(processor, "_user_base_path", None)
            )
            week_folder = (
                Path(week_folder_path)
                if week_folder_path
                else file_mgr.get_week_folder(week_of)
            )
            originals_dir = week_folder / "originals"
            originals_dir.mkdir(parents=True, exist_ok=True)

            safe_week = (
                week_of.replace("/", "-")
                .replace("\\", "-")
                .replace(" ", "_")
            )
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"combined_originals_{safe_week}_{timestamp}.docx"
            output_path = originals_dir / filename

            logger.info(
                "combined_originals_processing_started",
                extra={"total_plans": len(plans)},
            )

            renderer = DOCXRenderer(settings.DOCX_TEMPLATE_PATH)
            renderer.is_originals = True

            style_master = Document(settings.DOCX_TEMPLATE_PATH)

            unique_slot_files = set()
            deduplicated_plans = []
            removed_plans = []

            for plan in plans:
                source_key = getattr(plan, "source_file_path", None)
                slot_key = plan.slot_number
                subject_key = plan.subject
                dedup_key = (source_key, slot_key, subject_key)

                if dedup_key not in unique_slot_files:
                    unique_slot_files.add(dedup_key)
                    deduplicated_plans.append(plan)
                else:
                    removed_plans.append({
                        "slot_number": plan.slot_number,
                        "subject": plan.subject,
                        "source_file_name": plan.source_file_name,
                    })

            logger.info(
                "combined_originals_deduplication",
                extra={
                    "original_count": len(plans),
                    "deduplicated_count": len(deduplicated_plans),
                    "unique_keys_count": len(unique_slot_files),
                    "removed_plans": removed_plans,
                    "deduplicated_plans": [
                        {
                            "slot_number": p.slot_number,
                            "subject": p.subject,
                            "source_file_name": p.source_file_name,
                        }
                        for p in deduplicated_plans
                    ],
                },
            )

            original_slot_subjects = {
                (p.slot_number, p.subject) for p in plans
            }
            dedup_slot_subjects = {
                (p.slot_number, p.subject) for p in deduplicated_plans
            }
            missing_combinations = (
                original_slot_subjects - dedup_slot_subjects
            )
            if missing_combinations:
                logger.warning(
                    "combined_originals_missing_after_dedup",
                    extra={
                        "missing_combinations": [
                            {"slot_number": s, "subject": subj}
                            for s, subj in missing_combinations
                        ]
                    },
                )

            for plan in deduplicated_plans:
                temp_filename = f"_temp_orig_slot{plan.slot_number}_{plan.subject.replace('/', '_')}.docx"
                temp_path = str(originals_dir / temp_filename)

                try:
                    logger.info(
                        "combined_originals_rendering_plan",
                        extra={
                            "slot_number": plan.slot_number,
                            "subject": plan.subject,
                            "source_file_name": plan.source_file_name,
                            "temp_path": temp_path,
                        },
                    )

                    slot_json = processor._convert_originals_to_json([plan])

                    if slot_json.get("metadata", {}).get(
                        "slot_number"
                    ) != plan.slot_number:
                        logger.warning(
                            "combined_originals_slot_number_mismatch",
                            extra={
                                "expected": plan.slot_number,
                                "actual": slot_json.get("metadata", {}).get(
                                    "slot_number"
                                ),
                            },
                        )
                    if slot_json.get("metadata", {}).get(
                        "subject"
                    ) != plan.subject:
                        logger.warning(
                            "combined_originals_subject_mismatch",
                            extra={
                                "expected": plan.subject,
                                "actual": slot_json.get("metadata", {}).get(
                                    "subject"
                                ),
                            },
                        )

                    render_result = renderer.render(
                        slot_json,
                        temp_path,
                        plan_id=plan_id,
                        skip_fallback_sections=True,
                    )
                    success, unplaced_hl, unplaced_img = (
                        render_result
                        if isinstance(render_result, tuple)
                        else (render_result, [], [])
                    )

                    if success:
                        sub_doc = Document(temp_path)
                        from tools.docx_utils import (
                            remove_headers_footers,
                            strip_custom_styles,
                            strip_problematic_elements,
                            strip_sections,
                        )

                        strip_problematic_elements(sub_doc)
                        remove_headers_footers(sub_doc)
                        strip_sections(sub_doc)
                        strip_custom_styles(sub_doc, style_master)

                        normalize_styles_from_master(style_master, sub_doc)
                        if hasattr(sub_doc, "_normalized_stream"):
                            sub_doc._normalized_stream.seek(0)
                            sub_doc = Document(sub_doc._normalized_stream)

                        sub_doc.save(temp_path)
                        temp_files.append(temp_path)
                        rendering_results[plan.slot_number] = {
                            "success": True,
                            "subject": plan.subject,
                            "temp_path": temp_path,
                        }
                        logger.info(
                            "combined_originals_plan_rendered_successfully",
                            extra={
                                "slot_number": plan.slot_number,
                                "subject": plan.subject,
                            },
                        )
                    else:
                        rendering_results[plan.slot_number] = {
                            "success": False,
                            "subject": plan.subject,
                            "error": "Render returned False",
                        }
                        logger.error(
                            "combined_originals_plan_render_failed",
                            extra={
                                "slot_number": plan.slot_number,
                                "subject": plan.subject,
                                "unplaced_hyperlinks": len(unplaced_hl),
                                "unplaced_images": len(unplaced_img),
                            },
                        )
                except Exception as e:
                    rendering_results[plan.slot_number] = {
                        "success": False,
                        "subject": plan.subject,
                        "error": str(e),
                    }
                    logger.error(
                        "combined_originals_plan_render_exception",
                        extra={
                            "slot_number": plan.slot_number,
                            "subject": plan.subject,
                            "error": str(e),
                        },
                        exc_info=True,
                    )

            successful_slots = [
                slot_num
                for slot_num, result in rendering_results.items()
                if result.get("success")
            ]
            failed_slots = [
                slot_num
                for slot_num, result in rendering_results.items()
                if not result.get("success")
            ]

            logger.info(
                "combined_originals_rendering_summary",
                extra={
                    "total_plans": len(deduplicated_plans),
                    "successful_renders": len(successful_slots),
                    "failed_renders": len(failed_slots),
                    "successful_slots": successful_slots,
                    "failed_slots": [
                        {
                            "slot_number": slot_num,
                            "subject": rendering_results[slot_num].get(
                                "subject"
                            ),
                            "error": rendering_results[slot_num].get(
                                "error"
                            ),
                        }
                        for slot_num in failed_slots
                    ],
                },
            )

            if not temp_files:
                logger.warning(
                    "combined_originals_no_files_generated",
                    extra={
                        "total_plans": len(deduplicated_plans),
                        "rendering_results": rendering_results,
                    },
                )
                return None

            expected_slots = {
                p.slot_number for p in deduplicated_plans
            }
            rendered_slots = {
                p.slot_number
                for p in deduplicated_plans
                if rendering_results.get(p.slot_number, {}).get("success")
            }
            missing_slots = expected_slots - rendered_slots
            if missing_slots:
                logger.warning(
                    "combined_originals_missing_slots",
                    extra={
                        "missing_slots": [
                            {
                                "slot_number": slot_num,
                                "subject": next(
                                    (
                                        p.subject
                                        for p in deduplicated_plans
                                        if p.slot_number == slot_num
                                    ),
                                    "Unknown",
                                ),
                            }
                            for slot_num in missing_slots
                        ]
                    },
                )

            logger.info(
                "combined_originals_merging_files",
                extra={
                    "file_count": len(temp_files),
                    "file_paths": [Path(f).name for f in temp_files],
                },
            )
            processor._merge_docx_files(
                temp_files,
                str(output_path),
                master_template_path=settings.DOCX_TEMPLATE_PATH,
            )

            try:
                merged_doc = Document(str(output_path))
                table_count = len(merged_doc.tables)
                logger.info(
                    "combined_originals_merged_structure",
                    extra={
                        "table_count": table_count,
                        "estimated_slots": table_count // 2,
                    },
                )
            except Exception as e:
                logger.warning(
                    "combined_originals_structure_check_failed",
                    extra={"error": str(e)},
                )

            logger.info(
                "combined_originals_generated_successfully",
                extra={
                    "user_id": user_id,
                    "week": week_of,
                    "output_path": str(output_path.absolute()),
                    "size_bytes": (
                        output_path.stat().st_size
                        if output_path.exists()
                        else 0
                    ),
                    "successful_slots": successful_slots,
                    "failed_slots": failed_slots,
                },
            )
            return str(output_path)

        except Exception as e:
            logger.error(f"Error rendering originals doc: {e}")
            traceback.print_exc()
            return None
        finally:
            for tf in temp_files:
                try:
                    if os.path.exists(tf):
                        os.remove(tf)
                except Exception:
                    pass

    return await asyncio.to_thread(_render_originals)
