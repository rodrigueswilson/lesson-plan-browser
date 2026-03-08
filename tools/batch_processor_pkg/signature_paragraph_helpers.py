"""
Helpers for building signature paragraphs: font info and date section.
Used by signatures.py when modifying teacher signature cells.
"""

from typing import Optional

from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from tools.docx_renderer.style import sanitize_xml_text


def get_paragraph_font_info(para: Paragraph) -> tuple[Optional[object], Optional[str]]:
    """Return (font_size, font_name) from the first run that has either."""
    font_size = None
    font_name = None
    for run in para.runs:
        if run.font.size:
            font_size = run.font.size
        if run.font.name:
            font_name = run.font.name
        if font_size is not None and font_name is not None:
            break
    return (font_size, font_name)


def add_date_section_to_paragraph(
    para: Paragraph,
    date_value: Optional[str],
    font_size: Optional[object],
    font_name: Optional[str],
) -> None:
    """Add tab stop, tab, 'Date: ' label, and date value or placeholder to paragraph."""
    date_tab_position = Inches(5.5)
    para.paragraph_format.tab_stops.add_tab_stop(
        date_tab_position, WD_TAB_ALIGNMENT.LEFT
    )
    para.add_run("\t")

    date_label_run = para.add_run("Date: ")
    date_label_run.font.bold = True
    if font_size:
        date_label_run.font.size = font_size
    if font_name:
        date_label_run.font.name = font_name

    if date_value:
        date_value_run = para.add_run(sanitize_xml_text(date_value))
        date_value_run.font.underline = True
        if font_size:
            date_value_run.font.size = font_size
        if font_name:
            date_value_run.font.name = font_name
    else:
        placeholder_run = para.add_run("__________________")
        if font_size:
            placeholder_run.font.size = font_size
        if font_name:
            placeholder_run.font.name = font_name
