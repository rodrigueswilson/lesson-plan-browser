"""
Test TableStructureDetector on real files.
"""

import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.table_structure import TableStructureDetector

TEST_FILE = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Davies Lesson Plans.docx'

def test_structure_detector():
    print("="*80)
    print("TESTING TABLE STRUCTURE DETECTOR")
    print("="*80)
    print()
    
    doc = Document(TEST_FILE)
    detector = TableStructureDetector()
    
    print(f"File: {Path(TEST_FILE).name}")
    print(f"Total tables: {len(doc.tables)}")
    print()
    
    for table_idx, table in enumerate(doc.tables):
        print(f"{'─'*80}")
        print(f"TABLE {table_idx}:")
        print(f"{'─'*80}")
        
        structure = detector.detect_structure(table)
        
        print(f"Structure type: {structure.structure_type}")
        print(f"Dimensions: {structure.num_rows}x{structure.num_cols}")
        print(f"Row offset: {structure.row_offset}")
        print(f"Has day row: {structure.has_day_row}")
        print()
        
        print(f"Row label map ({len(structure.row_label_map)} entries):")
        for label, idx in sorted(structure.row_label_map.items(), key=lambda x: x[1]):
            print(f"  '{label}' → row {idx}")
        print()
        
        print(f"Column header map ({len(structure.col_header_map)} entries):")
        for day, idx in sorted(structure.col_header_map.items(), key=lambda x: x[1]):
            print(f"  '{day}' → col {idx}")
        print()
        
        # Test lookups
        print("Testing lookups:")
        test_labels = ['Objective:', 'objective', 'OBJECTIVE', 'Tailored Instruction:', 'instruction']
        for label in test_labels:
            row_idx = structure.get_row_index(label)
            print(f"  get_row_index('{label}') → {row_idx}")
        
        test_days = ['monday', 'MONDAY', 'Monday', 'friday']
        for day in test_days:
            col_idx = structure.get_col_index(day)
            print(f"  get_col_index('{day}') → {col_idx}")
        
        print()
    
    print("="*80)
    print("TEST COMPLETE")
    print("="*80)

if __name__ == '__main__':
    test_structure_detector()
