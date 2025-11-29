"""Utility to inspect inline vs. fallback hyperlinks in a DOCX output."""

import argparse
import sys
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.oxml.ns import qn


def extract_links(doc: Document) -> List[Dict]:
    """Extract hyperlinks with location metadata and resolved URLs."""

    links: List[Dict] = []
    rels = doc.part.rels

    # Tables first (primary placement area)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for hyperlink in para._element.xpath(".//w:hyperlink"):
                        rel_id = hyperlink.get(qn("r:id"))
                        url = rels.get(rel_id).target_ref if rel_id in rels else None
                        text = "".join(
                            node.text for node in hyperlink.xpath(".//w:t") if node.text
                        )

                        if text and url:
                            links.append(
                                {
                                    "text": text,
                                    "url": url,
                                    "location_type": "table",
                                    "table_idx": table_idx,
                                    "row_idx": row_idx,
                                    "cell_idx": cell_idx,
                                    "location": f"T{table_idx}R{row_idx}C{cell_idx}",
                                }
                            )

    # Paragraphs (fallback sections)
    for para_idx, para in enumerate(doc.paragraphs):
        for hyperlink in para._element.xpath(".//w:hyperlink"):
            rel_id = hyperlink.get(qn("r:id"))
            url = rels.get(rel_id).target_ref if rel_id in rels else None
            text = "".join(node.text for node in hyperlink.xpath(".//w:t") if node.text)

            if text and url:
                links.append(
                    {
                        "text": text,
                        "url": url,
                        "location_type": "paragraph",
                        "paragraph_idx": para_idx,
                        "paragraph_text": para.text[:100],
                    }
                )

    return links


def categorize_links(doc: Document, links: List[Dict]) -> Dict[str, List[Dict]]:
    """Split links into inline vs. fallback using table/paragraph heuristics."""

    inline: List[Dict] = []
    fallback: List[Dict] = []

    # Detect if a fallback table exists
    fallback_table_idx = None
    for table_idx, table in enumerate(doc.tables):
        if table.rows and table.rows[0].cells:
            header = table.rows[0].cells[0].text
            if "Referenced Links" in header or "Referenced Media" in header:
                fallback_table_idx = table_idx
                break

    # Detect fallback paragraph start
    fallback_paragraph_idx = None
    for idx, para in enumerate(doc.paragraphs):
        if "Referenced Links" in para.text or "Referenced Media" in para.text:
            fallback_paragraph_idx = idx
            break

    for link in links:
        if link["location_type"] == "table":
            if (
                fallback_table_idx is not None
                and link["table_idx"] >= fallback_table_idx
            ):
                fallback.append(link)
            else:
                inline.append(link)
        else:  # paragraph
            if (
                fallback_paragraph_idx is not None
                and link["paragraph_idx"] >= fallback_paragraph_idx
            ):
                fallback.append(link)
            else:
                # Conservative: assume paragraph links near end are fallback
                fallback.append(link)

    return {"inline": inline, "fallback": fallback}


def main(path: Path) -> int:
    if not path.exists():
        print(f"❌ File not found: {path}")
        return 1

    doc = Document(path)
    links = extract_links(doc)
    categorized = categorize_links(doc, links)

    print(f"Analyzing: {path.name}")
    print(f"Total hyperlinks: {len(links)}")
    print(f"Inline (tables before fallback section): {len(categorized['inline'])}")
    print(f"Fallback (Referenced Links section): {len(categorized['fallback'])}\n")

    if categorized["fallback"]:
        print("Fallback links:")
        for idx, link in enumerate(categorized["fallback"], 1):
            snippet = link.get("paragraph_text") or link.get("location", "")
            print(
                f"  [{idx}] {link['text'][:80]}\n      URL: {link['url']}\n      Location: {snippet}"
            )

    return 0


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Inspect inline vs fallback hyperlinks in DOCX output"
    )
    parser.add_argument("docx_path", type=str, help="Path to the DOCX file to inspect")
    args = parser.parse_args()

    sys.exit(main(Path(args.docx_path)))
