"""
Link extraction and categorization for threshold validation v2.
Extracts hyperlinks from output DOCX with URL resolution and categorizes as inline vs fallback.
"""

from typing import Dict, List, Tuple
from docx import Document
from docx.oxml.ns import qn


def extract_output_links_with_urls(doc: Document) -> List[Dict]:
    """Extract hyperlinks from document with proper URL resolution."""
    links = []
    rels = doc.part.rels

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    hyperlinks = para._element.xpath('.//w:hyperlink')
                    for hyperlink in hyperlinks:
                        r_id = hyperlink.get(qn('r:id'))
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        url = None
                        if r_id and r_id in rels:
                            url = rels[r_id].target_ref
                        if text and url:
                            links.append({
                                'text': text,
                                'url': url,
                                'location_type': 'table',
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'location': f"T{table_idx}R{row_idx}C{cell_idx}"
                            })

    for para_idx, para in enumerate(doc.paragraphs):
        hyperlinks = para._element.xpath('.//w:hyperlink')
        for hyperlink in hyperlinks:
            r_id = hyperlink.get(qn('r:id'))
            text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
            url = None
            if r_id and r_id in rels:
                url = rels[r_id].target_ref
            if text and url:
                links.append({
                    'text': text,
                    'url': url,
                    'location_type': 'paragraph',
                    'paragraph_idx': para_idx,
                    'paragraph_text': para.text[:100]
                })

    return links


def categorize_links(doc: Document, output_links: List[Dict]) -> Tuple[List, List]:
    """Categorize links as inline (in table) or fallback (in paragraphs / referenced links table)."""
    inline = []
    fallback = []

    fallback_table_idx = None
    for table_idx, table in enumerate(doc.tables):
        if table.rows and table.rows[0].cells:
            first_cell_text = table.rows[0].cells[0].text
            if 'Referenced Links' in first_cell_text or 'Referenced Media' in first_cell_text:
                fallback_table_idx = table_idx
                break

    for link in output_links:
        if link['location_type'] == 'table':
            if fallback_table_idx is not None and link['table_idx'] >= fallback_table_idx:
                fallback.append(link)
            else:
                inline.append(link)
        else:
            fallback.append(link)

    return inline, fallback
