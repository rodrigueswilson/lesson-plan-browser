"""Search for ANY occurrence of 'Referenced' in the entire document."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)

print("Searching for 'Referenced' in all text...\n")

# Search in all paragraphs
print("In paragraphs:")
for para_idx, para in enumerate(doc.paragraphs):
    if 'Referenced' in para.text or 'Links:' in para.text:
        print(f"  Paragraph {para_idx}: {para.text[:100]}")

# Search in all table cells
print("\nIn tables:")
for table_idx, table in enumerate(doc.tables):
    found = False
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            if 'Referenced' in cell_text or 'Links:' in cell_text:
                if not found:
                    print(f"\nTable {table_idx}:")
                    found = True
                print(f"  Row {row_idx}, Cell {cell_idx}: {cell_text[:100]}")
                
                # Count links in this cell
                link_count = 0
                for para in cell.paragraphs:
                    hyperlinks = para._element.xpath('.//w:hyperlink')
                    link_count += len(hyperlinks)
                if link_count > 0:
                    print(f"    → Contains {link_count} hyperlinks")

print("\n\nSearching for bullet points (•) which might indicate a list:")
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if '•' in cell.text:
                # Count links
                link_count = 0
                for para in cell.paragraphs:
                    hyperlinks = para._element.xpath('.//w:hyperlink')
                    link_count += len(hyperlinks)
                
                if link_count > 0:
                    print(f"\nTable {table_idx}, Row {row_idx}, Cell {cell_idx}:")
                    print(f"  {link_count} links in bulleted list")
                    print(f"  First 100 chars: {cell.text[:100]}")
