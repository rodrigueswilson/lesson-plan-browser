"""
Utility functions for DOCX manipulation.
Follows DRY principle - reusable across modules.
"""

import logging

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, Dict, Any
from io import BytesIO

from tools.docx_table_utils import (
    get_table_info,
    normalize_all_tables,
    normalize_table_column_widths,
)

try:
    import structlog
    _log = structlog.get_logger(__name__)
except ImportError:
    _log = logging.getLogger(__name__)

__all__ = [
    "diagnose_style_conflicts",
    "ensure_hyperlink_style",
    "get_table_info",
    "normalize_all_tables",
    "normalize_styles_from_master",
    "normalize_styles_via_file",
    "normalize_table_column_widths",
    "remove_headers_footers",
    "strip_custom_styles",
    "strip_problematic_elements",
    "strip_sections",
]


def strip_custom_styles(doc: Document, master_doc: Document) -> None:
    """
    Map all paragraphs and runs in doc to standard styles if their 
    current style is not present in master_doc.
    This prevents "unreadable content" errors where the document 
    references a style ID that no longer exists after normalization.
    """
    master_style_ids = {style.style_id for style in master_doc.styles}
    
    # 1. Map paragraphs
    for p in doc.paragraphs:
        if p.style.style_id not in master_style_ids:
            try:
                p.style = doc.styles['Normal']
            except:
                pass
                
    # 2. Map runs (character styles)
    for p in doc.paragraphs:
        for r in p.runs:
            if r.style and r.style.style_id not in master_style_ids:
                try:
                    r.style = doc.styles['Default Paragraph Font']
                except:
                    r.style = None
                    
    # 3. Map tables
    for t in doc.tables:
        if t.style and t.style.style_id not in master_style_ids:
            try:
                t.style = doc.styles['Table Grid']
            except:
                try:
                    t.style = doc.styles['Normal Table']
                except:
                    pass

def strip_problematic_elements(doc: Document) -> None:
    """
    Remove bookmarks, comments, and other elements that often cause 
    "unreadable content" errors after merging with docxcompose.
    """
    from docx.oxml.ns import qn
    
    # Use structlog if available
    _log.debug("stripping_problematic_elements_started")
    
    # 1. Remove all bookmarks (w:bookmarkStart, w:bookmarkEnd)
    # These often have duplicate IDs or names after merging
    count_bookmarks = 0
    for element in doc.element.xpath('.//w:bookmarkStart | .//w:bookmarkEnd'):
        element.getparent().remove(element)
        count_bookmarks += 1
        
    # 2. Remove all comments (w:commentReference, w:commentRangeStart, w:commentRangeEnd)
    count_comments = 0
    for element in doc.element.xpath('.//w:commentReference | .//w:commentRangeStart | .//w:commentRangeEnd'):
        element.getparent().remove(element)
        count_comments += 1
        
    # 3. Remove all proofing errors (w:proofErr)
    count_proof = 0
    for element in doc.element.xpath('.//w:proofErr'):
        element.getparent().remove(element)
        count_proof += 1

    if count_bookmarks > 0 or count_comments > 0 or count_proof > 0:
        _log.info("stripped_elements", extra={
            "bookmarks": count_bookmarks,
            "comments": count_comments,
            "proofing_errors": count_proof
        })
    else:
        _log.debug("no_problematic_elements_found_to_strip")

def remove_headers_footers(doc: Document) -> None:
    """
    Remove all headers and footers from a document to prevent 
    conflicts during merging.
    """
    for section in doc.sections:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        # Also clear the content if possible
        try:
            for p in section.header.paragraphs:
                p.text = ""
            for p in section.footer.paragraphs:
                p.text = ""
        except:
            pass

def ensure_hyperlink_style(doc: Document) -> None:
    """
    Ensure the "Hyperlink" style exists in the document.
    This is critical because Word often expects this style for hyperlinks,
    and removing it during style normalization causes "unreadable content" errors.
    """
    if "Hyperlink" not in doc.styles:
        try:
            from docx.enum.style import WD_STYLE_TYPE
            from docx.shared import Pt, RGBColor
            style = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
            style.font.name = "Times New Roman"
            style.font.size = Pt(8)
            style.font.color.rgb = RGBColor(0, 0, 255)
            style.font.underline = True
        except Exception:
            pass

def strip_sections(doc: Document) -> None:
    """
    Remove all section properties from the document body to prevent 
    Word from creating multiple sections after merging.
    The document will inherit the section properties of the master document.
    """
    from docx.oxml.ns import qn
    
    # Remove sectPr from the body
    body = doc.element.body
    sectPrs = body.findall(qn('w:sectPr'))
    for sectPr in sectPrs:
        # Don't remove the last one if it's the only one, as Word needs it
        # But we want to remove any that are NOT the last one (which are section breaks)
        pass
    
    # Actually, the simplest way is to remove all sectPr from paragraphs
    for p in body.findall(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                pPr.remove(sectPr)

def normalize_styles_via_file(master_doc: Document, target_doc: Document) -> Optional[BytesIO]:
    """
    Replace styles, numbering, font table, and docProps (custom.xml, core.xml) by using
    temporary files and zipfile. Replacing docProps avoids "Word found unreadable content"
    when merged DOCX have corrupt or conflicting custom properties.
    
    This is the most reliable method as it directly replaces the XML files in the DOCX package.
    Only parts present in the master are replaced; missing parts are left unchanged in target.
    
    Returns a BytesIO stream containing the normalized document, or None if failed.
    The caller should reload the document from this stream.
    """
    import tempfile
    import zipfile
    from pathlib import Path
    from io import BytesIO
    
    _log.warning("file_based_normalization_started")
    
    try:
        # Save both documents to temporary files
        _log.debug("saving_documents_to_temp_files")
        with tempfile.TemporaryDirectory() as temp_dir:
            master_path = Path(temp_dir) / "master.docx"
            target_path = Path(temp_dir) / "target.docx"
            output_path = Path(temp_dir) / "target_normalized.docx"
            
            master_doc.save(str(master_path))
            target_doc.save(str(target_path))
            _log.debug("documents_saved_to_temp")
            
            # Files to replace for consistency (styles/numbering/font + docProps to avoid "unreadable content")
            files_to_replace = [
                'word/styles.xml',
                'word/numbering.xml',
                'word/fontTable.xml',
                'docProps/custom.xml',
                'docProps/core.xml',
            ]
            
            replacement_data = {}
            
            # Extract replacement files from master
            with zipfile.ZipFile(master_path, 'r') as master_zip:
                master_files = master_zip.namelist()
                for xml_file in files_to_replace:
                    if xml_file in master_files:
                        replacement_data[xml_file] = master_zip.read(xml_file)
                        _log.debug("extracted_from_master", extra={"file": xml_file})
            
            missing = [f for f in files_to_replace if f not in replacement_data]
            if missing:
                _log.debug(
                    "replacement_files_not_in_master",
                    extra={"missing": missing, "note": "only parts present in template are replaced"},
                )
            if not replacement_data:
                _log.warning("no_replacement_files_found_in_master")
                return None
            
            # Replace files in target
            with zipfile.ZipFile(target_path, 'r') as target_zip:
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as output_zip:
                    # Copy all files from target except those being replaced
                    files_copied = 0
                    for item in target_zip.infolist():
                        if item.filename not in replacement_data:
                            data = target_zip.read(item.filename)
                            output_zip.writestr(item, data)
                            files_copied += 1
                    
                    # Add replacement files
                    for filename, data in replacement_data.items():
                        output_zip.writestr(filename, data)
                    
                    _log.debug("files_replaced_in_zip", extra={
                        "files_copied": files_copied,
                        "files_replaced": list(replacement_data.keys()),
                    })
            
            # Read the normalized file into BytesIO
            with open(output_path, 'rb') as f:
                normalized_data = BytesIO(f.read())
            
            _log.warning("styles_and_numbering_replaced_via_file", extra={
                "replaced_files": list(replacement_data.keys()),
                "success": True,
            })
            
            return normalized_data
            
    except Exception as e:
        _log.error("file_based_style_replacement_failed", extra={
            "error": str(e),
            "error_type": type(e).__name__
        }, exc_info=True)
        return None


def normalize_styles_from_master(master_doc: Document, target_doc: Document) -> None:
    """
    Replace entire styles part in target_doc with master_doc's styles.
    This prevents style conflicts when merging documents with docxcompose.
    
    This is more aggressive than normalizing individual styles - it completely
    replaces the styles part to ensure 100% consistency and prevent "Styles 1" errors.
    
    Args:
        master_doc: Document whose styles will be used as the source of truth
        target_doc: Document whose styles will be replaced with master_doc's styles
        
    Example:
        >>> from docx import Document
        >>> master = Document('template.docx')
        >>> target = Document('to_merge.docx')
        >>> normalize_styles_from_master(master, target)
        >>> # Now target can be safely merged with master
    """
    # Log function entry immediately
    _log.info("normalize_styles_from_master_called")
    
    try:
        # Log initial state - use INFO level to ensure visibility
        master_style_count = len(master_doc.styles)
        target_style_count = len(target_doc.styles)
        _log.info("style_normalization_start", extra={
            "master_style_count": master_style_count,
            "target_style_count": target_style_count
        })
        
        # Try multiple methods to access styles part
        master_styles_part = None
        target_styles_part = None
        
        # Method 1: Direct attribute access (using private _styles_part which is more reliable)
        try:
            master_styles_part = getattr(master_doc.part, '_styles_part', None)
            target_styles_part = getattr(target_doc.part, '_styles_part', None)
            if master_styles_part and target_styles_part:
                _log.debug("styles_part_accessed", extra={"method": "direct_private_attribute"})
        except Exception:
            pass

        if not master_styles_part or not target_styles_part:
            # Try public part.styles_part if available
            try:
                master_styles_part = getattr(master_doc.part, 'styles_part', None)
                target_styles_part = getattr(target_doc.part, 'styles_part', None)
                if master_styles_part and target_styles_part:
                    _log.debug("styles_part_accessed", extra={"method": "direct_attribute"})
            except Exception:
                pass
        
        if not master_styles_part or not target_styles_part:
            # Method 2: Try through part relationships
            try:
                from docx.opc.part import Part
                # Look for styles part in relationships
                master_rel = master_doc.part.rels.get('http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles')
                target_rel = target_doc.part.rels.get('http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles')
                
                if master_rel:
                    master_styles_part = master_rel.target_part
                if target_rel:
                    target_styles_part = target_rel.target_part
                    
                if master_styles_part or target_styles_part:
                    _log.debug("styles_part_accessed", extra={"method": "relationships"})
            except Exception as rel_e:
                _log.debug("styles_part_access_failed", extra={"method": "relationships", "error": str(rel_e)})
        
        # Also try to access and replace numbering part if possible
        try:
            master_numbering_part = getattr(master_doc.part, '_numbering_part', None) or getattr(master_doc.part, 'numbering_part', None)
            target_numbering_part = getattr(target_doc.part, '_numbering_part', None) or getattr(target_doc.part, 'numbering_part', None)
            if master_numbering_part and target_numbering_part:
                # Replace numbering XML similarly
                from copy import deepcopy
                master_num_xml_copy = deepcopy(master_numbering_part.element)
                target_num_xml = target_numbering_part.element
                parent = target_num_xml.getparent()
                if parent is not None:
                    parent.remove(target_num_xml)
                    parent.append(master_num_xml_copy)
                    _log.debug("numbering_xml_replaced_in_memory")
        except Exception as num_e:
            _log.debug("numbering_replacement_failed", extra={"error": str(num_e)})

        if master_styles_part is None or target_styles_part is None:
            # styles_part is not accessible - use file-based method as primary fallback
            # This is more reliable than the API fallback since it directly replaces styles.xml
            _log.warning("styles_part_not_accessible", extra={
                "master_styles_part_available": master_styles_part is not None,
                "target_styles_part_available": target_styles_part is not None,
                "fallback": "trying_file_based_method"
            })
            
            # Try file-based normalization first (most reliable)
            _log.warning("attempting_file_based_normalization", extra={
                "reason": "styles_part_not_accessible"
            })
            try:
                normalized_stream = normalize_styles_via_file(master_doc, target_doc)
                if normalized_stream:
                    # Store the stream in target_doc for caller to use
                    # The caller will reload the document from this stream
                    target_doc._normalized_stream = normalized_stream
                    _log.warning("file_based_normalization_prepared", extra={
                        "note": "caller_will_reload_document_from_stream",
                        "stream_size": len(normalized_stream.getvalue())
                    })
                    return
                else:
                    _log.warning("file_based_normalization_returned_none", extra={
                        "fallback": "using_document_styles_api"
                    })
            except Exception as file_e:
                _log.error("file_based_normalization_exception", extra={
                    "error": str(file_e),
                    "error_type": type(file_e).__name__,
                    "fallback": "using_document_styles_api"
                }, exc_info=True)
            
            # Final fallback: Copy styles using Document.styles API
            _log.info("using_document_styles_api_fallback", extra={
                "note": "api_fallback_has_limitations_may_not_prevent_styles_1_error"
            })
            _normalize_styles_via_api(master_doc, target_doc, _log)
            return
        
        # Get the styles XML elements
        master_styles_xml = master_styles_part.element
        target_styles_xml = target_styles_part.element
        
        # Log style IDs before replacement
        master_style_ids = [style.get(qn('w:styleId')) for style in master_styles_xml.findall(qn('w:style'))]
        target_style_ids_before = [style.get(qn('w:styleId')) for style in target_styles_xml.findall(qn('w:style'))]
        _log.debug("style_ids_before_normalization", extra={
            "master_style_ids": master_style_ids[:20],  # Limit to first 20 for log size
            "target_style_ids_before": target_style_ids_before[:20]
        })
        
        # Replace the entire styles XML in target with master's styles XML
        # This is more reliable than copying individual styles
        from lxml import etree
        
        # Deep copy the master styles XML element
        # Use deepcopy to ensure we get a complete, independent copy
        from copy import deepcopy
        master_xml_copy = deepcopy(master_styles_xml)
        
        # Get the parent of target_styles_xml to replace it
        parent = target_styles_xml.getparent()
        if parent is not None:
            # Remove the old styles element
            parent.remove(target_styles_xml)
            # Insert the master's styles element
            parent.append(master_xml_copy)
            _log.debug("styles_xml_replaced_via_parent")
        else:
            # Fallback: Clear and copy children if no parent
            target_styles_xml.clear()
            for child in master_styles_xml:
                target_styles_xml.append(deepcopy(child))
            
            # Copy all attributes from master
            for attr_name, attr_value in master_styles_xml.attrib.items():
                target_styles_xml.set(attr_name, attr_value)
            _log.debug("styles_xml_replaced_via_children")
        
        # Verify replacement
        target_style_ids_after = [style.get(qn('w:styleId')) for style in target_styles_xml.findall(qn('w:style'))]
        _log.info("style_normalization_complete", extra={
            "target_style_count_before": len(target_style_ids_before),
            "target_style_count_after": len(target_style_ids_after),
            "method": "xml_replacement",
            "target_style_ids_after": target_style_ids_after[:20]  # Limit to first 20
        })
                    
    except AttributeError as e:
        # Handle case where styles_part attribute doesn't exist
        _log.debug("styles_part_not_available", extra={"error": str(e), "action": "trying_api_fallback"})
        try:
            _normalize_styles_via_api(master_doc, target_doc, _log)
        except Exception as api_e:
            _log.warning("api_fallback_failed", extra={"error": str(api_e)})
    except Exception as e:
        # If normalization fails, log but don't crash
        # The document might still be mergeable, just with potential style conflicts
        _log.warning("style_normalization_failed", extra={"error": str(e)}, exc_info=True)


def _normalize_styles_via_api(master_doc: Document, target_doc: Document, log) -> None:
    """
    Fallback method to normalize styles using Document.styles API.
    This is used when direct styles_part access is not available.
    
    NOTE: This method has limitations - it can only copy style properties
    that are accessible through python-docx's API. Some style properties
    may not be copyable, which is why the "Styles 1" error may persist.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor
    
    log.info("using_document_styles_api_fallback", extra={"method": "fallback"})
    
    # Get all styles from master
    master_styles = {}
    for style in master_doc.styles:
        master_styles[style.name] = style
    
    log.info("api_fallback_processing_styles", extra={
        "master_style_count": len(master_styles),
        "target_style_count_before": len(target_doc.styles)
    })
    
    copied_count = 0
    added_count = 0
    failed_count = 0
    
    # Copy each style from master to target
    for style_name, master_style in master_styles.items():
        try:
            # Check if style exists in target
            if style_name in target_doc.styles:
                target_style = target_doc.styles[style_name]
                # Try to copy properties (limited by python-docx API)
                try:
                    if hasattr(master_style, 'font') and hasattr(target_style, 'font'):
                        if hasattr(master_style.font, 'name'):
                            target_style.font.name = master_style.font.name
                        if hasattr(master_style.font, 'size'):
                            target_style.font.size = master_style.font.size
                        if hasattr(master_style.font, 'bold'):
                            target_style.font.bold = master_style.font.bold
                        if hasattr(master_style.font, 'italic'):
                            target_style.font.italic = master_style.font.italic
                        if hasattr(master_style.font, 'underline'):
                            target_style.font.underline = master_style.font.underline
                        if hasattr(master_style.font, 'color') and hasattr(master_style.font.color, 'rgb'):
                            if master_style.font.color.rgb:
                                target_style.font.color.rgb = master_style.font.color.rgb
                    copied_count += 1
                except Exception as prop_e:
                    log.debug("style_property_copy_failed", extra={
                        "style_name": style_name,
                        "error": str(prop_e)
                    })
                    failed_count += 1
            else:
                # Style doesn't exist in target - try to add it
                try:
                    new_style = target_doc.styles.add_style(style_name, master_style.type)
                    if hasattr(master_style, 'font') and hasattr(new_style, 'font'):
                        if hasattr(master_style.font, 'name'):
                            new_style.font.name = master_style.font.name
                        if hasattr(master_style.font, 'size'):
                            new_style.font.size = master_style.font.size
                        if hasattr(master_style.font, 'bold'):
                            new_style.font.bold = master_style.font.bold
                        if hasattr(master_style.font, 'italic'):
                            new_style.font.italic = master_style.font.italic
                    added_count += 1
                except Exception as add_e:
                    log.debug("style_add_failed", extra={
                        "style_name": style_name,
                        "error": str(add_e)
                    })
                    failed_count += 1
        except Exception as e:
            log.debug("style_copy_exception", extra={
                "style_name": style_name,
                "error": str(e)
            })
            failed_count += 1
    
    log.info("styles_normalized_via_api", extra={
        "total_master_styles": len(master_styles),
        "copied_count": copied_count,
        "added_count": added_count,
        "failed_count": failed_count,
        "target_style_count_after": len(target_doc.styles),
        "method": "document_styles_api",
        "note": "API fallback has limitations - may not prevent Styles 1 error"
    })


def diagnose_style_conflicts(master_doc: Document, target_doc: Document) -> Dict[str, Any]:
    """
    Diagnose potential style conflicts between two documents.
    Returns a dictionary with diagnostic information.
    
    Args:
        master_doc: Master document
        target_doc: Target document to be merged
        
    Returns:
        Dictionary with diagnostic information including:
        - style_counts: Count of styles in each document
        - conflicting_styles: List of style names that exist in both but may differ
        - master_only: Styles only in master
        - target_only: Styles only in target
        - styles_part_available: Whether styles_part is accessible
    """
    diagnosis = {
        "style_counts": {
            "master": len(master_doc.styles),
            "target": len(target_doc.styles)
        },
        "conflicting_styles": [],
        "master_only": [],
        "target_only": [],
        "styles_part_available": {
            "master": False,
            "target": False
        }
    }
    
    # Check styles_part availability
    try:
        master_styles_part = getattr(master_doc.part, 'styles_part', None)
        diagnosis["styles_part_available"]["master"] = master_styles_part is not None
    except:
        pass
    
    try:
        target_styles_part = getattr(target_doc.part, 'styles_part', None)
        diagnosis["styles_part_available"]["target"] = target_styles_part is not None
    except:
        pass
    
    # Get style names
    master_style_names = {style.name for style in master_doc.styles}
    target_style_names = {style.name for style in target_doc.styles}
    
    # Find conflicts (styles that exist in both)
    common_styles = master_style_names & target_style_names
    diagnosis["conflicting_styles"] = list(common_styles)
    
    # Find unique styles
    diagnosis["master_only"] = list(master_style_names - target_style_names)
    diagnosis["target_only"] = list(target_style_names - master_style_names)
    
    _log.info("style_diagnosis_complete", extra={
        "master_style_count": diagnosis['style_counts']['master'],
        "target_style_count": diagnosis['style_counts']['target'],
        "common_styles_count": len(diagnosis['conflicting_styles']),
        "master_only_count": len(diagnosis['master_only']),
        "target_only_count": len(diagnosis['target_only']),
        "styles_part_available_master": diagnosis['styles_part_available']['master'],
        "styles_part_available_target": diagnosis['styles_part_available']['target']
    })
    
    return diagnosis
