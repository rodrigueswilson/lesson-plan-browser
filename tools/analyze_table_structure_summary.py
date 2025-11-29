"""
Quick summary analysis of table structure in lesson plan files.
Only shows statistics, not detailed per-file analysis.
"""

import sys
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))


def analyze_file_quick(docx_path: str) -> dict:
    """Quick analysis of a single file."""
    try:
        doc = Document(docx_path)
        
        # Count tables
        table_count = len(doc.tables)
        
        # Detect signature table
        has_signature = False
        if table_count > 0:
            last_table = doc.tables[-1]
            if last_table.rows and last_table.rows[0].cells:
                first_cell = last_table.rows[0].cells[0].text.strip().lower()
                if "signature" in first_cell:
                    has_signature = True
        
        # Estimate slots
        effective_tables = table_count - 1 if has_signature else table_count
        estimated_slots = effective_tables // 2
        
        return {
            "success": True,
            "table_count": table_count,
            "has_signature": has_signature,
            "estimated_slots": estimated_slots,
            "error": None
        }
    except Exception as e:
        return {
            "success": False,
            "table_count": 0,
            "has_signature": False,
            "estimated_slots": 0,
            "error": str(e)
        }


def main():
    """Analyze all lesson plan files - summary only."""
    # Directories to scan
    base_dirs = [
        Path("F:/rodri/Documents/OneDrive/AS/Daniela LP"),
        Path("F:/rodri/Documents/OneDrive/AS/Lesson Plan")
    ]
    
    print("Scanning directories...")
    all_docx_files = []
    
    # Find all DOCX files
    for base_dir in base_dirs:
        if not base_dir.exists():
            print(f"WARNING: Directory not found: {base_dir}")
            continue
        
        print(f"  {base_dir}")
        docx_files = list(base_dir.rglob("*.docx"))
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]
        print(f"    Found {len(docx_files)} files")
        all_docx_files.extend(docx_files)
    
    if not all_docx_files:
        print("No DOCX files found")
        return
    
    print(f"\nAnalyzing {len(all_docx_files)} files...")
    
    # Track statistics
    stats = {
        "total": len(all_docx_files),
        "success": 0,
        "errors": 0,
        "table_counts": {},
        "slot_counts": {},
        "has_signature": 0,
        "no_signature": 0
    }
    
    # Analyze each file
    for idx, docx_file in enumerate(all_docx_files, 1):
        if idx % 10 == 0:
            print(f"  Progress: {idx}/{len(all_docx_files)}")
        
        result = analyze_file_quick(str(docx_file))
        
        if result["success"]:
            stats["success"] += 1
            
            # Track table count
            tc = result["table_count"]
            stats["table_counts"][tc] = stats["table_counts"].get(tc, 0) + 1
            
            # Track slot count
            sc = result["estimated_slots"]
            stats["slot_counts"][sc] = stats["slot_counts"].get(sc, 0) + 1
            
            # Track signature
            if result["has_signature"]:
                stats["has_signature"] += 1
            else:
                stats["no_signature"] += 1
        else:
            stats["errors"] += 1
    
    # Print summary
    print(f"\n{'='*80}")
    print("ANALYSIS SUMMARY")
    print(f"{'='*80}")
    print(f"Total files: {stats['total']}")
    print(f"Successfully analyzed: {stats['success']}")
    print(f"Errors: {stats['errors']}")
    
    print(f"\nTable count distribution:")
    for count in sorted(stats["table_counts"].keys()):
        freq = stats["table_counts"][count]
        pct = (freq / stats["success"] * 100) if stats["success"] > 0 else 0
        print(f"  {count:2d} tables: {freq:3d} files ({pct:5.1f}%)")
    
    print(f"\nEstimated slot distribution:")
    for count in sorted(stats["slot_counts"].keys()):
        freq = stats["slot_counts"][count]
        pct = (freq / stats["success"] * 100) if stats["success"] > 0 else 0
        print(f"  {count:2d} slots: {freq:3d} files ({pct:5.1f}%)")
    
    print(f"\nSignature table detection:")
    print(f"  With signature: {stats['has_signature']} files")
    print(f"  Without signature: {stats['no_signature']} files")
    
    print(f"{'='*80}")


if __name__ == "__main__":
    main()
