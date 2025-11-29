"""
Detailed analysis of Daniela's output to understand:
1. Why there are 48 hyperlinks (vs 16 in input)
2. Why there's a Referenced Links section
3. Where the extra hyperlinks came from
"""

import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_parser import DOCXParser

OUTPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Daniela_Silva_Weekly_W43_10-20-10-24_20251019_150909.docx'


def analyze_all_hyperlinks():
    """Get detailed view of all hyperlinks in output."""
    
    print("="*80)
    print("DETAILED HYPERLINK ANALYSIS")
    print("="*80)
    print()
    
    parser = DOCXParser(OUTPUT_FILE)
    hyperlinks = parser.extract_hyperlinks()
    
    print(f"Total hyperlinks: {len(hyperlinks)}")
    print()
    
    # Group by table
    by_table = {}
    para_links = []
    
    for link in hyperlinks:
        table_idx = link.get('table_idx')
        if table_idx is not None:
            if table_idx not in by_table:
                by_table[table_idx] = []
            by_table[table_idx].append(link)
        else:
            para_links.append(link)
    
    print(f"Hyperlinks by location:")
    print(f"  Paragraph links: {len(para_links)}")
    for table_idx in sorted(by_table.keys()):
        print(f"  Table {table_idx}: {len(by_table[table_idx])} links")
    print()
    
    # Show paragraph links (these are likely in Referenced Links section)
    if para_links:
        print("="*80)
        print("PARAGRAPH LINKS (Likely in Referenced Links section):")
        print("="*80)
        print()
        for i, link in enumerate(para_links, 1):
            print(f"{i}. '{link['text'][:60]}'")
            print(f"   URL: {link['url'][:80]}")
            print()
    
    # Show links in each table
    for table_idx in sorted(by_table.keys()):
        print("="*80)
        print(f"TABLE {table_idx} HYPERLINKS ({len(by_table[table_idx])} links):")
        print("="*80)
        print()
        
        # Group by row
        by_row = {}
        for link in by_table[table_idx]:
            row_idx = link.get('row_idx')
            if row_idx not in by_row:
                by_row[row_idx] = []
            by_row[row_idx].append(link)
        
        for row_idx in sorted(by_row.keys()):
            row_label = by_row[row_idx][0].get('row_label', 'N/A')
            print(f"Row {row_idx} ({row_label}):")
            
            for link in by_row[row_idx]:
                col_header = link.get('col_header', 'N/A')
                print(f"  - '{link['text'][:50]}' in {col_header}")
        
        print()


def check_referenced_links_section():
    """Check what's in the Referenced Links section."""
    
    print("="*80)
    print("REFERENCED LINKS SECTION CONTENT")
    print("="*80)
    print()
    
    doc = Document(OUTPUT_FILE)
    
    in_referenced_section = False
    referenced_links = []
    
    for para in doc.paragraphs:
        if "Referenced Links" in para.text:
            in_referenced_section = True
            print("Found 'Referenced Links' heading")
            print()
            continue
        
        if in_referenced_section:
            # Check if paragraph has hyperlinks
            for hyperlink in para._element.xpath('.//w:hyperlink'):
                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                if text:
                    referenced_links.append(text)
    
    print(f"Links in Referenced Links section: {len(referenced_links)}")
    print()
    
    if referenced_links:
        print("Links:")
        for i, text in enumerate(referenced_links, 1):
            print(f"  {i}. {text[:60]}")
    
    print()
    return len(referenced_links)


def analyze_table_structure():
    """Understand the multi-slot structure."""
    
    print("="*80)
    print("TABLE STRUCTURE ANALYSIS")
    print("="*80)
    print()
    
    doc = Document(OUTPUT_FILE)
    
    print(f"Total tables in document: {len(doc.tables)}")
    print()
    
    # Analyze each table
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if table.rows else 0
        
        # Get first cell text to identify table type
        first_cell = table.rows[0].cells[0].text.strip()[:50] if table.rows and table.rows[0].cells else ""
        
        print(f"Table {i}:")
        print(f"  Dimensions: {rows}x{cols}")
        print(f"  First cell: '{first_cell}'")
        
        # Count hyperlinks
        link_count = 0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for hyperlink in para._element.xpath('.//w:hyperlink'):
                        link_count += 1
        
        print(f"  Hyperlinks: {link_count}")
        print()


if __name__ == '__main__':
    analyze_all_hyperlinks()
    fallback_count = check_referenced_links_section()
    analyze_table_structure()
    
    print("="*80)
    print("SUMMARY")
    print("="*80)
    print()
    print("The output has multiple lesson plan tables (4 slots)")
    print("Each slot has its own table with hyperlinks")
    print(f"Referenced Links section has {fallback_count} links")
    print()
    print("This is expected for multi-slot processing!")
