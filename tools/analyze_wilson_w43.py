"""
Detailed analysis of Wilson's W43 lesson plan to validate coordinate-based placement.
Adapted from Daniela analysis script.
"""

import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_parser import DOCXParser

OUTPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_152559.docx'


def analyze_all_hyperlinks():
    """Get detailed view of all hyperlinks in output."""
    
    print("="*80)
    print("WILSON W43 LESSON PLAN - DETAILED HYPERLINK ANALYSIS")
    print("="*80)
    print()
    
    parser = DOCXParser(OUTPUT_FILE)
    hyperlinks = parser.extract_hyperlinks()
    
    print(f"Total hyperlinks: {len(hyperlinks)}")
    print()
    
    # Group by table
    by_table = {}
    para_links = []
    
    for link in hyperlinks:
        table_idx = link.get('table_idx')
        if table_idx is not None:
            if table_idx not in by_table:
                by_table[table_idx] = []
            by_table[table_idx].append(link)
        else:
            para_links.append(link)
    
    print(f"Hyperlinks by location:")
    print(f"  Paragraph links: {len(para_links)}")
    for table_idx in sorted(by_table.keys()):
        print(f"  Table {table_idx}: {len(by_table[table_idx])} links")
    print()
    
    # Show paragraph links (these are likely in Referenced Links section)
    if para_links:
        print("="*80)
        print(f"PARAGRAPH LINKS ({len(para_links)} - Likely in Referenced Links section):")
        print("="*80)
        print()
        for i, link in enumerate(para_links, 1):
            print(f"{i}. '{link['text'][:60]}'")
            print(f"   URL: {link['url'][:80]}")
            print()
    
    # Show links in main lesson plan tables
    lesson_tables = [k for k in by_table.keys() if len(by_table[k]) > 5]  # Tables with many links
    
    for table_idx in sorted(lesson_tables):
        print("="*80)
        print(f"TABLE {table_idx} HYPERLINKS ({len(by_table[table_idx])} links):")
        print("="*80)
        print()
        
        # Group by row
        by_row = {}
        for link in by_table[table_idx]:
            row_idx = link.get('row_idx')
            if row_idx not in by_row:
                by_row[row_idx] = []
            by_row[row_idx].append(link)
        
        for row_idx in sorted(by_row.keys()):
            row_label = by_row[row_idx][0].get('row_label', 'N/A')
            print(f"Row {row_idx} ({row_label}): {len(by_row[row_idx])} links")
            
            # Group by column
            by_col = {}
            for link in by_row[row_idx]:
                col_header = link.get('col_header', 'N/A')
                if col_header not in by_col:
                    by_col[col_header] = []
                by_col[col_header].append(link)
            
            for col_header in sorted(by_col.keys()):
                print(f"  {col_header}: {len(by_col[col_header])} links")
                for link in by_col[col_header][:3]:  # Show first 3
                    print(f"    - '{link['text'][:50]}'")
                if len(by_col[col_header]) > 3:
                    print(f"    ... and {len(by_col[col_header]) - 3} more")
        
        print()
    
    return len(hyperlinks), len(para_links), len(hyperlinks) - len(para_links)


def check_referenced_links_section():
    """Check what's in the Referenced Links section."""
    
    print("="*80)
    print("REFERENCED LINKS SECTION")
    print("="*80)
    print()
    
    doc = Document(OUTPUT_FILE)
    
    in_referenced_section = False
    referenced_links = []
    
    for para in doc.paragraphs:
        if "Referenced Links" in para.text:
            in_referenced_section = True
            print("✓ Found 'Referenced Links' heading")
            continue
        
        if in_referenced_section:
            # Check if paragraph has hyperlinks
            for hyperlink in para._element.xpath('.//w:hyperlink'):
                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                if text:
                    referenced_links.append(text)
    
    if not in_referenced_section:
        print("✓ No 'Referenced Links' section found")
        print("  All hyperlinks were placed inline!")
        return 0
    
    print(f"\nLinks in Referenced Links section: {len(referenced_links)}")
    
    if referenced_links:
        print("\nLinks:")
        for i, text in enumerate(referenced_links, 1):
            print(f"  {i}. {text[:60]}")
    
    print()
    return len(referenced_links)


def analyze_table_structure():
    """Understand the multi-slot structure."""
    
    print("="*80)
    print("TABLE STRUCTURE")
    print("="*80)
    print()
    
    doc = Document(OUTPUT_FILE)
    
    print(f"Total tables in document: {len(doc.tables)}")
    print()
    
    # Identify lesson plan tables (8x6 structure)
    lesson_tables = []
    
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if table.rows else 0
        
        # Count hyperlinks
        link_count = 0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for hyperlink in para._element.xpath('.//w:hyperlink'):
                        link_count += 1
        
        # Get first cell to identify table type
        first_cell = ""
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()[:30]
        
        is_lesson_table = (rows == 8 and cols == 6)
        
        if is_lesson_table:
            lesson_tables.append(i)
        
        print(f"Table {i}: {rows}x{cols}, {link_count} links", end="")
        if is_lesson_table:
            print(" ← LESSON PLAN TABLE")
        else:
            print(f" ({first_cell})")
    
    print()
    print(f"Lesson plan tables found: {len(lesson_tables)}")
    print(f"Table indices: {lesson_tables}")
    print()
    
    return len(lesson_tables)


def calculate_placement_stats(total_links, fallback_links, table_links):
    """Calculate and display placement statistics."""
    
    print("="*80)
    print("PLACEMENT STATISTICS")
    print("="*80)
    print()
    
    inline_links = table_links
    inline_rate = (inline_links / total_links * 100) if total_links > 0 else 0
    fallback_rate = (fallback_links / total_links * 100) if total_links > 0 else 0
    
    print(f"Total hyperlinks: {total_links}")
    print(f"  ✅ Inline (in tables): {inline_links} ({inline_rate:.1f}%)")
    print(f"  ⚠️  Fallback (Referenced Links): {fallback_links} ({fallback_rate:.1f}%)")
    print()
    
    # Comparison with baseline
    baseline = 84.2
    improvement = inline_rate - baseline
    
    print(f"Baseline (fuzzy matching): {baseline}%")
    print(f"Current (coordinate-based): {inline_rate:.1f}%")
    print(f"Improvement: {improvement:+.1f} percentage points")
    print()
    
    # Target assessment
    if inline_rate >= 93:
        print(f"✅ TARGET ACHIEVED (93-97% inline)")
        if inline_rate >= 97:
            print(f"   EXCEEDED TARGET!")
    else:
        print(f"⚠️  Target not met (need {93 - inline_rate:.1f}% more)")
    
    print()


def main():
    print("\n")
    print("="*80)
    print("WILSON W43 PRODUCTION TEST - COORDINATE-BASED PLACEMENT ANALYSIS")
    print("="*80)
    print()
    print(f"File: {Path(OUTPUT_FILE).name}")
    print(f"Generated: 2025-10-19 15:25:59")
    print()
    
    # Analyze hyperlinks
    total_links, para_links, table_links = analyze_all_hyperlinks()
    
    # Check for fallback section
    fallback_count = check_referenced_links_section()
    
    # Analyze structure
    slot_count = analyze_table_structure()
    
    # Calculate statistics
    calculate_placement_stats(total_links, fallback_count, table_links)
    
    # Summary
    print("="*80)
    print("SUMMARY")
    print("="*80)
    print()
    print(f"✓ Slots processed: {slot_count}")
    print(f"✓ Total hyperlinks: {total_links}")
    print(f"✓ Table links: {table_links} (placed via coordinates)")
    print(f"✓ Paragraph links: {para_links} (in fallback)")
    print(f"✓ Fallback section: {fallback_count} links")
    print()
    
    if fallback_count == 0:
        print("🎉 PERFECT! All hyperlinks placed inline (100%)")
    elif fallback_count == para_links:
        print("✅ EXCELLENT! All table links placed inline")
        print("   Only paragraph links in fallback (expected)")
    else:
        print("⚠️  Some table links in fallback (investigate)")
    
    print()
    print("="*80)
    print("ANALYSIS COMPLETE")
    print("="*80)


if __name__ == '__main__':
    main()
