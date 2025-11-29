"""
Detailed inspection of district template structure.
"""

from docx import Document
from pathlib import Path
import sys


def inspect_template_detailed(template_path: str):
    """Detailed inspection of DOCX template structure."""
    doc = Document(template_path)
    
    print(f"=== Detailed Template Inspection ===\n")
    
    # Table 0 - Metadata
    print("TABLE 0 - METADATA:")
    table = doc.tables[0]
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {row_idx}: {cells}")
    
    # Table 1 - Daily Plans
    print("\nTABLE 1 - DAILY PLANS:")
    table = doc.tables[1]
    print(f"  Dimensions: {len(table.rows)} rows x {len(table.columns)} cols")
    
    # Show all rows
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip()[:30] for cell in row.cells]
        print(f"  Row {row_idx}: {cells}")
    
    # Table 2 - Signatures
    print("\nTABLE 2 - SIGNATURES:")
    table = doc.tables[2]
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {row_idx}: {cells}")
    
    # Check cell formatting
    print("\nCELL FORMATTING (Table 1, Row 0):")
    table = doc.tables[1]
    if table.rows:
        for cell_idx, cell in enumerate(table.rows[0].cells):
            para = cell.paragraphs[0] if cell.paragraphs else None
            if para and para.runs:
                run = para.runs[0]
                print(f"  Cell {cell_idx}: font={run.font.name}, size={run.font.size}, bold={run.font.bold}")


if __name__ == "__main__":
    template_path = sys.argv[1] if len(sys.argv) > 1 else "input/Lesson Plan Template SY'25-26.docx"
    inspect_template_detailed(template_path)
