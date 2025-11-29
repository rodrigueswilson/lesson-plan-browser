"""
Multi-Slot DOCX Renderer - Creates separate table sets for each slot.

For multi-slot lesson plans, this renderer:
1. Creates a complete table set for each slot (metadata + daily plans + signatures)
2. Each slot has its own teacher name, grade, and subject
3. Maintains template formatting for each slot
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict, List
import sys

# Handle imports for both CLI and module usage
try:
    from tools.docx_renderer import DOCXRenderer
except ImportError:
    from docx_renderer import DOCXRenderer


class MultiSlotDOCXRenderer:
    """Render multi-slot lesson plans with separate table sets per slot."""
    
    def __init__(self, template_path: str):
        """
        Initialize renderer with district template.
        
        Args:
            template_path: Path to district DOCX template
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Use single-slot renderer for each slot
        self.single_renderer = DOCXRenderer(str(template_path))
    
    def render(self, json_data: Dict, output_path: str) -> bool:
        """
        Render multi-slot JSON data to DOCX with separate table sets.
        
        Args:
            json_data: Multi-slot lesson plan JSON
            output_path: Path to save DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Check if this is multi-slot data
            if not self._is_multi_slot(json_data):
                # Use single-slot renderer
                return self.single_renderer.render(json_data, output_path)
            
            # Create new document
            doc = Document()
            
            # Extract slots from multi-slot structure
            slots_data = self._extract_slots(json_data)
            
            # Render each slot with its own table set
            for i, slot_data in enumerate(slots_data):
                if i > 0:
                    # Add page break between slots
                    doc.add_page_break()
                
                # Add slot header
                heading = doc.add_heading(f"Slot {slot_data['metadata']['slot_number']}: {slot_data['metadata']['subject']}", level=1)
                
                # Load template for this slot
                template_doc = Document(self.template_path)
                
                # Copy only metadata and daily plans tables (not signature)
                # Template has 3 tables: [0]=metadata, [1]=daily_plans, [2]=signatures
                for table_idx in [0, 1]:  # Only copy first 2 tables
                    self._copy_table(doc, template_doc.tables[table_idx])
                
                # Fill in the tables we just added
                self._fill_slot_tables(doc, slot_data, i)
            
            # Add ONE signature table at the very end
            template_doc = Document(self.template_path)
            if len(template_doc.tables) >= 3:
                # Copy signature table (table index 2)
                self._copy_table(doc, template_doc.tables[2])
            
            # Save
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            
            print(f"Successfully rendered multi-slot DOCX: {output_path}")
            return True
            
        except Exception as e:
            print(f"Error rendering DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _is_multi_slot(self, json_data: Dict) -> bool:
        """Check if JSON data is multi-slot format."""
        # Multi-slot format has 'slots' arrays in each day
        if 'days' in json_data:
            first_day = next(iter(json_data['days'].values()))
            return isinstance(first_day, dict) and 'slots' in first_day
        return False
    
    def _extract_slots(self, multi_slot_data: Dict) -> List[Dict]:
        """
        Extract individual slot data from multi-slot structure.
        
        Converts:
        {
          "days": {
            "monday": {"slots": [slot1, slot2, ...]},
            ...
          }
        }
        
        To:
        [
          {
            "metadata": {...},
            "days": {"monday": slot1_data, "tuesday": slot1_data, ...}
          },
          {
            "metadata": {...},
            "days": {"monday": slot2_data, "tuesday": slot2_data, ...}
          }
        ]
        """
        slots_by_number = {}
        
        # Collect all slots across all days
        for day_name, day_data in multi_slot_data['days'].items():
            if 'slots' in day_data:
                for slot in day_data['slots']:
                    slot_num = slot.get('slot_number', 0)
                    
                    if slot_num not in slots_by_number:
                        slots_by_number[slot_num] = {
                            'metadata': {
                                'week_of': multi_slot_data['metadata'].get('week_of', ''),
                                'grade': slot.get('grade', multi_slot_data['metadata'].get('grade', '')),
                                'subject': slot.get('subject', ''),
                                'teacher_name': slot.get('teacher_name', ''),
                                'homeroom': slot.get('homeroom', ''),
                                'slot_number': slot_num
                            },
                            'days': {}
                        }
                    
                    # Add this day's data for this slot
                    slots_by_number[slot_num]['days'][day_name] = slot
        
        # Convert to list sorted by slot number
        return [slots_by_number[num] for num in sorted(slots_by_number.keys())]
    
    def _copy_table(self, doc: Document, source_table):
        """Copy a table from source to destination document."""
        from copy import deepcopy
        
        # Add new table with same dimensions
        new_table = doc.add_table(
            rows=len(source_table.rows),
            cols=len(source_table.columns)
        )
        
        # Copy cell content
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.rows[i].cells[j]
                
                # Copy text content
                new_cell.text = cell.text
                
                # Copy cell formatting by cloning the element
                if cell._element.tcPr is not None:
                    try:
                        # Clone the tcPr element instead of assigning
                        new_cell._element.get_or_add_tcPr()
                        for child in cell._element.tcPr:
                            new_cell._element.tcPr.append(deepcopy(child))
                    except:
                        pass  # Skip if formatting copy fails
        
        # Copy table style
        if source_table.style:
            new_table.style = source_table.style
        
        return new_table
    
    def _fill_slot_tables(self, doc: Document, slot_data: Dict, slot_index: int):
        """Fill tables for a specific slot."""
        # Calculate table indices (each slot has 2 tables: metadata + daily_plans)
        # No signature table per slot
        base_idx = slot_index * 2
        
        # Use single-slot renderer logic
        metadata_table = doc.tables[base_idx]
        daily_plans_table = doc.tables[base_idx + 1]
        
        # Fill metadata
        self._fill_metadata_table(metadata_table, slot_data['metadata'])
        
        # Fill daily plans
        for day_name, day_data in slot_data['days'].items():
            self._fill_day_in_table(daily_plans_table, day_name, day_data)
    
    def _fill_metadata_table(self, table, metadata: Dict):
        """Fill metadata table (teacher, grade, subject, etc.)."""
        # This mirrors the logic from DOCXRenderer
        # Row 0: Teacher name, Grade, Subject
        # Row 1: Week of, Homeroom
        
        if len(table.rows) >= 2:
            # Assuming standard template layout
            # Adjust based on actual template structure
            cells = table.rows[0].cells
            if len(cells) >= 6:
                cells[1].text = metadata.get('teacher_name', '')
                cells[3].text = metadata.get('grade', '')
                cells[5].text = metadata.get('subject', '')
            
            cells = table.rows[1].cells
            if len(cells) >= 4:
                cells[1].text = metadata.get('week_of', '')
                cells[3].text = metadata.get('homeroom', '')
    
    def _fill_day_in_table(self, table, day_name: str, day_data: Dict):
        """Fill one day's data in the daily plans table."""
        # Use the single-slot renderer's logic
        col_idx = self.single_renderer.DAY_TO_COL.get(day_name.lower())
        if col_idx is None:
            return
        
        # Fill each row using single-slot format
        self.single_renderer._fill_single_slot_day(table, col_idx, day_data)


if __name__ == "__main__":
    """CLI usage for testing."""
    if len(sys.argv) < 3:
        print("Usage: python docx_renderer_multi_slot.py <input_json> <output_docx>")
        sys.exit(1)
    
    import json
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    template_file = "input/Lesson Plan Template SY'25-26.docx"
    
    # Load JSON
    with open(input_file, 'r') as f:
        json_data = json.load(f)
    
    # Render
    renderer = MultiSlotDOCXRenderer(template_file)
    success = renderer.render(json_data, output_file)
    
    sys.exit(0 if success else 1)
