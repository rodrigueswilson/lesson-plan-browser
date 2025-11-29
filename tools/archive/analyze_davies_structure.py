"""
Analyze Davies' file structure to understand table organization.
"""

from docx import Document

# Open Davies' file
doc = Document('F:/rodri/Documents/OneDrive/AS/Lesson Plan/25 W41/10_6-10_10 Davies Lesson Plans.docx')

print("="*70)
print("DAVIES FILE STRUCTURE ANALYSIS")
print("="*70)
print()

print(f"Total tables: {len(doc.tables)}")
print()

# Analyze each table
for i, table in enumerate(doc.tables):
    print(f"TABLE {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    
    # Check if it's a header table (1 row, multiple columns)
    if len(table.rows) == 1:
        print("  Type: HEADER TABLE")
        print("  Content:")
        for cell in table.rows[0].cells:
            if cell.text.strip():
                print(f"    - {cell.text.strip()}")
    
    # Check if it's a lesson table (multiple rows)
    elif len(table.rows) > 1:
        print("  Type: LESSON TABLE")
        # Show column headers (first row)
        headers = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
        print(f"  Columns: {', '.join(headers[:6])}")  # First 6 columns
        
        # Show first content row
        if len(table.rows) > 1:
            first_row_label = table.rows[1].cells[0].text.strip()
            if first_row_label:
                print(f"  First row: {first_row_label[:50]}")
    
    print()

# Extract subjects from header tables
print("="*70)
print("SUBJECTS IDENTIFIED:")
print("="*70)
print()

subjects = []
for i in range(0, len(doc.tables), 2):
    if i >= len(doc.tables):
        break
    
    header_table = doc.tables[i]
    
    # Look for subject in header table
    for cell in header_table.rows[0].cells:
        text = cell.text.strip()
        if 'Subject:' in text:
            subject = text.replace('Subject:', '').strip()
            subjects.append({
                'table_index': i,
                'subject': subject,
                'lesson_table_index': i + 1 if i + 1 < len(doc.tables) else None
            })
            break

for subj_info in subjects:
    print(f"Subject: {subj_info['subject']}")
    print(f"  Header Table: {subj_info['table_index']}")
    print(f"  Lesson Table: {subj_info['lesson_table_index']}")
    print()

# Show Math table specifically
print("="*70)
print("MATH LESSON TABLE DETAILS:")
print("="*70)
print()

math_info = next((s for s in subjects if 'math' in s['subject'].lower()), None)
if math_info and math_info['lesson_table_index'] is not None:
    math_table = doc.tables[math_info['lesson_table_index']]
    
    print(f"Table {math_info['lesson_table_index']}: {len(math_table.rows)} rows x {len(math_table.columns)} cols")
    print()
    
    # Show structure
    print("Column Headers:")
    for j, cell in enumerate(math_table.rows[0].cells):
        if cell.text.strip():
            print(f"  Col {j}: {cell.text.strip()}")
    print()
    
    print("Row Labels:")
    for i, row in enumerate(math_table.rows[:5]):  # First 5 rows
        label = row.cells[0].text.strip()
        if label:
            print(f"  Row {i}: {label[:60]}")
else:
    print("Math subject not found!")
