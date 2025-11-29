import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(
        "ERROR: python-docx is not installed. Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(2)


def cell_text(cell):
    return " ".join(p.text for p in cell.paragraphs).strip()


def dump_header(name, header):
    print(f"\n--- {name} ---")
    if header is None:
        print("(none)")
        return
    # Paragraphs
    for i, p in enumerate(header.paragraphs):
        t = p.text.strip()
        if t:
            print(f"para[{i}]: {t}")
    # Tables
    for ti, table in enumerate(header.tables):
        print(f"table[{ti}]: {len(table.rows)}x{len(table.columns)}")
        for ri, row in enumerate(table.rows):
            row_text = [cell_text(c) for c in row.cells]
            print(f"  r{ri}: | " + " | ".join(row_text))
    # Bookmarks / placeholders (best-effort)
    try:
        starts = header._element.xpath(
            ".//w:bookmarkStart", namespaces=header._element.nsmap
        )
        for s in starts:
            name = s.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name"
            )
            if name:
                print(f"bookmark: {name}")
    except Exception:
        pass
    # Placeholder patterns like {{FIELD}}
    try:
        for p in header.paragraphs:
            if "{{" in p.text and "}}" in p.text:
                print(f"placeholder: {p.text}")
        for t in header.tables:
            for r in t.rows:
                for c in r.cells:
                    tx = cell_text(c)
                    if "{{" in tx and "}}" in tx:
                        print(f"placeholder: {tx}")
    except Exception:
        pass


def main():
    if len(sys.argv) < 2:
        print("Usage: python tools/inspect_docx_headers.py <path-to-docx>")
        sys.exit(1)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"ERROR: File not found: {path}", file=sys.stderr)
        sys.exit(1)
    doc = Document(str(path))
    print(f"File: {path}")
    print(f"Sections: {len(doc.sections)}")
    # Dump headers for first section (most templates use one)
    sec = doc.sections[0]
    dump_header("header (default)", sec.header)
    # First/even page headers if present
    if hasattr(sec, "first_page_header"):
        dump_header("first_page_header", sec.first_page_header)
    if hasattr(sec, "even_page_header"):
        dump_header("even_page_header", sec.even_page_header)


if __name__ == "__main__":
    main()
