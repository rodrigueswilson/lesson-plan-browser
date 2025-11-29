"""
Script to review the signature image position in the generated document.
Inspects the document structure to verify the image is correctly positioned.
"""

from docx import Document
from docx.shared import Inches

# Configuration
OUTPUT_PATH = "output/test_signature.docx"


def review_signature_position():
    """Review the signature table structure and image positioning."""
    print("=" * 60)
    print("Reviewing Signature Position")
    print("=" * 60)
    print()
    
    # Load the document
    print(f"Loading document: {OUTPUT_PATH}")
    doc = Document(OUTPUT_PATH)
    
    # Find the signature table
    signature_table = None
    for i, table in enumerate(doc.tables):
        table_text = "\n".join(
            [cell.text for row in table.rows for cell in row.cells]
        )
        if "Required Signatures" in table_text:
            signature_table = table
            print(f"Found signature table at index {i}")
            break
    
    if not signature_table:
        print("ERROR: Could not find signature table")
        return False
    
    print(f"\nSignature table has {len(signature_table.rows)} rows")
    print()
    
    # Inspect each row
    for row_idx, row in enumerate(signature_table.rows):
        print(f"--- Row {row_idx} ---")
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            print(f"  Cell {cell_idx}: {cell_text[:80]}..." if len(cell_text) > 80 else f"  Cell {cell_idx}: {cell_text}")
            
            # Check for images in this cell
            image_count = 0
            for para in cell.paragraphs:
                for run in para.runs:
                    # Check if run has inline shapes (images)
                    if run._element.xpath('.//a:blip'):
                        image_count += 1
                        # Try to get image dimensions
                        try:
                            for drawing in run._element.xpath('.//wp:inline'):
                                # Get image dimensions
                                extent = drawing.find('.//wp:extent')
                                if extent is not None:
                                    cx = int(extent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cx', 0))
                                    cy = int(extent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cy', 0))
                                    # Convert EMU to inches (1 inch = 914400 EMU)
                                    width_inches = cx / 914400.0
                                    height_inches = cy / 914400.0
                                    print(f"    -> Found image: {width_inches:.3f}\" x {height_inches:.3f}\"")
                        except Exception as e:
                            print(f"    -> Found image (could not read dimensions: {e})")
            
            if image_count > 0:
                print(f"    -> Total images in cell: {image_count}")
            
            # Check paragraph structure for "Teacher Signature:"
            if "Teacher Signature:" in cell_text:
                print(f"\n  *** DETAILED INSPECTION: Teacher Signature Cell ***")
                for para_idx, para in enumerate(cell.paragraphs):
                    print(f"    Paragraph {para_idx}:")
                    print(f"      Text: '{para.text}'")
                    print(f"      Runs: {len(para.runs)}")
                    for run_idx, run in enumerate(para.runs):
                        run_text = run.text
                        has_image = bool(run._element.xpath('.//a:blip'))
                        font_size = run.font.size.pt if run.font.size else "default"
                        print(f"        Run {run_idx}: '{run_text[:50]}...' (font: {font_size}pt, has_image: {has_image})")
                        
                        if has_image:
                            # Get image dimensions and position
                            try:
                                from docx.oxml.ns import qn
                                for drawing in run._element.xpath('.//a:blip/..'):
                                    # Find the extent element
                                    extent_elem = drawing.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}extent')
                                    if extent_elem is not None:
                                        cx = int(extent_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cx', 0))
                                        cy = int(extent_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cy', 0))
                                        # Convert EMU to inches (1 inch = 914400 EMU)
                                        width_inches = cx / 914400.0
                                        height_inches = cy / 914400.0
                                        print(f"          Image dimensions: {width_inches:.3f}\" x {height_inches:.3f}\"")
                                        
                                        # Check where image appears in the run text
                                        run_text_before = run.text
                                        # Split run text to see structure
                                        if "Teacher Signature:" in run_text_before:
                                            parts = run_text_before.split("Teacher Signature:")
                                            if len(parts) > 1:
                                                after_label = parts[1]
                                                print(f"          Text after 'Teacher Signature:': '{after_label[:50]}...'")
                                                print(f"          Image position: In same run as text (inline)")
                            except Exception as e:
                                print(f"          Image found (error reading details: {e})")
                                import traceback
                                traceback.print_exc()
        print()
    
    print("\n" + "=" * 60)
    print("Review Complete")
    print("=" * 60)
    print("\nCheck the output above to verify:")
    print("1. Image is in the same paragraph as 'Teacher Signature:'")
    print("2. Image appears in the same run or immediately after the text")
    print("3. Image size is proportional to text (should be ~0.214 inches height)")
    
    return True


if __name__ == "__main__":
    review_signature_position()

