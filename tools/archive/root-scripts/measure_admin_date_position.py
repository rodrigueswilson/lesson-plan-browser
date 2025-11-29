"""Measure the exact position of 'Date:' in Administrator row."""

from docx import Document
from docx.shared import Inches, Pt, Twips

# Load template
doc = Document("input/Lesson Plan Template SY'25-26.docx")

# Find signature table
signature_table = None
for table in doc.tables:
    table_text = "\n".join([cell.text for row in table.rows for cell in row.cells])
    if "Required Signatures" in table_text:
        signature_table = table
        break

print("=" * 80)
print("MEASURING ADMINISTRATOR ROW")
print("=" * 80)

# Find Administrator row
for row in signature_table.rows:
    for cell in row.cells:
        # Get cell width
        tc = cell._element
        tcPr = tc.xpath('.//w:tcW')
        if tcPr:
            cell_width_twips = int(tcPr[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w'))
            cell_width_inches = Twips(cell_width_twips).inches
            print(f"\nCell width: {cell_width_twips} twips = {cell_width_inches:.2f} inches")
        
        for para in cell.paragraphs:
            para_text = para.text
            if "Administrator Signature:" in para_text:
                print(f"\nParagraph text: {repr(para_text)}")
                print(f"Text length: {len(para_text)} characters")
                
                # Find position of "Date:"
                admin_pos = para_text.find("Administrator Signature:")
                date_pos = para_text.find("Date:", admin_pos)
                
                print(f"\n'Administrator Signature:' at position: {admin_pos}")
                print(f"'Date:' at position: {date_pos}")
                print(f"Characters between: {date_pos - admin_pos - len('Administrator Signature:')}")
                
                # Analyze runs to understand spacing
                print(f"\nRuns in paragraph:")
                current_pos = 0
                for i, run in enumerate(para.runs):
                    run_text = run.text
                    run_len = len(run_text)
                    print(f"  Run {i}: pos {current_pos}-{current_pos + run_len}, text={repr(run_text[:50])}")
                    if run.font.size:
                        print(f"         Font size: {run.font.size.pt}pt")
                    if run.font.name:
                        print(f"         Font name: {run.font.name}")
                    current_pos += run_len
                
                # Calculate approximate position in inches
                # Typical character width in 11pt Times New Roman: ~0.07 inches
                # But underscores are different width than regular characters
                
                # Better approach: measure from cell start
                # "Administrator Signature:" = 24 chars
                # Then 1 space + 49 underscores = 50 chars
                # Total to "Date:" = 74 chars
                
                # In 11pt Times New Roman:
                # Regular char ≈ 0.07 inches
                # Underscore ≈ 0.08 inches
                
                # Approximate calculation:
                # 24 chars (label) * 0.07 = 1.68 inches
                # 50 chars (space + underscores) * 0.08 = 4.0 inches
                # Total ≈ 5.68 inches from cell start
                
                chars_to_date = date_pos
                estimated_inches = chars_to_date * 0.075  # Average char width
                
                print(f"\nEstimated 'Date:' position: {estimated_inches:.2f} inches from cell start")
                print(f"Cell width: {cell_width_inches:.2f} inches")
                print(f"'Date:' is at {(estimated_inches/cell_width_inches)*100:.1f}% of cell width")
                
                # Better estimate using actual measurements
                # Let's try different positions
                print(f"\nRecommended tab stop positions to test:")
                for test_inches in [5.0, 5.2, 5.4, 5.6, 5.8, 6.0]:
                    print(f"  {test_inches} inches")

print("\n" + "=" * 80)
