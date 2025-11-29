from docx import Document

doc = Document(r"d:\LP\input\Ms. Savoca- 9_15_24-9_19_24 Lesson plans.docx")

print("="*70)
print("SAVOCA FILE STRUCTURE ANALYSIS")
print("="*70)
print()

print(f"Total tables: {len(doc.tables)}")
print()

for i, t in enumerate(doc.tables):
    print(f"Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    
    # Show first row
    if len(t.rows) > 0:
        print("  First row:")
        for cell in t.rows[0].cells:
            if cell.text.strip():
                print(f"    - {cell.text.strip()[:60]}")
    
    # If more than 1 row, show second row
    if len(t.rows) > 1:
        print("  Second row:")
        for j, cell in enumerate(t.rows[1].cells[:3]):  # First 3 cells
            if cell.text.strip():
                print(f"    Col {j}: {cell.text.strip()[:60]}")
    
    print()
