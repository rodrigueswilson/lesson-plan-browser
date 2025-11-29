from docx import Document

doc = Document(r"d:\LP\input\Piret Lesson Plans 9_22_25-9_26_25.docx")

print("="*70)
print("PIRET FILE STRUCTURE ANALYSIS")
print("="*70)
print()

print(f"Total tables: {len(doc.tables)}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print()

# Check for text-based structure
print("First 10 paragraphs:")
for i, para in enumerate(doc.paragraphs[:10]):
    if para.text.strip():
        print(f"{i}: {para.text.strip()[:70]}")
print()

# Check tables
for i, t in enumerate(doc.tables):
    print(f"Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    
    # Show first row
    if len(t.rows) > 0:
        print("  First row:")
        for cell in t.rows[0].cells:
            if cell.text.strip():
                print(f"    - {cell.text.strip()[:60]}")
    
    # If more than 1 row, show second row
    if len(t.rows) > 1:
        print("  Second row:")
        for j, cell in enumerate(t.rows[1].cells[:3]):  # First 3 cells
            if cell.text.strip():
                print(f"    Col {j}: {cell.text.strip()[:60]}")
    
    print()
