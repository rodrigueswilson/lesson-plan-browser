from docx import Document

template_path = r"d:\LP\input\Lesson Plan Template SY'25-26.docx"
doc = Document(template_path)

print(f"Total tables: {len(doc.tables)}")
print()

for i, t in enumerate(doc.tables):
    print(f"Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")

print()
print("Table 0 first row:")
for cell in doc.tables[0].rows[0].cells:
    if cell.text.strip():
        print(f"  {cell.text.strip()}")
