"""
Inspect district template structure to understand layout.
"""

from docx import Document
from pathlib import Path
import sys


def inspect_template(template_path: str):
    """Inspect DOCX template structure."""
    doc = Document(template_path)
    
    print(f"=== Template Inspection: {template_path} ===\n")
    
    # Headers
    print("HEADERS:")
    for section in doc.sections:
        header = section.header
        print(f"  Header paragraphs: {len(header.paragraphs)}")
        for i, para in enumerate(header.paragraphs):
            print(f"    [{i}] {para.text[:100]}")
        print(f"  Header tables: {len(header.tables)}")
        for i, table in enumerate(header.tables):
            print(f"    Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    
    # Footers
    print("\nFOOTERS:")
    for section in doc.sections:
        footer = section.footer
        print(f"  Footer paragraphs: {len(footer.paragraphs)}")
        for i, para in enumerate(footer.paragraphs):
            print(f"    [{i}] {para.text[:100]}")
        print(f"  Footer tables: {len(footer.tables)}")
    
    # Body
    print(f"\nBODY:")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs[:10]):  # First 10
        text = para.text.strip()
        if text:
            print(f"    [{i}] {text[:100]}")
    
    print(f"\n  Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
        # Show first row
        if table.rows:
            first_row = table.rows[0]
            cells_text = [cell.text.strip()[:50] for cell in first_row.cells]
            print(f"    First row: {cells_text}")
    
    # Bookmarks (if any)
    print("\nBOOKMARKS/FIELDS:")
    # Note: python-docx doesn't directly expose bookmarks, but we can check for field codes
    for para in doc.paragraphs:
        if '{' in para.text or 'MERGEFIELD' in para.text:
            print(f"  Found potential field: {para.text[:100]}")


if __name__ == "__main__":
    template_path = sys.argv[1] if len(sys.argv) > 1 else "input/Lesson Plan Template SY'25-26.docx"
    inspect_template(template_path)
