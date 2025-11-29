"""
Analyze patterns of well-placed vs fallback links.

Questions to answer:
1. Where are the well-placed links? (which cells, which sections)
2. Where are the fallback links from? (which input sections)
3. What metadata do well-placed links have vs fallback?
4. What text patterns distinguish them?
"""

import sys
from pathlib import Path
from collections import defaultdict
import json

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.docx_parser import DOCXParser
from docx import Document
from docx.oxml.ns import qn

OUTPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx'

# We'll focus on Davies Grade 3 ELA since it has the most fallback links
INPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Davies Lesson Plans.docx'


def analyze_input_links(input_path):
    """Analyze all links in input file with their metadata."""
    
    print("="*100)
    print("ANALYZING INPUT LINKS")
    print("="*100)
    print()
    
    parser = DOCXParser(input_path)
    links = parser.extract_hyperlinks()
    
    print(f"Total links in input: {len(links)}\n")
    
    # Group by metadata completeness
    with_both = []
    with_section_only = []
    with_day_only = []
    with_neither = []
    
    for link in links:
        has_section = bool(link.get('section_hint'))
        has_day = bool(link.get('day_hint'))
        
        if has_section and has_day:
            with_both.append(link)
        elif has_section:
            with_section_only.append(link)
        elif has_day:
            with_day_only.append(link)
        else:
            with_neither.append(link)
    
    print("METADATA COVERAGE:")
    print(f"  With section_hint AND day_hint: {len(with_both)} ({len(with_both)/len(links)*100:.1f}%)")
    print(f"  With section_hint only: {len(with_section_only)} ({len(with_section_only)/len(links)*100:.1f}%)")
    print(f"  With day_hint only: {len(with_day_only)} ({len(with_day_only)/len(links)*100:.1f}%)")
    print(f"  With neither: {len(with_neither)} ({len(with_neither)/len(links)*100:.1f}%)")
    
    # Analyze text patterns
    print("\n\nTEXT PATTERNS:")
    
    # Count duplicates
    text_counts = defaultdict(int)
    for link in links:
        text_counts[link['text']] += 1
    
    duplicates = {text: count for text, count in text_counts.items() if count > 1}
    print(f"  Unique link texts: {len(text_counts)}")
    print(f"  Duplicate link texts: {len(duplicates)}")
    
    if duplicates:
        print(f"\n  Most duplicated:")
        for text, count in sorted(duplicates.items(), key=lambda x: x[1], reverse=True)[:10]:
            print(f"    '{text[:50]}': {count} times")
    
    # Analyze section hints
    print("\n\nSECTION HINTS:")
    section_counts = defaultdict(int)
    for link in links:
        hint = link.get('section_hint', 'NONE')
        section_counts[hint] += 1
    
    for section, count in sorted(section_counts.items(), key=lambda x: x[1], reverse=True):
        print(f"  {section}: {count} links")
    
    return links


def analyze_output_placement(output_path):
    """Analyze where links ended up in output."""
    
    print("\n\n" + "="*100)
    print("ANALYZING OUTPUT PLACEMENT")
    print("="*100)
    print()
    
    doc = Document(output_path)
    rels = doc.part.rels
    
    # Find Grade 3 ELA section (first slot)
    inline_links = []
    fallback_links = []
    
    # Table 1 is Grade 3 ELA daily plans
    table = doc.tables[1]
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            # Get row label
            row_label = row.cells[0].text.strip() if row.cells else ""
            
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                for hyperlink in hyperlinks:
                    r_id = hyperlink.get(qn('r:id'))
                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                    
                    url = None
                    if r_id and r_id in rels:
                        url = rels[r_id].target_ref
                    
                    if text and url:
                        inline_links.append({
                            'text': text,
                            'url': url,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'row_label': row_label,
                            'cell_text': cell.text.strip()[:100]
                        })
    
    # Find fallback links in paragraphs
    in_fallback = False
    for para_idx, para in enumerate(doc.paragraphs):
        if 'Referenced Links' in para.text:
            in_fallback = True
            continue
        
        if in_fallback:
            hyperlinks = para._element.xpath('.//w:hyperlink')
            for hyperlink in hyperlinks:
                r_id = hyperlink.get(qn('r:id'))
                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                
                url = None
                if r_id and r_id in rels:
                    url = rels[r_id].target_ref
                
                if text and url:
                    fallback_links.append({
                        'text': text,
                        'url': url,
                        'para_idx': para_idx
                    })
            
            # Stop at next section
            if 'Name:' in para.text or 'Required Signatures' in para.text:
                break
    
    print(f"Inline links (in table): {len(inline_links)}")
    print(f"Fallback links (in paragraphs): {len(fallback_links)}")
    
    # Analyze inline link placement
    print("\n\nINLINE LINKS BY ROW:")
    row_counts = defaultdict(int)
    for link in inline_links:
        row_counts[link['row_label']] += 1
    
    for row_label, count in sorted(row_counts.items(), key=lambda x: x[1], reverse=True):
        print(f"  {row_label[:40]}: {count} links")
    
    return inline_links, fallback_links


def compare_inline_vs_fallback(input_links, inline_links, fallback_links):
    """Compare characteristics of inline vs fallback links."""
    
    print("\n\n" + "="*100)
    print("COMPARING INLINE VS FALLBACK CHARACTERISTICS")
    print("="*100)
    print()
    
    # Create URL mappings
    input_by_url = {link['url']: link for link in input_links}
    inline_urls = set(link['url'] for link in inline_links)
    fallback_urls = set(link['url'] for link in fallback_links)
    
    # Analyze inline links
    print("INLINE LINKS (47 links that placed correctly):")
    print("-" * 80)
    
    inline_with_section = 0
    inline_with_day = 0
    inline_with_both = 0
    inline_text_lengths = []
    inline_unique_texts = set()
    
    for url in inline_urls:
        if url in input_by_url:
            input_link = input_by_url[url]
            has_section = bool(input_link.get('section_hint'))
            has_day = bool(input_link.get('day_hint'))
            
            if has_section:
                inline_with_section += 1
            if has_day:
                inline_with_day += 1
            if has_section and has_day:
                inline_with_both += 1
            
            inline_text_lengths.append(len(input_link['text']))
            inline_unique_texts.add(input_link['text'])
    
    print(f"  With section_hint: {inline_with_section}/{len(inline_urls)} ({inline_with_section/len(inline_urls)*100:.1f}%)")
    print(f"  With day_hint: {inline_with_day}/{len(inline_urls)} ({inline_with_day/len(inline_urls)*100:.1f}%)")
    print(f"  With both: {inline_with_both}/{len(inline_urls)} ({inline_with_both/len(inline_urls)*100:.1f}%)")
    print(f"  Avg text length: {sum(inline_text_lengths)/len(inline_text_lengths):.1f} chars")
    print(f"  Unique texts: {len(inline_unique_texts)}/{len(inline_urls)} ({len(inline_unique_texts)/len(inline_urls)*100:.1f}%)")
    
    # Show examples
    print(f"\n  Example inline links:")
    for url in list(inline_urls)[:5]:
        if url in input_by_url:
            link = input_by_url[url]
            print(f"    • {link['text'][:60]}")
            print(f"      section: {link.get('section_hint', 'NONE')}, day: {link.get('day_hint', 'NONE')}")
    
    # Analyze fallback links
    print("\n\nFALLBACK LINKS (77 links that failed to place):")
    print("-" * 80)
    
    fallback_with_section = 0
    fallback_with_day = 0
    fallback_with_both = 0
    fallback_text_lengths = []
    fallback_unique_texts = set()
    
    for url in fallback_urls:
        if url in input_by_url:
            input_link = input_by_url[url]
            has_section = bool(input_link.get('section_hint'))
            has_day = bool(input_link.get('day_hint'))
            
            if has_section:
                fallback_with_section += 1
            if has_day:
                fallback_with_day += 1
            if has_section and has_day:
                fallback_with_both += 1
            
            fallback_text_lengths.append(len(input_link['text']))
            fallback_unique_texts.add(input_link['text'])
    
    print(f"  With section_hint: {fallback_with_section}/{len(fallback_urls)} ({fallback_with_section/len(fallback_urls)*100:.1f}%)")
    print(f"  With day_hint: {fallback_with_day}/{len(fallback_urls)} ({fallback_with_day/len(fallback_urls)*100:.1f}%)")
    print(f"  With both: {fallback_with_both}/{len(fallback_urls)} ({fallback_with_both/len(fallback_urls)*100:.1f}%)")
    print(f"  Avg text length: {sum(fallback_text_lengths)/len(fallback_text_lengths):.1f} chars")
    print(f"  Unique texts: {len(fallback_unique_texts)}/{len(fallback_urls)} ({len(fallback_unique_texts)/len(fallback_urls)*100:.1f}%)")
    
    # Show examples
    print(f"\n  Example fallback links:")
    for url in list(fallback_urls)[:10]:
        if url in input_by_url:
            link = input_by_url[url]
            print(f"    • {link['text'][:60]}")
            print(f"      section: {link.get('section_hint', 'NONE')}, day: {link.get('day_hint', 'NONE')}")
    
    # Key differences
    print("\n\nKEY DIFFERENCES:")
    print("-" * 80)
    print(f"  Metadata completeness:")
    print(f"    Inline with both hints: {inline_with_both/len(inline_urls)*100:.1f}%")
    print(f"    Fallback with both hints: {fallback_with_both/len(fallback_urls)*100:.1f}%")
    print(f"    Difference: {(inline_with_both/len(inline_urls) - fallback_with_both/len(fallback_urls))*100:.1f}%")
    
    print(f"\n  Text uniqueness:")
    print(f"    Inline unique: {len(inline_unique_texts)/len(inline_urls)*100:.1f}%")
    print(f"    Fallback unique: {len(fallback_unique_texts)/len(fallback_urls)*100:.1f}%")
    print(f"    Difference: {(len(inline_unique_texts)/len(inline_urls) - len(fallback_unique_texts)/len(fallback_urls))*100:.1f}%")


def main():
    print("="*100)
    print("LINK PATTERN ANALYSIS - UNDERSTANDING SUCCESS VS FAILURE")
    print("="*100)
    print("\nFocusing on Davies Grade 3 ELA (47 inline, 77 fallback)\n")
    
    # Analyze input
    input_links = analyze_input_links(INPUT_FILE)
    
    # Analyze output
    inline_links, fallback_links = analyze_output_placement(OUTPUT_FILE)
    
    # Compare
    compare_inline_vs_fallback(input_links, inline_links, fallback_links)
    
    print("\n\n" + "="*100)
    print("CONCLUSIONS")
    print("="*100)
    print("\nThis analysis shows:")
    print("1. What metadata patterns distinguish successful vs failed placements")
    print("2. Whether duplicate text is the main issue")
    print("3. Which sections/rows are most problematic")
    print("4. Whether parser improvements could help")


if __name__ == '__main__':
    sys.exit(main())
