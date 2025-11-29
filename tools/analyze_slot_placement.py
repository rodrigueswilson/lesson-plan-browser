"""
Analyze hyperlink placement per slot.

For each slot:
1. Count links in daily plans table (inline)
2. Count links in Referenced Links section (fallback)
3. Compare metadata of inline vs fallback links
4. Identify patterns: why did some links fail to place inline?
"""

import sys
from pathlib import Path
from collections import defaultdict
import json

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.docx_parser import DOCXParser
from docx import Document
from docx.oxml.ns import qn


def extract_links_with_location(doc: Document):
    """Extract all hyperlinks with their location info."""
    
    rels = doc.part.rels
    all_links = []
    
    for table_idx, table in enumerate(doc.tables):
        # Identify table type
        table_type = "unknown"
        section_name = None
        
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()
            
            # Metadata table
            if "Name:" in first_cell:
                table_type = "metadata"
                # Extract subject and grade to identify slot
                subject = ""
                grade = ""
                for cell in table.rows[0].cells:
                    cell_text = cell.text.strip()
                    if "Subject:" in cell_text:
                        subject = cell_text.replace("Subject:", "").strip()
                    elif "Grade:" in cell_text:
                        grade = cell_text.replace("Grade:", "").strip()
                
                if subject:
                    section_name = f"Grade {grade} - {subject}" if grade else subject
            
            # Daily plans table
            elif "MONDAY" in first_cell or "Monday" in first_cell:
                table_type = "daily_plans"
            
            # Referenced Links table
            elif "Referenced Links" in first_cell or "Referenced Media" in first_cell:
                table_type = "referenced_links"
        
        # Extract links from this table
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    hyperlinks = para._element.xpath('.//w:hyperlink')
                    for hyperlink in hyperlinks:
                        r_id = hyperlink.get(qn('r:id'))
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        
                        # Resolve URL
                        url = None
                        if r_id and r_id in rels:
                            url = rels[r_id].target_ref
                        
                        if text and url:
                            # Get cell text for context
                            cell_text = cell.text.strip()[:200]
                            
                            # Get row label (first cell of row)
                            row_label = ""
                            if row.cells:
                                row_label = row.cells[0].text.strip()[:50]
                            
                            all_links.append({
                                'text': text,
                                'url': url,
                                'table_idx': table_idx,
                                'table_type': table_type,
                                'section_name': section_name,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'row_label': row_label,
                                'cell_text': cell_text,
                                'location': f"T{table_idx}R{row_idx}C{cell_idx}"
                            })
    
    return all_links


def group_by_slot(links):
    """Group links by slot (teacher section)."""
    
    slots = defaultdict(lambda: {'inline': [], 'fallback': []})
    current_slot = None
    
    for link in links:
        # Track current slot based on metadata tables
        if link['table_type'] == 'metadata' and link['section_name']:
            current_slot = link['section_name']
        
        # Categorize link
        if link['table_type'] == 'daily_plans':
            if current_slot:
                slots[current_slot]['inline'].append(link)
        elif link['table_type'] == 'referenced_links':
            if current_slot:
                slots[current_slot]['fallback'].append(link)
    
    return slots


def analyze_metadata_patterns(input_path: str, output_links: list):
    """
    Compare input metadata with output placement.
    Find patterns in what succeeded vs failed.
    """
    
    # Parse input
    parser = DOCXParser(input_path)
    input_links = parser.extract_hyperlinks()
    
    # Create URL -> metadata mapping
    input_metadata = {}
    for link in input_links:
        if link.get('url'):
            input_metadata[link['url']] = {
                'text': link['text'],
                'section_hint': link.get('section_hint'),
                'day_hint': link.get('day_hint'),
                'context': link.get('context_snippet', '')[:100]
            }
    
    # Analyze output links
    inline_with_meta = []
    inline_without_meta = []
    fallback_with_meta = []
    fallback_without_meta = []
    
    for link in output_links:
        url = link['url']
        meta = input_metadata.get(url, {})
        
        has_section = bool(meta.get('section_hint'))
        has_day = bool(meta.get('day_hint'))
        has_meta = has_section or has_day
        
        link_with_meta = {**link, **meta, 'has_metadata': has_meta}
        
        if link['table_type'] == 'daily_plans':
            if has_meta:
                inline_with_meta.append(link_with_meta)
            else:
                inline_without_meta.append(link_with_meta)
        elif link['table_type'] == 'referenced_links':
            if has_meta:
                fallback_with_meta.append(link_with_meta)
            else:
                fallback_without_meta.append(link_with_meta)
    
    return {
        'inline_with_meta': inline_with_meta,
        'inline_without_meta': inline_without_meta,
        'fallback_with_meta': fallback_with_meta,
        'fallback_without_meta': fallback_without_meta
    }


def main():
    """Analyze W43 output file."""
    
    print("="*100)
    print("HYPERLINK PLACEMENT ANALYSIS - PER SLOT")
    print("="*100)
    
    output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')
    
    if not output_file.exists():
        print(f"❌ File not found: {output_file}")
        return 1
    
    print(f"\nAnalyzing: {output_file.name}\n")
    
    # Extract all links
    doc = Document(output_file)
    all_links = extract_links_with_location(doc)
    
    print(f"Total hyperlinks found: {len(all_links)}\n")
    
    # Group by slot
    slots = group_by_slot(all_links)
    
    print(f"{'='*100}")
    print("PER-SLOT SUMMARY")
    print(f"{'='*100}\n")
    
    print(f"{'Slot':<30} {'Inline':<10} {'Fallback':<10} {'Total':<10} {'Inline %':<10}")
    print(f"{'-'*80}")
    
    total_inline = 0
    total_fallback = 0
    
    for slot_name in sorted(slots.keys()):
        slot_data = slots[slot_name]
        inline_count = len(slot_data['inline'])
        fallback_count = len(slot_data['fallback'])
        total = inline_count + fallback_count
        inline_pct = (inline_count / total * 100) if total > 0 else 0
        
        print(f"{slot_name:<30} {inline_count:<10} {fallback_count:<10} {total:<10} {inline_pct:<10.1f}%")
        
        total_inline += inline_count
        total_fallback += fallback_count
    
    print(f"{'-'*80}")
    total_all = total_inline + total_fallback
    overall_pct = (total_inline / total_all * 100) if total_all > 0 else 0
    print(f"{'TOTAL':<30} {total_inline:<10} {total_fallback:<10} {total_all:<10} {overall_pct:<10.1f}%")
    
    # Analyze fallback links in detail
    print(f"\n\n{'='*100}")
    print("FALLBACK LINKS ANALYSIS")
    print(f"{'='*100}\n")
    
    for slot_name in sorted(slots.keys()):
        fallback_links = slots[slot_name]['fallback']
        
        if fallback_links:
            print(f"\n{slot_name} - {len(fallback_links)} fallback links:")
            print(f"{'-'*80}")
            
            for i, link in enumerate(fallback_links[:10], 1):  # Show first 10
                print(f"{i}. {link['text'][:60]}")
                print(f"   Row: {link['row_label'][:40]}")
                print(f"   Context: {link['cell_text'][:60]}...")
                print()
            
            if len(fallback_links) > 10:
                print(f"   ... and {len(fallback_links) - 10} more\n")
    
    # Analyze with input metadata (for first slot as example)
    print(f"\n\n{'='*100}")
    print("METADATA ANALYSIS (Davies as example)")
    print(f"{'='*100}\n")
    
    input_files = {
        'Davies': Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Davies Lesson Plans.docx'),
        'Lang': Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Lang Lesson Plans 10_20_25-10_24_25.docx'),
        'Savoca': Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Ms. Savoca-10_20_25-10_25_25 Lesson plans.docx')
    }
    
    # Analyze Davies
    if input_files['Davies'].exists():
        davies_output_links = []
        for link in all_links:
            if link.get('section_name') and 'Davies' in link['section_name']:
                davies_output_links.append(link)
        
        if davies_output_links:
            patterns = analyze_metadata_patterns(str(input_files['Davies']), davies_output_links)
            
            print(f"Links WITH metadata (section_hint or day_hint):")
            print(f"  Inline: {len(patterns['inline_with_meta'])}")
            print(f"  Fallback: {len(patterns['fallback_with_meta'])}")
            
            print(f"\nLinks WITHOUT metadata:")
            print(f"  Inline: {len(patterns['inline_without_meta'])}")
            print(f"  Fallback: {len(patterns['fallback_without_meta'])}")
            
            # Show examples of fallback links WITH metadata (these are the failures)
            if patterns['fallback_with_meta']:
                print(f"\n\n⚠️  LINKS WITH METADATA BUT STILL IN FALLBACK ({len(patterns['fallback_with_meta'])}):")
                print(f"These SHOULD have been placed inline but weren't:\n")
                
                for i, link in enumerate(patterns['fallback_with_meta'][:10], 1):
                    print(f"{i}. {link['text'][:60]}")
                    print(f"   section_hint: {link.get('section_hint', 'NONE')}")
                    print(f"   day_hint: {link.get('day_hint', 'NONE')}")
                    print(f"   context: {link.get('context', '')[:60]}...")
                    print()
    
    # Save detailed results
    results_path = Path('d:/LP/docs/validation/slot_analysis_results.json')
    results_path.parent.mkdir(parents=True, exist_ok=True)
    
    results = {
        'summary': {
            'total_inline': total_inline,
            'total_fallback': total_fallback,
            'inline_rate': overall_pct
        },
        'per_slot': {
            slot_name: {
                'inline': len(data['inline']),
                'fallback': len(data['fallback']),
                'fallback_links': [{'text': l['text'], 'row': l['row_label']} for l in data['fallback']]
            }
            for slot_name, data in slots.items()
        }
    }
    
    with open(results_path, 'w') as f:
        json.dump(results, f, indent=2)
    
    print(f"\n\n💾 Detailed results saved to: {results_path}")
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
