"""
Markdown to DOCX Converter - Convert markdown formatting to DOCX runs.

Handles:
- Bold: **text** or __text__
- Italic: *text* or _text_
- Bullet lists: - item or * item
- Numbered lists: 1. item
- Line breaks and paragraphs
"""

import re
from typing import List, Tuple
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def sanitize_xml_text(text: str) -> str:
    """
    Sanitize text for XML compatibility in DOCX.
    
    Removes NULL bytes and control characters that are not allowed in XML.
    Preserves common whitespace characters (space, tab, newline, carriage return).
    
    Args:
        text: Input text that may contain invalid XML characters
        
    Returns:
        Sanitized text safe for XML/DOCX
    """
    if not text:
        return text
    
    # Characters allowed in XML 1.0:
    # - #x9 (tab)
    # - #xA (line feed/newline)
    # - #xD (carriage return)
    # - #x20-#xD7FF (most printable characters)
    # - #xE000-#xFFFD (extended characters)
    # - #x10000-#x10FFFF (supplementary characters)
    
    # Remove NULL bytes and other invalid control characters
    # Keep: tab (0x09), newline (0x0A), carriage return (0x0D)
    result = []
    for char in text:
        code = ord(char)
        # Allow: tab, newline, carriage return, and printable characters
        if code == 0x09 or code == 0x0A or code == 0x0D:
            result.append(char)
        elif code >= 0x20:  # Printable characters and above
            # Exclude surrogate pairs range (0xD800-0xDFFF) which are invalid
            if not (0xD800 <= code <= 0xDFFF):
                result.append(char)
        # Skip: NULL bytes (0x00) and other control characters (0x01-0x08, 0x0B-0x0C, 0x0E-0x1F)
    
    return ''.join(result)


class MarkdownToDocx:
    """Convert markdown text to DOCX formatted paragraphs."""
    
    # Regex patterns for markdown
    BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*|__(.+?)__')
    ITALIC_PATTERN = re.compile(r'\*(.+?)\*|_(.+?)_')
    BULLET_PATTERN = re.compile(r'^[\-\*]\s+(.+)$')
    NUMBERED_PATTERN = re.compile(r'^\d+\.\s+(.+)$')
    HYPERLINK_PATTERN = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    
    @staticmethod
    def add_formatted_text(paragraph, text: str):
        """
        Add formatted text to a paragraph, handling markdown formatting.
        
        Args:
            paragraph: python-docx paragraph object
            text: Text with markdown formatting
        """
        if not text:
            return
        
        # Sanitize text for XML compatibility
        text = sanitize_xml_text(text)
        
        # Split by bold markers first
        parts = []
        last_end = 0
        
        for match in MarkdownToDocx.BOLD_PATTERN.finditer(text):
            # Add text before bold
            if match.start() > last_end:
                parts.append(('normal', text[last_end:match.start()]))
            
            # Add bold text
            bold_text = match.group(1) or match.group(2)
            parts.append(('bold', bold_text))
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            parts.append(('normal', text[last_end:]))
        
        # Now process each part for hyperlinks and italics
        for style, part_text in parts:
            is_bold = (style == 'bold')
            MarkdownToDocx._add_text_with_hyperlinks(paragraph, part_text, bold=is_bold)

    @staticmethod
    def _add_text_with_hyperlinks(paragraph, text: str, bold: bool = False):
        """Add text with hyperlink formatting."""
        parts = []
        last_end = 0
        
        for match in MarkdownToDocx.HYPERLINK_PATTERN.finditer(text):
            # Add text before link
            if match.start() > last_end:
                parts.append(('normal', text[last_end:match.start()]))
            
            # Add link
            link_text = match.group(1)
            link_url = match.group(2)
            parts.append(('link', (link_text, link_url)))
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            parts.append(('normal', text[last_end:]))
        
        # Proactive Whitespace Enforcement:
        # We rewrite the parts list to ensure spaces around links where appropriate.
        enforced_parts = []
        for i, (style, val) in enumerate(parts):
            if style == 'link':
                link_text, link_url = val
                # 1. Check for leading space
                if enforced_parts:
                    prev_style, prev_val = enforced_parts[-1]
                    if prev_style == 'normal' and prev_val:
                        # Add space if previous text doesn't end with space or opening bracket
                        if not prev_val[-1].isspace() and prev_val[-1] not in '([{':
                            enforced_parts[-1] = ('normal', prev_val + ' ')
                    elif prev_style == 'link':
                        # Two adjacent links: add a separating space
                        enforced_parts.append(('normal', ' '))
                
                enforced_parts.append((style, val))
            else:
                # This is normal text. If the part before it was a link, check for leading space.
                if enforced_parts and enforced_parts[-1][0] == 'link':
                    # Add space if current text doesn't start with space or closing punctuation
                    if val and not val[0].isspace() and val[0] not in '.,;:)!?]}':
                        val = ' ' + val
                enforced_parts.append((style, val))

        # Process each part in order to maintain correct element order
        for style, val in enforced_parts:
            if style == 'link':
                link_text, link_url = val
                MarkdownToDocx._add_hyperlink(paragraph, link_text, link_url, bold=bold)
            else:
                MarkdownToDocx._add_text_with_italic(paragraph, val, bold=bold)
    
    @staticmethod
    def _add_text_with_italic(paragraph, text: str, bold: bool = False):
        """Add text with italic formatting."""
        parts = []
        last_end = 0
        
        for match in MarkdownToDocx.ITALIC_PATTERN.finditer(text):
            # Add text before italic
            if match.start() > last_end:
                parts.append(('normal', text[last_end:match.start()]))
            
            # Add italic text
            italic_text = match.group(1) or match.group(2)
            parts.append(('italic', italic_text))
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            parts.append(('normal', text[last_end:]))
        
        # Add runs
        for style, part_text in parts:
            if part_text:
                run = paragraph.add_run(part_text)
                if style == 'italic':
                    run.italic = True
                if bold:
                    run.bold = True
    
    @staticmethod
    def _add_bold_with_italic(paragraph, text: str):
        """Add bold text that may contain italic."""
        parts = []
        last_end = 0
        
        for match in MarkdownToDocx.ITALIC_PATTERN.finditer(text):
            # Add text before italic
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end:match.start()])
                run.bold = True
            
            # Add bold+italic text
            italic_text = match.group(1) or match.group(2)
            run = paragraph.add_run(italic_text)
            run.bold = True
            run.italic = True
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.bold = True
    
    @staticmethod
    def add_paragraph(cell, text: str, style=None):
        """
        Add a paragraph to a cell, handling markdown formatting.
        
        Args:
            cell: python-docx cell object
            text: Text with markdown formatting
            style: Optional paragraph style
            
        Returns:
            The created paragraph
        """
        # Check if it's a bullet list item
        bullet_match = MarkdownToDocx.BULLET_PATTERN.match(text)
        if bullet_match:
            para = cell.add_paragraph(style=style)
            # Add bullet character manually since template may not have List Bullet style
            # Use a real bullet character and a tab for proper alignment if possible, 
            # or just a bullet and space for simplicity.
            para.add_run('• ')  # Real bullet character
            MarkdownToDocx.add_formatted_text(para, bullet_match.group(1))
            return para
        
        # Check if it's a numbered list item
        numbered_match = MarkdownToDocx.NUMBERED_PATTERN.match(text)
        if numbered_match:
            para = cell.add_paragraph(style=style)
            # Keep the number prefix
            MarkdownToDocx.add_formatted_text(para, text)
            return para
        
        # Regular paragraph
        para = cell.add_paragraph(style=style)
        MarkdownToDocx.add_formatted_text(para, text)
        return para
    
    @staticmethod
    def add_multiline_text(cell, text: str, preserve_empty_lines: bool = True):
        """
        Add multi-line text to a cell, handling markdown formatting.
        
        Args:
            cell: python-docx cell object
            text: Multi-line text with markdown formatting
            preserve_empty_lines: Whether to preserve empty lines
        """
        if not text:
            return
        
        lines = text.split('\n')
        
        # Find first empty paragraph or create one
        # Skip non-empty paragraphs (like headers) to avoid merging content
        first_empty_para = None
        for para in cell.paragraphs:
            if not para.text.strip() and not para.runs:
                first_empty_para = para
                break
        
        # If no empty paragraph found, create one
        if first_empty_para is None:
            first_empty_para = cell.add_paragraph()
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines unless preserving
            if not line and not preserve_empty_lines:
                continue
            # Add paragraph
            if i == 0:
                # Use the first empty paragraph (found or created)
                MarkdownToDocx.add_formatted_text(first_empty_para, line)
            else:
                # Add new paragraph
                MarkdownToDocx.add_paragraph(cell, line)
    
    @staticmethod
    def clean_markdown(text: str) -> str:
        """
        Remove markdown formatting from text, leaving plain text.
        
        Args:
            text: Text with markdown formatting
            
        Returns:
            Plain text without markdown
        """
        if not text:
            return ""
        
        # Remove bold
        text = MarkdownToDocx.BOLD_PATTERN.sub(r'\1\2', text)
        
        # Remove italic
        text = MarkdownToDocx.ITALIC_PATTERN.sub(r'\1\2', text)
        
        # Remove bullet markers
        text = MarkdownToDocx.BULLET_PATTERN.sub(r'\1', text)
        
        # Remove numbered markers
        text = MarkdownToDocx.NUMBERED_PATTERN.sub(r'\1', text)
        
        # Remove hyperlinks (leave text)
        text = MarkdownToDocx.HYPERLINK_PATTERN.sub(r'\1', text)
        
        return text

    @staticmethod
    def _add_hyperlink(paragraph, text: str, url: str, insert_at: int = None, bold: bool = False):
        """Add a hyperlink to a paragraph.
        
        Args:
            paragraph: python-docx paragraph object
            text: Link text
            url: Link URL
            insert_at: Optional index to insert at (for maintaining order). If None, appends to end.
            bold: Whether to make the hyperlink text bold
        """
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Times New Roman 8pt so links match the rest of the lesson plan (avoid Arial 11pt inheritance)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rPr.append(rFonts)

        # Color and Underline (traditional link style)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "16")
        rPr.append(szCs)

        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        
        # Insert at specified position or append to end
        if insert_at is not None and insert_at < len(paragraph._p):
            paragraph._p.insert(insert_at, hyperlink)
        else:
            paragraph._p.append(hyperlink)
