"""
DOCX Renderer - Convert validated JSON to DOCX using district template.

This renderer:
1. Loads the district template (input/Lesson Plan Template SY'25-26.docx)
2. Fills in metadata (teacher name, grade, subject, week, homeroom)
3. Fills in daily plans for Monday-Friday
4. Preserves all original formatting
"""

import base64
import json
import sys
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.run import Run
import re
import time
import json

from backend.performance_tracker import get_tracker
from backend.telemetry import logger
from backend.services.sorting_utils import sort_slots
from backend.utils.metadata_utils import (
    get_homeroom,
    get_subject,
    get_teacher_name,
)
from tools.table_structure import StructureMetadata, TableStructureDetector

# Feature flag for fuzzy matching threshold
# ADJUSTED: Lowered to 0.50 to handle bilingual/rephrased content better
# Combined with hint-based boosting for more accurate placement
FUZZY_MATCH_THRESHOLD = 0.50  # Lowered from 0.65 for bilingual support

# Handle imports for both CLI and module usage
try:
    from tools.markdown_to_docx import MarkdownToDocx
except ImportError:
    from markdown_to_docx import MarkdownToDocx


def sanitize_xml_text(text: str) -> str:
    """
    Sanitize text for XML compatibility in DOCX.
    
    Removes NULL bytes and control characters that are not allowed in XML.
    Preserves common whitespace characters (space, tab, newline, carriage return).
    
    Args:
        text: Input text that may contain invalid XML characters
        
    Returns:
        Sanitized text safe for XML/DOCX
    """
    if not text:
        return text
    
    # Characters allowed in XML 1.0:
    # - #x9 (tab)
    # - #xA (line feed/newline)
    # - #xD (carriage return)
    # - #x20-#xD7FF (most printable characters)
    # - #xE000-#xFFFD (extended characters)
    # - #x10000-#x10FFFF (supplementary characters)
    
    # Remove NULL bytes and other invalid control characters
    # Keep: tab (0x09), newline (0x0A), carriage return (0x0D)
    result = []
    for char in text:
        code = ord(char)
        # Allow: tab, newline, carriage return, and printable characters
        if code == 0x09 or code == 0x0A or code == 0x0D:
            result.append(char)
        elif code >= 0x20:  # Printable characters and above
            # Exclude surrogate pairs range (0xD800-0xDFFF) which are invalid
            if not (0xD800 <= code <= 0xDFFF):
                result.append(char)
        # Skip: NULL bytes (0x00) and other control characters (0x01-0x08, 0x0B-0x0C, 0x0E-0x1F)
    
    return ''.join(result)


def is_signature_table(table) -> bool:
    """Check if a table is a signature table based on common headers."""
    if not table.rows:
        return False
    # Check first few cells for signature-related text
    try:
        first_row_text = "".join(cell.text for cell in table.rows[0].cells).lower()
        return any(x in first_row_text for x in ["signature", "approver", "approved"])
    except Exception:
        return False

class DOCXRenderer:
    """Render validated JSON lesson plans to DOCX format."""

    # Table indices in the template document
    METADATA_TABLE_IDX = 0  # First table contains metadata (Name, Grade, etc.)
    DAILY_PLANS_TABLE_IDX = 1  # Second table contains daily lesson plans

    # Row indices in the daily plans table
    UNIT_LESSON_ROW = 1  # Unit/Lesson row (first data row after headers)

    def __init__(self, template_path: str):
        """Initialize renderer with template path.

        Args:
            template_path: Path to the DOCX template file
        """
        self.template_path = template_path
        self.is_originals = False # Initialized to False, set by renderer logic
        self.logger = logger.bind(component="docx_renderer")
        self.structure_detector = TableStructureDetector()
        
        # Cache template in memory to avoid repeated disk I/O when reused
        self.template_buffer = None
        try:
            with open(template_path, "rb") as f:
                self.template_buffer = BytesIO(f.read())
        except Exception as e:
            self.logger.error(f"Failed to cache template buffer: {e}")

        # Initialize structure metadata
        self.structure_metadata = None
        self._initialize_structure()

        # Reset per-render state
        self._reset_state()

    def _reset_state(self):
        """Reset internal state before each render call."""
        self.placement_stats = {
            "coordinate": 0,
            "label_day": 0,
            "fuzzy": 0,
            "fallback": 0,
        }
        self.current_file = None
        self.current_teacher = None
        self._current_metadata = {}
        self.is_originals = False

    def _initialize_structure(self):
        """Initialize table structure metadata from template."""
        try:
            # Use buffer if available for structure detection too
            source = self.template_buffer if self.template_buffer else self.template_path
            if source:
                if hasattr(source, "seek"):
                    source.seek(0)
                doc = Document(source)
                if len(doc.tables) > 1:
                
                    # Find the daily plans table dynamically
                    # It's usually the first table after the metadata table (index 0)
                    # that is not a signature table.
                    daily_plans_table_found = False
                    for i in range(1, len(doc.tables)): # Start from index 1, assuming metadata is at 0
                        table = doc.tables[i]
                        if not is_signature_table(table):
                            self.DAILY_PLANS_TABLE_IDX = i
                            self.structure_metadata = self.structure_detector.detect_structure(table)
                            daily_plans_table_found = True
                            break
                    
                    if daily_plans_table_found:
                        self.logger.info(
                            "template_structure_detected",
                            extra={
                                "type": self.structure_metadata.structure_type,
                                "rows": self.structure_metadata.num_rows,
                                "cols": self.structure_metadata.num_cols,
                                "has_day_row": self.structure_metadata.has_day_row,
                                "daily_plans_table_idx": self.DAILY_PLANS_TABLE_IDX,
                            },
                        )
                    else:
                        self.logger.error("Could not locate daily plans table in template")
                else:
                    self.logger.warning(
                        "Template has fewer than 2 tables, cannot detect structure"
                    )
        except Exception as e:
            self.logger.error(f"Failed to initialize structure: {e}")

    def _get_row_index(self, label: str) -> int:
        """Get row index for a given label using structure metadata."""
        if self.structure_metadata:
            idx = self.structure_metadata.get_row_index(label)
            if idx is not None:
                # Apply offset if needed (e.g. for day row)
                return idx + self.structure_metadata.row_offset
        return -1

    def _get_col_index(self, day: str) -> int:
        """Get column index for a given day using structure metadata."""
        if self.structure_metadata:
            idx = self.structure_metadata.get_col_index(day)
            return idx if idx is not None else -1
        return -1

    def render(
        self,
        json_data: Dict,
        output_path: Any,  # Can be str, Path, or BytesIO
        plan_id: str = None,
        skip_fallback_sections: bool = False,
    ) -> Any:
        """
        Render JSON data to DOCX with semantic anchoring for media.

        Args:
            json_data: Validated lesson plan JSON (supports both single-slot and multi-slot)
            output_path: Path to save DOCX file or BytesIO stream
            plan_id: Optional plan ID for performance tracking
            skip_fallback_sections: If True, don't append "Referenced Links" or "Attached Images" sections
                                   (used when rendering slots for later merging)

        Returns:
            If skip_fallback_sections=True: Tuple (success: bool, pending_hyperlinks: List, pending_images: List)
            Otherwise: True if successful, False otherwise
        """
        tracker = get_tracker()
        
        # Reset per-render state
        self._reset_state()

        # Set context for logging
        is_stream = hasattr(output_path, "write")
        self.current_file = "stream.docx" if is_stream else Path(output_path).name
        
        metadata = json_data.get("metadata", {})
        self.current_teacher = get_teacher_name(metadata)
        # Store metadata for use in helper methods
        self._current_metadata = metadata
        self.is_originals = (metadata.get("source_type") == "originals")

        try:
            # Load template
            if self.template_buffer:
                self.template_buffer.seek(0)
            doc = Document(self.template_buffer if self.template_buffer else self.template_path)
            
            # NOTE: We no longer modify global styles or docDefaults here to avoid 
            # corruption during docxcompose merging. All formatting is applied 
            # per-run in the daily grid cells.

            # CRITICAL: Ensure "Hyperlink" style exists to prevent "Styles 1" error
            # Word sometimes requires this even if we use direct formatting.
            if "Hyperlink" not in doc.styles:
               try:
                   from docx.enum.style import WD_STYLE_TYPE
                   from docx.shared import Pt, RGBColor
                   style = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
                   style.font.name = "Times New Roman"
                   style.font.size = Pt(8)
                   style.font.color.rgb = RGBColor(0, 0, 255)
                   style.font.underline = True
               except Exception as e:
                   # Ignore if style already exists (race condition)
                   pass

            # Prepare structure metadata
            # Calculate table width once for use throughout rendering
            section = doc.sections[0]
            available_width_emus = (
                section.page_width - section.left_margin - section.right_margin
            )
            available_width_inches = available_width_emus / 914400

            # Track metadata filling
            if plan_id:
                with tracker.track_operation(plan_id, "render_fill_metadata"):
                    self._fill_metadata(doc, json_data)
            else:
                self._fill_metadata(doc, json_data)

            # Normalize metadata table immediately after filling to prevent Word auto-resizing
            from tools.docx_utils import normalize_table_column_widths

            normalize_table_column_widths(
                doc.tables[self.METADATA_TABLE_IDX],
                total_width_inches=available_width_inches,
            )

            # Prepare media for semantic anchoring
            schema_version = json_data.get("_media_schema_version", "1.0")
            # v1.1 and v2.0 both use semantic anchoring (v2.0 adds coordinates)
            pending_hyperlinks = (
                json_data.get("_hyperlinks", []).copy()
                if schema_version in ["1.1", "2.0"]
                else []
            )
            pending_images = (
                json_data.get("_images", []).copy()
                if schema_version in ["1.1", "2.0"]
                else []
            )

            # v2.0: Hyperlink placement is now handled during cell filling to ensure 
            # that inline replacement is always attempted before top-of-cell placement.
            # Coordination and fuzzy matching have been unified within _fill_cell.
            is_multi_slot = any(
                "slots" in day_data for day_data in json_data.get("days", {}).values()
            )

            # Extract slot metadata for filtering (if present)
            # In multi-slot batch processing, each lesson JSON has slot metadata
            metadata = json_data.get("metadata", {})
            slot_number = metadata.get("slot_number")
            subject = metadata.get("subject")

            # DEBUG: Log slot metadata extraction
            logger.info(
                "renderer_slot_metadata_extracted",
                extra={
                    "slot_number": slot_number,
                    "subject": subject,
                    "has_hyperlinks": len(json_data.get("_hyperlinks", [])),
                    "teacher": get_teacher_name(json_data.get("metadata", {})),
                },
            )

            # DIAGNOSTIC: Log renderer metadata extraction
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_renderer_extracted_metadata(
                slot_number,
                subject,
                len(json_data.get("_hyperlinks", [])),
                get_teacher_name(json_data.get("metadata", {})),
            )

            # Fill daily plans
            table = doc.tables[self.DAILY_PLANS_TABLE_IDX]
            
            def _fill_all_days():
                for day_name, day_data in json_data["days"].items():
                    # Skip if day_data is None
                    if not day_data:
                        continue
                        
                    col_idx = self._get_col_index(day_name)
                    # print(f"  [DEBUG RENDERER] Day: {day_name}, col_idx: {col_idx}")
                    if col_idx == -1:
                        continue

                    # Check if this day has multiple slots
                    if "slots" in day_data:
                        self._fill_multi_slot_day(
                            table,
                            col_idx,
                            day_data["slots"],
                            metadata=json_data.get("metadata", {}),
                            day_name=day_name,
                            pending_hyperlinks=pending_hyperlinks,
                            pending_images=pending_images,
                        )
                    else:
                        self._fill_single_slot_day(
                            table,
                            col_idx,
                            day_data,
                            day_name=day_name,
                            pending_hyperlinks=pending_hyperlinks,
                            pending_images=pending_images,
                            slot_number=slot_number,
                            subject=subject,
                        )  # Append unmatched media to fallback sections

            if plan_id:
                with tracker.track_operation(plan_id, "render_fill_days"):
                    _fill_all_days()
            else:
                _fill_all_days()

            if not skip_fallback_sections:
                if pending_hyperlinks or pending_images:
                    self._append_unmatched_media(doc, pending_hyperlinks, pending_images)
                    logger.info(
                        "unmatched_media_appended",
                        extra={
                            "hyperlinks": len(pending_hyperlinks),
                            "images": len(pending_images),
                        },
                    )

                # Legacy behavior for schema v1.0
                if schema_version == "1.0":
                    if "_images" in json_data and json_data["_images"]:
                        if plan_id:
                            with tracker.track_operation(plan_id, "render_insert_images"):
                                self._insert_images(doc, json_data["_images"])
                                logger.info(
                                    "images_inserted",
                                    extra={"count": len(json_data["_images"])},
                                )
                        else:
                            self._insert_images(doc, json_data["_images"])
                            logger.info(
                                "images_inserted",
                                extra={"count": len(json_data["_images"])},
                            )

                    if "_hyperlinks" in json_data and json_data["_hyperlinks"]:
                        if plan_id:
                            with tracker.track_operation(
                                plan_id, "render_restore_hyperlinks"
                            ):
                                self._restore_hyperlinks(doc, json_data["_hyperlinks"])
                                logger.info(
                                    "hyperlinks_restored",
                                    extra={"count": len(json_data["_hyperlinks"])},
                                )
                        else:
                            self._restore_hyperlinks(doc, json_data["_hyperlinks"])
                            logger.info(
                                    "hyperlinks_restored",
                                extra={"count": len(json_data["_hyperlinks"])},
                            )
            else:
                # Still log remaining count even if skipped
                logger.info(
                    "unmatched_media_skipped",
                    extra={
                        "hyperlinks": len(pending_hyperlinks),
                        "images": len(pending_images),
                        "note": "consolidation_active",
                    },
                )

            if self.is_originals:
                # Remove "Required Signatures" table and paragraphs
                # Instead of removing XML elements which can be unstable after merging,
                # we just clear the content of target sections.
                
                # CRITICAL: Use list(enumerate(doc.tables)) and reverse iteration 
                # to safely remove tables by index.
                all_tables = list(enumerate(doc.tables))
                for idx, table in reversed(all_tables):
                    # Only remove tables that are not metadata (0) or daily plans (1)
                    if idx > self.DAILY_PLANS_TABLE_IDX and is_signature_table(table):
                        # Clear cell content first (safer for some Word versions)
                        for row in table.rows:
                            for cell in row.cells:
                                cell.text = ""
                        # Remove table element from document XML
                        tbl = table._element
                        tbl.getparent().remove(tbl)

                # 2. Remove Signature paragraphs
                paras_to_remove = []
                for para in doc.paragraphs:
                    if "Required Signatures" in para.text:
                        paras_to_remove.append(para._element)
                
                for p_element in paras_to_remove:
                    try:
                        p_element.getparent().remove(p_element)
                    except Exception:
                        pass

            # Track table normalization
            # Use the same width calculated earlier for consistency

            if plan_id:
                with tracker.track_operation(plan_id, "render_normalize_tables"):
                    from tools.docx_utils import normalize_all_tables

                    table_count = normalize_all_tables(
                        doc, total_width_inches=available_width_inches
                    )
                    logger.info(
                        "tables_normalized",
                        extra={
                            "count": table_count,
                            "width_inches": available_width_inches,
                        },
                    )
            else:
                from tools.docx_utils import normalize_all_tables

                table_count = normalize_all_tables(
                    doc, total_width_inches=available_width_inches
                )
                logger.info(
                    "tables_normalized",
                    extra={
                        "count": table_count,
                        "width_inches": available_width_inches,
                    },
                )

            # Track file saving
            if plan_id:
                with tracker.track_operation(plan_id, "render_save_docx"):
                    if not is_stream:
                        output_path_obj = Path(output_path)
                        output_path_obj.parent.mkdir(parents=True, exist_ok=True)
                    doc.save(output_path)
            else:
                if not is_stream:
                    output_path_obj = Path(output_path)
                    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
                doc.save(output_path)

            logger.info("docx_render_success", extra={"output_path": str(output_path)})
            
            # #region agent log
            import json
            with open(r'd:\LP\.cursor\debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "run1",
                    "hypothesisId": "RENDER_SUCCESS",
                    "location": "docx_renderer.py:render:success",
                    "message": "Render completed successfully",
                    "data": {
                        "output_path": str(output_path),
                        "skip_fallback_sections": skip_fallback_sections,
                        "is_stream": hasattr(output_path, "write"),
                    },
                    "timestamp": int(__import__('time').time() * 1000)
                }) + '\n')
            # #endregion
            
            # RETURN VALUE CHANGE: Return unplaced media if skip_fallback_sections is True
            if skip_fallback_sections:
                return True, pending_hyperlinks, pending_images
            
            return True

        except Exception as e:
            # #region agent log
            import json
            import traceback
            with open(r'd:\LP\.cursor\debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "run1",
                    "hypothesisId": "RENDER_EXCEPTION",
                    "location": "docx_renderer.py:render:exception",
                    "message": "Renderer exception caught",
                    "data": {
                        "error": str(e),
                        "error_type": type(e).__name__,
                        "output_path": str(output_path) if "output_path" in locals() else None,
                        "traceback": traceback.format_exc()[:500],
                    },
                    "timestamp": int(__import__('time').time() * 1000)
                }) + '\n')
            # #endregion
            logger.exception(
                "docx_render_error",
                extra={
                    "output_path": str(output_path) if "output_path" in locals() else None,
                    "error": str(e),
                    "error_type": type(e).__name__,
                },
            )
            
            if skip_fallback_sections:
                return False, [], []
            
            # Re-raise the exception instead of returning False
            # This allows callers to see the actual error
            raise RuntimeError(
                f"Renderer failed to create DOCX file '{output_path}': {str(e)}"
            ) from e

    def _fill_metadata(self, doc: Document, json_data: Dict):
        """
        Fill metadata table (Table 0).

        Template structure:
        | Name: | Grade: | Homeroom: | Subject: | Week of: | Room: |

        For multi-slot lessons, extracts slot-specific metadata from the first slot
        found across all days, using standardized metadata utilities.

        Args:
            doc: Document object
            json_data: Full lesson plan JSON (supports both single-slot and multi-slot)
        """
        metadata = json_data.get("metadata", {})

        # Detect if this is a multi-slot lesson
        is_multi_slot = any(
            day_data and "slots" in day_data for day_data in json_data.get("days", {}).values()
        )

        # Extract representative slot for multi-slot lessons
        # Use first slot found across all days (sorted by slot_number)
        representative_slot = None
        if "days" in json_data and isinstance(json_data["days"], dict):
            all_slots = []
            for day_name, day_data in json_data["days"].items():
                if day_data and "slots" in day_data and isinstance(day_data["slots"], list):
                    all_slots.extend(day_data["slots"])

            if all_slots:
                # Sort by slot_number and use first slot
                sorted_slots = sorted(all_slots, key=lambda x: x.get("slot_number", 0))
                representative_slot = sorted_slots[0]

        table = doc.tables[self.METADATA_TABLE_IDX]
        row = table.rows[0]

        # Teacher Name (cell 0)
        # Always use get_teacher_name() which prioritizes primary_teacher_name from metadata
        # This ensures slot-specific teachers are shown, not the combined teacher_name
        teacher_name = get_teacher_name(metadata, slot=representative_slot)
        print(f"[DEBUG] DOCX_RENDERER: Metadata table teacher name")
        print(f"  metadata.primary_teacher_name: {metadata.get('primary_teacher_name')}")
        print(f"  metadata.teacher_name: {metadata.get('teacher_name')}")
        print(f"  get_teacher_name() result: {teacher_name}")
            
        cell = row.cells[0]
        # Use safer cell content clearing
        cell.text = ""
        para = cell.paragraphs[0]
        run1 = para.add_run(sanitize_xml_text("Name: "))
        self._force_font_arial10(run1, is_bold=True)
        run2 = para.add_run(sanitize_xml_text(teacher_name if teacher_name and teacher_name != "Unknown" else "Unknown"))
        self._force_font_arial10(run2, is_bold=False)

        # Grade (cell 1) - prioritize slot grade if available
        grade = None
        if representative_slot:
            grade = representative_slot.get("grade")
        if not grade or grade == "N/A":
            grade = metadata.get("grade")
        if grade and grade != "Unknown" and grade != "N/A":
            cell = row.cells[1]
            # Use safer cell content clearing
            cell.text = ""
            para = cell.paragraphs[0]
            run1 = para.add_run(sanitize_xml_text("Grade: "))
            self._force_font_arial10(run1, is_bold=True)
            run2 = para.add_run(sanitize_xml_text(grade))
            self._force_font_arial10(run2, is_bold=False)

        # Homeroom (cell 2) - use standardized helper to prevent leakage
        homeroom = get_homeroom(metadata, slot=representative_slot)
        if homeroom and homeroom != "Unknown":
            cell = row.cells[2]
            # Use safer cell content clearing
            cell.text = ""
            para = cell.paragraphs[0]
            run1 = para.add_run(sanitize_xml_text("Homeroom: "))
            self._force_font_arial10(run1, is_bold=True)
            run2 = para.add_run(sanitize_xml_text(homeroom))
            self._force_font_arial10(run2, is_bold=False)

        # Subject (cell 3) - use standardized helper
        subject = get_subject(metadata, slot=representative_slot)
        if subject and subject != "Unknown":
            cell = row.cells[3]
            # Use safer cell content clearing
            cell.text = ""
            para = cell.paragraphs[0]
            run1 = para.add_run(sanitize_xml_text("Subject: "))
            self._force_font_arial10(run1, is_bold=True)
            run2 = para.add_run(sanitize_xml_text(subject))
            self._force_font_arial10(run2, is_bold=False)

        # Week of (cell 4)
        week_of = metadata.get("week_of", "Unknown")
        if week_of and week_of != "Unknown":
            cell = row.cells[4]
            # Use safer cell content clearing
            cell.text = ""
            para = cell.paragraphs[0]
            run1 = para.add_run(sanitize_xml_text("Week of: "))
            self._force_font_arial10(run1, is_bold=True)
            run2 = para.add_run(sanitize_xml_text(week_of))
            self._force_font_arial10(run2, is_bold=False)

        # Room (cell 5, if template has it) - prioritize slot room if available
        if len(row.cells) > 5:
            room = None
            if representative_slot:
                room = representative_slot.get("room")
            if (
                not room
                or room == "N/A"
                or (isinstance(room, str) and not room.strip())
            ):
                room = metadata.get("room", "")
            if room and room.strip() and room != "N/A" and room != "Unknown":
                cell = row.cells[5]
                # Use safer cell content clearing
                cell.text = ""
                para = cell.paragraphs[0]
                run1 = para.add_run(sanitize_xml_text("Room: "))
                self._force_font_arial10(run1, is_bold=True)
                run2 = para.add_run(sanitize_xml_text(room))
                self._force_font_arial10(run2, is_bold=False)

    def _extract_unique_teachers(self, json_data: Dict) -> List[str]:
        """
        Extract unique teacher names from all slots in a multi-slot structure.

        Args:
            json_data: Full lesson plan JSON with days/slots structure

        Returns:
            Sorted list of unique teacher names
        """
        teachers = set()
        metadata = json_data.get("metadata", {})
        if "days" in json_data:
            for day_data in json_data["days"].values():
                if day_data and "slots" in day_data:
                    for slot in day_data["slots"]:
                        slot_teacher = get_teacher_name(metadata, slot=slot)
                        if slot_teacher and slot_teacher != "Unknown":
                            teachers.add(slot_teacher)
        return sorted(teachers)

    def _extract_unique_subjects(self, json_data: Dict) -> List[str]:
        """
        Extract unique subjects from all slots in a multi-slot structure.

        Args:
            json_data: Full lesson plan JSON with days/slots structure

        Returns:
            Sorted list of unique subjects
        """
        subjects = set()
        if "days" in json_data:
            for day_data in json_data["days"].values():
                if day_data and "slots" in day_data:
                    for slot in day_data["slots"]:
                        if slot.get("subject"):
                            subjects.add(slot["subject"])
        return sorted(subjects)

    def _abbreviate_content(
        self, content: str, num_slots: int, max_length: int = None
    ) -> str:
        """
        Abbreviate content based on number of slots to ensure it fits in template cells.

        Args:
            content: Original content text
            num_slots: Number of slots being displayed
            max_length: Optional max length override

        Returns:
            Abbreviated content with ellipsis if truncated
        """
        if not content:
            return content

        # Calculate max length based on number of slots
        if max_length is None:
            if num_slots <= 2:
                max_length = 500  # Full content
            elif num_slots == 3:
                max_length = 300  # Moderate truncation
            elif num_slots == 4:
                max_length = 200  # More truncation
            else:  # 5+ slots
                max_length = 150  # Aggressive truncation

        # Truncate if needed
        if len(content) > max_length:
            # Try to truncate at sentence boundary
            truncated = content[:max_length]
            last_period = truncated.rfind(".")
            last_newline = truncated.rfind("\n")

            # Use sentence boundary if found in last 30% of text
            boundary = max(last_period, last_newline)
            if boundary > max_length * 0.7:
                return truncated[: boundary + 1] + "..."
            else:
                # Just truncate at word boundary
                last_space = truncated.rfind(" ")
                if last_space > 0:
                    return truncated[:last_space] + "..."
                return truncated + "..."

        return content

    def _fill_day(
        self,
        doc,
        day_name: str,
        day_data: Dict,
        pending_hyperlinks: List[Dict] = None,
        pending_images: List[Dict] = None,
        slot_number: int = None,
        subject: str = None,
    ):
        """
        Fill a single day's data with media injection.

        Args:
            doc: Document object
            day_name: Day name (monday, tuesday, etc.)
            day_data: Day's lesson plan data (can be single-slot or multi-slot)
            pending_hyperlinks: List of hyperlinks awaiting placement
            pending_images: List of images awaiting placement
            slot_number: Slot number for filtering (if rendering single slot)
            subject: Subject name for filtering (if rendering single slot)
        """
        table = doc.tables[self.DAILY_PLANS_TABLE_IDX]
        col_idx = self.DAY_TO_COL[day_name.lower()]

        # Track hyperlinks before placement for observability
        hyperlinks_before = len(pending_hyperlinks) if pending_hyperlinks else 0

        # Check if this is multi-slot structure (has 'slots' array)
        if "slots" in day_data and isinstance(day_data["slots"], list):
            # Multi-slot: combine all slots for this day
            # Get metadata from instance variable set in render() method
            metadata = getattr(self, "_current_metadata", {})
            self._fill_multi_slot_day(
                table,
                col_idx,
                day_data["slots"],
                metadata=metadata,
                day_name=day_name,
                pending_hyperlinks=pending_hyperlinks,
                pending_images=pending_images,
            )
        else:
            # Single-slot: use existing logic with slot filtering
            self._fill_single_slot_day(
                table,
                col_idx,
                day_data,
                day_name=day_name,
                pending_hyperlinks=pending_hyperlinks,
                pending_images=pending_images,
                slot_number=slot_number,
                subject=subject,
            )

        # Log placement outcomes for observability
        hyperlinks_after = len(pending_hyperlinks) if pending_hyperlinks else 0
        placed_count = hyperlinks_before - hyperlinks_after

        if hyperlinks_before > 0:
            logger.info(
                "hyperlink_placement_outcome",
                extra={
                    "day": day_name,
                    "slot": slot_number,
                    "subject": subject,
                    "total_links": hyperlinks_before,
                    "placed_count": placed_count,
                    "remaining_count": hyperlinks_after,
                    "placement_rate": placed_count / hyperlinks_before
                    if hyperlinks_before > 0
                    else 0.0,
                },
            )

    def _fill_single_slot_day(
        self,
        table,
        col_idx: int,
        day_data: Dict,
        day_name: str = None,
        pending_hyperlinks: List[Dict] = None,
        pending_images: List[Dict] = None,
        slot_number: int = None,
        subject: str = None,
    ):
        """
        Fill a single slot's data for one day with media injection.

        Args:
            table: Table object
            col_idx: Column index for the day
            day_data: Day's lesson plan data
            day_name: Day name for media matching
            pending_hyperlinks: List of hyperlinks awaiting placement
            pending_images: List of images awaiting placement
            slot_number: Slot number for filtering hyperlinks/images
            subject: Subject name for filtering hyperlinks/images
        """
        # Check if this is a "No School" day
        unit_lesson = day_data.get("unit_lesson", "")
        if unit_lesson and "no school" in unit_lesson.lower():
            # For "No School" days, only fill the unit/lesson cell and clear others
            self._fill_cell(
                table,
                self.UNIT_LESSON_ROW,
                col_idx,
                unit_lesson,
                day_name=day_name,
                section_name="unit_lesson",
                pending_hyperlinks=pending_hyperlinks,
                pending_images=pending_images,
                current_slot_number=slot_number,
                current_subject=subject,
            )
            # Clear all other cells for this day
            for row_label in [
                "objective",
                "anticipatory",
                "instruction",
                "misconception",
                "assessment",
                "homework",
            ]:
                row_idx = self._get_row_index(row_label)
                if row_idx != -1:
                    table.rows[row_idx].cells[col_idx].text = ""
            return

        # Unit/Lesson (Row 1)
        self._fill_cell(
            table,
            self._get_row_index("unit"),
            col_idx,
            unit_lesson,
            day_name=day_name,
            section_name="unit_lesson",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

        # Objective (Row 2)
        objective_text = self._format_objective(day_data.get("objective", {}))
        self._fill_cell(
            table,
            self._get_row_index("objective"),
            col_idx,
            objective_text,
            day_name=day_name,
            section_name="objective",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

        # Anticipatory Set (Row 3)
        anticipatory_text = self._format_anticipatory_set(
            day_data.get("anticipatory_set", {})
        )
        self._fill_cell(
            table,
            self._get_row_index("anticipatory"),
            col_idx,
            anticipatory_text,
            day_name=day_name,
            section_name="anticipatory_set",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

        # Tailored Instruction (Row 4)
        instruction_text = self._format_tailored_instruction(
            day_data.get("tailored_instruction", {}),
            day_data.get("vocabulary_cognates"),
            day_data.get("sentence_frames"),
        )
        self._fill_cell(
            table,
            self._get_row_index("instruction"),
            col_idx,
            instruction_text,
            day_name=day_name,
            section_name="instruction",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

        # Misconceptions (Row 5)
        misconceptions_text = self._format_misconceptions(
            day_data.get("misconceptions", {})
        )
        self._fill_cell(
            table,
            self._get_row_index("misconception"),
            col_idx,
            misconceptions_text,
            day_name=day_name,
            section_name="misconceptions",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

        # Assessment (Row 6)
        assessment_text = self._format_assessment(day_data.get("assessment", {}))
        self._fill_cell(
            table,
            self._get_row_index("assessment"),
            col_idx,
            assessment_text,
            day_name=day_name,
            section_name="assessment",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

        # Homework (Row 7)
        homework_text = self._format_homework(day_data.get("homework", ""))
        self._fill_cell(
            table,
            self._get_row_index("homework"),
            col_idx,
            homework_text,
            day_name=day_name,
            section_name="homework",
            pending_hyperlinks=pending_hyperlinks,
            pending_images=pending_images,
            current_slot_number=slot_number,
            current_subject=subject,
        )

    def _fill_multi_slot_day(
        self,
        table,
        col_idx: int,
        slots: List[Dict],
        metadata: Dict = None,
        day_name: str = None,
        pending_hyperlinks: List[Dict] = None,
        pending_images: List[Dict] = None,
    ):
        """
        Fill multiple slots' data for one day, with per-slot hyperlink placement.

        Strategy:
        - For each row (Unit/Lesson, Objective, etc.):
          - Clear cell once (first _fill_cell call will do this)
          - For each slot:
            - Build slot text with header
            - Call _fill_cell with append_mode=True (after first)
            - Pass shared pending_hyperlinks/pending_images lists
            - Let _fill_cell filter by current_slot_number
            - Matched items removed from shared lists automatically

        Args:
            table: Table object
            col_idx: Column index for the day
            slots: List of slot data dicts
            day_name: Day name for media matching
            pending_hyperlinks: List of hyperlinks awaiting placement
            pending_images: List of images awaiting placement
        """
        if not slots:
            return

        # Sort slots by start_time (chronological) with slot_number as fallback
        sorted_slots = sort_slots(slots)
        num_slots = len(sorted_slots)

        # Check if any slot has content (for placeholder logic)
        slots_have_content = []
        for slot in sorted_slots:
            has_content = any(
                [
                    slot.get("unit_lesson"),
                    slot.get("objective"),
                    slot.get("anticipatory_set"),
                    slot.get("tailored_instruction"),
                    slot.get("misconceptions"),
                    slot.get("assessment"),
                    slot.get("homework"),
                ]
            )
            slots_have_content.append(has_content)

        # Define rows to fill with correct field names and format functions
        # Use dynamic row indices from structure detector
        rows_config = [
            # (field_name, row_idx, format_func, placeholder_text, max_length)
            (
                "unit_lesson",
                self._get_row_index("unit"),
                None,
                "[No unit/lesson specified]",
                100,
            ),
            (
                "objective",
                self._get_row_index("objective"),
                self._format_objective,
                "[No objective specified]",
                None,
            ),
            (
                "anticipatory_set",
                self._get_row_index("anticipatory"),
                self._format_anticipatory_set,
                None,
                None,
            ),
            (
                "tailored_instruction",
                self._get_row_index("instruction"),
                self._format_tailored_instruction,
                None,
                None,
            ),
            (
                "misconceptions",
                self._get_row_index("misconception"),
                self._format_misconceptions,
                None,
                None,
            ),
            (
                "assessment",
                self._get_row_index("assessment"),
                self._format_assessment,
                None,
                None,
            ),
            (
                "homework",
                self._get_row_index("homework"),
                self._format_homework,
                None,
                100,
            ),
        ]

        # Fill each row
        for (
            field_name,
            row_idx,
            format_func,
            placeholder_text,
            max_length,
        ) in rows_config:
            # NOTE: We don't manually clear the cell here. The first _fill_cell call
            # (with append_mode=False) will clear it automatically.

            # Track whether we've written any content to this row yet
            written_any = False

            # Fill each slot
            if metadata is None:
                metadata = {}
            for i, slot in enumerate(sorted_slots):
                slot_num = slot.get("slot_number", "?")
                subject = slot.get("subject", "Unknown")
                # Use get_teacher_name() which prioritizes slot-specific teacher names
                # This ensures each slot shows its own primary teacher, not the lesson-level teacher
                teacher = get_teacher_name(metadata, slot=slot)
                # Debug: Log teacher name resolution for troubleshooting
                slot_primary = slot.get("primary_teacher_name")
                slot_first = slot.get("primary_teacher_first_name")
                slot_last = slot.get("primary_teacher_last_name")
                metadata_teacher = metadata.get("teacher_name")
                print(f"[DEBUG] DOCX_RENDERER: Slot {slot_num} ({subject})")
                print(f"  slot.primary_teacher_name: {slot_primary}")
                print(f"  slot.primary_teacher_first_name: {slot_first}")
                print(f"  slot.primary_teacher_last_name: {slot_last}")
                print(f"  metadata.teacher_name: {metadata_teacher}")
                print(f"  get_teacher_name() result: {teacher}")
                print(f"  Final slot_header will use: {teacher}")

                has_content = slots_have_content[i]

                # Check for "No School" status
                # If this slot has "No School" in unit_lesson, we only show it in the unit_lesson row
                # and skip all other rows for this slot.
                is_no_school = "no school" in (slot.get("unit_lesson") or "").lower()
                if is_no_school and field_name != "unit_lesson":
                    continue

                # Build slot header
                slot_header = f"**Slot {slot_num}: {subject}**"
                if teacher:
                    slot_header += f" ({teacher})"

                # Get slot content
                slot_content = slot.get(field_name)

                # Build slot text
                if slot_content:
                    # Format if format function provided
                    if format_func:
                        slot_text = format_func(slot_content)
                    else:
                        # Raw string (e.g., unit_lesson)
                        slot_text = slot_content

                    # Abbreviate for multi-slot
                    if slot_text:
                        slot_text = self._abbreviate_content(
                            slot_text, num_slots, max_length=max_length
                        )
                        slot_text = f"{slot_header}\n{slot_text}"
                    else:
                        slot_text = None
                elif has_content and placeholder_text:
                    # Add placeholder if other fields exist but not this one
                    slot_text = f"{slot_header}\n{placeholder_text}"
                else:
                    # No content and no placeholder needed
                    slot_text = None

                # Skip if no content
                if not slot_text:
                    continue

                # Add separator only if there's another slot with content after this one
                # Look ahead to see if any remaining slots will produce content
                has_next_slot_with_content = False
                for j in range(i + 1, len(sorted_slots)):
                    next_slot = sorted_slots[j]
                    next_slot_content = next_slot.get(field_name)
                    next_has_content = slots_have_content[j]

                    # Check if next slot will produce text (content or placeholder)
                    if next_slot_content:
                        # Has actual content
                        if format_func:
                            next_text = format_func(next_slot_content)
                        else:
                            next_text = next_slot_content
                        if next_text:
                            has_next_slot_with_content = True
                            break
                    elif next_has_content and placeholder_text:
                        # Will show placeholder
                        has_next_slot_with_content = True
                        break

                # Only add separator if there's content coming after
                if has_next_slot_with_content:
                    slot_text += "\n\n---"

                # Fill cell with this slot's content
                # IMPORTANT: Pass shared pending_hyperlinks and pending_images lists
                # _fill_cell will filter by current_slot_number and remove matched items
                self._fill_cell(
                    table,
                    row_idx,
                    col_idx,
                    slot_text,
                    day_name=day_name,
                    section_name=field_name,
                    pending_hyperlinks=pending_hyperlinks,  # Shared list
                    pending_images=pending_images,  # Shared list
                    current_slot_number=slot_num,
                    current_subject=subject,
                    append_mode=written_any,  # Append only if we've already written to this row
                )

                # Mark that we've written content to this row
                written_any = True

    def _fill_cell(
        self,
        table,
        row_idx: int,
        col_idx: int,
        text: str,
        day_name: str = None,
        section_name: str = None,
        pending_hyperlinks: List[Dict] = None,
        pending_images: List[Dict] = None,
        current_slot_number: int = None,
        current_subject: str = None,
        append_mode: bool = False,
    ):
        """
        Fill a cell with formatted text and inject matched media inline.

        Args:
            table: Table object
            row_idx: Row index
            col_idx: Column index
            text: Text to fill (may contain markdown)
            day_name: Day name for media matching
            section_name: Section name for media matching
            pending_hyperlinks: List of hyperlinks awaiting placement
            pending_images: List of images awaiting placement
            current_slot_number: Slot number for filtering hyperlinks/images
            current_subject: Subject name for filtering hyperlinks/images
            append_mode: If True, append to cell without clearing existing content
        """
        from backend.config import settings
        import re
        # Explicit reference to ensure Python treats it as module-level/global, not local
        _ = (re, qn)  # Prevent UnboundLocalError by ensuring these are looked up in globals

        cell = table.rows[row_idx].cells[col_idx]
        print(f"DEBUG: _fill_cell(row={row_idx}, col={col_idx}, day={day_name}, section={section_name})")
        print(f"DEBUG: text[:50] = {repr(text[:50]) if text else 'None'}")
        
        # Check if cell already has hyperlinks (from coordinate placement)
        existing_hyperlinks = cell._element.xpath(".//w:hyperlink")
        has_coordinate_hyperlinks = len(existing_hyperlinks) > 0
        
        # If coordinate-placed hyperlinks exist, check if text already contains them in markdown
        # If so, remove the coordinate-placed hyperlinks to avoid duplicates
        if has_coordinate_hyperlinks and text:
            hyperlinks_to_remove = []
            for hl_elem in existing_hyperlinks:
                try:
                    r_id = hl_elem.get(qn("r:id"))
                    if r_id and cell.paragraphs:
                        # Get URL from relationship
                        para = cell.paragraphs[0]  # Use first paragraph to access part
                        if r_id in para.part.rels:
                            url = para.part.rels[r_id].target_ref
                            # Get link text
                            link_text = "".join(node.text for node in hl_elem.xpath(".//w:t") if node.text)
                            if link_text and url:
                                # Check if text contains this link in markdown format
                                markdown_pattern = rf"\[{re.escape(link_text)}\]\({re.escape(url)}\)"
                                if re.search(markdown_pattern, text, re.IGNORECASE):
                                    # Text already has this link, mark for removal
                                    hyperlinks_to_remove.append(hl_elem)
                except Exception:
                    pass  # Skip if we can't process this hyperlink
            
            # Remove duplicate hyperlinks
            for hl_elem in hyperlinks_to_remove:
                try:
                    # Remove the hyperlink element from its parent
                    parent = hl_elem.getparent()
                    if parent is not None:
                        parent.remove(hl_elem)
                        # Also remove the paragraph if it only contained this hyperlink
                        try:
                            para_elem = parent.getparent()
                            if para_elem is not None and para_elem.tag == qn("w:p"):
                                # Check if paragraph is now empty (only formatting, no text)
                                runs = para_elem.xpath(".//w:r")
                                has_text = any(run.xpath(".//w:t") for run in runs)
                                if not has_text:
                                    para_parent = para_elem.getparent()
                                    if para_parent is not None:
                                        para_parent.remove(para_elem)
                        except Exception:
                            # If we can't remove the paragraph, that's okay - just continue
                            pass
                except Exception as e:
                    # Log but don't fail - continue processing
                    logger.warning(
                        "failed_to_remove_duplicate_hyperlink",
                        extra={"error": str(e), "cell": f"{day_name}_{section_name}" if day_name and section_name else "unknown"},
                    )
                    pass
            
            # Re-check after removal
            existing_hyperlinks = cell._element.xpath(".//w:hyperlink")
            has_coordinate_hyperlinks = len(existing_hyperlinks) > 0
            
            if hyperlinks_to_remove:
                logger.info(
                    "removed_duplicate_coordinate_hyperlinks",
                    extra={
                        "removed_count": len(hyperlinks_to_remove),
                        "cell": f"{day_name}_{section_name}" if day_name and section_name else "unknown",
                        "slot": current_slot_number,
                    },
                )

        # Clear existing content ONLY if:
        # 1. No coordinate-placed hyperlinks exist, AND
        # 2. Not in append mode
        if not has_coordinate_hyperlinks and not append_mode:
            # CRITICAL: Manually remove all paragraphs instead of cell.text = ""
            # cell.text = "" leaves an empty paragraph with default formatting (Arial 11)
            # which causes blank lines. Manual removal is cleaner.
            paras_to_remove = list(cell.paragraphs)
            for para in paras_to_remove:
                p = para._element
                p.getparent().remove(p)
            # Word requires at least one paragraph - add a fresh one with proper formatting
            new_para = cell.add_paragraph()
            new_para.paragraph_format.space_after = Pt(0)
            new_para.paragraph_format.space_before = Pt(0)
            new_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Collect matching hyperlinks FIRST (before adding text)
        matching_hyperlinks = []
        if pending_hyperlinks and text:
            # DEBUG: Log filtering context
            logger.info(
                "hyperlink_filtering_context",
                extra={
                    "current_slot_number": current_slot_number,
                    "current_subject": current_subject,
                    "pending_hyperlinks_count": len(pending_hyperlinks),
                    "cell": f"{day_name}_{section_name}"
                    if day_name and section_name
                    else "unknown",
                },
            )

            # DIAGNOSTIC: Log filtering context
            from tools.diagnostic_logger import get_diagnostic_logger

            diag = get_diagnostic_logger()
            diag.log_filtering_context(
                current_slot_number,
                current_subject,
                len(pending_hyperlinks),
                f"{day_name}_{section_name}"
                if day_name and section_name
                else "unknown",
            )

            # Collect all matching hyperlinks first
            for hyperlink in pending_hyperlinks[:]:
                # CRITICAL: Filter hyperlinks by slot to prevent cross-contamination
                # Only consider hyperlinks from the current slot being rendered
                if current_slot_number is not None:
                    link_slot = hyperlink.get("_source_slot")

                    # STRICT: In multi-slot mode, REQUIRE slot metadata
                    if link_slot is None:
                        # Missing metadata - skip to prevent cross-contamination
                        logger.warning(
                            "hyperlink_missing_slot_metadata",
                            extra={
                                "text": hyperlink["text"][:50],
                                "url": hyperlink["url"],
                                "current_slot": current_slot_number,
                            },
                        )
                        continue

                    if link_slot != current_slot_number:
                        # Skip hyperlinks from other slots
                        logger.debug(
                            "hyperlink_filtered_by_slot",
                            extra={
                                "text": hyperlink["text"][:50],
                                "link_slot": link_slot,
                                "current_slot": current_slot_number,
                            },
                        )

                        # DIAGNOSTIC: Log filtered hyperlink
                        diag.log_hyperlink_filtered(
                            hyperlink["text"],
                            link_slot,
                            current_slot_number,
                            "slot_mismatch",
                        )
                        continue

                # Additional filter by subject for extra safety
                if current_subject is not None:
                    link_subject = hyperlink.get("_source_subject")

                    # STRICT: In multi-slot mode, REQUIRE subject metadata
                    if link_subject is None:
                        # Missing metadata - skip to prevent cross-contamination
                        logger.warning(
                            "hyperlink_missing_subject_metadata",
                            extra={
                                "text": hyperlink["text"][:50],
                                "url": hyperlink["url"],
                                "current_subject": current_subject,
                            },
                        )
                        continue

                    if link_subject != current_subject:
                        # Skip hyperlinks from other subjects
                        logger.debug(
                            "hyperlink_filtered_by_subject",
                            extra={
                                "text": hyperlink["text"][:50],
                                "link_subject": link_subject,
                                "current_subject": current_subject,
                            },
                        )
                        continue

                # UNIFIED PLACEMENT: Strategy 1 & 2: Structural matching (Coordinate or Label+Day)
                # This ensures that Schema 2.0 links also benefit from Smart Inline Replacement.
                is_structural_match = False
                match_type = "none"
                confidence = 0.0

                # Case 1: Coordinates (v2.0)
                if hyperlink.get("row_idx") is not None and hyperlink.get("cell_idx") is not None:
                    # Target row index (accounting for offset from structure metadata)
                    row_offset = self.structure_metadata.row_offset if self.structure_metadata else 0
                    target_row = hyperlink["row_idx"] + row_offset
                    if target_row == row_idx and hyperlink["cell_idx"] == col_idx:
                        is_structural_match = True
                        match_type = "coordinate"
                        confidence = 1.0
                
                # Case 2: Label + Day matching (v2.0)
                if not is_structural_match and hyperlink.get("row_label") and hyperlink.get("day_hint"):
                    # Use existing structure metadata for robust label matching if available
                    if self.structure_metadata:
                        target_row_idx = self.structure_metadata.get_row_index(hyperlink["row_label"])
                        target_col_idx = self.structure_metadata.get_col_index(hyperlink["day_hint"])
                        # Apply row offset
                        if target_row_idx is not None:
                            target_row_idx += self.structure_metadata.row_offset
                        
                        if target_row_idx == row_idx and target_col_idx == col_idx:
                            is_structural_match = True
                            match_type = "label_day"
                            confidence = 1.0

                if not is_structural_match:
                    # Case 3: Fuzzy matching (v1.1/v2.0)
                    confidence, match_type = self._calculate_match_confidence(
                        text, hyperlink, day_name, section_name
                    )

                if confidence >= settings.MEDIA_MATCH_CONFIDENCE_THRESHOLD:
                    matching_hyperlinks.append((hyperlink, confidence, match_type))

        # NEW: Smart Inline Replacement
        # If we found matching links, try to wrap their text in markdown format 
        # so MarkdownToDocx will place them inline instead of at the top.
        if matching_hyperlinks and text:
            # Sort hyperlinks by text length descending to avoid partial matches
            # (e.g., matching "Lenni Lenape" before "Lenni")
            sorted_matching = sorted(matching_hyperlinks, key=lambda m: len(m[0].get("text", "")), reverse=True)
            
            for hyperlink, confidence, match_type in sorted_matching:
                link_text = hyperlink.get("text", "")
                link_url = hyperlink.get("url", "")
                if not link_text or not link_url:
                    continue

                # NEW: Strict Day Enforcement during Smart Inline Replacement
                # Prevent cross-day contamination if text matches but link is for another day
                hl_day = hyperlink.get("day_hint") or hyperlink.get("day")
                if hl_day and day_name:
                    if hl_day.lower().strip() != day_name.lower().strip():
                        # Log the skip for debugging
                        logger.debug(
                            "hyperlink_skipped_day_mismatch_inline",
                            extra={
                                "text": link_text,
                                "link_day": hl_day,
                                "current_day": day_name,
                                "cell": f"{day_name}_{section_name}"
                            },
                        )
                        continue

                # Case-insensitive search with robust matching (flexible whitespace/punctuation)
                # We normalize search pattern to handle different spacing or minor rephrasing
                search_pattern_raw = re.escape(link_text)
                # Allow any amount of whitespace where input has space (escaped or literal)
                search_pattern_flexible = re.sub(r'(\\ )+', r'\\s+', search_pattern_raw)
                # Allow optional trailing punctuation
                search_pattern_final = f"{search_pattern_flexible}[.,;:]?"
                
                match = re.search(search_pattern_final, text, re.IGNORECASE)
                
                # DIAGNOSTIC: Log match results
                with open(r'd:\LP\hyperlink_debug.log', 'a', encoding='utf-8') as f:
                    try:
                        f.write(json.dumps({
                            "sessionId": "debug-session",
                            "runId": "investigation",
                            "location": "docx_renderer.py:_fill_cell:replacement_attempt",
                            "data": {
                                "cell": f"{day_name}_{section_name}",
                                "link_text": link_text,
                                "regex": search_pattern_final,
                                "match_found": match is not None,
                                "text_preview": (text[:200] if text else "None"),
                                "found_text": (match.group(0) if match else None)
                            },
                            "timestamp": int(time.time() * 1000)
                        }) + '\n')
                    except Exception as e:
                        # Don't let diagnostic logging crash the renderer
                        self.logger.warning(f"Diagnostic logging failed: {e}")
                
                if match:
                    # Match may include leading/trailing whitespace if regex was flexible
                    # We want to keep the whitespace OUTSIDE the link brackets for better readability
                    raw_match_text = match.group(0)
                    stripped_match_text = raw_match_text.strip()
                    
                    # NEW: Robust substitution that only replaces text NOT already inside a link
                    # total_pattern captures [existing link](url) in group 1 OR the plain text match in group 2
                    total_pattern = rf"(\[\s*{re.escape(stripped_match_text)}\s*\]\([^)]*\))|({search_pattern_final})"
                    
                    found_and_replaced = [False]
                    
                    def replacement_logic(m):
                        # If we matched an existing link (group 1), return it as-is
                        if m.group(1):
                            return m.group(1)
                        
                        # If we matched plain text (group 2) AND we haven't replaced yet, link it!
                        if not found_and_replaced[0]:
                            found_and_replaced[0] = True
                            plain_text = m.group(2)
                            # Preserve surrounding whitespace
                            leading_ws = plain_text[:plain_text.find(stripped_match_text)]
                            trailing_ws = plain_text[plain_text.find(stripped_match_text) + len(stripped_match_text):]
                            return f"{leading_ws}[{stripped_match_text}]({link_url}){trailing_ws}"
                        
                        # Otherwise return original text
                        return m.group(0)

                    # Apply substitution (case-insensitive)
                    new_text = re.sub(total_pattern, replacement_logic, text, flags=re.IGNORECASE)
                    
                    if found_and_replaced[0]:
                        text = new_text
                        logger.info(
                            "hyperlink_placed_smart_inline",
                            extra={
                                "text": stripped_match_text,
                                "url": link_url,
                                "cell": f"{day_name}_{section_name}",
                                "slot": current_slot_number,
                            },
                        )
                    else:
                        # If not replaced, it might be already inline or not found as plain text
                        logger.info(
                            "hyperlink_already_inline_or_unmatched",
                            extra={
                                "text": stripped_match_text,
                                "cell": f"{day_name}_{section_name}",
                                "slot": current_slot_number,
                            },
                        )
                    
                    # ALWAYS remove from matching/pending if found in text (either as plain text or markdown)
                    matching_hyperlinks = [m for m in matching_hyperlinks if m[0] != hyperlink]
                    if hyperlink in pending_hyperlinks:
                        pending_hyperlinks.remove(hyperlink)

        # Calculate total hyperlink count (coordinate + fuzzy)
        coordinate_link_count = len(existing_hyperlinks) if existing_hyperlinks else 0
        fuzzy_link_count = len(matching_hyperlinks) if matching_hyperlinks else 0
        total_link_count = coordinate_link_count + fuzzy_link_count


        # Consistently handled by fallback logic below to avoid duplicates and clutter

        # Now add the text content (which may now contain our new [text](url) markdown)
        if not has_coordinate_hyperlinks and not append_mode:
            # Add formatted text
            if text:
                MarkdownToDocx.add_multiline_text(cell, text)
        else:
            # Cell has coordinate-placed hyperlinks OR append mode - append text without clearing
            # Add text with markdown formatting to preserve existing hyperlinks
            if text:
                # Add formatted text line by line, preserving Markdown
                lines = text.split("\n")
                for i, line in enumerate(lines):
                    if line.strip():
                        if i == 0:
                            # First line - use MarkdownToDocx for proper formatting
                            new_para = cell.add_paragraph()
                            MarkdownToDocx.add_formatted_text(new_para, line)
                        else:
                            # Subsequent lines
                            MarkdownToDocx.add_paragraph(cell, line)

        # CRITICAL: Final formatting pass for ALL paragraphs and runs in cell
        if cell.paragraphs:
            for para in cell.paragraphs:
                # Force spacing
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # Format direct runs
                for run in para.runs:
                    self._force_font_tnr8(run, is_bold=(row_idx == self.UNIT_LESSON_ROW))
                
                # Format runs inside hyperlinks
                for hl in para._p.findall(qn("w:hyperlink")):
                    for r_elem in hl.findall(qn("w:r")):
                        hl_run = Run(r_elem, para)
                        is_hl_bold = (row_idx == self.UNIT_LESSON_ROW) or hl_run.font.bold
                        self._force_font_tnr8(hl_run, is_bold=is_hl_bold, is_hyperlink=True)

        # FALLBACK: Append media that belongs to this section but wasn't placed inline
        # This ensures "orphan" media stays with its relevant section
        if section_name and (pending_hyperlinks or pending_images):
            # Check for hyperlinks with matching section hint
            for hyperlink in pending_hyperlinks[:]:
                # Check if this hyperlink belongs to this slot/subject
                if current_slot_number is not None:
                    link_slot = hyperlink.get("_source_slot")
                    if link_slot is not None and link_slot != current_slot_number:
                        continue

                # STRICT DAY CHECK: Do not place fallback links in the wrong day
                # This prevents "orphan" links from Monday appearing as bullets on Tuesday
                hl_day = hyperlink.get("day_hint")
                if hl_day and day_name:
                    if hl_day.lower().strip() != day_name.lower().strip():
                        continue

                # Check section hint using flexible mapping
                hint = hyperlink.get("section_hint", "").lower()
                is_section_match = False
                
                # Flexible section mapping
                section_mappings = {
                    "unit_lesson": ["unit", "lesson", "module"],
                    "objective": ["objective", "goal", "swbat"],
                    "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                    "tailored_instruction": ["instruction", "activity", "procedure", "lesson", "tailored", "differentiation"],
                    "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                    "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                    "homework": ["homework", "assignment", "practice"]
                }

                if hint == section_name:
                    is_section_match = True
                elif section_name in section_mappings:
                    if hint in section_mappings[section_name]:
                        is_section_match = True
                    # Also check partial matches if hint is long enough
                    elif any(kw in hint for kw in section_mappings[section_name] if len(kw) > 3):
                        is_section_match = True

                if is_section_match:
                    # Append as a new paragraph
                    # Append as a new paragraph with bullet
                    p = cell.add_paragraph()
                    MarkdownToDocx.add_formatted_text(
                        p, f"• [{hyperlink['text']}]({hyperlink['url']})"
                    )
                    # Format the new paragraph
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(8)

                    pending_hyperlinks.remove(hyperlink)
                    logger.info(
                        "hyperlink_placed_fallback",
                        extra={"text": hyperlink["text"], "section": section_name},
                    )

            # Check for images with matching section hint
            for image in pending_images[:]:
                # Check slot/subject (similar to hyperlinks)
                if current_slot_number is not None:
                    img_slot = image.get("_source_slot")
                    if img_slot is not None and img_slot != current_slot_number:
                        continue

                hint = image.get("section_hint", "").lower()
                is_section_match = False
                
                # Flexible section mapping (shared with hyperlinks)
                section_mappings = {
                    "unit_lesson": ["unit", "lesson", "module"],
                    "objective": ["objective", "goal", "swbat"],
                    "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                    "tailored_instruction": ["instruction", "activity", "procedure", "lesson", "tailored", "differentiation"],
                    "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                    "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                    "homework": ["homework", "assignment", "practice"]
                }

                if hint == section_name:
                    is_section_match = True
                elif section_name in section_mappings:
                    if hint in section_mappings[section_name]:
                        is_section_match = True
                    # Also check partial matches if hint is long enough
                    elif any(kw in hint for kw in section_mappings[section_name] if len(kw) > 3):
                        is_section_match = True

                if is_section_match:
                    self._inject_image_inline(
                        cell, image, max_width=1.3
                    )  # approx column width
                    pending_images.remove(image)
                    logger.info(
                        "image_placed_fallback",
                        extra={"filename": image["filename"], "section": section_name},
                    )

        # CRITICAL: Remove placeholder paragraph if no content was actually added
        # This prevents empty cells from showing a placeholder when content is missing
        # Check if cell has any real content (text, hyperlinks, or images)
        has_content = False
        if cell.paragraphs:
            for para in cell.paragraphs:
                # Check if paragraph has content (not just whitespace)
                if para.text.strip() or para.runs:
                    # Check if runs contain actual content (not just formatting)
                    for run in para.runs:
                        if run.text and run.text.strip():
                            has_content = True
                            break
                    if has_content:
                        break
                    # Check for hyperlinks (they might be in XML, not runs)
                    if para._element.xpath(".//w:hyperlink"):
                        has_content = True
                        break

        # If no content was added and we have placeholder paragraphs, remove them
        if not has_content and not has_coordinate_hyperlinks and not append_mode:
            # Check if cell only has empty placeholder paragraphs
            empty_paras = []
            for para in cell.paragraphs:
                if not para.text.strip() and not para.runs:
                    empty_paras.append(para)

            # Remove all empty paragraphs (they're just placeholders)
            for para in empty_paras:
                p = para._element
                p.getparent().remove(p)

        # Match and inject images (existing logic continues)
        if pending_images:
            # Estimate column width (5 columns = ~1.3 inches each)
            estimated_col_width = 6.5 / 5

            if estimated_col_width >= settings.IMAGE_INLINE_MIN_COLUMN_WIDTH_INCHES:
                for image in pending_images[:]:
                    # CRITICAL: Filter images by slot to prevent cross-contamination
                    # Only consider images from the current slot being rendered
                    if current_slot_number is not None:
                        image_slot = image.get("_source_slot")

                        # STRICT: In multi-slot mode, REQUIRE slot metadata
                        if image_slot is None:
                            # Missing metadata - skip to prevent cross-contamination
                            logger.warning(
                                "image_missing_slot_metadata",
                                extra={
                                    "filename": image.get("filename", "unknown"),
                                    "current_slot": current_slot_number,
                                },
                            )
                            continue

                        if image_slot != current_slot_number:
                            # Skip images from other slots
                            logger.debug(
                                "image_filtered_by_slot",
                                extra={
                                    "filename": image["filename"],
                                    "image_slot": image_slot,
                                    "current_slot": current_slot_number,
                                },
                            )
                            continue

                    # Additional filter by subject for extra safety
                    if current_subject is not None:
                        image_subject = image.get("_source_subject")

                        # STRICT: In multi-slot mode, REQUIRE subject metadata
                        if image_subject is None:
                            # Missing metadata - skip to prevent cross-contamination
                            logger.warning(
                                "image_missing_subject_metadata",
                                extra={
                                    "filename": image.get("filename", "unknown"),
                                    "current_subject": current_subject,
                                },
                            )
                            continue

                        if image_subject != current_subject:
                            # Skip images from other subjects
                            logger.debug(
                                "image_filtered_by_subject",
                                extra={
                                    "filename": image["filename"],
                                    "image_subject": image_subject,
                                    "current_subject": current_subject,
                                },
                            )
                            continue

                    # Try structure-based placement first (exact location)
                    if self._try_structure_based_placement(
                        image, day_name, section_name, col_idx
                    ):
                        self._inject_image_inline(
                            cell, image, max_width=estimated_col_width
                        )
                        pending_images.remove(image)

                        logger.info(
                            "image_placed_inline",
                            extra={
                                "filename": image["filename"],
                                "cell": f"{day_name}_{section_name}",
                                "confidence": 1.0,
                                "match_type": "structure_based",
                                "slot": current_slot_number,
                                "subject": current_subject,
                            },
                        )
                    # Fall back to context-based matching
                    elif text:
                        confidence, match_type = self._calculate_match_confidence(
                            text, image, day_name, section_name
                        )

                        if confidence >= settings.MEDIA_MATCH_CONFIDENCE_THRESHOLD:
                            self._inject_image_inline(
                                cell, image, max_width=estimated_col_width
                            )
                            pending_images.remove(image)

                            logger.info(
                                "image_placed_inline",
                                extra={
                                    "filename": image["filename"],
                                    "cell": f"{day_name}_{section_name}",
                                    "confidence": confidence,
                                    "match_type": match_type,
                                    "slot": current_slot_number,
                                    "subject": current_subject,
                                },
                            )
                        else:
                            logger.debug(
                                "image_match_rejected",
                                extra={
                                    "filename": image["filename"],
                                    "cell": f"{day_name}_{section_name}",
                                    "confidence": confidence,
                                },
                            )

    def _format_objective(self, objective: Dict) -> str:
        """Format objective section."""
        if not objective:
            return ""

        # Originals mode: just return the raw text
        if self.is_originals:
            return objective.get("content_objective") or objective.get("student_goal") or ""

        parts = []

        if "content_objective" in objective:
            parts.append(f"**Content:** {objective['content_objective']}")

        if "student_goal" in objective:
            parts.append(f"**Student Goal:** {objective['student_goal']}")

        if "wida_objective" in objective:
            parts.append(f"**WIDA/Bilingual:** {objective['wida_objective']}")

        return "\n\n".join(parts)

    def _format_anticipatory_set(self, anticipatory: Dict) -> str:
        """Format anticipatory set section."""
        if not anticipatory:
            return ""

        # Originals mode: just return the raw text
        if self.is_originals:
            return anticipatory.get("original_content") or ""

        parts = []

        if "original_content" in anticipatory:
            parts.append(anticipatory["original_content"])

        if "bilingual_bridge" in anticipatory:
            parts.append(f"\n**Bilingual Bridge:** {anticipatory['bilingual_bridge']}")

        return "\n".join(parts)

    def _filter_valid_vocabulary_pairs(
        self, vocabulary_cognates: List[Dict]
    ) -> List[Dict]:
        """Filter vocabulary pairs to only include valid ones with both english and portuguese."""
        return [
            pair
            for pair in vocabulary_cognates
            if isinstance(pair, dict)
            and str(pair.get("english", "")).strip()
            and str(pair.get("portuguese", "")).strip()
        ]

    def _filter_valid_sentence_frames(self, sentence_frames: List[Dict]) -> List[Dict]:
        """Filter sentence frames to only include well-formed ones."""
        return [
            frame
            for frame in sentence_frames
            if isinstance(frame, dict)
            and str(frame.get("english", "")).strip()
            and str(frame.get("portuguese", "")).strip()
            and str(frame.get("proficiency_level", "")).strip()
        ]

    def _format_tailored_instruction(
        self,
        instruction: Dict,
        vocabulary_cognates: Optional[List[Dict]] = None,
        sentence_frames: Optional[List[Dict]] = None,
    ) -> str:
        """Format tailored instruction section."""
        if not instruction:
            return ""

        # Originals mode: just return the raw text
        if self.is_originals:
            return instruction.get("original_content") or ""

        parts = []

        # Original content
        if "original_content" in instruction:
            parts.append(instruction["original_content"])

        # Co-teaching model
        if "co_teaching_model" in instruction:
            co_teaching = instruction["co_teaching_model"]
            parts.append(
                f"\n**Co-Teaching Model:** {co_teaching.get('model_name', '')}"
            )

            if "phase_plan" in co_teaching:
                parts.append("\n**Phase Plan:**")
                for phase in co_teaching["phase_plan"]:
                    phase_name = phase.get("phase_name", "")
                    minutes = phase.get("minutes", "")
                    parts.append(f"- **{phase_name}** ({minutes} min)")

                    if "bilingual_teacher_role" in phase:
                        parts.append(
                            f"  - Bilingual: {phase['bilingual_teacher_role']}"
                        )
                    if "primary_teacher_role" in phase:
                        parts.append(f"  - Primary: {phase['primary_teacher_role']}")

        # ELL Support
        ell_support = instruction.get("ell_support") or []
        if ell_support:
            parts.append("\n**ELL Support:**")
            has_cognate_awareness = False

            for strategy in ell_support:
                strategy_name = strategy.get("strategy_name", "")
                implementation = strategy.get("implementation", "")
                levels = strategy.get("proficiency_levels", "")
                strategy_id = strategy.get("strategy_id", "")

                parts.append(f"- **{strategy_name}** ({levels}): {implementation}")

                if strategy_id == "cognate_awareness":
                    has_cognate_awareness = True

            # If we have structured vocabulary_cognates data, append a formatted vocabulary block.
            # Note: We render vocabulary_cognates if the data exists, similar to sentence_frames,
            # regardless of whether cognate_awareness strategy is explicitly in ell_support.
            # This ensures consistency with how sentence_frames are rendered.
            if vocabulary_cognates:
                valid_pairs = self._filter_valid_vocabulary_pairs(vocabulary_cognates)
                if valid_pairs:
                    parts.append("\n**Vocabulary / Cognate Awareness:**")
                    for pair in valid_pairs:
                        english = str(pair.get("english", "")).strip()
                        portuguese = str(pair.get("portuguese", "")).strip()
                        parts.append(f"- **{english}** → *{portuguese}*")

            # Sentence Frames grouped by proficiency level (if provided)
            if sentence_frames:
                valid_frames = self._filter_valid_sentence_frames(sentence_frames)

                if valid_frames:
                    parts.append("\n**Sentence Frames / Stems / Questions:**")

                    level_order = [
                        ("levels_1_2", "Levels 1-2"),
                        ("levels_3_4", "Levels 3-4"),
                        ("levels_5_6", "Levels 5-6"),
                    ]

                    for level_key, level_label in level_order:
                        frames_for_level = [
                            f
                            for f in valid_frames
                            if f.get("proficiency_level") == level_key
                        ]

                        if not frames_for_level:
                            continue

                        parts.append(f"\n- **{level_label}:**")

                        for frame in frames_for_level:
                            english = str(frame.get("english", "")).strip()
                            portuguese = str(frame.get("portuguese", "")).strip()
                            language_function = str(
                                frame.get("language_function", "")
                            ).strip()

                            parts.append(f"  - PT: *{portuguese}*")

                            if language_function:
                                parts.append(
                                    f"    EN: **{english}** (function: {language_function})"
                                )
                            else:
                                parts.append(f"    EN: **{english}**")

        # Materials
        if "materials" in instruction and instruction["materials"]:
            parts.append("\n**Materials:** " + ", ".join(instruction["materials"]))

        return "\n".join(parts)

    def _format_misconceptions(self, misconceptions: Dict) -> str:
        """Format misconceptions section."""
        if not misconceptions:
            return ""

        # Originals mode: just return the raw text
        if self.is_originals:
            return misconceptions.get("original_content") or ""

        parts = []

        if "original_content" in misconceptions:
            parts.append(misconceptions["original_content"])

        if "linguistic_note" in misconceptions:
            ling = misconceptions["linguistic_note"]
            if "note" in ling:
                parts.append(f"\n**Linguistic Note:** {ling['note']}")
            if "prevention_tip" in ling:
                parts.append(f"**Prevention:** {ling['prevention_tip']}")

        return "\n".join(parts)

    def _format_assessment(self, assessment: Dict) -> str:
        """Format assessment section."""
        if not assessment:
            return ""

        # Originals mode: just return the raw text
        if self.is_originals:
            return assessment.get("primary_assessment") or ""

        parts = []

        if "primary_assessment" in assessment:
            parts.append(f"**Assessment:** {assessment['primary_assessment']}")

        if "bilingual_overlay" in assessment:
            overlay = assessment["bilingual_overlay"]

            if "instrument" in overlay:
                parts.append(f"\n**Instrument:** {overlay['instrument']}")

            if "supports_by_level" in overlay:
                parts.append("\n**Supports by Level:**")
                supports = overlay["supports_by_level"]
                if "levels_1_2" in supports:
                    parts.append(f"- **Levels 1-2:** {supports['levels_1_2']}")
                if "levels_3_4" in supports:
                    parts.append(f"- **Levels 3-4:** {supports['levels_3_4']}")
                if "levels_5_6" in supports:
                    parts.append(f"- **Levels 5-6:** {supports['levels_5_6']}")

        return "\n".join(parts)

    def _try_structure_based_placement(
        self, image: Dict, day_name: str, section_name: str, col_idx: int
    ) -> bool:
        """
        Try to place image using structure-based matching (row label + cell index).

        Args:
            image: Image dictionary with row_label and cell_index
            day_name: Current day being rendered
            section_name: Current section being rendered
            col_idx: Current column index

        Returns:
            True if structure matches (should place here), False otherwise
        """
        # Check if image has structure information
        if not image.get("row_label") or image.get("cell_index") is None:
            return False

        # Map section names to match variations
        section_matches = {
            "unit_lesson": ["unit", "lesson"],
            "objective": ["objective", "goal"],
            "anticipatory_set": ["anticipatory", "warm", "hook"],
            "instruction": ["instruction", "activity", "lesson", "tailored"],
            "misconceptions": ["misconception", "error"],
            "assessment": ["assessment", "check", "evaluate"],
            "homework": ["homework", "assignment"],
        }

        # Check if section matches
        section_match = False
        image_section = image.get("section_hint", "")

        if section_name and image_section:
            # Direct match
            if section_name == image_section:
                section_match = True
            # Check variations
            elif section_name in section_matches:
                keywords = section_matches[section_name]
                if any(kw in image_section for kw in keywords):
                    section_match = True

        # Check if day/column matches
        day_match = image.get("cell_index") == col_idx

        # Both must match for structure-based placement
        return section_match and day_match

    def _calculate_match_confidence(
        self,
        cell_text: str,
        media: Dict,
        day_name: str = None,
        section_name: str = None,
    ) -> tuple:
        """
        Calculate match confidence with multiple strategies.

        Matching strategies (in order):
        1. Exact text match
        2. Semantic similarity (if available)
        3. Fuzzy context match
        4. Hint-based match

        Args:
            cell_text: Text content of the cell
            media: Media dictionary (hyperlink or image)
            day_name: Day name for hint matching
            section_name: Section name for hint matching

        Returns:
            (confidence_score, match_type) tuple
        """
        try:
            from rapidfuzz import fuzz
        except ImportError:
            logger.warning(
                "rapidfuzz_not_installed", extra={"fallback": "exact_match_only"}
            )
            # Fallback to exact matching only
            if "text" in media and media["text"] in cell_text:
                return (1.0, "exact_text")
            return (0.0, "no_match")

        # Normalize day hints for case-insensitive comparison
        day_hint_normalized = None
        if media.get("day_hint"):
            day_hint_normalized = media["day_hint"].lower().strip()

        day_name_normalized = None
        if day_name:
            day_name_normalized = day_name.lower().strip()

        # CRITICAL for Schema 2.0: If day_hint is provided, MUST match the current day
        # This prevents links leaking to other days with similar text
        if day_hint_normalized and day_name_normalized:
            if day_hint_normalized != day_name_normalized:
                return (0.0, "day_mismatch")

        # Strategy 1: Exact text match (hyperlinks only)
        if "text" in media and media["text"] in cell_text:
            return (1.0, "exact_text")

        # Strategy 2: Context fuzzy match with hint-based boosting
        context = media.get("context_snippet", "")
        context_score = 0.0
        hint_matches = 0

        # Count hint matches (case-insensitive)
        if day_name_normalized and day_hint_normalized == day_name_normalized:
            hint_matches += 1
            
        if section_name:
            hint = media.get("section_hint", "").lower()
            # Flexible section mapping (shared logic)
            section_mappings = {
                "unit_lesson": ["unit", "lesson", "module"],
                "objective": ["objective", "goal", "swbat"],
                "anticipatory_set": ["anticipatory", "warm up", "hook", "do now", "entry"],
                "tailored_instruction": ["instruction", "activity", "procedure", "lesson", "tailored", "differentiation"],
                "misconceptions": ["misconception", "misconceptions", "error", "pitfall"],
                "assessment": ["assessment", "check", "evaluate", "exit ticket"],
                "homework": ["homework", "assignment", "practice"]
            }
            
            if hint == section_name:
                hint_matches += 1
            elif section_name in section_mappings:
                if hint in section_mappings[section_name]:
                    hint_matches += 1
                elif any(kw in hint for kw in section_mappings[section_name] if len(kw) > 3):
                    hint_matches += 1

        if context:
            context_score = fuzz.partial_ratio(context, cell_text) / 100.0

            # NEW: Lower threshold if hints match (bilingual content support)
            if hint_matches == 2:
                # Both hints match - accept lower text similarity
                if context_score >= 0.40:  # Lower bar with strong hints
                    boosted_score = min(1.0, context_score + 0.15)
                    return (
                        boosted_score,
                        f"context_bilingual_with_{hint_matches}_hints",
                    )
            elif hint_matches == 1:
                # One hint matches - slightly lower threshold
                if context_score >= 0.45:
                    boosted_score = min(1.0, context_score + 0.10)
                    return (boosted_score, f"context_with_{hint_matches}_hint")

            # Standard threshold (no hint boost)
            if context_score >= FUZZY_MATCH_THRESHOLD:
                if hint_matches > 0:
                    boosted_score = min(1.0, context_score + (hint_matches * 0.05))
                    return (boosted_score, f"context_with_{hint_matches}_hints")
                return (context_score, "fuzzy_context")

        # Strategy 3: Section + day hint match only (weak signal)
        if hint_matches == 2:
            # Both hints match but no context - moderate confidence
            return (0.5, "hints_only")

        return (0.0, "no_match")

    def _inject_hyperlink_inline(self, cell, hyperlink: Dict, row_idx: int = None):
        """Inject hyperlink into cell on its own line.

        Args:
            cell: Cell object
            hyperlink: Hyperlink dictionary with text and url
            row_idx: Row index (for applying bold to unit/lesson row)
        """
        # Check if cell text already contains this link in markdown format
        # This prevents duplicate links when text already has [text](url) format
        link_text = hyperlink.get("text", "")
        link_url = hyperlink.get("url", "")
        cell_text = cell.text
        
        # Check for markdown link pattern [link_text](link_url) in cell text
        import re
        markdown_pattern = rf"\[{re.escape(link_text)}\]\({re.escape(link_url)}\)"
        if re.search(markdown_pattern, cell_text, re.IGNORECASE):
            # #region agent log
            import json
            import time
            with open(r'd:\LP\hyperlink_debug.log', 'a', encoding='utf-8') as f:
                f.write(json.dumps({
                    "sessionId": "debug-session",
                    "runId": "post-fix",
                    "hypothesisId": "FIX",
                    "location": "docx_renderer.py:_inject_hyperlink_inline:skip_duplicate",
                    "message": "Skipping duplicate hyperlink injection",
                    "data": {
                        "link_text": link_text,
                        "link_url": link_url[:50],
                        "cell_text_preview": cell_text[:200],
                        "reason": "already_in_markdown"
                    },
                    "timestamp": int(time.time() * 1000)
                }) + '\n')
            # #endregion
            logger.debug(
                f"Skipping duplicate hyperlink injection: '{link_text}' already in markdown format in cell"
            )
            return

        # CRITICAL: Each hyperlink must start on its own line
        # Create a new paragraph for the hyperlink
        para = cell.add_paragraph()

        # Remove spacing to avoid blank lines
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Add the hyperlink to the new paragraph
        # Pass bold=True if it's the unit/lesson row
        is_bold = (row_idx == self.UNIT_LESSON_ROW)
        self._add_hyperlink(para, hyperlink["text"], hyperlink["url"], bold=is_bold)

    def _place_hyperlink_hybrid(
        self, link: Dict, table, structure: StructureMetadata
    ) -> str:
        """
        Place hyperlink using hybrid strategy (v2.0) with detailed diagnostics.

        Strategies (in order):
        1. Coordinate-based (if standard structure)
        2. Label + Day matching (adaptive)
        3. Fuzzy text matching (fallback)
        4. Referenced Links section (last resort)

        Args:
            link: Hyperlink dictionary with v2.0 schema
            table: Output table object
            structure: Detected structure metadata

        Returns:
            Strategy used: 'coordinate', 'label_day', 'fuzzy', or 'fallback'
        """

        # Create diagnostic record
        diagnostic = {
            "link_text": link.get("text", "")[:50],
            "url": link.get("url", "")[:50],
            "input_coords": f"table={link.get('table_idx')}, row={link.get('row_idx')}, cell={link.get('cell_idx')}",
            "row_label": link.get("row_label", ""),
            "col_header": link.get("col_header", ""),
            "day_hint": link.get("day_hint", ""),
            "section_hint": link.get("section_hint", ""),
            "structure_type": structure.structure_type,
            "strategy_attempted": [],
        }

        # Skip non-table links
        if link.get("table_idx") is None:
            diagnostic["strategy_attempted"].append("skipped_non_table")
            diagnostic["result"] = "fallback"
            logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
            return "fallback"

        # Strategy 1: Coordinate-based (if standard structure)
        if structure.structure_type == "standard_8x6":
            diagnostic["strategy_attempted"].append("coordinate")
            if self._try_coordinate_placement(link, table, structure):
                diagnostic["result"] = "success_coordinate"
                logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
                return "coordinate"
            diagnostic["coordinate_failure"] = "mismatch_or_bounds"

        # Strategy 2: Label + Day matching
        diagnostic["strategy_attempted"].append("label_day")
        diagnostic["label_lookup"] = structure.get_row_index(link.get("row_label", ""))
        diagnostic["day_lookup"] = structure.get_col_index(link.get("day_hint", ""))

        if self._try_label_day_placement(link, table, structure):
            diagnostic["result"] = "success_label_day"
            logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
            return "label_day"
        diagnostic["label_day_failure"] = (
            f"row={diagnostic['label_lookup']}, col={diagnostic['day_lookup']}"
        )

        # Strategy 3: Fuzzy matching
        diagnostic["strategy_attempted"].append("fuzzy")
        if self._try_fuzzy_placement(link, table, threshold=FUZZY_MATCH_THRESHOLD):
            diagnostic["result"] = "success_fuzzy"
            logger.info("hyperlink_placement_diagnostic", extra=diagnostic)
            return "fuzzy"
        diagnostic["fuzzy_failure"] = "no_match_above_threshold"

        # Strategy 4: Fallback (will be added to Referenced Links)
        diagnostic["result"] = "fallback"
        logger.warning("hyperlink_placement_fallback", extra=diagnostic)
        return "fallback"

    def _try_coordinate_placement(
        self, link: Dict, table, structure: StructureMetadata
    ) -> bool:
        """Try to place link at exact coordinates."""

        row_idx = link.get("row_idx")
        cell_idx = link.get("cell_idx")

        if row_idx is None or cell_idx is None:
            return False

        # Apply row offset if needed
        target_row = row_idx + structure.row_offset

        # Guard against invalid coordinates
        try:
            if target_row >= len(table.rows):
                logger.warning(
                    f"Row {target_row} out of bounds (table has {len(table.rows)} rows)"
                )
                return False

            row = table.rows[target_row]

            if cell_idx >= len(row.cells):
                logger.warning(
                    f"Cell {cell_idx} out of bounds (row has {len(row.cells)} cells)"
                )
                return False

            cell = row.cells[cell_idx]

            # Place hyperlink
            self._inject_hyperlink_inline(cell, link)

            logger.debug(
                f"Placed '{link['text']}' at coordinates ({target_row}, {cell_idx})"
            )
            return True

        except (IndexError, AttributeError) as e:
            logger.warning(f"Coordinate placement failed for '{link['text']}': {e}")
            return False

    def _try_label_day_placement(
        self, link: Dict, table, structure: StructureMetadata
    ) -> bool:
        """Try to place link using row label + day matching."""

        row_label = link.get("row_label", "").strip().lower().rstrip(":")
        day_hint = link.get("day_hint", "").strip().lower()

        if not row_label or not day_hint:
            return False

        # Find target row
        target_row = structure.get_row_index(row_label)

        # Find target column
        target_col = structure.get_col_index(day_hint)

        if target_row is None or target_col is None:
            logger.debug(
                f"Could not find row/col for '{link['text']}': row={row_label}, day={day_hint}"
            )
            return False

        # Guard against invalid coordinates
        try:
            if target_row >= len(table.rows) or target_col >= len(
                table.rows[target_row].cells
            ):
                logger.warning(
                    f"Label/day placement out of bounds: ({target_row}, {target_col})"
                )
                return False

            cell = table.rows[target_row].cells[target_col]
            self._inject_hyperlink_inline(cell, link)

            logger.debug(
                f"Placed '{link['text']}' via label/day at ({target_row}, {target_col})"
            )
            return True

        except (IndexError, AttributeError) as e:
            logger.warning(f"Label/day placement failed for '{link['text']}': {e}")
            return False

    def _try_fuzzy_placement(self, link: Dict, table, threshold=0.65) -> bool:
        """Try to place link using fuzzy text matching with detailed scoring."""

        best_match = {
            "confidence": 0.0,
            "location": None,
            "cell_text_preview": "",
            "match_type": "none",
        }

        # Iterate through all cells in table
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Use existing _calculate_match_confidence method
                # Extract day/section from table structure
                day_name = None
                section_name = None

                # Get day from column header
                if table.rows and cell_idx < len(table.rows[0].cells):
                    col_header = table.rows[0].cells[cell_idx].text.strip().lower()
                    days = ["monday", "tuesday", "wednesday", "thursday", "friday"]
                    for day in days:
                        if day in col_header:
                            day_name = day
                            break

                # Get section from row label
                if row.cells:
                    row_label = row.cells[0].text.strip().lower()
                    if "objective" in row_label:
                        section_name = "objective"
                    elif "instruction" in row_label or "tailored" in row_label:
                        section_name = "instruction"
                    elif "assessment" in row_label:
                        section_name = "assessment"
                    elif "homework" in row_label:
                        section_name = "homework"

                confidence, match_type = self._calculate_match_confidence(
                    cell.text, link, day_name=day_name, section_name=section_name
                )

                # Track best match even if below threshold
                if confidence > best_match["confidence"]:
                    best_match["confidence"] = confidence
                    best_match["location"] = f"({row_idx}, {cell_idx})"
                    best_match["cell_text_preview"] = cell.text[:80]
                    best_match["match_type"] = match_type

                if confidence >= threshold:
                    self._inject_hyperlink_inline(cell, link)
                    logger.debug(
                        f"Placed '{link['text']}' via fuzzy matching "
                        f"at ({row_idx}, {cell_idx}), confidence={confidence:.2f}, "
                        f"match_type={match_type}"
                    )
                    return True

        # Log best match even if it failed
        logger.debug(
            "fuzzy_match_best_attempt",
            extra={
                "link_text": link.get("text", "")[:50],
                "best_confidence": best_match["confidence"],
                "best_location": best_match["location"],
                "threshold": threshold,
                "cell_preview": best_match["cell_text_preview"],
                "match_type": best_match["match_type"],
            },
        )

        return False

    def _inject_image_inline(self, cell, image: Dict, max_width: float):
        """Inject image into cell with width constraints.

        Args:
            cell: Cell object
            image: Image dictionary with base64 data
            max_width: Maximum width in inches
        """
        try:
            image_data = base64.b64decode(image["data"])
            image_stream = BytesIO(image_data)

            # Add to new paragraph
            para = cell.add_paragraph()
            run = para.add_run()
            run.add_picture(
                image_stream, width=Inches(max_width * 0.9)
            )  # 90% of column width

            # Optional: tiny caption
            caption = cell.add_paragraph()
            caption_run = caption.add_run(f"[{image.get('filename', 'image')}]")
            caption_run.font.size = Pt(7)
            caption_run.italic = True

        except Exception as e:
            logger.warning(
                "inline_image_failed",
                extra={"filename": image.get("filename"), "error": str(e)},
            )

    def _append_unmatched_media(
        self, doc: Document, hyperlinks: List[Dict], images: List[Dict]
    ):
        """Append unmatched media to end with context for traceability.

        Args:
            doc: Document object
            hyperlinks: List of unmatched hyperlinks
            images: List of unmatched images
        """
        if hyperlinks:
            doc.add_page_break()
            heading = doc.add_paragraph("Referenced Links")
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(12)

            for link in hyperlinks:
                para = doc.add_paragraph()
                try:
                    para.style = "List Bullet"
                except KeyError:
                    para.add_run("• ")

                self._add_hyperlink(para, link["text"], link["url"])

                # Add context hint
                if link.get("context_snippet"):
                    context_para = doc.add_paragraph()
                    context_run = context_para.add_run(
                        f'  Context: "{link["context_snippet"][:80]}..."'
                    )
                    context_run.font.size = Pt(9)
                    context_run.italic = True

        if images:
            if not hyperlinks:
                doc.add_page_break()

            heading = doc.add_paragraph("Attached Images")
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(12)

            for i, image in enumerate(images, 1):
                try:
                    image_data = base64.b64decode(image["data"])
                    image_stream = BytesIO(image_data)

                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(4.0))

                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(
                        f"Image {i}: {image.get('filename', 'Untitled')}"
                    )
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)

                    # Add context hint
                    if image.get("context_snippet"):
                        context_para = doc.add_paragraph()
                        context_run = context_para.add_run(
                            f'Context: "{image["context_snippet"][:80]}..."'
                        )
                        context_run.font.size = Pt(9)
                        context_run.italic = True

                except Exception as e:
                    logger.warning(
                        "fallback_image_failed",
                        extra={"filename": image.get("filename"), "error": str(e)},
                    )

    def _insert_images(self, doc: Document, images: List[Dict]):
        """Insert images at the end of the document.

        Args:
            doc: Document object
            images: List of image dictionaries with base64-encoded data
        """
        if not images:
            return

        try:
            # Add a page break before images section
            doc.add_page_break()

            # Add section header
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Attached Images")
            heading_run.bold = True
            heading_run.font.size = Pt(14)

            # Insert each image
            for i, image in enumerate(images, 1):
                try:
                    # Decode base64 image data
                    image_data = base64.b64decode(image["data"])
                    image_stream = BytesIO(image_data)

                    # Add image with caption
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()

                    # Insert image with reasonable width (4 inches)
                    run.add_picture(image_stream, width=Inches(4.0))

                    # Add caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(
                        f"Image {i}: {image.get('filename', 'Untitled')}"
                    )
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)

                except Exception as e:
                    logger.warning(
                        "image_insertion_failed",
                        extra={
                            "image_index": i,
                            "filename": image.get("filename", "unknown"),
                            "error": str(e),
                        },
                    )
        except Exception as e:
            logger.warning("images_insertion_error", extra={"error": str(e)})

    def _restore_hyperlinks(self, doc: Document, hyperlinks: List[Dict]):
        """Restore hyperlinks by adding them at the end of the document.

        Args:
            doc: Document object
            hyperlinks: List of hyperlink dictionaries with text and url
        """
        if not hyperlinks:
            return

        # Log fallback placement for observability
        logger.info(
            "hyperlink_fallback_placement",
            extra={
                "total_fallback_links": len(hyperlinks),
                "links": [
                    {
                        "text": link.get("text", "")[:50],
                        "url": link.get("url", "")[:50],
                        "day_hint": link.get("day_hint"),
                        "section_hint": link.get("section_hint"),
                        "slot": link.get("_source_slot"),
                        "subject": link.get("_source_subject"),
                    }
                    for link in hyperlinks
                ],
            },
        )

        try:
            # Add section header
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Referenced Links")
            heading_run.bold = True
            heading_run.font.size = Pt(12)

            # Add each hyperlink
            for link in hyperlinks:
                try:
                    paragraph = doc.add_paragraph()
                    # Try to use List Bullet style if available, otherwise use default
                    try:
                        paragraph.style = "List Bullet"
                    except KeyError:
                        # Style doesn't exist, add bullet manually
                        paragraph.style = "Normal"
                        paragraph.paragraph_format.left_indent = Inches(0.25)
                        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                        paragraph.add_run("• ")  # Add bullet manually

                    # Add hyperlink text and URL
                    self._add_hyperlink(paragraph, link["text"], link["url"])

                except Exception as e:
                    logger.warning(
                        "hyperlink_restoration_failed",
                        extra={
                            "text": link.get("text", "unknown"),
                            "url": link.get("url", "unknown"),
                            "error": str(e),
                        },
                    )
        except Exception as e:
            logger.warning("hyperlinks_restoration_error", extra={"error": str(e)})

    def _force_font_tnr8(self, run, is_bold: bool = False, is_hyperlink: bool = False):
        """Force Times New Roman 8pt on a Run object.
        Reconstructs w:rPr from scratch to ensure correct XML schema order and fix corruption.
        """
        r_elem = run._element
        rPr = r_elem.get_or_add_rPr()
        
        # Capture existing formatting from OXML if possible
        existing_rStyle = rPr.find(qn("w:rStyle"))
        existing_color = rPr.find(qn("w:color"))
        existing_u = rPr.find(qn("w:u"))
        existing_b = rPr.find(qn("w:b"))
        existing_bCs = rPr.find(qn("w:bCs"))
        existing_i = rPr.find(qn("w:i"))
        existing_iCs = rPr.find(qn("w:iCs"))
        existing_sz = rPr.find(qn("w:sz"))
        existing_szCs = rPr.find(qn("w:szCs"))
        existing_vertAlign = rPr.find(qn("w:vertAlign"))
        
        # Clear ALL children of rPr to ensure perfect order
        for child in list(rPr):
            rPr.remove(child)
            
        # 1. rStyle (Must be first)
        if existing_rStyle is not None:
            rPr.append(existing_rStyle)
            
        # 2. rFonts (Set all 4 attributes to prevent Fallback)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rPr.append(rFonts)
        
        # 3. Bold
        if is_bold:
            rPr.append(OxmlElement("w:b"))
            rPr.append(OxmlElement("w:bCs"))
        elif existing_b is not None:
            rPr.append(existing_b)
            if existing_bCs is not None:
                rPr.append(existing_bCs)
            else:
                rPr.append(OxmlElement("w:bCs"))
            
        # 4. Color
        if is_hyperlink:
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0000FF")
            rPr.append(color)
        elif existing_color is not None:
            rPr.append(existing_color)
            
        # 5. Size (8pt = 16 half-points)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "16")
        rPr.append(szCs)
        
        # 6. Underline
        if is_hyperlink:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        elif existing_u is not None:
            rPr.append(existing_u)

    def _force_font_arial10(self, run, is_bold: bool = False):
        """Force Arial 10pt on a Run object.
        Used for metadata table formatting.
        Reconstructs w:rPr from scratch to ensure correct XML schema order.
        """
        r_elem = run._element
        rPr = r_elem.get_or_add_rPr()
        
        # Capture existing formatting from OXML if possible
        existing_rStyle = rPr.find(qn("w:rStyle"))
        existing_color = rPr.find(qn("w:color"))
        existing_u = rPr.find(qn("w:u"))
        
        # Clear ALL children of rPr to ensure perfect order
        for child in list(rPr):
            rPr.remove(child)
            
        # 1. rStyle (Must be first)
        if existing_rStyle is not None:
            rPr.append(existing_rStyle)
            
        # 2. rFonts (Set all 4 attributes to prevent Fallback)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        rFonts.set(qn("w:cs"), "Arial")
        rFonts.set(qn("w:eastAsia"), "Arial")
        rPr.append(rFonts)
        
        # 3. Bold
        if is_bold:
            rPr.append(OxmlElement("w:b"))
            rPr.append(OxmlElement("w:bCs"))
            
        # 4. Color
        if existing_color is not None:
            rPr.append(existing_color)
            
        # 5. Size (10pt = 20 half-points)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "20")
        rPr.append(szCs)
        
        # 6. Underline
        if existing_u is not None:
            rPr.append(existing_u)

    def _add_hyperlink(self, paragraph, text: str, url: str, bold: bool = False, insert_at: int = None):
        """Add a hyperlink to a paragraph.
        Using direct OxmlElement construction to avoid out-of-sync paragraph.runs
        which causes DOCX corruption.
        
        Args:
            paragraph: Paragraph object
            text: Link text
            url: Link URL
            bold: Whether to bold the link text
            insert_at: Optional index to insert at (for maintaining order). If None, appends to end.
        """
        # Get the paragraph's part
        part = paragraph.part

        # Create relationship to the URL
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create the hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create the run element directly
        r_elem = OxmlElement("w:r")
        
        # We wrap it in a Run object temporarily to use our _force_font_tnr8 helper
        from docx.text.run import Run
        temp_run = Run(r_elem, paragraph)
        self._force_font_tnr8(temp_run, is_bold=bold, is_hyperlink=True)
        
        # Add the text with preservation of spaces
        t_elem = OxmlElement("w:t")
        t_elem.set(qn("xml:space"), "preserve")
        t_elem.text = text
        r_elem.append(t_elem)
        
        # Add run to hyperlink
        hyperlink.append(r_elem)
        
        # Insert at specified position or append to end
        if insert_at is not None and insert_at < len(paragraph._p):
            paragraph._p.insert(insert_at, hyperlink)
        else:
            paragraph._p.append(hyperlink)

    def _format_homework(self, homework: Dict) -> str:
        """Format homework section."""
        if not homework:
            return ""

        # Originals mode: just return the raw text
        if self.is_originals:
            return homework.get("original_content") or ""

        parts = []

        if "original_content" in homework and homework["original_content"]:
            parts.append(homework["original_content"])

        if "family_connection" in homework:
            parts.append(f"\n**Family Connection:** {homework['family_connection']}")

        return "\n".join(parts)


def main():
    """CLI entry point."""
    if len(sys.argv) < 3:
        print(
            "Usage: python docx_renderer.py <input.json> <output.docx> [template.docx]"
        )
        print("\nExample:")
        print(
            "  python docx_renderer.py tests/fixtures/valid_lesson_minimal.json output/lesson.docx"
        )
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    template_path = (
        sys.argv[3] if len(sys.argv) > 3 else "input/Lesson Plan Template SY'25-26.docx"
    )

    # Load JSON
    with open(input_path, "r", encoding="utf-8") as f:
        json_data = json.load(f)

    # Render
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
