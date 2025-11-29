"""Quick check of hyperlink placement in output file."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

print(f"Checking: {output_file.name}\n")

doc = Document(output_file)

# Count links in daily plans table (table 1)
daily_table = doc.tables[1] if len(doc.tables) > 1 else None
inline_count = 0

if daily_table:
    for row in daily_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                inline_count += len(hyperlinks)

print(f"Links in daily plans table: {inline_count}")

# Count links in Referenced Links section
fallback_count = 0
for table in doc.tables:
    # Check if this is the Referenced Links table
    if table.rows and table.rows[0].cells:
        first_cell = table.rows[0].cells[0].text
        if 'Referenced Links' in first_cell or 'Referenced Media' in first_cell:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        hyperlinks = para._element.xpath('.//w:hyperlink')
                        fallback_count += len(hyperlinks)

print(f"Links in Referenced Links: {fallback_count}")

total = inline_count + fallback_count
inline_rate = (inline_count / total * 100) if total > 0 else 0

print(f"\nTotal links: {total}")
print(f"Inline rate: {inline_rate:.1f}%")
print(f"Fallback rate: {100-inline_rate:.1f}%")

print(f"\n{'='*60}")
if inline_rate >= 45:
    print("✓ PASS: Inline rate ≥ 45%")
else:
    print("✗ FAIL: Inline rate < 45%")
