"""
Automated Validation Script for Threshold Change

Analyzes input and output files to validate hyperlink placement.
Compares inline vs. fallback rates and detects false positives.
"""

import sys
from pathlib import Path
from collections import defaultdict
import json

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.docx_parser import DOCXParser
from docx import Document

class ThresholdValidator:
    """Validate threshold change by comparing input/output files."""
    
    def __init__(self):
        self.results = {
            'files': [],
            'aggregate': {
                'total_links': 0,
                'inline': 0,
                'fallback': 0,
                'broken': 0,
                'false_positives': []
            }
        }
    
    def validate_file_pair(self, input_path: str, output_path: str, file_name: str):
        """
        Validate one input/output pair.
        
        Args:
            input_path: Path to input DOCX
            output_path: Path to output DOCX
            file_name: Friendly name for reporting
        """
        print(f"\n{'='*80}")
        print(f"Validating: {file_name}")
        print(f"{'='*80}\n")
        
        # Parse input to get all hyperlinks
        print("1. Parsing input file...")
        parser = DOCXParser(input_path)
        input_links = parser.extract_hyperlinks()
        print(f"   Found {len(input_links)} hyperlinks in input")
        
        # Parse output to categorize links
        print("\n2. Analyzing output file...")
        output_doc = Document(output_path)
        
        # Extract all hyperlinks from output
        output_links = self._extract_output_links(output_doc)
        
        # Categorize links
        inline_links = []
        fallback_links = []
        
        # Check if there's a "Referenced Links" section
        has_fallback_section = self._has_referenced_links_section(output_doc)
        
        if has_fallback_section:
            # Find where fallback section starts
            fallback_start = self._find_fallback_section_start(output_doc)
            
            for link in output_links:
                if link['paragraph_index'] >= fallback_start:
                    fallback_links.append(link)
                else:
                    inline_links.append(link)
        else:
            # No fallback section - all links are inline
            inline_links = output_links
        
        print(f"   Inline links: {len(inline_links)}")
        print(f"   Fallback links: {len(fallback_links)}")
        
        # Calculate rates
        total = len(input_links)
        inline_count = len(inline_links)
        fallback_count = len(fallback_links)
        
        inline_rate = (inline_count / total * 100) if total > 0 else 0
        fallback_rate = (fallback_count / total * 100) if total > 0 else 0
        
        # Check for broken links (links in input but not in output)
        output_urls = set(link['url'] for link in output_links)
        input_urls = set(link['url'] for link in input_links)
        broken_urls = input_urls - output_urls
        
        # Detect potential false positives
        false_positives = self._detect_false_positives(inline_links, output_doc)
        
        # Store results
        result = {
            'file': file_name,
            'input_path': input_path,
            'output_path': output_path,
            'total_links': total,
            'inline': inline_count,
            'fallback': fallback_count,
            'inline_rate': inline_rate,
            'fallback_rate': fallback_rate,
            'broken': len(broken_urls),
            'broken_urls': list(broken_urls),
            'false_positives': false_positives,
            'false_positive_rate': (len(false_positives) / total * 100) if total > 0 else 0
        }
        
        self.results['files'].append(result)
        
        # Update aggregate
        self.results['aggregate']['total_links'] += total
        self.results['aggregate']['inline'] += inline_count
        self.results['aggregate']['fallback'] += fallback_count
        self.results['aggregate']['broken'] += len(broken_urls)
        self.results['aggregate']['false_positives'].extend(false_positives)
        
        # Print summary
        print(f"\n3. Results:")
        print(f"   Total links: {total}")
        print(f"   Inline: {inline_count} ({inline_rate:.1f}%)")
        print(f"   Fallback: {fallback_count} ({fallback_rate:.1f}%)")
        print(f"   Broken: {len(broken_urls)}")
        print(f"   Potential false positives: {len(false_positives)}")
        
        if broken_urls:
            print(f"\n   ⚠️  BROKEN LINKS:")
            for url in list(broken_urls)[:5]:
                print(f"      - {url[:80]}")
        
        if false_positives:
            print(f"\n   ⚠️  POTENTIAL FALSE POSITIVES:")
            for fp in false_positives[:5]:
                print(f"      - {fp['text'][:60]} (confidence: {fp['confidence']:.2f})")
        
        return result
    
    def _extract_output_links(self, doc: Document) -> list:
        """Extract all hyperlinks from output document with locations."""
        links = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Check for hyperlinks in paragraph
            for run in paragraph.runs:
                if run._element.xpath('.//w:hyperlink'):
                    hyperlinks = run._element.xpath('.//w:hyperlink')
                    for hyperlink in hyperlinks:
                        url = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        
                        if url and text:
                            links.append({
                                'text': text,
                                'url': url,
                                'paragraph_index': para_idx,
                                'paragraph_text': paragraph.text[:100]
                            })
        
        # Also check tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run in paragraph.runs:
                            if run._element.xpath('.//w:hyperlink'):
                                hyperlinks = run._element.xpath('.//w:hyperlink')
                                for hyperlink in hyperlinks:
                                    url = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                    
                                    if url and text:
                                        links.append({
                                            'text': text,
                                            'url': url,
                                            'table_idx': table_idx,
                                            'row_idx': row_idx,
                                            'cell_idx': cell_idx,
                                            'location': f"T{table_idx}R{row_idx}C{cell_idx}"
                                        })
        
        return links
    
    def _has_referenced_links_section(self, doc: Document) -> bool:
        """Check if document has a 'Referenced Links' section."""
        for paragraph in doc.paragraphs:
            if 'Referenced Links' in paragraph.text or 'Referenced Media' in paragraph.text:
                return True
        return False
    
    def _find_fallback_section_start(self, doc: Document) -> int:
        """Find paragraph index where fallback section starts."""
        for idx, paragraph in enumerate(doc.paragraphs):
            if 'Referenced Links' in paragraph.text or 'Referenced Media' in paragraph.text:
                return idx
        return len(doc.paragraphs)  # No fallback section
    
    def _detect_false_positives(self, inline_links: list, doc: Document) -> list:
        """
        Detect potential false positives.
        
        Heuristics:
        - Duplicate link text in different locations (might be wrong cell)
        - Generic link text ("Stage 6", "Cool Down") with low confidence
        """
        false_positives = []
        
        # Group by link text
        by_text = defaultdict(list)
        for link in inline_links:
            by_text[link['text']].append(link)
        
        # Check for duplicates with generic text
        generic_patterns = ['stage', 'level', 'cool down', 'activity', 'worksheet']
        
        for text, links in by_text.items():
            if len(links) > 1:  # Duplicate text
                # Check if it's generic
                is_generic = any(pattern in text.lower() for pattern in generic_patterns)
                
                if is_generic:
                    # Potential false positive
                    for link in links:
                        false_positives.append({
                            'text': text,
                            'location': link.get('location', 'unknown'),
                            'reason': 'duplicate_generic_text',
                            'confidence': 0.5  # Low confidence
                        })
        
        return false_positives
    
    def print_summary(self):
        """Print validation summary."""
        print(f"\n\n{'='*80}")
        print("VALIDATION SUMMARY")
        print(f"{'='*80}\n")
        
        agg = self.results['aggregate']
        
        if agg['total_links'] == 0:
            print("No links found to validate!")
            return
        
        inline_rate = (agg['inline'] / agg['total_links'] * 100)
        fallback_rate = (agg['fallback'] / agg['total_links'] * 100)
        fp_rate = (len(agg['false_positives']) / agg['total_links'] * 100)
        
        print(f"Total links analyzed: {agg['total_links']}")
        print(f"\nPlacement:")
        print(f"  Inline: {agg['inline']} ({inline_rate:.1f}%)")
        print(f"  Fallback: {agg['fallback']} ({fallback_rate:.1f}%)")
        print(f"\nIssues:")
        print(f"  Broken links: {agg['broken']}")
        print(f"  Potential false positives: {len(agg['false_positives'])} ({fp_rate:.1f}%)")
        
        # Per-file breakdown
        print(f"\n\nPER-FILE RESULTS:")
        print(f"{'File':<30} {'Total':<8} {'Inline':<10} {'Fallback':<10} {'FP':<8}")
        print(f"{'-'*70}")
        
        for result in self.results['files']:
            print(f"{result['file']:<30} {result['total_links']:<8} "
                  f"{result['inline']:<10} {result['fallback']:<10} "
                  f"{len(result['false_positives']):<8}")
        
        # Success criteria
        print(f"\n\nSUCCESS CRITERIA:")
        
        criteria = {
            'Inline rate ≥ 45%': inline_rate >= 45,
            'FP rate ≤ 5%': fp_rate <= 5,
            'Zero broken links': agg['broken'] == 0
        }
        
        all_pass = all(criteria.values())
        
        for criterion, passed in criteria.items():
            status = '✓ PASS' if passed else '✗ FAIL'
            print(f"  {status}: {criterion}")
        
        print(f"\n\nFINAL RESULT:")
        if all_pass:
            print(f"  ✓ VALIDATION PASSED")
            print(f"  → Threshold change is working well")
            print(f"  → Safe to deploy to production")
        else:
            print(f"  ✗ VALIDATION FAILED")
            print(f"  → Issues found that need attention")
            print(f"  → Review results before deploying")
        
        return all_pass


def main():
    """Run validation on W43 files."""
    
    print("="*80)
    print("AUTOMATED THRESHOLD VALIDATION")
    print("="*80)
    print("\nValidating threshold change (0.65 → 0.55)")
    print("Analyzing W43 lesson plans\n")
    
    base_dir = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43')
    
    # Define input/output pairs
    # You'll need to update the output paths after processing
    pairs = [
        {
            'name': 'Davies W43',
            'input': base_dir / '10_20-10_24 Davies Lesson Plans.docx',
            'output': base_dir / 'Wilson_Rodrigues_Weekly_W43_10-20-10-24_LATEST.docx'  # Update this
        },
        {
            'name': 'Lang W43',
            'input': base_dir / 'Lang Lesson Plans 10_20_25-10_24_25.docx',
            'output': base_dir / 'Wilson_Rodrigues_Weekly_W43_10-20-10-24_LATEST.docx'  # Update this
        },
        {
            'name': 'Savoca W43',
            'input': base_dir / 'Ms. Savoca-10_20_25-10_25_25 Lesson plans.docx',
            'output': base_dir / 'Wilson_Rodrigues_Weekly_W43_10-20-10-24_LATEST.docx'  # Update this
        }
    ]
    
    validator = ThresholdValidator()
    
    for pair in pairs:
        if pair['input'].exists() and pair['output'].exists():
            try:
                validator.validate_file_pair(
                    str(pair['input']),
                    str(pair['output']),
                    pair['name']
                )
            except Exception as e:
                print(f"\n❌ Error validating {pair['name']}: {e}")
        else:
            print(f"\n⚠️  Skipping {pair['name']}: Files not found")
            if not pair['input'].exists():
                print(f"   Input missing: {pair['input']}")
            if not pair['output'].exists():
                print(f"   Output missing: {pair['output']}")
    
    # Print summary
    passed = validator.print_summary()
    
    # Save results
    results_path = Path('d:/LP/docs/validation/W43_validation_results.json')
    results_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(results_path, 'w') as f:
        json.dump(validator.results, f, indent=2)
    
    print(f"\n💾 Detailed results saved to: {results_path}")
    
    return 0 if passed else 1


if __name__ == '__main__':
    sys.exit(main())
