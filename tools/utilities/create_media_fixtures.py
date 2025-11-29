"""
Create test fixtures with actual images and hyperlinks.
"""

import sys
from pathlib import Path
from io import BytesIO

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Inches
from PIL import Image
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_test_image():
    """Create a simple test image (100x100 red square)."""
    img = Image.new('RGB', (100, 100), color='red')
    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_lesson_with_image():
    """Create a lesson plan DOCX with an actual embedded image."""
    print("Creating lesson_with_image.docx...")
    
    doc = Document()
    
    # Add title
    doc.add_heading('Lesson Plan - Math', level=1)
    
    # Add some content
    doc.add_paragraph('Teacher: Ms. Smith')
    doc.add_paragraph('Grade: 3')
    doc.add_paragraph('Week of: 10/14-10/18')
    
    # Add lesson content
    doc.add_heading('Monday - Fractions', level=2)
    doc.add_paragraph('Objective: Students will understand basic fractions.')
    
    # Add materials section with image
    doc.add_heading('Materials', level=3)
    doc.add_paragraph('Visual aids for fraction demonstration:')
    
    # Create and add test image
    img_bytes = create_test_image()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_bytes, width=Inches(2.0))
    
    # Add caption
    caption = doc.add_paragraph()
    caption.add_run('Figure 1: Fraction circles').italic = True
    
    # Add more content
    doc.add_paragraph('Activities: Students will work with fraction manipulatives.')
    
    # Save
    output_path = Path('tests/fixtures/lesson_with_image.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"  Saved: {output_path}")


def create_lesson_with_hyperlinks():
    """Create a lesson plan DOCX with actual hyperlinks."""
    print("Creating lesson_with_hyperlinks.docx...")
    
    doc = Document()
    
    # Add title
    doc.add_heading('Lesson Plan - Science', level=1)
    
    # Add some content
    doc.add_paragraph('Teacher: Mr. Johnson')
    doc.add_paragraph('Grade: 4')
    doc.add_paragraph('Week of: 10/14-10/18')
    
    # Add lesson content
    doc.add_heading('Wednesday - Solar System', level=2)
    doc.add_paragraph('Objective: Students will learn about planets.')
    
    # Add resources section with hyperlinks
    doc.add_heading('Online Resources', level=3)
    
    # Add first hyperlink
    p1 = doc.add_paragraph('Visit ')
    add_hyperlink(p1, 'NASA Solar System Exploration', 'https://solarsystem.nasa.gov/')
    p1.add_run(' for interactive content.')
    
    # Add second hyperlink
    p2 = doc.add_paragraph('Watch ')
    add_hyperlink(p2, 'Educational Video on YouTube', 'https://www.youtube.com/watch?v=example')
    p2.add_run(' for visual learning.')
    
    # Add third hyperlink
    p3 = doc.add_paragraph('Download ')
    add_hyperlink(p3, 'Worksheet from Teachers Pay Teachers', 'https://www.teacherspayteachers.com/example')
    p3.add_run(' for practice.')
    
    # Add more content
    doc.add_paragraph('Assessment: Students will complete a planet identification quiz.')
    
    # Save
    output_path = Path('tests/fixtures/lesson_with_hyperlinks.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"  Saved: {output_path}")


def create_lesson_with_both():
    """Create a lesson plan DOCX with both images and hyperlinks."""
    print("Creating lesson_with_both.docx...")
    
    doc = Document()
    
    # Add title
    doc.add_heading('Lesson Plan - ELA', level=1)
    
    # Add some content
    doc.add_paragraph('Teacher: Mrs. Davis')
    doc.add_paragraph('Grade: 5')
    doc.add_paragraph('Week of: 10/14-10/18')
    
    # Add lesson content
    doc.add_heading('Thursday - Story Elements', level=2)
    doc.add_paragraph('Objective: Students will identify story elements.')
    
    # Add image
    doc.add_heading('Visual Aid', level=3)
    img_bytes = create_test_image()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_bytes, width=Inches(1.5))
    caption = doc.add_paragraph()
    caption.add_run('Story elements diagram').italic = True
    
    # Add hyperlinks
    doc.add_heading('Additional Resources', level=3)
    p1 = doc.add_paragraph('Read ')
    add_hyperlink(p1, 'Story Elements Guide', 'https://www.readwritethink.org/example')
    p1.add_run(' for more information.')
    
    p2 = doc.add_paragraph('Practice with ')
    add_hyperlink(p2, 'Interactive Story Map', 'https://www.scholastic.com/example')
    p2.add_run(' online.')
    
    # Add more content
    doc.add_paragraph('Assessment: Students will create their own story map.')
    
    # Save
    output_path = Path('tests/fixtures/lesson_with_both.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"  Saved: {output_path}")


def main():
    """Create all media fixtures."""
    print("=" * 60)
    print("Creating Media Test Fixtures")
    print("=" * 60)
    print()
    
    try:
        create_lesson_with_image()
        create_lesson_with_hyperlinks()
        create_lesson_with_both()
        
        print()
        print("=" * 60)
        print("SUCCESS: All fixtures created")
        print("=" * 60)
        return True
    except Exception as e:
        print()
        print("=" * 60)
        print(f"ERROR: {e}")
        print("=" * 60)
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
