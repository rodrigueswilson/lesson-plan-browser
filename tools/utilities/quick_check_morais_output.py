"""Quick check: Count hyperlinks in most recent Morais output."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')

# Get most recent output file
output_files = sorted(
    [f for f in folder.glob('Daniela_Silva_Weekly_*.docx')],
    key=lambda f: f.stat().st_mtime,
    reverse=True
)

if not output_files:
    print("❌ No output files found!")
else:
    latest = output_files[0]
    print("=" * 80)
    print(f"📄 Latest: {latest.name}")
    print(f"   Modified: {latest.stat().st_mtime}")
    print("=" * 80)
    
    doc = Document(str(latest))
    
    # Count hyperlinks
    para_links = 0
    table_links = 0
    
    for para in doc.paragraphs:
        para_links += len(para._element.xpath('.//w:hyperlink'))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_links += len(para._element.xpath('.//w:hyperlink'))
    
    total = para_links + table_links
    
    print(f"\n📊 RESULTS:")
    print(f"   Paragraph links: {para_links}")
    print(f"   Table links: {table_links}")
    print(f"   TOTAL: {total}")
    
    if total == 16:
        print(f"\n✅ SUCCESS! All 16 hyperlinks preserved!")
    elif total == 0:
        print(f"\n❌ FAILED! No hyperlinks in output (json_merger still broken)")
    else:
        print(f"\n⚠️  PARTIAL! Expected 16, got {total}")
    
    # Check for Referenced Links section
    has_ref_section = any('referenced links' in p.text.lower() for p in doc.paragraphs)
    print(f"\n   'Referenced Links' section: {'YES' if has_ref_section else 'NO'}")
    
    print("\n" + "=" * 80)
