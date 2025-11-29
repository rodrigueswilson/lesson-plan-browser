"""Find all Referenced Links sections."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)

print("Searching for 'Referenced Links' sections...\n")

for table_idx, table in enumerate(doc.tables):
    # Check all cells for "Referenced"
    found_referenced = False
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if 'Referenced Links' in cell.text or 'Referenced Media' in cell.text:
                found_referenced = True
                print(f"Table {table_idx}: Found 'Referenced Links' section")
                
                # Count links
                link_count = 0
                for r in table.rows:
                    for c in r.cells:
                        for para in c.paragraphs:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            link_count += len(hyperlinks)
                
                print(f"  Links in this table: {link_count}")
                
                # Show first few links
                shown = 0
                for r in table.rows:
                    for c in r.cells:
                        for para in c.paragraphs:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            for hyperlink in hyperlinks:
                                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                if shown < 5:
                                    print(f"    {shown+1}. {text[:70]}")
                                    shown += 1
                
                if link_count > 5:
                    print(f"    ... and {link_count - 5} more")
                print()
                break
        if found_referenced:
            break

print("\nSummary:")
print(f"Total tables: {len(doc.tables)}")
print(f"Tables with links: {sum(1 for t in doc.tables if any(len(p._element.xpath('.//w:hyperlink')) > 0 for r in t.rows for c in r.cells for p in c.paragraphs))}")
