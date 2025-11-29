"""Check every table for links."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)

print("Checking all tables for hyperlinks:\n")

total_links = 0

for table_idx, table in enumerate(doc.tables):
    link_count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                link_count += len(hyperlinks)
    
    if link_count > 0:
        print(f"Table {table_idx}: {link_count} links")
        total_links += link_count
        
        # Show what type of table this is
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()[:50]
            second_cell = table.rows[0].cells[1].text.strip()[:50] if len(table.rows[0].cells) > 1 else ""
            print(f"  First cell: {first_cell}")
            if second_cell:
                print(f"  Second cell: {second_cell}")
        print()

print(f"Total links across all tables: {total_links}")
