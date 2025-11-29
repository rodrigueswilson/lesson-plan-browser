"""Find EVERY hyperlink in the output file."""

from docx import Document
from pathlib import Path
from docx.oxml.ns import qn

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)
rels = doc.part.rels

all_links = []

# Check EVERY table
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                for hyperlink in hyperlinks:
                    r_id = hyperlink.get(qn('r:id'))
                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                    
                    url = None
                    if r_id and r_id in rels:
                        url = rels[r_id].target_ref
                    
                    if text and url:
                        all_links.append({
                            'text': text,
                            'url': url,
                            'table': table_idx,
                            'row': row_idx,
                            'cell': cell_idx
                        })

# Check paragraphs (outside tables)
for para_idx, para in enumerate(doc.paragraphs):
    hyperlinks = para._element.xpath('.//w:hyperlink')
    for hyperlink in hyperlinks:
        r_id = hyperlink.get(qn('r:id'))
        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
        
        url = None
        if r_id and r_id in rels:
            url = rels[r_id].target_ref
        
        if text and url:
            all_links.append({
                'text': text,
                'url': url,
                'paragraph': para_idx
            })

print(f"TOTAL HYPERLINKS IN DOCUMENT: {len(all_links)}")
print()

# Group by table
by_table = {}
for link in all_links:
    if 'table' in link:
        table_idx = link['table']
        if table_idx not in by_table:
            by_table[table_idx] = []
        by_table[table_idx].append(link)

print("Links per table:")
for table_idx in sorted(by_table.keys()):
    print(f"  Table {table_idx}: {len(by_table[table_idx])} links")

print()
print(f"Links in paragraphs: {sum(1 for l in all_links if 'paragraph' in l)}")
