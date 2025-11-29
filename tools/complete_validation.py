"""
Complete validation: properly handle multi-slot inputs and find all fallback sections.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.docx_parser import DOCXParser
from docx import Document

OUTPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx'

INPUT_FILES = [
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Davies Lesson Plans.docx',
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Lang Lesson Plans 10_20_25-10_24_25.docx',
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Ms. Savoca-10_20_25-10_25_25 Lesson plans.docx'
]


def analyze_output_structure(output_path):
    """Analyze complete output structure including fallback sections."""
    
    doc = Document(output_path)
    
    structure = []
    current_item = None
    
    for idx, table in enumerate(doc.tables):
        if not table.rows or not table.rows[0].cells:
            continue
        
        first_cell = table.rows[0].cells[0].text.strip()
        second_cell = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        
        # Count links in this table
        link_count = 0
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    hyperlinks = para._element.xpath('.//w:hyperlink')
                    link_count += len(hyperlinks)
        
        # Identify table type
        if "Name:" in first_cell:
            # Metadata table - start of new slot
            if current_item:
                structure.append(current_item)
            
            # Extract slot info
            subject = ""
            grade = ""
            for cell in table.rows[0].cells:
                text = cell.text.strip()
                if "Subject:" in text:
                    subject = text.split("Subject:")[-1].strip()
                elif "Grade:" in text:
                    grade = text.split("Grade:")[-1].strip()
            
            current_item = {
                'type': 'slot',
                'name': f"Grade {grade} - {subject}",
                'metadata_table': idx,
                'daily_table': None,
                'inline_links': 0,
                'fallback_table': None,
                'fallback_links': 0
            }
        
        elif "MONDAY" in second_cell.upper() and current_item:
            # Daily plans table
            current_item['daily_table'] = idx
            current_item['inline_links'] = link_count
        
        elif ("Referenced Links" in first_cell or "Referenced Media" in first_cell) and current_item:
            # Fallback section for current slot
            current_item['fallback_table'] = idx
            current_item['fallback_links'] = link_count
    
    # Don't forget the last slot
    if current_item:
        structure.append(current_item)
    
    return structure


def count_input_links(input_path):
    """Count total links in input file."""
    parser = DOCXParser(input_path)
    links = parser.extract_hyperlinks()
    return len(links)


def main():
    print("="*100)
    print("COMPLETE VALIDATION - INPUT vs OUTPUT")
    print("="*100)
    print()
    
    # Analyze output structure
    print("1. Analyzing output structure...")
    output_structure = analyze_output_structure(OUTPUT_FILE)
    
    print(f"   Found {len(output_structure)} slots in output\n")
    
    # Count input links
    print("2. Counting input links...")
    input_totals = []
    for input_file in INPUT_FILES:
        path = Path(input_file)
        if path.exists():
            count = count_input_links(input_file)
            input_totals.append((path.name, count))
            print(f"   {path.name}: {count} links")
        else:
            print(f"   ⚠️  Not found: {path.name}")
    
    total_input = sum(count for _, count in input_totals)
    print(f"\n   Total input links: {total_input}")
    
    # Show output structure
    print(f"\n\n{'='*100}")
    print("OUTPUT STRUCTURE")
    print(f"{'='*100}\n")
    
    print(f"{'Slot':<40} {'Inline':<10} {'Fallback':<10} {'Total':<10} {'Inline %':<10}")
    print("-"*90)
    
    total_inline = 0
    total_fallback = 0
    
    for slot in output_structure:
        inline = slot['inline_links']
        fallback = slot['fallback_links']
        total = inline + fallback
        inline_pct = (inline / total * 100) if total > 0 else 0
        
        print(f"{slot['name']:<40} {inline:<10} {fallback:<10} {total:<10} {inline_pct:<10.1f}%")
        
        total_inline += inline
        total_fallback += fallback
    
    print("-"*90)
    total_output = total_inline + total_fallback
    overall_pct = (total_inline / total_output * 100) if total_output > 0 else 0
    print(f"{'TOTAL':<40} {total_inline:<10} {total_fallback:<10} {total_output:<10} {overall_pct:<10.1f}%")
    
    # Show fallback details
    if total_fallback > 0:
        print(f"\n\n{'='*100}")
        print("FALLBACK LINKS DETAILS")
        print(f"{'='*100}\n")
        
        doc = Document(OUTPUT_FILE)
        
        for slot in output_structure:
            if slot['fallback_links'] > 0:
                print(f"\n{slot['name']}: {slot['fallback_links']} links in fallback")
                print("-"*80)
                
                # Get the fallback table
                fallback_table = doc.tables[slot['fallback_table']]
                
                # Show links
                link_num = 0
                for row in fallback_table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            for hyperlink in hyperlinks:
                                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                link_num += 1
                                if link_num <= 10:
                                    print(f"  {link_num}. {text[:70]}")
                
                if link_num > 10:
                    print(f"  ... and {link_num - 10} more")
    
    # Final assessment
    print(f"\n\n{'='*100}")
    print("FINAL ASSESSMENT")
    print(f"{'='*100}\n")
    
    print(f"Total input links: {total_input}")
    print(f"Total output links: {total_output}")
    
    if total_input == total_output:
        print(f"✓ Link counts match")
    else:
        diff = abs(total_input - total_output)
        print(f"⚠️  Link count mismatch: {diff} links difference")
    
    print(f"\nPlacement:")
    print(f"  Inline: {total_inline} ({overall_pct:.1f}%)")
    print(f"  Fallback: {total_fallback} ({100-overall_pct:.1f}%)")
    
    print(f"\n{'='*100}")
    if overall_pct == 100:
        print("✓ PERFECT: All links placed inline!")
    elif overall_pct >= 80:
        print(f"✓ EXCELLENT: {overall_pct:.1f}% inline (far exceeds 45% target)")
    elif overall_pct >= 45:
        print(f"✓ PASS: {overall_pct:.1f}% inline (meets 45% target)")
    else:
        print(f"✗ FAIL: {overall_pct:.1f}% inline (below 45% target)")
    print(f"{'='*100}")


if __name__ == '__main__':
    sys.exit(main())
