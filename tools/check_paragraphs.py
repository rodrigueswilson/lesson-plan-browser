"""Check paragraphs for Referenced Links sections."""

from docx import Document
from pathlib import Path
from docx.oxml.ns import qn

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)
rels = doc.part.rels

print("Checking paragraphs around 'Referenced Links'...\n")

for para_idx, para in enumerate(doc.paragraphs):
    if 'Referenced Links' in para.text:
        print(f"{'='*80}")
        print(f"Found 'Referenced Links' at paragraph {para_idx}")
        print(f"{'='*80}\n")
        
        # Show context (10 paragraphs after this one)
        print("Following paragraphs:")
        link_count = 0
        for i in range(para_idx, min(para_idx + 50, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            
            # Count links in this paragraph
            hyperlinks = p._element.xpath('.//w:hyperlink')
            if hyperlinks:
                link_count += len(hyperlinks)
                for hyperlink in hyperlinks:
                    r_id = hyperlink.get(qn('r:id'))
                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                    
                    url = None
                    if r_id and r_id in rels:
                        url = rels[r_id].target_ref
                    
                    if text:
                        print(f"  • {text[:70]}")
            
            # Stop if we hit another section or empty paragraphs
            if i > para_idx and ('Referenced Links' in p.text or 'Required Signatures' in p.text):
                break
        
        print(f"\nTotal links in this section: {link_count}\n")

print("\n" + "="*80)
print("SUMMARY")
print("="*80)

# Count all links in paragraphs
total_para_links = 0
for para in doc.paragraphs:
    hyperlinks = para._element.xpath('.//w:hyperlink')
    total_para_links += len(hyperlinks)

print(f"Total links in paragraphs (outside tables): {total_para_links}")

# Count links in tables
total_table_links = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                total_table_links += len(hyperlinks)

print(f"Total links in tables: {total_table_links}")
print(f"Grand total: {total_para_links + total_table_links}")
