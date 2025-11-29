"""
Validate input-output matching per slot.

Match each slot in output to its input file and compare:
1. Number of links (should match)
2. Which links are inline vs fallback
3. Links with metadata that should be inline but aren't
"""

import sys
from pathlib import Path
from collections import defaultdict

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.docx_parser import DOCXParser
from docx import Document

# Define input files in the order they were processed
INPUT_FILES = [
    {
        'name': 'Davies (Grade 3 ELA)',
        'path': r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Davies Lesson Plans.docx',
        'expected_slot': 'Grade 3 - ELA'
    },
    {
        'name': 'Lang (Grade 2 ELA/SS)',
        'path': r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Lang Lesson Plans 10_20_25-10_24_25.docx',
        'expected_slot': 'Grade 2 - ELA/SS'
    },
    {
        'name': 'Savoca (Grade 2 Math)',
        'path': r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Ms. Savoca-10_20_25-10_25_25 Lesson plans.docx',
        'expected_slot': 'Grade 2 - Math'
    }
]

OUTPUT_FILE = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx'


def get_output_slots(output_path):
    """Extract slots from output file."""
    doc = Document(output_path)
    
    slots = []
    current_slot = None
    
    for idx, table in enumerate(doc.tables):
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()
            second_cell = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
            
            if "Name:" in first_cell:
                # Extract subject and grade
                subject = ""
                grade = ""
                for cell in table.rows[0].cells:
                    text = cell.text.strip()
                    if "Subject:" in text:
                        subject = text.split("Subject:")[-1].strip()
                    elif "Grade:" in text:
                        grade = text.split("Grade:")[-1].strip()
                
                slot_name = f"Grade {grade} - {subject}"
                current_slot = {
                    'name': slot_name,
                    'metadata_table': idx,
                    'daily_table': None,
                    'inline_links': 0,
                    'fallback_links': 0
                }
                slots.append(current_slot)
            
            elif ("MONDAY" in second_cell.upper() or "Monday" in second_cell) and current_slot:
                current_slot['daily_table'] = idx
                
                # Count links
                link_count = 0
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            link_count += len(hyperlinks)
                
                current_slot['inline_links'] = link_count
            
            elif ("Referenced Links" in first_cell or "Referenced Media" in first_cell) and current_slot:
                link_count = 0
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            link_count += len(hyperlinks)
                
                current_slot['fallback_links'] = link_count
    
    return slots


def main():
    print("="*100)
    print("INPUT-OUTPUT VALIDATION")
    print("="*100)
    print()
    
    # Get output slots
    output_slots = get_output_slots(OUTPUT_FILE)
    
    print(f"Found {len(output_slots)} slots in output\n")
    
    # Validate each input against its output slot
    print("="*100)
    print("VALIDATION RESULTS")
    print("="*100)
    print()
    
    for i, input_file in enumerate(INPUT_FILES):
        print(f"\n{'='*100}")
        print(f"SLOT {i+1}: {input_file['name']}")
        print(f"{'='*100}\n")
        
        input_path = Path(input_file['path'])
        
        if not input_path.exists():
            print(f"⚠️  Input file not found: {input_path}")
            continue
        
        # Parse input
        print(f"1. Parsing input file...")
        parser = DOCXParser(str(input_path))
        input_links = parser.extract_hyperlinks()
        
        print(f"   Input links: {len(input_links)}")
        print(f"   With section_hint: {sum(1 for l in input_links if l.get('section_hint'))}")
        print(f"   With day_hint: {sum(1 for l in input_links if l.get('day_hint'))}")
        
        # Find matching output slot
        output_slot = None
        if i < len(output_slots):
            output_slot = output_slots[i]
        
        if not output_slot:
            print(f"\n⚠️  No matching output slot found!")
            continue
        
        print(f"\n2. Matching output slot: {output_slot['name']}")
        print(f"   Output inline links: {output_slot['inline_links']}")
        print(f"   Output fallback links: {output_slot['fallback_links']}")
        
        # Compare counts
        print(f"\n3. Comparison:")
        
        input_count = len(input_links)
        output_total = output_slot['inline_links'] + output_slot['fallback_links']
        
        if input_count == output_total:
            print(f"   ✓ Link count matches: {input_count} input = {output_total} output")
        else:
            print(f"   ⚠️  Link count mismatch: {input_count} input ≠ {output_total} output")
            print(f"      Difference: {abs(input_count - output_total)} links")
        
        # Calculate rates
        inline_rate = (output_slot['inline_links'] / input_count * 100) if input_count > 0 else 0
        fallback_rate = (output_slot['fallback_links'] / input_count * 100) if input_count > 0 else 0
        
        print(f"\n4. Placement rates:")
        print(f"   Inline: {inline_rate:.1f}% ({output_slot['inline_links']}/{input_count})")
        print(f"   Fallback: {fallback_rate:.1f}% ({output_slot['fallback_links']}/{input_count})")
        
        # Analyze metadata
        with_section = sum(1 for l in input_links if l.get('section_hint'))
        with_day = sum(1 for l in input_links if l.get('day_hint'))
        with_both = sum(1 for l in input_links if l.get('section_hint') and l.get('day_hint'))
        
        print(f"\n5. Metadata coverage:")
        print(f"   With section_hint: {with_section}/{input_count} ({with_section/input_count*100:.1f}%)")
        print(f"   With day_hint: {with_day}/{input_count} ({with_day/input_count*100:.1f}%)")
        print(f"   With both: {with_both}/{input_count} ({with_both/input_count*100:.1f}%)")
        
        # Success assessment
        print(f"\n6. Assessment:")
        if inline_rate == 100:
            print(f"   ✓ EXCELLENT: 100% inline placement")
        elif inline_rate >= 80:
            print(f"   ✓ GOOD: {inline_rate:.1f}% inline (target: ≥45%)")
        elif inline_rate >= 45:
            print(f"   ✓ PASS: {inline_rate:.1f}% inline (target: ≥45%)")
        else:
            print(f"   ✗ FAIL: {inline_rate:.1f}% inline (target: ≥45%)")
        
        if output_slot['fallback_links'] > 0:
            print(f"   ⚠️  {output_slot['fallback_links']} links in fallback section")
            
            # Analyze which links failed
            links_with_meta = with_both
            if output_slot['fallback_links'] <= links_with_meta:
                print(f"      Some links with metadata failed to place inline")
            else:
                print(f"      Mostly links without metadata (expected)")
    
    # Overall summary
    print(f"\n\n{'='*100}")
    print("OVERALL SUMMARY")
    print(f"{'='*100}\n")
    
    total_input = 0
    total_inline = 0
    total_fallback = 0
    
    for i, input_file in enumerate(INPUT_FILES):
        input_path = Path(input_file['path'])
        if input_path.exists():
            parser = DOCXParser(str(input_path))
            input_links = parser.extract_hyperlinks()
            total_input += len(input_links)
            
            if i < len(output_slots):
                total_inline += output_slots[i]['inline_links']
                total_fallback += output_slots[i]['fallback_links']
    
    overall_inline_rate = (total_inline / total_input * 100) if total_input > 0 else 0
    
    print(f"Total input links: {total_input}")
    print(f"Total output inline: {total_inline} ({overall_inline_rate:.1f}%)")
    print(f"Total output fallback: {total_fallback} ({100-overall_inline_rate:.1f}%)")
    
    print(f"\n{'='*100}")
    if overall_inline_rate == 100:
        print("✓ PERFECT: All links placed inline!")
    elif overall_inline_rate >= 80:
        print(f"✓ EXCELLENT: {overall_inline_rate:.1f}% inline (far exceeds 45% target)")
    elif overall_inline_rate >= 45:
        print(f"✓ PASS: {overall_inline_rate:.1f}% inline (meets 45% target)")
    else:
        print(f"✗ FAIL: {overall_inline_rate:.1f}% inline (below 45% target)")
    print(f"{'='*100}")


if __name__ == '__main__':
    sys.exit(main())
