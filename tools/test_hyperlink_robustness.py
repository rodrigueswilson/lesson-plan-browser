"""
Test script for verifying robust hyperlink matching in DOCXRenderer.
"""
import sys
from pathlib import Path
from docx import Document

# Add project root to sys.path
sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_renderer import DOCXRenderer

TEMPLATE_FILE = r"d:\LP\input\Lesson Plan Template SY'25-26.docx"
OUTPUT_FILE = r"d:\LP\test_hyperlink_robustness.docx"

def test_matching():
    # Test cases: (cell_text, link_text, expected_inline)
    test_cases = [
        # Exact match
        ("Unit 3: Lenni Lenape - Lesson 10", "Lenni Lenape", True),
        # Space mismatch (multiple spaces)
        ("Unit 3: Lenni  Lenape - Lesson 10", "Lenni Lenape", True),
        # Punctuation at end
        ("Unit 3: Lenni Lenape, Lesson 10", "Lenni Lenape", True),
        # Case mismatch (should already work)
        ("Unit 3: lenni lenape - Lesson 10", "Lenni Lenape", True),
        # Slight rephrasing (should NOT match inline if too different, but our fuzzy logic might handle it)
        # Here we just test the Smart Inline Replacement regex
    ]

    renderer = DOCXRenderer(TEMPLATE_FILE)
    
    # We'll mock a small lesson plan for each case
    days_data = {}
    pending_hyperlinks = []
    
    for i, (cell_text, link_text, _) in enumerate(test_cases):
        day_key = f"day_{i}"
        days_data[day_key] = {
            "unit_lesson": cell_text,
            "objective": {"content": "Test"},
            "anticipatory_set": {"content": "Test"},
            "tailored_instruction": {"content": "Test"},
            "misconceptions": {"content": "Test"},
            "assessment": {"content": "Test"},
            "homework": {"content": "Test"}
        }
        pending_hyperlinks.append({
            "text": link_text,
            "url": f"http://example.com/{i}",
            "day_hint": day_key,
            "section_hint": "unit_lesson",
            "_source_slot": 1,
            "_source_subject": "Test"
        })

    # Prepare complete data
    json_data = {
        "metadata": {
            "teacher_name": "Test",
            "week_dates": "10/20-10/24",
            "grade": "3",
            "subject": "Test"
        },
        "days": days_data,
        "_hyperlinks": pending_hyperlinks
    }

    print(f"Rendering {len(test_cases)} cases to {OUTPUT_FILE}...")
    success = renderer.render(json_data, OUTPUT_FILE)
    
    if success:
        print("✅ Success! Please verify the output in Word.")
        # Check if any links were fallback'd (meaning they didn't match inline)
        # Note: renderer.render clears pending_hyperlinks as it processes them
        # but we can check if it outputted them at the top of the cell.
        
        # We can use python-docx to check for hyperlinks in the cells
        doc = Document(OUTPUT_FILE)
        table = doc.tables[1]
        
        # Define namespaces
        from docx.oxml.ns import qn
        
        for i in range(len(test_cases)):
            cell = table.rows[renderer._get_row_index("unit")].cells[i+1]
            print(f"\nCase {i}: '{test_cases[i][0]}'")
            print(f"  Cell text: '{cell.text.strip()}'")
            
            # Count hyperlinks in cell using iter()
            hyperlink_tag = qn("w:hyperlink")
            links = [child for child in cell._element.iter() if child.tag == hyperlink_tag]
            print(f"  Found {len(links)} links in cell.")
            
            if len(links) > 0:
                # Check if it's inline
                para_with_link = None
                for p_idx, para in enumerate(cell.paragraphs):
                    para_links = [child for child in para._element.iter() if child.tag == hyperlink_tag]
                    if para_links:
                        para_with_link = p_idx
                        break
                
                if para_with_link == 0:
                    print(f"  ✅ Link is INLINE in paragraph 0")
                elif para_with_link is not None:
                    print(f"  ❌ Link is in paragraph {para_with_link} (expected paragraph 0)")
                else:
                    print(f"  ❌ Link element found but not in any paragraph?")
            else:
                print(f"  ❌ FAILED: Link not found inline or at top.")
    else:
        print("❌ Rendering failed")

if __name__ == "__main__":
    test_matching()
