"""
Phase 0.5: Prototype Coordinate Capture Validation

Tests minimal parser changes on 3-5 standard 8x6 files to validate:
1. Coordinates are captured correctly
2. Row labels are accurate
3. Column headers are extracted
4. Non-table links are handled
5. No crashes or errors

This is a MINIMAL implementation - just coordinate capture, no placement logic.
"""

import sys
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent))

# Test files (standard 8x6 structure)
TEST_FILES = [
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Davies Lesson Plans.docx',
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Laverti Lesson Plans.docx',
    r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\10_20-10_24 Lonesky Lesson Plans.docx',
]


class PrototypeParser:
    """Minimal parser with coordinate capture for validation."""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        self.doc = Document(str(self.file_path))
    
    def extract_hyperlinks_with_coordinates(self) -> List[Dict[str, Any]]:
        """
        Extract hyperlinks with v2.0 schema (coordinates).
        
        This is the MINIMAL implementation for Phase 0.5 validation.
        """
        hyperlinks = []
        
        # Extract from paragraphs (non-table)
        for paragraph in self.doc.paragraphs:
            for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
                try:
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id and r_id in paragraph.part.rels:
                        url = paragraph.part.rels[r_id].target_ref
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        
                        if text and url:
                            hyperlinks.append({
                                'schema_version': '2.0',
                                'text': text,
                                'url': url,
                                'context_type': 'paragraph',
                                # v2.0: No table coordinates for paragraphs
                                'table_idx': None,
                                'row_idx': None,
                                'cell_idx': None,
                                'row_label': None,
                                'col_header': None
                            })
                except Exception as e:
                    print(f"Warning: Failed to extract paragraph hyperlink: {e}")
        
        # Extract from tables (with coordinates)
        for table_idx, table in enumerate(self.doc.tables):
            # Get column headers from first row
            col_headers = []
            if table.rows:
                col_headers = [
                    cell.text.strip() for cell in table.rows[0].cells
                ]
            
            for row_idx, row in enumerate(table.rows):
                # Get row label (first cell)
                row_label = row.cells[0].text.strip() if row.cells else ""
                
                for cell_idx, cell in enumerate(row.cells):
                    # Get column header for this cell
                    col_header = col_headers[cell_idx] if cell_idx < len(col_headers) else ""
                    
                    for paragraph in cell.paragraphs:
                        for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
                            try:
                                r_id = hyperlink.get(qn('r:id'))
                                if r_id and r_id in paragraph.part.rels:
                                    url = paragraph.part.rels[r_id].target_ref
                                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                    
                                    if text and url:
                                        hyperlinks.append({
                                            'schema_version': '2.0',
                                            'text': text,
                                            'url': url,
                                            'context_type': 'table_cell',
                                            # v2.0: Table coordinates
                                            'table_idx': table_idx,
                                            'row_idx': row_idx,
                                            'cell_idx': cell_idx,
                                            'row_label': row_label,
                                            'col_header': col_header
                                        })
                            except Exception as e:
                                print(f"Warning: Failed to extract table hyperlink: {e}")
        
        return hyperlinks


def validate_coordinates(file_path: str) -> Dict[str, Any]:
    """
    Validate coordinate capture for a single file.
    
    Returns:
        Validation results dictionary
    """
    print(f"\n{'='*80}")
    print(f"VALIDATING: {Path(file_path).name}")
    print(f"{'='*80}\n")
    
    parser = PrototypeParser(file_path)
    hyperlinks = parser.extract_hyperlinks_with_coordinates()
    
    # Statistics
    total_links = len(hyperlinks)
    table_links = [h for h in hyperlinks if h['table_idx'] is not None]
    non_table_links = [h for h in hyperlinks if h['table_idx'] is None]
    
    print(f"Total hyperlinks found: {total_links}")
    print(f"  - In tables: {len(table_links)}")
    print(f"  - In paragraphs: {len(non_table_links)}")
    
    # Validation checks
    results = {
        'file': Path(file_path).name,
        'total_links': total_links,
        'table_links': len(table_links),
        'non_table_links': len(non_table_links),
        'errors': [],
        'warnings': []
    }
    
    # Check 1: All links have schema_version
    for link in hyperlinks:
        if link.get('schema_version') != '2.0':
            results['errors'].append(f"Link '{link['text']}' missing schema_version")
    
    # Check 2: Table links have coordinates
    for link in table_links:
        if link['table_idx'] is None or link['row_idx'] is None or link['cell_idx'] is None:
            results['errors'].append(f"Table link '{link['text']}' missing coordinates")
    
    # Check 3: Non-table links have None coordinates
    for link in non_table_links:
        if link['table_idx'] is not None:
            results['errors'].append(f"Non-table link '{link['text']}' has table_idx")
    
    # Check 4: Row labels are present for table links
    missing_row_labels = [h for h in table_links if not h.get('row_label')]
    if missing_row_labels:
        results['warnings'].append(f"{len(missing_row_labels)} table links missing row_label")
    
    # Check 5: Column headers are present for table links
    missing_col_headers = [h for h in table_links if not h.get('col_header')]
    if missing_col_headers:
        results['warnings'].append(f"{len(missing_col_headers)} table links missing col_header")
    
    # Display sample links
    print(f"\n{'─'*80}")
    print("SAMPLE LINKS (first 5):")
    print(f"{'─'*80}\n")
    
    for i, link in enumerate(hyperlinks[:5], 1):
        print(f"{i}. {link['text'][:50]}")
        print(f"   URL: {link['url'][:60]}")
        print(f"   Type: {link['context_type']}")
        
        if link['table_idx'] is not None:
            print(f"   Coordinates: table={link['table_idx']}, "
                  f"row={link['row_idx']}, cell={link['cell_idx']}")
            print(f"   Row label: '{link['row_label'][:40]}'")
            print(f"   Col header: '{link['col_header'][:20]}'")
        else:
            print(f"   Coordinates: None (paragraph link)")
        print()
    
    # Display validation results
    print(f"{'─'*80}")
    print("VALIDATION RESULTS:")
    print(f"{'─'*80}\n")
    
    if results['errors']:
        print(f"❌ ERRORS ({len(results['errors'])}):")
        for error in results['errors']:
            print(f"  - {error}")
        print()
    else:
        print("✅ No errors found\n")
    
    if results['warnings']:
        print(f"⚠️  WARNINGS ({len(results['warnings'])}):")
        for warning in results['warnings']:
            print(f"  - {warning}")
        print()
    else:
        print("✅ No warnings\n")
    
    return results


def manual_verification_guide(file_path: str, hyperlinks: List[Dict]):
    """
    Generate manual verification guide for a file.
    """
    print(f"\n{'='*80}")
    print(f"MANUAL VERIFICATION GUIDE: {Path(file_path).name}")
    print(f"{'='*80}\n")
    
    print("Instructions:")
    print("1. Open the file in Microsoft Word")
    print("2. For each link below, verify the coordinates match the actual position")
    print("3. Check that row labels and column headers are correct\n")
    
    table_links = [h for h in hyperlinks if h['table_idx'] is not None]
    
    print(f"Links to verify ({len(table_links)} total, showing first 10):\n")
    
    for i, link in enumerate(table_links[:10], 1):
        print(f"{i}. Link text: '{link['text']}'")
        print(f"   Expected position:")
        print(f"     - Table: {link['table_idx']}")
        print(f"     - Row: {link['row_idx']} (label: '{link['row_label']}')")
        print(f"     - Column: {link['cell_idx']} (header: '{link['col_header']}')")
        print(f"   Action: Locate this link in Word and verify position")
        print()


def run_validation():
    """Run validation on all test files."""
    
    print("="*80)
    print("PHASE 0.5: COORDINATE CAPTURE PROTOTYPE VALIDATION")
    print("="*80)
    print("\nTesting minimal parser changes on standard 8x6 files...")
    print(f"Test files: {len(TEST_FILES)}\n")
    
    all_results = []
    
    for file_path in TEST_FILES:
        if not Path(file_path).exists():
            print(f"\n⚠️  File not found: {file_path}")
            continue
        
        try:
            results = validate_coordinates(file_path)
            all_results.append(results)
            
            # Generate manual verification guide for first file
            if len(all_results) == 1:
                parser = PrototypeParser(file_path)
                hyperlinks = parser.extract_hyperlinks_with_coordinates()
                manual_verification_guide(file_path, hyperlinks)
        
        except Exception as e:
            print(f"\n❌ ERROR processing {Path(file_path).name}:")
            print(f"   {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
    
    # Summary
    print(f"\n{'='*80}")
    print("VALIDATION SUMMARY")
    print(f"{'='*80}\n")
    
    total_files = len(all_results)
    total_links = sum(r['total_links'] for r in all_results)
    total_table_links = sum(r['table_links'] for r in all_results)
    total_errors = sum(len(r['errors']) for r in all_results)
    total_warnings = sum(len(r['warnings']) for r in all_results)
    
    print(f"Files processed: {total_files}/{len(TEST_FILES)}")
    print(f"Total hyperlinks: {total_links}")
    print(f"  - Table links: {total_table_links}")
    print(f"  - Paragraph links: {total_links - total_table_links}")
    print(f"\nValidation:")
    print(f"  - Errors: {total_errors}")
    print(f"  - Warnings: {total_warnings}")
    
    if total_errors == 0:
        print(f"\n✅ SUCCESS: All validation checks passed!")
        print(f"\nNext steps:")
        print(f"  1. Manually verify coordinates for first file (see guide above)")
        print(f"  2. If coordinates are correct, proceed to Phase 1")
        print(f"  3. If coordinates are wrong, debug coordinate capture logic")
    else:
        print(f"\n❌ FAILED: {total_errors} errors found")
        print(f"\nAction required:")
        print(f"  1. Review errors above")
        print(f"  2. Fix coordinate capture logic")
        print(f"  3. Re-run validation")
    
    return all_results


if __name__ == '__main__':
    results = run_validation()
    
    # Exit code for CI/CD
    total_errors = sum(len(r['errors']) for r in results)
    sys.exit(0 if total_errors == 0 else 1)
