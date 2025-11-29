"""
Fix table widths to correct 10 inches for landscape 11" page with 0.5" margins.
"""

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil
from datetime import datetime

# File paths
TEMPLATE_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template.docx'
BACKUP_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template_BACKUP_10INCH_' + datetime.now().strftime('%Y%m%d_%H%M%S') + '.docx'

# Correct width for 11" landscape page with 0.5" margins
TARGET_WIDTH_INCHES = 10.0
TARGET_WIDTH_EMUS = 9144000  # 10 inches * 914400 EMUs per inch
TARGET_WIDTH_DXA = 14400     # 10 inches * 1440 DXA per inch

def set_table_width_xml(table, width_dxa):
    """Set table width in XML."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblW_list = tblPr.findall(qn('w:tblW'))
    for tblW in tblW_list:
        tblPr.remove(tblW)
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_table_indent(table, indent_dxa=0):
    """Set table indent to 0."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblInd_list = tblPr.findall(qn('w:tblInd'))
    for tblInd in tblInd_list:
        tblPr.remove(tblInd)
    
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(indent_dxa))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def set_table_alignment(table, alignment='left'):
    """Set table alignment."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    jc_list = tblPr.findall(qn('w:jc'))
    for jc in jc_list:
        tblPr.remove(jc)
    
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    tblPr.append(jc)

def set_fixed_layout(table):
    """Set table to fixed layout."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    layout_list = tblPr.findall(qn('w:tblLayout'))
    for layout in layout_list:
        tblPr.remove(layout)
    
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

print("=" * 80)
print("FIXING TABLE WIDTHS TO 10 INCHES (LANDSCAPE PAGE)")
print("=" * 80)

# Create backup
print(f"\n1. Creating backup...")
shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)
print(f"   Backup saved to: {BACKUP_PATH}")

# Load document
print(f"\n2. Loading template...")
doc = Document(TEMPLATE_PATH)

# Verify page setup
section = doc.sections[0]
print(f"\n   Page setup:")
print(f"     Page size: {section.page_width.inches:.1f}\" x {section.page_height.inches:.1f}\"")
print(f"     Margins: L={section.left_margin.inches:.1f}\" R={section.right_margin.inches:.1f}\"")
print(f"     Available width: {(section.page_width - section.left_margin - section.right_margin) / 914400:.1f}\"")

print(f"\n   Found {len(doc.tables)} table(s)")

# Fix each table
print(f"\n3. Setting table widths to {TARGET_WIDTH_INCHES} inches...")
for i, table in enumerate(doc.tables):
    print(f"\n   Table {i+1}:")
    
    # Set table width XML property
    set_table_width_xml(table, TARGET_WIDTH_DXA)
    print(f"     ✓ tblW set to {TARGET_WIDTH_DXA} DXA ({TARGET_WIDTH_INCHES} inches)")
    
    # Set zero indent
    set_table_indent(table, indent_dxa=0)
    print(f"     ✓ tblInd set to 0 DXA")
    
    # Set left alignment
    set_table_alignment(table, alignment='left')
    print(f"     ✓ Alignment set to LEFT")
    
    # Set fixed layout
    set_fixed_layout(table)
    print(f"     ✓ Layout set to FIXED")
    
    # Set column widths to match total
    col_width = TARGET_WIDTH_EMUS // len(table.columns)
    remainder = TARGET_WIDTH_EMUS % len(table.columns)
    
    for j, column in enumerate(table.columns):
        if j == len(table.columns) - 1:
            column.width = col_width + remainder
        else:
            column.width = col_width
    
    # Verify total
    total = sum(col.width for col in table.columns)
    print(f"     ✓ Column widths: {total} EMUs = {total/914400:.4f} inches")

# Save document
print(f"\n4. Saving fixed template...")
doc.save(TEMPLATE_PATH)
print(f"   Saved to: {TEMPLATE_PATH}")

print("\n" + "=" * 80)
print("COMPLETE!")
print("=" * 80)
print(f"\nAll tables now:")
print(f"  - Width: {TARGET_WIDTH_INCHES} inches")
print(f"  - Fit perfectly in 11\" landscape page with 0.5\" margins")
print(f"  - Left and right edges aligned")
print(f"\nBackup: {BACKUP_PATH}")
print("=" * 80)
