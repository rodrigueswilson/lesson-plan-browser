"""
Fix table alignment by setting consistent width, indent, and alignment properties.

The issue: Tables can have different tblInd (indent) values causing misalignment.
Solution: Set all tables to:
  - Same width (7.5 inches)
  - Zero indent (tblInd = 0)
  - Left alignment
  - Fixed layout (disable autofit)
"""

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil
from datetime import datetime

# File paths
TEMPLATE_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template.docx'
BACKUP_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template_BACKUP_ALIGNED_' + datetime.now().strftime('%Y%m%d_%H%M%S') + '.docx'

# Target width for 0.5-inch margins
TARGET_WIDTH_INCHES = 7.5
TARGET_WIDTH_EMUS = 6858000
TARGET_WIDTH_DXA = 10800

def set_table_indent(table, indent_dxa=0):
    """
    Set table indent (tblInd) to ensure consistent left alignment.
    
    Args:
        table: python-docx Table object
        indent_dxa: Indent in DXA (twentieths of a point). Default 0 for no indent.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing tblInd if present
    tblInd_list = tblPr.findall(qn('w:tblInd'))
    for tblInd in tblInd_list:
        tblPr.remove(tblInd)
    
    # Add new tblInd with specified value
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(indent_dxa))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def set_table_alignment(table, alignment='left'):
    """
    Set table alignment (jc) to left, center, or right.
    
    Args:
        table: python-docx Table object
        alignment: 'left', 'center', or 'right'
    """
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing jc if present
    jc_list = tblPr.findall(qn('w:jc'))
    for jc in jc_list:
        tblPr.remove(jc)
    
    # Add new jc with specified alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    tblPr.append(jc)

def set_fixed_layout(table):
    """
    Set table to fixed layout (disable autofit).
    
    Args:
        table: python-docx Table object
    """
    table.allow_autofit = False

print("=" * 80)
print("FIXING TABLE ALIGNMENT IN LESSON PLAN TEMPLATE")
print("=" * 80)

# Create backup
print(f"\n1. Creating backup...")
shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)
print(f"   Backup saved to: {BACKUP_PATH}")

# Load document
print(f"\n2. Loading template...")
doc = Document(TEMPLATE_PATH)
print(f"   Found {len(doc.tables)} table(s)")

# Fix each table
print(f"\n3. Fixing table alignment properties...")
for i, table in enumerate(doc.tables):
    print(f"\n   Table {i+1}:")
    print(f"     Rows: {len(table.rows)}, Columns: {len(table.columns)}")
    
    # Set table width
    table.width = Inches(TARGET_WIDTH_INCHES)
    print(f"     ✓ Width set to {TARGET_WIDTH_INCHES} inches")
    
    # Set zero indent (align to left margin)
    set_table_indent(table, indent_dxa=0)
    print(f"     ✓ Indent set to 0 (no left indent)")
    
    # Set left alignment
    set_table_alignment(table, alignment='left')
    print(f"     ✓ Alignment set to LEFT")
    
    # Set fixed layout
    set_fixed_layout(table)
    print(f"     ✓ Layout set to FIXED (autofit disabled)")
    
    # Distribute width evenly across columns
    col_width = TARGET_WIDTH_EMUS // len(table.columns)
    remainder = TARGET_WIDTH_EMUS % len(table.columns)
    
    for j, column in enumerate(table.columns):
        if j == len(table.columns) - 1:
            column.width = col_width + remainder
        else:
            column.width = col_width
    
    print(f"     ✓ Column widths distributed evenly")

# Save document
print(f"\n4. Saving fixed template...")
doc.save(TEMPLATE_PATH)
print(f"   Saved to: {TEMPLATE_PATH}")

print("\n" + "=" * 80)
print("COMPLETE!")
print("=" * 80)
print(f"\nAll tables are now:")
print(f"  - Width: {TARGET_WIDTH_INCHES} inches (fits 0.5\" margins)")
print(f"  - Indent: 0 (aligned to left margin)")
print(f"  - Alignment: LEFT")
print(f"  - Layout: FIXED (no autofit)")
print(f"\nTables should now be perfectly aligned!")
print(f"\nBackup location: {BACKUP_PATH}")
print("=" * 80)
