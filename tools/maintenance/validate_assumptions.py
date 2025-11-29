"""
Validate ALL assumptions before implementing hyperlink placement strategy.
Answers the other AI's critical questions with REAL DATA.
"""

from pathlib import Path
import sys
import json
import re
from collections import defaultdict, Counter

sys.path.insert(0, str(Path(__file__).parent))
from tools.docx_parser import DOCXParser
from docx import Document

class AssumptionValidator:
    """Validate all assumptions with real data."""
    
    def __init__(self):
        self.results = {
            'grouping_accuracy': {},
            'template_structures': [],
            'links_per_cell': [],
            'section_hint_sources': defaultdict(int),
            'classification_confusion': defaultdict(lambda: defaultdict(int))
        }
    
    def validate_grouping_accuracy(self, all_links):
        """
        Test: Can we reliably classify links as Curriculum/Activity/Assessment?
        
        Returns: Confusion matrix, accuracy metrics
        """
        print(f"\n{'='*80}")
        print("VALIDATION 1: Grouping Classification Accuracy")
        print(f"{'='*80}\n")
        
        # Classification rules
        RULES = {
            'Curriculum': [
                r'LESSON\s+\d+',
                r'UNIT\s+\d+',
                r'Teacher\s+Guide',
                r'\d+\.\d+\.\d+\s+Teacher\s+Guide',  # e.g., "3.2.5 Teacher Guide"
                r'Curriculum'
            ],
            'Activity': [
                r'activity',
                r'worksheet',
                r'practice',
                r'carousel',
                r'gallery\s+walk'
            ],
            'Assessment': [
                r'assessment',
                r'test',
                r'quiz',
                r'evaluation'
            ]
        }
        
        def classify(text):
            for category, patterns in RULES.items():
                for pattern in patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        return category
            return 'Other'
        
        # Classify all links
        classifications = Counter()
        examples = defaultdict(list)
        
        for link in all_links:
            category = classify(link['text'])
            classifications[category] += 1
            if len(examples[category]) < 3:
                examples[category].append(link['text'][:60])
        
        # Print results
        total = len(all_links)
        print(f"Total links: {total}\n")
        print(f"Classification Results:")
        for category, count in classifications.most_common():
            pct = (count / total) * 100
            print(f"  {category}: {count} ({pct:.1f}%)")
            print(f"    Examples: {examples[category]}")
            print()
        
        # Calculate "Other" rate (unclassified)
        other_pct = (classifications['Other'] / total) * 100
        
        print(f"KEY METRIC: {other_pct:.1f}% of links are UNCLASSIFIED")
        if other_pct > 30:
            print(f"   ⚠️  WARNING: >30% unclassified - grouping may not work!")
        elif other_pct > 20:
            print(f"   ⚠️  CAUTION: >20% unclassified - needs refinement")
        else:
            print(f"   ✓ GOOD: <20% unclassified")
        
        self.results['grouping_accuracy'] = {
            'total': total,
            'classified': total - classifications['Other'],
            'unclassified': classifications['Other'],
            'unclassified_pct': other_pct,
            'breakdown': dict(classifications)
        }
        
        return classifications
    
    def validate_template_structure(self, input_files):
        """
        Test: Do all files use the same template structure?
        
        Returns: Template variations, consistency metrics
        """
        print(f"\n{'='*80}")
        print("VALIDATION 2: Template Structure Consistency")
        print(f"{'='*80}\n")
        
        structures = []
        
        for file_path in input_files:
            if not file_path.exists():
                continue
            
            try:
                doc = Document(str(file_path))
                if not doc.tables:
                    continue
                
                table = doc.tables[0]
                
                # Extract structure
                structure = {
                    'file': file_path.name,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'row_labels': []
                }
                
                # Get row labels (first cell of each row)
                for i in range(min(10, len(table.rows))):  # First 10 rows
                    label = table.rows[i].cells[0].text.strip()
                    structure['row_labels'].append(label)
                
                structures.append(structure)
                
                print(f"FILE: {file_path.name}")
                print(f"   Rows: {structure['rows']}, Columns: {structure['columns']}")
                print(f"   Labels: {structure['row_labels'][:5]}")
                print()
                
            except Exception as e:
                print(f"❌ Error analyzing {file_path.name}: {e}\n")
        
        # Check consistency
        if not structures:
            print("⚠️  No structures analyzed!")
            return
        
        # Compare row counts
        row_counts = [s['rows'] for s in structures]
        col_counts = [s['columns'] for s in structures]
        
        print(f"\nANALYSIS: CONSISTENCY CHECK:")
        print(f"   Row counts: {set(row_counts)}")
        print(f"   Column counts: {set(col_counts)}")
        
        if len(set(row_counts)) == 1 and len(set(col_counts)) == 1:
            print(f"   ✓ ALL files have identical table dimensions")
        else:
            print(f"   ⚠️  WARNING: Template dimensions vary across files!")
        
        # Compare row labels
        label_sets = [tuple(s['row_labels']) for s in structures]
        unique_label_sets = set(label_sets)
        
        if len(unique_label_sets) == 1:
            print(f"   ✓ ALL files have identical row labels")
        else:
            print(f"   ⚠️  WARNING: {len(unique_label_sets)} different row label patterns found!")
            for i, label_set in enumerate(unique_label_sets, 1):
                print(f"      Pattern {i}: {label_set[:3]}...")
        
        self.results['template_structures'] = structures
        self.results['template_consistent'] = len(unique_label_sets) == 1
        
        return structures
    
    def validate_links_per_cell(self, input_files):
        """
        Test: How many links per cell? What's the distribution?
        
        Returns: Distribution, max, average
        """
        print(f"\n{'='*80}")
        print("VALIDATION 3: Links Per Cell Distribution")
        print(f"{'='*80}\n")
        
        cell_link_counts = defaultdict(int)
        
        for file_path in input_files:
            if not file_path.exists():
                continue
            
            try:
                parser = DOCXParser(str(file_path))
                hyperlinks = parser.extract_hyperlinks()
                
                # Group by cell (section + day)
                for link in hyperlinks:
                    section = link.get('section_hint', 'unknown')
                    day = link.get('day_hint', 'unknown')
                    cell_key = f"{section}_{day}"
                    cell_link_counts[cell_key] += 1
                
            except Exception as e:
                print(f"❌ Error: {e}")
        
        # Analyze distribution
        counts = list(cell_link_counts.values())
        if not counts:
            print("⚠️  No data!")
            return
        
        distribution = Counter(counts)
        
        print(f"Total cells with links: {len(counts)}")
        print(f"Total links: {sum(counts)}")
        print(f"Average links per cell: {sum(counts) / len(counts):.1f}")
        print(f"Max links in one cell: {max(counts)}")
        print()
        
        print(f"Distribution:")
        for link_count in sorted(distribution.keys()):
            cell_count = distribution[link_count]
            pct = (cell_count / len(counts)) * 100
            bar = '█' * int(pct / 2)
            print(f"  {link_count:2d} links: {cell_count:3d} cells ({pct:5.1f}%) {bar}")
        
        # Identify problematic cells
        print(f"\n🎯 OVERFLOW ANALYSIS:")
        cells_with_5plus = sum(1 for c in counts if c >= 5)
        cells_with_8plus = sum(1 for c in counts if c >= 8)
        
        print(f"   Cells with 5+ links: {cells_with_5plus} ({cells_with_5plus/len(counts)*100:.1f}%)")
        print(f"   Cells with 8+ links: {cells_with_8plus} ({cells_with_8plus/len(counts)*100:.1f}%)")
        
        if cells_with_8plus > 0:
            print(f"   ⚠️  {cells_with_8plus} cells will have SEVERE clutter!")
        
        self.results['links_per_cell'] = {
            'distribution': dict(distribution),
            'max': max(counts),
            'average': sum(counts) / len(counts),
            'cells_5plus': cells_with_5plus,
            'cells_8plus': cells_with_8plus
        }
        
        return distribution
    
    def validate_section_hint_accuracy(self, input_files):
        """
        Test: Where do links WITHOUT section_hint actually appear?
        
        Returns: Actual locations of "missing hint" links
        """
        print(f"\n{'='*80}")
        print("VALIDATION 4: Section Hint Accuracy")
        print(f"{'='*80}\n")
        
        missing_hint_links = []
        
        for file_path in input_files:
            if not file_path.exists():
                continue
            
            try:
                parser = DOCXParser(str(file_path))
                hyperlinks = parser.extract_hyperlinks()
                
                for link in hyperlinks:
                    if not link.get('section_hint'):
                        missing_hint_links.append({
                            'text': link['text'],
                            'context': link.get('context_snippet', '')[:80],
                            'file': file_path.name
                        })
                
            except Exception as e:
                pass
        
        print(f"Links without section_hint: {len(missing_hint_links)}")
        print(f"\nExamples (first 10):")
        for i, link in enumerate(missing_hint_links[:10], 1):
            print(f"  [{i}] {link['text'][:50]}")
            print(f"      Context: {link['context']}")
            print()
        
        # Analyze patterns
        print(f"ANALYSIS: PATTERN ANALYSIS:")
        curriculum_pattern = sum(1 for l in missing_hint_links 
                                if re.search(r'LESSON|UNIT|GUIDE', l['text'], re.I))
        
        print(f"   Curriculum-like: {curriculum_pattern} ({curriculum_pattern/len(missing_hint_links)*100:.1f}%)")
        print(f"   → These SHOULD go in Unit/Lesson row")
        
        return missing_hint_links
    
    def print_summary(self):
        """Print validation summary with recommendations."""
        print(f"\n\n{'='*80}")
        print("SUMMARY: VALIDATION SUMMARY & RECOMMENDATIONS")
        print(f"{'='*80}\n")
        
        # Grouping
        if 'grouping_accuracy' in self.results:
            unclass_pct = self.results['grouping_accuracy']['unclassified_pct']
            print(f"1. GROUPING CLASSIFICATION:")
            print(f"   Unclassified: {unclass_pct:.1f}%")
            if unclass_pct > 30:
                print(f"   ❌ REJECT: Too many unclassified - grouping won't work")
            elif unclass_pct > 20:
                print(f"   ⚠️  CAUTION: Needs better rules before using")
            else:
                print(f"   ✓ ACCEPT: Grouping is viable")
        
        # Template
        if 'template_consistent' in self.results:
            print(f"\n2. TEMPLATE CONSISTENCY:")
            if self.results['template_consistent']:
                print(f"   ✓ ACCEPT: All files use same template")
                print(f"   → Can hardcode row mapping")
            else:
                print(f"   ❌ REJECT: Templates vary")
                print(f"   → MUST implement dynamic detection")
        
        # Overflow
        if 'links_per_cell' in self.results:
            cells_8plus = self.results['links_per_cell']['cells_8plus']
            print(f"\n3. CELL OVERFLOW:")
            print(f"   Cells with 8+ links: {cells_8plus}")
            if cells_8plus > 0:
                print(f"   ⚠️  WARNING: Need overflow strategy")
                print(f"   → Recommend: Limit 5-6 per cell, rest to Resources row")
        
        print(f"\n\n🎯 FINAL RECOMMENDATION:")
        print(f"   Based on validation results:")
        
        # Decision logic
        can_proceed = True
        blockers = []
        
        if self.results.get('grouping_accuracy', {}).get('unclassified_pct', 100) > 30:
            can_proceed = False
            blockers.append("Grouping classification too inaccurate")
        
        if not self.results.get('template_consistent', False):
            blockers.append("Template structure varies - need dynamic detection")
        
        if can_proceed and not blockers:
            print(f"   ✓ PROCEED with implementation")
            print(f"   → Use validated grouping rules")
            print(f"   → Hardcode template mapping (all consistent)")
            print(f"   → Implement overflow handling for 8+ links")
        else:
            print(f"   ⚠️  DO NOT PROCEED yet")
            for blocker in blockers:
                print(f"   → Fix: {blocker}")


def main():
    """Run all validations."""
    
    # Input files
    input_dir = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan')
    
    input_files = [
        input_dir / '25 W43' / 'Lang Lesson Plans 10_20_25-10_24_25.docx',
        input_dir / '25 W43' / '10_20-10_24 Davies Lesson Plans.docx',
        input_dir / '25 W43' / 'Ms. Savoca-10_20_25-10_25_25 Lesson plans.docx',
        input_dir / '25 W42' / 'Lang Lesson Plans 10_13_25-10_17_25.docx',
        input_dir / '25 W41' / 'Lang Lesson Plans 10_6_25-10_10_25.docx',
    ]
    
    # Collect all links
    all_links = []
    for file_path in input_files:
        if file_path.exists():
            try:
                parser = DOCXParser(str(file_path))
                all_links.extend(parser.extract_hyperlinks())
            except:
                pass
    
    print(f"Loaded {len(all_links)} hyperlinks from {len(input_files)} files\n")
    
    # Run validations
    validator = AssumptionValidator()
    
    validator.validate_grouping_accuracy(all_links)
    validator.validate_template_structure(input_files)
    validator.validate_links_per_cell(input_files)
    validator.validate_section_hint_accuracy(input_files)
    
    validator.print_summary()
    
    # Save results
    results_path = Path('d:/LP/validation_results.json')
    with open(results_path, 'w') as f:
        json.dump(validator.results, f, indent=2)
    
    print(f"\n💾 Results saved to: {results_path}")


if __name__ == '__main__':
    main()
