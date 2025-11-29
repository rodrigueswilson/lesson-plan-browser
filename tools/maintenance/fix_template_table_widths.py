"""
Fix table widths in the Lesson Plan Template to fit within 0.5-inch margins.
All tables will be set to 7.5 inches (6,858,000 EMUs) to align perfectly.
"""

from docx import Document
from docx.shared import Inches
import shutil
from datetime import datetime

# File paths
TEMPLATE_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template.docx'
BACKUP_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template_BACKUP_' + datetime.now().strftime('%Y%m%d_%H%M%S') + '.docx'

# Target width for 0.5-inch margins
TARGET_WIDTH_INCHES = 7.5
TARGET_WIDTH_EMUS = 6858000

print("=" * 80)
print("FIXING LESSON PLAN TEMPLATE TABLE WIDTHS")
print("=" * 80)

# Create backup
print(f"\n1. Creating backup...")
shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)
print(f"   Backup saved to: {BACKUP_PATH}")

# Load document
print(f"\n2. Loading template...")
doc = Document(TEMPLATE_PATH)
print(f"   Found {len(doc.tables)} table(s)")

# Fix each table
print(f"\n3. Fixing table widths to {TARGET_WIDTH_INCHES} inches...")
for i, table in enumerate(doc.tables):
    # Get current width
    current_width = None
    total_col_width = sum(col.width for col in table.columns if col.width)
    if total_col_width > 0:
        current_width = total_col_width / 914400  # Convert EMUs to inches
    
    print(f"\n   Table {i+1}:")
    print(f"     Current width: {current_width:.2f} inches" if current_width else "     Current width: Not specified")
    print(f"     Rows: {len(table.rows)}, Columns: {len(table.columns)}")
    
    # Set table width
    table.width = Inches(TARGET_WIDTH_INCHES)
    
    # Distribute width evenly across columns
    col_width = TARGET_WIDTH_EMUS // len(table.columns)
    remainder = TARGET_WIDTH_EMUS % len(table.columns)
    
    for j, column in enumerate(table.columns):
        # Add remainder to last column to ensure exact total
        if j == len(table.columns) - 1:
            column.width = col_width + remainder
        else:
            column.width = col_width
    
    # Verify
    new_total = sum(col.width for col in table.columns)
    new_width_inches = new_total / 914400
    print(f"     New width: {new_width_inches:.2f} inches ({new_total:,} EMUs)")
    print(f"     Column widths: {[f'{col.width/914400:.2f}\"' for col in table.columns]}")

# Save document
print(f"\n4. Saving fixed template...")
doc.save(TEMPLATE_PATH)
print(f"   Saved to: {TEMPLATE_PATH}")

print("\n" + "=" * 80)
print("COMPLETE!")
print("=" * 80)
print(f"\nAll tables are now {TARGET_WIDTH_INCHES} inches wide")
print(f"This fits perfectly within 0.5-inch margins on 8.5\" letter paper")
print(f"\nBackup location: {BACKUP_PATH}")
print("=" * 80)
