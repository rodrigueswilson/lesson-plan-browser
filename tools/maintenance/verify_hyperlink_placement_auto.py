"""Verify hyperlinks are in correct positions with proper context - Auto-detect latest files."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import json

def find_most_recent_output(user_name, week="25 W43"):
    """Find the most recent output file for a user."""
    if "Daniela" in user_name or "Silva" in user_name:
        folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP') / week
        pattern = "Daniela_Silva_*.docx"
    elif "Wilson" in user_name or "Rodrigues" in user_name:
        folder = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan') / week
        pattern = "Wilson_*.docx"
    else:
        return None
    
    if not folder.exists():
        return None
    
    output_files = sorted(
        [f for f in folder.glob(pattern)],
        key=lambda f: f.stat().st_mtime,
        reverse=True
    )
    
    return output_files[0] if output_files else None

def find_input_files(user_name, week="25 W43"):
    """Find all input DOCX files for a user."""
    if "Daniela" in user_name or "Silva" in user_name:
        folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP') / week
    elif "Wilson" in user_name or "Rodrigues" in user_name:
        folder = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan') / week
    else:
        return []
    
    if not folder.exists():
        return []
    
    # Get all DOCX files that are NOT output files
    all_files = list(folder.glob('*.docx'))
    input_files = [f for f in all_files if not (f.name.startswith('Daniela_Silva_') or f.name.startswith('Wilson_Rodrigues_'))]
    
    return input_files

def extract_hyperlinks(doc, source_file=None):
    """Extract all hyperlinks from a document."""
    hyperlinks = []
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for hyperlink in para._element.xpath('.//w:hyperlink'):
                        r_id = hyperlink.get(qn('r:id'))
                        if r_id and r_id in para.part.rels:
                            url = para.part.rels[r_id].target_ref
                            text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                            
                            # Get cell text context (more context for better matching)
                            cell_text = cell.text.strip()
                            
                            hyperlinks.append({
                                'text': text,
                                'url': url,
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'cell_text': cell_text[:200],  # More context
                                'source_file': source_file,
                            })
    
    return hyperlinks

def verify_user(user_name):
    """Verify hyperlinks for a user."""
    print("\n" + "=" * 80)
    print(f"{user_name.upper()}")
    print("=" * 80)
    
    # Find output file
    output_file = find_most_recent_output(user_name)
    if not output_file:
        print(f"❌ No output file found for {user_name}")
        return
    
    # Find input files
    input_files = find_input_files(user_name)
    if not input_files:
        print(f"❌ No input files found for {user_name}")
        return
    
    print(f"\n📄 Output: {output_file.name}")
    print(f"📄 Input files: {len(input_files)}")
    for f in input_files:
        print(f"   - {f.name}")
    
    print("\n" + "-" * 80)
    
    # Extract hyperlinks from all input files
    print("\n📥 EXTRACTING INPUT HYPERLINKS...")
    input_hyperlinks = []
    for input_file in input_files:
        doc = Document(str(input_file))
        links = extract_hyperlinks(doc, source_file=input_file.name)
        input_hyperlinks.extend(links)
        print(f"   {input_file.name}: {len(links)} hyperlinks")
    
    print(f"\n   Total input hyperlinks: {len(input_hyperlinks)}")
    
    # Extract hyperlinks from output
    print("\n📤 EXTRACTING OUTPUT HYPERLINKS...")
    output_doc = Document(str(output_file))
    output_hyperlinks = extract_hyperlinks(output_doc, source_file="output")
    print(f"   Found {len(output_hyperlinks)} hyperlinks in output")
    
    # Compare - use coordinate-based matching for duplicates
    print("\n" + "=" * 80)
    print("COMPARISON RESULTS")
    print("=" * 80)
    
    # Create a copy of output list for matching
    output_remaining = output_hyperlinks.copy()
    matched_pairs = []
    unmatched_input = []
    
    # First pass: Match by coordinates + cell context
    # For multi-slot outputs, ignore table_idx (it changes during merge)
    # Focus on row_idx, cell_idx, and cell context for better matching
    for inp_link in input_hyperlinks:
        best_match = None
        best_score = 0
        
        for i, out_link in enumerate(output_remaining):
            # Check if same URL and text
            if inp_link['url'] == out_link['url'] and inp_link['text'] == out_link['text']:
                # Calculate match score
                score = 0
                
                # Position match (most important)
                if (inp_link['row_idx'] == out_link['row_idx'] and 
                    inp_link['cell_idx'] == out_link['cell_idx']):
                    score += 100
                
                # Cell context similarity (helps distinguish duplicates)
                inp_context = inp_link['cell_text'].lower()
                out_context = out_link['cell_text'].lower()
                
                # Check if contexts are similar (simple substring check)
                if inp_context and out_context:
                    # Remove the hyperlink text itself from context for comparison
                    inp_clean = inp_context.replace(inp_link['text'].lower(), '')
                    out_clean = out_context.replace(out_link['text'].lower(), '')
                    
                    # Count matching words
                    inp_words = set(inp_clean.split())
                    out_words = set(out_clean.split())
                    if inp_words and out_words:
                        common_words = inp_words & out_words
                        context_similarity = len(common_words) / max(len(inp_words), len(out_words))
                        score += int(context_similarity * 50)
                
                if score > best_score:
                    best_score = score
                    best_match = (i, out_link, 'exact' if score >= 100 else 'context_match')
        
        if best_match and best_score >= 100:
            # Strong match (position + context)
            idx, out_link, match_type = best_match
            matched_pairs.append((inp_link, out_link, 'exact'))
            output_remaining.pop(idx)
        elif best_match and best_score >= 50:
            # Good context match even if position differs slightly
            idx, out_link, match_type = best_match
            matched_pairs.append((inp_link, out_link, 'context_match'))
            output_remaining.pop(idx)
        else:
            # Fallback: Match by URL+text only
            for i, out_link in enumerate(output_remaining):
                if inp_link['url'] == out_link['url'] and inp_link['text'] == out_link['text']:
                    matched_pairs.append((inp_link, out_link, 'moved'))
                    output_remaining.pop(i)
                    break
            else:
                unmatched_input.append(inp_link)
    
    if unmatched_input:
        print(f"\n❌ MISSING HYPERLINKS: {len(unmatched_input)}")
        for link in unmatched_input[:5]:
            print(f"   - {link['text'][:40]} → {link['url'][:50]}")
    else:
        print(f"\n✅ All {len(input_hyperlinks)} hyperlinks preserved!")
    
    if output_remaining:
        print(f"\n⚠️  EXTRA HYPERLINKS in output: {len(output_remaining)}")
        for link in output_remaining[:3]:
            print(f"   + {link['text'][:40]}")
    
    # Check placement accuracy
    print("\n" + "=" * 80)
    print("PLACEMENT ACCURACY CHECK")
    print("=" * 80)
    
    exact_position = sum(1 for _, _, match_type in matched_pairs if match_type == 'exact')
    context_match = sum(1 for _, _, match_type in matched_pairs if match_type == 'context_match')
    moved_position = sum(1 for _, _, match_type in matched_pairs if match_type == 'moved')
    
    print(f"\n📊 Position Accuracy:")
    print(f"   ✅ Exact position: {exact_position}/{len(input_hyperlinks)} ({exact_position/len(input_hyperlinks)*100:.1f}%)")
    if context_match > 0:
        print(f"   ✅ Context match: {context_match}/{len(input_hyperlinks)} ({context_match/len(input_hyperlinks)*100:.1f}%)")
    print(f"   ⚠️  Moved position: {moved_position}/{len(input_hyperlinks)} ({moved_position/len(input_hyperlinks)*100:.1f}%)")
    print(f"   ❌ Missing: {len(unmatched_input)}/{len(input_hyperlinks)}")
    
    # Show details for moved links (first 5)
    if moved_position > 0:
        print(f"\n📋 Moved hyperlinks (first 5):")
        count = 0
        for inp_link, out_link, match_type in matched_pairs:
            if match_type == 'moved' and count < 5:
                print(f"   ⚠️  {inp_link['text'][:30]}")
                print(f"      Input:  T{inp_link['table_idx']}R{inp_link['row_idx']}C{inp_link['cell_idx']}")
                print(f"      Output: T{out_link['table_idx']}R{out_link['row_idx']}C{out_link['cell_idx']}")
                count += 1
    
    # Show exact matches (first 5)
    if exact_position > 0:
        print(f"\n📋 Exact position matches (first 5):")
        count = 0
        for inp_link, out_link, match_type in matched_pairs:
            if match_type == 'exact' and count < 5:
                print(f"   ✅ {inp_link['text'][:40]}")
                print(f"      Position: T{inp_link['table_idx']}R{inp_link['row_idx']}C{inp_link['cell_idx']}")
                count += 1
    
    # Check for text deletion using matched pairs
    print("\n" + "=" * 80)
    print("TEXT PRESERVATION CHECK")
    print("=" * 80)
    
    text_issues = []
    for inp_link, out_link, match_type in matched_pairs:
        # Check if cell has text (not just the hyperlink)
        input_has_text = len(inp_link['cell_text']) > len(inp_link['text']) + 10
        output_has_text = len(out_link['cell_text']) > len(out_link['text']) + 10
        
        if input_has_text and not output_has_text:
            text_issues.append({
                'link': inp_link['text'][:30],
                'input_cell': inp_link['cell_text'][:50],
                'output_cell': out_link['cell_text'][:50],
                'position': f"T{inp_link['table_idx']}R{inp_link['row_idx']}C{inp_link['cell_idx']}"
            })
    
    if text_issues:
        print(f"\n⚠️  POTENTIAL TEXT DELETION: {len(text_issues)} cells")
        for issue in text_issues[:3]:
            print(f"\n   Link: {issue['link']}")
            print(f"   Position: {issue['position']}")
            print(f"   Input cell:  {issue['input_cell']}")
            print(f"   Output cell: {issue['output_cell']}")
    else:
        print(f"\n✅ No obvious text deletion detected")
    
    # Final summary
    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    print(f"\n✅ Hyperlinks preserved: {len(matched_pairs)}/{len(input_hyperlinks)}")
    print(f"✅ Exact position: {exact_position}/{len(input_hyperlinks)} ({exact_position/len(input_hyperlinks)*100:.1f}%)")
    if context_match > 0:
        print(f"✅ Context match: {context_match}/{len(input_hyperlinks)} ({context_match/len(input_hyperlinks)*100:.1f}%)")
    print(f"⚠️  Moved position: {moved_position}/{len(input_hyperlinks)} ({moved_position/len(input_hyperlinks)*100:.1f}%)")
    print(f"❌ Missing: {len(unmatched_input)}/{len(input_hyperlinks)}")
    print(f"⚠️  Text issues: {len(text_issues)}")
    
    total_good_matches = exact_position + context_match
    
    if len(matched_pairs) == len(input_hyperlinks) and total_good_matches == len(input_hyperlinks) and len(text_issues) == 0:
        print(f"\n🎉 PERFECT! All hyperlinks preserved in correct positions with text intact!")
    elif len(matched_pairs) == len(input_hyperlinks) and len(text_issues) == 0:
        print(f"\n✅ GOOD! All hyperlinks preserved with text intact")
        if moved_position > 0:
            print(f"   Note: {moved_position} hyperlinks in different positions (may be intentional)")
    elif len(matched_pairs) == len(input_hyperlinks):
        print(f"\n⚠️  All hyperlinks preserved but check text preservation")
    else:
        print(f"\n❌ Some hyperlinks missing")
    
    print("=" * 80)

# Main execution
print("=" * 80)
print("HYPERLINK PLACEMENT VERIFICATION - AUTO-DETECT")
print("=" * 80)

# Verify both users
verify_user("Daniela Silva")
verify_user("Wilson Rodrigues")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)
