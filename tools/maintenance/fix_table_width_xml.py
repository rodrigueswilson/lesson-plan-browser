"""
Fix table alignment by properly setting tblW in XML.
The issue: Column widths are correct but tblW XML property has wrong values.
"""

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil
from datetime import datetime

# File paths
TEMPLATE_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template.docx'
BACKUP_PATH = r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template_BACKUP_FINAL_' + datetime.now().strftime('%Y%m%d_%H%M%S') + '.docx'

# Target width for 0.5-inch margins
TARGET_WIDTH_INCHES = 7.5
TARGET_WIDTH_EMUS = 6858000
TARGET_WIDTH_DXA = 10800  # 7.5 inches * 1440 DXA per inch

def set_table_width_xml(table, width_dxa):
    """
    Set table width (tblW) in XML to exact DXA value.
    
    Args:
        table: python-docx Table object
        width_dxa: Width in DXA (twentieths of a point)
    """
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing tblW if present
    tblW_list = tblPr.findall(qn('w:tblW'))
    for tblW in tblW_list:
        tblPr.remove(tblW)
    
    # Add new tblW with exact value
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_table_indent(table, indent_dxa=0):
    """Set table indent to 0."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblInd_list = tblPr.findall(qn('w:tblInd'))
    for tblInd in tblInd_list:
        tblPr.remove(tblInd)
    
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(indent_dxa))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def set_table_alignment(table, alignment='left'):
    """Set table alignment."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    jc_list = tblPr.findall(qn('w:jc'))
    for jc in jc_list:
        tblPr.remove(jc)
    
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    tblPr.append(jc)

def set_fixed_layout(table):
    """Set table to fixed layout."""
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    layout_list = tblPr.findall(qn('w:tblLayout'))
    for layout in layout_list:
        tblPr.remove(layout)
    
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

print("=" * 80)
print("FIXING TABLE WIDTH XML PROPERTIES")
print("=" * 80)

# Create backup
print(f"\n1. Creating backup...")
shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)
print(f"   Backup saved to: {BACKUP_PATH}")

# Load document
print(f"\n2. Loading template...")
doc = Document(TEMPLATE_PATH)
print(f"   Found {len(doc.tables)} table(s)")

# Fix each table
print(f"\n3. Setting exact table widths in XML...")
for i, table in enumerate(doc.tables):
    print(f"\n   Table {i+1}:")
    
    # Set table width XML property to exact DXA value
    set_table_width_xml(table, TARGET_WIDTH_DXA)
    print(f"     ✓ tblW set to {TARGET_WIDTH_DXA} DXA ({TARGET_WIDTH_INCHES} inches)")
    
    # Set zero indent
    set_table_indent(table, indent_dxa=0)
    print(f"     ✓ tblInd set to 0 DXA")
    
    # Set left alignment
    set_table_alignment(table, alignment='left')
    print(f"     ✓ Alignment set to LEFT")
    
    # Set fixed layout
    set_fixed_layout(table)
    print(f"     ✓ Layout set to FIXED")
    
    # Set column widths to match total
    col_width = TARGET_WIDTH_EMUS // len(table.columns)
    remainder = TARGET_WIDTH_EMUS % len(table.columns)
    
    for j, column in enumerate(table.columns):
        if j == len(table.columns) - 1:
            column.width = col_width + remainder
        else:
            column.width = col_width
    
    # Verify total
    total = sum(col.width for col in table.columns)
    print(f"     ✓ Column widths: {total} EMUs = {total/914400:.4f} inches")

# Save document
print(f"\n4. Saving fixed template...")
doc.save(TEMPLATE_PATH)
print(f"   Saved to: {TEMPLATE_PATH}")

print("\n" + "=" * 80)
print("COMPLETE!")
print("=" * 80)
print(f"\nAll tables now have:")
print(f"  - tblW (XML): {TARGET_WIDTH_DXA} DXA = {TARGET_WIDTH_INCHES} inches")
print(f"  - Column widths: {TARGET_WIDTH_EMUS} EMUs = {TARGET_WIDTH_INCHES} inches")
print(f"  - tblInd: 0 (no indent)")
print(f"  - Alignment: LEFT")
print(f"  - Layout: FIXED")
print(f"\nTables should now be perfectly aligned on both left AND right edges!")
print(f"\nBackup: {BACKUP_PATH}")
print("=" * 80)
