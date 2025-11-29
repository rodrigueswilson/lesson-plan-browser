"""Verify hyperlinks are in correct positions with proper context."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import json

# Get the most recent output file
folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')
output_files = sorted(
    [f for f in folder.glob('Daniela_Silva_Lesson_plan_*.docx')],
    key=lambda f: f.stat().st_mtime,
    reverse=True
)

if not output_files:
    print("❌ No output files found!")
    exit(1)

output_file = output_files[0]
input_file = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43') / "Morais 10_20_25 - 10_24_25.docx"

print("=" * 80)
print("HYPERLINK PLACEMENT VERIFICATION")
print("=" * 80)
print(f"\n📄 Input:  {input_file.name}")
print(f"📄 Output: {output_file.name}")
print("=" * 80)

# Load both documents
input_doc = Document(str(input_file))
output_doc = Document(str(output_file))

# Extract hyperlinks from input
print("\n📥 EXTRACTING INPUT HYPERLINKS...")
input_hyperlinks = []

for table_idx, table in enumerate(input_doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for hyperlink in para._element.xpath('.//w:hyperlink'):
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id and r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        
                        # Get cell text context
                        cell_text = cell.text.strip()
                        
                        input_hyperlinks.append({
                            'text': text,
                            'url': url,
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'cell_text': cell_text[:100],  # First 100 chars
                        })

print(f"   Found {len(input_hyperlinks)} hyperlinks in input")

# Extract hyperlinks from output
print("\n📤 EXTRACTING OUTPUT HYPERLINKS...")
output_hyperlinks = []

for table_idx, table in enumerate(output_doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for hyperlink in para._element.xpath('.//w:hyperlink'):
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id and r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        
                        # Get cell text context
                        cell_text = cell.text.strip()
                        
                        output_hyperlinks.append({
                            'text': text,
                            'url': url,
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx,
                            'cell_text': cell_text[:100],
                        })

print(f"   Found {len(output_hyperlinks)} hyperlinks in output")

# Compare - use coordinate-based matching for duplicates
print("\n" + "=" * 80)
print("COMPARISON RESULTS")
print("=" * 80)

# Create a copy of output list for matching
output_remaining = output_hyperlinks.copy()
matched_pairs = []
unmatched_input = []

# First pass: Match by coordinates (exact position)
for inp_link in input_hyperlinks:
    best_match = None
    for i, out_link in enumerate(output_remaining):
        # Check if same URL and text
        if inp_link['url'] == out_link['url'] and inp_link['text'] == out_link['text']:
            # Check if same position
            if (inp_link['table_idx'] == out_link['table_idx'] and 
                inp_link['row_idx'] == out_link['row_idx'] and 
                inp_link['cell_idx'] == out_link['cell_idx']):
                best_match = (i, out_link, 'exact')
                break
    
    if best_match:
        idx, out_link, match_type = best_match
        matched_pairs.append((inp_link, out_link, match_type))
        output_remaining.pop(idx)
    else:
        # Second pass: Match by URL+text only (for moved links)
        for i, out_link in enumerate(output_remaining):
            if inp_link['url'] == out_link['url'] and inp_link['text'] == out_link['text']:
                matched_pairs.append((inp_link, out_link, 'moved'))
                output_remaining.pop(i)
                break
        else:
            unmatched_input.append(inp_link)

if unmatched_input:
    print(f"\n❌ MISSING HYPERLINKS: {len(unmatched_input)}")
    for link in unmatched_input:
        print(f"   - {link['text'][:40]} → {link['url'][:50]}")
else:
    print(f"\n✅ All {len(input_hyperlinks)} hyperlinks preserved!")

if output_remaining:
    print(f"\n⚠️  EXTRA HYPERLINKS in output: {len(output_remaining)}")
    for link in output_remaining[:3]:
        print(f"   + {link['text'][:40]}")

# Check placement accuracy
print("\n" + "=" * 80)
print("PLACEMENT ACCURACY CHECK")
print("=" * 80)

exact_position = sum(1 for _, _, match_type in matched_pairs if match_type == 'exact')
moved_position = sum(1 for _, _, match_type in matched_pairs if match_type == 'moved')

print(f"\n📊 Position Accuracy:")
print(f"   ✅ Exact position: {exact_position}/{len(input_hyperlinks)} ({exact_position/len(input_hyperlinks)*100:.1f}%)")
print(f"   ⚠️  Moved position: {moved_position}/{len(input_hyperlinks)} ({moved_position/len(input_hyperlinks)*100:.1f}%)")
print(f"   ❌ Missing: {len(unmatched_input)}/{len(input_hyperlinks)}")

# Show details for moved links
if moved_position > 0:
    print(f"\n📋 Moved hyperlinks (first 5):")
    count = 0
    for inp_link, out_link, match_type in matched_pairs:
        if match_type == 'moved' and count < 5:
            print(f"   ⚠️  {inp_link['text'][:30]}")
            print(f"      Input:  T{inp_link['table_idx']}R{inp_link['row_idx']}C{inp_link['cell_idx']}")
            print(f"      Output: T{out_link['table_idx']}R{out_link['row_idx']}C{out_link['cell_idx']}")
            count += 1

# Show exact matches (first 5)
if exact_position > 0:
    print(f"\n📋 Exact position matches (first 5):")
    count = 0
    for inp_link, out_link, match_type in matched_pairs:
        if match_type == 'exact' and count < 5:
            print(f"   ✅ {inp_link['text'][:40]}")
            print(f"      Position: T{inp_link['table_idx']}R{inp_link['row_idx']}C{inp_link['cell_idx']}")
            count += 1

# Check for text deletion using matched pairs
print("\n" + "=" * 80)
print("TEXT PRESERVATION CHECK")
print("=" * 80)

text_issues = []
for inp_link, out_link, match_type in matched_pairs:
    # Check if cell has text (not just the hyperlink)
    input_has_text = len(inp_link['cell_text']) > len(inp_link['text']) + 10  # Allow some whitespace
    output_has_text = len(out_link['cell_text']) > len(out_link['text']) + 10
    
    if input_has_text and not output_has_text:
        text_issues.append({
            'link': inp_link['text'][:30],
            'input_cell': inp_link['cell_text'][:50],
            'output_cell': out_link['cell_text'][:50],
            'position': f"T{inp_link['table_idx']}R{inp_link['row_idx']}C{inp_link['cell_idx']}"
        })

if text_issues:
    print(f"\n⚠️  POTENTIAL TEXT DELETION: {len(text_issues)} cells")
    for issue in text_issues[:3]:
        print(f"\n   Link: {issue['link']}")
        print(f"   Position: {issue['position']}")
        print(f"   Input cell:  {issue['input_cell']}")
        print(f"   Output cell: {issue['output_cell']}")
else:
    print(f"\n✅ No obvious text deletion detected")

# Final summary
print("\n" + "=" * 80)
print("SUMMARY")
print("=" * 80)
print(f"\n✅ Hyperlinks preserved: {len(matched_pairs)}/{len(input_hyperlinks)}")
print(f"✅ Exact position: {exact_position}/{len(input_hyperlinks)} ({exact_position/len(input_hyperlinks)*100:.1f}%)")
print(f"⚠️  Moved position: {moved_position}/{len(input_hyperlinks)} ({moved_position/len(input_hyperlinks)*100:.1f}%)")
print(f"❌ Missing: {len(unmatched_input)}/{len(input_hyperlinks)}")
print(f"⚠️  Text issues: {len(text_issues)}")

if len(matched_pairs) == len(input_hyperlinks) and exact_position == len(input_hyperlinks) and len(text_issues) == 0:
    print(f"\n🎉 PERFECT! All hyperlinks preserved in exact positions with text intact!")
elif len(matched_pairs) == len(input_hyperlinks) and len(text_issues) == 0:
    print(f"\n✅ GOOD! All hyperlinks preserved with text intact")
    if moved_position > 0:
        print(f"   Note: {moved_position} hyperlinks in different positions (may be intentional)")
elif len(matched_pairs) == len(input_hyperlinks):
    print(f"\n⚠️  All hyperlinks preserved but check text preservation")
else:
    print(f"\n❌ Some hyperlinks missing")

print("=" * 80)
