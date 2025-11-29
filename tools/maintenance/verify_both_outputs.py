"""Verify the most recent outputs for both users."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

print("=" * 80)
print("VERIFICATION OF MOST RECENT LESSON PLANS")
print("=" * 80)

# Daniela's folder
daniela_folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')
daniela_files = sorted(
    [f for f in daniela_folder.glob('Daniela_Silva_*.docx')],
    key=lambda f: f.stat().st_mtime,
    reverse=True
)

# Wilson's folder
wilson_folder = Path(r'F:\rodri\Documents\OneDrive\AS\Wilson LP\25 W43')
wilson_files = sorted(
    [f for f in wilson_folder.glob('Wilson_*.docx')],
    key=lambda f: f.stat().st_mtime,
    reverse=True
)

def verify_file(file_path):
    """Verify a single output file."""
    print(f"\n📄 File: {file_path.name}")
    print("-" * 80)
    
    try:
        doc = Document(str(file_path))
        
        # Count tables
        table_count = len(doc.tables)
        print(f"   Tables: {table_count}")
        
        # Count hyperlinks
        hyperlink_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        hyperlinks = para._element.xpath('.//w:hyperlink')
                        hyperlink_count += len(hyperlinks)
        
        print(f"   Hyperlinks: {hyperlink_count}")
        
        # Check for "Required Signatures"
        signature_count = 0
        for table in doc.tables:
            table_text = "\n".join([cell.text for row in table.rows for cell in row.cells])
            if "Required Signatures" in table_text:
                signature_count += 1
        
        print(f"   Signature tables: {signature_count}")
        
        # Check table widths
        table_widths = []
        for i, table in enumerate(doc.tables):
            if hasattr(table, 'width') and table.width:
                width_inches = table.width.inches
                table_widths.append(width_inches)
        
        if table_widths:
            min_width = min(table_widths)
            max_width = max(table_widths)
            print(f"   Table widths: {min_width:.2f}\" to {max_width:.2f}\"")
            if abs(max_width - min_width) < 0.1:
                print(f"   ✅ All tables same width")
            else:
                print(f"   ⚠️  Table widths vary")
        
        # Check for "Slot 1:" in content
        slot_mentions = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Slot 1:" in cell.text:
                        slot_mentions += 1
        
        if slot_mentions > 2:  # Allow for header row
            print(f"   ⚠️  'Slot 1:' appears {slot_mentions} times (may be excessive)")
        else:
            print(f"   ✅ No excessive 'Slot 1:' repetition")
        
        # Check for markdown artifacts
        markdown_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if '**' in text and not any(run.bold for para in cell.paragraphs for run in para.runs):
                        markdown_count += 1
        
        if markdown_count > 0:
            print(f"   ⚠️  Found {markdown_count} cells with unrendered markdown")
        else:
            print(f"   ✅ Markdown properly rendered")
        
        # Summary
        print(f"\n   Summary:")
        if signature_count == 1:
            print(f"   ✅ Exactly 1 signature table")
        else:
            print(f"   ❌ {signature_count} signature tables (should be 1)")
        
        if hyperlink_count > 0:
            print(f"   ✅ {hyperlink_count} hyperlinks preserved")
        
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        return False

# Verify Daniela's file
if daniela_files:
    print("\n" + "=" * 80)
    print("DANIELA SILVA")
    print("=" * 80)
    verify_file(daniela_files[0])
else:
    print("\n❌ No Daniela output files found")

# Verify Wilson's file
if wilson_files:
    print("\n" + "=" * 80)
    print("WILSON")
    print("=" * 80)
    verify_file(wilson_files[0])
else:
    print("\n❌ No Wilson output files found")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)
