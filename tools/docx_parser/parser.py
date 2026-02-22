"""
DOCX Parser: main DOCXParser class delegating to structure, no_school, table_extraction, content_sections, slot_extraction, images_metadata.
"""

import hashlib
import re
import string
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml.ns import qn

from backend.performance_tracker import get_tracker
from backend.telemetry import logger

from . import content_sections
from . import images_metadata
from . import no_school
from . import slot_extraction
from . import table_extraction
from .structure import validate_slot_structure


class DOCXParser:
    """Parser for primary teacher DOCX lesson plan files."""

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        self.doc = Document(str(self.file_path))
        self.paragraphs = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        self.tables = table_extraction.extract_tables(self.doc)

    def get_full_text(self, exclude_signatures: bool = True, strip_urls: bool = False) -> str:
        paragraphs = self.paragraphs
        if exclude_signatures:
            paragraphs = [p for p in paragraphs if not table_extraction.is_signature_content(p)]
        if strip_urls:
            url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
            paragraphs = [re.sub(url_pattern, "", p) for p in paragraphs]
        return "\n".join(paragraphs)

    def _get_no_school_patterns(self) -> list:
        return no_school.get_no_school_patterns()

    def is_no_school_day(self) -> bool:
        patterns = self._get_no_school_patterns()
        full_text = self.get_full_text(exclude_signatures=True, strip_urls=False)
        for pattern in patterns:
            if re.search(pattern, full_text, re.IGNORECASE):
                logger.info("no_school_detected", extra={"pattern": pattern, "file": str(self.file_path.name)})
                return True
        return False

    def is_day_no_school(self, day_text: str) -> bool:
        return no_school.is_day_no_school(day_text)

    def find_subject_sections(self) -> Dict[str, Any]:
        return content_sections.find_subject_sections(self.doc, self.tables, self.get_full_text)

    def extract_by_heading(self, heading_text: str, case_sensitive: bool = False) -> List[str]:
        return content_sections.extract_by_heading(self.paragraphs, heading_text, case_sensitive)

    def extract_table_by_header(self, header_text: str) -> Optional[List[List[str]]]:
        return content_sections.extract_table_by_header(self.tables, header_text)

    def extract_table_lesson(self, table_index: int) -> Dict[str, Any]:
        return content_sections.extract_table_lesson(self.doc, table_index, table_extraction.get_cell_text_with_hyperlinks)

    def extract_subject_content(self, subject: str, strip_urls: bool = True) -> Dict[str, Any]:
        return content_sections.extract_subject_content(
            self.doc,
            self.tables,
            self.paragraphs,
            subject,
            self.get_full_text,
            self.is_day_no_school,
            self.extract_table_lesson,
            table_extraction.get_cell_text_with_hyperlinks,
            strip_urls=strip_urls,
        )

    def list_available_subjects(self) -> List[str]:
        return [s.lower() for s in content_sections.list_available_subjects(self.find_subject_sections())]

    def extract_week_info(self) -> Optional[str]:
        return images_metadata.extract_week_info_from_text(self.get_full_text())

    def extract_images(self, exclude_signatures: bool = True) -> List[Dict[str, Any]]:
        def find_image_context_fn(rel_id):
            return images_metadata.find_image_context(self.doc, rel_id, images_metadata.infer_section)
        return images_metadata.extract_images(
            self.doc,
            exclude_signatures=exclude_signatures,
            is_signature_content_fn=table_extraction.is_signature_content,
            find_image_context_fn=find_image_context_fn,
        )

    def extract_hyperlinks(self) -> List[Dict[str, str]]:
        return images_metadata.extract_hyperlinks(
            self.doc,
            images_metadata.get_context_snippet,
            images_metadata.infer_section,
            images_metadata.extract_day_from_header,
            images_metadata.normalize_day_hint,
        )

    def get_metadata(self) -> Dict[str, Any]:
        return images_metadata.get_metadata(
            self.doc,
            self.paragraphs,
            self.tables,
            self.list_available_subjects,
            self.extract_week_info,
        )

    def find_slot_by_subject(
        self, subject: str, teacher_name: str = None, homeroom: str = None, grade: str = None
    ) -> int:
        return slot_extraction.find_slot_by_subject(
            self.doc, subject, teacher_name, homeroom, grade, validate_slot_structure
        )

    def _get_context_snippet(self, paragraph, link_text: str, window: int = None) -> str:
        return images_metadata.get_context_snippet(paragraph, link_text, window)

    def _infer_section(self, text: str) -> Optional[str]:
        return images_metadata.infer_section(text)

    def _detect_day_from_table(self, table) -> Optional[str]:
        return images_metadata.detect_day_from_table(table)

    def _extract_day_from_header(self, col_header: str) -> Optional[str]:
        return images_metadata.extract_day_from_header(col_header)

    def _normalize_day_hint(self, day_hint: str) -> Optional[str]:
        return images_metadata.normalize_day_hint(day_hint)

    def _find_image_context(self, rel_id: str) -> Dict[str, Any]:
        return images_metadata.find_image_context(self.doc, rel_id, images_metadata.infer_section)

    def extract_hyperlinks_for_slot(
        self, slot_number: int = None, subject: str = None, teacher_name: str = None
    ) -> List[Dict[str, str]]:
        if subject:
            slot_number = self.find_slot_by_subject(subject, teacher_name)
            logger.debug("hyperlinks_extraction_via_subject", extra={"subject": subject, "teacher": teacher_name, "resolved_slot": slot_number})
        elif slot_number is None:
            raise ValueError("Must provide either slot_number or subject for slot-aware extraction")

        table_start, table_end = validate_slot_structure(self.doc, slot_number)
        logger.info("extracting_slot_hyperlinks", extra={"slot_number": slot_number, "table_start": table_start, "table_end": table_end})

        hyperlinks = []
        seen_links = set()
        for table_idx in range(table_start, table_end + 1):
            table = self.doc.tables[table_idx]
            col_headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            for row_idx, row in enumerate(table.rows):
                row_label = row.cells[0].text.strip() if row.cells else ""
                for cell_idx, cell in enumerate(row.cells):
                    col_header = col_headers[cell_idx] if cell_idx < len(col_headers) else ""
                    day_hint = self._extract_day_from_header(col_header)
                    if day_hint:
                        day_hint = self._normalize_day_hint(day_hint)
                    elif cell_idx > 0:
                        col_header_lower = col_header.lower() if col_header else ""
                        is_metadata_column = any(p in col_header_lower for p in ["grade:", "homeroom:", "subject:", "week of:", "week:"])
                        if not is_metadata_column:
                            logger.warning("hyperlink_extraction_missing_day_hint", extra={"slot": slot_number, "table_idx": table_idx, "row": row_idx, "col": cell_idx, "header": col_header[:50] if col_header else "empty"})
                    for paragraph in cell.paragraphs:
                        for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                            try:
                                r_id = hyperlink.get(qn("r:id"))
                                if r_id and r_id in paragraph.part.rels:
                                    url = paragraph.part.rels[r_id].target_ref
                                    text = "".join(node.text for node in hyperlink.xpath(".//w:t") if node.text)
                                    if text and url:
                                        text_clean = text.strip()
                                        url_clean = url.strip()
                                        link_key = (text_clean, url_clean, table_idx, row_idx, cell_idx)
                                        if link_key in seen_links:
                                            continue
                                        seen_links.add(link_key)
                                        hyperlinks.append({
                                            "schema_version": "2.0",
                                            "text": text_clean,
                                            "url": url_clean,
                                            "context_snippet": self._get_context_snippet(paragraph, text),
                                            "context_type": "table",
                                            "section_hint": self._infer_section(row_label),
                                            "day_hint": day_hint,
                                            "table_idx": table_idx,
                                            "row_idx": row_idx,
                                            "cell_idx": cell_idx,
                                            "row_label": row_label,
                                            "col_header": col_header,
                                        })
                            except Exception as e:
                                logger.warning("hyperlink_extraction_failed", extra={"error": str(e), "table_idx": table_idx})
        logger.info("slot_hyperlinks_extracted", extra={"slot_number": slot_number, "hyperlink_count": len(hyperlinks)})
        return hyperlinks

    def extract_images_for_slot(
        self, slot_number: int = None, subject: str = None, teacher_name: str = None
    ) -> List[Dict]:
        if subject:
            slot_number = self.find_slot_by_subject(subject, teacher_name)
            logger.debug("images_extraction_via_subject", extra={"subject": subject, "teacher": teacher_name, "resolved_slot": slot_number})
        elif slot_number is None:
            raise ValueError("Must provide either slot_number or subject for slot-aware extraction")

        table_start, table_end = validate_slot_structure(self.doc, slot_number)
        all_images = self.extract_images()
        slot_images = []
        seen_images = set()
        for img in all_images:
            if img.get("table_idx") is not None and table_start <= img["table_idx"] <= table_end:
                data_hash = hashlib.md5(img["data"].encode()).hexdigest()
                img_key = (data_hash, img["table_idx"], img.get("row_idx", -1), img.get("cell_index", -1))
                if img_key in seen_images:
                    continue
                seen_images.add(img_key)
                slot_images.append(img)
        logger.info("slot_images_extracted", extra={"slot_number": slot_number, "total_images": len(all_images), "slot_images": len(slot_images), "table_start": table_start, "table_end": table_end})
        return slot_images

    def extract_subject_content_for_slot(
        self,
        slot_number: int,
        subject: str,
        teacher_name: str = None,
        strip_urls: bool = True,
    ) -> Dict[str, Any]:
        def normalize_text(text):
            if not text:
                return ""
            return text.translate(str.maketrans("", "", string.punctuation)).lower().strip()

        try:
            actual_slot_number = self.find_slot_by_subject(subject, teacher_name=teacher_name)
            if actual_slot_number != slot_number:
                try:
                    table_start, table_end = validate_slot_structure(self.doc, slot_number)
                    meta_table = self.doc.tables[table_start]
                    slot_subject = slot_teacher = slot_homeroom = None
                    for row in meta_table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            cell_lower = cell_text.lower()
                            if "subject:" in cell_lower:
                                slot_subject = cell_text.split(":", 1)[-1].strip()
                            if "name:" in cell_lower or "teacher:" in cell_lower:
                                slot_teacher = cell_text.split(":", 1)[-1].strip()
                            if "homeroom:" in cell_lower:
                                slot_homeroom = cell_text.split(":", 1)[-1].strip()
                    subject_normalized = normalize_text(subject)
                    slot_subject_normalized = normalize_text(slot_subject) if slot_subject else ""
                    subject_matches = subject_normalized in slot_subject_normalized or slot_subject_normalized in subject_normalized
                    teacher_matches = False
                    if teacher_name and slot_teacher:
                        teacher_matches = normalize_text(teacher_name) in normalize_text(slot_teacher) or normalize_text(slot_teacher) in normalize_text(teacher_name)
                    if subject_matches and (not teacher_name or teacher_matches):
                        logger.info("using_requested_slot_number", extra={"requested_slot": slot_number, "found_slot": actual_slot_number, "subject": subject, "teacher_name": teacher_name, "reason": "Requested slot contains subject and matches teacher (if provided)"})
                    else:
                        logger.warning("slot_number_mismatch", extra={"requested_slot": slot_number, "actual_slot": actual_slot_number, "subject": subject, "teacher_name": teacher_name, "message": f"Slot {slot_number} requested but subject '{subject}' found in slot {actual_slot_number}. Teacher disambiguation overrides."})
                        slot_number = actual_slot_number
                except ValueError:
                    logger.warning("slot_number_mismatch", extra={"requested_slot": slot_number, "actual_slot": actual_slot_number, "subject": subject, "teacher_name": teacher_name, "message": f"Slot {slot_number} requested but subject '{subject}' found in slot {actual_slot_number}"})
                    slot_number = actual_slot_number
        except ValueError as e:
            logger.warning("subject_slot_detection_failed", extra={"slot_number": slot_number, "subject": subject, "teacher_name": teacher_name, "error": str(e), "fallback": "using requested slot number"})

        try:
            table_start, table_end = validate_slot_structure(self.doc, slot_number)
        except ValueError:
            total_tables = len(self.doc.tables)
            if total_tables >= 2:
                table_start = (slot_number - 1) * 2
                table_end = table_start + 1
            else:
                raise

        meta_table = self.doc.tables[table_start]
        meta_data = {}
        for row in meta_table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                lower = text.lower()
                if "subject:" in lower:
                    meta_data["subject"] = text.split(":", 1)[-1].strip()
                elif "name:" in lower or "teacher:" in lower:
                    meta_data["primary_teacher_name"] = text.split(":", 1)[-1].strip()
                elif "homeroom:" in lower:
                    meta_data["homeroom"] = text.split(":", 1)[-1].strip()
                elif "grade:" in lower:
                    meta_data["grade"] = text.split(":", 1)[-1].strip()
                elif "week of:" in lower:
                    meta_data["week_of"] = text.split(":", 1)[-1].strip()

        daily_table_idx = table_end
        logger.info("extracting_slot_content", extra={"slot_number": slot_number, "subject": subject, "table_idx": daily_table_idx})

        is_single_lesson = False
        if daily_table_idx < len(self.doc.tables):
            daily_table = self.doc.tables[daily_table_idx]
            if daily_table.rows:
                first_row_text = " ".join(cell.text.strip() for cell in daily_table.rows[0].cells)
                first_row_clean = first_row_text.translate(str.maketrans("", "", string.punctuation)).upper()
                metadata_indicators = ["UNIT", "LESSON", "MODULE", "SUBJECT", "OBJECTIVE"]
                weekdays = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY"]
                has_metadata_indicator = any(indicator in first_row_clean for indicator in metadata_indicators)
                has_weekday = any(day in first_row_clean for day in weekdays)
                if has_metadata_indicator and not has_weekday:
                    is_single_lesson = True

        if is_single_lesson:
            full_text_parts = []
            meta_table = self.doc.tables[table_start]
            for row in meta_table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_parts.append(row_text)
            lesson_table = self.doc.tables[table_end]
            for row in lesson_table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_parts.append(row_text)
            table_content = {"All Days": {"Lesson Content": "\n".join(full_text_parts)}}
            no_school_days = []
            logger.info("single_lesson_content_extracted", extra={"slot_number": slot_number, "subject": subject, "content_length": len("\n".join(full_text_parts))})
        else:
            table_content = self.extract_table_lesson(daily_table_idx)
            full_text_parts = []
            no_school_days = []
            for day, day_content in table_content.items():
                day_text = " ".join(day_content.values())
                if self.is_day_no_school(day_text):
                    no_school_days.append(day)
                    full_text_parts.append(f"\n{day.upper()}")
                    full_text_parts.append("Unit/Lesson: No School")
                    for label, text in day_content.items():
                        label_lower = label.lower()
                        if "unit" not in label_lower and "lesson" not in label_lower:
                            full_text_parts.append(f"{label} {text}")
                        elif text.strip() and text.strip().lower() != "no school":
                            full_text_parts.append(f"{label} {text}")
                else:
                    full_text_parts.append(f"\n{day.upper()}")
                    for label, text in day_content.items():
                        full_text_parts.append(f"{label} {text}")
            logger.info("slot_content_extracted", extra={"slot_number": slot_number, "subject": subject, "content_length": len("\n".join(full_text_parts)), "no_school_days": len(no_school_days)})

        return {
            "subject": subject,
            "full_text": "\n".join(full_text_parts),
            "table_content": table_content,
            "no_school_days": no_school_days,
            "found": True,
            "format": "table",
            "slot_number": slot_number,
            "table_idx": daily_table_idx,
            "metadata": meta_data,
        }


def parse_docx(
    file_path: str, subject: Optional[str] = None, plan_id: Optional[str] = None
) -> Dict[str, Any]:
    """Convenience function to parse DOCX file."""
    tracker = get_tracker()
    if plan_id:
        with tracker.track_operation(plan_id, "parse_locate_file"):
            file_path_obj = Path(file_path)
            if not file_path_obj.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")

    if plan_id:
        with tracker.track_operation(plan_id, "parse_open_docx"):
            parser = DOCXParser(file_path)
    else:
        parser = DOCXParser(file_path)

    if subject:
        if plan_id:
            with tracker.track_operation(plan_id, "parse_extract_subject"):
                result = parser.extract_subject_content(subject)
        else:
            result = parser.extract_subject_content(subject)
        return result

    result = {}
    if plan_id:
        with tracker.track_operation(plan_id, "parse_extract_text"):
            result["full_text"] = parser.get_full_text()
        with tracker.track_operation(plan_id, "parse_extract_metadata"):
            result["metadata"] = parser.get_metadata()
        with tracker.track_operation(plan_id, "parse_list_subjects"):
            result["available_subjects"] = parser.list_available_subjects()
        result["tables"] = parser.tables
    else:
        result = {
            "full_text": parser.get_full_text(),
            "metadata": parser.get_metadata(),
            "available_subjects": parser.list_available_subjects(),
            "tables": parser.tables,
        }
    return result
