"""
Image and hyperlink extraction, week info, and metadata for DOCX parsing.
"""

import base64
import re
from typing import Any, Dict, List, Optional

from docx.oxml.ns import qn


def get_context_snippet(paragraph, link_text: str, window: Optional[int] = None) -> str:
    """Extract text context around element for semantic matching."""
    if window is None:
        try:
            from backend.config import settings
            window = settings.MEDIA_CONTEXT_WINDOW_CHARS
        except Exception:
            window = 100

    full_text = paragraph.text
    if link_text in full_text:
        pos = full_text.index(link_text)
        start = max(0, pos - window // 2)
        end = min(len(full_text), pos + len(link_text) + window // 2)
        return full_text[start:end].strip()
    return full_text[:window].strip()


def infer_section(text: str) -> Optional[str]:
    """Infer lesson plan section from text content."""
    text_lower = text.lower()
    if (
        re.search(r"\bunit\b", text_lower)
        or "unit/lesson" in text_lower
        or "lesson #" in text_lower
        or re.search(r"\bmodule\b", text_lower)
    ):
        return "unit_lesson"
    elif any(kw in text_lower for kw in ["objective", "goal", "swbat", "students will be able"]):
        return "objective"
    elif any(kw in text_lower for kw in ["warm-up", "hook", "anticipatory", "do now"]):
        return "anticipatory_set"
    elif any(kw in text_lower for kw in ["instruction", "activity", "procedure", "tailored"]):
        return "instruction"
    elif any(kw in text_lower for kw in ["misconception", "common error"]):
        return "misconceptions"
    elif any(kw in text_lower for kw in ["assessment", "check", "evaluate", "exit ticket"]):
        return "assessment"
    elif any(kw in text_lower for kw in ["homework", "assignment", "practice"]):
        return "homework"
    return None


def detect_day_from_table(table) -> Optional[str]:
    """Try to detect day from table header row."""
    if not table.rows:
        return None
    header_text = " ".join(cell.text.lower() for cell in table.rows[0].cells)
    days = ["monday", "tuesday", "wednesday", "thursday", "friday"]
    for day in days:
        if day in header_text:
            return day
    return None


def extract_day_from_header(col_header: str) -> Optional[str]:
    """Extract day name from column header with enhanced format support."""
    if not col_header:
        return None
    header_lower = col_header.lower().strip()
    days_full = ["monday", "tuesday", "wednesday", "thursday", "friday"]
    for day in days_full:
        if day in header_lower:
            return day
    day_abbrev_3 = {
        "mon": "monday", "tue": "tuesday", "wed": "wednesday",
        "thu": "thursday", "fri": "friday",
    }
    for abbrev, full in day_abbrev_3.items():
        if abbrev in header_lower:
            return full
    day_abbrev_1 = {"m": "monday", "t": "tuesday", "w": "wednesday", "r": "thursday", "f": "friday"}
    words = header_lower.split()
    if words and len(words[0]) == 1 and words[0] in day_abbrev_1:
        return day_abbrev_1[words[0]]
    return None


def normalize_day_hint(day_hint: str) -> Optional[str]:
    """Normalize day hint from compound formats to a single valid day name."""
    if not day_hint:
        return None
    day_hint_lower = day_hint.lower().strip()
    valid_days = ["monday", "tuesday", "wednesday", "thursday", "friday"]
    if day_hint_lower in valid_days:
        return day_hint_lower
    for day in valid_days:
        if day in day_hint_lower:
            return day
    day_abbrev_3 = {"mon": "monday", "tue": "tuesday", "wed": "wednesday", "thu": "thursday", "fri": "friday"}
    parts = re.split(r"[/\-,]|\band\b", day_hint_lower)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        for day in valid_days:
            if day in part:
                return day
        for abbrev, full in day_abbrev_3.items():
            if abbrev in part:
                return full
    return None


def find_image_context(doc, rel_id: str, infer_section_fn) -> Dict[str, Any]:
    """Find context for an image by locating it in document structure."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            try:
                if rel_id in str(run._element.xml):
                    return {
                        "context": paragraph.text[:100],
                        "type": "paragraph",
                        "section": infer_section_fn(paragraph.text),
                        "day": None,
                        "row_label": None,
                        "cell_index": None,
                        "table_idx": None,
                    }
            except Exception:
                pass

    for table_idx, table in enumerate(doc.tables):
        day_hint = detect_day_from_table(table)
        for row_idx, row in enumerate(table.rows):
            row_label = row.cells[0].text.strip() if row.cells else ""
            for cell_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        try:
                            if rel_id in str(run._element.xml):
                                section = infer_section_fn(row_label) if row_label else infer_section_fn(cell.text)
                                day_map = {1: "monday", 2: "tuesday", 3: "wednesday", 4: "thursday", 5: "friday"}
                                day_from_cell = day_map.get(cell_idx, day_hint)
                                return {
                                    "context": cell.text[:100],
                                    "type": "table_cell",
                                    "section": section,
                                    "day": day_from_cell,
                                    "row_label": row_label,
                                    "row_idx": row_idx,
                                    "cell_index": cell_idx,
                                    "table_idx": table_idx,
                                }
                        except Exception:
                            pass

    return {
        "context": "",
        "type": "unknown",
        "section": None,
        "day": None,
        "row_label": None,
        "cell_index": None,
        "table_idx": None,
    }


def extract_week_info_from_text(full_text: str) -> Optional[str]:
    """Try to extract week information from full text."""
    week_patterns = [
        r"week\s+of\s+(\d{1,2}/\d{1,2})",
        r"week\s+(\d{1,2})",
        r"(\d{1,2}/\d{1,2}\s*-\s*\d{1,2}/\d{1,2})",
    ]
    for pattern in week_patterns:
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            return match.group(1)
    return None


def extract_images(
    doc,
    exclude_signatures: bool = True,
    is_signature_content_fn=None,
    find_image_context_fn=None,
) -> List[Dict[str, Any]]:
    """Extract all images from document with context for semantic anchoring."""
    from backend.telemetry import logger

    images = []
    try:
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref.lower():
                try:
                    context_info = find_image_context_fn(rel_id) if find_image_context_fn else {}
                    if exclude_signatures and is_signature_content_fn:
                        context = context_info.get("context", "")
                        if is_signature_content_fn(context):
                            continue
                    image_data = {
                        "data": base64.b64encode(rel.target_part.blob).decode("utf-8"),
                        "content_type": rel.target_part.content_type,
                        "filename": rel.target_ref.split("/")[-1],
                        "rel_id": rel_id,
                        "context_snippet": context_info.get("context", ""),
                        "context_type": context_info.get("type", "unknown"),
                        "section_hint": context_info.get("section"),
                        "day_hint": context_info.get("day"),
                        "row_label": context_info.get("row_label"),
                        "row_idx": context_info.get("row_idx"),
                        "cell_index": context_info.get("cell_index"),
                        "table_idx": context_info.get("table_idx"),
                    }
                    images.append(image_data)
                except Exception as e:
                    logger.warning("image_extraction_failed", extra={"rel_id": rel_id, "error": str(e)})
    except Exception as e:
        logger.warning("images_extraction_error", extra={"error": str(e)})
    return images


def extract_hyperlinks(
    doc,
    get_context_snippet_fn,
    infer_section_fn,
    extract_day_from_header_fn,
    normalize_day_hint_fn,
) -> List[Dict[str, str]]:
    """Extract all hyperlinks from document with context for semantic anchoring."""
    from backend.telemetry import logger

    hyperlinks = []
    try:
        for paragraph in doc.paragraphs:
            for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                try:
                    r_id = hyperlink.get(qn("r:id"))
                    if r_id and r_id in paragraph.part.rels:
                        url = paragraph.part.rels[r_id].target_ref
                        text = "".join(node.text for node in hyperlink.xpath(".//w:t") if node.text)
                        if text and url:
                            hyperlinks.append({
                                "schema_version": "2.0",
                                "text": text,
                                "url": url,
                                "context_snippet": get_context_snippet_fn(paragraph, text),
                                "context_type": "paragraph",
                                "section_hint": infer_section_fn(paragraph.text),
                                "day_hint": None,
                                "table_idx": None,
                                "row_idx": None,
                                "cell_idx": None,
                                "row_label": None,
                                "col_header": None,
                            })
                except Exception as e:
                    logger.warning("hyperlink_extraction_failed", extra={"error": str(e)})

        for table_idx, table in enumerate(doc.tables):
            col_headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            for row_idx, row in enumerate(table.rows):
                row_label = row.cells[0].text.strip() if row.cells else ""
                for cell_idx, cell in enumerate(row.cells):
                    col_header = col_headers[cell_idx] if cell_idx < len(col_headers) else ""
                    day_hint = extract_day_from_header_fn(col_header) if col_header else None
                    if day_hint:
                        day_hint = normalize_day_hint_fn(day_hint)
                    for paragraph in cell.paragraphs:
                        for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                            try:
                                r_id = hyperlink.get(qn("r:id"))
                                if r_id and r_id in paragraph.part.rels:
                                    url = paragraph.part.rels[r_id].target_ref
                                    text = "".join(node.text for node in hyperlink.xpath(".//w:t") if node.text)
                                    if text and url:
                                        hyperlinks.append({
                                            "schema_version": "2.0",
                                            "text": text,
                                            "url": url,
                                            "context_snippet": get_context_snippet_fn(paragraph, text),
                                            "context_type": "table_cell",
                                            "section_hint": infer_section_fn(cell.text),
                                            "day_hint": day_hint,
                                            "table_idx": table_idx,
                                            "row_idx": row_idx,
                                            "cell_idx": cell_idx,
                                            "row_label": row_label,
                                            "col_header": col_header,
                                        })
                            except Exception as e:
                                logger.warning("table_hyperlink_extraction_failed", extra={"error": str(e)})

        validated_hyperlinks = []
        for link in hyperlinks:
            if link.get("row_label"):
                link["row_label"] = link["row_label"].strip()
            if link.get("day_hint"):
                original = link["day_hint"]
                link["day_hint"] = link["day_hint"].lower().strip()
                normalized = normalize_day_hint_fn(link["day_hint"])
                if normalized:
                    link["day_hint"] = normalized
                else:
                    valid_days = ["monday", "tuesday", "wednesday", "thursday", "friday"]
                    if link["day_hint"] not in valid_days:
                        logger.warning("invalid_day_hint", extra={"day_hint": original, "link_text": link.get("text", "")[:50]})
                        link["day_hint"] = None
            validated_hyperlinks.append(link)
        return validated_hyperlinks
    except Exception as e:
        logger.warning("hyperlinks_extraction_error", extra={"error": str(e)})
        return hyperlinks


def get_metadata(
    doc,
    paragraphs: List[str],
    tables: List,
    list_available_subjects_fn,
    extract_week_info_fn,
) -> Dict[str, Any]:
    """Extract metadata from document."""
    core_props = doc.core_properties
    return {
        "title": core_props.title or "",
        "author": core_props.author or "",
        "created": core_props.created.isoformat() if core_props.created else None,
        "modified": core_props.modified.isoformat() if core_props.modified else None,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "available_subjects": list_available_subjects_fn(),
        "week_info": extract_week_info_fn(),
    }
