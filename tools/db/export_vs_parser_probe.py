"""
Compare lightweight structure signals between Google Docs JSON and DOCX export.

Usage:
  python tools/db/export_vs_parser_probe.py --doc-id <id> --label <slug> --out-dir docs/curriculum/research-notes/export-vs-parser
"""
from __future__ import annotations

import argparse
import json
import os
import tempfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[2]

import sys

_SCRAPER_DIR = _REPO_ROOT / "tools" / "scraper"
if str(_SCRAPER_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRAPER_DIR))

from docs_client import DocsClient


ANCHORS = [
    "learning intentions",
    "materials",
    "activity",
    "warm-up",
    "cool-down",
    "standards",
    "narrative",
    "purpose",
    "daily instructional task",
    "success criteria",
]


def _iter_json_text_runs(node: Any):
    if isinstance(node, dict):
        tr = node.get("textRun")
        if isinstance(tr, dict):
            content = tr.get("content")
            if isinstance(content, str):
                yield content
        for v in node.values():
            yield from _iter_json_text_runs(v)
    elif isinstance(node, list):
        for it in node:
            yield from _iter_json_text_runs(it)


def _count_json_tables(node: Any) -> int:
    if isinstance(node, dict):
        n = 1 if "table" in node else 0
        return n + sum(_count_json_tables(v) for v in node.values())
    if isinstance(node, list):
        return sum(_count_json_tables(it) for it in node)
    return 0


def _count_json_hyperlinks(node: Any) -> int:
    if isinstance(node, dict):
        n = 0
        ts = node.get("textStyle")
        if isinstance(ts, dict):
            link = ts.get("link")
            if isinstance(link, dict) and link.get("url"):
                n += 1
        return n + sum(_count_json_hyperlinks(v) for v in node.values())
    if isinstance(node, list):
        return sum(_count_json_hyperlinks(it) for it in node)
    return 0


def _docx_table_count_with_nesting(doc: Document) -> int:
    def cell_nested_tables(cell):
        n = len(cell.tables)
        for t in cell.tables:
            for r in t.rows:
                for c in r.cells:
                    n += cell_nested_tables(c)
        return n

    total = 0
    for t in doc.tables:
        total += 1
        for r in t.rows:
            for c in r.cells:
                total += cell_nested_tables(c)
    return total


def _docx_hyperlink_count(doc: Document) -> int:
    rels = doc.part.rels.values()
    return sum(1 for r in rels if "hyperlink" in r.reltype and getattr(r, "target_ref", None))


def _docx_text(doc: Document) -> str:
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def _anchor_hits(text: str) -> dict[str, int]:
    low = text.lower()
    return {a: low.count(a) for a in ANCHORS}


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--doc-id", required=True)
    ap.add_argument("--label", required=True)
    ap.add_argument("--out-dir", default=str(_REPO_ROOT / "docs" / "curriculum" / "research-notes" / "export-vs-parser"))
    args = ap.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    client = DocsClient()
    doc_json = client.get_document(args.doc_id)
    if not doc_json:
        raise RuntimeError(f"Failed to fetch docs JSON for {args.doc_id}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        docx_path = tmp.name
    try:
        ok = client.export_document(
            args.doc_id,
            docx_path,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        if not ok:
            raise RuntimeError(f"Failed to export DOCX for {args.doc_id}")

        doc = Document(docx_path)
        json_text = "".join(_iter_json_text_runs(doc_json))
        docx_text = _docx_text(doc)

        metrics = {
            "label": args.label,
            "doc_id": args.doc_id,
            "json_table_count": _count_json_tables(doc_json),
            "docx_table_count_with_nesting": _docx_table_count_with_nesting(doc),
            "json_hyperlink_count": _count_json_hyperlinks(doc_json),
            "docx_hyperlink_count": _docx_hyperlink_count(doc),
            "json_text_len": len(json_text),
            "docx_text_len": len(docx_text),
            "json_anchor_hits": _anchor_hits(json_text),
            "docx_anchor_hits": _anchor_hits(docx_text),
        }

        out_path = out_dir / f"{args.label}-{args.doc_id}-metrics.json"
        out_path.write_text(json.dumps(metrics, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        print(str(out_path))
        return 0
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


if __name__ == "__main__":
    raise SystemExit(main())
