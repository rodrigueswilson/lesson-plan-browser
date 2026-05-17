import docx

path = r'd:\LP\temp_u2\originals\Unit_2__Area_and_Multiplication.docx'
doc = docx.Document(path)

print(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables.")

# Print the first 20 paragraphs to see the preamble
for i, p in enumerate(doc.paragraphs[:20]):
    print(f"P{i}: {p.text[:100]} (Style: {p.style.name})")

# Print the first few tables and their sizes
for i, t in enumerate(doc.tables[:5]):
    print(f"Table {i}: {len(t.rows)} rows, {len(t.columns)} columns")
    if len(t.rows) > 0 and len(t.columns) > 0:
        print(f"  First cell: {t.cell(0,0).text[:100]}")
