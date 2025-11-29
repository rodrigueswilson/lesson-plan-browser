"""Find all Referenced Links sections and show their content."""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

print(f"Finding Referenced Links sections in: {output_file.name}\n")

doc = Document(output_file)

# Check each table
referenced_tables = []

for idx, table in enumerate(doc.tables):
    # Check all cells for "Referenced" text
    has_referenced = False
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if 'Referenced' in cell_text or 'Links:' in cell_text:
                has_referenced = True
                print(f"\n{'='*80}")
                print(f"TABLE {idx} - POTENTIAL FALLBACK SECTION")
                print(f"{'='*80}")
                print(f"Cell text: {cell_text[:200]}")
                
                # Count links in this table
                link_count = 0
                sample_links = []
                for r in table.rows:
                    for c in r.cells:
                        for para in c.paragraphs:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            link_count += len(hyperlinks)
                            for hyperlink in hyperlinks:
                                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                if len(sample_links) < 5:
                                    sample_links.append(text)
                
                print(f"\nLinks in this table: {link_count}")
                if sample_links:
                    print(f"Sample links:")
                    for link in sample_links:
                        print(f"  - {link[:80]}")
                
                referenced_tables.append((idx, link_count))
                break
        if has_referenced:
            break

print(f"\n\n{'='*80}")
print("SUMMARY")
print(f"{'='*80}\n")

if referenced_tables:
    print(f"Found {len(referenced_tables)} 'Referenced Links' sections:")
    total_fallback = 0
    for table_idx, link_count in referenced_tables:
        print(f"  Table {table_idx}: {link_count} links")
        total_fallback += link_count
    print(f"\nTotal fallback links: {total_fallback}")
else:
    print("No 'Referenced Links' sections found with that exact text.")
    print("\nLet me check table 9 (which has 77 links)...")
    
    table = doc.tables[9]
    print(f"\nTable 9 first few cells:")
    for row_idx in range(min(3, len(table.rows))):
        for cell_idx in range(min(3, len(table.rows[row_idx].cells))):
            cell_text = table.rows[row_idx].cells[cell_idx].text.strip()
            if cell_text:
                print(f"  Row {row_idx}, Cell {cell_idx}: {cell_text[:100]}")
