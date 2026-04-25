from docx import Document
import re

def find_lessons_anywhere(docx_path):
    doc = Document(docx_path)
    lesson_pattern = re.compile(r"LESSON\s+(\d+)", re.IGNORECASE)
    
    print(f"Searching for lessons in: {docx_path}")
    
    # 1. Search in top-level tables
    for i, table in enumerate(doc.tables):
        text = table.rows[0].cells[0].text.strip()
        if lesson_pattern.search(text):
            print(f"Top-level Table {i} starts with: {text[:50]}")
            
    # 2. Search in all paragraphs
    for i, p in enumerate(doc.paragraphs):
        if lesson_pattern.search(p.text):
            print(f"Paragraph {i} contains: {p.text.strip()[:50]}")

if __name__ == "__main__":
    import sys
    path = r"D:\LP\reference_docs\scraped\Unit_2__Area_and_Multiplication.docx"
    if len(sys.argv) > 1:
        path = sys.argv[1]
    find_lessons_anywhere(path)
