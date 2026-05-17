"""
Build a local .docx from one Google Docs navigation tab using Docs API JSON.

Use when Drive files.export hits exportSizeLimitExceeded on the whole document.
Output is best-effort (paragraphs, basic bold/italic, tables); complex layouts may differ
from a native Google export but are usually enough for table_extractor.

Usage (repo root):
  python tools/scraper/gdoc_tab_to_docx.py --json-path path/to/file.docs-api.json \\
    --tab-title "Grade 3 Unit 8" --out reference_docs/scraped/grade3_ela_exports/Unit8.docx

  python tools/scraper/gdoc_tab_to_docx.py DOCUMENT_ID --tab-id t.2plykckkv6ev --out Unit8.docx
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def _safe_docx_text(s: str) -> str:
    """Strip characters illegal in XML text nodes (lxml / python-docx reject NULL and most C0 controls)."""
    if not s:
        return s
    return "".join(
        c
        for c in s
        if c != "\x00" and (ord(c) >= 32 or c in "\t\n\r")
    )


def _collect_tabs_flat(tabs: List[Dict[str, Any]], out: List[Dict[str, Any]]) -> None:
    for tab in tabs or []:
        out.append(tab)
        _collect_tabs_flat(tab.get("childTabs") or [], out)


def _find_tab(
    doc: Dict[str, Any],
    tab_id: Optional[str],
    tab_title: Optional[str],
    tab_title_contains: Optional[str],
) -> Optional[Dict[str, Any]]:
    tabs = doc.get("tabs") or []
    flat: List[Dict[str, Any]] = []
    _collect_tabs_flat(tabs, flat)
    if tab_id:
        for t in flat:
            tid = (t.get("tabProperties") or {}).get("tabId")
            if tid == tab_id:
                return t
        return None
    if tab_title:
        for t in flat:
            title = (t.get("tabProperties") or {}).get("title") or ""
            if title == tab_title:
                return t
        return None
    if tab_title_contains:
        needle = tab_title_contains.lower()
        for t in flat:
            title = ((t.get("tabProperties") or {}).get("title") or "").lower()
            if needle in title:
                return t
        return None
    return None


def _paragraph_plain_text(paragraph: Dict[str, Any]) -> str:
    parts: List[str] = []
    for el in paragraph.get("elements") or []:
        if "textRun" in el:
            parts.append(el["textRun"].get("content") or "")
        elif "inlineObjectElement" in el:
            parts.append("\n[image]\n")
        elif "equation" in el:
            parts.append("\n[equation]\n")
    return "".join(parts)


def _heading_level(named_style: Optional[str]) -> Optional[int]:
    if not named_style:
        return None
    m = re.match(r"HEADING_(\d+)", named_style.upper())
    if m:
        return min(int(m.group(1)), 9)
    if named_style.upper() == "TITLE":
        return 0
    return None


def _add_paragraph_from_api(document: Document, p_el: Dict[str, Any]) -> None:
    para = p_el.get("paragraph") or {}
    style_type = (para.get("paragraphStyle") or {}).get("namedStyleType")
    level = _heading_level(style_type)
    text = _safe_docx_text(_paragraph_plain_text(para).replace("\r", "")).strip("\n")
    if not text.strip() and not any(
        "textRun" in x for x in (para.get("elements") or [])
    ):
        document.add_paragraph("")
        return

    if level is not None and level >= 0:
        h = document.add_heading(text if text else " ", level=min(level, 9) if level > 0 else 0)
        if level == 0:
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return

    doc_p = document.add_paragraph()
    for el in para.get("elements") or []:
        if "textRun" not in el:
            continue
        tr = el["textRun"]
        chunk = _safe_docx_text(tr.get("content") or "")
        if not chunk:
            continue
        run = doc_p.add_run(chunk)
        ts = tr.get("textStyle") or {}
        run.bold = bool(ts.get("bold"))
        run.italic = bool(ts.get("italic"))
        if ts.get("fontSize") and isinstance(ts["fontSize"], dict):
            sz = ts["fontSize"].get("magnitude")
            if sz is not None:
                run.font.size = Pt(float(sz))


def _max_table_columns(rows: List[Dict[str, Any]]) -> int:
    m = 0
    for row in rows:
        cells = row.get("tableCells") or []
        m = max(m, len(cells))
    return max(m, 1)


def _cell_text(cell: Dict[str, Any]) -> str:
    lines: List[str] = []
    for chunk in cell.get("content") or []:
        if "paragraph" in chunk:
            t = _paragraph_plain_text(chunk["paragraph"]).strip()
            if t:
                lines.append(t)
    return "\n".join(lines)


def _add_table_from_api(document: Document, tbl: Dict[str, Any]) -> None:
    rows = tbl.get("tableRows") or []
    if not rows:
        return
    ncols = _max_table_columns(rows)
    nrows = len(rows)
    table = document.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = row.get("tableCells") or []
        for ci, cell in enumerate(cells):
            if ci >= ncols:
                break
            txt = _safe_docx_text(_cell_text(cell))
            cell_obj = table.rows[ri].cells[ci]
            cell_obj.text = ""
            if txt:
                cell_obj.paragraphs[0].add_run(txt)


def _walk_content(document: Document, elements: List[Dict[str, Any]]) -> None:
    for el in elements or []:
        if "paragraph" in el:
            _add_paragraph_from_api(document, el)
        elif "table" in el:
            _add_table_from_api(document, el["table"])
        elif "sectionBreak" in el:
            document.add_paragraph("")
        elif "tableOfContents" in el:
            document.add_paragraph("[Table of contents]")


def tab_body_content(tab: Dict[str, Any]) -> List[Dict[str, Any]]:
    body = (tab.get("documentTab") or {}).get("body") or {}
    return body.get("content") or []


def build_docx_from_tab(tab: Dict[str, Any], source_title: str) -> Document:
    document = Document()
    document.core_properties.title = (tab.get("tabProperties") or {}).get("title") or "Tab export"
    document.core_properties.comments = f"Synthesized from Google Docs API JSON; source doc: {source_title}"
    _walk_content(document, tab_body_content(tab))
    return document


def _sanitize_file_stem(name: str, max_len: int = 100) -> str:
    base = "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()
    base = re.sub(r"_+", "_", base).strip("_")
    return (base or "tab")[:max_len]


def list_all_tabs(doc: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Flat preorder list of all navigation tabs (includes nested childTabs)."""
    flat: List[Dict[str, Any]] = []
    _collect_tabs_flat(doc.get("tabs") or [], flat)
    return flat


def export_all_tabs_to_directory(
    doc: Dict[str, Any],
    output_dir: str,
    basename_prefix: str = "",
) -> List[str]:
    """
    Write one .docx per tab (synthesized via Docs API structure).
    Returns absolute paths of files written.
    """
    os.makedirs(output_dir, exist_ok=True)
    source_title = doc.get("title") or ""
    flat = list_all_tabs(doc)
    written: List[str] = []
    used_stems: Set[str] = set()
    for idx, tab in enumerate(flat):
        props = tab.get("tabProperties") or {}
        ttitle = props.get("title") or f"tab_{idx}"
        tid = props.get("tabId") or str(idx)
        stem = _sanitize_file_stem(ttitle)
        tid_part = _sanitize_file_stem(tid.replace(".", "_"), max_len=48)
        fname_stem = f"{stem}__{tid_part}"
        if fname_stem in used_stems:
            fname_stem = f"{fname_stem}_{idx}"
        used_stems.add(fname_stem)
        prefix = _sanitize_file_stem(basename_prefix, max_len=60) if basename_prefix else ""
        fname = f"{prefix}__{fname_stem}.docx" if prefix else f"{fname_stem}.docx"
        path = os.path.join(output_dir, fname)
        built = build_docx_from_tab(tab, source_title)
        built.save(path)
        written.append(os.path.abspath(path))
    return written


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("document_id", nargs="?", help="Google Doc ID (if not using --json-path)")
    ap.add_argument("--json-path", metavar="PATH", help="Path to .docs-api.json snapshot")
    ap.add_argument("--out", required=True, help="Output .docx path")
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--tab-id", metavar="ID", help="Exact tabId (e.g. t.2plykckkv6ev)")
    g.add_argument("--tab-title", metavar="TITLE", help="Exact tab title match")
    g.add_argument("--tab-title-contains", metavar="SUBSTR", help="Substring match on tab title (lowercase)")
    args = ap.parse_args()

    if args.json_path:
        p = Path(args.json_path)
        if not p.is_file():
            print(f"File not found: {p}", file=sys.stderr)
            return 1
        doc = json.loads(p.read_text(encoding="utf-8"))
    elif args.document_id:
        repo_root = Path(__file__).resolve().parents[2]
        if str(repo_root) not in sys.path:
            sys.path.insert(0, str(repo_root))
        from tools.scraper.docs_client import DocsClient

        client = DocsClient()
        doc = client.get_document(args.document_id)
        if not doc:
            print("Failed to fetch document.", file=sys.stderr)
            return 1
    else:
        ap.print_help()
        print("\nProvide --json-path or DOCUMENT_ID.", file=sys.stderr)
        return 2

    tab = _find_tab(doc, args.tab_id, args.tab_title, args.tab_title_contains)
    if not tab:
        print("Tab not found. Run: python tools/scraper/inspect_gdoc_tabs.py ...", file=sys.stderr)
        return 1

    title = (tab.get("tabProperties") or {}).get("title", "tab")
    print(f"Building DOCX from tab: {title!r}", flush=True)

    out_path = Path(args.out).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    built = build_docx_from_tab(tab, doc.get("title") or "")
    built.save(str(out_path))
    print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
