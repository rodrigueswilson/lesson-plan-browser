import os
import json
from docx import Document
from table_extractor import RecursiveTableParser


def create_complex_docx(path):
    doc = Document()
    doc.add_paragraph("Test Document for Recursive Table Parser")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Standard Cell"
    table.cell(0, 1).text = "Merged Header"
    a = table.cell(0, 1)._tc
    b = table.cell(1, 1)._tc
    a.merge(b)

    table2 = doc.add_table(rows=1, cols=1)
    cell = table2.cell(0, 0)
    cell.text = "Parent Cell"
    nested_table = cell.add_table(rows=2, cols=2)
    nested_table.cell(0, 0).text = "Nested 1,1"
    nested_table.cell(1, 1).text = "Nested 2,2"

    table3 = doc.add_table(rows=2, cols=2)
    table3.cell(0, 0).text = "Monday Content"
    table3.cell(0, 1).text = "Tuesday Content"
    table3.cell(1, 0).text = "Math Lesson"
    table3.cell(1, 1).text = "Reading Lesson"

    p = doc.add_paragraph()
    run = p.add_run("Bold and Highlighed")
    run.bold = True
    run.font.highlight_color = 7

    doc.save(path)
    print(f"Created test doc: {path}")


def test_parser():
    test_path = "complex_test.docx"
    create_complex_docx(test_path)

    parser = RecursiveTableParser()
    result = parser.parse_document(test_path)
    flat = parser.flatten_elements(result)

    print("\n--- PARSED JSON (summary) ---")
    print("top_level_items", len(result))
    print("flattened_elements", len(flat))

    print("\n--- FLATTENED TEXT ---")
    print(parser.flatten_for_llm(result))

    assert len(result) > 0
    assert any(e.get("type") == "table" for e in result)

    if os.path.exists(test_path):
        os.remove(test_path)


if __name__ == "__main__":
    test_parser()
