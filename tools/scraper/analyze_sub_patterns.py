import sys
import re
from docx import Document

def analyze_sub_patterns(docx_path):
    doc = Document(docx_path)
    lesson_pattern = re.compile(r"LESSON\s+(\d+)", re.IGNORECASE)
    
    # Track patterns for specific sections across all lessons
    # section -> list of bold sub-headers found
    patterns = {
        "Teacher Objectives (R2 C1)": [],
        "Lesson Purpose (R2 C2)": [],
        "Row 8 Left (Supplemental)": [],
        "Row 8 Right (Formative)": []
    }
    
    for table in doc.tables:
        match = lesson_pattern.search(table.rows[0].cells[0].text.strip())
        if not match: continue
        
        # Row 2 (idx 1)
        r2 = table.rows[1]
        r2_cells = []
        for c in r2.cells:
            if c not in r2_cells: r2_cells.append(c)
        
        # R2 C1: Teacher Objectives
        bold_r2c1 = [p.text.strip() for p in r2_cells[0].paragraphs if any(run.bold for run in p.runs) and p.text.strip()]
        patterns["Teacher Objectives (R2 C1)"].append(bold_r2c1)
        
        # R2 C2: Lesson Purpose
        bold_r2c2 = [p.text.strip() for p in r2_cells[1].paragraphs if any(run.bold for run in p.runs) and p.text.strip()]
        patterns["Lesson Purpose (R2 C2)"].append(bold_r2c2)
        
        # Row 8 (idx 7)
        if len(table.rows) > 7:
            r8 = table.rows[7]
            r8_cells = []
            for c in r8.cells:
                if c not in r8_cells: r8_cells.append(c)
            
            bold_r8c1 = [p.text.strip() for p in r8_cells[0].paragraphs if any(run.bold for run in p.runs) and p.text.strip()]
            patterns["Row 8 Left (Supplemental)"].append(bold_r8c1)
            
            bold_r8c2 = [p.text.strip() for p in r8_cells[1].paragraphs if any(run.bold for run in p.runs) and p.text.strip()]
            patterns["Row 8 Right (Formative)"].append(bold_r8c2)

    print(f"Internal Cell Pattern Analysis: {docx_path}")
    print("-" * 50)
    for section, occurrences in patterns.items():
        print(f"{section}:")
        # Count identical bold header lists
        unique_patterns = {}
        for occ in occurrences:
            key = tuple(occ)
            unique_patterns[key] = unique_patterns.get(key, 0) + 1
        
        for p_tuple, count in unique_patterns.items():
            print(f"  - Headers: {list(p_tuple)} (Count: {count})")
        print("-" * 20)

if __name__ == "__main__":
    path = r"D:\LP\reference_docs\scraped\3rd grade_unit 2_docx\Copy of Unit_2__Area_and_Multiplication.docx"
    analyze_sub_patterns(path)
