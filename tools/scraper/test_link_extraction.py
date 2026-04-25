import os
from docx import Document
from docx.oxml.ns import qn

def extract_hyperlinks(paragraph):
    """
    Extracts hyperlinks from a paragraph by traversing its XML elements.
    Returns a list of dictionaries with text and url.
    """
    hyperlinks = []
    
    # Traverse all child elements of the paragraph's XML
    for child in paragraph._element:
        tag = child.tag.split('}')[-1]
        
        if tag == "hyperlink":
            # Extract rId and visible text
            r_id = child.get(qn('r:id'))
            if r_id in paragraph.part.rels:
                url = paragraph.part.rels[r_id].target_ref
                # Hyperlinks contain runs themselves
                link_text = "".join([t.text for t in child.xpath('.//w:t')])
                hyperlinks.append({
                    "text": link_text,
                    "url": url,
                    "rId": r_id
                })
        elif tag == "r":
            # Standard run text can be extracted if needed, 
            # but here we focus on hyperlinks.
            pass
            
    return hyperlinks

def test_on_file(docx_path):
    doc = Document(docx_path)
    found_links = []
    
    # Check paragraphs
    for p in doc.paragraphs:
        links = extract_hyperlinks(p)
        if links:
            found_links.extend(links)
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    links = extract_hyperlinks(p)
                    if links:
                        found_links.extend(links)
                        
    return found_links

if __name__ == "__main__":
    path = r"D:\LP\reference_docs\scraped\3rd grade_unit 2_docx\Copy of Unit_2__Area_and_Multiplication.docx"
    if os.path.exists(path):
        links = test_on_file(path)
        print(f"Found {len(links)} links:")
        for l in links[:10]: # Print first 10
            print(f"- {l['text']}: {l['url']}")
    else:
        print(f"File not found: {path}")
