import sys
import re
from docx import Document

def debug_tables(docx_path):
    doc = Document(docx_path)
    print(f"Total tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        try:
            first_cell = table.rows[0].cells[0].text.strip().replace('\n', ' ')
            snippet = first_cell[:100]
            print(f"Table {i}: {snippet}")
        except Exception as e:
            print(f"Table {i}: [Error: {e}]")

if __name__ == "__main__":
    path = r"D:\LP\reference_docs\scraped\Unit_2__Area_and_Multiplication.docx"
    debug_tables(path)
