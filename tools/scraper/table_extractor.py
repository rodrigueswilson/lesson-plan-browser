import os
import sys
import json
import hashlib
import re
import tempfile
import uuid
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Set, Tuple, Union

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl, CT_Row, CT_Tc
from docx.oxml.ns import qn
try:
    from tools.scraper.subject_config import SubjectConfig
except ImportError:
    from subject_config import SubjectConfig
try:
    from tools.scraper.docs_client import DocsClient
except ImportError:
    from docs_client import DocsClient

try:
    from tools.scraper.ingest_failure_codes import apply_ingest_failure_code
except ImportError:
    from ingest_failure_codes import apply_ingest_failure_code

try:
    from tools.scraper.ela_summary_table import is_summary_key_learning_table, parse_summary_rows
except ImportError:
    from ela_summary_table import is_summary_key_learning_table, parse_summary_rows

try:
    from tools.scraper.ela_lesson_plan_table import is_ela_lesson_plan_table, parse_ela_lesson_plan_table
except ImportError:
    from ela_lesson_plan_table import is_ela_lesson_plan_table, parse_ela_lesson_plan_table

try:
    from tools.scraper.cell_content_format import (
        collect_hyperlinks_from_buffer,
        split_leading_bold_runs_for_paragraph,
    )
except ImportError:
    from cell_content_format import (  # type: ignore[no-redef]
        collect_hyperlinks_from_buffer,
        split_leading_bold_runs_for_paragraph,
    )

# Add parent directory to path for backend imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from backend.database.sqlite_impl import SQLiteDatabase
from backend.database.curriculum import CurriculumDatabase

def log(msg, level="INFO"):
    print(f"[{level}] {msg}")


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _extract_source_doc_id(source_url: str) -> Optional[str]:
    if not source_url:
        return None
    m = re.search(r"/document/d/([a-zA-Z0-9_-]+)", source_url)
    return m.group(1) if m else None


def _extract_google_doc_id_from_url(url: str) -> Optional[str]:
    if not url:
        return None
    m = re.search(r"/document/d/([a-zA-Z0-9_-]+)", url)
    return m.group(1) if m else None


_LESSON_START_TO_INGEST_WARN_RATIO = 2.0


def _looks_like_equation_token_loss(text: str) -> bool:
    """Heuristic: detect standards lines where math/equation placeholders appear blank."""
    t = (text or "").lower()
    if not t:
        return False
    signals = (
        "as a multiple of .",
        "as the product ,",
        "by the equation .",
        "in general, .",
        "will eat  of a pound",
    )
    return any(s in t for s in signals)


def _xml_local_name(tag: str) -> str:
    return tag.split("}")[-1] if "}" in tag else tag


def _omml_node_to_text(node: Any) -> str:
    """Best-effort OMML -> text for curriculum fidelity (fractions, superscripts, roots)."""
    name = _xml_local_name(getattr(node, "tag", ""))
    if name in {"oMath", "oMathPara", "r", "box", "d"}:
        return "".join(_omml_node_to_text(c) for c in node)
    if name == "t":
        return (node.text or "").strip()
    if name == "f":
        num = ""
        den = ""
        for c in node:
            c_name = _xml_local_name(getattr(c, "tag", ""))
            if c_name == "num":
                num = _omml_node_to_text(c)
            elif c_name == "den":
                den = _omml_node_to_text(c)
        if num and den:
            return f"{num}/{den}"
        return "".join(_omml_node_to_text(c) for c in node)
    if name == "sSup":
        base = ""
        sup = ""
        for c in node:
            c_name = _xml_local_name(getattr(c, "tag", ""))
            if c_name == "e":
                base = _omml_node_to_text(c)
            elif c_name == "sup":
                sup = _omml_node_to_text(c)
        if base and sup:
            return f"{base}^{sup}"
    if name == "sSub":
        base = ""
        sub = ""
        for c in node:
            c_name = _xml_local_name(getattr(c, "tag", ""))
            if c_name == "e":
                base = _omml_node_to_text(c)
            elif c_name == "sub":
                sub = _omml_node_to_text(c)
        if base and sub:
            return f"{base}_{sub}"
    if name == "rad":
        deg = ""
        expr = ""
        for c in node:
            c_name = _xml_local_name(getattr(c, "tag", ""))
            if c_name == "deg":
                deg = _omml_node_to_text(c)
            elif c_name == "e":
                expr = _omml_node_to_text(c)
        if expr:
            return f"root({deg},{expr})" if deg else f"sqrt({expr})"
    if name == "nary":
        op = ""
        body = ""
        lower = ""
        upper = ""
        for c in node:
            c_name = _xml_local_name(getattr(c, "tag", ""))
            if c_name == "chr":
                op = c.get("{http://schemas.openxmlformats.org/officeDocument/2006/math}val", "")
            elif c_name == "sub":
                lower = _omml_node_to_text(c)
            elif c_name == "sup":
                upper = _omml_node_to_text(c)
            elif c_name == "e":
                body = _omml_node_to_text(c)
        if op or body:
            lim = ""
            if lower and upper:
                lim = f"_({lower})^({upper})"
            elif lower:
                lim = f"_({lower})"
            elif upper:
                lim = f"^({upper})"
            return f"{op}{lim}{body}"
    if name == "d":
        begin = ""
        end = ""
        expr = ""
        for c in node:
            c_name = _xml_local_name(getattr(c, "tag", ""))
            if c_name == "begChr":
                begin = c.get("{http://schemas.openxmlformats.org/officeDocument/2006/math}val", "")
            elif c_name == "endChr":
                end = c.get("{http://schemas.openxmlformats.org/officeDocument/2006/math}val", "")
            elif c_name == "e":
                expr = _omml_node_to_text(c)
        if expr:
            return f"{begin}{expr}{end}"
    if name == "m":
        rows: List[str] = []
        for c in node:
            if _xml_local_name(getattr(c, "tag", "")) == "mr":
                cells = []
                for cc in c:
                    if _xml_local_name(getattr(cc, "tag", "")) == "e":
                        cells.append(_omml_node_to_text(cc))
                if cells:
                    rows.append(",".join(cells))
        if rows:
            return "[" + ";".join(rows) + "]"
    return "".join(_omml_node_to_text(c) for c in node)


def _join_runs_for_paragraph_text(runs: List[Dict[str, Any]]) -> str:
    """Join run text while avoiding token glue around recovered math expressions."""
    out: List[str] = []
    prev_last = ""
    for run in runs:
        text = run.get("text") or ""
        if not text:
            continue
        cur_first = text[0]
        # Insert a spacer when two alphanumeric boundaries collide (e.g., "a/bas").
        if out and prev_last.isalnum() and cur_first.isalnum():
            out.append(" ")
        out.append(text)
        prev_last = text[-1]
    return "".join(out).strip()


def _annotate_ingest_report(
    report: Dict[str, Any],
    ingest_stats: Dict[str, Any],
    ingested_count: int,
) -> None:
    lessons_started_raw = int(ingest_stats.get("lessons_started") or 0)
    lessons_started_unique = int(ingest_stats.get("lessons_started_unique") or lessons_started_raw)
    report["ingest_stats"] = {
        "lessons_started": lessons_started_raw,
        "lessons_started_unique": lessons_started_unique,
        "paragraphs_appended": ingest_stats.get("paragraphs_appended"),
        "section_hits": dict(ingest_stats.get("section_hits") or {}),
        "suspected_equation_token_loss_standards": ingest_stats.get(
            "suspected_equation_token_loss_standards", 0
        ),
    }
    if ingested_count == 0 and lessons_started_unique == 0:
        apply_ingest_failure_code(report, "ANCHOR_MISS")
        report["warnings"].append("Ingest: no lesson titles matched the subject lesson pattern.")
    elif ingested_count == 0:
        apply_ingest_failure_code(report, "EMPTY_FIELD")
        report["warnings"].append(
            "Ingest: lessons started in stream but none had persistable body fields."
        )
    elif ingested_count > 0:
        start_to_ingest_ratio = lessons_started_unique / float(ingested_count)
        report["ingest_stats"]["lessons_started_to_ingested_ratio"] = round(
            start_to_ingest_ratio, 3
        )
        if start_to_ingest_ratio > _LESSON_START_TO_INGEST_WARN_RATIO:
            apply_ingest_failure_code(report, "ANCHOR_MISS")
            report["warnings"].append(
                "Ingest: lessons_started_unique/lessons_ingested ratio exceeds "
                f"{_LESSON_START_TO_INGEST_WARN_RATIO:.1f} "
                f"({lessons_started_unique}/{ingested_count}); review lesson title/anchor matching."
            )
    eq_loss_hits = int(ingest_stats.get("suspected_equation_token_loss_standards") or 0)
    if eq_loss_hits > 0:
        apply_ingest_failure_code(report, "EXPORT_LOSS")
        report["warnings"].append(
            "Ingest: detected standards lines with likely equation-token loss "
            f"({eq_loss_hits} hit(s)); review Docs/Export math fidelity."
        )


def _build_ingest_report(
    run_id: str,
    target_unit: str,
    source_doc_id: Optional[str],
    source_url: Optional[str],
    started_at: str,
    ended_at: str,
    parser_version: str,
    lessons_ingested: int,
) -> Dict[str, Any]:
    return {
        "run_id": run_id,
        "target_unit": target_unit,
        "source_doc_id": source_doc_id,
        "source_url": source_url,
        "started_at": started_at,
        "ended_at": ended_at,
        "parser_version": parser_version,
        "lessons_ingested": lessons_ingested,
        "primary_failure_code": None,
        "secondary_failure_codes": [],
        "warnings": [],
        "fidelity_checks": {
            "sample_lines_checked": 0,
            "mismatches": 0,
        },
    }


def _write_ingest_report(report: Dict[str, Any]) -> str:
    repo_root = Path(__file__).resolve().parents[2]
    out_dir = repo_root / "ingest_reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{report['run_id']}.json"
    with out_path.open("w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)
    return str(out_path)


def _normalize_anchor_text(text: str) -> str:
    t = (text or "").lower().strip()
    t = re.sub(r"\s+", " ", t)
    return t


def _resolve_field_from_anchors(
    normalized_text: str,
    field_map: Dict[str, str],
    max_len_for_substring: int = 120,
    short_key_len: int = 8,
) -> Optional[str]:
    if not normalized_text:
        return None
    if normalized_text in field_map:
        return field_map[normalized_text]
    if len(normalized_text) > max_len_for_substring:
        return None
    best_field: Optional[str] = None
    best_key_len = -1
    for key, field in field_map.items():
        if not key:
            continue
        if len(key) < short_key_len:
            if normalized_text == key:
                matched = True
            elif (
                normalized_text.startswith(key + " ")
                or normalized_text.startswith(key + ":")
                or normalized_text.startswith(key + "(")
                or normalized_text.startswith(key + ".")
            ):
                matched = True
            else:
                matched = False
        else:
            matched = key in normalized_text
        if matched and len(key) > best_key_len:
            best_key_len = len(key)
            best_field = field
    return best_field

def _extract_vocab_terms(text: str) -> List[str]:
    cleaned = (text or "").strip()
    if not cleaned:
        return []
    # Remove common bullet glyphs/prefixes and split lightweight term lists.
    cleaned = re.sub(r"^[\u2022\-\*\u25CF\u25A0\s]+", "", cleaned)
    parts = re.split(r"[,\n;]+", cleaned)
    return [p.strip() for p in parts if p and p.strip()]

def _is_procedure_subheader(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    return bool(
        re.match(
            r"^(warm-?up|lesson\s+activity|activity\s*\d+|cool-?down|lesson synthesis|activity synthesis)\b",
            t,
        )
    )

def _is_non_standard_heading(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    return bool(
        re.match(
            r"^(warm-?up|lesson\s+activity|activity\s*\d+|cool-?down|lesson synthesis|activity synthesis|math workshop resources|homework|suggested centers|gifted and talented|amplify desmos math)\b",
            t,
        )
    )


_COOLDOWN_TITLE_RE = re.compile(r"\b\d+\.\d+\.\d+\s+cool-?down\b", re.IGNORECASE)
_COOLDOWN_LESSON_RE = re.compile(r"\b\d+\.\d+\.(\d+)\s+cool-?down\b", re.IGNORECASE)
_DRIVE_FILE_ID_RE = re.compile(r"https://drive\.google\.com/file/d/([a-zA-Z0-9_-]+)", re.IGNORECASE)
_COOLDOWN_CACHE_BY_SOURCE_DOC: Dict[str, Dict[int, List[Dict[str, str]]]] = {}


def _html_escape(text: str) -> str:
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _cooldown_links_from_resources(resources: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    seen_urls: Set[str] = set()
    for res in resources or []:
        url = (res.get("url") or "").strip()
        text = (res.get("text") or "").strip()
        if not url or url in seen_urls:
            continue
        # Keep this strict to avoid mixing unrelated links into formative assessment.
        if _COOLDOWN_TITLE_RE.search(text):
            out.append({"text": text, "url": url})
            seen_urls.add(url)
    return out


def _extract_drive_file_id(url: str) -> Optional[str]:
    m = _DRIVE_FILE_ID_RE.search(url or "")
    return m.group(1) if m else None


def _iter_rich_link_props(node: Any) -> Iterator[Dict[str, Any]]:
    if isinstance(node, dict):
        rich = node.get("richLinkProperties")
        if isinstance(rich, dict):
            yield rich
        for v in node.values():
            yield from _iter_rich_link_props(v)
    elif isinstance(node, list):
        for item in node:
            yield from _iter_rich_link_props(item)


def _cooldown_links_from_source_json(source_doc_id: Optional[str]) -> Dict[int, List[Dict[str, str]]]:
    if not source_doc_id:
        return {}
    cached = _COOLDOWN_CACHE_BY_SOURCE_DOC.get(source_doc_id)
    if cached is not None:
        return cached

    repo_root = Path(__file__).resolve().parents[2]
    scraped_root = repo_root / "reference_docs" / "scraped"
    lesson_map: Dict[int, List[Dict[str, str]]] = {}
    if not scraped_root.is_dir():
        _COOLDOWN_CACHE_BY_SOURCE_DOC[source_doc_id] = lesson_map
        return lesson_map

    doc_id_re = re.compile(r'"documentId"\s*:\s*"' + re.escape(source_doc_id) + r'"')
    for json_path in scraped_root.rglob("originals/*.json"):
        try:
            raw = json_path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue
        if not doc_id_re.search(raw):
            continue
        try:
            payload = json.loads(raw)
        except Exception:
            continue
        for rich in _iter_rich_link_props(payload):
            title = str(rich.get("title") or "").strip()
            uri = str(rich.get("uri") or "").strip()
            mime = str(rich.get("mimeType") or "").strip().lower()
            m = _COOLDOWN_LESSON_RE.search(title)
            if not m or not uri:
                continue
            if mime and "pdf" not in mime:
                continue
            lesson_num = int(m.group(1))
            lesson_map.setdefault(lesson_num, [])
            if uri not in {x["url"] for x in lesson_map[lesson_num]}:
                lesson_map[lesson_num].append({"text": title, "url": uri})
        break

    _COOLDOWN_CACHE_BY_SOURCE_DOC[source_doc_id] = lesson_map
    return lesson_map


def _merge_lesson_cooldown_resources(
    lesson_data: Dict[str, Any],
    source_doc_id: Optional[str],
) -> None:
    lesson_num = int(lesson_data.get("lesson_number") or 0)
    if lesson_num <= 0:
        return
    source_map = _cooldown_links_from_source_json(source_doc_id)
    additions = source_map.get(lesson_num, [])
    if not additions:
        return
    existing_urls = {str((r or {}).get("url") or "") for r in (lesson_data.get("_resources") or [])}
    for item in additions:
        url = item["url"]
        if url in existing_urls:
            continue
        res: Dict[str, Any] = {"url": url, "text": item["text"]}
        file_id = _extract_drive_file_id(url)
        if file_id:
            res["google_id"] = file_id
        lesson_data["_resources"].append(res)
        existing_urls.add(url)


def _append_formative_assessment_resources_section(
    instructional_resources_html: str,
    resources: List[Dict[str, Any]],
) -> str:
    cooldown_links = _cooldown_links_from_resources(resources)
    if not cooldown_links:
        return instructional_resources_html or ""

    existing = instructional_resources_html or ""
    missing = [r for r in cooldown_links if r["url"] not in existing]
    if not missing:
        return existing

    has_heading = bool(re.search(r"formative assessment resources", existing, re.IGNORECASE))
    parts: List[str] = []
    if existing.strip():
        parts.append(existing.strip())
    if not has_heading:
        parts.append("<p><b>Formative Assessment Resources</b></p>")
    parts.append("<ul>")
    for res in missing:
        label = _html_escape(res["text"] or "Cool-down")
        url = _html_escape(res["url"])
        parts.append(
            f'<li><a href="{url}" target="_blank" rel="noopener noreferrer">{label}</a></li>'
        )
    parts.append("</ul>")
    return "".join(parts)


_ELA_RESOURCE_LINK_CACHE: Dict[Tuple[str, int, int], List[Dict[str, str]]] = {}
_ELA_RESOURCE_BY_UNIT_CACHE: Dict[Tuple[str, int], Dict[int, List[Dict[str, str]]]] = {}


def _unit_number_from_unit_id(unit_id: str) -> Optional[int]:
    m = re.search(r"_U(\d+)", unit_id or "", re.IGNORECASE)
    if not m:
        return None
    try:
        return int(m.group(1))
    except ValueError:
        return None


def _extract_markdown_links(text: str) -> List[Dict[str, str]]:
    links: List[Dict[str, str]] = []
    for m in re.finditer(r"\[([^\]]+)\]\((https?://[^)]+)\)", text or "", re.IGNORECASE):
        label = (m.group(1) or "").strip()
        url = (m.group(2) or "").strip()
        if not label or not url:
            continue
        links.append({"text": label, "url": url})
    return links


def _load_ela_instructional_links_from_reference_markdown(
    source_doc_id: Optional[str],
    unit_id: str,
    lesson_number: int,
) -> List[Dict[str, str]]:
    if not source_doc_id or lesson_number <= 0:
        return []
    unit_number = _unit_number_from_unit_id(unit_id)
    if unit_number is None:
        return []
    cache_key = (source_doc_id, unit_number, lesson_number)
    if cache_key in _ELA_RESOURCE_LINK_CACHE:
        return _ELA_RESOURCE_LINK_CACHE[cache_key]
    unit_cache_key = (source_doc_id, unit_number)
    if unit_cache_key in _ELA_RESOURCE_BY_UNIT_CACHE:
        links = _ELA_RESOURCE_BY_UNIT_CACHE[unit_cache_key].get(lesson_number, [])
        _ELA_RESOURCE_LINK_CACHE[cache_key] = links
        return links

    repo_root = Path(__file__).resolve().parents[2]
    roots = [
        repo_root / "reference_docs" / "curriculum",
        repo_root / "reference_docs" / "scraped_final",
    ]
    target_file = f"Grade_2_Unit_{unit_number}.md"
    lesson_links_map: Dict[int, List[Dict[str, str]]] = {}

    for root in roots:
        if not root.is_dir():
            continue
        for md in root.rglob(target_file):
            try:
                raw = md.read_text(encoding="utf-8", errors="ignore")
            except OSError:
                continue
            if source_doc_id not in raw:
                continue
            lines = raw.splitlines()
            starts: List[Tuple[int, int]] = []
            for i, line in enumerate(lines):
                m = re.match(r"^\|\s*Lesson\s+(\d+)\s*:", line.strip(), re.IGNORECASE)
                if m:
                    try:
                        starts.append((i, int(m.group(1))))
                    except ValueError:
                        continue
            if not starts:
                continue
            for idx, (line_idx, ln) in enumerate(starts):
                end_idx = starts[idx + 1][0] if idx + 1 < len(starts) else len(lines)
                block = "\n".join(lines[line_idx:end_idx])
                links: List[Dict[str, str]] = []
                seen_urls: Set[str] = set()
                for link in _extract_markdown_links(block):
                    if link["url"] in seen_urls:
                        continue
                    seen_urls.add(link["url"])
                    links.append(link)
                if links:
                    lesson_links_map[ln] = links
            if lesson_links_map:
                _ELA_RESOURCE_BY_UNIT_CACHE[unit_cache_key] = lesson_links_map
                links = lesson_links_map.get(lesson_number, [])
                _ELA_RESOURCE_LINK_CACHE[cache_key] = links
                return links

    _ELA_RESOURCE_BY_UNIT_CACHE[unit_cache_key] = lesson_links_map
    links = lesson_links_map.get(lesson_number, [])
    _ELA_RESOURCE_LINK_CACHE[cache_key] = links
    return links


def _inject_links_into_instructional_resources_html(
    html: str,
    links: List[Dict[str, str]],
) -> str:
    rendered = html or ""
    if not rendered or "<a href=" in rendered or not links:
        return rendered
    for link in sorted(links, key=lambda x: len(x.get("text") or ""), reverse=True):
        label = (link.get("text") or "").strip()
        url = (link.get("url") or "").strip()
        if not label or not url:
            continue
        anchor = (
            f'<a href="{_html_escape(url)}" target="_blank" rel="noopener noreferrer">{_html_escape(label)}</a>'
        )
        rendered = re.sub(
            re.escape(label),
            anchor,
            rendered,
            count=1,
            flags=re.IGNORECASE,
        )
    return rendered


def _inject_links_into_html_by_label(html: str, links: List[Dict[str, str]]) -> str:
    """Best-effort label-based link hydration for plain-text HTML blocks."""
    rendered = html or ""
    if not rendered or "<a href=" in rendered or not links:
        return rendered
    out = rendered
    for link in sorted(links, key=lambda x: len(x.get("text") or ""), reverse=True):
        label = (link.get("text") or "").strip()
        url = (link.get("url") or "").strip()
        if not label or not url:
            continue
        anchor = (
            f'<a href="{_html_escape(url)}" target="_blank" rel="noopener noreferrer">{_html_escape(label)}</a>'
        )
        out = re.sub(re.escape(label), anchor, out, count=1, flags=re.IGNORECASE)
    return out


def _collect_links_from_ela_structured_plan(plan: Dict[str, Any]) -> List[Dict[str, str]]:
    """Extract all anchors from all *_html string fields in structured ELA lesson payload."""
    out: List[Dict[str, str]] = []
    seen: Set[Tuple[str, str]] = set()
    for key, value in (plan or {}).items():
        if not key.endswith("_html"):
            continue
        if not isinstance(value, str) or not value.strip():
            continue
        for m in re.finditer(
            r'<a\s+[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
            value,
            re.IGNORECASE | re.DOTALL,
        ):
            url = (m.group(1) or "").strip()
            label = re.sub(r"<[^>]+>", "", m.group(2) or "").strip()
            if not url:
                continue
            identity = (url, label.lower())
            if identity in seen:
                continue
            seen.add(identity)
            out.append({"url": url, "text": label or url})
        # Safety net: capture any href attributes, including malformed/nested anchors.
        for m in re.finditer(r'href="([^"]+)"', value, re.IGNORECASE):
            url = (m.group(1) or "").strip()
            if not url:
                continue
            identity = (url, url.lower())
            if identity in seen:
                continue
            seen.add(identity)
            out.append({"url": url, "text": url})
    return out


def _collect_links_from_ela_structured_json(raw: Any) -> List[Dict[str, str]]:
    """Extract anchors from persisted `ela_lesson_plan_structured` JSON string."""
    if not isinstance(raw, str) or not raw.strip():
        return []
    try:
        plan = json.loads(raw)
    except Exception:
        return []
    if not isinstance(plan, dict):
        return []
    return _collect_links_from_ela_structured_plan(plan)


def _collect_links_from_lesson_html_fields(lesson_data: Dict[str, Any]) -> List[Dict[str, str]]:
    """Extract anchors from non-structured lesson HTML columns persisted on lessons table."""
    out: List[Dict[str, str]] = []
    seen: Set[Tuple[str, str]] = set()
    html_fields = (
        "instructional_resources",
        "procedure_html",
        "narrative_html",
        "materials",
        "daily_instructional_task",
        "purpose",
    )
    for field in html_fields:
        html = lesson_data.get(field)
        if not isinstance(html, str) or not html.strip():
            continue
        for m in re.finditer(
            r'<a\s+[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
            html,
            re.IGNORECASE | re.DOTALL,
        ):
            url = (m.group(1) or "").strip()
            label = re.sub(r"<[^>]+>", "", m.group(2) or "").strip()
            if not url:
                continue
            identity = (url, label.lower())
            if identity in seen:
                continue
            seen.add(identity)
            out.append({"url": url, "text": label or url})
        # Safety net for malformed HTML where nested anchors hide inner hrefs.
        for m in re.finditer(r'href="([^"]+)"', html, re.IGNORECASE):
            url = (m.group(1) or "").strip()
            if not url:
                continue
            identity = (url, url.lower())
            if identity in seen:
                continue
            seen.add(identity)
            out.append({"url": url, "text": url})
    return out


def _merge_links_into_resources(
    lesson_data: Dict[str, Any],
    links: List[Dict[str, str]],
) -> None:
    existing_urls = {str((r or {}).get("url") or "") for r in (lesson_data.get("_resources") or [])}
    for link in links:
        url = (link.get("url") or "").strip()
        text = (link.get("text") or "").strip()
        if not url or url in existing_urls:
            continue
        res: Dict[str, Any] = {"url": url, "text": text or url}
        google_id = _extract_google_doc_id_from_url(url)
        if google_id:
            res["google_id"] = google_id
        lesson_data["_resources"].append(res)
        existing_urls.add(url)


def _backfill_ela_structured_instructional_links(
    lesson_data: Dict[str, Any],
    unit_id: str,
    source_doc_id: Optional[str],
) -> None:
    raw = (lesson_data.get("ela_lesson_plan_structured") or "").strip()
    if not raw:
        return
    try:
        plan = json.loads(raw)
    except Exception:
        return
    if not isinstance(plan, dict):
        return
    html = str(plan.get("instructional_resources_cell_html") or "").strip()
    if not html or "<a href=" in html:
        return
    lesson_number = int(lesson_data.get("lesson_number") or 0)
    links = _load_ela_instructional_links_from_reference_markdown(
        source_doc_id=source_doc_id,
        unit_id=unit_id,
        lesson_number=lesson_number,
    )
    if not links:
        return
    for key, value in list(plan.items()):
        if not key.endswith("_html"):
            continue
        if not isinstance(value, str) or not value.strip():
            continue
        plan[key] = _inject_links_into_html_by_label(value, links)
    lesson_data["ela_lesson_plan_structured"] = json.dumps(plan, ensure_ascii=False)
    # All-cells SSOT: persist all links found in structured plan, not just one cell.
    collected = _collect_links_from_ela_structured_plan(plan)
    _merge_links_into_resources(lesson_data, collected)

_STANDARDS_SECTION_META: Dict[str, Dict[str, str]] = {
    "new jersey state learning standards": {
        "panel": "left",
        "section": "New Jersey State Learning Standards",
    },
    "mathematical practice standards": {
        "panel": "left",
        "section": "Mathematical Practice Standards",
    },
    "national council of teachers of mathematics content standards": {
        "panel": "right",
        "section": "National Council of Teachers of Mathematics Content Standards",
    },
    "national council of teachers of mathematics process standards": {
        "panel": "right",
        "section": "National Council of Teachers of Mathematics Process Standards",
    },
}


def _ends_with_sentence_terminal(text: str) -> bool:
    t = (text or "").rstrip()
    if not t:
        return False
    return bool(re.search(r"(?:\.{3}|[.!?]|\u2026)\s*[\)\]\"'”’]*\s*$", t))


def _starts_with_lowercase_continuation(text: str) -> bool:
    s = (text or "").lstrip()
    if not s:
        return False
    return s[0].isalpha() and s[0].islower()


def _should_merge_docx_soft_break(prev: Dict[str, Any], nxt: Dict[str, Any]) -> bool:
    if prev.get("type") != "paragraph" or nxt.get("type") != "paragraph":
        return False
    if prev.get("is_bullet") and nxt.get("is_bullet"):
        return False
    p = (prev.get("text") or "").strip()
    n = (nxt.get("text") or "").strip()
    if not p or not n:
        return False
    if len(p) > 6000 or len(n) > 6000:
        return False
    if _ends_with_sentence_terminal(p):
        return False
    if not _starts_with_lowercase_continuation(n):
        return False
    if _is_procedure_subheader(n) or _is_non_standard_heading(n):
        return False
    nh = _normalize_anchor_text(n)
    if nh in _STANDARDS_SECTION_META:
        return False
    return True


def _merge_two_paragraph_items(a: Dict[str, Any], b: Dict[str, Any]) -> Dict[str, Any]:
    text_a = (a.get("text") or "").rstrip()
    text_b = (b.get("text") or "").lstrip()
    merged_text = f"{text_a} {text_b}".strip()
    runs_a = list(a.get("runs") or [])
    runs_b = list(b.get("runs") or [])
    runs: List[Dict[str, Any]] = []
    runs.extend(runs_a)
    if runs_a and runs_b and text_a and text_b:
        runs.append({"type": "run", "text": " ", "style": None})
    runs.extend(runs_b)
    return {
        "type": "paragraph",
        "text": merged_text,
        "is_bullet": False,
        "ilvl": 0,
        "runs": runs,
    }


# Lesson fields where DOCX often inserts a stray paragraph break mid-sentence.
_SOFT_BREAK_MERGE_SECTIONS: frozenset = frozenset({
    "narrative_html",
    "purpose",
    "procedure_html",
    "learning_intentions",
    "objectives_student",
    "mlr",
    "materials",
    "vocabulary",
    "instructional_resources",
    "daily_instructional_task",
    "success_criteria",
    "essential_questions",
    "lesson_narrative",
})


def merge_docx_soft_break_paragraphs(buffer: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Join adjacent paragraphs that Word split mid-sentence (stray Enter).
    Heuristic: previous chunk lacks sentence-ending punctuation; next starts with a lowercase letter.
    """
    if not buffer:
        return buffer
    out: List[Dict[str, Any]] = []
    i = 0
    while i < len(buffer):
        item = buffer[i]
        if item.get("type") != "paragraph":
            out.append(item)
            i += 1
            continue
        cur: Dict[str, Any] = item
        j = i + 1
        while j < len(buffer):
            nxt = buffer[j]
            if not _should_merge_docx_soft_break(cur, nxt):
                break
            cur = _merge_two_paragraph_items(cur, nxt)
            j += 1
        out.append(cur)
        i = j
    return out


class RecursiveTableParser:
    """
    Extracted data from complex and nested tables in .docx files.
    Identifies 'Lesson X' structures for curriculum ingestion.
    """
    
    FRIENDLY_NAMES = {
        "p": "paragraph",
        "tbl": "table",
        "tr": "table-row",
        "tc": "table-cell",
        "r": "run",
        "t": "text"
    }

    def __init__(self, db_path: str = r"d:\LP\data\curriculum.db"):
        self.db_path = db_path
        self.db = SQLiteDatabase(db_path=db_path) # Always initialize db
        self.depth_limit = 10
        self.doc = None
        self._current_lesson = None
        self._current_section = None
        self._buffer = []
        self._docs_client = None # Lazy-load docs client
        # First top-level paragraph per json_to_html() root gets leading bold wrapped in .rich-cell-label
        self.enrich_leading_cell_labels = True

    def parse_document(self, docx_path: str) -> List[Dict[str, Any]]:
        """Entry point for parsing a document."""
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")
            
        doc = Document(docx_path)
        self.doc = doc
        content = []
        
        # In docx, top-level elements can be paragraphs or tables
        for element in doc.element.body:
            parsed = self.parse_element(element)
            if parsed:
                content.append(parsed)
                
        return content

    def parse_element(self, element: Any, depth: int = 0) -> Optional[Dict[str, Any]]:
        """Generic recursive handler for Docx elements."""
        if depth > self.depth_limit:
            return None
            
        tag_name = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        
        if tag_name == "p":
            return self._parse_paragraph(element)
        elif tag_name == "tbl":
            return self._parse_table(element, depth)
        
        return None

    def _parse_paragraph(self, p_element: Any) -> Dict[str, Any]:
        """Parses a paragraph element into JSON with bullet detection and link extraction."""
        p = Paragraph(p_element, None)
        
        # Bullet/Numbering detection
        ilvl = 0
        num_pr = p._element.xpath('.//w:numPr')
        if num_pr:
            ilvl_xpath = num_pr[0].xpath('./w:ilvl/@w:val')
            if ilvl_xpath:
                try: ilvl = int(ilvl_xpath[0])
                except ValueError: ilvl = 0
            
        runs = []
        # XML Traversal for high-fidelity link and run extraction
        for child in p._element:
            tag = child.tag.split('}')[-1]
            
            if tag == "hyperlink":
                r_id = child.get(qn('r:id'))
                url = self.doc.part.rels[r_id].target_ref if self.doc and r_id in self.doc.part.rels else None
                
                # Extract text within hyperlink (can have multiple runs)
                link_text = "".join([t.text for t in child.xpath('.//w:t')])
                
                style = {"link": True}
                if url:
                    style["url"] = url
                    # Extract Google ID if applicable
                    g_id_match = re.search(r"/document/d/([a-zA-Z0-9-_]+)", url)
                    if g_id_match:
                        style["google_id"] = g_id_match.group(1)

                runs.append({
                    "type": "run",
                    "text": link_text,
                    "style": style
                })
            elif tag == "r":
                # Standard run
                run_obj = None
                # Low-level run extraction to keep it simple but accurate
                for r in p.runs:
                    if r._element is child:
                        run_obj = r
                        break
                
                if run_obj:
                    style = {}
                    if run_obj.bold: style["bold"] = True
                    if run_obj.italic: style["italic"] = True
                    runs.append({
                        "type": "run",
                        "text": run_obj.text,
                        "style": style if style else None
                    })
            elif tag in {"oMath", "oMathPara"}:
                eq_text = _omml_node_to_text(child)
                if eq_text:
                    runs.append({
                        "type": "run",
                        "text": eq_text,
                        "style": {"equation": True}
                    })
        paragraph_text = _join_runs_for_paragraph_text(runs)
        return {
            "type": "paragraph",
            # Use run-composed text so OMML equations are preserved for ingest and LLM feed.
            "text": paragraph_text or p.text,
            "is_bullet": bool(num_pr),
            "ilvl": ilvl,
            "runs": runs
        }

    def _parse_table(self, tbl_element: CT_Tbl, depth: int) -> Dict[str, Any]:
        """Parses a table element recursively with robust merge handling."""
        rows = []
        
        # 1. Get max grid columns (Bypassing python-docx Twips conversion error)
        try:
            grid_cols = [int(float(gc.get(qn('w:w')))) for gc in tbl_element.tblGrid.gridCol_lst if gc.get(qn('w:w'))]
            total_grid_width = sum(grid_cols)
        except Exception as e:
            log(f"DEBUG: Grid parsing failed. {e}", "WARNING")
            grid_cols = []
            total_grid_width = 0
        
        for tr_xml in tbl_element.tr_lst:
            row_cells = []
            for tc_xml in tr_xml.tc_lst:
                v_merge = tc_xml.vMerge
                
                cell_json = {
                    "type": "table-cell",
                    "grid_span": tc_xml.grid_span,
                    "v_merge": v_merge,
                    "content": []
                }
                
                # Recurse through cell elements
                for child in tc_xml:
                    child_parsed = self.parse_element(child, depth + 1)
                    if child_parsed:
                        cell_json["content"].append(child_parsed)
                
                row_cells.append(cell_json)
                
            rows.append({
                "type": "table-row",
                "cells": row_cells
            })
            
        return {
            "type": "table",
            "rows": rows,
            "depth": depth
        }

    def flatten_for_llm(self, content: List[Dict[str, Any]]) -> str:
        """Converts structured JSON to a flattened text representation.

        NOTE (fidelity): keep raw Unicode/math-like tokens as-is here.
        Do not normalize/remove special characters in this path, because LLM
        prompts rely on exact curriculum text fidelity when symbols are present.
        """
        lines = []
        for item in content:
            if item["type"] == "paragraph":
                lines.append(item["text"])
            elif item["type"] == "table":
                lines.append(self._table_to_text(item))
        return "\n".join(lines)

    def _table_to_text(self, table: Dict[str, Any]) -> str:
        """Flatten table to a Markdown-style string."""
        if not table.get("rows"):
            return ""
            
        lines = []
        for row in table["rows"]:
            cells = []
            for cell in row["cells"]:
                cell_text = []
                for item in cell.get("content", []):
                    if item["type"] == "paragraph":
                        cell_text.append(item["text"])
                    elif item["type"] == "table":
                        cell_text.append(self._table_to_text(item))
                
                content = " ".join(cell_text).replace("\n", " ").strip()
                cells.append(content if content else " ")
            
            lines.append("| " + " | ".join(cells) + " |")
            
        return "\n".join(lines)
    def json_to_html(
        self,
        elements: Union[Dict[str, Any], List[Dict[str, Any]]],
        label_state: Optional[Dict[str, bool]] = None,
    ) -> str:
        """Converts structured JSON back to basic HTML, now with nested list support."""
        if self.enrich_leading_cell_labels and label_state is None:
            label_state = {"used": False}
        elif not self.enrich_leading_cell_labels:
            label_state = None

        if isinstance(elements, list):
            html = ""
            list_stack = []  # Stack of current ilvl

            for e in elements:
                if e.get("is_bullet"):
                    target_lvl = e.get("ilvl", 0)

                    if not list_stack:
                        html += "<ul>"
                        list_stack.append(target_lvl)
                    elif target_lvl > list_stack[-1]:
                        while target_lvl > list_stack[-1]:
                            html += "<ul>"
                            list_stack.append(list_stack[-1] + 1)
                        list_stack[-1] = target_lvl
                    elif target_lvl < list_stack[-1]:
                        while list_stack and target_lvl < list_stack[-1]:
                            html += "</ul>"
                            list_stack.pop()
                        if not list_stack:
                            html += "<ul>"
                            list_stack.append(target_lvl)

                    html += f"<li>{self._get_paragraph_inner_html(e, enrich_leading=False)}</li>"
                else:
                    while list_stack:
                        html += "</ul>"
                        list_stack.pop()
                    html += self.json_to_html(e, label_state)

            while list_stack:
                html += "</ul>"
                list_stack.pop()
            return html

        el_type = elements.get("type")
        if el_type == "paragraph":
            use_enrich = (
                label_state is not None
                and not label_state["used"]
                and not elements.get("is_bullet")
            )
            inner: str
            if use_enrich:
                leading, _tail = split_leading_bold_runs_for_paragraph(elements)
                if leading:
                    inner = self._get_paragraph_inner_html(elements, enrich_leading=True)
                    label_state["used"] = True
                else:
                    inner = self._get_paragraph_inner_html(elements, enrich_leading=False)
            else:
                inner = self._get_paragraph_inner_html(elements, enrich_leading=False)
            if elements.get("is_bullet"):
                return inner
            return f"<p>{inner}</p>" if inner else ""

        if el_type == "table":
            html = '<table border="1">'
            for row in elements.get("rows", []):
                html += "<tr>"
                for cell in row.get("cells", []):
                    span = f' colspan="{cell["grid_span"]}"' if cell.get("grid_span") and cell["grid_span"] > 1 else ""
                    html += f"<td{span}>"
                    cell_state = {"used": False} if self.enrich_leading_cell_labels else None
                    html += self.json_to_html(cell.get("content", []), cell_state)
                    html += "</td>"
                html += "</tr>"
            html += "</table>"
            return html

        return ""

    def _get_paragraph_inner_html(
        self,
        p_json: Dict[str, Any],
        *,
        enrich_leading: bool = False,
    ) -> str:
        """Helper to get HTML inside a paragraph (handles runs and links)."""
        if enrich_leading:
            leading, tail = split_leading_bold_runs_for_paragraph(p_json)
            if leading:
                head_para = {**p_json, "runs": leading}
                tail_para = {**p_json, "runs": tail}
                inner_head = self._get_paragraph_inner_html(head_para, enrich_leading=False)
                inner_tail = (
                    self._get_paragraph_inner_html(tail_para, enrich_leading=False) if tail else ""
                )
                combined = f'<span class="rich-cell-label">{inner_head}</span>{inner_tail}'
                if "\n" in combined:
                    combined = combined.replace("\n", "<br>")
                return combined

        inner = ""
        for run in p_json.get("runs", []):
            text = run["text"]
            if not text:
                continue
            style = run.get("style") or {}

            if style.get("bold"):
                text = f"<b>{text}</b>"
            if style.get("italic"):
                text = f"<i>{text}</i>"

            if style.get("link") and style.get("url"):
                url = style["url"]
                g_id = style.get("google_id")
                attr = f' data-resource-id="{g_id}"' if g_id else ""
                text = (
                    f'<a href="{url}" target="_blank" rel="noopener noreferrer"{attr}>'
                    f"{text}</a>"
                )

            inner += text
        if inner and "\n" in inner:
            inner = inner.replace("\n", "<br>")
        return inner

    def flatten_elements(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Recursively flattens all paragraphs and sub-tables into a single sequence."""
        flat = []
        for e in elements:
            if e["type"] == "paragraph":
                flat.append(e)
            elif e["type"] == "table":
                flat.append(e)
                for row in e.get("rows", []):
                    for cell in row.get("cells", []):
                        flat.extend(self.flatten_elements(cell.get("content", [])))
        return flat

    def parse_to_stream(self, docx_path: str):
        """Generates a linearized stream of semantic elements from the document."""
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")
            
        doc = Document(docx_path)
        self.doc = doc
        for element in doc.element.body:
            tag_name = element.tag.split('}')[-1] if '}' in element.tag else element.tag
            
            if tag_name == "p":
                yield self._parse_paragraph(element)
            elif tag_name == "tbl":
                table_json = self._parse_table(element, 0)
                for item in self._table_to_semantic_items(table_json):
                    yield item

    def _stream_docx_for_curriculum_ingest(self, docx_path: str, subject: str) -> Iterator[Dict[str, Any]]:
        """Semantic stream for ingest_to_curriculum; ELA skips first Summary of Key Learning table."""
        self._ela_summary_by_lesson = {}
        self._ela_lesson_plan_by_lesson = {}
        if subject.upper() != "ELA":
            yield from self.parse_to_stream(docx_path)
            return

        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")

        doc = Document(docx_path)
        self.doc = doc
        summary_consumed = False
        cfg = SubjectConfig.get_config("ELA")
        pat = cfg["patterns"].get("standards_summary_mentions") or cfg["patterns"]["standards"]

        for element in doc.element.body:
            tag_name = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag_name == "p":
                yield self._parse_paragraph(element)
            elif tag_name == "tbl":
                table_json = self._parse_table(element, 0)
                if is_summary_key_learning_table(table_json):
                    if not summary_consumed:
                        parsed = parse_summary_rows(
                            table_json,
                            lambda c: self.json_to_html(c),
                            pat,
                        )
                        self._ela_summary_by_lesson.update(parsed)
                        summary_consumed = True
                        log(
                            f"ELA: captured Summary of Key Learning table ({len(parsed)} lesson row(s))",
                            "INFO",
                        )
                        continue
                    log(
                        "ELA: additional Summary of Key Learning-shaped table; "
                        "flattening into stream (only first table is structured)",
                        "WARNING",
                    )
                elif is_ela_lesson_plan_table(table_json):
                    plan = parse_ela_lesson_plan_table(
                        table_json,
                        lambda c: self.json_to_html(c),
                    )
                    if plan:
                        n = plan.get("lesson_number")
                        if n is not None:
                            if n in self._ela_lesson_plan_by_lesson:
                                log(
                                    f"ELA: replacing structured lesson plan for lesson {n} "
                                    "(duplicate detailed table in DOCX)",
                                    "WARNING",
                                )
                            self._ela_lesson_plan_by_lesson[n] = plan
                            log(
                                f"ELA: parsed detailed lesson plan table (lesson {n})",
                                "INFO",
                            )
                for item in self._table_to_semantic_items(table_json):
                    yield item

    def _table_to_semantic_items(self, table_json: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Converts a table into a sequence of elements, handling side-by-side metadata."""
        items = []
        rows = table_json.get("rows", [])
        
        # Known headers that can appear side-by-side
        HEADERS = ["supplemental resources", "formative assessment", "materials", "vocabulary", "standards"]
        
        for row in rows:
            cells = row.get("cells", [])
            is_processed = False
            
            if len(cells) == 2:
                c0_cont = cells[0].get("content", [])
                c1_cont = cells[1].get("content", [])
                c0_text = "".join([p.get("text", "") for p in c0_cont if p["type"] == "paragraph"]).strip()
                c1_text = "".join([p.get("text", "") for p in c1_cont if p["type"] == "paragraph"]).strip()
                
                # Check if both cells are headers
                is0 = any(h in c0_text.lower() for h in HEADERS) and len(c0_text) < 60
                is1 = any(h in c1_text.lower() for h in HEADERS) and len(c1_text) < 60
                
                if is0 and is1:
                    # Both are headers, but they might contain the values too
                    items.append({"type": "property", "key": c0_text, "content": c0_cont})
                    items.append({"type": "property", "key": c1_text, "content": c1_cont})
                    is_processed = True
                elif is0:
                    items.append({"type": "property", "key": c0_text, "content": c1_cont})
                    is_processed = True
            
            if not is_processed:
                for cell in cells:
                    for content_item in cell.get("content", []):
                        if content_item["type"] == "table":
                            items.extend(self._table_to_semantic_items(content_item))
                        else:
                            items.append(content_item)
        return items

    def flush_buffer(self, lesson_data: Dict[str, Any]):
        """Flushes the current buffer into the appropriate section of the lesson."""
        if not self._buffer or not lesson_data: return
        section = lesson_data.get("_current_section") or "procedure_html"
        buf_for_html: List[Dict[str, Any]] = self._buffer
        if section in _SOFT_BREAK_MERGE_SECTIONS:
            buf_for_html = merge_docx_soft_break_paragraphs(list(self._buffer))
        lesson_data[section] += self.json_to_html(buf_for_html)
        
        # Extract links and standards from buffer
        standards_section_headers = set(_STANDARDS_SECTION_META.keys())

        current_standard_code: Optional[str] = None
        current_structured_entry: Optional[Dict[str, Any]] = None
        current_structured_section = lesson_data.get("_standards_section", "")
        current_structured_panel = lesson_data.get("_standards_panel", "")
        for b in self._buffer:
            if b["type"] == "paragraph":
                p_text = (b.get("text") or "").strip()
                if p_text:
                    # Standards + descriptions
                    config = SubjectConfig.get_config(lesson_data["subject"]) # Assuming subject is available
                    standard_pattern = config["patterns"]["standards"]
                    normalized_line = _normalize_anchor_text(p_text)
                    if normalized_line in standards_section_headers:
                        section_meta = _STANDARDS_SECTION_META[normalized_line]
                        current_structured_section = section_meta["section"]
                        current_structured_panel = section_meta["panel"]
                        current_standard_code = None
                        current_structured_entry = None
                        continue
                    found_codes = standard_pattern.findall(p_text)

                    if found_codes:
                        current_standard_code = None
                        current_structured_entry = None
                        for code in found_codes:
                            normalized_code = (code or "").strip()
                            if not normalized_code:
                                continue
                            if normalized_code not in lesson_data["_standards"]:
                                lesson_data["_standards"][normalized_code] = ""

                            # Handle one-line forms like "MP3. Construct viable arguments..."
                            inline = re.match(
                                rf"^\s*{re.escape(normalized_code)}[\.\:\-\s]*(.+)$",
                                p_text,
                                flags=re.IGNORECASE,
                            )
                            if inline:
                                desc = inline.group(1).strip()
                                if desc:
                                    prev = lesson_data["_standards"].get(normalized_code, "")
                                    lesson_data["_standards"][normalized_code] = (
                                        f"{prev}\n{desc}".strip() if prev else desc
                                    )
                            entry: Dict[str, Any] = {
                                "panel": current_structured_panel or "left",
                                "section": current_structured_section or "Standards",
                                "code": normalized_code,
                                "description_lines": [],
                            }
                            if inline:
                                desc = inline.group(1).strip()
                                if desc:
                                    entry["description_lines"].append(desc)
                            lesson_data["_standards_structured"].append(entry)
                            current_structured_entry = entry
                            current_standard_code = normalized_code
                    elif (
                        current_standard_code
                        and normalized_line not in standards_section_headers
                    ):
                        if _is_non_standard_heading(p_text):
                            current_standard_code = None
                            current_structured_entry = None
                            continue
                        # Continuation line for the most recent standard code.
                        prev = lesson_data["_standards"].get(current_standard_code, "")
                        lesson_data["_standards"][current_standard_code] = (
                            f"{prev}\n{p_text}".strip() if prev else p_text
                        )
                        if current_structured_entry is not None:
                            current_structured_entry["description_lines"].append(p_text)

        for res in collect_hyperlinks_from_buffer(self._buffer):
            url = res.get("url")
            if url:
                lesson_data["_resources"].append(res)
        lesson_data["_standards_section"] = current_structured_section
        lesson_data["_standards_panel"] = current_structured_panel
        self._buffer.clear()

    def _ela_title_from_summary_payload(self, lesson_num: int, payload: Dict[str, Any]) -> str:
        daily = (payload.get("daily_task_title") or "").strip()
        daily_low = daily.lower().rstrip(":").strip()
        skip_daily = not daily or daily.isdigit() or daily_low in {
            "daily instructional task",
            "daily instructional tasks",
            "essential question",
            "perspective",
            "central idea (portfolio artifact)",
        }
        if not skip_daily:
            return daily[:500]
        li = (payload.get("learning_intention") or "").strip()
        if li:
            one = li.split("\n")[0].strip()
            if len(one) > 120:
                one = one[:120].rstrip() + "..."
            return one
        return f"Lesson {lesson_num}"

    def _ela_new_lesson_shell(
        self,
        *,
        unit_id: str,
        lesson_number: int,
        title: str,
        subject: str,
        source_doc_id: Optional[str],
        source_url: Optional[str],
        started_at: str,
        run_id: str,
        parser_version: str,
    ) -> Dict[str, Any]:
        return {
            "id": f"{unit_id}_L{lesson_number}",
            "unit_id": unit_id,
            "lesson_number": lesson_number,
            "title": title,
            "subject": subject,
            "narrative_html": "",
            "lesson_narrative": "",
            "learning_intentions": "",
            "procedure_html": "",
            "daily_instructional_task": "",
            "success_criteria": "",
            "essential_questions": "",
            "procedure": "",
            "instructional_resources": "",
            "materials": "",
            "objectives_student": "",
            "purpose": "",
            "mlr": "",
            "vocabulary": "",
            "practices": "",
            "procedures": "",
            "differentiation": "",
            "standards_structured": "",
            "source_doc_id": source_doc_id,
            "source_url": source_url,
            "ingested_at": started_at,
            "ingest_run_id": run_id,
            "ingest_parser_version": parser_version,
            "_current_section": None,
            "_standards": {},
            "_standards_structured": [],
            "_vocabulary": [],
            "_resources": [],
            "_standards_section": "",
            "_standards_panel": "",
            "_standards_temp": "",
            "_vocab_temp": "",
        }

    def _populate_ela_lesson_from_summary_payload(
        self, le: Dict[str, Any], payload: Dict[str, Any]
    ) -> None:
        if payload.get("learning_intention"):
            le["learning_intentions"] = payload["learning_intention"]
        if payload.get("success_criteria"):
            le["success_criteria"] = payload["success_criteria"]
        dt_title = (payload.get("daily_task_title") or "").strip()
        dt_body = payload.get("daily_task_body") or ""
        if dt_body:
            le["daily_instructional_task"] = (
                f"<p><b>{dt_title}</b></p>{dt_body}" if dt_title else dt_body
            )
        elif dt_title:
            le["daily_instructional_task"] = dt_title
        if payload.get("content_and_strategies"):
            le["procedure_html"] = payload["content_and_strategies"]

    def _ela_coalesce_lessons_with_summary(
        self,
        lessons_data: List[Dict[str, Any]],
        summary_map: Dict[int, Dict[str, Any]],
        unit_id: str,
        subject: str,
        source_doc_id: Optional[str],
        source_url: Optional[str],
        started_at: str,
        run_id: str,
        parser_version: str,
    ) -> List[Dict[str, Any]]:
        """
        When paragraph lesson anchors are missing or sparse vs the Summary of Key Learning
        matrix (common for API-built DOCX), synthesize or replace lessons from summary rows.
        """
        expected = len(summary_map)
        if expected == 0:
            return lessons_data
        parsed = len(lessons_data)
        ratio = parsed / expected if expected else 1.0
        replace = parsed == 0 or (expected >= 5 and ratio < 0.35)

        def one_synth(n: int) -> Dict[str, Any]:
            payload = summary_map[n]
            title = self._ela_title_from_summary_payload(n, payload)
            le = self._ela_new_lesson_shell(
                unit_id=unit_id,
                lesson_number=n,
                title=title,
                subject=subject,
                source_doc_id=source_doc_id,
                source_url=source_url,
                started_at=started_at,
                run_id=run_id,
                parser_version=parser_version,
            )
            self._populate_ela_lesson_from_summary_payload(le, payload)
            return le

        if replace:
            log(
                "ELA: coalescing to summary-table lessons only "
                f"(parsed={parsed}, summary_rows={expected})",
                "INFO",
            )
            return [one_synth(n) for n in sorted(summary_map.keys())]

        # Do not partially backfill missing summary rows when paragraph parsing already
        # captured most lessons (would duplicate flex/culminating rows mis-labeled in col A).
        return list(lessons_data)

    def ingest_to_curriculum(
        self,
        docx_path: str,
        unit_id: str,
        subject: str = "Math",
        source_url: Optional[str] = None,
        ingest_run_id: Optional[str] = None,
        parser_version: str = "table_extractor@v1",
    ):
        """Specialized ingestion using the Semantic Stream."""
        curr_db = CurriculumDatabase(self.db_path)
        started_at = _utc_now_iso()
        run_id = ingest_run_id or f"{started_at.replace(':', '-').replace('.', '-')}_{uuid.uuid4().hex[:8]}"
        source_doc_id = _extract_source_doc_id(source_url or "")
        stream = self._stream_docx_for_curriculum_ingest(docx_path, subject)
        
        lessons_data = []
        current_lesson = None
        
        config = SubjectConfig.get_config(subject)
        lesson_pattern = config["patterns"]["lesson_title"]
        standard_pattern = config["patterns"]["standards"]
        
        # Build reverse field map from anchors (normalized keys)
        FIELD_MAP: Dict[str, str] = {}
        for field, anchors in config["anchors"].items():
            for a in anchors:
                FIELD_MAP[_normalize_anchor_text(a)] = field

        # Special internal buffers for standards/vocab if not explicitly in anchors
        if "standards" not in FIELD_MAP:
            FIELD_MAP[_normalize_anchor_text("standards")] = "_standards_temp"
        if "vocabulary" not in FIELD_MAP:
            FIELD_MAP[_normalize_anchor_text("vocabulary")] = "_vocab_temp"

        ingest_stats: Dict[str, Any] = {
            "lessons_started": 0,
            "lessons_started_unique": 0,
            "paragraphs_appended": 0,
            "section_hits": Counter(),
            "suspected_equation_token_loss_standards": 0,
        }
        seen_lesson_numbers: Set[int] = set()

        for item in stream:
            item_type = item.get("type")
            if item_type == "paragraph":
                text = item.get("text", "").strip()
                lesson_match = lesson_pattern.match(text)
                if lesson_match:
                    # Math pattern: groups (n|x.x.x title, title); ELA pattern: (n, title) only.
                    m = lesson_match
                    if m.lastindex is not None and m.lastindex >= 3 and m.group(3) is not None:
                        lnum_str = m.group(1) or m.group(2)
                        ltitle = m.group(3).strip().split("   ")[0]
                    else:
                        lnum_str = m.group(1)
                        ltitle = (m.group(2) or "").strip().split("   ")[0]
                    lnum = int(lnum_str.split('.')[-1]) if '.' in lnum_str else int(lnum_str)
                    ingest_stats["lessons_started"] += 1
                    if lnum not in seen_lesson_numbers:
                        seen_lesson_numbers.add(lnum)
                        ingest_stats["lessons_started_unique"] += 1
                    # Store current links before starting new lesson
                    lesson_links = []
                    if current_lesson:
                        lesson_links = [
                            r["url"] for r in current_lesson.get("_resources", [])
                            if "docs.google.com/document/d/" in r.get("url", "")
                        ]
                    
                    self.flush_buffer(current_lesson)
                    
                    # Recursive ingestion for the PREVIOUS lesson before moving to next
                    if current_lesson and lesson_links:
                        self.process_recursive_links(current_lesson, lesson_links, subject)
                    
                    if current_lesson: lessons_data.append(current_lesson)
                    current_lesson = {
                        "id": f"{unit_id}_L{lnum}", "unit_id": unit_id, "lesson_number": lnum, "title": ltitle,
                        "subject": subject,
                        "narrative_html": "", "lesson_narrative": "", "learning_intentions": "", "procedure_html": "",
                        "daily_instructional_task": "", "success_criteria": "", "essential_questions": "",
                        "procedure": "", "instructional_resources": "", "materials": "", "objectives_student": "",
                        "purpose": "", "mlr": "", "vocabulary": "", "practices": "", "procedures": "",
                        "differentiation": "", "standards_structured": "",
                        "source_doc_id": source_doc_id,
                        "source_url": source_url,
                        "ingested_at": started_at,
                        "ingest_run_id": run_id,
                        "ingest_parser_version": parser_version,
                        "_current_section": None, "_standards": {}, "_standards_structured": [], "_vocabulary": [], "_resources": [],
                        "_standards_section": "", "_standards_panel": "",
                        "_standards_temp": "", "_vocab_temp": ""
                    }
                    continue

                if not current_lesson: continue
                low_text = _normalize_anchor_text(text)
                new_section = _resolve_field_from_anchors(low_text, FIELD_MAP)
                if not new_section and _is_procedure_subheader(text):
                    new_section = "procedure_html"
                if new_section:
                    self.flush_buffer(current_lesson)
                    if new_section == "_standards_temp" and low_text in _STANDARDS_SECTION_META:
                        section_meta = _STANDARDS_SECTION_META[low_text]
                        current_lesson["_standards_section"] = section_meta["section"]
                        current_lesson["_standards_panel"] = section_meta["panel"]
                    current_lesson["_current_section"] = new_section
                    ingest_stats["section_hits"][new_section] += 1
                    # Preserve activity headers (e.g., "Warm-up: ...") as visible content.
                    if new_section == "procedure_html" and _is_procedure_subheader(text):
                        self._buffer.append(item)
                        ingest_stats["paragraphs_appended"] += 1
                    continue
                if current_lesson.get("_current_section") in {"_vocab_temp", "vocabulary"}:
                    current_lesson["_vocabulary"].extend(_extract_vocab_terms(text))
                    continue
                self._buffer.append(item)
                ingest_stats["paragraphs_appended"] += 1

            elif item_type == "property":
                if not current_lesson: continue
                key = _normalize_anchor_text(item.get("key", ""))
                field = _resolve_field_from_anchors(key, FIELD_MAP, max_len_for_substring=200)
                if not field:
                    field = next((f for k, f in FIELD_MAP.items() if k in key), None)
                if field:
                    self.flush_buffer(current_lesson)
                    if field in {"_vocab_temp", "vocabulary"}:
                        txt = "".join([p.get("text", "") for p in item["content"] if p["type"] == "paragraph"])
                        current_lesson["_vocabulary"].extend(_extract_vocab_terms(txt))
                    elif field == "instructional_resources":
                        current_lesson["instructional_resources"] += self.json_to_html(item["content"])
                    else:
                        current_lesson[field] += self.json_to_html(item["content"])
                else:
                    current_lesson["procedure_html"] += self.json_to_html(item["content"])
            
            # Note: item_type == "table" is no longer yielded by parse_to_stream
            # Contents are streamed as paragraphs/properties with table_context

        self.flush_buffer(current_lesson)
        if current_lesson:
            # Ensure the final lesson also gets recursive link enrichment.
            final_links = [
                r["url"] for r in current_lesson.get("_resources", [])
                if "docs.google.com/document/d/" in r.get("url", "")
            ]
            if final_links:
                self.process_recursive_links(current_lesson, final_links, subject)
            lessons_data.append(current_lesson)

        log(
            "INGEST_STATS: lessons_started=%s paragraphs_buffered=%s section_hits=%s"
            % (
                ingest_stats["lessons_started"],
                ingest_stats["paragraphs_appended"],
                dict(ingest_stats["section_hits"]),
            ),
        )

        summary_map = getattr(self, "_ela_summary_by_lesson", None) or {}
        plan_map = getattr(self, "_ela_lesson_plan_by_lesson", None) or {}
        if subject.upper() == "ELA" and summary_map:
            lessons_data = self._ela_coalesce_lessons_with_summary(
                lessons_data,
                summary_map,
                unit_id,
                subject,
                source_doc_id,
                source_url,
                started_at,
                run_id,
                parser_version,
            )
        for le in lessons_data:
            n = le.get("lesson_number")
            if n in summary_map:
                # SSOT: when a per-lesson structured plan exists, do not persist the summary matrix
                # JSON for the same lesson (explorer uses structured grid; avoids duplicate sources).
                if n not in plan_map:
                    payload = {"schema_version": 1, **summary_map[n]}
                    le["ela_key_learning_summary"] = json.dumps(payload, ensure_ascii=False)
            if n in plan_map:
                le["ela_lesson_plan_structured"] = json.dumps(
                    plan_map[n],
                    ensure_ascii=False,
                )

        unique_lessons = {
            l["id"]: l for l in lessons_data
            if any([
                l.get("narrative_html"),
                l.get("procedure_html"),
                l.get("instructional_resources"),
                l.get("learning_intentions"),
                l.get("daily_instructional_task"),
                l.get("success_criteria"),
                l.get("materials"),
                l.get("purpose"),
                l.get("mlr"),
                l.get("objectives_student"),
                l.get("vocabulary"),
                l.get("practices"),
                l.get("procedures"),
                l.get("differentiation"),
                l.get("_standards"),
                l.get("_vocabulary"),
                l.get("_resources"),
                l.get("ela_key_learning_summary"),
                l.get("ela_lesson_plan_structured"),
            ])
        }
        
        ingested_count = 0
        with curr_db._get_conn() as conn:
            conn.execute(
                """
                UPDATE units
                SET source_doc_id = ?,
                    source_url = ?,
                    ingested_at = ?,
                    ingest_run_id = ?,
                    ingest_parser_version = ?
                WHERE id = ?
                """,
                (source_doc_id, source_url, started_at, run_id, parser_version, unit_id),
            )
            conn.commit()
        for data in unique_lessons.values():
            _merge_lesson_cooldown_resources(data, source_doc_id)
            if subject.upper() == "ELA":
                _backfill_ela_structured_instructional_links(
                    lesson_data=data,
                    unit_id=unit_id,
                    source_doc_id=source_doc_id,
                )
                # Persist anchors already present inside structured-plan JSON.
                _merge_links_into_resources(
                    data,
                    _collect_links_from_ela_structured_json(
                        data.get("ela_lesson_plan_structured"),
                    ),
                )
            data["instructional_resources"] = _append_formative_assessment_resources_section(
                data.get("instructional_resources", ""),
                data.get("_resources", []),
            )
            # Global SSOT safety net: capture anchors already present in persisted HTML fields.
            _merge_links_into_resources(
                data,
                _collect_links_from_lesson_html_fields(data),
            )
            log(f"Upserting Lesson {data['lesson_number']}: {data['title']}")
            data["standards_structured"] = json.dumps(
                data.get("_standards_structured", []),
                ensure_ascii=False,
            )
            
            # 1. Clear existing links (individual transactions or one per lesson)
            with curr_db._get_conn() as conn:
                conn.execute("DELETE FROM lesson_standards WHERE lesson_id = ?", (data["id"],))
                conn.execute("DELETE FROM lesson_vocabulary WHERE lesson_id = ?", (data["id"],))
                conn.execute("DELETE FROM lesson_resources WHERE lesson_id = ?", (data["id"],))
                conn.commit()
                
            # 2. Save Standards
            for code, desc in data["_standards"].items():
                if _looks_like_equation_token_loss(desc):
                    ingest_stats["suspected_equation_token_loss_standards"] += 1
                try:
                    curr_db.upsert_standard(code, desc, subject)
                except TypeError:
                    curr_db.upsert_standard(code, desc)
                curr_db.link_lesson_to_standard(data["id"], code)
            
            # 3. Save Vocabulary
            for term in data["_vocabulary"]:
                curr_db.link_lesson_to_vocabulary(data["id"], term)
            
            # 4. Save Resources (Double Strategy)
            for res in data["_resources"]:
                curr_db.upsert_lesson_resource(data["id"], res)

            # 5. Save Lesson Core
            clean_data = {k: v for k, v in data.items() if not k.startswith("_")}
            curr_db.upsert_lesson(clean_data)
            ingested_count += 1
        ended_at = _utc_now_iso()
        report = _build_ingest_report(
            run_id=run_id,
            target_unit=unit_id,
            source_doc_id=source_doc_id,
            source_url=source_url,
            started_at=started_at,
            ended_at=ended_at,
            parser_version=parser_version,
            lessons_ingested=ingested_count,
        )
        _annotate_ingest_report(report, ingest_stats, ingested_count)
        report_path = _write_ingest_report(report)
        log(f"Ingest report written: {report_path}")
        return ingested_count

    def ingest_to_db(self, docx_path: str, metadata: Dict[str, Any]) -> str:
        """Parses DOCX and stores extraction cache in original_lesson_plans."""
        required = ["user_id", "week_of", "slot_number", "subject", "grade"]
        missing = [k for k in required if metadata.get(k) in (None, "")]
        if missing:
            raise ValueError(f"Missing required metadata for ingestion: {', '.join(missing)}")

        content = self.parse_document(docx_path)
        full_text = self.flatten_for_llm(content)
        content_hash = hashlib.md5(full_text.encode("utf-8")).hexdigest() if full_text else None

        source_file_path = os.path.abspath(docx_path)
        source_file_name = os.path.basename(source_file_path)
        identity = f"{metadata['user_id']}|{metadata['week_of']}|{metadata['slot_number']}|{source_file_path}"
        plan_id = hashlib.md5(identity.encode("utf-8")).hexdigest()

        plan_data = {
            "id": plan_id,
            "user_id": metadata["user_id"],
            "week_of": metadata["week_of"],
            "slot_number": int(metadata["slot_number"]),
            "subject": metadata["subject"],
            "grade": metadata["grade"],
            "homeroom": metadata.get("homeroom"),
            "source_file_path": source_file_path,
            "source_file_name": source_file_name,
            "primary_teacher_name": metadata.get("primary_teacher_name"),
            "content_json": content,
            "full_text": full_text,
            "available_days": metadata.get("available_days"),
            "has_no_school": bool(metadata.get("has_no_school", False)),
            "content_hash": content_hash,
            "status": metadata.get("status", "extracted"),
            "error_message": metadata.get("error_message"),
        }

        # original_lesson_plans lives on the app lesson-plans DB, not curriculum.db
        db = SQLiteDatabase()
        return db.create_original_lesson_plan(plan_data)

    def process_recursive_links(self, lesson: Dict[str, Any], links: List[str], subject: str):
        """Processes external links to fetch detailed content."""
        if not self._docs_client:
            try:
                self._docs_client = DocsClient()
            except Exception as e:
                print(f"[WARN] Failed to init DocsClient: {e}")
                return

        seen = set()
        for link in links:
            if link in seen:
                continue
            seen.add(link)
            # Extract Google Doc ID
            match = re.search(r"/d/([a-zA-Z0-9_-]+)", link)
            if not match:
                continue
            doc_id = match.group(1)

            print(f"[RECURSION] Fetching details for {lesson['title']} from ID: {doc_id}")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name

            try:
                success = self._docs_client.export_document(
                    doc_id,
                    tmp_path,
                    mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

                if success:
                    # Parse the detailed doc
                    # We create a new parser instance to avoid state pollution, or reset state
                    sub_parser = RecursiveTableParser(self.db_path)
                    detailed_stream = list(sub_parser.parse_to_stream(tmp_path))

                    # Merge content into our current lesson
                    # We primarily want the 'Procedure' from the detailed doc
                    detailed_html = ""
                    for item in detailed_stream:
                        if item["type"] == "paragraph":
                            detailed_html += self.json_to_html(item)
                        elif item["type"] == "property":
                            # If the detailed doc has its own headers/anchors, process them
                            key = item["key"].lower()
                            # Check if matches procedure anchors
                            if any(anchor in key for anchor in ["activity", "step", "procedure", "synthesis", "launch", "cool-down"]):
                                detailed_html += f"<h3>{item['key']}</h3>"
                                detailed_html += self.json_to_html(item["content"])

                    if detailed_html:
                        print(f"  -> Merged {len(detailed_html)} chars of detailed procedure.")
                        lesson["procedure_html"] = (lesson["procedure_html"] or "") + detailed_html
            except Exception as e:
                print(f"[WARN] Recursive link processing failed for {link}: {e}")
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)

if __name__ == "__main__":
    if len(sys.argv) > 2:
        # Mode: curriculum
        docx_path = sys.argv[1]
        unit_id = sys.argv[2]
        parser = RecursiveTableParser()
        count = parser.ingest_to_curriculum(docx_path, unit_id)
        log(f"Successfully processed {count} lessons into curriculum.db")
    elif len(sys.argv) > 1:
        # Mode: test/print
        path = sys.argv[1]
        parser = RecursiveTableParser()
        result = parser.parse_document(path)
        print(json.dumps(result, indent=2))
