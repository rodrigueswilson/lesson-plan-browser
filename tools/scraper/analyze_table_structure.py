import sys
import re
from docx import Document

def analyze_header_variations(docx_path):
    doc = Document(docx_path)
    lesson_pattern = re.compile(r"LESSON\s+(\d+)", re.IGNORECASE)
    
    KNOWN_HEADERS = [
        "teacher-facing", "student-facing", "lesson purpose", "lesson narrative",
        "mathematical language", "standards", "new jersey", "national council",
        "supplemental", "formative assessment"
    ]
    
    # row_idx -> col_idx -> list of found text
    variations = {
        2: {0: [], 1: []},
        3: {0: [], 1: []},
        4: {0: [], 1: []},
        8: {0: [], 1: []}
    }
    
    for table in doc.tables:
        first_cell_text = table.rows[0].cells[0].text.strip()
        if not lesson_pattern.search(first_cell_text): continue
        
        for r_offset in [1, 2, 3, 7]: # Rows 2, 3, 4, 8
            if r_offset >= len(table.rows): continue
            row = table.rows[r_offset]
            unique_cells = []
            for c in row.cells:
                if c not in unique_cells: unique_cells.append(c)
            
            for c_idx, cell in enumerate(unique_cells):
                if c_idx >= 2: continue
                # Find the first bold paragraph or matching a header word
                found_header = ""
                for p in cell.paragraphs[:3]:
                    txt = p.text.strip()
                    if not txt: continue
                    is_bold = any(run.bold is True for run in p.runs)
                    is_known = any(k in txt.lower() for k in KNOWN_HEADERS)
                    if is_bold or is_known:
                        found_header = txt
                        break
                
                if found_header:
                    variations[r_offset+1][c_idx].append(found_header)

    print(f"Header Content Analysis across ALL Lessons in: {docx_path}")
    print("-" * 50)
    for r_idx in [2, 3, 4, 8]:
        for c_idx in [0, 1]:
            coord = f"Row {r_idx}, Col {c_idx+1}"
            v_list = variations[r_idx][c_idx]
            unique_v = sorted(list(set(v_list)))
            print(f"{coord}:")
            for h in unique_v:
                count = v_list.count(h)
                print(f"  - '{h}' (Count: {count})")
            if not unique_v:
                print("  - [NONE FOUND]")
        print("-" * 20)

if __name__ == "__main__":
    path = r"D:\LP\reference_docs\scraped\3rd grade_unit 2_docx\Copy of Unit_2__Area_and_Multiplication.docx"
    if len(sys.argv) > 1:
        path = sys.argv[1]
    analyze_header_variations(path)
