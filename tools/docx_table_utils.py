"""
Table-related DOCX utilities: column width normalization and table info.
Re-exported from tools.docx_utils for backward compatibility.
"""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx import Document


def normalize_table_column_widths(
    table: Table, total_width_inches: float = 6.5
) -> None:
    """
    Set all columns in table to equal width.

    Args:
        table: python-docx Table object
        total_width_inches: Total width to distribute (default 6.5" for backward compatibility)
            IMPORTANT: Should be calculated from document page setup:
            available_width = (page_width - left_margin - right_margin) / 914400

    Example:
        >>> from docx import Document
        >>> doc = Document('input.docx')
        >>> for table in doc.tables:
        >>>     normalize_table_column_widths(table)
        >>> doc.save('output.docx')

    Note:
        - Table preferred width uses dxa/twips: 1 inch = 1,440 dxa
        - Column widths use python-docx Length values (EMU internally)
        - Handles merged cells correctly
        - Empty tables are skipped
        - Sets table preferred width type to 'dxa' to prevent auto-sizing
    """
    if not table.columns:
        return

    tbl_element = table._element
    tblPr = tbl_element.tblPr

    for tblW in tblPr.findall(qn("w:tblW")):
        tblPr.remove(tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(total_width_inches * 1440)))
    _insert_tbl_pr_element(
        tblPr,
        tblW,
        (
            "w:jc",
            "w:tblCellSpacing",
            "w:tblInd",
            "w:tblBorders",
            "w:shd",
            "w:tblLayout",
            "w:tblCellMar",
            "w:tblLook",
        ),
    )

    layout_list = tblPr.findall(qn("w:tblLayout"))
    for layout in layout_list:
        tblPr.remove(layout)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _insert_tbl_pr_element(
        tblPr,
        layout,
        (
            "w:tblCellMar",
            "w:tblLook",
        ),
    )

    tblInd_list = tblPr.findall(qn("w:tblInd"))
    for tblInd in tblInd_list:
        tblPr.remove(tblInd)

    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    _insert_tbl_pr_element(
        tblPr,
        tblInd,
        (
            "w:tblBorders",
            "w:shd",
            "w:tblLayout",
            "w:tblCellMar",
            "w:tblLook",
        ),
    )

    col_width = int(Inches(total_width_inches) / len(table.columns))

    for column in table.columns:
        column.width = col_width

    for cell in tbl_element.findall(f".//{qn('w:tc')}"):
        _ensure_cell_has_block_content(cell)


def _insert_tbl_pr_element(tblPr, element, successor_tags: tuple[str, ...]) -> None:
    successor_qnames = {qn(tag) for tag in successor_tags}
    for index, child in enumerate(tblPr):
        if child.tag in successor_qnames:
            tblPr.insert(index, element)
            return
    tblPr.append(element)


def _ensure_cell_has_block_content(cell) -> None:
    for child in cell:
        if child.tag != qn("w:tcPr"):
            return
    cell.append(OxmlElement("w:p"))


def normalize_all_tables(doc: Document, total_width_inches: float = 6.5) -> int:
    """
    Normalize column widths for all tables in document.

    Args:
        doc: python-docx Document object
        total_width_inches: Total width to distribute (default 6.5" for backward compatibility)
            IMPORTANT: Should be calculated from document page setup dynamically

    Returns:
        Number of tables normalized

    Example:
        >>> from docx import Document
        >>> doc = Document('input.docx')
        >>> count = normalize_all_tables(doc)
        >>> print(f"Normalized {count} tables")
        >>> doc.save('output.docx')
    """
    count = 0
    for table in doc.tables:
        normalize_table_column_widths(table, total_width_inches)
        count += 1

    return count


def get_table_info(table: Table) -> dict:
    """
    Get information about a table's structure.

    Args:
        table: python-docx Table object

    Returns:
        Dictionary with table information

    Example:
        >>> info = get_table_info(table)
        >>> print(f"Table has {info['num_columns']} columns")
    """
    info = {
        "num_rows": len(table.rows),
        "num_columns": len(table.columns),
        "column_widths": [col.width for col in table.columns] if table.columns else [],
        "has_merged_cells": False,
    }

    if table.rows:
        for row in table.rows:
            if len(row.cells) != len(table.columns):
                info["has_merged_cells"] = True
                break

    return info
