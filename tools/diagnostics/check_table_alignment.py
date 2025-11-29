"""Check table alignment/positioning in documents."""

from docx import Document
from pathlib import Path

files_to_check = [
    "input/Lesson Plan Template SY'25-26.docx",
    "output/test_signature.docx",
    "output/Maria Garcia_Lesson plan_W06_10-07-10-11.docx",
]

print("=" * 80)
print("CHECKING TABLE ALIGNMENT/POSITIONING")
print("=" * 80)

for file_path_str in files_to_check:
    file_path = Path(file_path_str)
    
    if not file_path.exists():
        print(f"\n✗ SKIPPED: {file_path.name} (not found)")
        continue
    
    print(f"\n{'=' * 80}")
    print(f"FILE: {file_path.name}")
    print(f"{'=' * 80}")
    
    try:
        doc = Document(str(file_path))
        
        print(f"\nTables ({len(doc.tables)} total):")
        
        for i, table in enumerate(doc.tables):
            # Get table alignment from XML
            tbl_element = table._element
            tblPr = tbl_element.tblPr
            
            # Check table justification (jc)
            jc = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
            if jc is not None:
                alignment = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                print(f"  Table {i+1}: Alignment = {alignment}")
            else:
                print(f"  Table {i+1}: Alignment = (default/left)")
            
            # Check table indent (tblInd)
            tblInd = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblInd')
            if tblInd is not None:
                indent_value = tblInd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                indent_type = tblInd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                print(f"           Indent = {indent_value} ({indent_type})")
            
            # Check table positioning (tblpPr) - for floating tables
            tblpPr = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblpPr')
            if tblpPr is not None:
                print(f"           Positioning = FLOATING TABLE")
            
    except Exception as e:
        print(f"  ERROR: {e}")
        import traceback
        traceback.print_exc()

print("\n" + "=" * 80)
print("EXPLANATION")
print("=" * 80)
print("\nTable Alignment Values:")
print("  - 'left' or (default): Table aligned to left margin")
print("  - 'center': Table centered between margins")
print("  - 'right': Table aligned to right margin")
print("\nIf tables have different alignments:")
print("  - Template tables should all have same alignment")
print("  - Output tables should match template alignment")
print("  - Misalignment causes visual inconsistency")
