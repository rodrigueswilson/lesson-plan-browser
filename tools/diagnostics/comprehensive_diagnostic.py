"""
Comprehensive Diagnostic: Input → Output Mapping with Telemetry
Implements all recommendations from the other AI.
"""

from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import json
import csv
from collections import defaultdict
from rapidfuzz import fuzz
import sys

sys.path.insert(0, str(Path(__file__).parent))
from tools.docx_parser import DOCXParser

class ComprehensiveDiagnostic:
    """Map input hyperlinks to output placement with detailed telemetry."""
    
    def __init__(self):
        self.results = []
        self.summary_stats = defaultdict(int)
        self.duplicate_texts = defaultdict(list)
        self.threshold_analysis = []
    
    def analyze_input_output_pair(self, input_path, output_path):
        """
        Map each input hyperlink to its output location.
        Categorize failures with detailed telemetry.
        """
        print(f"\n{'='*80}")
        print(f"Input:  {Path(input_path).name}")
        print(f"Output: {Path(output_path).name}")
        print(f"{'='*80}\n")
        
        # Extract input hyperlinks
        parser = DOCXParser(str(input_path))
        input_links = parser.extract_hyperlinks()
        
        print(f"📥 Input: {len(input_links)} hyperlinks")
        
        # Extract output hyperlinks (inline + referenced)
        output_doc = Document(output_path)
        inline_links = self._extract_inline_hyperlinks(output_doc)
        referenced_links = self._extract_referenced_links(output_doc)
        
        print(f"📤 Output: {len(inline_links)} inline, {len(referenced_links)} referenced")
        
        # Map each input link to output
        for input_link in input_links:
            result = self._map_link(input_link, inline_links, referenced_links, output_doc)
            self.results.append(result)
            self.summary_stats[result['status']] += 1
            
            # Track duplicates
            if result['duplicate_risk']:
                self.duplicate_texts[input_link['text']].append({
                    'url': input_link['url'],
                    'section': input_link.get('section_hint'),
                    'day': input_link.get('day_hint')
                })
        
        # Print summary
        print(f"\n📊 Mapping Results:")
        for status, count in sorted(self.summary_stats.items()):
            pct = (count / len(input_links) * 100) if input_links else 0
            print(f"  {status}: {count} ({pct:.1f}%)")
    
    def _extract_inline_hyperlinks(self, doc):
        """Extract inline hyperlinks with full context."""
        hyperlinks = []
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    
                    for para in cell.paragraphs:
                        for hyperlink in para._element.xpath('.//w:hyperlink'):
                            try:
                                r_id = hyperlink.get(qn('r:id'))
                                if r_id and r_id in para.part.rels:
                                    url = para.part.rels[r_id].target_ref
                                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                    
                                    if text and url:
                                        hyperlinks.append({
                                            'text': text,
                                            'url': url,
                                            'cell_text': cell_text,
                                            'location': f"T{table_idx}R{row_idx}C{cell_idx}",
                                            'table_idx': table_idx,
                                            'row_idx': row_idx,
                                            'cell_idx': cell_idx
                                        })
                            except:
                                pass
        
        return hyperlinks
    
    def _extract_referenced_links(self, doc):
        """Extract referenced links."""
        referenced = []
        in_section = False
        
        for para in doc.paragraphs:
            if "Referenced Links" in para.text:
                in_section = True
                continue
            
            if in_section:
                for hyperlink in para._element.xpath('.//w:hyperlink'):
                    try:
                        r_id = hyperlink.get(qn('r:id'))
                        if r_id and r_id in para.part.rels:
                            url = para.part.rels[r_id].target_ref
                            text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                            
                            if text and url:
                                referenced.append({'text': text, 'url': url})
                    except:
                        pass
        
        return referenced
    
    def _map_link(self, input_link, inline_links, referenced_links, output_doc):
        """
        Map a single input link to output location.
        Returns detailed diagnostic info.
        """
        link_text = input_link['text']
        link_url = input_link['url']
        
        # Check if placed inline
        inline_match = self._find_exact_match(link_text, link_url, inline_links)
        if inline_match:
            # Simulate matching strategy
            telemetry = self._simulate_matching(input_link, inline_match)
            
            return {
                'input_text': link_text,
                'input_url': link_url,
                'input_section': input_link.get('section_hint'),
                'input_day': input_link.get('day_hint'),
                'status': 'inline_placed',
                'output_location': inline_match['location'],
                'output_cell_text': inline_match['cell_text'][:100],
                'match_strategy': telemetry['strategy'],
                'match_confidence': telemetry['confidence'],
                'duplicate_risk': self._check_duplicate_risk(link_text, inline_links),
                'threshold_analysis': self._analyze_thresholds(input_link, inline_match)
            }
        
        # Check if in referenced links
        ref_match = self._find_exact_match(link_text, link_url, referenced_links)
        if ref_match:
            # Analyze why it failed inline placement
            failure_reason = self._diagnose_failure(input_link, inline_links, output_doc)
            
            return {
                'input_text': link_text,
                'input_url': link_url,
                'input_section': input_link.get('section_hint'),
                'input_day': input_link.get('day_hint'),
                'status': 'referenced_section',
                'failure_reason': failure_reason['reason'],
                'failure_details': failure_reason['details'],
                'best_fuzzy_score': failure_reason['best_score'],
                'duplicate_risk': self._check_duplicate_risk(link_text, inline_links),
                'threshold_analysis': failure_reason['threshold_analysis']
            }
        
        # Link completely missing!
        return {
            'input_text': link_text,
            'input_url': link_url,
            'input_section': input_link.get('section_hint'),
            'input_day': input_link.get('day_hint'),
            'status': 'MISSING',
            'failure_reason': 'Link not found in output',
            'duplicate_risk': False
        }
    
    def _find_exact_match(self, text, url, link_list):
        """Find exact match by text and URL."""
        for link in link_list:
            if link['text'] == text and link['url'] == url:
                return link
        return None
    
    def _simulate_matching(self, input_link, output_link):
        """
        Simulate the _calculate_match_confidence logic.
        Returns which strategy would have matched.
        """
        link_text = input_link['text']
        cell_text = output_link['cell_text']
        
        # Strategy 1: Exact text match
        if link_text in cell_text:
            return {'strategy': 'exact_text', 'confidence': 1.0}
        
        # Strategy 2: Case-insensitive
        if link_text.lower() in cell_text.lower():
            return {'strategy': 'exact_case_insensitive', 'confidence': 0.95}
        
        # Strategy 3: Normalized whitespace
        link_norm = ' '.join(link_text.split()).strip()
        cell_norm = ' '.join(cell_text.split()).strip()
        if link_norm.lower() in cell_norm.lower():
            return {'strategy': 'exact_normalized', 'confidence': 0.90}
        
        # Strategy 4: Fuzzy match
        context = input_link.get('context_snippet', '')
        if context:
            score = fuzz.partial_ratio(context, cell_text) / 100.0
            if score >= 0.65:
                return {'strategy': 'fuzzy_context', 'confidence': score}
        
        return {'strategy': 'unknown', 'confidence': 0.0}
    
    def _diagnose_failure(self, input_link, inline_links, output_doc):
        """
        Diagnose WHY a link failed to place inline.
        """
        link_text = input_link['text']
        context = input_link.get('context_snippet', '')
        expected_section = input_link.get('section_hint')
        expected_day = input_link.get('day_hint')
        
        # Get expected cell
        expected_cell = self._get_expected_cell(output_doc, expected_section, expected_day)
        
        if not expected_cell:
            return {
                'reason': 'expected_cell_not_found',
                'details': f"Section: {expected_section}, Day: {expected_day}",
                'best_score': 0.0,
                'threshold_analysis': {}
            }
        
        expected_cell_text = expected_cell.text
        
        # Check exact match
        if link_text in expected_cell_text:
            return {
                'reason': 'exact_match_but_not_placed',
                'details': 'Link text found in expected cell but not placed (bug?)',
                'best_score': 1.0,
                'threshold_analysis': {}
            }
        
        # Check case-insensitive
        if link_text.lower() in expected_cell_text.lower():
            return {
                'reason': 'case_mismatch',
                'details': f"Link: '{link_text}', Cell has different case",
                'best_score': 0.95,
                'threshold_analysis': {}
            }
        
        # Check whitespace
        link_norm = ' '.join(link_text.split()).strip()
        cell_norm = ' '.join(expected_cell_text.split()).strip()
        if link_norm.lower() in cell_norm.lower():
            return {
                'reason': 'whitespace_mismatch',
                'details': f"Link: '{link_text}' (has whitespace issues)",
                'best_score': 0.90,
                'threshold_analysis': {}
            }
        
        # Fuzzy match analysis
        fuzzy_score = fuzz.partial_ratio(context, expected_cell_text) / 100.0 if context else 0.0
        
        # Threshold sensitivity analysis
        threshold_analysis = self._threshold_sensitivity(context, expected_cell_text, inline_links)
        
        if fuzzy_score >= 0.50 and fuzzy_score < 0.65:
            return {
                'reason': 'fuzzy_below_threshold',
                'details': f"Score: {fuzzy_score:.2f} (threshold: 0.65)",
                'best_score': fuzzy_score,
                'threshold_analysis': threshold_analysis
            }
        
        # Text was rewritten
        return {
            'reason': 'text_rewritten',
            'details': f"Link text not found in expected cell (LLM rephrased?)",
            'best_score': fuzzy_score,
            'threshold_analysis': threshold_analysis
        }
    
    def _threshold_sensitivity(self, context, target_cell, all_cells):
        """
        Analyze how many cells would match at different thresholds.
        """
        if not context:
            return {}
        
        scores = []
        for cell in all_cells:
            score = fuzz.partial_ratio(context, cell['cell_text']) / 100.0
            scores.append(score)
        
        return {
            'target_score': fuzz.partial_ratio(context, target_cell) / 100.0 if target_cell else 0.0,
            'cells_above_0.60': sum(1 for s in scores if s >= 0.60),
            'cells_above_0.55': sum(1 for s in scores if s >= 0.55),
            'cells_above_0.50': sum(1 for s in scores if s >= 0.50),
            'max_competing_score': max(scores) if scores else 0.0
        }
    
    def _check_duplicate_risk(self, link_text, inline_links):
        """Check if link text appears multiple times."""
        count = sum(1 for link in inline_links if link['text'] == link_text)
        return count > 1
    
    def _get_expected_cell(self, doc, section_hint, day_hint):
        """Get the expected cell based on structural hints."""
        # Simplified - would need actual row/column mapping
        if not section_hint or not day_hint:
            return None
        
        # Map section to row index (example)
        section_map = {
            'unit_lesson': 1,
            'objective': 2,
            'anticipatory_set': 3,
            'instruction': 4,
            'misconceptions': 5,
            'assessment': 6,
            'homework': 7
        }
        
        day_map = {
            'monday': 1,
            'tuesday': 2,
            'wednesday': 3,
            'thursday': 4,
            'friday': 5
        }
        
        row_idx = section_map.get(section_hint)
        col_idx = day_map.get(day_hint)
        
        if row_idx and col_idx and doc.tables:
            try:
                return doc.tables[0].rows[row_idx].cells[col_idx]
            except:
                pass
        
        return None
    
    def _analyze_thresholds(self, input_link, output_link):
        """Analyze threshold sensitivity for this specific link."""
        context = input_link.get('context_snippet', '')
        if not context:
            return {}
        
        cell_text = output_link['cell_text']
        score = fuzz.partial_ratio(context, cell_text) / 100.0
        
        return {
            'fuzzy_score': score,
            'would_match_0.60': score >= 0.60,
            'would_match_0.55': score >= 0.55,
            'would_match_0.50': score >= 0.50
        }
    
    def export_to_csv(self, output_path):
        """Export results to CSV for analysis."""
        if not self.results:
            return
        
        # Flatten all results first to get all possible fields
        flattened_results = []
        for result in self.results:
            row = result.copy()
            # Flatten threshold_analysis
            if 'threshold_analysis' in row and isinstance(row['threshold_analysis'], dict):
                for k, v in row['threshold_analysis'].items():
                    row[f'threshold_{k}'] = v
                del row['threshold_analysis']
            # Convert other dicts to strings
            for k, v in row.items():
                if isinstance(v, dict):
                    row[k] = str(v)
            flattened_results.append(row)
        
        # Get all unique fieldnames
        all_fields = set()
        for row in flattened_results:
            all_fields.update(row.keys())
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=sorted(all_fields))
            writer.writeheader()
            for row in flattened_results:
                writer.writerow(row)
        
        print(f"\n💾 Exported {len(self.results)} results to: {output_path}")
    
    def print_summary(self):
        """Print comprehensive summary."""
        print(f"\n\n{'='*80}")
        print("📈 COMPREHENSIVE DIAGNOSTIC SUMMARY")
        print(f"{'='*80}\n")
        
        total = len(self.results)
        print(f"Total hyperlinks analyzed: {total}\n")
        
        # Status breakdown
        print("📊 Status Breakdown:")
        for status, count in sorted(self.summary_stats.items()):
            pct = (count / total * 100) if total > 0 else 0
            print(f"  {status}: {count} ({pct:.1f}%)")
        
        # Failure reason breakdown
        print(f"\n🔍 Failure Reasons (for referenced links):")
        failure_reasons = defaultdict(int)
        for result in self.results:
            if result['status'] == 'referenced_section':
                failure_reasons[result['failure_reason']] += 1
        
        for reason, count in sorted(failure_reasons.items(), key=lambda x: x[1], reverse=True):
            print(f"  {reason}: {count}")
        
        # Duplicate risks
        if self.duplicate_texts:
            print(f"\n⚠️  Duplicate Link Texts:")
            for text, occurrences in list(self.duplicate_texts.items())[:5]:
                print(f"  '{text}': {len(occurrences)} occurrences")


def main():
    """Run comprehensive diagnostic."""
    
    # Example: Map one input/output pair
    input_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Lang Lesson Plans 10_20_25-10_24_25.docx')
    output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_025901.docx')
    
    if not input_file.exists() or not output_file.exists():
        print(f"❌ Files not found!")
        print(f"Input: {input_file}")
        print(f"Output: {output_file}")
        return
    
    diagnostic = ComprehensiveDiagnostic()
    diagnostic.analyze_input_output_pair(str(input_file), str(output_file))
    diagnostic.print_summary()
    
    # Export to CSV
    csv_path = Path('d:/LP/hyperlink_diagnostic.csv')
    diagnostic.export_to_csv(str(csv_path))
    
    # Export to JSON
    json_path = Path('d:/LP/hyperlink_diagnostic.json')
    with open(json_path, 'w') as f:
        json.dump(diagnostic.results, f, indent=2)
    
    print(f"💾 Exported to JSON: {json_path}")


if __name__ == '__main__':
    main()
