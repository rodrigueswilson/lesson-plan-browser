"""Check hyperlinks in the single-slot Morais output."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# Get the most recent file
folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')
output_files = sorted(
    [f for f in folder.glob('Daniela_Silva_Lesson_plan_*.docx')],
    key=lambda f: f.stat().st_mtime,
    reverse=True
)

if not output_files:
    print("❌ No output files found!")
    exit(1)

file_path = output_files[0]

print("=" * 80)
print(f"📄 File: {file_path.name}")
print("=" * 80)

if not file_path.exists():
    print("❌ File not found!")
else:
    doc = Document(str(file_path))
    
    # Count hyperlinks
    para_links = 0
    table_links = 0
    
    for para in doc.paragraphs:
        para_links += len(para._element.xpath('.//w:hyperlink'))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_links += len(para._element.xpath('.//w:hyperlink'))
    
    total = para_links + table_links
    
    print(f"\n📊 HYPERLINK COUNT:")
    print(f"   Paragraph links: {para_links}")
    print(f"   Table links: {table_links}")
    print(f"   TOTAL: {total}")
    
    if total == 16:
        print(f"\n✅ SUCCESS! All 16 hyperlinks preserved!")
    elif total == 0:
        print(f"\n❌ FAILED! No hyperlinks in output")
    else:
        print(f"\n⚠️  PARTIAL! Expected 16, got {total}")
    
    # Check for Referenced Links section
    has_ref_section = False
    for para in doc.paragraphs:
        if 'referenced links' in para.text.lower():
            has_ref_section = True
            break
    
    print(f"   'Referenced Links' section: {'YES' if has_ref_section else 'NO'}")
    
    # Show first few hyperlinks if any
    if total > 0:
        print(f"\n📋 First 3 hyperlinks found:")
        count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for hl in para._element.xpath('.//w:hyperlink'):
                            r_id = hl.get(qn('r:id'))
                            if r_id and r_id in para.part.rels:
                                url = para.part.rels[r_id].target_ref
                                text = ''.join(node.text for node in hl.xpath('.//w:t') if node.text)
                                print(f"   {count+1}. {text[:40]} → {url[:50]}")
                                count += 1
                                if count >= 3:
                                    break
                        if count >= 3:
                            break
                    if count >= 3:
                        break
                if count >= 3:
                    break
            if count >= 3:
                break
    
    print("\n" + "=" * 80)
