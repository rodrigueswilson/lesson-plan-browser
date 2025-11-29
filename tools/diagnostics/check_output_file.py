"""Check the most recent Daniela output file for hyperlinks."""
from pathlib import Path
from docx import Document

folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')

# Get most recent output file
output_files = sorted(
    [f for f in folder.glob('Daniela_Silva_Weekly_*.docx')],
    key=lambda f: f.stat().st_mtime,
    reverse=True
)

if not output_files:
    print("❌ No output files found!")
else:
    latest = output_files[0]
    print("=" * 80)
    print(f"📄 Latest output: {latest.name}")
    print(f"   Modified: {latest.stat().st_mtime}")
    print("=" * 80)
    
    try:
        doc = Document(str(latest))
        
        # Count hyperlinks
        total_links = 0
        para_links = 0
        table_links = 0
        
        # Check paragraphs
        for para in doc.paragraphs:
            links = para._element.xpath('.//w:hyperlink')
            para_links += len(links)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        links = para._element.xpath('.//w:hyperlink')
                        table_links += len(links)
        
        total_links = para_links + table_links
        
        print(f"\n📊 Hyperlinks in output:")
        print(f"   Paragraph links: {para_links}")
        print(f"   Table links: {table_links}")
        print(f"   TOTAL: {total_links}")
        
        # Check for "Referenced Links" section
        has_ref_section = False
        for para in doc.paragraphs:
            if 'referenced links' in para.text.lower():
                has_ref_section = True
                print(f"\n✅ Found 'Referenced Links' section")
                break
        
        if not has_ref_section:
            print(f"\n⚠️  No 'Referenced Links' section found")
        
        if total_links == 0:
            print(f"\n❌ NO HYPERLINKS IN OUTPUT FILE!")
            print(f"   This means hyperlinks were NOT inserted during rendering")
        
    except Exception as e:
        print(f"❌ Error: {e}")
