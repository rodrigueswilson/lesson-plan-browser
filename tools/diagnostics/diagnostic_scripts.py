"""
Diagnostic scripts to trace hyperlink loss in the pipeline.
Run these to identify where hyperlinks are being lost.
"""

import json
from pathlib import Path
from datetime import datetime

# ============================================================================
# DIAGNOSTIC 1: Check if backend is using the fixed json_merger code
# ============================================================================

def check_json_merger_code():
    """Verify the json_merger.py file has the fix."""
    print("=" * 80)
    print("DIAGNOSTIC 1: Checking json_merger.py code")
    print("=" * 80)
    
    merger_file = Path("d:/LP/tools/json_merger.py")
    content = merger_file.read_text(encoding='utf-8')
    
    # Check for the fix
    has_fix = "all_hyperlinks.extend(lesson_json['_hyperlinks'])" in content
    has_debug = "JSON_MERGER DEBUG" in content
    
    print(f"\n✅ Fix present: {has_fix}")
    print(f"✅ Debug logging present: {has_debug}")
    
    if has_fix and has_debug:
        print("\n✅ Code file has the fix!")
    else:
        print("\n❌ Code file is missing the fix!")
    
    print("=" * 80)
    return has_fix and has_debug


# ============================================================================
# DIAGNOSTIC 2: Check for .pyc cache files
# ============================================================================

def check_pyc_cache():
    """Find all .pyc files that might be cached."""
    print("\n" + "=" * 80)
    print("DIAGNOSTIC 2: Checking for cached .pyc files")
    print("=" * 80)
    
    tools_dir = Path("d:/LP/tools")
    pyc_files = list(tools_dir.rglob("*.pyc"))
    
    if pyc_files:
        print(f"\n⚠️  Found {len(pyc_files)} .pyc files:")
        for f in pyc_files:
            print(f"   - {f}")
        print("\n❌ These cached files might be preventing the fix from loading!")
        print("   Run: Remove-Item -Path 'd:\\LP\\tools\\__pycache__' -Recurse -Force")
    else:
        print("\n✅ No .pyc cache files found")
    
    print("=" * 80)
    return len(pyc_files) == 0


# ============================================================================
# DIAGNOSTIC 3: Inspect the most recent output file
# ============================================================================

def inspect_output_file():
    """Check the most recent output file for hyperlinks."""
    print("\n" + "=" * 80)
    print("DIAGNOSTIC 3: Inspecting output file")
    print("=" * 80)
    
    from docx import Document
    
    folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')
    output_files = sorted(
        [f for f in folder.glob('Daniela_Silva_*.docx')],
        key=lambda f: f.stat().st_mtime,
        reverse=True
    )
    
    if not output_files:
        print("\n❌ No output files found!")
        return False
    
    latest = output_files[0]
    print(f"\n📄 Latest file: {latest.name}")
    print(f"   Modified: {datetime.fromtimestamp(latest.stat().st_mtime)}")
    
    doc = Document(str(latest))
    
    # Count hyperlinks
    total = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += len(para._element.xpath('.//w:hyperlink'))
    
    print(f"\n📊 Hyperlinks in output: {total}")
    
    if total == 0:
        print("❌ No hyperlinks found in output!")
        return False
    elif total == 16:
        print("✅ All 16 hyperlinks preserved!")
        return True
    else:
        print(f"⚠️  Partial preservation: {total}/16")
        return False
    
    print("=" * 80)


# ============================================================================
# DIAGNOSTIC 4: Check if json_merger is actually being called
# ============================================================================

def check_merger_import():
    """Test if we can import and call the merger with the fix."""
    print("\n" + "=" * 80)
    print("DIAGNOSTIC 4: Testing json_merger import")
    print("=" * 80)
    
    try:
        # Force reload
        import sys
        if 'tools.json_merger' in sys.modules:
            del sys.modules['tools.json_merger']
        
        from tools.json_merger import merge_lesson_jsons
        
        # Test with sample data
        test_lessons = [
            {
                'slot_number': 1,
                'subject': 'Test',
                'lesson_json': {
                    'metadata': {},
                    'days': {
                        'monday': {}, 'tuesday': {}, 'wednesday': {}, 
                        'thursday': {}, 'friday': {}
                    },
                    '_hyperlinks': [{'text': 'Test', 'url': 'http://test.com'}],
                    '_media_schema_version': '2.0'
                }
            }
        ]
        
        merged = merge_lesson_jsons(test_lessons)
        
        has_hyperlinks = '_hyperlinks' in merged
        has_version = '_media_schema_version' in merged
        
        print(f"\n✅ Import successful")
        print(f"✅ Hyperlinks in merged: {has_hyperlinks}")
        print(f"✅ Schema version in merged: {has_version}")
        
        if has_hyperlinks and has_version:
            print("\n✅ json_merger is working correctly when imported!")
            return True
        else:
            print("\n❌ json_merger is NOT preserving hyperlinks!")
            return False
            
    except Exception as e:
        print(f"\n❌ Error importing json_merger: {e}")
        return False
    
    print("=" * 80)


# ============================================================================
# DIAGNOSTIC 5: Create a test to add to batch_processor
# ============================================================================

def generate_batch_processor_patch():
    """Generate code to add to batch_processor for debugging."""
    print("\n" + "=" * 80)
    print("DIAGNOSTIC 5: Batch Processor Debug Patch")
    print("=" * 80)
    
    patch = '''
# ADD THIS AFTER LINE 557 in tools/batch_processor.py (after return lesson_json)
# This will dump the lesson_json to a file for inspection

import json
from pathlib import Path

debug_file = Path("d:/LP/DEBUG_lesson_json.json")
with open(debug_file, "w") as f:
    json.dump(lesson_json, f, indent=2, default=str)
print(f"🔍 DEBUG: Dumped lesson_json to {debug_file}")
print(f"🔍 DEBUG: lesson_json has _hyperlinks: {'_hyperlinks' in lesson_json}")
if '_hyperlinks' in lesson_json:
    print(f"🔍 DEBUG: Number of hyperlinks: {len(lesson_json['_hyperlinks'])}")
'''
    
    print("\nAdd this code to batch_processor.py to dump lesson_json:")
    print(patch)
    print("=" * 80)


# ============================================================================
# DIAGNOSTIC 6: Create a test to add to json_merger
# ============================================================================

def generate_merger_patch():
    """Generate code to add to json_merger for debugging."""
    print("\n" + "=" * 80)
    print("DIAGNOSTIC 6: JSON Merger Debug Patch")
    print("=" * 80)
    
    patch = '''
# ADD THIS AFTER LINE 113 in tools/json_merger.py (after setting media_schema_version)
# This will dump the merged JSON to a file for inspection

import json
from pathlib import Path

debug_file = Path("d:/LP/DEBUG_merged_json.json")
with open(debug_file, "w") as f:
    json.dump(merged, f, indent=2, default=str)
print(f"🔍 DEBUG: Dumped merged JSON to {debug_file}")
print(f"🔍 DEBUG: merged has _hyperlinks: {'_hyperlinks' in merged}")
if '_hyperlinks' in merged:
    print(f"🔍 DEBUG: Number of hyperlinks in merged: {len(merged['_hyperlinks'])}")
'''
    
    print("\nAdd this code to json_merger.py to dump merged JSON:")
    print(patch)
    print("=" * 80)


# ============================================================================
# RUN ALL DIAGNOSTICS
# ============================================================================

if __name__ == "__main__":
    print("\n" + "=" * 80)
    print("HYPERLINK BUG DIAGNOSTIC SUITE")
    print("=" * 80)
    
    results = {
        "Code has fix": check_json_merger_code(),
        "No cache files": check_pyc_cache(),
        "Output has hyperlinks": inspect_output_file(),
        "Merger works when imported": check_merger_import(),
    }
    
    print("\n" + "=" * 80)
    print("DIAGNOSTIC SUMMARY")
    print("=" * 80)
    
    for test, passed in results.items():
        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"{status}: {test}")
    
    print("\n" + "=" * 80)
    print("RECOMMENDED PATCHES")
    print("=" * 80)
    
    generate_batch_processor_patch()
    generate_merger_patch()
    
    print("\n" + "=" * 80)
    print("NEXT STEPS")
    print("=" * 80)
    print("""
1. Run this diagnostic script: python diagnostic_scripts.py
2. Apply the patches to batch_processor.py and json_merger.py
3. Restart the backend
4. Regenerate the lesson plan
5. Check for DEBUG_lesson_json.json and DEBUG_merged_json.json files
6. Inspect those files to see if hyperlinks are present
""")
    print("=" * 80)
