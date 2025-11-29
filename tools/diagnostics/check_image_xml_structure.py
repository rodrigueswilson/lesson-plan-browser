"""
Check the XML structure to see exactly where the image is positioned.
"""

from docx import Document

OUTPUT_PATH = "output/test_signature.docx"

doc = Document(OUTPUT_PATH)

# Find signature table
for table in doc.tables:
    table_text = "\n".join([cell.text for row in table.rows for cell in row.cells])
    if "Required Signatures" in table_text:
        # Find Teacher Signature cell
        for row in table.rows:
            for cell in row.cells:
                if "Teacher Signature:" in cell.text:
                    print("Found Teacher Signature cell")
                    print("=" * 60)
                    
                    for para_idx, para in enumerate(cell.paragraphs):
                        if "Teacher Signature:" in para.text:
                            print(f"\nParagraph {para_idx}:")
                            print(f"Full text: '{para.text}'")
                            print(f"Number of runs: {len(para.runs)}")
                            
                            for run_idx, run in enumerate(para.runs):
                                print(f"\n  Run {run_idx}:")
                                print(f"    Text: '{run.text}'")
                                
                                # Check XML structure
                                run_xml = run._element
                                
                                # Count text elements and drawing elements
                                from docx.oxml.ns import qn
                                
                                # Get all child elements
                                children = list(run_xml)
                                print(f"    Number of child elements: {len(children)}")
                                
                                for i, child in enumerate(children):
                                    tag = child.tag
                                    # Check if it's text or drawing
                                    if 't}' in tag:  # text element
                                        text_content = child.text if hasattr(child, 'text') else ''
                                        print(f"      Element {i}: TEXT - '{text_content}'")
                                    elif 'drawing' in tag.lower() or 'pict' in tag.lower():
                                        print(f"      Element {i}: IMAGE/DRAWING")
                                        # Try to get dimensions
                                        try:
                                            extent = child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}extent')
                                            if extent is not None:
                                                cx = int(extent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cx', 0))
                                                cy = int(extent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cy', 0))
                                                width_inches = cx / 914400.0
                                                height_inches = cy / 914400.0
                                                print(f"        Dimensions: {width_inches:.3f}\" x {height_inches:.3f}\"")
                                        except Exception as e:
                                            print(f"        (Could not read dimensions: {e})")
                                    
                                # Check if image comes before or after text
                                has_image = bool(run_xml.xpath('.//a:blip'))
                                if has_image:
                                    # Get the order of elements
                                    text_before_image = []
                                    for child in run_xml:
                                        if 't}' in child.tag:
                                            text_before_image.append(child.text or '')
                                        elif 'drawing' in child.tag.lower() or 'pict' in child.tag.lower():
                                            print(f"\n    *** IMAGE POSITION ANALYSIS ***")
                                            print(f"    Text elements before image: {''.join(text_before_image)}")
                                            print(f"    Full run text: '{run.text}'")
                                            if "Teacher Signature:" in ''.join(text_before_image):
                                                print(f"    [OK] Image appears AFTER 'Teacher Signature:' text")
                                            else:
                                                print(f"    [WARNING] Image position unclear")
                                            break
                    break
        break

print("\n" + "=" * 60)
print("Analysis Complete")

