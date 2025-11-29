from docx import Document
from docx.shared import Inches

# Load the template
doc = Document(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\Lesson Plan Template.docx')

print(f"Found {len(doc.tables)} table(s) in the template\n")

for i, table in enumerate(doc.tables):
    print(f"=" * 80)
    print(f"TABLE {i+1}")
    print(f"=" * 80)
    print(f"Rows: {len(table.rows)}")
    print(f"Columns: {len(table.columns)}")
    
    # Get table width if available
    try:
        if table._element.tblPr is not None:
            tbl_w = table._element.tblPr.xpath('.//w:tblW')
            if tbl_w:
                print(f"Table width type: {tbl_w[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')}")
                print(f"Table width value: {tbl_w[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')}")
    except Exception as e:
        print(f"Could not read table width: {e}")
    
    print(f"\nColumn Widths:")
    total_width = 0
    for j, column in enumerate(table.columns):
        if column.width:
            width_inches = column.width / 914400  # Convert EMUs to inches
            total_width += column.width
            print(f"  Column {j+1}: {column.width} EMUs = {width_inches:.2f} inches")
        else:
            print(f"  Column {j+1}: No width specified")
    
    if total_width > 0:
        total_inches = total_width / 914400
        print(f"\nTotal table width: {total_width} EMUs = {total_inches:.2f} inches")
    
    # Sample first row to check cell widths
    if len(table.rows) > 0:
        print(f"\nFirst Row Cell Widths:")
        for j, cell in enumerate(table.rows[0].cells):
            if cell.width:
                width_inches = cell.width / 914400
                print(f"  Cell {j+1}: {cell.width} EMUs = {width_inches:.2f} inches")
            else:
                print(f"  Cell {j+1}: No width specified")
    
    print()
