"""
Analyze only relevant teacher lesson plan files.
"""

import sys
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent))

from tools.docx_parser import DOCXParser


# Relevant teacher last names
TEACHERS = ["Lang", "Savoca", "Davies", "Morais", "Santiago", "Grande"]


def is_teacher_file(filename):
    """Check if filename belongs to a relevant teacher."""
    filename_lower = filename.lower()
    for teacher in TEACHERS:
        if teacher.lower() in filename_lower:
            return True
    return False


def analyze_file_quick(file_path):
    """Quick analysis of file structure."""
    try:
        doc = Document(file_path)
        table_count = len(doc.tables)
        
        # Check for signature
        has_signature = False
        if table_count > 0:
            last_table = doc.tables[-1]
            if last_table.rows and last_table.rows[0].cells:
                first_cell = last_table.rows[0].cells[0].text.strip().lower()
                if "signature" in first_cell or "required signatures" in first_cell:
                    has_signature = True
        
        # Count metadata tables
        metadata_count = 0
        for table in doc.tables:
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0].text.strip()
                if first_cell.startswith("Name:"):
                    metadata_count += 1
        
        return {
            "success": True,
            "table_count": table_count,
            "has_signature": has_signature,
            "metadata_count": metadata_count,
            "estimated_slots": metadata_count
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


def main():
    """Analyze teacher files."""
    base_dirs = [
        Path("F:/rodri/Documents/OneDrive/AS/Lesson Plan"),
        Path("F:/rodri/Documents/OneDrive/AS/Daniela LP")
    ]
    
    teacher_files = []
    
    # Find teacher files
    for base_dir in base_dirs:
        if not base_dir.exists():
            continue
        
        for docx_file in base_dir.rglob("*.docx"):
            if docx_file.name.startswith("~$"):
                continue
            if "old" in str(docx_file).lower():
                continue
            
            if is_teacher_file(docx_file.name):
                teacher_files.append(docx_file)
    
    print(f"Found {len(teacher_files)} teacher lesson plan files\n")
    
    # Analyze each
    patterns = {
        "1_slot": [],
        "4_slots": [],
        "5_slots": [],
        "other": []
    }
    
    for file in teacher_files:
        result = analyze_file_quick(str(file))
        
        if result["success"]:
            slots = result["estimated_slots"]
            table_count = result["table_count"]
            has_sig = result["has_signature"]
            
            print(f"{file.name}")
            print(f"  Tables: {table_count}, Slots: {slots}, Signature: {has_sig}")
            
            # Categorize
            if slots == 1:
                patterns["1_slot"].append(file.name)
            elif slots == 4:
                patterns["4_slots"].append(file.name)
            elif slots == 5:
                patterns["5_slots"].append(file.name)
            else:
                patterns["other"].append((file.name, table_count, slots))
    
    # Summary
    print(f"\n{'='*80}")
    print("SUMMARY")
    print(f"{'='*80}")
    print(f"Total teacher files: {len(teacher_files)}")
    print(f"\n1 slot files: {len(patterns['1_slot'])}")
    print(f"4 slot files: {len(patterns['4_slots'])}")
    print(f"5 slot files: {len(patterns['5_slots'])}")
    print(f"Other: {len(patterns['other'])}")
    
    if patterns["other"]:
        print(f"\nUnexpected structures:")
        for name, tables, slots in patterns["other"]:
            print(f"  {name}: {tables} tables, {slots} slots")
    
    # Conclusion
    total_standard = len(patterns["1_slot"]) + len(patterns["4_slots"]) + len(patterns["5_slots"])
    pct_standard = (total_standard / len(teacher_files) * 100) if teacher_files else 0
    
    print(f"\n{'='*80}")
    print(f"Standard layouts: {total_standard}/{len(teacher_files)} ({pct_standard:.1f}%)")
    print(f"{'='*80}")
    
    if pct_standard >= 95:
        print("\n✅ CONCLUSION: Structure is highly consistent!")
        print("   Safe to implement slot-aware extraction with simple validation.")
    elif pct_standard >= 80:
        print("\n⚠️ CONCLUSION: Mostly consistent with some variations.")
        print("   Implement with validation and fallback for edge cases.")
    else:
        print("\n❌ CONCLUSION: Too much variation.")
        print("   Need dynamic slot detection or content-based mapping.")


if __name__ == "__main__":
    main()
