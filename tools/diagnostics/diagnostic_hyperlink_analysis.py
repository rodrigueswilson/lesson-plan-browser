"""
Comprehensive Diagnostic Analysis for Hyperlink Placement
Answers all open questions with actual data from output files.
"""

from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import json
from collections import defaultdict
import re

class HyperlinkDiagnostic:
    """Diagnose hyperlink placement issues with detailed categorization."""
    
    def __init__(self):
        self.results = {
            'files_analyzed': 0,
            'total_hyperlinks': 0,
            'inline_count': 0,
            'referenced_count': 0,
            'bilingual_cells': 0,
            'english_only_cells': 0,
            'portuguese_only_cells': 0,
            'duplicate_link_texts': defaultdict(list),
            'case_variations': defaultdict(set),
            'whitespace_issues': [],
            'markdown_patterns': [],
            'failure_categories': defaultdict(int),
        }
    
    def analyze_file(self, output_path, input_path=None):
        """Analyze a single output file."""
        print(f"\n{'='*80}")
        print(f"Analyzing: {Path(output_path).name}")
        print(f"{'='*80}\n")
        
        doc = Document(output_path)
        
        # Extract inline hyperlinks
        inline_links = self._extract_inline_hyperlinks(doc)
        
        # Extract referenced links
        referenced_links = self._extract_referenced_links(doc)
        
        # Analyze bilingual content
        bilingual_stats = self._analyze_bilingual_content(doc)
        
        # Detect patterns
        self._detect_duplicate_texts(inline_links)
        self._detect_case_variations(inline_links)
        self._detect_whitespace_issues(inline_links)
        
        # Update results
        self.results['files_analyzed'] += 1
        self.results['total_hyperlinks'] += len(inline_links) + len(referenced_links)
        self.results['inline_count'] += len(inline_links)
        self.results['referenced_count'] += len(referenced_links)
        self.results['bilingual_cells'] += bilingual_stats['bilingual']
        self.results['english_only_cells'] += bilingual_stats['english_only']
        self.results['portuguese_only_cells'] += bilingual_stats['portuguese_only']
        
        # Print file summary
        total = len(inline_links) + len(referenced_links)
        if total > 0:
            inline_pct = (len(inline_links) / total) * 100
            print(f"📊 Hyperlinks: {len(inline_links)} inline, {len(referenced_links)} referenced ({inline_pct:.1f}% inline)")
        
        print(f"🌐 Bilingual: {bilingual_stats['bilingual']} cells, "
              f"English-only: {bilingual_stats['english_only']}, "
              f"Portuguese-only: {bilingual_stats['portuguese_only']}")
        
        return {
            'inline': inline_links,
            'referenced': referenced_links,
            'bilingual_stats': bilingual_stats
        }
    
    def _extract_inline_hyperlinks(self, doc):
        """Extract all inline hyperlinks from table cells."""
        hyperlinks = []
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    
                    for para in cell.paragraphs:
                        for hyperlink in para._element.xpath('.//w:hyperlink'):
                            try:
                                r_id = hyperlink.get(qn('r:id'))
                                if r_id and r_id in para.part.rels:
                                    url = para.part.rels[r_id].target_ref
                                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                    
                                    if text and url:
                                        hyperlinks.append({
                                            'text': text,
                                            'url': url,
                                            'cell_text': cell_text,
                                            'location': f"T{table_idx}R{row_idx}C{cell_idx}",
                                            'table_idx': table_idx,
                                            'row_idx': row_idx,
                                            'cell_idx': cell_idx
                                        })
                            except Exception as e:
                                pass
        
        return hyperlinks
    
    def _extract_referenced_links(self, doc):
        """Extract hyperlinks from 'Referenced Links' section."""
        referenced = []
        in_section = False
        
        for para in doc.paragraphs:
            if "Referenced Links" in para.text or "Hyperlinks" in para.text:
                in_section = True
                continue
            
            if in_section:
                for hyperlink in para._element.xpath('.//w:hyperlink'):
                    try:
                        r_id = hyperlink.get(qn('r:id'))
                        if r_id and r_id in para.part.rels:
                            url = para.part.rels[r_id].target_ref
                            text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                            
                            if text and url:
                                referenced.append({
                                    'text': text,
                                    'url': url
                                })
                    except:
                        pass
        
        return referenced
    
    def _analyze_bilingual_content(self, doc):
        """Check if cells contain bilingual content."""
        stats = {
            'bilingual': 0,
            'english_only': 0,
            'portuguese_only': 0,
            'empty': 0
        }
        
        portuguese_chars = set('ãõçáéíóúâêôàè')
        english_indicators = {'the', 'and', 'to', 'of', 'in', 'for', 'is', 'on', 'with'}
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.lower()
                    
                    if not text.strip():
                        stats['empty'] += 1
                        continue
                    
                    has_portuguese = any(char in text for char in portuguese_chars)
                    has_english = any(word in text.split() for word in english_indicators)
                    
                    if has_portuguese and has_english:
                        stats['bilingual'] += 1
                    elif has_english:
                        stats['english_only'] += 1
                    elif has_portuguese:
                        stats['portuguese_only'] += 1
        
        return stats
    
    def _detect_duplicate_texts(self, hyperlinks):
        """Find hyperlinks with identical text but different URLs."""
        text_to_urls = defaultdict(set)
        
        for link in hyperlinks:
            text_to_urls[link['text']].add(link['url'])
        
        for text, urls in text_to_urls.items():
            if len(urls) > 1:
                self.results['duplicate_link_texts'][text] = list(urls)
    
    def _detect_case_variations(self, hyperlinks):
        """Find link texts that differ only in case."""
        normalized = defaultdict(set)
        
        for link in hyperlinks:
            normalized[link['text'].lower()].add(link['text'])
        
        for lower, variations in normalized.items():
            if len(variations) > 1:
                self.results['case_variations'][lower] = variations
    
    def _detect_whitespace_issues(self, hyperlinks):
        """Find link texts with unusual whitespace."""
        for link in hyperlinks:
            text = link['text']
            
            # Check for multiple spaces
            if '  ' in text:
                self.results['whitespace_issues'].append({
                    'text': text,
                    'issue': 'multiple_spaces',
                    'location': link['location']
                })
            
            # Check for leading/trailing whitespace
            if text != text.strip():
                self.results['whitespace_issues'].append({
                    'text': repr(text),
                    'issue': 'leading_trailing',
                    'location': link['location']
                })
    
    def print_summary(self):
        """Print comprehensive diagnostic summary."""
        print(f"\n\n{'='*80}")
        print("📈 DIAGNOSTIC SUMMARY")
        print(f"{'='*80}\n")
        
        # Overall stats
        total = self.results['total_hyperlinks']
        inline = self.results['inline_count']
        referenced = self.results['referenced_count']
        
        print(f"Files analyzed: {self.results['files_analyzed']}")
        print(f"Total hyperlinks: {total}")
        print(f"  - Inline: {inline} ({inline/total*100 if total > 0 else 0:.1f}%)")
        print(f"  - Referenced: {referenced} ({referenced/total*100 if total > 0 else 0:.1f}%)")
        
        # Bilingual analysis
        print(f"\n🌐 BILINGUAL CONTENT ANALYSIS:")
        total_cells = (self.results['bilingual_cells'] + 
                      self.results['english_only_cells'] + 
                      self.results['portuguese_only_cells'])
        
        if total_cells > 0:
            print(f"  Bilingual cells: {self.results['bilingual_cells']} "
                  f"({self.results['bilingual_cells']/total_cells*100:.1f}%)")
            print(f"  English-only: {self.results['english_only_cells']} "
                  f"({self.results['english_only_cells']/total_cells*100:.1f}%)")
            print(f"  Portuguese-only: {self.results['portuguese_only_cells']} "
                  f"({self.results['portuguese_only_cells']/total_cells*100:.1f}%)")
        
        # Duplicate link texts
        if self.results['duplicate_link_texts']:
            print(f"\n⚠️  DUPLICATE LINK TEXTS (same text, different URLs):")
            for text, urls in list(self.results['duplicate_link_texts'].items())[:5]:
                print(f"  '{text}' → {len(urls)} different URLs")
                for url in urls[:2]:
                    print(f"    - {url[:60]}...")
        
        # Case variations
        if self.results['case_variations']:
            print(f"\n🔤 CASE VARIATIONS (same text, different case):")
            for lower, variations in list(self.results['case_variations'].items())[:5]:
                print(f"  {variations}")
        
        # Whitespace issues
        if self.results['whitespace_issues']:
            print(f"\n␣  WHITESPACE ISSUES:")
            for issue in self.results['whitespace_issues'][:5]:
                print(f"  {issue['issue']}: {issue['text']} at {issue['location']}")
        
        # Key findings
        print(f"\n\n🔍 KEY FINDINGS:")
        
        # Finding 1: Bilingual transformation
        english_only_pct = (self.results['english_only_cells'] / total_cells * 100 
                           if total_cells > 0 else 0)
        if english_only_pct > 50:
            print(f"  ⚠️  {english_only_pct:.1f}% of cells are English-only!")
            print(f"      → Bilingual transformation may not be working properly")
        else:
            print(f"  ✓ Bilingual transformation appears to be working")
        
        # Finding 2: Duplicate texts
        if self.results['duplicate_link_texts']:
            print(f"  ⚠️  {len(self.results['duplicate_link_texts'])} link texts appear with multiple URLs")
            print(f"      → Case-insensitive matching could place links incorrectly")
        
        # Finding 3: Case variations
        if self.results['case_variations']:
            print(f"  ℹ️  {len(self.results['case_variations'])} link texts have case variations")
            print(f"      → Case-insensitive matching would help")
        
        # Finding 4: Whitespace
        if self.results['whitespace_issues']:
            print(f"  ℹ️  {len(self.results['whitespace_issues'])} links have whitespace issues")
            print(f"      → Whitespace normalization would help")
        
        # Recommendations
        print(f"\n\n💡 RECOMMENDATIONS:")
        
        potential_improvement = 0
        
        if self.results['case_variations']:
            case_impact = len(self.results['case_variations']) / referenced * 100 if referenced > 0 else 0
            print(f"  1. Add case-insensitive matching → +{case_impact:.1f}% inline rate")
            potential_improvement += case_impact
        
        if self.results['whitespace_issues']:
            ws_impact = len(self.results['whitespace_issues']) / referenced * 100 if referenced > 0 else 0
            print(f"  2. Add whitespace normalization → +{ws_impact:.1f}% inline rate")
            potential_improvement += ws_impact
        
        if potential_improvement > 0:
            current_rate = inline / total * 100 if total > 0 else 0
            projected_rate = min(95, current_rate + potential_improvement)
            print(f"\n  Projected inline rate: {current_rate:.1f}% → {projected_rate:.1f}%")


def main():
    """Run diagnostic analysis on output files."""
    
    # Find output files
    output_dir = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan')
    
    output_files = []
    for week_dir in ['25 W41', '25 W42', '25 W43']:
        week_path = output_dir / week_dir
        if week_path.exists():
            output_files.extend(week_path.glob('Wilson_Rodrigues_Weekly_*.docx'))
    
    # Limit to recent files
    output_files = sorted(output_files, key=lambda x: x.stat().st_mtime, reverse=True)[:10]
    
    print(f"Found {len(output_files)} output files to analyze\n")
    
    # Run diagnostic
    diagnostic = HyperlinkDiagnostic()
    
    for file_path in output_files:
        try:
            diagnostic.analyze_file(str(file_path))
        except Exception as e:
            print(f"❌ Error analyzing {file_path.name}: {e}\n")
    
    # Print summary
    diagnostic.print_summary()
    
    # Save results
    results_path = Path('d:/LP/diagnostic_results.json')
    with open(results_path, 'w') as f:
        # Convert defaultdicts to regular dicts for JSON
        results_json = {
            **diagnostic.results,
            'duplicate_link_texts': dict(diagnostic.results['duplicate_link_texts']),
            'case_variations': {k: list(v) for k, v in diagnostic.results['case_variations'].items()},
        }
        json.dump(results_json, f, indent=2)
    
    print(f"\n\n💾 Results saved to: {results_path}")


if __name__ == '__main__':
    main()
