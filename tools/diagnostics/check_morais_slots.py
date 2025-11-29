from docx import Document

filepath = r"F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W44\Morais 10_27-31.docx"
doc = Document(filepath)

print("=== Morais 10_27-31.docx Structure ===\n")

for i, table in enumerate(doc.tables):
    print(f"Table {i}:")
    for row_idx, row in enumerate(table.rows[:2]):  # First 2 rows
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text and ('subject' in text.lower() or 'name:' in text.lower() or 'grade' in text.lower()):
                print(f"  Row {row_idx}, Cell {cell_idx}: {text[:80]}")
    print()
    
    if i >= 5:  # Check first 6 tables
        break

print("\nTotal tables:", len(doc.tables))
