"""
Analyze the layout of an objectives DOCX file to understand space distribution.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from collections import defaultdict

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from tools.diagnostics.objectives_layout_heights import calculate_estimated_heights


def analyze_docx_layout(docx_path: str):
    """Analyze the layout of an objectives DOCX file."""
    # Convert to Path and resolve
    docx_file = Path(docx_path).resolve()
    
    if not docx_file.exists():
        raise FileNotFoundError(f"File not found: {docx_file}")
    
    print(f"Opening file: {docx_file}")
    doc = Document(str(docx_file))
    
    # Get page dimensions
    section = doc.sections[0]
    page_width = section.page_width.inches
    page_height = section.page_height.inches
    orientation = section.orientation
    top_margin = section.top_margin.inches
    bottom_margin = section.bottom_margin.inches
    left_margin = section.left_margin.inches
    right_margin = section.right_margin.inches
    
    usable_width = page_width - left_margin - right_margin
    usable_height = page_height - top_margin - bottom_margin
    
    print("=" * 80)
    print("OBJECTIVES DOCX LAYOUT ANALYSIS")
    print("=" * 80)
    print(f"\nFile: {Path(docx_path).name}")
    print(f"\nPage Configuration:")
    print(f"  Orientation: {orientation} ({'LANDSCAPE' if orientation == 1 else 'PORTRAIT'})")
    print(f"  Page Size: {page_width}\" × {page_height}\"")
    print(f"  Margins: Top={top_margin}\", Bottom={bottom_margin}\", Left={left_margin}\", Right={right_margin}\"")
    print(f"  Usable Area: {usable_width}\" × {usable_height}\"")
    
    # Analyze paragraphs - group by objectives (each header starts a new objective)
    objectives = []
    current_objective_paras = []
    current_header = None
    
    for para in doc.paragraphs:
        para_text = ''.join(run.text for run in para.runs).strip() if para.runs else ''
        
        # Check if this is a header (starts a new objective)
        if is_header(para_text):
            # Save previous objective if exists
            if current_objective_paras and current_header:
                objectives.append({
                    'header': current_header,
                    'paragraphs': current_objective_paras
                })
            # Start new objective
            current_header = para_text
            current_objective_paras = [para]
        else:
            # Add to current objective
            if current_header is not None:
                current_objective_paras.append(para)
    
    # Add last objective
    if current_objective_paras and current_header:
        objectives.append({
            'header': current_header,
            'paragraphs': current_objective_paras
        })
    
    # If no headers found, treat entire document as one objective
    if not objectives:
        objectives.append({
            'header': 'Single Objective',
            'paragraphs': list(doc.paragraphs)
        })
    
    # Analyze each objective
    pages_data = []
    for obj_idx, obj in enumerate(objectives):
        page_num = obj_idx + 1
        page_info = analyze_page(obj['paragraphs'], page_num, usable_width, usable_height)
        page_info['header_text'] = obj['header']
        pages_data.append(page_info)
    
    # Print analysis for each page
    for page_info in pages_data:
        print_page_analysis(page_info, usable_height)
    
    # Summary statistics
    print("\n" + "=" * 80)
    print("SUMMARY STATISTICS")
    print("=" * 80)
    
    if pages_data:
        avg_student_goal_font = sum(p['student_goal_font_size'] for p in pages_data if p['student_goal_font_size']) / len([p for p in pages_data if p['student_goal_font_size']])
        avg_wida_font = sum(p['wida_font_size'] for p in pages_data if p['wida_font_size']) / len([p for p in pages_data if p['wida_font_size']])
        
        print(f"\nAverage Font Sizes:")
        print(f"  Student Goal: {avg_student_goal_font:.1f}pt")
        print(f"  WIDA Objective: {avg_wida_font:.1f}pt")
        
        print(f"\nPage Distribution:")
        print(f"  Total Pages: {len(pages_data)}")
        pages_with_both = sum(1 for p in pages_data if p['has_student_goal'] and p['has_wida'])
        pages_with_only_student = sum(1 for p in pages_data if p['has_student_goal'] and not p['has_wida'])
        pages_with_only_wida = sum(1 for p in pages_data if not p['has_student_goal'] and p['has_wida'])
        
        print(f"  Pages with both sections: {pages_with_both}")
        print(f"  Pages with only Student Goal: {pages_with_only_student}")
        print(f"  Pages with only WIDA: {pages_with_only_wida}")
        
        # Key findings - check each objective
        print(f"\n" + "=" * 80)
        print("KEY FINDINGS - OBJECTIVE FIT VERIFICATION")
        print("=" * 80)
        
        all_fit = True
        for page_info in pages_data:
            obj_num = page_info['page_num']
            header = page_info.get('header_text', f'Objective {obj_num}')
            
            if page_info['has_student_goal'] and page_info['has_wida']:
                heights = page_info['estimated_heights']
                student_h = heights.get('student_goal', 0)
                wida_h = heights.get('wida', 0)
                header_h = heights.get('header', 0)
                separator_h = heights.get('separator', 0)
                total_h = student_h + wida_h + header_h + separator_h
                
                fits = total_h <= usable_height
                if not fits:
                    all_fit = False
                
                status = "[OK] FITS" if fits else "[ERROR] OVERFLOW"
                print(f"\nObjective {obj_num}: {status}")
                print(f"  Header: {header[:60]}...")
                print(f"  Total Height: {total_h:.2f}\" / {usable_height}\" ({total_h/usable_height*100:.1f}%)")
                
                if not fits:
                    overflow = total_h - usable_height
                    print(f"  OVERFLOW: {overflow:.2f}\" ({overflow/usable_height*100:.1f}%)")
                    print(f"  Student Goal: {student_h:.2f}\" ({student_h/usable_height*100:.1f}%)")
                    print(f"  WIDA Objective: {wida_h:.2f}\" ({wida_h/usable_height*100:.1f}%)")
                else:
                    remaining = usable_height - total_h
                    print(f"  Remaining Space: {remaining:.2f}\" ({remaining/usable_height*100:.1f}%)")
                    print(f"  Student Goal: {student_h:.2f}\" ({student_h/usable_height*100:.1f}%)")
                    print(f"  WIDA Objective: {wida_h:.2f}\" ({wida_h/usable_height*100:.1f}%)")
                
                student_paras = page_info['paragraphs_by_type']['student_goal']
                wida_paras = page_info['paragraphs_by_type']['wida']
                
                print(f"\n  Details:")
                print(f"    - Student Goal: {page_info['student_goal_font_size']}pt, {len(student_paras)} paragraphs, {len(page_info['student_goal_text'])} chars")
                print(f"    - WIDA Objective: {page_info['wida_font_size']}pt, {len(wida_paras)} paragraphs, {len(page_info['wida_text'])} chars")
            
            elif page_info['has_student_goal'] and not page_info['has_wida']:
                all_fit = False
                print(f"\nObjective {obj_num}: [ERROR] INCOMPLETE")
                print(f"  Header: {header[:60]}...")
                print(f"  Student Goal present but WIDA Objective missing or on next page")
            
            elif not page_info['has_student_goal'] and page_info['has_wida']:
                all_fit = False
                print(f"\nObjective {obj_num}: [ERROR] INCOMPLETE")
                print(f"  Header: {header[:60]}...")
                print(f"  WIDA Objective present but Student Goal missing or on previous page")
            else:
                all_fit = False
                print(f"\nObjective {obj_num}: [ERROR] NO CONTENT")
                print(f"  Header: {header[:60]}...")
                print(f"  No Student Goal or WIDA Objective found")
        
        # Final summary
        print(f"\n" + "=" * 80)
        print("FINAL VERIFICATION RESULT")
        print("=" * 80)
        if all_fit:
            print(f"[SUCCESS] All {len(pages_data)} objective(s) fit on their respective page(s)!")
        else:
            print(f"[FAILURE] Some objectives do not fit on their page(s). See details above.")


def analyze_page(paragraphs, page_num, usable_width, usable_height):
    """Analyze a single page's content."""
    page_info = {
        'page_num': page_num,
        'total_paragraphs': len(paragraphs),
        'has_header': False,
        'has_student_goal': False,
        'has_separator': False,
        'has_wida': False,
        'header_text': '',
        'student_goal_text': '',
        'wida_text': '',
        'student_goal_font_size': None,
        'wida_font_size': None,
        'student_goal_lines': 0,
        'wida_lines': 0,
        'paragraphs_by_type': defaultdict(list),
        'estimated_heights': {},
    }
    
    current_section = None
    
    for para in paragraphs:
        if not para.runs:
            continue
        
        # Get text and formatting
        para_text = ''.join(run.text for run in para.runs).strip()
        if not para_text:
            continue
        
        # Check for page break
        if '\f' in para_text or '\x0c' in para_text:
            continue
        
        # Analyze paragraph
        para_info = analyze_paragraph(para)
        
        # Determine section type
        if is_header(para_text):
            current_section = 'header'
            page_info['has_header'] = True
            page_info['header_text'] = para_text
            page_info['paragraphs_by_type']['header'].append(para_info)
        elif is_separator(para):
            current_section = 'separator'
            page_info['has_separator'] = True
            page_info['paragraphs_by_type']['separator'].append(para_info)
        elif 'WIDA/Bilingual:' in para_text or current_section == 'wida':
            current_section = 'wida'
            page_info['has_wida'] = True
            if para_text != 'WIDA/Bilingual:':
                page_info['wida_text'] += para_text + ' '
            page_info['paragraphs_by_type']['wida'].append(para_info)
            if para_info.get('font_size'):
                page_info['wida_font_size'] = para_info['font_size']
        else:
            # Assume Student Goal if not header or separator
            if current_section != 'wida':
                current_section = 'student_goal'
                page_info['has_student_goal'] = True
                page_info['student_goal_text'] += para_text + ' '
                page_info['paragraphs_by_type']['student_goal'].append(para_info)
                if para_info.get('font_size') and not page_info['student_goal_font_size']:
                    page_info['student_goal_font_size'] = para_info['font_size']
    
    # Calculate estimated heights
    page_info['estimated_heights'] = calculate_estimated_heights(
        page_info, usable_height
    )
    
    return page_info


def analyze_paragraph(para):
    """Analyze a single paragraph."""
    info = {
        'text': ''.join(run.text for run in para.runs),
        'font_size': None,
        'font_name': None,
        'is_bold': False,
        'color': None,
        'line_spacing': None,
        'space_before': para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
        'space_after': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
    }
    
    # Get formatting from first run
    if para.runs:
        run = para.runs[0]
        if run.font.size:
            info['font_size'] = run.font.size.pt
        if run.font.name:
            info['font_name'] = run.font.name
        if run.font.bold is not None:
            info['is_bold'] = run.font.bold
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            # RGBColor has hex value, convert to RGB components
            hex_value = rgb
            if isinstance(hex_value, int):
                r = (hex_value >> 16) & 0xFF
                g = (hex_value >> 8) & 0xFF
                b = hex_value & 0xFF
                info['color'] = f"RGB({r}, {g}, {b})"
            else:
                info['color'] = str(hex_value)
    
    if para.paragraph_format.line_spacing:
        if para.paragraph_format.line_spacing is not None:
            info['line_spacing'] = para.paragraph_format.line_spacing
    
    return info


def is_header(text):
    """Check if paragraph is a header."""
    return '|' in text and ('Grade' in text or 'Subject' in text)


def is_separator(para):
    """Check if paragraph is a separator line."""
    # Check for border formatting
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
    return pBdr is not None


def print_page_analysis(page_info, usable_height):
    """Print detailed analysis for a page."""
    print(f"\n{'=' * 80}")
    print(f"OBJECTIVE {page_info['page_num']} ANALYSIS")
    print(f"{'=' * 80}")
    
    header = page_info.get('header_text', 'N/A')
    print(f"\nHeader: {header[:80]}...")
    
    print(f"\nSections Found:")
    print(f"  Header: {'Yes' if page_info['has_header'] else 'No'}")
    print(f"  Student Goal: {'Yes' if page_info['has_student_goal'] else 'No'}")
    print(f"  Separator: {'Yes' if page_info['has_separator'] else 'No'}")
    print(f"  WIDA Objective: {'Yes' if page_info['has_wida'] else 'No'}")
    
    # Header details
    if page_info['has_header']:
        header_paras = page_info['paragraphs_by_type']['header']
        if header_paras:
            para_info = header_paras[0]
            print(f"\nHeader Details:")
            print(f"  Font Size: {para_info.get('font_size', 'N/A')}pt")
            print(f"  Font: {para_info.get('font_name', 'N/A')}")
    
    # Student Goal details
    if page_info['has_student_goal']:
        print(f"\nStudent Goal Section:")
        student_paras = page_info['paragraphs_by_type']['student_goal']
        print(f"  Paragraphs: {len(student_paras)}")
        print(f"  Font Size: {page_info['student_goal_font_size']}pt" if page_info['student_goal_font_size'] else "  Font Size: N/A")
        
        if student_paras:
            para_info = student_paras[0]
            print(f"  Font: {para_info.get('font_name', 'N/A')}")
            print(f"  Bold: {para_info.get('is_bold', False)}")
            print(f"  Line Spacing: {para_info.get('line_spacing', 'N/A')}")
            text_preview = page_info['student_goal_text'][:100].replace('\n', ' ')
            print(f"  Text Preview: {text_preview}...")
            print(f"  Text Length: {len(page_info['student_goal_text'])} characters")
            
            # Analyze paragraph structure
            non_empty_paras = [p for p in student_paras if p['text'].strip()]
            print(f"  Non-empty paragraphs: {len(non_empty_paras)}")
            if len(non_empty_paras) > 1:
                avg_para_length = sum(len(p['text']) for p in non_empty_paras) / len(non_empty_paras)
                print(f"  Average paragraph length: {avg_para_length:.0f} characters")
                print(f"  [NOTE] Text appears to be split across multiple paragraphs")
            
            # Estimate actual wrapped lines more accurately
            combined_text = ' '.join(p['text'] for p in non_empty_paras if p['text'].strip())
            font_size = page_info['student_goal_font_size'] or 83
            # More accurate: account for word wrapping
            words = combined_text.split()
            chars_per_line = (9.4 * 72) / (font_size * 0.6) if font_size > 0 else 60
            # Estimate: average word length + space = ~6 chars
            words_per_line = max(1, chars_per_line / 6)
            estimated_wrapped_lines = max(1, len(words) / words_per_line)
            print(f"  Estimated wrapped lines: {estimated_wrapped_lines:.1f}")
            estimated_height = (font_size * 1.25 * estimated_wrapped_lines) / 72
            print(f"  Estimated height (based on wrapped lines): {estimated_height:.3f}\"")
    
    # Separator details
    if page_info['has_separator']:
        print(f"\nSeparator:")
        separator_paras = page_info['paragraphs_by_type']['separator']
        if separator_paras:
            para_info = separator_paras[0]
            print(f"  Space Before: {para_info.get('space_before', 0)}pt")
            print(f"  Space After: {para_info.get('space_after', 0)}pt")
    
    # WIDA details
    if page_info['has_wida']:
        print(f"\nWIDA Objective Section:")
        wida_paras = page_info['paragraphs_by_type']['wida']
        print(f"  Paragraphs: {len(wida_paras)}")
        print(f"  Font Size: {page_info['wida_font_size']}pt" if page_info['wida_font_size'] else "  Font Size: N/A")
        
        if wida_paras:
            para_info = wida_paras[0]
            print(f"  Font: {para_info.get('font_name', 'N/A')}")
            print(f"  Color: {para_info.get('color', 'N/A')}")
            text_preview = page_info['wida_text'][:100].replace('\n', ' ')
            print(f"  Text Preview: {text_preview}...")
            print(f"  Text Length: {len(page_info['wida_text'])} characters")
    
    # Estimated heights
    heights = page_info['estimated_heights']
    if heights:
        print(f"\nEstimated Heights:")
        if 'header' in heights:
            print(f"  Header: {heights['header']:.3f}\" ({heights['header']/usable_height*100:.1f}% of page)")
        if 'student_goal' in heights:
            print(f"  Student Goal: {heights['student_goal']:.3f}\" ({heights['student_goal']/usable_height*100:.1f}% of page)")
        if 'separator' in heights:
            print(f"  Separator: {heights['separator']:.3f}\" ({heights['separator']/usable_height*100:.1f}% of page)")
        if 'wida' in heights:
            print(f"  WIDA Objective: {heights['wida']:.3f}\" ({heights['wida']/usable_height*100:.1f}% of page)")
        
        if 'total_estimated' in heights:
            print(f"\n  Total Estimated: {heights['total_estimated']:.3f}\" / {usable_height}\" ({heights['percentage_of_page']:.1f}%)")
            
            if heights['total_estimated'] > usable_height:
                overflow = heights['total_estimated'] - usable_height
                print(f"  [WARNING] OVERFLOW: {overflow:.3f}\" ({overflow/usable_height*100:.1f}%)")
                print(f"     This would push content to the next page!")
            else:
                remaining = usable_height - heights['total_estimated']
                print(f"  [OK] Remaining Space: {remaining:.3f}\" ({remaining/usable_height*100:.1f}%)")
    
    # Check if both sections are on same page
    if page_info['has_student_goal'] and page_info['has_wida']:
        print(f"\n[OK] Both Student Goal and WIDA Objective are on the same page")
    elif page_info['has_student_goal'] and not page_info['has_wida']:
        print(f"\n[WARNING] Student Goal is on this page, but WIDA Objective is missing or on next page")
    elif not page_info['has_student_goal'] and page_info['has_wida']:
        print(f"\n[WARNING] WIDA Objective is on this page, but Student Goal is missing or on previous page")


def main():
    """Main function."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Analyze objectives DOCX layout')
    parser.add_argument('file', type=str, help='Path to DOCX file to analyze')
    
    args = parser.parse_args()
    
    docx_path = args.file
    
    if not Path(docx_path).exists():
        print(f"ERROR: File not found: {docx_path}")
        return
    
    try:
        analyze_docx_layout(docx_path)
    except Exception as e:
        print(f"ERROR: Failed to analyze file: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    # If no arguments, use default test file
    if len(sys.argv) == 1:
        default_file = r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W47\Wilson_Rodrigues_W47_11-17-11-21_objectives_TEST.docx"
        if Path(default_file).exists():
            print(f"Using default file: {Path(default_file).name}\n")
            analyze_docx_layout(default_file)
        else:
            print("ERROR: Default file not found. Please provide file path as argument.")
            print("Usage: python analyze_objectives_layout.py <path_to_docx_file>")
    else:
        main()

