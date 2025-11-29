"""Analyze the input file hyperlinks in detail."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

input_file = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Morais 10_20_25 - 10_24_25.docx')

print("=" * 80)
print("INPUT FILE HYPERLINK ANALYSIS")
print("=" * 80)

doc = Document(str(input_file))

# Focus on table 1 (the daily plans table)
table = doc.tables[1]

print(f"\nTable 1 has {len(table.rows)} rows, {len(table.rows[0].cells)} columns")

# Extract all hyperlinks from row 3 (where we see the duplicates)
print("\n" + "=" * 80)
print("ROW 3 ANALYSIS (where duplicates appear)")
print("=" * 80)

row = table.rows[3]
for cell_idx, cell in enumerate(row.cells):
    hyperlinks = []
    for para in cell.paragraphs:
        for hyperlink in para._element.xpath('.//w:hyperlink'):
            r_id = hyperlink.get(qn('r:id'))
            if r_id and r_id in para.part.rels:
                url = para.part.rels[r_id].target_ref
                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                hyperlinks.append({'text': text, 'url': url})
    
    if hyperlinks:
        print(f"\nCell {cell_idx}:")
        print(f"  Cell text: {cell.text[:80]}")
        for hl in hyperlinks:
            print(f"  Link: {hl['text']}")
            print(f"  URL: {hl['url'][:60]}...")

print("\n" + "=" * 80)
print("ALL HYPERLINKS IN TABLE 1")
print("=" * 80)

all_links = []
for row_idx, row in enumerate(table.rows):
    for cell_idx, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for hyperlink in para._element.xpath('.//w:hyperlink'):
                r_id = hyperlink.get(qn('r:id'))
                if r_id and r_id in para.part.rels:
                    url = para.part.rels[r_id].target_ref
                    text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                    all_links.append({
                        'row': row_idx,
                        'cell': cell_idx,
                        'text': text,
                        'url': url
                    })

# Group by text
from collections import defaultdict
by_text = defaultdict(list)
for link in all_links:
    by_text[link['text']].append(link)

print(f"\nFound {len(all_links)} total hyperlinks")
print(f"Unique texts: {len(by_text)}")

print("\nDuplicate hyperlink texts:")
for text, links in by_text.items():
    if len(links) > 1:
        print(f"\n'{text}' appears {len(links)} times:")
        for link in links:
            print(f"  Row {link['row']}, Cell {link['cell']}")
            print(f"    URL: {link['url'][:60]}...")
