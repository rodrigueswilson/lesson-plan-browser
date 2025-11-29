"""Check what the 'missing' links actually are."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

input_file = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Morais 10_20_25 - 10_24_25.docx')
output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Daniela_Silva_Weekly_W43_10-20-10-24_20251019_185016.docx')

print("=" * 80)
print("CHECKING 'MISSING' LINKS")
print("=" * 80)

# Extract from input
print("\n📥 INPUT FILE - All hyperlinks:")
input_doc = Document(str(input_file))
input_links = []

for table_idx, table in enumerate(input_doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for hyperlink in para._element.xpath('.//w:hyperlink'):
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id and r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        input_links.append({
                            'text': text,
                            'url': url,
                            'table': table_idx,
                            'row': row_idx,
                            'cell': cell_idx
                        })
                        print(f"\n{len(input_links)}. {text}")
                        print(f"   URL: {url[:80]}")
                        print(f"   Position: T{table_idx}R{row_idx}C{cell_idx}")

# Extract from output
print("\n" + "=" * 80)
print("📤 OUTPUT FILE - Hyperlinks with '2nd Grade' or 'Unit Description':")
output_doc = Document(str(output_file))
output_links = []

for table_idx, table in enumerate(output_doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for hyperlink in para._element.xpath('.//w:hyperlink'):
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id and r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        
                        # Filter for relevant links
                        if '2nd Grade' in text or 'Unit Description' in text or 'Science' in text:
                            output_links.append({
                                'text': text,
                                'url': url,
                                'table': table_idx,
                                'row': row_idx,
                                'cell': cell_idx
                            })
                            print(f"\n{len(output_links)}. {text}")
                            print(f"   URL: {url[:80]}")
                            print(f"   Position: T{table_idx}R{row_idx}C{cell_idx}")

print("\n" + "=" * 80)
print("ANALYSIS")
print("=" * 80)
print(f"\nTotal input links: {len(input_links)}")
print(f"Output links with '2nd Grade/Science/Unit Description': {len(output_links)}")

# Check for exact matches
print("\n" + "=" * 80)
print("MATCHING INPUT TO OUTPUT")
print("=" * 80)

for inp in input_links:
    found = False
    for out in output_links:
        if inp['url'] == out['url'] and inp['text'] == out['text']:
            if inp['row'] == out['row'] and inp['cell'] == out['cell']:
                print(f"\n✅ MATCHED: {inp['text'][:50]}")
                print(f"   Input:  T{inp['table']}R{inp['row']}C{inp['cell']}")
                print(f"   Output: T{out['table']}R{out['row']}C{out['cell']}")
                found = True
                break
    
    if not found:
        print(f"\n❌ NOT MATCHED: {inp['text'][:50]}")
        print(f"   Input position: T{inp['table']}R{inp['row']}C{inp['cell']}")
        print(f"   URL: {inp['url'][:60]}")
