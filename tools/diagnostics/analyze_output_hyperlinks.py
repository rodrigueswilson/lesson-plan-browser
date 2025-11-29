"""
Analyze hyperlinks in OUTPUT files (after LLM transformation).
This will show us what actually happens to hyperlinks after bilingual transformation.
"""

from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

def analyze_output_file(doc_path):
    """Analyze a single output file for hyperlink patterns."""
    print(f"\n{'='*80}")
    print(f"Analyzing OUTPUT: {Path(doc_path).name}")
    print(f"{'='*80}\n")
    
    doc = Document(doc_path)
    
    # Find hyperlinks in tables
    table_hyperlinks = []
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                for para in cell.paragraphs:
                    for hyperlink in para._element.xpath('.//w:hyperlink'):
                        try:
                            r_id = hyperlink.get(qn('r:id'))
                            if r_id and r_id in para.part.rels:
                                url = para.part.rels[r_id].target_ref
                                text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                if text and url:
                                    table_hyperlinks.append({
                                        'text': text,
                                        'url': url,
                                        'cell_text': cell_text[:200],
                                        'location': f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}"
                                    })
                        except Exception as e:
                            print(f"Error: {e}")
    
    # Find "Referenced Links" section
    referenced_links = []
    in_referenced_section = False
    for para in doc.paragraphs:
        if "Referenced Links" in para.text or "Hyperlinks" in para.text:
            in_referenced_section = True
            continue
        
        if in_referenced_section:
            for hyperlink in para._element.xpath('.//w:hyperlink'):
                try:
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id and r_id in para.part.rels:
                        url = para.part.rels[r_id].target_ref
                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                        if text and url:
                            referenced_links.append({
                                'text': text,
                                'url': url
                            })
                except:
                    pass
    
    # Print results
    print(f"📊 SUMMARY:")
    print(f"  Inline (in table cells): {len(table_hyperlinks)}")
    print(f"  Referenced Links section: {len(referenced_links)}")
    print(f"  Total: {len(table_hyperlinks) + len(referenced_links)}")
    
    if len(table_hyperlinks) + len(referenced_links) > 0:
        inline_pct = (len(table_hyperlinks) / (len(table_hyperlinks) + len(referenced_links))) * 100
        print(f"  Inline placement rate: {inline_pct:.1f}%")
    
    # Show examples of inline hyperlinks
    if table_hyperlinks:
        print(f"\n📋 INLINE HYPERLINKS (first 5):")
        for i, link in enumerate(table_hyperlinks[:5], 1):
            print(f"\n  [{i}] Link text: \"{link['text']}\"")
            print(f"      Location: {link['location']}")
            print(f"      Cell content: \"{link['cell_text'][:100]}...\"")
            
            # Check if Portuguese is present
            portuguese_chars = any(c in link['cell_text'] for c in ['ã', 'õ', 'ç', 'á', 'é', 'í', 'ó', 'ú'])
            print(f"      Has Portuguese: {'✓' if portuguese_chars else '✗'}")
            
            # Check if link text is in English
            english_indicators = ['lesson', 'unit', 'day', 'activity', 'guide']
            is_english = any(word in link['text'].lower() for word in english_indicators)
            print(f"      Link text in English: {'✓' if is_english else '✗'}")
    
    # Show examples of referenced links
    if referenced_links:
        print(f"\n📎 REFERENCED LINKS SECTION (first 5):")
        for i, link in enumerate(referenced_links[:5], 1):
            print(f"  [{i}] \"{link['text']}\"")
    
    return {
        'filename': Path(doc_path).name,
        'inline': len(table_hyperlinks),
        'referenced': len(referenced_links),
        'total': len(table_hyperlinks) + len(referenced_links),
        'inline_examples': table_hyperlinks[:3]
    }

if __name__ == '__main__':
    # Analyze recent OUTPUT files
    output_files = [
        r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_025901.docx',
        r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W42\Wilson_Rodrigues_Weekly_W42_10-13-10-17.docx',
        r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W41\Wilson_Rodrigues_Weekly_W41_10-06-10-10.docx',
    ]
    
    results = []
    for file_path in output_files:
        if Path(file_path).exists():
            try:
                result = analyze_output_file(file_path)
                results.append(result)
            except Exception as e:
                print(f"\n❌ Error analyzing {Path(file_path).name}: {e}\n")
    
    # Overall summary
    print(f"\n\n{'='*80}")
    print("📈 OVERALL SUMMARY (OUTPUT FILES)")
    print(f"{'='*80}")
    
    total_inline = sum(r['inline'] for r in results)
    total_referenced = sum(r['referenced'] for r in results)
    total_all = total_inline + total_referenced
    
    print(f"Files analyzed: {len(results)}")
    print(f"Total hyperlinks: {total_all}")
    print(f"  - Inline (in cells): {total_inline} ({(total_inline/total_all*100) if total_all > 0 else 0:.1f}%)")
    print(f"  - Referenced Links: {total_referenced} ({(total_referenced/total_all*100) if total_all > 0 else 0:.1f}%)")
    
    print(f"\n🔍 KEY FINDING:")
    if total_inline > 0:
        print("✓ Some hyperlinks ARE being placed inline!")
        print("  → Check examples above to see if link text stayed in English")
    else:
        print("✗ NO hyperlinks placed inline - all in 'Referenced Links'")
        print("  → This confirms the 10-20% placement rate issue")
