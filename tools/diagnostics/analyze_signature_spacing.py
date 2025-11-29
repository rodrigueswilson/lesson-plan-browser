"""Analyze the spacing in the signature table to determine correct alignment."""

from docx import Document

# Load template
doc = Document("input/Lesson Plan Template SY'25-26.docx")

# Find signature table
signature_table = None
for table in doc.tables:
    table_text = "\n".join([cell.text for row in table.rows for cell in row.cells])
    if "Required Signatures" in table_text:
        signature_table = table
        break

if not signature_table:
    print("Signature table not found!")
    exit(1)

print("=" * 80)
print("SIGNATURE TABLE ANALYSIS")
print("=" * 80)

# Analyze Administrator row
for row in signature_table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            para_text = para.text
            if "Administrator Signature:" in para_text:
                print("\n[ADMINISTRATOR ROW]")
                print(f"Full text: {repr(para_text)}")
                print(f"Total length: {len(para_text)}")
                
                admin_pos = para_text.find("Administrator Signature:")
                date_pos = para_text.find("Date:", admin_pos)
                
                print(f"\nPosition of 'Administrator Signature:': {admin_pos}")
                print(f"Length of 'Administrator Signature:': {len('Administrator Signature:')}")
                print(f"Position of 'Date:': {date_pos}")
                
                # Calculate spacing between them
                spacing_start = admin_pos + len("Administrator Signature:")
                spacing_end = date_pos
                spacing_chars = spacing_end - spacing_start
                spacing_text = para_text[spacing_start:spacing_end]
                
                print(f"\nSpacing between 'Administrator Signature:' and 'Date:':")
                print(f"  Characters: {spacing_chars}")
                print(f"  Content: {repr(spacing_text)}")
                print(f"  (Underscores: {spacing_text.count('_')}, Spaces: {spacing_text.count(' ')})")
                
                print(f"\n'Date:' starts at position: {date_pos}")

print("\n" + "=" * 80)
print("CALCULATION FOR TEACHER ROW")
print("=" * 80)

teacher_label = "Teacher Signature:"
admin_label = "Administrator Signature:"

print(f"\nAdministrator label length: {len(admin_label)} chars")
print(f"Teacher label length: {len(teacher_label)} chars")
print(f"Difference: {len(admin_label) - len(teacher_label)} chars")

print(f"\nTo align 'Date:' labels:")
print(f"  Administrator 'Date:' starts at position: {date_pos}")
print(f"  Teacher needs to reach position: {date_pos}")
print(f"  After 'Teacher Signature:' ({len(teacher_label)} chars), need: {date_pos - len(teacher_label)} more chars")

# Account for signature/name width
print(f"\n[WITH SIGNATURE IMAGE]")
print(f"  Signature image takes visual space (approximately 15-20 chars equivalent)")
print(f"  Recommended spacing: {date_pos - len(teacher_label) - 18} spaces")

print(f"\n[WITH NAME ONLY (e.g., 'Wilson Rodrigues' = 16 chars)]")
print(f"  Name takes: 16 chars")
print(f"  Recommended spacing: {date_pos - len(teacher_label) - 16} spaces")

print("\n" + "=" * 80)
