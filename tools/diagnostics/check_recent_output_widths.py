"""Check table widths in recently generated output files."""

from docx import Document
from docx.shared import Twips
from pathlib import Path

# Check output directory for recent files
output_dir = Path("output")
docx_files = list(output_dir.glob("*.docx"))

# Sort by modification time (most recent first)
docx_files.sort(key=lambda x: x.stat().st_mtime, reverse=True)

print("=" * 80)
print("CHECKING TABLE WIDTHS IN RECENT OUTPUT FILES")
print("=" * 80)

# Check the 5 most recent files
for file_path in docx_files[:10]:
    print(f"\n{'=' * 80}")
    print(f"FILE: {file_path.name}")
    print(f"{'=' * 80}")
    
    try:
        doc = Document(str(file_path))
        
        # Get page setup
        section = doc.sections[0]
        page_width = section.page_width.inches
        left_margin = section.left_margin.inches
        right_margin = section.right_margin.inches
        available_width = page_width - left_margin - right_margin
        
        print(f"\nPage Setup:")
        print(f"  Page width: {page_width:.2f} inches")
        print(f"  Left margin: {left_margin:.2f} inches")
        print(f"  Right margin: {right_margin:.2f} inches")
        print(f"  Available width: {available_width:.2f} inches")
        
        # Check each table
        print(f"\nTables ({len(doc.tables)} total):")
        for i, table in enumerate(doc.tables):
            # Get table width from XML
            tbl_element = table._element
            tblPr = tbl_element.tblPr
            tblW = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
            
            if tblW is not None:
                width_type = tblW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                width_value = tblW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                
                if width_type == 'dxa':
                    width_twips = int(width_value)
                    width_inches = Twips(width_twips).inches
                    
                    # Check if width matches available width
                    width_diff = abs(width_inches - available_width)
                    status = "✓ OK" if width_diff < 0.1 else "✗ MISMATCH"
                    
                    print(f"  Table {i+1}: {width_inches:.2f}\" ({len(table.columns)} cols) - {status}")
                    if width_diff >= 0.1:
                        print(f"           Expected: {available_width:.2f}\", Difference: {width_diff:.2f}\"")
                else:
                    print(f"  Table {i+1}: {width_type} type ({len(table.columns)} cols)")
            else:
                print(f"  Table {i+1}: No width specified ({len(table.columns)} cols)")
        
    except Exception as e:
        print(f"  ERROR: {e}")

print("\n" + "=" * 80)
print("SUMMARY")
print("=" * 80)
print("\nExpected behavior:")
print("  - All tables should be ~10.0 inches (landscape template)")
print("  - Width type should be 'dxa' (twentieths of a point)")
print("\nIf tables show different widths:")
print("  - Check if docx_renderer.py is calling normalize_all_tables()")
print("  - Verify available_width calculation is correct")
