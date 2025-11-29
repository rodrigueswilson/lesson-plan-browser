"""Check table widths in specific output files."""

from docx import Document
from docx.shared import Twips
from pathlib import Path

# Files to check
files_to_check = [
    "output/Maria Garcia_Lesson plan_W06_10-07-10-11.docx",  # Production file
    "output/test_merged.docx",  # Test merged file
    "output/test_full_flow.docx",  # Full flow test
    "output/test_combined_preservation.docx",  # Combined preservation
]

print("=" * 80)
print("CHECKING TABLE WIDTHS IN SPECIFIC FILES")
print("=" * 80)

for file_path_str in files_to_check:
    file_path = Path(file_path_str)
    
    if not file_path.exists():
        print(f"\n✗ SKIPPED: {file_path.name} (not found)")
        continue
    
    print(f"\n{'=' * 80}")
    print(f"FILE: {file_path.name}")
    print(f"{'=' * 80}")
    
    try:
        doc = Document(str(file_path))
        
        # Get page setup
        section = doc.sections[0]
        page_width = section.page_width.inches
        left_margin = section.left_margin.inches
        right_margin = section.right_margin.inches
        available_width = page_width - left_margin - right_margin
        
        print(f"\nPage Setup:")
        print(f"  Page width: {page_width:.2f} inches")
        print(f"  Left margin: {left_margin:.2f} inches")
        print(f"  Right margin: {right_margin:.2f} inches")
        print(f"  Available width: {available_width:.2f} inches")
        
        # Check each table
        print(f"\nTables ({len(doc.tables)} total):")
        
        width_issues = []
        for i, table in enumerate(doc.tables):
            # Get table width from XML
            tbl_element = table._element
            tblPr = tbl_element.tblPr
            tblW = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
            
            if tblW is not None:
                width_type = tblW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                width_value = tblW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                
                if width_type == 'dxa':
                    width_twips = int(width_value)
                    width_inches = Twips(width_twips).inches
                    
                    # Check if width matches available width
                    width_diff = abs(width_inches - available_width)
                    status = "✓ OK" if width_diff < 0.1 else "✗ MISMATCH"
                    
                    print(f"  Table {i+1}: {width_inches:.2f}\" ({len(table.columns)} cols) - {status}")
                    if width_diff >= 0.1:
                        print(f"           Expected: {available_width:.2f}\", Difference: {width_diff:.2f}\"")
                        width_issues.append((i+1, width_inches, available_width, width_diff))
                elif width_type == 'auto':
                    print(f"  Table {i+1}: AUTO-SIZED ({len(table.columns)} cols) - ✗ PROBLEM")
                    width_issues.append((i+1, "auto", available_width, "N/A"))
                else:
                    print(f"  Table {i+1}: {width_type} type ({len(table.columns)} cols)")
            else:
                print(f"  Table {i+1}: No width specified ({len(table.columns)} cols) - ✗ PROBLEM")
                width_issues.append((i+1, "unspecified", available_width, "N/A"))
        
        # Summary for this file
        if width_issues:
            print(f"\n  ⚠️  ISSUES FOUND: {len(width_issues)} table(s) with width problems")
            for table_num, actual, expected, diff in width_issues:
                print(f"      Table {table_num}: {actual} (expected {expected})")
        else:
            print(f"\n  ✓ All tables have correct widths")
        
    except Exception as e:
        print(f"  ERROR: {e}")
        import traceback
        traceback.print_exc()

print("\n" + "=" * 80)
print("DIAGNOSIS")
print("=" * 80)
print("\nIf width issues are found:")
print("  1. Check if normalize_all_tables() is being called in the rendering pipeline")
print("  2. Verify the available_width calculation is correct")
print("  3. Check if tables are being modified after normalization")
