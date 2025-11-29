"""
Simple script to count links per slot.
Assumes structure: Metadata table, then Daily Plans table, repeat.
"""

from docx import Document
from pathlib import Path

output_file = Path(r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_133850.docx')

doc = Document(output_file)

slots = []
current_slot = None

for idx, table in enumerate(doc.tables):
    # Check if metadata table
    if table.rows and table.rows[0].cells:
        first_cell = table.rows[0].cells[0].text.strip()
        # Also check second cell for daily plans tables
        second_cell = table.rows[0].cells[1].text.strip() if len(table.rows[0].cells) > 1 else ""
        
        if "Name:" in first_cell:
            # Extract subject and grade
            subject = ""
            grade = ""
            for cell in table.rows[0].cells:
                text = cell.text.strip()
                if "Subject:" in text:
                    subject = text.split("Subject:")[-1].strip()
                elif "Grade:" in text:
                    grade = text.split("Grade:")[-1].strip()
            
            slot_name = f"Grade {grade} - {subject}"
            current_slot = {
                'name': slot_name,
                'metadata_table': idx,
                'daily_table': None,
                'inline_links': 0,
                'fallback_links': 0
            }
            slots.append(current_slot)
        
        elif ("MONDAY" in second_cell.upper() or "Monday" in second_cell) and current_slot:
            # This is the daily plans table for current slot
            current_slot['daily_table'] = idx
            
            # Count links
            link_count = 0
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        hyperlinks = para._element.xpath('.//w:hyperlink')
                        link_count += len(hyperlinks)
            
            current_slot['inline_links'] = link_count
        
        elif ("Referenced Links" in first_cell or "Referenced Media" in first_cell) and current_slot:
            # This is fallback section for current slot
            link_count = 0
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        hyperlinks = para._element.xpath('.//w:hyperlink')
                        link_count += len(hyperlinks)
            
            current_slot['fallback_links'] = link_count

# Print results
print("="*100)
print("HYPERLINK PLACEMENT PER SLOT")
print("="*100)
print()

print(f"{'Slot':<40} {'Inline':<10} {'Fallback':<10} {'Total':<10} {'Inline %':<10}")
print("-"*90)

total_inline = 0
total_fallback = 0

for slot in slots:
    inline = slot['inline_links']
    fallback = slot['fallback_links']
    total = inline + fallback
    inline_pct = (inline / total * 100) if total > 0 else 0
    
    print(f"{slot['name']:<40} {inline:<10} {fallback:<10} {total:<10} {inline_pct:<10.1f}%")
    
    total_inline += inline
    total_fallback += fallback

print("-"*90)
total_all = total_inline + total_fallback
overall_pct = (total_inline / total_all * 100) if total_all > 0 else 0
print(f"{'TOTAL':<40} {total_inline:<10} {total_fallback:<10} {total_all:<10} {overall_pct:<10.1f}%")

print()
print()

# Show fallback details
print("="*100)
print("FALLBACK LINKS DETAILS")
print("="*100)
print()

for slot in slots:
    if slot['fallback_links'] > 0:
        print(f"\n{slot['name']}: {slot['fallback_links']} links in fallback")
        print("-"*80)
        
        # Find the fallback table
        for idx, table in enumerate(doc.tables):
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0].text.strip()
                if "Referenced Links" in first_cell or "Referenced Media" in first_cell:
                    # Check if this is after this slot's daily table
                    if idx > slot['daily_table']:
                        # Show first few links
                        link_num = 0
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    hyperlinks = para._element.xpath('.//w:hyperlink')
                                    for hyperlink in hyperlinks:
                                        text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
                                        link_num += 1
                                        if link_num <= 10:
                                            print(f"  {link_num}. {text[:70]}")
                        
                        if link_num > 10:
                            print(f"  ... and {link_num - 10} more")
                        break
