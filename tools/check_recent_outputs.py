"""
Check the most recent output files to see what's in them.
"""

import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_parser import DOCXParser

FILES = [
    ('Daniela (NEW)', r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43\Daniela_Silva_Weekly_W43_10-20-10-24_20251019_155023.docx'),
    ('Wilson (NEW)', r'F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W43\Wilson_Rodrigues_Weekly_W43_10-20-10-24_20251019_160328.docx'),
]


def check_file(label, file_path):
    """Check what's in a file."""
    
    print(f"\n{'='*80}")
    print(f"{label}")
    print(f"{'='*80}\n")
    
    path = Path(file_path)
    
    if not path.exists():
        print(f"❌ File not found: {file_path}")
        return
    
    print(f"✓ File exists")
    print(f"  Size: {path.stat().st_size:,} bytes")
    print()
    
    try:
        # Open with python-docx
        doc = Document(file_path)
        
        print(f"Document structure:")
        print(f"  Paragraphs: {len(doc.paragraphs)}")
        print(f"  Tables: {len(doc.tables)}")
        print()
        
        # Count hyperlinks manually
        total_hyperlinks = 0
        
        # In paragraphs
        para_links = 0
        for para in doc.paragraphs:
            for hyperlink in para._element.xpath('.//w:hyperlink'):
                para_links += 1
        
        # In tables
        table_links = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for hyperlink in para._element.xpath('.//w:hyperlink'):
                            table_links += 1
        
        total_hyperlinks = para_links + table_links
        
        print(f"Hyperlinks (manual count):")
        print(f"  In paragraphs: {para_links}")
        print(f"  In tables: {table_links}")
        print(f"  Total: {total_hyperlinks}")
        print()
        
        # Try parser
        print(f"Hyperlinks (via parser):")
        try:
            parser = DOCXParser(file_path)
            hyperlinks = parser.extract_hyperlinks()
            print(f"  Extracted: {len(hyperlinks)}")
            
            if hyperlinks:
                v2_links = [h for h in hyperlinks if h.get('schema_version') == '2.0']
                print(f"  Schema v2.0: {len(v2_links)}/{len(hyperlinks)}")
                
                # Show sample
                print()
                print(f"  Sample links:")
                for i, link in enumerate(hyperlinks[:3], 1):
                    print(f"    {i}. '{link['text'][:40]}'")
                    print(f"       Table: {link.get('table_idx')}, Row: {link.get('row_idx')}, Cell: {link.get('cell_idx')}")
        except Exception as e:
            print(f"  ❌ Parser error: {e}")
        
        print()
        
        # Check for "Referenced Links" section
        has_ref_links = False
        for para in doc.paragraphs:
            if "Referenced Links" in para.text:
                has_ref_links = True
                break
        
        if has_ref_links:
            print(f"⚠️  Has 'Referenced Links' section (fallback used)")
        else:
            print(f"✓ No 'Referenced Links' section (all inline)")
        
    except Exception as e:
        print(f"❌ Error opening file: {e}")
        import traceback
        traceback.print_exc()


def main():
    print("="*80)
    print("CHECKING MOST RECENT OUTPUT FILES")
    print("="*80)
    
    for label, file_path in FILES:
        check_file(label, file_path)
    
    print()
    print("="*80)
    print("CHECK COMPLETE")
    print("="*80)


if __name__ == '__main__':
    main()
