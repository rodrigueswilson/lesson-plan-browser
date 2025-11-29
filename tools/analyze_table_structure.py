"""
Analyze table structure in multi-slot DOCX files.

This script examines the actual table layout to understand:
1. How many tables exist
2. What's in each table's first cell
3. Which tables belong to which slots
4. Whether there are extra tables (summary, signature, etc.)
"""

import sys
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))


def analyze_document_structure(docx_path: str):
    """Analyze the structure of a multi-slot DOCX file."""
    print(f"\n{'='*80}")
    print(f"Analyzing: {Path(docx_path).name}")
    print(f"{'='*80}\n")
    
    doc = Document(docx_path)
    
    # Basic document info
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"\n{'='*80}")
    print("TABLE STRUCTURE")
    print(f"{'='*80}\n")
    
    # Analyze each table
    for idx, table in enumerate(doc.tables):
        print(f"Table {idx}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.columns) if table.rows else 0}")
        
        # Get first cell content (usually indicates table type)
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()
            print(f"  First cell: '{first_cell[:100]}'")
            
            # Try to identify table type
            if "Name:" in first_cell or "Teacher:" in first_cell:
                print(f"  -> Likely METADATA table")
            elif "Monday" in first_cell or "MONDAY" in first_cell:
                print(f"  -> Likely DAILY PLANS table")
            elif "Slot" in first_cell:
                print(f"  -> Contains slot identifier")
            elif "Signature" in first_cell or "signature" in first_cell:
                print(f"  -> Likely SIGNATURE table")
            elif "Referenced" in first_cell or "Links" in first_cell:
                print(f"  -> Likely REFERENCED LINKS table")
            
            # Check for slot number in first row
            first_row_text = " ".join(cell.text for cell in table.rows[0].cells)
            if "Slot" in first_row_text:
                print(f"  First row contains: '{first_row_text[:150]}'")
        
        print()
    
    # Analyze paragraphs with hyperlinks
    print(f"{'='*80}")
    print("PARAGRAPH HYPERLINKS")
    print(f"{'='*80}\n")
    
    para_hyperlinks = []
    for para_idx, para in enumerate(doc.paragraphs):
        # Check if paragraph has hyperlinks
        has_hyperlinks = False
        for run in para.runs:
            if hasattr(run, '_element'):
                hyperlinks = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
                if hyperlinks:
                    has_hyperlinks = True
                    break
        
        if has_hyperlinks:
            para_text = para.text.strip()
            if para_text:
                para_hyperlinks.append((para_idx, para_text[:100]))
    
    if para_hyperlinks:
        print(f"Found {len(para_hyperlinks)} paragraphs with hyperlinks:")
        for para_idx, text in para_hyperlinks[:10]:  # Show first 10
            print(f"  Paragraph {para_idx}: '{text}'")
        if len(para_hyperlinks) > 10:
            print(f"  ... and {len(para_hyperlinks) - 10} more")
    else:
        print("No hyperlinks found in paragraphs (all in tables)")
    
    print()


def analyze_slot_detection(docx_path: str):
    """Try to detect slot boundaries."""
    print(f"{'='*80}")
    print("SLOT DETECTION")
    print(f"{'='*80}\n")
    
    doc = Document(docx_path)
    
    # Strategy 1: Look for metadata tables with teacher names
    print("Strategy 1: Detect by metadata table pattern")
    metadata_tables = []
    
    for idx, table in enumerate(doc.tables):
        if not table.rows or len(table.rows) < 2:
            continue
        
        # Check if this looks like a metadata table
        first_row_text = " ".join(cell.text for cell in table.rows[0].cells).lower()
        
        is_metadata = any(keyword in first_row_text for keyword in [
            "name:", "teacher:", "grade:", "subject:", "week of:"
        ])
        
        if is_metadata:
            # Try to extract slot info
            slot_info = {
                "table_idx": idx,
                "teacher": None,
                "subject": None,
                "grade": None
            }
            
            # Parse metadata rows
            for row in table.rows[:5]:  # Check first 5 rows
                if len(row.cells) >= 2:
                    label = row.cells[0].text.strip().lower()
                    value = row.cells[1].text.strip()
                    
                    if "name" in label or "teacher" in label:
                        slot_info["teacher"] = value
                    elif "subject" in label:
                        slot_info["subject"] = value
                    elif "grade" in label:
                        slot_info["grade"] = value
            
            metadata_tables.append(slot_info)
            print(f"  Table {idx}: Teacher='{slot_info['teacher']}', Subject='{slot_info['subject']}', Grade='{slot_info['grade']}'")
    
    print(f"\nDetected {len(metadata_tables)} potential slots")
    
    # Strategy 2: Assume 2 tables per slot
    print(f"\nStrategy 2: Assume 2 tables per slot")
    print(f"  Total tables: {len(doc.tables)}")
    print(f"  Estimated slots: {len(doc.tables) // 2}")
    
    if len(metadata_tables) * 2 != len(doc.tables):
        print(f"  ⚠️ WARNING: Table count doesn't match 2-per-slot assumption!")
        print(f"     Expected: {len(metadata_tables) * 2} tables")
        print(f"     Actual: {len(doc.tables)} tables")
        print(f"     Difference: {len(doc.tables) - len(metadata_tables) * 2} extra tables")
    
    print()


def main():
    """Analyze all lesson plan files."""
    # Directories to scan
    base_dirs = [
        Path("F:/rodri/Documents/OneDrive/AS/Daniela LP"),
        Path("F:/rodri/Documents/OneDrive/AS/Lesson Plan")
    ]
    
    all_docx_files = []
    
    # Find all DOCX files in all directories and subdirectories
    for base_dir in base_dirs:
        if not base_dir.exists():
            print(f"WARNING: Directory not found: {base_dir}")
            continue
        
        print(f"Scanning: {base_dir}")
        docx_files = list(base_dir.rglob("*.docx"))
        
        # Filter out temp files
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]
        
        print(f"  Found {len(docx_files)} DOCX files")
        all_docx_files.extend(docx_files)
    
    if not all_docx_files:
        print("No DOCX files found in any directory")
        return
    
    print(f"\nTotal files to analyze: {len(all_docx_files)}")
    print(f"{'='*80}\n")
    
    # Track statistics
    stats = {
        "total_files": len(all_docx_files),
        "analyzed": 0,
        "errors": 0,
        "table_counts": {},
        "slot_counts": {}
    }
    
    # Analyze each file
    for idx, docx_file in enumerate(all_docx_files, 1):
        print(f"[{idx}/{len(all_docx_files)}] Analyzing: {docx_file.relative_to(docx_file.parents[3])}")
        
        try:
            analyze_document_structure(str(docx_file))
            analyze_slot_detection(str(docx_file))
            stats["analyzed"] += 1
            
            # Track table count
            doc = Document(str(docx_file))
            table_count = len(doc.tables)
            stats["table_counts"][table_count] = stats["table_counts"].get(table_count, 0) + 1
            
            # Estimate slot count (excluding signature table)
            estimated_slots = (table_count - 1) // 2 if table_count > 0 else 0
            stats["slot_counts"][estimated_slots] = stats["slot_counts"].get(estimated_slots, 0) + 1
            
        except Exception as e:
            print(f"ERROR: {e}")
            stats["errors"] += 1
            import traceback
            traceback.print_exc()
        
        print()  # Blank line between files
    
    # Print summary statistics
    print(f"\n{'='*80}")
    print("ANALYSIS SUMMARY")
    print(f"{'='*80}")
    print(f"Total files: {stats['total_files']}")
    print(f"Successfully analyzed: {stats['analyzed']}")
    print(f"Errors: {stats['errors']}")
    print(f"\nTable count distribution:")
    for count, freq in sorted(stats["table_counts"].items()):
        print(f"  {count} tables: {freq} files")
    print(f"\nEstimated slot distribution:")
    for count, freq in sorted(stats["slot_counts"].items()):
        print(f"  {count} slots: {freq} files")
    print(f"{'='*80}")


if __name__ == "__main__":
    main()
