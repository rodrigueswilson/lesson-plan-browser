"""
Robust DOCX Parser for Primary Teacher Lesson Plans
Handles multiple table formats and provides warnings for variations
"""

from docx import Document
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import re


class DOCXParser:
    """Parser for primary teacher DOCX files with multiple format support."""
    
    def __init__(self, file_path: str):
        """
        Initialize parser with DOCX file.
        
        Args:
            file_path: Path to DOCX file
        """
        self.file_path = Path(file_path)
        self.doc = Document(file_path)
        self.warnings = []
        
    def find_subject_tables(self) -> List[Dict[str, Any]]:
        """
        Find all subject sections in the document.
        
        Returns:
            List of dicts with subject info and table indices
        """
        subjects = []
        
        # Loop through tables looking for headers
        i = 0
        while i < len(self.doc.tables):
            table = self.doc.tables[i]
            
            # Check if this is a header table (1 row with "Subject:")
            if len(table.rows) == 1:
                subject_info = self._extract_header_info(table, i)
                if subject_info:
                    # Check if next table exists and is a lesson table
                    if i + 1 < len(self.doc.tables):
                        lesson_table = self.doc.tables[i + 1]
                        
                        # Validate lesson table structure
                        validation = self._validate_lesson_table(lesson_table, i + 1)
                        subject_info.update(validation)
                        
                        subjects.append(subject_info)
                        i += 2  # Skip to next pair
                        continue
            
            i += 1
        
        return subjects
    
    def _extract_header_info(self, table, table_index: int) -> Optional[Dict[str, Any]]:
        """
        Extract information from header table.
        
        Args:
            table: Header table
            table_index: Index of this table
            
        Returns:
            Dict with header info or None if not a valid header
        """
        header_info = {
            'header_table_index': table_index,
            'lesson_table_index': table_index + 1
        }
        
        # Extract fields from header row
        for cell in table.rows[0].cells:
            text = cell.text.strip()
            
            if 'Name:' in text:
                header_info['teacher_name'] = text.replace('Name:', '').strip()
            elif 'Grade:' in text:
                header_info['grade'] = text.replace('Grade:', '').strip()
            elif 'Homeroom:' in text:
                header_info['homeroom'] = text.replace('Homeroom:', '').strip()
            elif 'Subject:' in text:
                header_info['subject'] = text.replace('Subject:', '').strip()
            elif 'Week of:' in text or 'Week:' in text:
                week_text = text.replace('Week of:', '').replace('Week:', '').strip()
                header_info['week'] = week_text
        
        # Must have subject to be valid
        if 'subject' in header_info:
            return header_info
        
        return None
    
    def _validate_lesson_table(self, table, table_index: int) -> Dict[str, Any]:
        """
        Validate and analyze lesson table structure.
        
        Args:
            table: Lesson table
            table_index: Index of this table
            
        Returns:
            Dict with validation results and warnings
        """
        validation = {
            'rows': len(table.rows),
            'columns': len(table.columns),
            'format_type': 'unknown',
            'warnings': []
        }
        
        # Check if it's a signature table (should be skipped)
        if len(table.rows) <= 4 and len(table.columns) == 1:
            first_cell = table.rows[0].cells[0].text.strip()
            if 'signature' in first_cell.lower() or 'certify' in first_cell.lower():
                validation['format_type'] = 'signature'
                validation['warnings'].append(f"Table {table_index}: Signature table detected, should be skipped")
                return validation
        
        # Check column structure (should have Monday-Friday)
        if len(table.rows) > 0:
            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            days_found = sum(1 for cell in first_row if any(day in cell.upper() for day in ['MONDAY', 'TUESDAY', 'WEDNESDAY', 'THURSDAY', 'FRIDAY']))
            
            if days_found >= 3:  # At least 3 days found
                validation['has_day_columns'] = True
            else:
                validation['has_day_columns'] = False
                validation['warnings'].append(f"Table {table_index}: Day columns not clearly identified")
        
        # Determine format type based on row count
        if validation['rows'] == 8:
            validation['format_type'] = 'standard'  # Davies, Lang, Savoca format
        elif validation['rows'] == 13:
            validation['format_type'] = 'extended'  # Piret format
        elif validation['rows'] > 0:
            validation['format_type'] = 'custom'
            validation['warnings'].append(f"Table {table_index}: Non-standard row count ({validation['rows']} rows)")
        
        return validation
    
    def extract_subject_content(self, subject: str) -> Dict[str, Any]:
        """
        Extract content for a specific subject.
        
        Args:
            subject: Subject name to extract (case-insensitive)
            
        Returns:
            Dict with extracted content and metadata
        """
        subjects = self.find_subject_tables()
        
        # Find matching subject
        subject_lower = subject.lower()
        for subj_info in subjects:
            if subject_lower in subj_info.get('subject', '').lower():
                # Extract lesson content
                lesson_table = self.doc.tables[subj_info['lesson_table_index']]
                content = self._extract_table_content(lesson_table)
                
                # Combine warnings
                all_warnings = self.warnings + subj_info.get('warnings', [])
                
                return {
                    'subject': subj_info['subject'],
                    'found': True,
                    'format_type': subj_info.get('format_type', 'unknown'),
                    'table_content': content,
                    'full_text': self._table_to_text(content),
                    'metadata': {
                        'teacher_name': subj_info.get('teacher_name'),
                        'grade': subj_info.get('grade'),
                        'homeroom': subj_info.get('homeroom'),
                        'week': subj_info.get('week'),
                        'rows': subj_info.get('rows'),
                        'columns': subj_info.get('columns')
                    },
                    'warnings': all_warnings
                }
        
        # Subject not found
        return {
            'subject': subject,
            'found': False,
            'full_text': '',
            'warnings': [f"Subject '{subject}' not found in document. Available subjects: {[s.get('subject') for s in subjects]}"]
        }
    
    def _extract_table_content(self, table) -> Dict[str, Dict[str, str]]:
        """
        Extract content from lesson table.
        
        Args:
            table: Lesson table
            
        Returns:
            Dict mapping days to lesson components
        """
        content = {}
        
        # Get day columns from first row
        days = []
        for i, cell in enumerate(table.rows[0].cells):
            day_text = cell.text.strip()
            if day_text:  # Any non-empty cell in first row
                days.append((i, day_text))
        
        # Extract content for each row
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            row_label = row.cells[0].text.strip()
            
            if not row_label:
                continue
            
            # For each day column
            for col_idx, day_name in days:
                if day_name not in content:
                    content[day_name] = {}
                
                if col_idx < len(row.cells):
                    cell_text = row.cells[col_idx].text.strip()
                    if cell_text:
                        content[day_name][row_label] = cell_text
        
        return content
    
    def _table_to_text(self, table_content: Dict[str, Dict[str, str]]) -> str:
        """
        Convert table content to text format.
        
        Args:
            table_content: Dict from _extract_table_content
            
        Returns:
            Formatted text string
        """
        text_parts = []
        
        for day, components in table_content.items():
            text_parts.append(f"\n{day.upper()}")
            for label, content in components.items():
                text_parts.append(f"{label}: {content}")
        
        return "\n".join(text_parts)
    
    def get_all_subjects(self) -> List[str]:
        """
        Get list of all subjects in the document.
        
        Returns:
            List of subject names
        """
        subjects = self.find_subject_tables()
        return [s.get('subject', 'Unknown') for s in subjects]
    
    def get_warnings(self) -> List[str]:
        """
        Get all warnings from parsing.
        
        Returns:
            List of warning messages
        """
        return self.warnings
    
    def validate_structure(self) -> Dict[str, Any]:
        """
        Validate overall document structure.
        
        Returns:
            Dict with validation results
        """
        subjects = self.find_subject_tables()
        
        validation = {
            'total_tables': len(self.doc.tables),
            'subjects_found': len(subjects),
            'subjects': [s.get('subject') for s in subjects],
            'format_types': {},
            'warnings': [],
            'is_valid': True
        }
        
        # Count format types
        for subj in subjects:
            fmt = subj.get('format_type', 'unknown')
            validation['format_types'][fmt] = validation['format_types'].get(fmt, 0) + 1
        
        # Collect all warnings
        for subj in subjects:
            validation['warnings'].extend(subj.get('warnings', []))
        
        # Check for issues
        if len(subjects) == 0:
            validation['is_valid'] = False
            validation['warnings'].append("No valid subject sections found")
        
        if 'custom' in validation['format_types']:
            validation['warnings'].append(f"Document contains {validation['format_types']['custom']} custom format table(s)")
        
        return validation


def get_parser(file_path: str) -> DOCXParser:
    """
    Get a DOCX parser instance.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        DOCXParser instance
    """
    return DOCXParser(file_path)
