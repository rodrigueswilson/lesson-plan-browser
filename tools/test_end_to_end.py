"""
End-to-end test for multi-slot lesson plan generation.
Tests the complete pipeline from database to final DOCX output.
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.database import get_db
from backend.llm_service import get_llm_service
from tools.batch_processor import BatchProcessor


async def test_end_to_end():
    """Run end-to-end test of multi-slot processing."""
    
    print("=" * 70)
    print("END-TO-END TEST: Multi-Slot Lesson Plan Generation")
    print("=" * 70)
    
    # Setup
    db = get_db()
    user = db.get_user_by_name('Wilson Rodrigues')
    
    if not user:
        print("ERROR: User 'Wilson Rodrigues' not found")
        return False
    
    slots = db.get_user_slots(user['id'])
    
    if not slots:
        print("ERROR: No slots configured for user")
        return False
    
    print(f"\nUser: {user['name']} (ID: {user['id']})")
    print(f"Configured Slots: {len(slots)}")
    for slot in slots:
        print(f"  - Slot {slot['slot_number']}: {slot['subject']} ({slot.get('primary_teacher_name', 'N/A')})")
    
    # Test parameters
    week_of = "9/15-9/19"  # Week that matches our input files
    provider = "openai"
    
    print(f"\nTest Parameters:")
    print(f"  Week: {week_of}")
    print(f"  Provider: {provider}")
    
    # Check input files
    print(f"\nChecking input files for week {week_of}...")
    input_dir = Path('input')
    week_files = []
    
    for slot in slots:
        teacher_pattern = slot.get('primary_teacher_name', '')
        if teacher_pattern:
            matching_files = list(input_dir.glob(f"*{teacher_pattern}*.docx"))
            if matching_files:
                week_files.append((slot['slot_number'], slot['subject'], matching_files[0]))
                print(f"  ✓ Slot {slot['slot_number']} ({slot['subject']}): {matching_files[0].name}")
            else:
                print(f"  ✗ Slot {slot['slot_number']} ({slot['subject']}): No file found for pattern '{teacher_pattern}'")
    
    if not week_files:
        print("\nERROR: No input files found for configured slots")
        return False
    
    # Initialize LLM service
    print(f"\nInitializing LLM service...")
    llm_service = get_llm_service()
    
    # Check API key
    if not llm_service.openai_client:
        print("WARNING: OpenAI API key not configured. Test will use mock responses.")
        print("To test with real LLM, set OPENAI_API_KEY in .env file")
    
    # Create batch processor
    print(f"\nCreating batch processor...")
    processor = BatchProcessor(llm_service)
    
    # Run processing
    print(f"\n{'=' * 70}")
    print(f"STARTING BATCH PROCESSING")
    print(f"{'=' * 70}\n")
    
    start_time = datetime.now()
    
    try:
        result = await processor.process_user_week(
            user_id=user['id'],
            week_of=week_of,
            provider=provider
        )
        
        end_time = datetime.now()
        elapsed = (end_time - start_time).total_seconds()
        
        print(f"\n{'=' * 70}")
        print(f"PROCESSING COMPLETE")
        print(f"{'=' * 70}\n")
        
        # Display results
        print(f"Results:")
        print(f"  Success: {result['success']}")
        print(f"  Processed Slots: {result['processed_slots']}")
        print(f"  Failed Slots: {result['failed_slots']}")
        print(f"  Total Time: {elapsed:.2f}s ({elapsed/60:.1f} minutes)")
        
        if result['success']:
            print(f"  Output File: {result['output_file']}")
            
            # Verify output file
            output_path = Path(result['output_file'])
            if output_path.exists():
                file_size = output_path.stat().st_size
                print(f"  File Size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
                print(f"\n✅ SUCCESS: Multi-slot DOCX generated successfully!")
                
                # Check file content
                print(f"\nVerifying DOCX content...")
                from docx import Document
                doc = Document(output_path)
                
                # Count tables
                print(f"  Tables in document: {len(doc.tables)}")
                
                # Check daily plans table
                if len(doc.tables) >= 2:
                    daily_table = doc.tables[1]
                    print(f"  Daily plans table rows: {len(daily_table.rows)}")
                    print(f"  Daily plans table columns: {len(daily_table.columns)}")
                    
                    # Sample Monday content
                    monday_col = 1
                    unit_row = 1
                    monday_cell = daily_table.rows[unit_row].cells[monday_col]
                    monday_text = monday_cell.text[:200] if monday_cell.text else "(empty)"
                    print(f"\n  Sample Monday content (Unit/Lesson):")
                    print(f"    {monday_text}...")
                    
                    # Check for slot markers
                    if "Slot" in monday_text:
                        print(f"\n  ✅ Multi-slot markers found in content!")
                    else:
                        print(f"\n  ⚠️  No slot markers found - may be single slot or formatting issue")
                
                return True
            else:
                print(f"\n❌ ERROR: Output file not found at {output_path}")
                return False
        else:
            print(f"\n❌ PROCESSING FAILED")
            if result.get('errors'):
                print(f"\nErrors:")
                for error in result['errors']:
                    print(f"  - {error}")
            return False
            
    except Exception as e:
        print(f"\n❌ EXCEPTION DURING PROCESSING:")
        print(f"  {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main entry point."""
    success = asyncio.run(test_end_to_end())
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
