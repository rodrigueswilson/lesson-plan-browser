"""
Simple module to extract and print objectives from lesson plans.

This module helps teachers print objectives from their lesson plans,
either from JSON files or from the database (if lesson_json is stored).
"""

import json
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.services.objectives_utils import normalize_objective_payload
from backend.services.sorting_utils import sort_slots
from backend.telemetry import logger


def extract_subject_from_unit_lesson(unit_lesson: str) -> str:
    """
    Extract the actual subject from the unit_lesson text.

    This function parses patterns like:
    - "Unit 3 Lesson 9: MEASURE TO FIND THE AREA" -> "Math"
    - "Math Chapter 5: MULTIPLICATION FACTS" -> "Math"
    - "ELA Unit 2: READING COMPREHENSION" -> "ELA"
    - "Science Lab: PLANT GROWTH" -> "Science"

    Args:
        unit_lesson: The unit/lesson text from the lesson plan

    Returns:
        Detected subject name or "Unknown" if not found
    """
    import re

    if not unit_lesson or not isinstance(unit_lesson, str):
        return "Unknown"

    unit_lesson = unit_lesson.strip().upper()

    # Pattern 1: Explicit subject prefixes
    subject_patterns = {
        r"\bMATH\b": "Math",
        r"\bELA\b": "ELA",
        r"\bENGLISH\b": "ELA",
        r"\bREADING\b": "ELA",
        r"\bWRITING\b": "ELA",
        r"\bSCIENCE\b": "Science",
        r"\bSOCIAL STUDIES\b": "Social Studies",
        r"\bHISTORY\b": "Social Studies",
        r"\bGEOGRAPHY\b": "Social Studies",
        r"\bART\b": "Art",
        r"\bMUSIC\b": "Music",
        r"\bPE\b": "Physical Education",
        r"\bPHYSICAL EDUCATION\b": "Physical Education",
        r"\bHEALTH\b": "Health",
        r"\bLIBRARY\b": "Library",
        r"\bTECH\b": "Technology",
        r"\bTECHNOLOGY\b": "Technology",
        r"\bCOMPUTER\b": "Technology",
    }

    for pattern, subject in subject_patterns.items():
        if re.search(pattern, unit_lesson):
            return subject

    # Pattern 2: Math-specific keywords
    math_keywords = [
        "MULTIPLICATION",
        "DIVISION",
        "ADDITION",
        "SUBTRACTION",
        "MEASURE",
        "AREA",
        "PERIMETER",
        "GEOMETRY",
        "FRACTIONS",
        "DECIMALS",
        "PLACE VALUE",
        "NUMBER",
        "CALCULATION",
        "EQUATION",
        "PROBLEM SOLVING",
        "MATH CHAPTER",
        "UNIT .* LESSON",
        "UNIT",  # Add standalone UNIT pattern
    ]

    for keyword in math_keywords:
        if re.search(r"\b" + keyword + r"\b", unit_lesson):
            return "Math"

    # Pattern 3: ELA-specific keywords
    ela_keywords = [
        "READING",
        "COMPREHENSION",
        "GRAMMAR",
        "SPELLING",
        "VOCABULARY",
        "WRITING",
        "PHONICS",
        "LITERATURE",
        "STORY",
        "POEM",
        "ESSAY",
        "LANGUAGE ARTS",
    ]

    for keyword in ela_keywords:
        if re.search(r"\b" + keyword + r"\b", unit_lesson):
            return "ELA"

    # Pattern 4: Science-specific keywords
    science_keywords = [
        "EXPERIMENT",
        "LAB",
        "ORGANISM",
        "HABITAT",
        "ECOSYSTEM",
        "MATTER",
        "ENERGY",
        "FORCE",
        "MOTION",
        "PLANT",
        "ANIMAL",
    ]

    for keyword in science_keywords:
        if re.search(r"\b" + keyword + r"\b", unit_lesson):
            return "Science"

    # Pattern 5: Social Studies keywords
    ss_keywords = [
        "COMMUNITY",
        "GOVERNMENT",
        "HISTORY",
        "GEOGRAPHY",
        "CULTURE",
        "SOCIETY",
        "CIVICS",
        "ECONOMICS",
    ]

    for keyword in ss_keywords:
        if re.search(r"\b" + keyword + r"\b", unit_lesson):
            return "Social Studies"

    # Default: Return "Unknown"
    return "Unknown"


class ObjectivesPrinter:
    """Extract and format objectives from lesson plans for printing."""

    DAY_NAMES = ["monday", "tuesday", "wednesday", "thursday", "friday"]

    def __init__(self):
        """Initialize the objectives printer."""
        pass

    def load_from_json_file(self, json_path: str) -> Dict[str, Any]:
        """Load lesson plan from JSON file.

        Args:
            json_path: Path to lesson plan JSON file

        Returns:
            lesson_json dictionary
        """
        json_file = Path(json_path)
        if not json_file.exists():
            raise FileNotFoundError(f"JSON file not found: {json_path}")

        with open(json_file, "r", encoding="utf-8") as f:
            lesson_json = json.load(f)

        return lesson_json

    def load_from_database(self, plan_id: str) -> Optional[Dict[str, Any]]:
        """Load lesson plan from database.

        Note: This assumes lesson_json is stored in weekly_plans table.
        If not yet implemented, this will return None.

        Args:
            plan_id: Plan ID from weekly_plans table

        Returns:
            lesson_json dictionary or None if not found
        """
        from backend.database import get_db

        db = get_db()
        plan = db.get_weekly_plan(plan_id)

        if not plan:
            return None

        # Check if lesson_json column exists and has data
        # Note: This may not be implemented yet in the database
        lesson_json = plan.get("lesson_json")

        if isinstance(lesson_json, str):
            # If stored as JSON string, parse it
            return json.loads(lesson_json)
        elif isinstance(lesson_json, dict):
            # If already a dict (JSONB in Supabase)
            return lesson_json

        return None

    def extract_objectives(self, lesson_json: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract all objectives from lesson plan.

        Args:
            lesson_json: Lesson plan JSON structure

        Returns:
            List of objective dictionaries with day, slot, and objective data
        """
        objectives = []

        if "days" not in lesson_json:
            logger.warning("No 'days' key found in lesson_json")
            return objectives

        metadata = lesson_json.get("metadata", {})
        week_of = metadata.get("week_of", "Unknown")
        grade = metadata.get("grade", "Unknown")
        subject = metadata.get("subject", "Unknown")

        days = lesson_json["days"]

        for day_name in self.DAY_NAMES:
            if day_name not in days:
                continue

            day_data = days[day_name]
            slots = day_data.get("slots", [])

            # Sort slots by start_time (chronological) with slot_number as fallback
            sorted_slots = sort_slots(slots)

            for slot in sorted_slots:
                slot_number = slot.get("slot_number", 0)
                slot_subject = slot.get("subject", subject)
                unit_lesson = slot.get("unit_lesson", "")
                teacher_name = slot.get("teacher_name", "")

                # Skip "No School" entries
                if unit_lesson and unit_lesson.strip().lower() == "no school":
                    continue

                objective_data = normalize_objective_payload(
                    slot.get("objective", {}),
                    {
                        "day": day_name,
                        "slot_number": slot_number,
                        "subject": slot_subject,
                    },
                )

                if not objective_data:
                    continue

                # Check if all objective fields are "No School"
                content_obj = (
                    objective_data.get("content_objective", "").strip().lower()
                )
                student_goal = objective_data.get("student_goal", "").strip().lower()
                wida_obj = objective_data.get("wida_objective", "").strip().lower()

                if (
                    content_obj == "no school"
                    and student_goal == "no school"
                    and wida_obj == "no school"
                ):
                    continue

                # *** IMPROVEMENT: Use subject detection with proper fallback ***
                # Priority: 1) slot subject, 2) detected from unit_lesson, 3) metadata subject
                final_subject = slot_subject

                # If slot subject is Unknown/empty, try detection
                if final_subject in ["Unknown", ""]:
                    detected_subject = extract_subject_from_unit_lesson(unit_lesson)
                    if detected_subject != "Unknown":
                        final_subject = detected_subject
                    elif unit_lesson and unit_lesson.strip():
                        # Only fall back to metadata if we have valid unit_lesson but couldn't detect
                        final_subject = subject
                    else:
                        # unit_lesson is empty/None, keep Unknown
                        final_subject = "Unknown"

                objectives.append(
                    {
                        "week_of": week_of,
                        "day": day_name.capitalize(),
                        "slot_number": slot_number,
                        "subject": final_subject,  # Use improved subject detection
                        "unit_lesson": unit_lesson,
                        "teacher_name": teacher_name,
                        "content_objective": objective_data.get(
                            "content_objective", ""
                        ),
                        "student_goal": objective_data.get("student_goal", ""),
                        "wida_objective": objective_data.get("wida_objective", ""),
                    }
                )

        return objectives

    def format_for_print(
        self, objectives: List[Dict[str, Any]], format_type: str = "text"
    ) -> str:
        """Format objectives for printing.

        Args:
            objectives: List of objective dictionaries
            format_type: 'text', 'markdown', or 'html'

        Returns:
            Formatted string ready for printing
        """
        if not objectives:
            return "No objectives found."

        # Group by week
        week_of = objectives[0].get("week_of", "Unknown")

        if format_type == "markdown":
            return self._format_markdown(objectives, week_of)
        elif format_type == "html":
            return self._format_html(objectives, week_of)
        else:  # text
            return self._format_text(objectives, week_of)

    def _format_text(self, objectives: List[Dict[str, Any]], week_of: str) -> str:
        """Format as plain text."""
        lines = []
        lines.append("=" * 80)
        lines.append(f"LESSON PLAN OBJECTIVES - Week of {week_of}")
        lines.append("=" * 80)
        lines.append("")

        # Group by day
        current_day = None
        for obj in objectives:
            day = obj["day"]
            if day != current_day:
                if current_day is not None:
                    lines.append("")
                lines.append(f"{day.upper()}")
                lines.append("-" * 80)
                current_day = day

            lines.append(f"\nSlot {obj['slot_number']}: {obj['subject']}")
            if obj["unit_lesson"]:
                lines.append(f"Unit/Lesson: {obj['unit_lesson']}")
            if obj["teacher_name"]:
                lines.append(f"Teacher: {obj['teacher_name']}")
            lines.append("")

            if obj["content_objective"]:
                lines.append("Content Objective:")
                lines.append(f"  {obj['content_objective']}")
                lines.append("")

            if obj["student_goal"]:
                lines.append("Student Goal:")
                lines.append(f"  {obj['student_goal']}")
                lines.append("")

            if obj["wida_objective"]:
                lines.append("WIDA Objective:")
                lines.append(f"  {obj['wida_objective']}")
                lines.append("")

            lines.append("-" * 80)

        return "\n".join(lines)

    def _format_markdown(self, objectives: List[Dict[str, Any]], week_of: str) -> str:
        """Format as Markdown."""
        lines = []
        lines.append(f"# Lesson Plan Objectives - Week of {week_of}")
        lines.append("")

        # Group by day
        current_day = None
        for obj in objectives:
            day = obj["day"]
            if day != current_day:
                if current_day is not None:
                    lines.append("")
                lines.append(f"## {day}")
                lines.append("")
                current_day = day

            lines.append(f"### Slot {obj['slot_number']}: {obj['subject']}")
            if obj["unit_lesson"]:
                lines.append(f"**Unit/Lesson:** {obj['unit_lesson']}")
            if obj["teacher_name"]:
                lines.append(f"**Teacher:** {obj['teacher_name']}")
            lines.append("")

            if obj["content_objective"]:
                lines.append("**Content Objective:**")
                lines.append(f"{obj['content_objective']}")
                lines.append("")

            if obj["student_goal"]:
                lines.append("**Student Goal:**")
                lines.append(f"{obj['student_goal']}")
                lines.append("")

            if obj["wida_objective"]:
                lines.append("**WIDA Objective:**")
                lines.append(f"{obj['wida_objective']}")
                lines.append("")

        return "\n".join(lines)

    def _format_html(self, objectives: List[Dict[str, Any]], week_of: str) -> str:
        """Format as HTML."""
        lines = []
        lines.append("<!DOCTYPE html>")
        lines.append("<html><head>")
        lines.append("<meta charset='utf-8'>")
        lines.append("<title>Lesson Plan Objectives - Week of {week_of}</title>")
        lines.append("<style>")
        lines.append("body { font-family: Arial, sans-serif; margin: 20px; }")
        lines.append("h1 { color: #333; border-bottom: 2px solid #333; }")
        lines.append("h2 { color: #666; margin-top: 30px; }")
        lines.append("h3 { color: #888; }")
        lines.append(
            ".objective { margin: 15px 0; padding: 10px; background: #f5f5f5; }"
        )
        lines.append(".label { font-weight: bold; color: #555; }")
        lines.append("</style>")
        lines.append("</head><body>")
        lines.append(f"<h1>Lesson Plan Objectives - Week of {week_of}</h1>")

        # Group by day
        current_day = None
        for obj in objectives:
            day = obj["day"]
            if day != current_day:
                if current_day is not None:
                    lines.append("</div>")
                lines.append(f"<h2>{day}</h2>")
                lines.append("<div class='day-objectives'>")
                current_day = day

            lines.append("<div class='objective'>")
            lines.append(f"<h3>Slot {obj['slot_number']}: {obj['subject']}</h3>")
            if obj["unit_lesson"]:
                lines.append(
                    f"<p><span class='label'>Unit/Lesson:</span> {obj['unit_lesson']}</p>"
                )
            if obj["teacher_name"]:
                lines.append(
                    f"<p><span class='label'>Teacher:</span> {obj['teacher_name']}</p>"
                )

            if obj["content_objective"]:
                lines.append(
                    f"<p><span class='label'>Content Objective:</span><br>{obj['content_objective']}</p>"
                )

            if obj["student_goal"]:
                lines.append(
                    f"<p><span class='label'>Student Goal:</span><br>{obj['student_goal']}</p>"
                )

            if obj["wida_objective"]:
                lines.append(
                    f"<p><span class='label'>WIDA Objective:</span><br>{obj['wida_objective']}</p>"
                )

            lines.append("</div>")

        if current_day:
            lines.append("</div>")
        lines.append("</body></html>")

        return "\n".join(lines)

    def save_to_file(
        self,
        objectives: List[Dict[str, Any]],
        output_path: str,
        format_type: str = "text",
    ) -> str:
        """Save formatted objectives to file.

        Args:
            objectives: List of objective dictionaries
            output_path: Path to save file
            format_type: 'text', 'markdown', or 'html'

        Returns:
            Path to saved file
        """
        formatted = self.format_for_print(objectives, format_type)

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(formatted)

        logger.info(
            "objectives_saved",
            extra={
                "output_path": str(output_file),
                "format": format_type,
                "objective_count": len(objectives),
            },
        )

        return str(output_file)

    def print_objectives(
        self,
        source: str,
        source_type: str = "file",
        output_path: Optional[str] = None,
        format_type: str = "text",
    ) -> str:
        """Main method to extract and print objectives.

        Args:
            source: Path to JSON file or plan_id from database
            source_type: 'file' or 'database'
            output_path: Optional path to save formatted output
            format_type: 'text', 'markdown', or 'html'

        Returns:
            Formatted objectives string
        """
        # Load lesson plan
        if source_type == "database":
            lesson_json = self.load_from_database(source)
            if not lesson_json:
                raise ValueError(f"Plan not found in database: {source}")
        else:
            lesson_json = self.load_from_json_file(source)

        # Extract objectives
        objectives = self.extract_objectives(lesson_json)

        if not objectives:
            return "No objectives found in lesson plan."

        # Format
        formatted = self.format_for_print(objectives, format_type)

        # Save if output path provided
        if output_path:
            self.save_to_file(objectives, output_path, format_type)

        return formatted

    def calculate_font_size(
        self,
        text: str,
        available_width: float,
        available_height: float,
        base_font_size: int,
        min_font_size: int,
        max_font_size: int,
    ) -> int:
        """
        Calculate optimal font size to maximize text fill in available space.

        Uses iterative approach to find font size that fills target height (85%)
        while allowing text to wrap naturally within available width.

        Font metrics (Verdana Bold):
        - Average character width: ~0.6 * font_size (in points)
        - Line height with spacing: ~1.3 * font_size (in points)
        - Conversion: 72 points = 1 inch

        Args:
            text: Text to measure
            available_width: Available width in inches
            available_height: Available height in inches
            base_font_size: Base font size in points (starting point)
            min_font_size: Minimum font size in points
            max_font_size: Maximum font size in points

        Returns:
            Calculated font size in points
        """
        if not text or not text.strip():
            return base_font_size

        # Estimate text dimensions at base font size
        lines = text.split("\n")
        max_line_length = max(len(line) for line in lines) if lines else 0

        if max_line_length == 0:
            return base_font_size

        # Calculate maximum font size that fills target height while allowing natural wrapping
        # Based on typography research and python-docx best practices:
        # - Verdana Bold: average character width ≈ 0.6 * font_size (in points)
        # - Line height with 1.25 spacing: ≈ 1.3 * font_size (in points)
        # - Conversion: 72 points = 1 inch

        # Font metrics constants (based on Verdana Bold typography)
        CHAR_WIDTH_RATIO = (
            0.6  # Character width as fraction of font size (Verdana is wider)
        )
        LINE_HEIGHT_RATIO = (
            1.3  # Line height including spacing (1.25 spacing + baseline)
        )

        # Strategy: Fill 85% of available height for optimal visual fill
        height_target = available_height * 0.85

        # Get total character count (for wrapping calculations)
        text_without_newlines = text.replace("\n", " ")
        total_chars = len(text_without_newlines)
        existing_lines = len(lines)

        if total_chars == 0 or available_width <= 0 or available_height <= 0:
            return base_font_size

        if existing_lines > 1:
            # Multi-line text: Calculate based on existing line breaks
            # Formula: font_size = height_target * 72 / (lines * line_height_ratio)
            font_size = (height_target * 72) / (existing_lines * LINE_HEIGHT_RATIO)
        else:
            # Single line: Text will wrap, use iterative approach for accuracy
            # Formula derivation for wrapping text:
            #   chars_per_line = (available_width * 72) / (font_size * CHAR_WIDTH_RATIO)
            #   wrapped_lines = total_chars / chars_per_line
            #   wrapped_lines = total_chars * font_size * CHAR_WIDTH_RATIO / (available_width * 72)
            #   height = wrapped_lines * font_size * LINE_HEIGHT_RATIO / 72
            #   height = total_chars * font_size^2 * CHAR_WIDTH_RATIO * LINE_HEIGHT_RATIO / (available_width * 72^2)
            # Solving for font_size:
            #   font_size^2 = height_target * available_width * 72^2 / (total_chars * CHAR_WIDTH_RATIO * LINE_HEIGHT_RATIO)
            #   font_size = sqrt(height_target * available_width * 72^2 / (total_chars * CHAR_WIDTH_RATIO * LINE_HEIGHT_RATIO))

            numerator = height_target * available_width * 72 * 72
            denominator = total_chars * CHAR_WIDTH_RATIO * LINE_HEIGHT_RATIO

            if denominator > 0:
                font_size_squared = numerator / denominator
                font_size = font_size_squared**0.5
            else:
                font_size = base_font_size

            # Refine with iterative approach for better accuracy
            # Account for word boundaries (words don't break mid-word, adding ~10% overhead)
            # Iterate 2-3 times to converge on optimal size
            for iteration in range(3):
                # Calculate characters per line at current font size
                chars_per_line = (
                    (available_width * 72) / (font_size * CHAR_WIDTH_RATIO)
                    if font_size > 0
                    else 1
                )

                # Estimate wrapped lines (add 10% overhead for word boundaries)
                estimated_wrapped_lines = max(1, (total_chars / chars_per_line) * 1.1)

                # Calculate font size needed for these lines to fill height_target
                new_font_size = (height_target * 72) / (
                    estimated_wrapped_lines * LINE_HEIGHT_RATIO
                )

                # Converge using weighted average to avoid oscillation
                font_size = font_size * 0.6 + new_font_size * 0.4

        # Clamp to min/max bounds
        font_size = max(min_font_size, min(int(font_size), max_font_size))

        return font_size

    def calculate_font_size_to_fill_height(
        self,
        text: str,
        available_width: float,
        target_height: float,
        min_font_size: int,
        max_font_size: int,
        char_width_ratio: float = 0.6,
        line_height_ratio: float = 1.3,
    ) -> int:
        """
        Calculate maximum font size that fits within target height.
        Uses conservative approach to ensure content definitely fits.

        Args:
            text: Text to measure
            available_width: Available width in inches
            target_height: Maximum height in inches (content must fit within this)
            min_font_size: Minimum font size in points
            max_font_size: Maximum font size in points
            char_width_ratio: Character width ratio (default 0.6 for Verdana)
            line_height_ratio: Line height ratio (default 1.3 for 1.25 spacing)

        Returns:
            Calculated font size in points (guaranteed to fit within target_height)
        """
        if not text or not text.strip() or target_height <= 0:
            return min_font_size

        text_without_newlines = text.replace("\n", " ")
        total_chars = len(text_without_newlines)
        words = text.split()

        if total_chars == 0 or available_width <= 0:
            return min_font_size

        # Use conservative approach: calculate font size that ensures content fits
        # Start with a conservative estimate and verify it fits

        # Use 70% of target height as safety margin (very conservative)
        # Word's rendering can differ significantly from our estimates
        safe_height = target_height * 0.70

        # Initial estimate based on character count
        # Formula: height = total_chars * font_size^2 * char_width_ratio * line_height_ratio / (available_width * 72^2)
        # Solving for font_size:
        numerator = safe_height * available_width * 72 * 72
        denominator = total_chars * char_width_ratio * line_height_ratio

        if denominator > 0:
            font_size_squared = numerator / denominator
            font_size = font_size_squared**0.5

            # Refine iteratively for word boundaries (more conservative)
            for _ in range(5):  # More iterations for accuracy
                chars_per_line = (
                    (available_width * 72) / (font_size * char_width_ratio)
                    if font_size > 0
                    else 1
                )
                # Use word-based calculation for better accuracy
                words_per_line = max(1, chars_per_line / 6)  # Average 6 chars per word
                estimated_wrapped_lines = max(
                    1, (len(words) / words_per_line) * 1.15
                )  # 15% overhead for word boundaries
                new_font_size = (safe_height * 72) / (
                    estimated_wrapped_lines * line_height_ratio
                )
                font_size = (
                    font_size * 0.5 + new_font_size * 0.5
                )  # More conservative weighted average
        else:
            font_size = min_font_size

        # Clamp to bounds
        font_size = max(min_font_size, min(int(font_size), max_font_size))

        # Final verification: ensure it actually fits
        # Recalculate with final font size
        chars_per_line = (
            (available_width * 72) / (font_size * char_width_ratio)
            if font_size > 0
            else 1
        )
        words_per_line = max(1, chars_per_line / 6)
        final_estimated_lines = max(1, (len(words) / words_per_line) * 1.15)
        final_estimated_height = (
            font_size * line_height_ratio * final_estimated_lines
        ) / 72

        # If estimated height exceeds target, reduce font size further
        if final_estimated_height > target_height:
            # Reduce proportionally with very conservative safety margin
            # Word's rendering can differ, so we use 85% safety margin
            reduction_factor = (
                target_height / final_estimated_height * 0.85
            )  # 85% safety margin (very conservative)
            font_size = int(font_size * reduction_factor)
            font_size = max(min_font_size, min(font_size, max_font_size))

        return font_size

    def _parse_week_of_date(self, week_of: str) -> Optional[datetime]:
        """
        Parse week_of string to get Monday's date.

        Handles formats like:
        - "11/18/2024"
        - "Week of 11/18/2024"
        - "2024-11-18"
        - "10/06-10/10" (date range, use first date)
        - "11-17-11-21" (date range, use first date)
        """
        # Remove "Week of" prefix if present
        date_str = week_of.replace("Week of", "").strip()

        # Try MM/DD/YYYY format
        match = re.search(r"(\d{1,2})/(\d{1,2})/(\d{4})", date_str)
        if match:
            month, day, year = match.groups()
            try:
                return datetime(int(year), int(month), int(day))
            except ValueError:
                pass

        # Try MM/DD-MM/DD format (date range, use first date)
        match = re.search(r"(\d{1,2})/(\d{1,2})-", date_str)
        if match:
            month, day = match.groups()
            # Assume current year or previous year if month > current month
            current_year = datetime.now().year
            try:
                parsed_date = datetime(current_year, int(month), int(day))
                # If date is in the future, assume previous year
                if parsed_date > datetime.now():
                    parsed_date = datetime(current_year - 1, int(month), int(day))
                return parsed_date
            except ValueError:
                pass

        # Try MM-DD-MM-DD format (date range, use first date)
        match = re.search(r"(\d{1,2})-(\d{1,2})-", date_str)
        if match:
            month, day = match.groups()
            # Assume current year or previous year if month > current month
            current_year = datetime.now().year
            try:
                parsed_date = datetime(current_year, int(month), int(day))
                # If date is in the future, assume previous year
                if parsed_date > datetime.now():
                    parsed_date = datetime(current_year - 1, int(month), int(day))
                return parsed_date
            except ValueError:
                pass

        # Try YYYY-MM-DD format
        match = re.search(r"(\d{4})-(\d{1,2})-(\d{1,2})", date_str)
        if match:
            year, month, day = match.groups()
            try:
                return datetime(int(year), int(month), int(day))
            except ValueError:
                pass

        return None

    def _get_day_date(self, week_of: str, day_name: str) -> str:
        """
        Get the date for a specific day of the week.

        Args:
            week_of: Week identifier (e.g., "Week of 11/18/2024")
            day_name: Day name (e.g., "monday", "tuesday")

        Returns:
            Formatted date string (MM/DD/YYYY) or week_of if parsing fails
        """
        monday_date = self._parse_week_of_date(week_of)
        if not monday_date:
            return week_of

        # Calculate day offset
        day_index = (
            self.DAY_NAMES.index(day_name.lower())
            if day_name.lower() in self.DAY_NAMES
            else 0
        )
        target_date = monday_date + timedelta(days=day_index)

        # Format as MM/DD/YYYY
        return target_date.strftime("%m/%d/%Y")

    def _add_gray_line(self, paragraph):
        """Add a thin gray horizontal line to a paragraph."""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.75pt line
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "808080")  # 50% gray
        pBdr.append(bottom)

    def generate_docx(
        self,
        lesson_json: Dict[str, Any],
        output_path: str,
        user_name: Optional[str] = None,
        week_of: Optional[str] = None,
    ) -> str:
        """
        Generate DOCX file with objectives, one lesson per page.

        Each page contains:
        - Header: Date | Subject | Grade | Homeroom (10pt)
        - Student Goal: 3/4 of page (base 48pt, bold)
        - Separator: Thin gray line
        - WIDA Objective: 1/4 of page (base 14pt, 50% gray)

        Args:
            lesson_json: Lesson plan JSON structure
            output_path: Path to save DOCX file
            user_name: Optional teacher name for file naming
            week_of: Optional week identifier

        Returns:
            Path to generated DOCX file
        """
        # Extract objectives
        objectives = self.extract_objectives(lesson_json)

        if not objectives:
            raise ValueError("No objectives found in lesson plan")

        # Get metadata
        metadata = lesson_json.get("metadata", {})
        week_of = week_of or metadata.get("week_of", "Unknown")
        grade = metadata.get("grade", "Unknown")
        subject = metadata.get("subject", "Unknown")
        homeroom = metadata.get("homeroom", "Unknown")

        # Create document
        doc = Document()

        # Set page orientation to landscape
        # In python-docx, you must swap width and height to get landscape
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap page dimensions: US Letter portrait is 8.5" x 11", landscape is 11" x 8.5"
        section.page_width = Inches(11)  # Landscape width
        section.page_height = Inches(8.5)  # Landscape height

        # Set margins to 0.5 inches
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # Page dimensions (landscape US Letter: 11" × 8.5")
        # With 0.5" margins: usable area is 10" × 7.5"
        usable_width = 10.0  # inches (11" - 0.5" - 0.5")
        usable_height = 7.5  # inches (8.5" - 0.5" - 0.5")

        # Section dimensions - precise allocation for exact 75%/25% split
        # Header takes minimal space
        header_height_estimate = 0.3  # inches (10pt font + spacing - more conservative)

        # Calculate remaining height after header - this is the space for objectives
        objectives_area_height = usable_height - header_height_estimate  # ~7.2"

        # Account for separator between sections (~0.15" including line and spacing - more conservative)
        separator_height = 0.15
        available_for_objectives = objectives_area_height - separator_height  # ~7.05"

        # Precise allocation: Student Goal = 75%, WIDA = 25% of available_for_objectives
        # Use 95% of calculated heights to ensure we don't exceed page
        student_goal_height = (
            available_for_objectives * 0.75
        ) * 0.95  # ~5.02" (conservative)
        wida_height = (available_for_objectives * 0.25) * 0.95  # ~1.67" (conservative)

        # Content areas (with minimal padding for text readability)
        # Use conservative padding to ensure content fits
        student_goal_content_width = usable_width - 0.4  # 0.2" padding each side
        student_goal_content_height = (
            student_goal_height - 0.1
        )  # 0.05" padding top/bottom

        wida_content_width = usable_width - 0.4  # 0.2" padding each side
        wida_content_height = (
            wida_height - 0.1
        )  # 0.05" padding top/bottom (will account for label separately)

        # Generate one page per lesson (slot per day)
        # Use sections for more reliable page breaks - Word respects section breaks better
        for i, obj in enumerate(objectives):
            # Create a new section for each objective (except the first one uses the default section)
            if i > 0:
                # Add a new section that starts on a new page
                # This is more reliable than paragraph page breaks
                doc.add_section(WD_SECTION.NEW_PAGE)
                # Configure the new section with same settings
                section = doc.sections[-1]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11)
                section.page_height = Inches(8.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Get the current section
            section = doc.sections[-1]

            # Header: Date | Subject | Grade | Homeroom (10pt)
            # Get date for this specific day
            day_date = self._get_day_date(week_of, obj["day"])
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(
                f"{day_date} | {obj['subject']} | Grade {grade} | {homeroom}"
            )
            header_run.font.size = Pt(10)
            header_run.font.name = "Calibri"
            # Minimal spacing after header - Student Goal starts immediately after
            header_para.paragraph_format.space_after = Pt(2)
            # Prevent page breaks within this objective - keep header with next content
            header_para.paragraph_format.keep_with_next = True

            # STEP 1: Calculate WIDA Objective height first (fixed 12pt font)
            # This ensures we know exactly how much space WIDA will take
            wida_text = obj.get("wida_objective", "").strip()
            if not wida_text:
                wida_text = "No WIDA Objective specified"

            # WIDA font size is fixed at 12pt
            wida_font_size = 12

            # Calculate actual WIDA text height with 12pt font
            wida_words = wida_text.split()
            wida_chars_per_line = (
                (wida_content_width * 72) / (wida_font_size * 0.55)
                if wida_font_size > 0
                else 60
            )
            wida_words_per_line = max(1, wida_chars_per_line / 6)
            wida_estimated_lines = max(1, len(wida_words) / wida_words_per_line)
            actual_wida_text_height = (
                wida_font_size * 1.2 * wida_estimated_lines
            ) / 72  # in inches (1.2 line spacing)

            # Calculate label height (12pt font)
            label_font_size = 12
            label_height_inches = (
                label_font_size * 1.2
            ) / 72  # Approximate label height

            # Total WIDA content height (label + spacing + text)
            wida_label_spacing = 0.03  # Space between label and text
            total_wida_content_height = (
                label_height_inches + wida_label_spacing + actual_wida_text_height
            )

            # Add some padding to WIDA section for safety
            wida_section_padding = 0.1  # 0.1" total padding (top + bottom)
            total_wida_section_height = total_wida_content_height + wida_section_padding

            # STEP 2: Calculate Student Goal font size to fit in remaining space
            student_goal_text = obj.get("student_goal", "").strip()
            if not student_goal_text:
                student_goal_text = "No Student Goal specified"

            # Calculate available space for Student Goal
            # Total usable height = objectives_area_height (after header)
            # Student Goal gets: total - separator - WIDA
            separator_height = 0.15  # Separator line height
            available_for_student_goal = (
                objectives_area_height - separator_height - total_wida_section_height
            )

            # Use 75% of available space to ensure it fits (very conservative with large safety margin)
            # Word's rendering can differ from our estimates, so we need extra margin
            student_goal_target_height = available_for_student_goal * 0.75

            # Calculate font size to fit within available height
            student_goal_font_size = self.calculate_font_size_to_fill_height(
                student_goal_text,
                student_goal_content_width,
                student_goal_target_height,  # Use calculated available space
                min_font_size=12,  # Minimum font size
                max_font_size=60,  # Maximum font size
                char_width_ratio=0.6,  # Verdana Bold
                line_height_ratio=1.3,  # 1.25 spacing + baseline
            )

            # Add Student Goal paragraph - starts at top of objectives area
            student_goal_para = doc.add_paragraph()
            student_goal_run = student_goal_para.add_run(student_goal_text)
            student_goal_run.font.size = Pt(student_goal_font_size)
            student_goal_run.font.bold = True
            student_goal_run.font.name = "Verdana"
            # Enable word wrapping and set line spacing
            # Prevent page breaks to ensure content stays on same page
            # Use keep_together to prevent paragraph from splitting across pages
            student_goal_para.paragraph_format.widow_control = True
            student_goal_para.paragraph_format.keep_together = (
                True  # Keep all lines together
            )
            student_goal_para.paragraph_format.keep_with_next = (
                True  # Keep with separator
            )
            student_goal_para.paragraph_format.page_break_before = False
            student_goal_para.paragraph_format.line_spacing = 1.25
            # No space before - starts immediately after header (top-aligned)
            student_goal_para.paragraph_format.space_before = Pt(0)

            # Calculate actual Student Goal text height to verify it fits
            words = student_goal_text.split()
            chars_per_line = (
                (student_goal_content_width * 72) / (student_goal_font_size * 0.6)
                if student_goal_font_size > 0
                else 60
            )
            words_per_line = max(1, chars_per_line / 6)
            estimated_lines = max(1, len(words) / words_per_line)
            # Actual height = font_size * line_spacing * lines / 72
            actual_student_goal_height = (
                student_goal_font_size * 1.3 * estimated_lines
            ) / 72  # in inches

            # Calculate spacing to fill available space (but don't exceed it)
            # Available space = available_for_student_goal (calculated above)
            # If text is shorter than allocated space, add minimal spacing
            remaining_student_goal_space = (
                available_for_student_goal - actual_student_goal_height
            )
            if remaining_student_goal_space > 0.1:  # Only if significant space remains
                # Add spacing after to fill some of the space (very conservative)
                # Use 30% of remaining space to be extremely conservative and prevent overflow
                # Word's rendering can add extra space, so we use less
                student_goal_para.paragraph_format.space_after = Pt(
                    int(remaining_student_goal_space * 72 * 0.30)
                )
            else:
                # Text fills or nearly fills space, minimal spacing
                student_goal_para.paragraph_format.space_after = Pt(1)

            # Separator: Thin gray line (minimal height)
            # Separator is accounted for in the space calculations above
            separator_para = doc.add_paragraph()
            self._add_gray_line(separator_para)
            separator_para.paragraph_format.space_before = Pt(
                0
            )  # No space - part of Student Goal spacing
            separator_para.paragraph_format.space_after = Pt(
                0
            )  # No space - WIDA spacing handles positioning
            # Prevent page breaks - keep separator with both sections
            separator_para.paragraph_format.page_break_before = False
            separator_para.paragraph_format.keep_together = True
            separator_para.paragraph_format.keep_with_next = (
                True  # Keep with WIDA section
            )

            # WIDA Objective section - bottom-aligned, fixed 12pt font
            # Height already calculated above (total_wida_section_height)

            # Calculate spacing needed to position WIDA at bottom
            # Available space for WIDA = objectives_area_height - header - student_goal - separator
            # We want WIDA to be positioned at the bottom of the objectives area
            # Calculate remaining space after Student Goal and separator
            remaining_wida_space = (
                objectives_area_height
                - actual_student_goal_height
                - separator_height
                - total_wida_section_height
            )

            # Add "WIDA/Bilingual:" label
            wida_label_para = doc.add_paragraph()
            wida_label_run = wida_label_para.add_run("WIDA/Bilingual:")
            wida_label_run.font.size = Pt(wida_font_size)
            wida_label_run.font.bold = True
            wida_label_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            wida_label_run.font.name = "Calibri"

            # Position WIDA section at bottom: add space before label to push entire section down
            # Use 30% of remaining space to be extremely conservative and prevent overflow
            # Word's rendering can add extra space, so we use less
            if remaining_wida_space > 0.05:  # Only if significant space remains
                wida_label_para.paragraph_format.space_before = Pt(
                    int(remaining_wida_space * 72 * 0.30)
                )
            else:
                wida_label_para.paragraph_format.space_before = Pt(0)
            wida_label_para.paragraph_format.space_after = Pt(1)  # Reduced spacing

            # Add WIDA Objective text
            wida_para = doc.add_paragraph()
            wida_run = wida_para.add_run(wida_text)
            wida_run.font.size = Pt(wida_font_size)
            wida_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            wida_run.font.name = "Calibri"
            wida_para.paragraph_format.line_spacing = 1.0
            # Prevent page breaks to ensure content stays on same page
            # Keep WIDA content together and with the label
            wida_label_para.paragraph_format.keep_together = True
            wida_label_para.paragraph_format.keep_with_next = (
                True  # Keep label with WIDA text
            )
            wida_para.paragraph_format.widow_control = True
            wida_para.paragraph_format.keep_together = True  # Keep all lines together
            wida_para.paragraph_format.keep_with_next = False
            wida_para.paragraph_format.page_break_before = False
            wida_para.paragraph_format.space_after = Pt(0)  # No space after (at bottom)

        # Save document
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))

        logger.info(
            "objectives_docx_generated",
            extra={
                "output_path": str(output_file),
                "objective_count": len(objectives),
                "week_of": week_of,
            },
        )

        return str(output_file)


def print_objectives_from_file(
    json_path: str, output_path: Optional[str] = None
) -> str:
    """Convenience function to print objectives from JSON file.

    Args:
        json_path: Path to lesson plan JSON file
        output_path: Optional path to save formatted output

    Returns:
        Formatted objectives string
    """
    printer = ObjectivesPrinter()
    return printer.print_objectives(
        json_path, source_type="file", output_path=output_path
    )


def print_objectives_from_plan(plan_id: str, output_path: Optional[str] = None) -> str:
    """Convenience function to print objectives from database plan.

    Args:
        plan_id: Plan ID from weekly_plans table
        output_path: Optional path to save formatted output

    Returns:
        Formatted objectives string
    """
    printer = ObjectivesPrinter()
    return printer.print_objectives(
        plan_id, source_type="database", output_path=output_path
    )
