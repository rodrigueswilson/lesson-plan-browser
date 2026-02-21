"""HTML + PDF generator for sentence frames with precise layout control.

This module mirrors the ObjectivesPDFGenerator pattern but focuses on
printable, English-only sentence frame sheets (two pages per lesson,
with fold lines and thirds).
"""

from __future__ import annotations

import html as html_module
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.config import settings
from backend.file_manager import get_file_manager
from backend.services.sorting_utils import sort_slots
from backend.telemetry import logger
from backend.utils.metadata_utils import (
    build_document_header,
    get_day_date,
    get_homeroom,
    get_subject,
    get_teacher_name,
)


class SentenceFramesPDFGenerator:
    """Generate sentence frames as HTML and convert to PDF for printing."""

    def __init__(self) -> None:
        self.css_template = self._get_css_template()
        self.html_template = self._get_html_template()
        self.file_manager = get_file_manager()
        self._default_output_dir = Path(self.file_manager.base_path)

    # --- Path helpers (mirroring objectives generator) ---

    def _resolve_output_directory(self, lesson_json: Dict[str, Any]) -> Path:
        metadata = lesson_json.get("metadata", {})
        week_of = metadata.get("week_of")
        if week_of:
            try:
                week_folder = Path(self.file_manager.get_week_folder(week_of))
                week_folder.mkdir(parents=True, exist_ok=True)
                return week_folder
            except Exception as exc:
                logger.warning(
                    "sentence_frames_week_folder_resolution_failed",
                    extra={"week_of": week_of, "error": str(exc)},
                )

        self._default_output_dir.mkdir(parents=True, exist_ok=True)
        return self._default_output_dir

    def _sanitize_for_filename(self, value: str, fallback: str) -> str:
        import re

        clean = re.sub(r"[^A-Za-z0-9]+", "_", (value or "")).strip("_")
        return clean or fallback

    def _build_default_basename(
        self, lesson_json: Dict[str, Any], user_name: Optional[str]
    ) -> str:
        metadata = lesson_json.get("metadata", {})
        teacher = get_teacher_name(metadata, user_name=user_name) or "Teacher"
        week_label = metadata.get("week_of") or datetime.now().strftime("%m-%d")
        teacher_slug = self._sanitize_for_filename(teacher, "Teacher")
        week_slug = self._sanitize_for_filename(week_label.replace("/", "-"), "Week")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{teacher_slug}_SentenceFrames_{week_slug}_{timestamp}"

    def _resolve_html_path(
        self,
        lesson_json: Dict[str, Any],
        user_name: Optional[str],
        output_path: Optional[str],
    ) -> Path:
        if output_path:
            target = Path(output_path)
            target.parent.mkdir(parents=True, exist_ok=True)
            return target

        directory = self._resolve_output_directory(lesson_json)
        base_name = self._build_default_basename(lesson_json, user_name)
        target = directory / f"{base_name}.html"
        target.parent.mkdir(parents=True, exist_ok=True)
        return target

    def _resolve_pdf_and_html_paths(
        self,
        lesson_json: Dict[str, Any],
        user_name: Optional[str],
        pdf_path: Optional[str],
    ) -> Tuple[Path, Path]:
        if pdf_path:
            pdf_file = Path(pdf_path)
            html_file = pdf_file.with_suffix(".html")
            pdf_file.parent.mkdir(parents=True, exist_ok=True)
            html_file.parent.mkdir(parents=True, exist_ok=True)
            return html_file, pdf_file

        directory = self._resolve_output_directory(lesson_json)
        base_name = self._build_default_basename(lesson_json, user_name)
        html_file = directory / f"{base_name}.html"
        pdf_file = directory / f"{base_name}.pdf"
        html_file.parent.mkdir(parents=True, exist_ok=True)
        return html_file, pdf_file

    # --- CSS / HTML templates ---

    @staticmethod
    def _convert_markdown_to_html(text: str) -> str:
        """Convert markdown bold syntax (**word**) to HTML <strong> tags.

        This function handles vocabulary words that are wrapped in markdown
        double asterisks and converts them to proper HTML formatting.
        """
        # Escape any existing HTML to prevent XSS
        text = html_module.escape(text)

        # Convert **word** to <strong>word</strong>
        # Pattern matches **word** where word can contain spaces, punctuation, etc.
        # Uses non-greedy matching to handle multiple bold words in one sentence
        pattern = r"\*\*(.+?)\*\*"
        text = re.sub(pattern, r"<strong>\1</strong>", text)

        return text

    @staticmethod
    def _extract_bold_words_from_markdown(text: str) -> Tuple[str, List[str]]:
        """Extract bold words from markdown and return cleaned text with list of bold words.

        Returns:
            Tuple of (cleaned_text, list_of_bold_words)
        """
        bold_words: List[str] = []
        pattern = r"\*\*([^*]+?)\*\*"

        # Find all bold words
        matches = re.finditer(pattern, text)
        for match in matches:
            bold_words.append(match.group(1))

        # Remove markdown syntax
        cleaned_text = re.sub(pattern, r"\1", text)

        return cleaned_text, bold_words

    def _add_bold_text_to_paragraph(
        self,
        paragraph,
        text: str,
        bold_words: List[str],
        font_size: int,
        font_name: str = "Source Sans Pro",
    ):
        """Add text to a paragraph with specified words in bold.

        Args:
            paragraph: docx Paragraph object
            text: Text to add
            bold_words: List of words that should be bold (case-insensitive)
            font_size: Font size in points
            font_name: Font name
        """
        if not bold_words:
            run = paragraph.add_run(text)
            run.font.size = Pt(font_size)
            run.font.bold = True
            run.font.name = font_name
            return

        # Create a case-insensitive pattern for all bold words
        words_pattern = re.compile(
            r"\b(" + "|".join(re.escape(word) for word in bold_words) + r")\b",
            re.IGNORECASE,
        )

        last_index = 0
        for match in words_pattern.finditer(text):
            # Add text before the match
            if match.start() > last_index:
                run = paragraph.add_run(text[last_index : match.start()])
                run.font.size = Pt(font_size)
                run.font.bold = True
                run.font.name = font_name

            # Add bold word (extra bold)
            bold_run = paragraph.add_run(match.group(0))
            bold_run.font.size = Pt(font_size)
            bold_run.font.bold = True
            bold_run.font.name = font_name

            last_index = match.end()

        # Add remaining text
        if last_index < len(text):
            run = paragraph.add_run(text[last_index:])
            run.font.size = Pt(font_size)
            run.font.bold = True
            run.font.name = font_name

    def _get_css_template(self) -> str:
        """Return CSS for sentence frame pages.

        Layout assumptions:
        - Portrait US Letter (8.5" x 11") with no CSS margins (full page usage).
        - Content area: 8.5" wide x 11" tall (full page dimensions).
        - Header: 0.35" tall at the top of each page (positioned 1 cm from top/left for printer safety).
        - Three equal-height panels: each panel is exactly 3.67" tall.
          Calculation: 11" page height / 3 = 3.67" per panel (for accurate folding).
        - Panels use explicit heights (not flexbox) for reliable print output.
        - Fold lines are positioned at exact panel boundaries:
          - First fold: 3.67" from top of page (end of panel 1)
          - Second fold: 7.33" from top of page (end of panel 2)
        - Note: Some printers have non-printable margins (typically 0.1-0.2") which may slightly
          affect the exact measurements. For best results, print at 100% scale with no margins.
        - English-only frames, with a mandatory function label.
        """
        return """
        @import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@400;600&display=swap');

        @page {
            size: 8.5in 11in; /* Portrait US Letter - exact size for printer */
            margin: 0; /* No margins - use full page for maximum space */
            /* Ensure no scaling - printers should respect exact page size */
        }

        * {
            box-sizing: border-box;
        }

        body {
            font-family: 'Source Sans Pro', 'Inter', 'Arial', sans-serif;
            margin: 0;
            padding: 0;
            background: white;
        }

        .frames-page {
            width: 8.5in;  /* Full page width */
            height: 11in;  /* Full page height */
            position: relative;
            margin: 0;
            page-break-after: always;
            display: flex;
            flex-direction: column;
        }

        .header {
            position: absolute;
            top: 0.39in; /* 1 cm from top for printer safety */
            left: 0.39in; /* 1 cm from left for printer safety */
            height: 0.35in;
            font-size: 10pt;
            color: #808080; /* 50% gray */
            display: flex;
            align-items: center;
            justify-content: flex-start;
            z-index: 10; /* Ensure header appears above panels */
            border: none;
            border-bottom: none;
        }

        .panels {
            position: relative;
            height: 11in; /* Full page height for accurate three-way division */
            display: flex;
            flex-direction: column;
            border-top: none;
        }

        .panel {
            height: 3.67in; /* 11in / 3 = 3.67in per panel (exact equal division for folding) */
            box-sizing: border-box; /* Include padding in height calculation */
            padding: 0.25in 0.3in;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            text-align: center;
            flex-shrink: 0; /* Prevent panels from shrinking */
            flex-grow: 0; /* Prevent panels from growing */
        }

        .panel-title {
            font-size: 10pt;
            font-weight: 600;
            color: #808080;
            margin-bottom: 0.1in;
        }

        .frame-text {
            font-size: 28pt;
            font-weight: 600;
            color: #000000;
            line-height: 1.2;
            margin-bottom: 0.1in;
            word-wrap: break-word;
            max-width: 100%;
        }

        .frame-text strong {
            font-weight: 700; /* Ensure bold vocabulary words are clearly visible */
        }

        .frame-function {
            font-size: 11pt;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #808080;
        }

        .bundle-frame-row {
            margin: 0.05in 0;
        }

        .bundle-frame-text {
            font-size: 18pt;
            font-weight: 600;
            color: #000000;
            line-height: 1.2;
        }

        .bundle-frame-text strong {
            font-weight: 700; /* Ensure bold vocabulary words are clearly visible */
        }

        .bundle-frame-function {
            font-size: 9pt;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #808080;
        }

        .fold-line {
            position: absolute;
            left: 0;
            right: 0;
            height: 0;
            border-top: 1px solid rgba(128, 128, 128, 0.4); /* Thin gray line at 40% opacity */
            z-index: 1;
            /* Use fixed positions based on full 11" page height for printer accuracy */
        }

        /* Fold lines positioned to match exact panel divisions */
        /* Since .panels has position: relative, fold lines are positioned relative to .panels container */
        /* Each panel is exactly 3.67in tall (11in / 3) */
        /* First fold: at the end of first panel = 3.67in from top of .panels */
        .fold-line-1 {
            top: 3.67in; /* End of first panel, start of second panel */
        }

        /* Second fold: at the end of second panel = 3.67in + 3.67in = 7.33in from top of .panels */
        .fold-line-2 {
            top: 7.33in; /* End of second panel, start of third panel */
        }

        @media print {
            .frames-page {
                page-break-after: always;
                /* Ensure exact dimensions for print */
                width: 8.5in !important;
                height: 11in !important;
            }
            .frames-page:last-child {
                page-break-after: auto;
            }
            .panels {
                height: 11in !important; /* Enforce full page height for print */
            }
            .panel {
                height: 3.67in !important; /* Enforce exact panel height for print (11in / 3) */
                flex-shrink: 0 !important;
                flex-grow: 0 !important;
            }
        }
        """

    def _get_html_template(self) -> str:
        """Return outer HTML template for sentence frame pages."""
        return """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Sentence Frames - Week of {week_of}</title>
            <style>
                {css}
            </style>
        </head>
        <body>
            {pages}
        </body>
        </html>
        """

    # --- Extraction ---

    def extract_sentence_frames(
        self, lesson_json: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Extract sentence frame payloads from lesson JSON, ordered by schedule (day and slot).

        Each item is a dict with metadata and three level buckets:
        - levels_1_2: list of 3 frames
        - levels_3_4: list of 3 frames
        - levels_5_6: list of 2 frames

        Frames are ordered by day (Monday-Friday) and then by start_time (chronological) to follow schedule order.
        """
        results: List[Dict[str, Any]] = []

        days = lesson_json.get("days") or {}
        if not isinstance(days, dict):
            return results

        metadata = lesson_json.get("metadata", {})
        week_of = metadata.get("week_of", "Unknown")
        default_grade = metadata.get("grade", "Unknown")
        default_room = metadata.get("room", "")
        default_teacher_name = get_teacher_name(metadata)

        day_names = ["monday", "tuesday", "wednesday", "thursday", "friday"]

        for day_name in day_names:
            day_data = days.get(day_name)
            if not isinstance(day_data, dict):
                continue

            # Check if this is a multi-slot structure
            slots = day_data.get("slots") or []
            if isinstance(slots, list) and len(slots) > 0:
                # Multi-slot: process each slot individually, ordered by slot_number
                # Sort slots by start_time (chronological) with slot_number as fallback
                sorted_slots = sort_slots(slots)
                from backend.telemetry import logger

                logger.debug(
                    f"Extracting sentence frames for {day_name}: {len(sorted_slots)} slots found (slot_numbers: {[s.get('slot_number', 0) for s in sorted_slots]})"
                )

                for slot in sorted_slots:
                    if not isinstance(slot, dict):
                        continue

                    slot_num = slot.get("slot_number", 0)
                    slot_unit_lesson = slot.get("unit_lesson", "")

                    # Skip "No School" entries
                    if (
                        slot_unit_lesson
                        and slot_unit_lesson.strip().lower() == "no school"
                    ):
                        logger.debug(
                            f"Skipping slot {slot_num} ({day_name}) - No School entry"
                        )
                        continue

                    # Get slot-specific sentence frames
                    slot_frames = slot.get("sentence_frames") or []
                    if not isinstance(slot_frames, list) or len(slot_frames) == 0:
                        # Include slot even if it has no sentence frames - use empty structure
                        logger.debug(
                            f"Slot {slot_num} ({day_name}) has no sentence frames - including with empty frames"
                        )
                        slot_frames = []
                    else:
                        logger.debug(
                            f"Processing slot {slot_num} ({day_name}) for sentence frames extraction ({len(slot_frames)} frames)"
                        )

                    # Continue processing even with empty frames to include the slot

                    # Use slot-specific metadata, fallback to defaults
                    slot_grade = slot.get("grade", default_grade)
                    # Get subject using standardized helper (metadata only)
                    slot_subject = get_subject(metadata, slot=slot)
                    # Use standardized helper to prevent homeroom leakage
                    slot_homeroom = get_homeroom(metadata, slot=slot)
                    # Similar logic for room: use slot room if present, otherwise default
                    slot_room = slot.get("room")
                    if slot_room is None:
                        slot_room = default_room
                    slot_teacher = get_teacher_name(metadata, slot=slot)

                    # Organize frames by proficiency level
                    levels_1_2: List[Dict[str, Any]] = []
                    levels_3_4: List[Dict[str, Any]] = []
                    levels_5_6: List[Dict[str, Any]] = []

                    for frame in slot_frames:
                        if not isinstance(frame, dict):
                            continue
                        level = frame.get("proficiency_level")
                        if level == "levels_1_2":
                            levels_1_2.append(frame)
                        elif level == "levels_3_4":
                            levels_3_4.append(frame)
                        elif level == "levels_5_6":
                            levels_5_6.append(frame)

                    # Include slot even if it has no frames - this ensures slots that appear
                    # in DOCX also appear in sentence frames output
                    # if not (levels_1_2 or levels_3_4 or levels_5_6):
                    #     continue

                    # Homeroom already handled by get_homeroom helper (prevents leakage)

                    # Handle room similar to homeroom - avoid leakage between slots
                    final_room = slot_room
                    if (
                        not final_room
                        or final_room == "N/A"
                        or final_room.strip() == ""
                    ):
                        if default_room and default_room != "Unknown":
                            final_room = default_room
                        else:
                            final_room = ""

                    results.append(
                        {
                            "week_of": week_of,
                            "day": day_name.capitalize(),
                            "grade": slot_grade
                            if slot_grade and slot_grade != "N/A"
                            else default_grade,
                            "subject": slot_subject,  # Already uses get_subject helper with proper fallback
                            "homeroom": slot_homeroom,  # Already uses get_homeroom helper (prevents leakage)
                            "room": final_room,
                            "teacher_name": slot_teacher
                            if slot_teacher != "Unknown"
                            else default_teacher_name,
                            "unit_lesson": slot_unit_lesson,
                            "slot_number": slot.get("slot_number", 0),
                            "start_time": slot.get("start_time", ""),  # Include start_time for sorting
                            "levels_1_2": levels_1_2,
                            "levels_3_4": levels_3_4,
                            "levels_5_6": levels_5_6,
                        }
                    )
            else:
                # Single-slot or day-level frames: process day-level frames if present
                day_unit_lesson = day_data.get("unit_lesson", "")

                # Skip "No School" entries
                if day_unit_lesson and day_unit_lesson.strip().lower() == "no school":
                    continue

                day_level_frames = day_data.get("sentence_frames") or []
                if not isinstance(day_level_frames, list) or len(day_level_frames) == 0:
                    continue

                # Organize frames by proficiency level
                levels_1_2: List[Dict[str, Any]] = []
                levels_3_4: List[Dict[str, Any]] = []
                levels_5_6: List[Dict[str, Any]] = []

                for frame in day_level_frames:
                    if not isinstance(frame, dict):
                        continue
                    level = frame.get("proficiency_level")
                    if level == "levels_1_2":
                        levels_1_2.append(frame)
                    elif level == "levels_3_4":
                        levels_3_4.append(frame)
                    elif level == "levels_5_6":
                        levels_5_6.append(frame)

                if not (levels_1_2 or levels_3_4 or levels_5_6):
                    continue

                results.append(
                    {
                        "week_of": week_of,
                        "day": day_name.capitalize(),
                        "grade": default_grade,
                        "subject": get_subject(metadata),  # Use standardized helper
                        "homeroom": get_homeroom(metadata),  # Use standardized helper
                        "room": default_room if default_room else "",
                        "teacher_name": get_teacher_name(metadata),
                        "unit_lesson": day_data.get("unit_lesson", ""),
                        "slot_number": 0,
                        "levels_1_2": levels_1_2,
                        "levels_3_4": levels_3_4,
                        "levels_5_6": levels_5_6,
                        }
                    )

        # Sort results by day (Monday-Friday) and then by start_time (chronological)
        # This ensures consistent ordering matching the schedule and objectives
        day_order = {"Monday": 0, "Tuesday": 1, "Wednesday": 2, "Thursday": 3, "Friday": 4}
        
        def get_result_sort_key(result: Dict[str, Any]) -> tuple:
            day = result.get("day", "")
            day_idx = day_order.get(day, 99)
            start_time = result.get("start_time", "") or ""
            slot_num = result.get("slot_number", 0)
            
            # Convert time to sortable format (HH:MM -> minutes since midnight)
            time_sort = 0
            if start_time:
                try:
                    parts = str(start_time).split(":")
                    if len(parts) >= 2:
                        time_sort = int(parts[0]) * 60 + int(parts[1])
                except (ValueError, TypeError):
                    pass
            
            return (day_idx, time_sort, slot_num)
        
        results.sort(key=get_result_sort_key)
        return results

    # --- HTML generation ---

    def _calculate_day_date(self, week_of: str, day: str) -> str:
        """Calculate the specific date for a given day in the week using standardized utility.

        Args:
            week_of: Week range string like "11/17-11/21" or "11/17/2025-11/21/2025"
            day: Day name like "Monday", "Tuesday", etc.

        Returns:
            Formatted date string "MM/DD/YY" (e.g., "11/17/25")
        """
        # Get school year from config if available
        config_school_year = None
        if settings.SCHOOL_YEAR_START_YEAR and settings.SCHOOL_YEAR_END_YEAR:
            config_school_year = (
                settings.SCHOOL_YEAR_START_YEAR,
                settings.SCHOOL_YEAR_END_YEAR,
            )

        # Use standardized utility
        full_date = get_day_date(week_of, day, config_school_year=config_school_year)

        # Convert MM/DD/YYYY to MM/DD/YY format
        if full_date and full_date != week_of:
            try:
                date_obj = datetime.strptime(full_date, "%m/%d/%Y")
                return date_obj.strftime("%m/%d/%y")
            except ValueError:
                return full_date

        return week_of

    def _build_header_text(
        self, payload: Dict[str, Any], metadata: Dict[str, Any]
    ) -> str:
        """
        Build standardized header text using metadata_utils helper.

        Args:
            payload: Payload dictionary with day, week_of, subject, grade, homeroom, room
            metadata: Lesson-level metadata dictionary

        Returns:
            Formatted header string
        """
        # Calculate day date
        week_of = payload.get("week_of", "Unknown")
        day_name = payload.get("day", "Unknown")
        day_date = None
        if week_of and week_of != "Unknown" and day_name and day_name != "Unknown":
            day_date = self._calculate_day_date(week_of, day_name)

        # Reconstruct slot dict for header builder (if slot-specific)
        slot_dict = None
        if (
            payload.get("subject")
            or payload.get("grade")
            or payload.get("homeroom")
            or payload.get("room")
        ):
            slot_dict = {
                "subject": payload.get("subject"),
                "grade": payload.get("grade"),
                "homeroom": payload.get("homeroom"),
                "room": payload.get("room"),
            }

        return build_document_header(
            metadata=metadata,
            slot=slot_dict,
            day_date=day_date,
            day_name=day_name if day_name != "Unknown" else None,
            include_time=False,  # Sentence Frames PDF doesn't include time
            include_day_name=True,  # Include day name for consistency
        )

    @staticmethod
    def _pretty_function(label: str) -> str:
        if not label:
            return ""
        text = label.replace("_", " ").strip()
        return text[:1].upper() + text[1:]

    def _render_page_pair(
        self, payload: Dict[str, Any], metadata: Dict[str, Any]
    ) -> str:
        """Render the two pages (front/back) for a single day's frames."""
        header_text = self._build_header_text(payload, metadata)

        levels_3_4 = payload.get("levels_3_4") or []
        levels_1_2 = payload.get("levels_1_2") or []
        levels_5_6 = payload.get("levels_5_6") or []

        # Ensure predictable ordering
        levels_3_4 = list(levels_3_4)[:3]
        levels_1_2 = list(levels_1_2)[:3]
        levels_5_6 = list(levels_5_6)[:2]

        # Helper to get safe english/function
        def _frame_text(frame: Dict[str, Any]) -> Tuple[str, str]:
            english = str(frame.get("english", "")).strip()
            # Convert markdown bold (**word**) to HTML <strong> tags
            english = self._convert_markdown_to_html(english)
            func = self._pretty_function(
                str(frame.get("language_function", "")).strip()
            )
            return english, func

        # Page 1: three Levels 3-4 frames, one per panel
        page1_panels: List[str] = []
        for idx in range(3):
            english, func = ("", "")
            if idx < len(levels_3_4):
                english, func = _frame_text(levels_3_4[idx])
            panel_html = f"""
                <div class=\"panel panel-{idx + 1}\">
                    <div class=\"frame-text\">{english}</div>
                    <div class=\"frame-function\">{func}</div>
                </div>
            """
            page1_panels.append(panel_html)

        page1_html = f"""
        <div class=\"frames-page frames-page-1\">
            <div class=\"header\">{header_text}</div>
            <div class=\"panels\">
                <div class=\"fold-line fold-line-1\"></div>
                <div class=\"fold-line fold-line-2\"></div>
                {"".join(page1_panels)}
            </div>
        </div>
        """

        # Page 2: top third = all 1-2 frames, middle/bottom = 5-6 frames
        # Top panel: bundle
        bundle_rows: List[str] = []
        for frame in levels_1_2:
            english, func = _frame_text(frame)
            bundle_rows.append(
                f"""
                <div class=\"bundle-frame-row\">
                    <div class=\"bundle-frame-text\">{english}</div>
                    <div class=\"bundle-frame-function\">{func}</div>
                </div>
                """
            )

        bundle_html = "".join(bundle_rows) if bundle_rows else ""

        # Middle and bottom panels: 5-6 frames
        page2_panel2_text, page2_panel2_func = ("", "")
        page2_panel3_text, page2_panel3_func = ("", "")
        if len(levels_5_6) >= 1:
            page2_panel2_text, page2_panel2_func = _frame_text(levels_5_6[0])
        if len(levels_5_6) >= 2:
            page2_panel3_text, page2_panel3_func = _frame_text(levels_5_6[1])

        page2_html = f"""
        <div class=\"frames-page frames-page-2\">
            <div class=\"header\">{header_text}</div>
            <div class=\"panels\">
                <div class=\"fold-line fold-line-1\"></div>
                <div class=\"fold-line fold-line-2\"></div>
                <div class=\"panel panel-1\">
                    <div class=\"panel-title\">Levels 1-2</div>
                    {bundle_html}
                </div>
                <div class=\"panel panel-2\">
                    <div class=\"panel-title\">Levels 5-6</div>
                    <div class=\"frame-text\">{page2_panel2_text}</div>
                    <div class=\"frame-function\">{page2_panel2_func}</div>
                </div>
                <div class=\"panel panel-3\">
                    <div class=\"panel-title\">Levels 5-6</div>
                    <div class=\"frame-text\">{page2_panel3_text}</div>
                    <div class=\"frame-function\">{page2_panel3_func}</div>
                </div>
            </div>
        </div>
        """

        return page1_html + "\n" + page2_html

    def generate_html(
        self,
        lesson_json: Dict[str, Any],
        output_path: Optional[str] = None,
        user_name: Optional[str] = None,
    ) -> str:
        """Generate HTML file with sentence frames pages.

        For now, we generate pages for every day that has sentence_frames.
        Each day contributes a two-page pair (front/back) to the output.
        """
        payloads = self.extract_sentence_frames(lesson_json)
        if not payloads:
            raise ValueError("No sentence_frames found in lesson plan")

        metadata = lesson_json.get("metadata", {})
        pages: List[str] = []
        for payload in payloads:
            pages.append(self._render_page_pair(payload, metadata))

        final_html = self.html_template.format(
            css=self.css_template,
            pages="\n".join(pages),
            week_of=payloads[0]["week_of"],
        )

        output_file = self._resolve_html_path(lesson_json, user_name, output_path)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(final_html)

        logger.info(
            "sentence_frames_html_generated",
            extra={
                "output_path": str(output_file),
                "page_pairs": len(payloads),
                "week_of": payloads[0]["week_of"],
            },
        )

        return str(output_file)

    def convert_to_pdf(self, html_path: str, pdf_path: str) -> str:
        """Convert HTML file to PDF using WeasyPrint with Playwright fallback.

        This mirrors ObjectivesPDFGenerator.convert_to_pdf but uses the
        sentence-frames CSS template.
        """

        html_file = Path(html_path)
        pdf_file = Path(pdf_path)
        pdf_file.parent.mkdir(parents=True, exist_ok=True)
        if pdf_file.exists():
            try:
                pdf_file.unlink()
            except OSError:
                pass

        try:
            import weasyprint

            weasyprint.HTML(filename=str(html_file)).write_pdf(
                str(pdf_file), stylesheets=[weasyprint.CSS(string=self.css_template)]
            )
            logger.info(
                "sentence_frames_weasyprint_pdf_generated",
                extra={"html_path": str(html_file), "pdf_path": str(pdf_file)},
            )
            return str(pdf_file)
        except Exception as exc:
            logger.warning(
                "sentence_frames_weasyprint_failed",
                extra={
                    "html_path": str(html_file),
                    "pdf_path": str(pdf_file),
                    "error": str(exc),
                },
            )

        # Fallback to Playwright
        try:
            from playwright.sync_api import sync_playwright
        except ImportError as exc:
            logger.error(
                "sentence_frames_playwright_not_installed", extra={"error": str(exc)}
            )
            raise

        with sync_playwright() as p:  # type: ignore[name-defined]
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(f"file:///{html_file}", wait_until="networkidle")
            page.pdf(
                path=str(pdf_file),
                format="Letter",
                landscape=False,
                print_background=True,
                prefer_css_page_size=True,  # Use CSS @page size instead of default
                margin={
                    "top": "0",
                    "right": "0",
                    "bottom": "0",
                    "left": "0",
                },  # No margins - use full page
            )
            browser.close()

        logger.info(
            "sentence_frames_playwright_pdf_generated",
            extra={"html_path": str(html_file), "pdf_path": str(pdf_file)},
        )
        return str(pdf_file)

    def _add_gray_line(self, paragraph):
        """Add a thin gray horizontal line to a paragraph."""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.75pt line
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "808080")  # 50% gray
        pBdr.append(bottom)

    def generate_docx(
        self,
        lesson_json: Dict[str, Any],
        output_path: Optional[str] = None,
        user_name: Optional[str] = None,
    ) -> str:
        """Generate DOCX file with sentence frames, two pages per day (front/back).

        Each day gets two pages:
        - Page 1: Three Levels 3-4 frames, one per panel (top, middle, bottom)
        - Page 2: Top panel = Levels 1-2 frames (bundled), Middle/Bottom = Levels 5-6 frames

        Layout:
        - Portrait US Letter (8.5" x 11")
        - Header: Date | Day | Subject | Grade | Homeroom (10pt, gray)
        - Three equal-height panels per page (3.67" each = 11" / 3 for accurate folding)
        - Frame text: 28pt (Page 1) or 18pt (Page 2 bundle), bold
        - Function label: 11pt (Page 1) or 9pt (Page 2), uppercase, gray

        Args:
            lesson_json: Lesson plan JSON structure
            output_path: Optional path to save DOCX file
            user_name: Optional teacher name for file naming

        Returns:
            Path to generated DOCX file
        """
        payloads = self.extract_sentence_frames(lesson_json)
        if not payloads:
            raise ValueError("No sentence frames found in lesson plan")

        metadata = lesson_json.get("metadata", {})

        # Resolve output path
        if output_path:
            docx_file = Path(output_path)
        else:
            directory = self._resolve_output_directory(lesson_json)
            base_name = self._build_default_basename(lesson_json, user_name)
            docx_file = directory / f"{base_name}.docx"

        docx_file.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = Document()

        # Set page orientation to portrait
        section = doc.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)  # Portrait width
        section.page_height = Inches(11)  # Portrait height

        # No margins - use full page (matching PDF layout)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        # Page dimensions: 8.5" x 11" (full page)
        # Three equal panels: divide full page height by 3 for accurate folding
        # Each panel: 11" / 3 = 3.67" per panel
        # Note: Panel height is used implicitly in spacing calculations below

        # Generate two pages per payload (day)
        for payload_idx, payload in enumerate(payloads):
            header_text = self._build_header_text(payload, metadata)

            levels_3_4 = payload.get("levels_3_4") or []
            levels_1_2 = payload.get("levels_1_2") or []
            levels_5_6 = payload.get("levels_5_6") or []

            # Ensure predictable ordering
            levels_3_4 = list(levels_3_4)[:3]
            levels_1_2 = list(levels_1_2)[:3]
            levels_5_6 = list(levels_5_6)[:2]

            # Helper to get frame text and extract bold words
            def _get_frame_data(frame: Dict[str, Any]) -> Tuple[str, List[str], str]:
                english = str(frame.get("english", "")).strip()
                cleaned_text, bold_words = self._extract_bold_words_from_markdown(
                    english
                )
                func = self._pretty_function(
                    str(frame.get("language_function", "")).strip()
                )
                return cleaned_text, bold_words, func

            # PAGE 1: Three Levels 3-4 frames, one per panel
            if payload_idx > 0 or payload_idx == 0:
                # Add new section for each page (except first)
                if payload_idx > 0 or (payload_idx == 0 and len(payloads) > 1):
                    doc.add_section(WD_SECTION.NEW_PAGE)
                    section = doc.sections[-1]
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Inches(8.5)
                    section.page_height = Inches(11)
                    section.top_margin = Inches(0)
                    section.bottom_margin = Inches(0)
                    section.left_margin = Inches(0)
                    section.right_margin = Inches(0)

            # Header - positioned 1 cm from top and left (0.39 inches)
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_before = Inches(0.39)  # 1 cm top margin
            header_para.paragraph_format.left_indent = Inches(0.39)  # 1 cm left margin
            header_run = header_para.add_run(header_text)
            header_run.font.size = Pt(10)
            header_run.font.name = "Source Sans Pro"
            header_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            header_para.paragraph_format.space_after = Pt(0)
            header_para.paragraph_format.alignment = 0  # Left align

            # Three panels for Page 1
            for panel_idx in range(3):
                panel_para = doc.add_paragraph()
                panel_para.paragraph_format.alignment = 1  # Center align
                panel_para.paragraph_format.space_before = Pt(0)
                panel_para.paragraph_format.space_after = Pt(0)

                # Set panel height using spacing
                # Each panel should be 3.55" = 255.6pt
                if panel_idx == 0:
                    # First panel: minimal spacing
                    panel_para.paragraph_format.space_before = Pt(0)
                else:
                    # Subsequent panels: space to position them
                    panel_para.paragraph_format.space_before = Pt(0)

                english, bold_words, func = ("", [], "")
                if panel_idx < len(levels_3_4):
                    english, bold_words, func = _get_frame_data(levels_3_4[panel_idx])

                # Frame text (28pt, bold) - apply bold to vocabulary words
                self._add_bold_text_to_paragraph(
                    panel_para, english, bold_words, 28, "Source Sans Pro"
                )

                # Function label (11pt, uppercase, gray)
                if func:
                    func_para = doc.add_paragraph()
                    func_para.paragraph_format.alignment = 1  # Center align
                    func_run = func_para.add_run(func.upper())
                    func_run.font.size = Pt(11)
                    func_run.font.name = "Source Sans Pro"
                    func_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
                    func_para.paragraph_format.space_before = Pt(2)
                    func_para.paragraph_format.space_after = Pt(0)

                # Add spacing to position next panel (approximate 3.55" panel height)
                # Use line spacing and paragraph spacing to approximate panel height
                panel_para.paragraph_format.line_spacing = 1.2
                if panel_idx < 2:  # Not the last panel
                    # Add spacing to approximate panel separation
                    panel_para.paragraph_format.space_after = Pt(20)

            # PAGE 2: Top panel = Levels 1-2 frames (bundled), Middle/Bottom = Levels 5-6 frames
            doc.add_section(WD_SECTION.NEW_PAGE)
            section = doc.sections[-1]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)

            # Header - positioned 1 cm from top and left (0.39 inches)
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_before = Inches(0.39)  # 1 cm top margin
            header_para.paragraph_format.left_indent = Inches(0.39)  # 1 cm left margin
            header_run = header_para.add_run(header_text)
            header_run.font.size = Pt(10)
            header_run.font.name = "Source Sans Pro"
            header_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            header_para.paragraph_format.space_after = Pt(0)

            # Top panel: Levels 1-2 (bundled)
            panel1_para = doc.add_paragraph()
            panel1_para.paragraph_format.alignment = 1  # Center align

            # Panel title
            title_run = panel1_para.add_run("Levels 1-2")
            title_run.font.size = Pt(10)
            title_run.font.bold = True
            title_run.font.name = "Source Sans Pro"
            title_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            panel1_para.paragraph_format.space_after = Pt(5)

            # Bundle frames (18pt, bold)
            for frame in levels_1_2:
                english, bold_words, func = _get_frame_data(frame)

                frame_para = doc.add_paragraph()
                frame_para.paragraph_format.alignment = 1  # Center align
                self._add_bold_text_to_paragraph(
                    frame_para, english, bold_words, 18, "Source Sans Pro"
                )
                frame_para.paragraph_format.space_after = Pt(2)

                # Function label (9pt, uppercase, gray)
                if func:
                    func_para = doc.add_paragraph()
                    func_para.paragraph_format.alignment = 1  # Center align
                    func_run = func_para.add_run(func.upper())
                    func_run.font.size = Pt(9)
                    func_run.font.name = "Source Sans Pro"
                    func_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
                    func_para.paragraph_format.space_after = Pt(5)

            # Middle panel: Levels 5-6 (first frame)
            panel2_para = doc.add_paragraph()
            panel2_para.paragraph_format.alignment = 1  # Center align

            # Panel title
            title_run = panel2_para.add_run("Levels 5-6")
            title_run.font.size = Pt(10)
            title_run.font.bold = True
            title_run.font.name = "Source Sans Pro"
            title_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            panel2_para.paragraph_format.space_after = Pt(5)

            if len(levels_5_6) >= 1:
                english, bold_words, func = _get_frame_data(levels_5_6[0])

                frame_para = doc.add_paragraph()
                frame_para.paragraph_format.alignment = 1  # Center align
                self._add_bold_text_to_paragraph(
                    frame_para, english, bold_words, 28, "Source Sans Pro"
                )
                frame_para.paragraph_format.space_after = Pt(2)

                if func:
                    func_para = doc.add_paragraph()
                    func_para.paragraph_format.alignment = 1  # Center align
                    func_run = func_para.add_run(func.upper())
                    func_run.font.size = Pt(11)
                    func_run.font.name = "Source Sans Pro"
                    func_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray

            # Bottom panel: Levels 5-6 (second frame)
            panel3_para = doc.add_paragraph()
            panel3_para.paragraph_format.alignment = 1  # Center align

            # Panel title
            title_run = panel3_para.add_run("Levels 5-6")
            title_run.font.size = Pt(10)
            title_run.font.bold = True
            title_run.font.name = "Source Sans Pro"
            title_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray
            panel3_para.paragraph_format.space_after = Pt(5)

            if len(levels_5_6) >= 2:
                english, bold_words, func = _get_frame_data(levels_5_6[1])

                frame_para = doc.add_paragraph()
                frame_para.paragraph_format.alignment = 1  # Center align
                self._add_bold_text_to_paragraph(
                    frame_para, english, bold_words, 28, "Source Sans Pro"
                )
                frame_para.paragraph_format.space_after = Pt(2)

                if func:
                    func_para = doc.add_paragraph()
                    func_para.paragraph_format.alignment = 1  # Center align
                    func_run = func_para.add_run(func.upper())
                    func_run.font.size = Pt(11)
                    func_run.font.name = "Source Sans Pro"
                    func_run.font.color.rgb = RGBColor(128, 128, 128)  # 50% gray

        # Save document
        doc.save(str(docx_file))

        logger.info(
            "sentence_frames_docx_generated",
            extra={
                "output_path": str(docx_file),
                "page_pairs": len(payloads),
                "week_of": payloads[0]["week_of"] if payloads else "Unknown",
            },
        )

        return str(docx_file)

    def generate_pdf(
        self,
        lesson_json: Dict[str, Any],
        output_path: Optional[str] = None,
        user_name: Optional[str] = None,
        keep_html: bool = True,
    ) -> str:
        """Generate PDF file with sentence frames (HTML + PDF conversion)."""

        html_path, pdf_path = self._resolve_pdf_and_html_paths(
            lesson_json, user_name, output_path
        )
        generated_html = self.generate_html(
            lesson_json, str(html_path), user_name=user_name
        )
        final_pdf = self.convert_to_pdf(generated_html, str(pdf_path))
        if not keep_html:
            try:
                Path(generated_html).unlink()
            except OSError:
                pass
        return final_pdf


def generate_sentence_frames_html(
    lesson_json: Dict[str, Any],
    output_path: str,
    user_name: Optional[str] = None,
) -> str:
    """Convenience wrapper to generate sentence frames HTML."""
    generator = SentenceFramesPDFGenerator()
    return generator.generate_html(
        lesson_json, output_path=output_path, user_name=user_name
    )


def generate_sentence_frames_pdf(
    lesson_json: Dict[str, Any],
    output_path: str,
    user_name: Optional[str] = None,
    keep_html: bool = True,
) -> str:
    """Convenience wrapper to generate sentence frames PDF from lesson JSON."""
    generator = SentenceFramesPDFGenerator()
    return generator.generate_pdf(lesson_json, output_path, user_name, keep_html)


def generate_sentence_frames_docx(
    lesson_json: Dict[str, Any],
    output_path: str,
    user_name: Optional[str] = None,
) -> str:
    """Convenience wrapper to generate sentence frames DOCX from lesson JSON."""
    generator = SentenceFramesPDFGenerator()
    return generator.generate_docx(lesson_json, output_path, user_name)
