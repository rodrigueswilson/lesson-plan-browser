"""DOCX generation for sentence frames."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.telemetry import logger

from .html_builder import (
    build_header_text,
    extract_bold_words_from_markdown,
    pretty_function,
)


def add_gray_line(paragraph) -> None:
    """Add a thin gray horizontal line to a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.append(pBdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)


def add_bold_text_to_paragraph(
    paragraph,
    text: str,
    bold_words: List[str],
    font_size: int,
    font_name: str = "Source Sans Pro",
) -> None:
    """Add text to a paragraph with specified words in bold."""
    if not bold_words:
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.name = font_name
        return

    words_pattern = re.compile(
        r"\b(" + "|".join(re.escape(word) for word in bold_words) + r")\b",
        re.IGNORECASE,
    )

    last_index = 0
    for match in words_pattern.finditer(text):
        if match.start() > last_index:
            run = paragraph.add_run(text[last_index : match.start()])
            run.font.size = Pt(font_size)
            run.font.bold = True
            run.font.name = font_name

        bold_run = paragraph.add_run(match.group(0))
        bold_run.font.size = Pt(font_size)
        bold_run.font.bold = True
        bold_run.font.name = font_name

        last_index = match.end()

    if last_index < len(text):
        run = paragraph.add_run(text[last_index:])
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.name = font_name


def generate_docx(
    payloads: List[Dict[str, Any]],
    metadata: Dict[str, Any],
    docx_path: Path,
) -> str:
    """Generate DOCX file with sentence frames, two pages per day (front/back)."""
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    for payload_idx, payload in enumerate(payloads):
        header_text = build_header_text(payload, metadata)

        levels_3_4 = payload.get("levels_3_4") or []
        levels_1_2 = payload.get("levels_1_2") or []
        levels_5_6 = payload.get("levels_5_6") or []

        levels_3_4 = list(levels_3_4)[:3]
        levels_1_2 = list(levels_1_2)[:3]
        levels_5_6 = list(levels_5_6)[:2]

        def _get_frame_data(frame: Dict[str, Any]) -> Tuple[str, List[str], str]:
            english = str(frame.get("english", "")).strip()
            cleaned_text, bold_words = extract_bold_words_from_markdown(english)
            func = pretty_function(
                str(frame.get("language_function", "")).strip()
            )
            return cleaned_text, bold_words, func

        if payload_idx > 0 or (payload_idx == 0 and len(payloads) > 1):
            doc.add_section(WD_SECTION.NEW_PAGE)
            section = doc.sections[-1]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)

        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Inches(0.39)
        header_para.paragraph_format.left_indent = Inches(0.39)
        header_run = header_para.add_run(header_text)
        header_run.font.size = Pt(10)
        header_run.font.name = "Source Sans Pro"
        header_run.font.color.rgb = RGBColor(128, 128, 128)
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.alignment = 0

        for panel_idx in range(3):
            panel_para = doc.add_paragraph()
            panel_para.paragraph_format.alignment = 1
            panel_para.paragraph_format.space_before = Pt(0)
            panel_para.paragraph_format.space_after = Pt(0)

            english, bold_words, func = ("", [], "")
            if panel_idx < len(levels_3_4):
                english, bold_words, func = _get_frame_data(levels_3_4[panel_idx])

            add_bold_text_to_paragraph(
                panel_para, english, bold_words, 28, "Source Sans Pro"
            )

            if func:
                func_para = doc.add_paragraph()
                func_para.paragraph_format.alignment = 1
                func_run = func_para.add_run(func.upper())
                func_run.font.size = Pt(11)
                func_run.font.name = "Source Sans Pro"
                func_run.font.color.rgb = RGBColor(128, 128, 128)
                func_para.paragraph_format.space_before = Pt(2)
                func_para.paragraph_format.space_after = Pt(0)

            panel_para.paragraph_format.line_spacing = 1.2
            if panel_idx < 2:
                panel_para.paragraph_format.space_after = Pt(20)

        doc.add_section(WD_SECTION.NEW_PAGE)
        section = doc.sections[-1]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Inches(0.39)
        header_para.paragraph_format.left_indent = Inches(0.39)
        header_run = header_para.add_run(header_text)
        header_run.font.size = Pt(10)
        header_run.font.name = "Source Sans Pro"
        header_run.font.color.rgb = RGBColor(128, 128, 128)
        header_para.paragraph_format.space_after = Pt(0)

        panel1_para = doc.add_paragraph()
        panel1_para.paragraph_format.alignment = 1

        title_run = panel1_para.add_run("Levels 1-2")
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_run.font.name = "Source Sans Pro"
        title_run.font.color.rgb = RGBColor(128, 128, 128)
        panel1_para.paragraph_format.space_after = Pt(5)

        for frame in levels_1_2:
            english, bold_words, func = _get_frame_data(frame)

            frame_para = doc.add_paragraph()
            frame_para.paragraph_format.alignment = 1
            add_bold_text_to_paragraph(
                frame_para, english, bold_words, 18, "Source Sans Pro"
            )
            frame_para.paragraph_format.space_after = Pt(2)

            if func:
                func_para = doc.add_paragraph()
                func_para.paragraph_format.alignment = 1
                func_run = func_para.add_run(func.upper())
                func_run.font.size = Pt(9)
                func_run.font.name = "Source Sans Pro"
                func_run.font.color.rgb = RGBColor(128, 128, 128)
                func_para.paragraph_format.space_after = Pt(5)

        panel2_para = doc.add_paragraph()
        panel2_para.paragraph_format.alignment = 1

        title_run = panel2_para.add_run("Levels 5-6")
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_run.font.name = "Source Sans Pro"
        title_run.font.color.rgb = RGBColor(128, 128, 128)
        panel2_para.paragraph_format.space_after = Pt(5)

        if len(levels_5_6) >= 1:
            english, bold_words, func = _get_frame_data(levels_5_6[0])

            frame_para = doc.add_paragraph()
            frame_para.paragraph_format.alignment = 1
            add_bold_text_to_paragraph(
                frame_para, english, bold_words, 28, "Source Sans Pro"
            )
            frame_para.paragraph_format.space_after = Pt(2)

            if func:
                func_para = doc.add_paragraph()
                func_para.paragraph_format.alignment = 1
                func_run = func_para.add_run(func.upper())
                func_run.font.size = Pt(11)
                func_run.font.name = "Source Sans Pro"
                func_run.font.color.rgb = RGBColor(128, 128, 128)

        panel3_para = doc.add_paragraph()
        panel3_para.paragraph_format.alignment = 1

        title_run = panel3_para.add_run("Levels 5-6")
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_run.font.name = "Source Sans Pro"
        title_run.font.color.rgb = RGBColor(128, 128, 128)
        panel3_para.paragraph_format.space_after = Pt(5)

        if len(levels_5_6) >= 2:
            english, bold_words, func = _get_frame_data(levels_5_6[1])

            frame_para = doc.add_paragraph()
            frame_para.paragraph_format.alignment = 1
            add_bold_text_to_paragraph(
                frame_para, english, bold_words, 28, "Source Sans Pro"
            )
            frame_para.paragraph_format.space_after = Pt(2)

            if func:
                func_para = doc.add_paragraph()
                func_para.paragraph_format.alignment = 1
                func_run = func_para.add_run(func.upper())
                func_run.font.size = Pt(11)
                func_run.font.name = "Source Sans Pro"
                func_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(str(docx_path))

    logger.info(
        "sentence_frames_docx_generated",
        extra={
            "output_path": str(docx_path),
            "page_pairs": len(payloads),
            "week_of": payloads[0]["week_of"] if payloads else "Unknown",
        },
    )

    return str(docx_path)
