"""DOCX generation for objectives (one lesson per page)."""

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.telemetry import logger

from . import extraction
from . import font_calculation


def add_gray_line(paragraph) -> None:
    """Add a thin gray horizontal line to a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.append(pBdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)


def generate_docx(
    lesson_json: Dict[str, Any],
    output_path: str,
    user_name: Optional[str] = None,
    week_of: Optional[str] = None,
) -> str:
    """
    Generate DOCX file with objectives, one lesson per page.

    Each page contains:
    - Header: Date | Subject | Grade | Homeroom (10pt)
    - Student Goal: 3/4 of page (base 48pt, bold)
    - Separator: Thin gray line
    - WIDA Objective: 1/4 of page (base 14pt, 50% gray)
    """
    objectives = extraction.extract_objectives(lesson_json)

    if not objectives:
        raise ValueError("No objectives found in lesson plan")

    metadata = lesson_json.get("metadata", {})
    week_of = week_of or metadata.get("week_of", "Unknown")
    grade = metadata.get("grade", "Unknown")
    subject = metadata.get("subject", "Unknown")
    homeroom = metadata.get("homeroom", "Unknown")

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    usable_width = 10.0
    usable_height = 7.5

    header_height_estimate = 0.3
    objectives_area_height = usable_height - header_height_estimate

    separator_height = 0.15
    available_for_objectives = objectives_area_height - separator_height

    student_goal_height = (available_for_objectives * 0.75) * 0.95
    wida_height = (available_for_objectives * 0.25) * 0.95

    student_goal_content_width = usable_width - 0.4
    student_goal_content_height = student_goal_height - 0.1

    wida_content_width = usable_width - 0.4
    wida_content_height = wida_height - 0.1

    for i, obj in enumerate(objectives):
        if i > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
            section = doc.sections[-1]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        section = doc.sections[-1]

        day_date = font_calculation.get_day_date(week_of, obj["day"])
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(
            f"{day_date} | {obj['subject']} | Grade {grade} | {homeroom}"
        )
        header_run.font.size = Pt(10)
        header_run.font.name = "Calibri"
        header_para.paragraph_format.space_after = Pt(2)
        header_para.paragraph_format.keep_with_next = True

        wida_text = obj.get("wida_objective", "").strip()
        if not wida_text:
            wida_text = "No WIDA Objective specified"

        wida_font_size = 12

        wida_words = wida_text.split()
        wida_chars_per_line = (
            (wida_content_width * 72) / (wida_font_size * 0.55)
            if wida_font_size > 0
            else 60
        )
        wida_words_per_line = max(1, wida_chars_per_line / 6)
        wida_estimated_lines = max(1, len(wida_words) / wida_words_per_line)
        actual_wida_text_height = (
            wida_font_size * 1.2 * wida_estimated_lines
        ) / 72

        label_font_size = 12
        label_height_inches = (label_font_size * 1.2) / 72

        wida_label_spacing = 0.03
        total_wida_content_height = (
            label_height_inches + wida_label_spacing + actual_wida_text_height
        )

        wida_section_padding = 0.1
        total_wida_section_height = total_wida_content_height + wida_section_padding

        student_goal_text = obj.get("student_goal", "").strip()
        if not student_goal_text:
            student_goal_text = "No Student Goal specified"

        separator_height_local = 0.15
        available_for_student_goal = (
            objectives_area_height
            - separator_height_local
            - total_wida_section_height
        )

        student_goal_target_height = available_for_student_goal * 0.75

        student_goal_font_size = font_calculation.calculate_font_size_to_fill_height(
            student_goal_text,
            student_goal_content_width,
            student_goal_target_height,
            min_font_size=12,
            max_font_size=60,
            char_width_ratio=0.6,
            line_height_ratio=1.3,
        )

        student_goal_para = doc.add_paragraph()
        student_goal_run = student_goal_para.add_run(student_goal_text)
        student_goal_run.font.size = Pt(student_goal_font_size)
        student_goal_run.font.bold = True
        student_goal_run.font.name = "Verdana"
        student_goal_para.paragraph_format.widow_control = True
        student_goal_para.paragraph_format.keep_together = True
        student_goal_para.paragraph_format.keep_with_next = True
        student_goal_para.paragraph_format.page_break_before = False
        student_goal_para.paragraph_format.line_spacing = 1.25
        student_goal_para.paragraph_format.space_before = Pt(0)

        words = student_goal_text.split()
        chars_per_line = (
            (student_goal_content_width * 72) / (student_goal_font_size * 0.6)
            if student_goal_font_size > 0
            else 60
        )
        words_per_line = max(1, chars_per_line / 6)
        estimated_lines = max(1, len(words) / words_per_line)
        actual_student_goal_height = (
            student_goal_font_size * 1.3 * estimated_lines
        ) / 72

        remaining_student_goal_space = (
            available_for_student_goal - actual_student_goal_height
        )
        if remaining_student_goal_space > 0.1:
            student_goal_para.paragraph_format.space_after = Pt(
                int(remaining_student_goal_space * 72 * 0.30)
            )
        else:
            student_goal_para.paragraph_format.space_after = Pt(1)

        separator_para = doc.add_paragraph()
        add_gray_line(separator_para)
        separator_para.paragraph_format.space_before = Pt(0)
        separator_para.paragraph_format.space_after = Pt(0)
        separator_para.paragraph_format.page_break_before = False
        separator_para.paragraph_format.keep_together = True
        separator_para.paragraph_format.keep_with_next = True

        remaining_wida_space = (
            objectives_area_height
            - actual_student_goal_height
            - separator_height_local
            - total_wida_section_height
        )

        wida_label_para = doc.add_paragraph()
        wida_label_run = wida_label_para.add_run("WIDA/Bilingual:")
        wida_label_run.font.size = Pt(wida_font_size)
        wida_label_run.font.bold = True
        wida_label_run.font.color.rgb = RGBColor(128, 128, 128)
        wida_label_run.font.name = "Calibri"

        if remaining_wida_space > 0.05:
            wida_label_para.paragraph_format.space_before = Pt(
                int(remaining_wida_space * 72 * 0.30)
            )
        else:
            wida_label_para.paragraph_format.space_before = Pt(0)
        wida_label_para.paragraph_format.space_after = Pt(1)
        wida_label_para.paragraph_format.keep_together = True
        wida_label_para.paragraph_format.keep_with_next = True

        wida_para = doc.add_paragraph()
        wida_run = wida_para.add_run(wida_text)
        wida_run.font.size = Pt(wida_font_size)
        wida_run.font.color.rgb = RGBColor(128, 128, 128)
        wida_run.font.name = "Calibri"
        wida_para.paragraph_format.line_spacing = 1.0
        wida_para.paragraph_format.widow_control = True
        wida_para.paragraph_format.keep_together = True
        wida_para.paragraph_format.keep_with_next = False
        wida_para.paragraph_format.page_break_before = False
        wida_para.paragraph_format.space_after = Pt(0)

    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))

    logger.info(
        "objectives_docx_generated",
        extra={
            "output_path": str(output_file),
            "objective_count": len(objectives),
            "week_of": week_of,
        },
    )

    return str(output_file)
