"""
Test subject-based slot detection to handle misaligned slot numbers.

Validates that find_slot_by_subject() correctly identifies slots by subject
and handles edge cases like duplicate subjects, combined subjects, and aliases.
"""

import pytest
from pathlib import Path
from docx import Document
from tools.docx_parser import DOCXParser
import io
from PIL import Image as PILImage


def create_doc_with_subjects(tmp_path, subjects_and_teachers):
    """
    Create a test document with specified subjects and teachers.
    
    Args:
        subjects_and_teachers: List of (subject, teacher) tuples
    
    Returns:
        Path to created document
    """
    doc_path = tmp_path / "test_subjects.docx"
    doc = Document()
    
    for subject, teacher in subjects_and_teachers:
        # Metadata table
        meta = doc.add_table(rows=3, cols=2)
        meta.rows[0].cells[0].text = f"Name: {teacher}"
        meta.rows[1].cells[0].text = f"Subject: {subject}"
        meta.rows[1].cells[1].text = "Grade: 2"
        
        # Daily plans table
        daily = doc.add_table(rows=2, cols=3)
        daily.rows[0].cells[1].text = "Monday"
        daily.rows[1].cells[0].text = "Objective"
        daily.rows[1].cells[1].text = f"{subject} content for {teacher}"
    
    # Signature table
    sig = doc.add_table(rows=2, cols=2)
    sig.rows[0].cells[0].text = "Signature"
    
    doc.save(str(doc_path))
    return doc_path


class TestFindSlotBySubject:
    """Test subject-based slot detection."""
    
    def test_find_slot_exact_match(self, tmp_path):
        """Test finding slot by exact subject match."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Math", "Teacher A"),
            ("ELA", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Find Math slot
        slot_num = parser.find_slot_by_subject("Math")
        assert slot_num == 1
        
        # Find ELA slot
        slot_num = parser.find_slot_by_subject("ELA")
        assert slot_num == 2
    
    def test_find_slot_alias_match(self, tmp_path):
        """Test finding slot using subject alias."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Mathematics", "Teacher A"),
            ("English", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Request "Math" should match "Mathematics"
        slot_num = parser.find_slot_by_subject("Math")
        assert slot_num == 1
        
        # Request "ELA" should match "English"
        slot_num = parser.find_slot_by_subject("ELA")
        assert slot_num == 2
    
    def test_find_slot_combined_subject(self, tmp_path):
        """Test finding combined subject like ELA/SS."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("ELA/SS", "Teacher A"),
            ("Math", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Request "ELA/SS" should match
        slot_num = parser.find_slot_by_subject("ELA/SS")
        assert slot_num == 1
    
    def test_find_slot_no_cross_match(self, tmp_path):
        """Test that SS doesn't match ELA/SS."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("ELA/SS", "Teacher A"),
            ("Social Studies", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Request "Social Studies" should match slot 2, NOT slot 1 (ELA/SS)
        slot_num = parser.find_slot_by_subject("Social Studies")
        assert slot_num == 2
        
        # Request "SS" should also match slot 2
        slot_num = parser.find_slot_by_subject("SS")
        assert slot_num == 2
    
    def test_find_slot_misaligned_numbers(self, tmp_path):
        """Test misaligned slot numbers (Savoca scenario)."""
        # Create doc: Slot 1 = ELA/SS, Slot 2 = Math
        doc_path = create_doc_with_subjects(tmp_path, [
            ("ELA/SS", "Donna Savoca"),
            ("Math", "Donna Savoca")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Weekly plan config says Slot 2 wants ELA/SS
        # But in source file, ELA/SS is in Slot 1
        slot_num = parser.find_slot_by_subject("ELA/SS")
        assert slot_num == 1  # Finds it in slot 1, not 2
    
    def test_find_slot_multiple_math_with_teacher(self, tmp_path):
        """Test disambiguating duplicate subjects by teacher."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Math", "Teacher A"),
            ("Math", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Request Math with Teacher B
        slot_num = parser.find_slot_by_subject("Math", teacher_name="Teacher B")
        assert slot_num == 2
        
        # Request Math with Teacher A
        slot_num = parser.find_slot_by_subject("Math", teacher_name="Teacher A")
        assert slot_num == 1
    
    def test_find_slot_multiple_math_no_teacher(self, tmp_path):
        """Test multiple matches without teacher returns first."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Math", "Teacher A"),
            ("Math", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Request Math without teacher - should return first match
        slot_num = parser.find_slot_by_subject("Math")
        assert slot_num == 1  # First match
    
    def test_find_slot_not_found(self, tmp_path):
        """Test subject not found raises ValueError."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Math", "Teacher A")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Request Science (not in document)
        with pytest.raises(ValueError, match="Subject 'Science' not found"):
            parser.find_slot_by_subject("Science")
    
    def test_find_slot_case_insensitive(self, tmp_path):
        """Test case-insensitive matching."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("MATH", "Teacher A"),
            ("ela", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Lowercase request should match uppercase metadata
        slot_num = parser.find_slot_by_subject("math")
        assert slot_num == 1
        
        # Uppercase request should match lowercase metadata
        slot_num = parser.find_slot_by_subject("ELA")
        assert slot_num == 2
    
    def test_find_slot_daniela_bug_scenario(self, tmp_path):
        """Test Daniela's bug: same teacher (Grande), same subject (Science), multiple slots."""
        # Create doc simulating Daniela's scenario:
        # Slot 1: Science Grade 6 (Mariela Grande)
        # Slot 2: Science Grade 2 (Mariela Grande)
        doc_path = tmp_path / "test_daniela_scenario.docx"
        doc = Document()
        
        # Slot 1: Science Grade 6 (Grande)
        meta1 = doc.add_table(rows=3, cols=2)
        meta1.rows[0].cells[0].text = "Name: Mariela Grande"
        meta1.rows[1].cells[0].text = "Subject: Science"
        meta1.rows[1].cells[1].text = "Grade: 6"
        
        daily1 = doc.add_table(rows=2, cols=3)
        daily1.rows[0].cells[1].text = "Monday"
        daily1.rows[1].cells[0].text = "Objective"
        daily1.rows[1].cells[1].text = "GRADE6_SCIENCE_CONTENT"
        
        # Slot 2: Science Grade 2 (Grande)
        meta2 = doc.add_table(rows=3, cols=2)
        meta2.rows[0].cells[0].text = "Name: Mariela Grande"
        meta2.rows[1].cells[0].text = "Subject: Science"
        meta2.rows[1].cells[1].text = "Grade: 2"
        
        daily2 = doc.add_table(rows=2, cols=3)
        daily2.rows[0].cells[1].text = "Monday"
        daily2.rows[1].cells[0].text = "Objective"
        daily2.rows[1].cells[1].text = "GRADE2_SCIENCE_CONTENT"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        parser = DOCXParser(str(doc_path))
        
        # Without teacher_name, should return first match (slot 1)
        slot_num = parser.find_slot_by_subject("Science")
        assert slot_num == 1
        
        # With teacher_name, should still return first match if both match
        # (This tests the edge case where teacher name doesn't disambiguate)
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Mariela Grande")
        # Should return first match (slot 1) because both slots have same teacher
        assert slot_num == 1
    
    def test_find_slot_same_subject_different_teachers(self, tmp_path):
        """Test disambiguation when multiple teachers teach same subject."""
        # Create doc: Science taught by Morais and Grande
        doc_path = tmp_path / "test_science_morais_grande.docx"
        doc = Document()
        
        # Slot 1: Science (Morais)
        meta1 = doc.add_table(rows=3, cols=2)
        meta1.rows[0].cells[0].text = "Name: Catarina Morais"
        meta1.rows[1].cells[0].text = "Subject: Science"
        
        daily1 = doc.add_table(rows=2, cols=3)
        daily1.rows[0].cells[1].text = "Monday"
        daily1.rows[1].cells[0].text = "Objective"
        daily1.rows[1].cells[1].text = "MORAIS_SCIENCE_CONTENT"
        
        # Slot 2: Science (Grande)
        meta2 = doc.add_table(rows=3, cols=2)
        meta2.rows[0].cells[0].text = "Name: Mariela Grande"
        meta2.rows[1].cells[0].text = "Subject: Science"
        
        daily2 = doc.add_table(rows=2, cols=3)
        daily2.rows[0].cells[1].text = "Monday"
        daily2.rows[1].cells[0].text = "Objective"
        daily2.rows[1].cells[1].text = "GRANDE_SCIENCE_CONTENT"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        parser = DOCXParser(str(doc_path))
        
        # Request Science with Grande - should return Slot 2
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Mariela Grande")
        assert slot_num == 2
        
        # Request Science with Morais - should return Slot 1
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Catarina Morais")
        assert slot_num == 1
        
        # Request Science with just "Grande" - should match Slot 2
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Grande")
        assert slot_num == 2
        
        # Request Science with just "Morais" - should match Slot 1
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Morais")
        assert slot_num == 1
    
    def test_find_slot_teacher_name_partial_match(self, tmp_path):
        """Test teacher name matching with partial names."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Science", "Mariela Grande"),
            ("Science", "Catarina Morais")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Partial match should work
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Grande")
        assert slot_num == 1
        
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Morais")
        assert slot_num == 2
        
        # Full name should work
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Mariela Grande")
        assert slot_num == 1
    
    def test_find_slot_teacher_name_case_insensitive(self, tmp_path):
        """Test teacher name matching is case-insensitive."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Science", "Mariela Grande"),
            ("Science", "Catarina Morais")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Case-insensitive matching
        slot_num = parser.find_slot_by_subject("Science", teacher_name="mariela grande")
        assert slot_num == 1
        
        slot_num = parser.find_slot_by_subject("Science", teacher_name="MARIELA GRANDE")
        assert slot_num == 1
        
        slot_num = parser.find_slot_by_subject("Science", teacher_name="Catarina Morais")
        assert slot_num == 2


class TestSubjectAwareExtraction:
    """Test extraction using subject instead of slot number."""
    
    def test_extract_hyperlinks_by_subject(self, tmp_path):
        """Test hyperlink extraction using subject."""
        doc_path = tmp_path / "test_hyperlinks.docx"
        doc = Document()
        
        # Slot 1: ELA with 2 hyperlinks
        meta1 = doc.add_table(rows=2, cols=2)
        meta1.rows[0].cells[0].text = "Name: Teacher A"
        meta1.rows[1].cells[0].text = "Subject: ELA"
        
        daily1 = doc.add_table(rows=2, cols=3)
        daily1.rows[0].cells[1].text = "Monday"
        # Add hyperlinks (simplified - just text for now)
        daily1.rows[1].cells[1].text = "ELA Link 1"
        daily1.rows[1].cells[2].text = "ELA Link 2"
        
        # Slot 2: Math with 3 hyperlinks
        meta2 = doc.add_table(rows=2, cols=2)
        meta2.rows[0].cells[0].text = "Name: Teacher B"
        meta2.rows[1].cells[0].text = "Subject: Math"
        
        daily2 = doc.add_table(rows=2, cols=3)
        daily2.rows[0].cells[1].text = "Monday"
        daily2.rows[1].cells[1].text = "Math Link 1"
        daily2.rows[1].cells[2].text = "Math Link 2"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        parser = DOCXParser(str(doc_path))
        
        # Extract by subject (no hyperlinks in this simplified test, but method should work)
        ela_links = parser.extract_hyperlinks_for_slot(subject="ELA")
        math_links = parser.extract_hyperlinks_for_slot(subject="Math")
        
        # Should not raise errors
        assert isinstance(ela_links, list)
        assert isinstance(math_links, list)
    
    def test_extract_images_by_subject(self, tmp_path):
        """Test image extraction using subject."""
        doc_path = tmp_path / "test_images.docx"
        doc = Document()
        
        # Slot 1: ELA with image
        meta1 = doc.add_table(rows=2, cols=2)
        meta1.rows[0].cells[0].text = "Name: Teacher A"
        meta1.rows[1].cells[0].text = "Subject: ELA"
        
        daily1 = doc.add_table(rows=2, cols=3)
        daily1.rows[0].cells[1].text = "Monday"
        
        # Add image
        img1 = PILImage.new('RGB', (2, 2), color='red')
        img1_bytes = io.BytesIO()
        img1.save(img1_bytes, format='PNG')
        img1_bytes.seek(0)
        run1 = daily1.rows[1].cells[1].paragraphs[0].add_run()
        run1.add_picture(img1_bytes, width=1)
        
        # Slot 2: Math (no image)
        meta2 = doc.add_table(rows=2, cols=2)
        meta2.rows[0].cells[0].text = "Name: Teacher B"
        meta2.rows[1].cells[0].text = "Subject: Math"
        
        daily2 = doc.add_table(rows=2, cols=3)
        daily2.rows[0].cells[1].text = "Monday"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        parser = DOCXParser(str(doc_path))
        
        # Extract by subject
        ela_images = parser.extract_images_for_slot(subject="ELA")
        math_images = parser.extract_images_for_slot(subject="Math")
        
        # ELA should have 1 image, Math should have 0
        assert len(ela_images) == 1
        assert len(math_images) == 0
        assert ela_images[0]['table_idx'] == 1
    
    def test_extract_content_by_subject_misaligned(self, tmp_path):
        """Test content extraction with misaligned slots."""
        # Create doc: Slot 1 = ELA/SS, Slot 2 = Math
        doc_path = create_doc_with_subjects(tmp_path, [
            ("ELA/SS", "Donna Savoca"),
            ("Math", "Donna Savoca")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Weekly plan says Slot 2 wants ELA/SS
        # But in file, ELA/SS is Slot 1
        # extract_subject_content_for_slot should find it in Slot 1
        content = parser.extract_subject_content_for_slot(
            slot_number=2,  # What weekly plan says
            subject="ELA/SS"  # What we're looking for
        )
        
        # Should extract from Slot 1 (where ELA/SS actually is)
        assert content['found'] is True
        assert content['slot_number'] == 1  # Corrected to actual slot
        assert "ELA/SS content" in content['full_text']
    
    def test_extract_requires_slot_or_subject(self, tmp_path):
        """Test that extraction requires either slot_number or subject."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("Math", "Teacher A")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Should raise error if neither provided
        with pytest.raises(ValueError, match="Must provide either slot_number or subject"):
            parser.extract_hyperlinks_for_slot()
        
        with pytest.raises(ValueError, match="Must provide either slot_number or subject"):
            parser.extract_images_for_slot()


class TestAliasNormalization:
    """Test alias normalization and matching logic."""
    
    def test_punctuation_stripped(self, tmp_path):
        """Test that punctuation is stripped from subject names."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("ELA/SS:", "Teacher A"),  # Extra colon
            ("Math.", "Teacher B")  # Extra period
        ])
        parser = DOCXParser(str(doc_path))
        
        # Should match despite punctuation
        slot_num = parser.find_slot_by_subject("ELA/SS")
        assert slot_num == 1
        
        slot_num = parser.find_slot_by_subject("Math")
        assert slot_num == 2
    
    def test_whitespace_normalized(self, tmp_path):
        """Test that whitespace is normalized."""
        doc_path = create_doc_with_subjects(tmp_path, [
            ("  Math  ", "Teacher A"),  # Extra whitespace
            ("ELA", "Teacher B")
        ])
        parser = DOCXParser(str(doc_path))
        
        # Should match despite whitespace
        slot_num = parser.find_slot_by_subject("Math")
        assert slot_num == 1
