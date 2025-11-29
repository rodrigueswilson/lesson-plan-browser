"""
Check if page breaks exist between slots in the consolidated DOCX.
"""

import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

# Get the most recent Weekly file
output_dir = Path(r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W38")
weekly_files = list(output_dir.glob("*Weekly*.docx"))

if not weekly_files:
    print("No Weekly DOCX files found")
    sys.exit(1)

# Get most recent
weekly_file = max(weekly_files, key=lambda p: p.stat().st_mtime)

print("=" * 70)
print(f"CHECKING PAGE BREAKS IN: {weekly_file.name}")
print("=" * 70)

doc = Document(weekly_file)

# Count page breaks in the document
page_breaks = 0
page_break_locations = []

for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
            page_breaks += 1
            page_break_locations.append(i)

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Page breaks found: {page_breaks}")

if page_break_locations:
    print(f"\nPage break locations (paragraph indices):")
    for loc in page_break_locations:
        print(f"  - Paragraph {loc}")
else:
    print("\n⚠ WARNING: No explicit page breaks found!")
    print("\nNote: docxcompose may use section breaks instead of page breaks.")
    print("Let me check for section breaks...")
    
    # Check for section breaks
    sections = len(doc.sections)
    print(f"\nNumber of sections: {sections}")
    
    if sections > 1:
        print(f"✓ Document has {sections} sections (likely one per slot)")
        print("  Section breaks act as page breaks in Word.")
    else:
        print("✗ Only 1 section found - slots may not be on separate pages")

# Count tables to estimate slots
metadata_tables = sum(1 for t in doc.tables if len(t.rows) == 1 and len(t.columns) == 5)
print(f"\nEstimated slots (from metadata tables): {metadata_tables}")

if sections == metadata_tables:
    print(f"✓ Sections match slots - each slot is on its own page!")
elif page_breaks >= metadata_tables - 1:
    print(f"✓ Page breaks match expected count - each slot is on its own page!")
else:
    print(f"⚠ Page break count doesn't match - may need to add explicit breaks")
