"""
Tests for "No School" day detection functionality.
"""

import pytest
from pathlib import Path

from tools.docx_parser import DOCXParser


class TestNoSchoolDetection:
    """Test 'No School' day detection."""

    def test_no_school_detection_with_fixture(self):
        """Test 'No School' day detection with test fixture."""
        fixture_path = Path("tests/fixtures/no_school_day.docx")
        
        if not fixture_path.exists():
            pytest.skip(f"Test fixture not found: {fixture_path}")
        
        parser = DOCXParser(str(fixture_path))
        assert parser.is_no_school_day() == True, \
            "Should detect 'No School' in fixture file"

    def test_regular_lesson_not_no_school(self):
        """Test that regular lesson is not detected as 'No School'."""
        fixture_path = Path("tests/fixtures/regular_lesson.docx")
        
        if not fixture_path.exists():
            pytest.skip(f"Test fixture not found: {fixture_path}")
        
        parser = DOCXParser(str(fixture_path))
        assert parser.is_no_school_day() == False, \
            "Should NOT detect 'No School' in regular lesson"

    def test_no_school_pattern_no_school(self):
        """Test 'No School' pattern detection."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        # Create temporary document with "No School" text
        doc = Document()
        doc.add_paragraph("NO SCHOOL - Professional Development Day")
        
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_no_school_day() == True
        finally:
            Path(tmp_path).unlink()

    def test_no_school_pattern_holiday(self):
        """Test 'Holiday' pattern detection."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        doc = Document()
        doc.add_paragraph("HOLIDAY - Thanksgiving")
        
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_no_school_day() == True
        finally:
            Path(tmp_path).unlink()

    def test_no_school_pattern_school_closed(self):
        """Test 'School Closed' pattern detection."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        doc = Document()
        doc.add_paragraph("School Closed - Holiday Weekend")
        
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_no_school_day() == True
        finally:
            Path(tmp_path).unlink()

    def test_no_school_pattern_professional_development(self):
        """Test 'Professional Development' pattern detection."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        doc = Document()
        doc.add_paragraph("Professional Development Day - No Students")
        
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_no_school_day() == True
        finally:
            Path(tmp_path).unlink()

    def test_no_school_pattern_teacher_workday(self):
        """Test 'Teacher Workday' pattern detection."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        doc = Document()
        doc.add_paragraph("Teacher Workday - Planning")
        
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_no_school_day() == True
        finally:
            Path(tmp_path).unlink()

    def test_no_school_case_insensitive(self):
        """Test that detection is case-insensitive."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        test_cases = [
            "no school",
            "No School",
            "NO SCHOOL",
            "No ScHoOl",  # Mixed case
        ]
        
        for text in test_cases:
            doc = Document()
            doc.add_paragraph(text)
            
            with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                parser = DOCXParser(tmp_path)
                assert parser.is_no_school_day() == True, \
                    f"Should detect '{text}' as No School"
            finally:
                Path(tmp_path).unlink()

    def test_regular_content_not_detected(self):
        """Test that regular lesson content is not falsely detected."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        # Regular lesson content that should NOT trigger detection
        doc = Document()
        doc.add_paragraph("Unit: Ancient Rome - Lesson 1")
        doc.add_paragraph("Objective: Students will explain Roman law")
        doc.add_paragraph("Instruction: Review slideshow on Roman legal system")
        
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_no_school_day() == False, \
                "Should NOT detect regular lesson as No School"
        finally:
            Path(tmp_path).unlink()

    def test_no_school_with_extra_whitespace(self):
        """Test detection with various whitespace patterns."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        test_cases = [
            "no  school",      # Double space
            "no   school",     # Triple space
            "no\tschool",      # Tab
            "no\nschool",      # Newline (in same paragraph)
        ]
        
        for text in test_cases:
            doc = Document()
            doc.add_paragraph(text)
            
            with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                parser = DOCXParser(tmp_path)
                # Regex \s+ should match any whitespace
                assert parser.is_no_school_day() == True, \
                    f"Should detect 'no school' with whitespace: {repr(text)}"
            finally:
                Path(tmp_path).unlink()

    def test_staff_development_variations(self):
        """Test detection of Staff Development variations."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        test_cases = [
            "Staff Development",
            "No School- Staff Development",
            "No School - Staff Development",
            "Teacher Development Day",
            "PD Day",
            "In-Service Training",
            "Inservice Day",
        ]
        
        for text in test_cases:
            doc = Document()
            doc.add_paragraph(text)
            
            with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                parser = DOCXParser(tmp_path)
                assert parser.is_no_school_day() == True, \
                    f"Should detect '{text}' as No School"
            finally:
                Path(tmp_path).unlink()

    def test_conference_and_planning_days(self):
        """Test detection of conference and planning days."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        test_cases = [
            "Conference Day",
            "Parent-Teacher Conference",
            "Parent Teacher Conference",
            "Planning Day",
            "Prep Day",
        ]
        
        for text in test_cases:
            doc = Document()
            doc.add_paragraph(text)
            
            with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                parser = DOCXParser(tmp_path)
                assert parser.is_no_school_day() == True, \
                    f"Should detect '{text}' as No School"
            finally:
                Path(tmp_path).unlink()

    def test_early_dismissal_patterns(self):
        """Test detection of early dismissal patterns."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        test_cases = [
            "Early Dismissal",
            "Half Day",
            "Early Release",
        ]
        
        for text in test_cases:
            doc = Document()
            doc.add_paragraph(text)
            
            with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                parser = DOCXParser(tmp_path)
                assert parser.is_no_school_day() == True, \
                    f"Should detect '{text}' as No School"
            finally:
                Path(tmp_path).unlink()

    def test_hyphenated_no_school(self):
        """Test detection of hyphenated 'No School' variations."""
        from docx import Document
        from tempfile import NamedTemporaryFile
        
        test_cases = [
            "No-School",
            "No - School",
            "No-School Day",
        ]
        
        for text in test_cases:
            doc = Document()
            doc.add_paragraph(text)
            
            with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                parser = DOCXParser(tmp_path)
                assert parser.is_no_school_day() == True, \
                    f"Should detect '{text}' as No School"
            finally:
                Path(tmp_path).unlink()

    def test_is_day_no_school_staff_development_and_pd_day(self):
        """Test that is_day_no_school() detects Staff Development and PD Day for short cell text."""
        from docx import Document
        from tempfile import NamedTemporaryFile

        doc = Document()
        doc.add_paragraph("Placeholder")
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        try:
            parser = DOCXParser(tmp_path)
            assert parser.is_day_no_school("Staff Development") is True
            assert parser.is_day_no_school("PD Day") is True
            assert parser.is_day_no_school("Planning Day") is True
            assert parser.is_day_no_school("No School") is True
            assert parser.is_day_no_school("Conference Day") is True
            assert parser.is_day_no_school("In-Service") is True
            assert parser.is_day_no_school("Some real lesson content here for the day.") is False
        finally:
            Path(tmp_path).unlink()


if __name__ == "__main__":
    # Run tests
    pytest.main([__file__, "-v"])
