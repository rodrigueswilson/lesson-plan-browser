"""
Tests for table width normalization functionality.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Inches

from tools.docx_utils import (
    normalize_table_column_widths,
    normalize_all_tables,
    get_table_info
)


class TestTableWidthNormalization:
    """Test table width normalization utilities."""

    def test_normalize_single_table_basic(self):
        """Test normalizing a single table with equal widths."""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        
        # Set unequal widths
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(2.0)
        table.columns[2].width = Inches(3.0)
        
        # Normalize
        normalize_table_column_widths(table, total_width_inches=6.0)
        
        # Check all equal
        expected_width = int(Inches(6.0) / 3)
        for column in table.columns:
            assert column.width == expected_width, \
                f"Column width {column.width} != expected {expected_width}"

    def test_normalize_single_table_default_width(self):
        """Test normalizing with default 6.5 inch width."""
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        
        # Normalize with default width
        normalize_table_column_widths(table)
        
        # Check all equal
        expected_width = int(Inches(6.5) / 5)
        for column in table.columns:
            assert column.width == expected_width

    def test_normalize_empty_table(self):
        """Test that empty tables are handled gracefully."""
        doc = Document()
        table = doc.add_table(rows=0, cols=0)
        
        # Should not raise exception
        normalize_table_column_widths(table)

    def test_normalize_all_tables_in_document(self):
        """Test normalizing all tables in a document."""
        doc = Document()
        
        # Create 3 tables with different structures
        table1 = doc.add_table(rows=2, cols=3)
        doc.add_paragraph()  # Separator
        table2 = doc.add_table(rows=3, cols=5)
        doc.add_paragraph()  # Separator
        table3 = doc.add_table(rows=1, cols=2)
        
        # Set unequal widths
        for table in [table1, table2, table3]:
            for i, column in enumerate(table.columns):
                column.width = Inches(1.0 + i * 0.5)
        
        # Normalize all
        count = normalize_all_tables(doc, total_width_inches=6.0)
        
        # Check count
        assert count == 3, f"Expected 3 tables, got {count}"
        
        # Check all tables normalized
        for table in doc.tables:
            if table.columns:
                first_width = table.columns[0].width
                for column in table.columns:
                    assert column.width == first_width, \
                        "All columns in table should have equal width"

    def test_normalize_with_fixture(self):
        """Test normalizing tables in test fixture."""
        fixture_path = Path("tests/fixtures/lesson_with_tables.docx")
        
        if not fixture_path.exists():
            pytest.skip(f"Test fixture not found: {fixture_path}")
        
        doc = Document(str(fixture_path))
        
        # Get original widths
        original_widths = []
        for table in doc.tables:
            if table.columns:
                original_widths.append([col.width for col in table.columns])
        
        # Normalize
        count = normalize_all_tables(doc)
        
        assert count > 0, "Should have normalized at least one table"
        
        # Check all tables now have equal column widths
        for table in doc.tables:
            if table.columns:
                first_width = table.columns[0].width
                for column in table.columns:
                    assert column.width == first_width

    def test_get_table_info(self):
        """Test getting table information."""
        doc = Document()
        table = doc.add_table(rows=3, cols=4)
        
        # Set widths
        for i, column in enumerate(table.columns):
            column.width = Inches(1.0 + i * 0.5)
        
        info = get_table_info(table)
        
        assert info['num_rows'] == 3
        assert info['num_columns'] == 4
        assert len(info['column_widths']) == 4
        assert info['column_widths'][0] == Inches(1.0)
        assert info['column_widths'][1] == Inches(1.5)

    def test_normalize_preserves_content(self):
        """Test that normalization doesn't affect cell content."""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        
        # Add content
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(0, 2).text = "Header 3"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        table.cell(1, 2).text = "Data 3"
        
        # Normalize
        normalize_table_column_widths(table)
        
        # Check content preserved
        assert table.cell(0, 0).text == "Header 1"
        assert table.cell(0, 1).text == "Header 2"
        assert table.cell(0, 2).text == "Header 3"
        assert table.cell(1, 0).text == "Data 1"
        assert table.cell(1, 1).text == "Data 2"
        assert table.cell(1, 2).text == "Data 3"

    def test_normalize_with_merged_cells(self):
        """Test normalization with merged cells."""
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        
        # Merge some cells
        cell_a = table.cell(0, 0)
        cell_b = table.cell(0, 2)
        cell_a.merge(cell_b)  # Merge first row
        
        # Normalize - should not raise exception
        normalize_table_column_widths(table)
        
        # Check widths are set
        for column in table.columns:
            assert column.width > 0

    def test_normalize_different_widths(self):
        """Test normalization with different total widths."""
        test_widths = [5.0, 6.0, 6.5, 7.0, 8.0]
        
        for total_width in test_widths:
            doc = Document()
            table = doc.add_table(rows=2, cols=4)
            
            normalize_table_column_widths(table, total_width_inches=total_width)
            
            expected_width = int(Inches(total_width) / 4)
            for column in table.columns:
                assert column.width == expected_width, \
                    f"Width mismatch for total_width={total_width}"

    def test_normalize_single_column_table(self):
        """Test normalizing a table with single column."""
        doc = Document()
        table = doc.add_table(rows=3, cols=1)
        
        normalize_table_column_widths(table, total_width_inches=6.5)
        
        # Single column should get full width
        expected_width = int(Inches(6.5))
        assert table.columns[0].width == expected_width

    def test_normalize_many_columns(self):
        """Test normalizing a table with many columns."""
        doc = Document()
        table = doc.add_table(rows=2, cols=10)
        
        normalize_table_column_widths(table, total_width_inches=6.5)
        
        # Each column should get equal share
        expected_width = int(Inches(6.5) / 10)
        for column in table.columns:
            assert column.width == expected_width


if __name__ == "__main__":
    # Run tests
    pytest.main([__file__, "-v"])
