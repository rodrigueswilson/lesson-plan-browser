"""
Comprehensive tests for slot-aware extraction functionality.

Tests validate_slot_structure, extract_hyperlinks_for_slot, and extract_images_for_slot
to ensure no cross-contamination between slots.
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from tools.docx_parser import DOCXParser, validate_slot_structure


def add_hyperlink(paragraph, text, url):
    """
    Add a real hyperlink to a paragraph.
    
    This creates actual DOCX hyperlink relationships that will be extracted.
    """
    # Get the paragraph's parent part
    part = paragraph.part
    
    # Create relationship
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run with text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style as hyperlink (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    run.append(rPr)
    
    # Add text
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def create_1x1_png():
    """Create a minimal 1x1 pixel PNG image using PIL."""
    import io
    try:
        from PIL import Image as PILImage
        
        # Create a 1x1 red pixel image
        img = PILImage.new('RGB', (1, 1), color='red')
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return img_bytes
    except ImportError:
        # Fallback: use a proper minimal PNG
        png_data = (
            b'\x89PNG\r\n\x1a\n'
            b'\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01'
            b'\x08\x02\x00\x00\x00\x90wS\xde'
            b'\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0\x00\x00\x00\x03\x00\x01'
            b'\x00\x18\xdd\x8d\xb4'
            b'\x00\x00\x00\x00IEND\xaeB`\x82'
        )
        img_bytes = io.BytesIO(png_data)
        img_bytes.seek(0)
        return img_bytes


def add_image_to_cell(cell, image_stream, width_inches=1.0):
    """
    Add a real image to a table cell.
    
    This creates actual DOCX image relationships that will be extracted.
    """
    from docx.shared import Inches
    
    # Get or create paragraph in cell
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    
    # Add picture to paragraph
    run = paragraph.add_run()
    run.add_picture(image_stream, width=Inches(width_inches))
    
    return run


class TestValidateSlotStructure:
    """Test slot structure validation function."""
    
    def test_validate_slot_structure_valid_4_slots(self, tmp_path):
        """Test validation passes for valid 4-slot structure (9 tables)."""
        doc_path = tmp_path / "test_4_slots.docx"
        doc = Document()
        
        # Add 4 slots (8 tables) + 1 signature table = 9 tables
        for slot_num in range(1, 5):
            # Metadata table
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            meta_table.rows[1].cells[0].text = f"Grade: {slot_num}"
            meta_table.rows[2].cells[0].text = f"Subject: Subject {slot_num}"
            
            # Daily plans table
            daily_table = doc.add_table(rows=5, cols=6)
            daily_table.rows[0].cells[0].text = "Section"
            daily_table.rows[0].cells[1].text = "Monday"
            daily_table.rows[0].cells[2].text = "Tuesday"
            daily_table.rows[0].cells[3].text = "Wednesday"
            daily_table.rows[0].cells[4].text = "Thursday"
            daily_table.rows[0].cells[5].text = "Friday"
        
        # Signature table
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Required Signatures"
        
        doc.save(str(doc_path))
        
        # Test validation
        test_doc = Document(str(doc_path))
        
        # Slot 1: tables 0-1
        table_start, table_end = validate_slot_structure(test_doc, 1)
        assert table_start == 0
        assert table_end == 1
        
        # Slot 4: tables 6-7
        table_start, table_end = validate_slot_structure(test_doc, 4)
        assert table_start == 6
        assert table_end == 7
    
    def test_validate_slot_structure_valid_5_slots(self, tmp_path):
        """Test validation passes for valid 5-slot structure (11 tables)."""
        doc_path = tmp_path / "test_5_slots.docx"
        doc = Document()
        
        # Add 5 slots (10 tables) + 1 signature table = 11 tables
        for slot_num in range(1, 6):
            # Metadata table
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            
            # Daily plans table
            daily_table = doc.add_table(rows=5, cols=6)
            daily_table.rows[0].cells[1].text = "Monday"
            daily_table.rows[0].cells[2].text = "Tuesday"
        
        # Signature table
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        # Test validation
        test_doc = Document(str(doc_path))
        
        # Slot 5: tables 8-9
        table_start, table_end = validate_slot_structure(test_doc, 5)
        assert table_start == 8
        assert table_end == 9
    
    def test_validate_slot_structure_missing_signature(self, tmp_path):
        """Test validation succeeds without signature table (logs warning, proceeds with even table count)."""
        doc_path = tmp_path / "test_no_sig.docx"
        doc = Document()
        
        # Add 4 slots but NO signature table (8 tables total)
        for slot_num in range(1, 5):
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            
            daily_table = doc.add_table(rows=5, cols=6)
            daily_table.rows[0].cells[1].text = "Monday"
        
        doc.save(str(doc_path))
        
        # Validation proceeds without signature table (original behavior: log warning, available_slots = table_count // 2)
        test_doc = Document(str(doc_path))
        table_start, table_end = validate_slot_structure(test_doc, 1)
        assert table_start == 0
        assert table_end == 1
    
    def test_validate_slot_structure_invalid_table_count(self, tmp_path):
        """Test validation fails for unexpected table count."""
        doc_path = tmp_path / "test_8_tables.docx"
        doc = Document()
        
        # Add 8 tables (invalid - should be 3, 7, 9, or 11)
        # 8 tables: (8-1)//2 = 3 slots, expected = 3*2+1 = 7, but we have 8!
        # Make first two tables look valid to reach table count check
        for i in range(8):
            table = doc.add_table(rows=2, cols=6)
            if i == 7:
                table.rows[0].cells[0].text = "Signature"
            elif i == 0:
                table.rows[0].cells[0].text = "Name: Teacher"
            elif i == 1:
                table.rows[0].cells[1].text = "Monday"
                table.rows[0].cells[2].text = "Tuesday"
            else:
                table.rows[0].cells[0].text = f"Table {i}"
        
        doc.save(str(doc_path))
        
        # Test validation
        test_doc = Document(str(doc_path))
        
        # Will fail on table count check since 8 != 7 (expected for 3 slots)
        with pytest.raises(ValueError, match="Unexpected table count"):
            validate_slot_structure(test_doc, 1)
    
    def test_validate_slot_structure_slot_exceeds(self, tmp_path):
        """Test validation fails when slot > available."""
        doc_path = tmp_path / "test_4_slots.docx"
        doc = Document()
        
        # Add 4 slots (9 tables total)
        for slot_num in range(1, 5):
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            
            daily_table = doc.add_table(rows=5, cols=6)
            daily_table.rows[0].cells[1].text = "Monday"
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        # Test validation
        test_doc = Document(str(doc_path))
        
        with pytest.raises(ValueError, match="only 4 slots available"):
            validate_slot_structure(test_doc, 5)
    
    def test_validate_slot_structure_no_tables(self, tmp_path):
        """Test validation fails for document with no tables."""
        doc_path = tmp_path / "test_no_tables.docx"
        doc = Document()
        doc.add_paragraph("No tables here")
        doc.save(str(doc_path))
        
        test_doc = Document(str(doc_path))
        
        with pytest.raises(ValueError, match="Document has no tables"):
            validate_slot_structure(test_doc, 1)


class TestExtractHyperlinksForSlot:
    """Test slot-aware hyperlink extraction."""
    
    def test_extract_hyperlinks_for_slot_filters_by_table(self, tmp_path):
        """Test slot extraction gets only slot's links."""
        doc_path = tmp_path / "test_hyperlinks.docx"
        doc = Document()
        
        # Create 2 slots + signature
        for slot_num in range(1, 3):
            # Metadata table
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            
            # Daily plans table with REAL hyperlinks
            daily_table = doc.add_table(rows=3, cols=6)
            daily_table.rows[0].cells[1].text = "Monday"
            daily_table.rows[1].cells[0].text = "Objective"
            
            # Add REAL hyperlink to slot's table
            paragraph = daily_table.rows[1].cells[1].paragraphs[0]
            add_hyperlink(paragraph, f"Link for slot {slot_num}", f"https://example.com/slot{slot_num}")
        
        # Signature table
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        # Test extraction
        parser = DOCXParser(str(doc_path))
        
        # Extract for slot 1 (tables 0-1)
        slot1_links = parser.extract_hyperlinks_for_slot(1)
        
        # Should have at least one link
        assert len(slot1_links) > 0, "Slot 1 should have hyperlinks"
        
        # All links should be from tables 0-1
        for link in slot1_links:
            assert link['table_idx'] in [0, 1], f"Link from wrong table: {link['table_idx']}"
            assert 'slot1' in link['url'], f"Wrong URL for slot 1: {link['url']}"
        
        # Extract for slot 2 (tables 2-3)
        slot2_links = parser.extract_hyperlinks_for_slot(2)
        
        # Should have at least one link
        assert len(slot2_links) > 0, "Slot 2 should have hyperlinks"
        
        # All links should be from tables 2-3
        for link in slot2_links:
            assert link['table_idx'] in [2, 3], f"Link from wrong table: {link['table_idx']}"
            assert 'slot2' in link['url'], f"Wrong URL for slot 2: {link['url']}"
    
    def test_extract_hyperlinks_for_slot_no_paragraphs(self, tmp_path):
        """Test paragraph links are excluded."""
        doc_path = tmp_path / "test_no_paragraphs.docx"
        doc = Document()
        
        # Add paragraph with REAL hyperlink (should be excluded)
        para = doc.add_paragraph()
        add_hyperlink(para, "Paragraph link", "https://example.com/paragraph")
        
        # Add 1 slot + signature with table hyperlink
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.rows[0].cells[0].text = "Name: Teacher"
        
        daily_table = doc.add_table(rows=3, cols=6)
        daily_table.rows[0].cells[1].text = "Monday"
        daily_table.rows[1].cells[0].text = "Objective"
        
        # Add table hyperlink
        table_para = daily_table.rows[1].cells[1].paragraphs[0]
        add_hyperlink(table_para, "Table link", "https://example.com/table")
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        # Test extraction
        parser = DOCXParser(str(doc_path))
        
        # Slot-aware extraction should only get table links
        slot_links = parser.extract_hyperlinks_for_slot(1)
        
        # Should have the table link
        assert len(slot_links) > 0, "Should have table hyperlinks"
        
        # No paragraph links (context_type != 'paragraph')
        for link in slot_links:
            assert link['context_type'] == 'table', f"Found paragraph link: {link}"
            assert 'table' in link['url'], f"Wrong link extracted: {link['url']}"
        
        # Verify paragraph link exists in document but wasn't extracted
        all_links = parser.extract_hyperlinks()
        paragraph_links = [l for l in all_links if l['context_type'] == 'paragraph']
        assert len(paragraph_links) > 0, "Paragraph link should exist in full extraction"


class TestExtractImagesForSlot:
    """Test slot-aware image extraction."""
    
    def test_extract_images_for_slot_filters_by_table(self, tmp_path):
        """Test slot extraction gets only slot's images."""
        doc_path = tmp_path / "test_images.docx"
        doc = Document()
        
        # Create 2 slots + signature with REAL images
        for slot_num in range(1, 3):
            # Metadata table
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            
            # Daily plans table with REAL embedded image
            daily_table = doc.add_table(rows=3, cols=6)
            daily_table.rows[0].cells[1].text = "Monday"
            daily_table.rows[1].cells[0].text = "Objective"
            
            # Add UNIQUE image to this slot's table (different color per slot)
            # This ensures python-docx creates separate relationships
            import io
            from PIL import Image as PILImage
            color = 'red' if slot_num == 1 else 'blue'
            img = PILImage.new('RGB', (2, 2), color=color)  # 2x2 to make them different
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            img_bytes.seek(0)
            add_image_to_cell(daily_table.rows[1].cells[1], img_bytes, width_inches=0.5)
        
        # Signature table
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        # Test extraction
        parser = DOCXParser(str(doc_path))
        
        # Extract for slot 1 (tables 0-1)
        slot1_images = parser.extract_images_for_slot(1)
        
        # Should have at least one image
        assert len(slot1_images) > 0, "Slot 1 should have images"
        
        # All images should be from tables 0-1
        for img in slot1_images:
            assert img.get('table_idx') is not None, "Image missing table_idx"
            assert 0 <= img['table_idx'] <= 1, f"Image from wrong table: {img['table_idx']}"
            assert img.get('data'), "Image missing data"
            assert 'png' in img.get('content_type', '').lower(), "Image should be PNG"
        
        # Extract for slot 2 (tables 2-3)
        slot2_images = parser.extract_images_for_slot(2)
        
        # Should have at least one image
        assert len(slot2_images) > 0, "Slot 2 should have images"
        
        # All images should be from tables 2-3
        for img in slot2_images:
            assert img.get('table_idx') is not None, "Image missing table_idx"
            assert 2 <= img['table_idx'] <= 3, f"Image from wrong table: {img['table_idx']}"
            assert img.get('data'), "Image missing data"
        
        # Verify no overlap
        slot1_tables = set(img['table_idx'] for img in slot1_images)
        slot2_tables = set(img['table_idx'] for img in slot2_images)
        assert len(slot1_tables & slot2_tables) == 0, "Image table indices overlap between slots"


class TestIntegrationNoCrossContamination:
    """Integration tests to verify no cross-contamination."""
    
    def test_no_cross_contamination_between_slots(self, tmp_path):
        """Test that slots don't share hyperlinks or images."""
        doc_path = tmp_path / "test_integration.docx"
        doc = Document()
        
        # Create 2 distinct slots
        for slot_num in range(1, 3):
            # Metadata table
            meta_table = doc.add_table(rows=3, cols=2)
            meta_table.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            meta_table.rows[1].cells[0].text = f"Grade: {slot_num}"
            meta_table.rows[2].cells[0].text = f"Subject: Subject {slot_num}"
            
            # Daily plans table
            daily_table = doc.add_table(rows=5, cols=6)
            daily_table.rows[0].cells[0].text = "Section"
            daily_table.rows[0].cells[1].text = "Monday"
            daily_table.rows[1].cells[0].text = "Objective"
            daily_table.rows[1].cells[1].text = f"Objective for slot {slot_num}"
        
        # Signature table
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = "Required Signatures"
        
        doc.save(str(doc_path))
        
        # Test extraction
        parser = DOCXParser(str(doc_path))
        
        # Extract for both slots
        slot1_links = parser.extract_hyperlinks_for_slot(1)
        slot2_links = parser.extract_hyperlinks_for_slot(2)
        
        slot1_images = parser.extract_images_for_slot(1)
        slot2_images = parser.extract_images_for_slot(2)
        
        # Verify no overlap in table indices
        slot1_tables = set()
        for link in slot1_links:
            slot1_tables.add(link['table_idx'])
        for img in slot1_images:
            if img.get('table_idx') is not None:
                slot1_tables.add(img['table_idx'])
        
        slot2_tables = set()
        for link in slot2_links:
            slot2_tables.add(link['table_idx'])
        for img in slot2_images:
            if img.get('table_idx') is not None:
                slot2_tables.add(img['table_idx'])
        
        # No overlap
        assert len(slot1_tables & slot2_tables) == 0, \
            f"Table overlap detected: {slot1_tables & slot2_tables}"
        
        # Slot 1 should only have tables 0-1
        assert slot1_tables <= {0, 1}, f"Slot 1 has wrong tables: {slot1_tables}"
        
        # Slot 2 should only have tables 2-3
        assert slot2_tables <= {2, 3}, f"Slot 2 has wrong tables: {slot2_tables}"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
