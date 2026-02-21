
import asyncio
import os
import sys
from pathlib import Path
from datetime import datetime
from unittest.mock import MagicMock, patch
import json

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.batch_processor import BatchProcessor

async def test_originals_export():
    print("Testing Originals Export refined layout...")
    
    # Mock DB and LLM
    db = MagicMock()
    llm = MagicMock()
    processor = BatchProcessor(llm)
    processor.db = db
    processor._user_base_path = str(Path("./test_output").absolute())
    
    user_id = "test_user"
    week_of = "2025-12-22"
    plan_id = "test_plan_id"
    
    # Create mock original plans with structured day content
    # Slot 1: Science (Mon, Tue)
    plan1 = MagicMock()
    plan1.slot_number = 1
    plan1.subject = "Science"
    plan1.source_file_name = "science_source.docx"
    plan1.primary_teacher_name = "Dr. Bio"
    plan1.grade = "9"
    plan1.homeroom = "101"
    plan1.extracted_at = datetime.now()
    plan1.has_no_school = False
    
    plan1.monday_content = {
        "unit_lesson": "Biology 101",
        "objective": {"content_objective": "Understand cell structure"},
        "instruction": {"activities": "Microscope lab"},
        "materials": {"root": ["Microscope", "Slides"]},
        "assessment": {"primary_assessment": "Lab report"},
        "hyperlinks": {"root": [
            {"text": "Cell Video", "url": "https://example.com/cell", "row_label": "instruction", "row_idx": 4, "cell_idx": 1, "day_hint": "monday"},
            {"text": "Unit Plan", "url": "https://example.com/unit", "row_label": "unit_lesson", "row_idx": 1, "cell_idx": 1, "day_hint": "monday"}
        ]}
    }
    plan1.tuesday_content = {
        "unit_lesson": "Biology 101",
        "objective": {"content_objective": "Understand DNA"},
        "instruction": {"activities": "DNA modeling"}
    }
    plan1.wednesday_content = None
    plan1.thursday_content = None
    plan1.friday_content = None
    
    # Slot 2: Math (Mon, Wed)
    plan2 = MagicMock()
    plan2.slot_number = 2
    plan2.subject = "Math"
    plan2.source_file_name = "math_source.docx"
    plan2.primary_teacher_name = "Dr. Bio" # Same teacher
    plan2.grade = "9"
    plan2.homeroom = "101"
    plan2.extracted_at = datetime.now()
    plan2.has_no_school = False
    
    plan2.monday_content = {
        "unit_lesson": "Algebra 1",
        "objective": {"content_objective": "Solve for X"},
        "instruction": {"activities": "Practice problems"}
    }
    plan2.tuesday_content = None
    plan2.wednesday_content = {
        "unit_lesson": "Algebra 1",
        "objective": {"content_objective": "Y-intercept"},
        "instruction": {"activities": "Graphing"}
    }
    plan2.thursday_content = None
    plan2.friday_content = None
    
    db.get_original_lesson_plans_for_week.return_value = [plan1, plan2]
    
    # Mock file manager and settings
    with patch("tools.batch_processor.get_file_manager") as mock_fm_get, \
         patch("backend.config.settings.DOCX_TEMPLATE_PATH", "input/Lesson Plan Template SY'25-26.docx"):
        fm = MagicMock()
        mock_fm_get.return_value = fm
        fm.get_week_folder.return_value = Path("./test_output/2025-12-22")
        
        # Run the generation
        output_path_str = await processor._generate_combined_original_docx(
            user_id, week_of, plan_id, week_folder_path="./test_output/2025-12-22"
        )
        
        print(f"Generated DOCX at: {output_path_str}")
        
        if output_path_str and os.path.exists(output_path_str):
            print("SUCCESS: Originals DOCX generated successfully.")
            from docx import Document
            doc = Document(output_path_str)
            print(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables.")
            
            # Expecting 4 tables TOTAL: 2 tables per slot template (meta + grid)
            if len(doc.tables) == 4:
                print(f"SUCCESS: Found {len(doc.tables)} tables.")
                
                # DEBUG: Print all table content
                for i, t in enumerate(doc.tables):
                    t_text = ""
                    for r in t.rows:
                        for c in r.cells:
                            t_text += c.text + " "
                    print(f"Table {i} text snippet: {t_text[:200]}...")

                # DEBUG: Print all paragraphs
                print("--- ALL PARAGRAPHS ---")
                for i, p in enumerate(doc.paragraphs):
                    if p.text.strip():
                        print(f"Para {i}: {p.text}")
                print("-----------------------")

                # Check Science Plan (Tables 0-2)
                sci_meta = doc.tables[0]
                sci_plan = doc.tables[1]
                
                sci_full_text = ""
                for row in sci_plan.rows:
                    for cell in row.cells:
                        sci_full_text += cell.text + " "
                
                if "Biology 101" in sci_full_text and "Science" in sci_meta.rows[0].cells[3].text:
                    print("SUCCESS: Slot 1 (Science) has its own table.")
                else:
                    print("FAILURE: Slot 1 (Science) table structure or content mismatch.")

                # VERIFY: No "Slot 1" header and no "**Content:**" labels
                if "Slot 1:" in sci_full_text:
                    print("FAILURE: Unexpected 'Slot 1:' header found in cell.")
                else:
                    print("SUCCESS: No unexpected 'Slot' headers found.")
                
                if "**Content:**" in sci_full_text or "Content:" in sci_full_text:
                    print("FAILURE: Unexpected 'Content:' label found in cell.")
                else:
                    print("SUCCESS: No unexpected formatting labels found.")

                if "| " in sci_full_text:
                    print(f"FAILURE: Science Full Text contains separators: {sci_full_text[:100]}...")
                
                if "Cell Video" in sci_full_text:
                    print("SUCCESS: Hyperlink found in Science table.")
                else:
                    print("FAILURE: Hyperlink NOT found in Science table.")

                # Check Math Plan (Tables 3-5)
                # Slot 2 Meta should be at index 3, Grid at index 4
                # SLOT 2: MATH
                math_meta = doc.tables[2]
                math_plan = doc.tables[3]
                
                math_full_text = " ".join([c.text for r in math_plan.rows for c in r.cells])
                if "Algebra 1" in math_full_text and "Math" in math_meta.rows[0].cells[3].text:
                    print("SUCCESS: Slot 2 (Math) has its own table.")
                else:
                    print("FAILURE: Slot 2 (Math) table structure or content mismatch.")
                    print(f"Slot 2 meta subject: {math_meta.rows[0].cells[3].text}")
                    print(f"Slot 2 text snippet: {math_full_text[:100]}...")

                # Verify cross-contamination (Bio should NOT be in Math table)
                if "Biology 101" in math_full_text or "Cell Video" in math_full_text:
                    print("FAILURE: Contamination detected! Biology content found in Math table.")
                else:
                    print("SUCCESS: No cross-contamination detected.")

            else:
                print(f"FAILURE: Expected at least 6 tables, found {len(doc.tables)}")
        else:
            print("FAILURE: Originals DOCX was not generated.")

if __name__ == "__main__":
    # Ensure test_output exists
    os.makedirs("./test_output/2025-12-22", exist_ok=True)
    asyncio.run(test_originals_export())
