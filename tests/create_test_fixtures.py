"""
Create test DOCX fixtures for validation and testing.
Generates 5 test files with different characteristics.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path
import sys

def create_regular_lesson():
    """Create a regular lesson plan without special features."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Weekly Lesson Plan")
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # Metadata
    doc.add_paragraph("Week: 10/6-10/10")
    doc.add_paragraph("Grade: 7")
    doc.add_paragraph("Subject: Social Studies")
    doc.add_paragraph()
    
    # Monday
    monday = doc.add_heading("Monday", level=2)
    doc.add_paragraph("Unit: Ancient Rome - Lesson 1")
    doc.add_paragraph("Objective: Students will explain how Roman law and banking systems enabled economic growth.")
    doc.add_paragraph("Instruction: Review slideshow on Roman legal system. Discuss banking innovations.")
    doc.add_paragraph("Assessment: Exit ticket on cause-effect relationships.")
    doc.add_paragraph("Homework: Read pages 45-50 on Roman economy.")
    doc.add_paragraph()
    
    # Tuesday
    doc.add_heading("Tuesday", level=2)
    doc.add_paragraph("Unit: Ancient Rome - Lesson 2")
    doc.add_paragraph("Objective: Students will analyze the impact of Pax Romana on trade.")
    doc.add_paragraph("Instruction: Map activity showing trade routes during Pax Romana.")
    doc.add_paragraph("Assessment: Group presentation on trade goods.")
    doc.add_paragraph()
    
    # Save
    output_path = Path("tests/fixtures/regular_lesson.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path

def create_no_school_day():
    """Create a document indicating 'No School' day."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Weekly Lesson Plan")
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # Metadata
    doc.add_paragraph("Week: 11/25-11/29")
    doc.add_paragraph("Grade: 7")
    doc.add_paragraph("Subject: Social Studies")
    doc.add_paragraph()
    
    # Monday - Regular
    doc.add_heading("Monday", level=2)
    doc.add_paragraph("Unit: Colonial America - Lesson 5")
    doc.add_paragraph("Objective: Students will analyze colonial trade patterns.")
    doc.add_paragraph()
    
    # Tuesday - No School
    doc.add_heading("Tuesday", level=2)
    doc.add_paragraph("NO SCHOOL - Professional Development Day")
    doc.add_paragraph()
    
    # Wednesday - No School
    doc.add_heading("Wednesday", level=2)
    doc.add_paragraph("No School - Thanksgiving Break")
    doc.add_paragraph()
    
    # Thursday - Holiday
    doc.add_heading("Thursday", level=2)
    doc.add_paragraph("HOLIDAY - Thanksgiving")
    doc.add_paragraph()
    
    # Friday - School Closed
    doc.add_heading("Friday", level=2)
    doc.add_paragraph("School Closed - Holiday Weekend")
    doc.add_paragraph()
    
    # Save
    output_path = Path("tests/fixtures/no_school_day.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path

def create_lesson_with_tables():
    """Create a lesson with multiple tables of various structures."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Weekly Lesson Plan - Table Format")
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    doc.add_paragraph()
    
    # Table 1: Simple 3x3 table
    doc.add_heading("Schedule Table", level=3)
    table1 = doc.add_table(rows=3, cols=3)
    table1.style = 'Table Grid'
    
    # Header row
    table1.cell(0, 0).text = "Day"
    table1.cell(0, 1).text = "Topic"
    table1.cell(0, 2).text = "Assessment"
    
    # Data rows
    table1.cell(1, 0).text = "Monday"
    table1.cell(1, 1).text = "Roman Law"
    table1.cell(1, 2).text = "Exit Ticket"
    
    table1.cell(2, 0).text = "Tuesday"
    table1.cell(2, 1).text = "Banking Systems"
    table1.cell(2, 2).text = "Quiz"
    
    doc.add_paragraph()
    
    # Table 2: Wide table with unequal columns (5 columns)
    doc.add_heading("Detailed Lesson Plan", level=3)
    table2 = doc.add_table(rows=2, cols=5)
    table2.style = 'Table Grid'
    
    # Set unequal widths
    table2.columns[0].width = Inches(0.8)
    table2.columns[1].width = Inches(1.5)
    table2.columns[2].width = Inches(2.0)
    table2.columns[3].width = Inches(1.2)
    table2.columns[4].width = Inches(1.0)
    
    # Headers
    headers = ["Day", "Objective", "Instruction", "Materials", "Assessment"]
    for i, header in enumerate(headers):
        table2.cell(0, i).text = header
    
    # Data
    data = ["Monday", "Explain Roman law", "Slideshow and discussion", "Projector, handouts", "Exit ticket"]
    for i, value in enumerate(data):
        table2.cell(1, i).text = value
    
    doc.add_paragraph()
    
    # Table 3: Table with merged cells
    doc.add_heading("Weekly Overview", level=3)
    table3 = doc.add_table(rows=3, cols=3)
    table3.style = 'Table Grid'
    
    # Merge first row (title)
    cell = table3.cell(0, 0)
    cell.merge(table3.cell(0, 2))
    cell.text = "Week of 10/6-10/10: Ancient Rome Unit"
    
    # Regular cells
    table3.cell(1, 0).text = "Monday-Wednesday"
    table3.cell(1, 1).text = "Roman Law & Banking"
    table3.cell(1, 2).text = "Lectures"
    
    table3.cell(2, 0).text = "Thursday-Friday"
    table3.cell(2, 1).text = "Pax Romana"
    table3.cell(2, 2).text = "Group Work"
    
    # Save
    output_path = Path("tests/fixtures/lesson_with_tables.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path

def create_lesson_with_image():
    """Create a lesson with an embedded image placeholder."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Weekly Lesson Plan with Visual Aids")
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    doc.add_paragraph()
    
    # Lesson content
    doc.add_heading("Monday - Roman Architecture", level=2)
    doc.add_paragraph("Objective: Students will identify key features of Roman architecture.")
    doc.add_paragraph()
    
    doc.add_paragraph("Materials:")
    doc.add_paragraph("- Images of Roman buildings (Colosseum, Pantheon, aqueducts)")
    doc.add_paragraph("- Architecture comparison chart")
    doc.add_paragraph()
    
    # Note: Creating actual image requires PIL/image file
    # For now, add placeholder text indicating where image would be
    doc.add_paragraph("[IMAGE PLACEHOLDER: Colosseum diagram with labeled architectural features]")
    doc.add_paragraph("Image caption: The Colosseum demonstrates Roman innovations in arch construction.")
    doc.add_paragraph()
    
    doc.add_paragraph("Instruction: Students will analyze the image and identify three architectural innovations.")
    doc.add_paragraph("Assessment: Label diagram of Roman arch structure.")
    
    # Save
    output_path = Path("tests/fixtures/lesson_with_image.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    print("  Note: Image is placeholder text (python-docx requires actual image file)")
    return output_path

def create_lesson_with_hyperlinks():
    """Create a lesson with hyperlinks."""
    doc = Document()
    
    # Title
    title = doc.add_paragraph("Weekly Lesson Plan with Online Resources")
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    doc.add_paragraph()
    
    # Lesson content
    doc.add_heading("Monday - Roman Empire Research", level=2)
    doc.add_paragraph("Objective: Students will research Roman Empire using primary sources.")
    doc.add_paragraph()
    
    # Add hyperlinks
    doc.add_paragraph("Online Resources:")
    
    # Hyperlink 1
    p1 = doc.add_paragraph()
    p1.add_run("1. ")
    # Note: python-docx hyperlinks require complex XML manipulation
    # For now, add URL as text
    run1 = p1.add_run("Ancient Rome Primary Sources")
    run1.font.color.rgb = RGBColor(0, 0, 255)
    run1.font.underline = True
    p1.add_run(" - https://www.example.com/rome-sources")
    
    # Hyperlink 2
    p2 = doc.add_paragraph()
    p2.add_run("2. ")
    run2 = p2.add_run("Roman Law Database")
    run2.font.color.rgb = RGBColor(0, 0, 255)
    run2.font.underline = True
    p2.add_run(" - https://www.example.com/roman-law")
    
    doc.add_paragraph()
    doc.add_paragraph("Instruction: Students will select one primary source and summarize key points.")
    doc.add_paragraph("Assessment: Written summary (1 paragraph) with source citation.")
    
    # Save
    output_path = Path("tests/fixtures/lesson_with_hyperlinks.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    print("  Note: Hyperlinks are styled text (python-docx hyperlinks require complex XML)")
    return output_path

def create_fixtures_readme(fixtures):
    """Create README documenting the test fixtures."""
    readme_content = """# Test Fixtures

This directory contains test DOCX files for validation and testing.

## Files

### regular_lesson.docx
- **Purpose**: Baseline test with normal lesson plan structure
- **Features**: Standard text-based lesson plan for Monday-Tuesday
- **Use Cases**: 
  - Test basic parsing
  - Test normal processing flow
  - Baseline for comparison

### no_school_day.docx
- **Purpose**: Tests "No School" day detection
- **Features**: Contains multiple "No School" indicators:
  - "NO SCHOOL - Professional Development Day"
  - "No School - Thanksgiving Break"
  - "HOLIDAY - Thanksgiving"
  - "School Closed - Holiday Weekend"
- **Use Cases**:
  - Test `is_no_school_day()` detection
  - Test copy-without-processing workflow
  - Test various "No School" patterns

### lesson_with_tables.docx
- **Purpose**: Tests table handling and width normalization
- **Features**: 
  - Table 1: Simple 3x3 table
  - Table 2: Wide 5-column table with unequal widths
  - Table 3: Table with merged cells
- **Use Cases**:
  - Test table width normalization
  - Test merged cell handling
  - Test multiple table structures

### lesson_with_image.docx
- **Purpose**: Tests image handling (when implemented)
- **Features**: 
  - Image placeholder text
  - Image caption
  - Context: Materials section
- **Use Cases**:
  - Future: Test image extraction
  - Future: Test image preservation
  - Currently: Placeholder for future feature

### lesson_with_hyperlinks.docx
- **Purpose**: Tests hyperlink handling (when implemented)
- **Features**:
  - 2 styled hyperlinks (blue, underlined text)
  - URLs as text
  - Context: Online resources section
- **Use Cases**:
  - Future: Test hyperlink extraction
  - Future: Test hyperlink preservation
  - Currently: Placeholder for future feature

## Notes

- **Images**: Actual image embedding requires image files. Current fixture uses placeholder text.
- **Hyperlinks**: True hyperlinks require complex XML manipulation. Current fixture uses styled text.
- **Created**: 2025-10-18 during validation phase
- **Generator**: `tests/create_test_fixtures.py`

## Regenerating Fixtures

To regenerate all fixtures:

```bash
python tests/create_test_fixtures.py
```

This will overwrite existing fixtures with fresh versions.
"""
    
    readme_path = Path("tests/fixtures/README.md")
    readme_path.parent.mkdir(parents=True, exist_ok=True)
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(readme_content)
    
    print(f"\nCreated: {readme_path}")

if __name__ == '__main__':
    print("=" * 60)
    print("CREATING TEST FIXTURES")
    print("=" * 60)
    print()
    
    fixtures = []
    
    try:
        fixtures.append(create_regular_lesson())
        fixtures.append(create_no_school_day())
        fixtures.append(create_lesson_with_tables())
        fixtures.append(create_lesson_with_image())
        fixtures.append(create_lesson_with_hyperlinks())
        
        create_fixtures_readme(fixtures)
        
        print()
        print("=" * 60)
        print("SUCCESS: All 5 test fixtures created!")
        print("=" * 60)
        print()
        print("Files created:")
        for fixture in fixtures:
            print(f"  - {fixture}")
        print(f"  - tests/fixtures/README.md")
        
        sys.exit(0)
        
    except Exception as e:
        print()
        print("=" * 60)
        print(f"ERROR: {e}")
        print("=" * 60)
        import traceback
        traceback.print_exc()
        sys.exit(1)
