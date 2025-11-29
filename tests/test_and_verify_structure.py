"""
Generate a fresh multi-slot file and verify its structure.
"""

import sys
import asyncio
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.database import get_db
from backend.mock_llm_service import get_mock_llm_service
from tools.batch_processor import BatchProcessor


async def main():
    print("=" * 70)
    print("GENERATE AND VERIFY MULTI-SLOT STRUCTURE")
    print("=" * 70)
    
    db = get_db()
    llm_service = get_mock_llm_service()
    processor = BatchProcessor(llm_service)
    
    # Create test user with 2 slots
    user_id = db.create_user("Structure Test", "test@example.com")
    
    # Create slot 1
    slot1_id = db.create_class_slot(
        user_id=user_id,
        slot_number=1,
        subject="ELA",
        grade="3",
        homeroom="3-1",
        proficiency_levels='["Entering"]',
        primary_teacher_file="input/Lang Lesson Plans 9_15_25-9_19_25.docx"
    )
    db.update_class_slot(slot1_id, primary_teacher_name="Lang")
    
    # Create slot 2
    slot2_id = db.create_class_slot(
        user_id=user_id,
        slot_number=2,
        subject="Math",
        grade="3",
        homeroom="3-1",
        proficiency_levels='["Emerging"]',
        primary_teacher_file="input/9_15-9_19 Davies Lesson Plans.docx"
    )
    db.update_class_slot(slot2_id, primary_teacher_name="Davies")
    
    # Process week
    print("\nGenerating 2-slot consolidated file...")
    result = await processor.process_user_week(
        user_id=user_id,
        week_of="9/15-9/19",
        provider="mock"
    )
    
    print(f"\nGeneration result: {result['success']}")
    print(f"Output file: {result.get('output_file')}")
    
    # Verify structure
    if result.get('output_file'):
        output_path = Path(result['output_file'])
        if output_path.exists():
            print("\n" + "=" * 70)
            print("VERIFYING STRUCTURE")
            print("=" * 70)
            
            doc = Document(output_path)
            
            print(f"\nTotal tables: {len(doc.tables)}")
            
            # Analyze tables
            metadata_tables = 0
            daily_tables = 0
            
            for i, table in enumerate(doc.tables):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                
                if rows == 1 and cols == 5:
                    metadata_tables += 1
                    teacher_name = table.rows[0].cells[0].text
                    print(f"Table {i}: Metadata - {teacher_name[:50]}")
                elif rows >= 7 and cols == 6:
                    # Daily plans table (may have extra row for signature)
                    daily_tables += 1
                    print(f"Table {i}: Daily Plans ({rows}x{cols})")
                else:
                    print(f"Table {i}: Other ({rows}x{cols})")
            
            print(f"\n" + "=" * 70)
            print("RESULTS:")
            print("=" * 70)
            print(f"Metadata tables: {metadata_tables}")
            print(f"Daily plans tables: {daily_tables}")
            print(f"Total slots: {min(metadata_tables, daily_tables)}")
            
            if metadata_tables == 2 and daily_tables == 2:
                print("\n✓ SUCCESS: Document has 2 complete slots as expected!")
            else:
                print(f"\n✗ FAIL: Expected 2 slots, found {min(metadata_tables, daily_tables)}")
        else:
            print(f"\nERROR: File not found at {output_path}")
    
    # Cleanup
    db.delete_user(user_id)


if __name__ == "__main__":
    asyncio.run(main())
