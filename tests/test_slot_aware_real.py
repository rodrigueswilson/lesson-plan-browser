"""
Test slot-aware extraction with real lesson plan files.
Skips when no sample DOCX is provided.
"""

from pathlib import Path

import pytest
from docx import Document

from tools.docx_parser import DOCXParser, validate_slot_structure

# Optional: add paths to real DOCX files; test skips if none exist
SLOT_AWARE_SAMPLE_PATHS = [
    Path("tests/fixtures/sample_lesson_plan.docx"),
]


@pytest.mark.parametrize("file_path", SLOT_AWARE_SAMPLE_PATHS, ids=[p.name for p in SLOT_AWARE_SAMPLE_PATHS])
def test_file(file_path):
    """Test slot-aware extraction on a real file. Skips if file missing."""
    path = Path(file_path)
    if not path.exists():
        pytest.skip(f"No sample DOCX at {path}")
    print(f"\n{'='*80}")
    print(f"Testing: {path.name}")
    print(f"{'='*80}")

    doc = Document(str(path))
    table_count = len(doc.tables)
    print(f"Total tables: {table_count}")
    
    # Check signature table
    if table_count > 0:
        last_table = doc.tables[-1]
        if last_table.rows and last_table.rows[0].cells:
            first_cell = last_table.rows[0].cells[0].text.strip()
            print(f"Last table first cell: '{first_cell[:50]}'")
    
    # Calculate expected slots
    if table_count > 0:
        available_slots = (table_count - 1) // 2
        expected_table_count = (available_slots * 2) + 1
        print(f"Available slots: {available_slots}")
        print(f"Expected table count: {expected_table_count}")
        print(f"Match: {table_count == expected_table_count}")
    
    # Try to validate each slot
    parser = DOCXParser(str(path))
    
    for slot_num in range(1, min(available_slots + 1, 6)):  # Test up to 5 slots
        print(f"\n--- Slot {slot_num} ---")
        try:
            table_start, table_end = validate_slot_structure(doc, slot_num)
            print(f"✅ Validation passed: tables {table_start}-{table_end}")
            
            # Check metadata table
            meta_table = doc.tables[table_start]
            if meta_table.rows and meta_table.rows[0].cells:
                first_cell = meta_table.rows[0].cells[0].text.strip()
                print(f"   Metadata: '{first_cell[:40]}'")
            
            # Check daily plans table
            daily_table = doc.tables[table_end]
            if daily_table.rows and daily_table.rows[0].cells:
                headers = " | ".join(cell.text.strip()[:10] for cell in daily_table.rows[0].cells[:6])
                print(f"   Headers: {headers}")
            
            # Extract hyperlinks
            hyperlinks = parser.extract_hyperlinks_for_slot(slot_num)
            print(f"   Hyperlinks: {len(hyperlinks)}")
            
            # Show table distribution
            table_indices = set(link['table_idx'] for link in hyperlinks)
            print(f"   Table indices: {sorted(table_indices)}")
            
            # Extract images
            images = parser.extract_images_for_slot(slot_num)
            print(f"   Images: {len(images)}")
            
        except ValueError as e:
            print(f"❌ Validation failed: {e}")
        except Exception as e:
            print(f"❌ Error: {type(e).__name__}: {e}")
    
    print(f"\n{'='*80}\n")


def main():
    """Test all available files."""
    input_dir = Path("d:/LP/input")
    
    # Test files with multiple slots
    test_files = [
        "Lang Lesson Plans 9_15_25-9_19_25.docx",
        "Ms. Savoca- 9_15_24-9_19_24 Lesson plans.docx",
        "Piret Lesson Plans 9_22_25-9_26_25.docx",
    ]
    
    for filename in test_files:
        file_path = input_dir / filename
        if file_path.exists():
            try:
                test_file(str(file_path))
            except Exception as e:
                print(f"❌ Failed to test {filename}: {e}\n")
        else:
            print(f"⚠️  File not found: {filename}\n")


if __name__ == "__main__":
    main()
