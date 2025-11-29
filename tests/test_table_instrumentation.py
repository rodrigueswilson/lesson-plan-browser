"""
Test table structure instrumentation without processing through LLM.

This script directly parses DOCX files and logs their table structure,
bypassing the full batch processing pipeline.
"""

import sys
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent))

from tools.docx_parser import DOCXParser


def analyze_file_structure(file_path: str):
    """Analyze and log table structure of a single file."""
    print(f"\n{'='*80}")
    print(f"File: {Path(file_path).name}")
    print(f"{'='*80}")
    
    try:
        # Parse file
        parser = DOCXParser(file_path)
        doc = parser.doc
        
        print(f"Total tables: {len(doc.tables)}")
        print()
        
        # Log each table
        for table_idx, table in enumerate(doc.tables):
            first_cell = ""
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0].text.strip()
            
            # Get first row text
            first_row_text = ""
            if table.rows:
                first_row_text = " | ".join(
                    cell.text.strip()[:30] for cell in table.rows[0].cells[:5]
                )
            
            print(f"Table {table_idx}:")
            print(f"  First cell: {first_cell[:100]}")
            print(f"  First row:  {first_row_text[:150]}")
            print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns) if table.rows else 0}")
            
            # Detect table type
            first_cell_lower = first_cell.lower()
            if "name:" in first_cell_lower or "teacher:" in first_cell_lower:
                print(f"  → METADATA table")
            elif any(day in first_row_text.lower() for day in ["monday", "tuesday", "wednesday"]):
                print(f"  → DAILY PLANS table")
            elif "signature" in first_cell_lower:
                print(f"  → SIGNATURE table")
            
            print()
        
        # Suggest slot mapping
        print(f"{'='*80}")
        print("Suggested Slot Mapping:")
        print(f"{'='*80}")
        
        metadata_tables = []
        for idx, table in enumerate(doc.tables):
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0].text.strip().lower()
                if "name:" in first_cell or "teacher:" in first_cell:
                    metadata_tables.append(idx)
        
        for i, meta_idx in enumerate(metadata_tables, 1):
            daily_idx = meta_idx + 1
            print(f"Slot {i}: Tables [{meta_idx}, {daily_idx}]")
        
        # Check for signature table
        if len(doc.tables) > 0:
            last_table = doc.tables[-1]
            if last_table.rows and last_table.rows[0].cells:
                first_cell = last_table.rows[0].cells[0].text.strip().lower()
                if "signature" in first_cell:
                    print(f"Signature: Table [{len(doc.tables)-1}]")
        
        print()
        
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()


def main():
    """Analyze files from multiple weeks and teachers."""
    # Base directories to scan
    base_dirs = [
        Path("F:/rodri/Documents/OneDrive/AS/Lesson Plan"),
        Path("F:/rodri/Documents/OneDrive/AS/Daniela LP")
    ]
    
    all_files = []
    
    # Find all DOCX files in subdirectories
    for base_dir in base_dirs:
        if not base_dir.exists():
            print(f"WARNING: Directory not found: {base_dir}")
            continue
        
        print(f"Scanning: {base_dir}")
        
        # Find all DOCX files recursively
        for docx_file in base_dir.rglob("*.docx"):
            # Skip temp files and output files
            if docx_file.name.startswith("~$"):
                continue
            if docx_file.name.startswith("Wilson_Rodrigues_Weekly"):
                continue
            if "old" in str(docx_file).lower():
                continue
            
            all_files.append(docx_file)
        
        print(f"  Found {len([f for f in all_files if base_dir in f.parents])} files")
    
    if not all_files:
        print("No DOCX files found")
        return
    
    print(f"\nTotal files to analyze: {len(all_files)}")
    print(f"{'='*80}\n")
    
    # Track patterns
    patterns = {
        "9_tables_4_slots": 0,
        "11_tables_5_slots": 0,
        "3_tables_1_slot": 0,
        "other": 0,
        "has_signature": 0,
        "no_signature": 0,
        "unexpected_structure": []
    }
    
    # Analyze each file
    for idx, docx_file in enumerate(all_files, 1):
        print(f"[{idx}/{len(all_files)}] {docx_file.relative_to(docx_file.parents[3])}")
        
        try:
            analyze_file_structure(str(docx_file))
            
            # Track pattern
            doc = Document(str(docx_file))
            table_count = len(doc.tables)
            
            # Check for signature
            has_sig = False
            if table_count > 0:
                last_table = doc.tables[-1]
                if last_table.rows and last_table.rows[0].cells:
                    first_cell = last_table.rows[0].cells[0].text.strip().lower()
                    if "signature" in first_cell:
                        has_sig = True
                        patterns["has_signature"] += 1
                    else:
                        patterns["no_signature"] += 1
            
            # Categorize
            if table_count == 9 and has_sig:
                patterns["9_tables_4_slots"] += 1
            elif table_count == 11 and has_sig:
                patterns["11_tables_5_slots"] += 1
            elif table_count == 3 and has_sig:
                patterns["3_tables_1_slot"] += 1
            else:
                patterns["other"] += 1
                patterns["unexpected_structure"].append({
                    "file": docx_file.name,
                    "tables": table_count,
                    "has_signature": has_sig
                })
        
        except Exception as e:
            print(f"ERROR: {e}\n")
            patterns["other"] += 1
    
    # Print summary
    print(f"\n{'='*80}")
    print("PATTERN SUMMARY")
    print(f"{'='*80}")
    print(f"Total files analyzed: {len(all_files)}")
    print(f"\nCommon patterns:")
    print(f"  9 tables (4 slots + signature): {patterns['9_tables_4_slots']} files")
    print(f"  11 tables (5 slots + signature): {patterns['11_tables_5_slots']} files")
    print(f"  3 tables (1 slot + signature): {patterns['3_tables_1_slot']} files")
    print(f"  Other structures: {patterns['other']} files")
    print(f"\nSignature table:")
    print(f"  With signature: {patterns['has_signature']} files")
    print(f"  Without signature: {patterns['no_signature']} files")
    
    if patterns["unexpected_structure"]:
        print(f"\nUnexpected structures:")
        for item in patterns["unexpected_structure"][:10]:
            print(f"  {item['file']}: {item['tables']} tables, signature={item['has_signature']}")
    
    print(f"{'='*80}")


if __name__ == "__main__":
    main()
