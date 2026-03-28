"""Smoke test: Grade 3 ELA subject config ingest path (minimal DOCX, isolated DB copy)."""

from __future__ import annotations

import os
import shutil
import sqlite3
import sys
from pathlib import Path

import pytest
from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[1]
_CURRICULUM_DB = _REPO_ROOT / "data" / "curriculum.db"
_SCRAPER_DIR = str(_REPO_ROOT / "tools" / "scraper")

UNIT_ID = "ELA_3_U8_sample"


@pytest.mark.skipif(not _CURRICULUM_DB.is_file(), reason="data/curriculum.db not present")
def test_grade3_ela_minimal_docx_ingests_sections_and_standards(tmp_path: Path) -> None:
    db_copy = tmp_path / "curriculum_test.db"
    shutil.copy(_CURRICULUM_DB, db_copy)

    docx_path = tmp_path / "ela_sample.docx"
    doc = Document()
    doc.add_paragraph("Lesson 1: Sample ELA Lesson Title")
    doc.add_paragraph("Learning Intention")
    doc.add_paragraph("Students will identify key details in informational text.")
    doc.add_paragraph("Success Criteria")
    doc.add_paragraph("I can explain the main idea using evidence from the text.")
    doc.add_paragraph("Standards")
    doc.add_paragraph("NJSLS ELA.R.3.1 Ask and answer questions to demonstrate understanding.")
    doc.save(docx_path)

    conn = sqlite3.connect(db_copy)
    conn.execute(
        """INSERT OR REPLACE INTO units
        (id, grade, subject, unit_number, title, description, source)
        VALUES (?, ?, ?, ?, ?, ?, ?)""",
        (
            UNIT_ID,
            3,
            "ELA",
            8,
            "Grade 3 Unit 8: Reading and Writing (ELA sample)",
            "",
            "test_grade3_ela_ingest_smoke",
        ),
    )
    conn.execute(
        """INSERT OR IGNORE INTO units_intro (unit_id, essential_questions, enduring_understandings)
        VALUES (?, '[]', '[]')""",
        (UNIT_ID,),
    )
    conn.commit()
    conn.close()

    sys.path.insert(0, _SCRAPER_DIR)
    from table_extractor import RecursiveTableParser

    parser = RecursiveTableParser(db_path=str(db_copy))
    n = parser.ingest_to_curriculum(
        str(docx_path),
        UNIT_ID,
        subject="ELA",
        source_url="https://docs.google.com/document/d/test_el_sample/edit",
        parser_version="table_extractor@test",
    )
    assert n >= 1

    conn = sqlite3.connect(db_copy)
    conn.row_factory = sqlite3.Row
    row = conn.execute(
        "SELECT learning_intentions, success_criteria, standards_structured FROM lessons WHERE unit_id = ? AND lesson_number = 1",
        (UNIT_ID,),
    ).fetchone()
    conn.close()
    assert row is not None
    assert row["learning_intentions"] and "key details" in row["learning_intentions"].lower()
    assert row["success_criteria"] and "main idea" in row["success_criteria"].lower()
    assert row["standards_structured"] and row["standards_structured"] != "[]"
