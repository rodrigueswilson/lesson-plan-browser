"""
Verify the structure of a multi-slot consolidated DOCX file.
"""

import sys
from pathlib import Path
from docx import Document

# Get the most recent Weekly file
output_dir = Path(r"F:\rodri\Documents\OneDrive\AS\Lesson Plan\25 W38")
weekly_files = list(output_dir.glob("*Weekly*.docx"))

if not weekly_files:
    print("No Weekly DOCX files found")
    sys.exit(1)

# Get most recent
weekly_file = max(weekly_files, key=lambda p: p.stat().st_mtime)

print("=" * 70)
print(f"VERIFYING: {weekly_file.name}")
print("=" * 70)

doc = Document(weekly_file)

print(f"\nTotal tables: {len(doc.tables)}")
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Count page breaks
page_breaks = 0
for para in doc.paragraphs:
    for run in para.runs:
        if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
            page_breaks += 1

print(f"Page breaks found: {page_breaks}")

# Analyze tables
print("\n" + "=" * 70)
print("TABLE STRUCTURE:")
print("=" * 70)

for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns) if table.rows else 0
    
    # Check if it's a metadata table (1 row, 5 cols)
    if rows == 1 and cols == 5:
        print(f"\nTable {i}: METADATA TABLE ({rows} row x {cols} cols)")
        # Show teacher name
        if table.rows[0].cells[0].text:
            print(f"  Teacher: {table.rows[0].cells[0].text[:50]}")
    # Check if it's a daily plans table (7 rows, 6 cols)
    elif rows == 7 and cols == 6:
        print(f"\nTable {i}: DAILY PLANS TABLE ({rows} rows x {cols} cols)")
    else:
        print(f"\nTable {i}: UNKNOWN ({rows} rows x {cols} cols)")

# Expected structure for N slots:
# - N metadata tables (1 row x 5 cols each)
# - N daily plans tables (7 rows x 6 cols each)
# - Total: 2N tables

metadata_tables = sum(1 for t in doc.tables if len(t.rows) == 1 and len(t.columns) == 5)
daily_tables = sum(1 for t in doc.tables if len(t.rows) == 7 and len(t.columns) == 6)

print("\n" + "=" * 70)
print("SUMMARY:")
print("=" * 70)
print(f"Metadata tables: {metadata_tables}")
print(f"Daily plans tables: {daily_tables}")
print(f"Estimated slots: {min(metadata_tables, daily_tables)}")

if metadata_tables == daily_tables and metadata_tables > 1:
    print(f"\n✓ SUCCESS: Document contains {metadata_tables} complete slots!")
elif metadata_tables == daily_tables == 1:
    print(f"\n⚠ Single slot document")
else:
    print(f"\n✗ WARNING: Mismatch between metadata and daily tables")
