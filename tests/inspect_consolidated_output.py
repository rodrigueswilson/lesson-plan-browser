"""
Inspect the consolidated DOCX output to verify content structure.
"""

import sys
import asyncio
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.database import get_db
from backend.mock_llm_service import get_mock_llm_service
from tools.batch_processor import BatchProcessor


async def main():
    print("=" * 70)
    print("INSPECT CONSOLIDATED OUTPUT")
    print("=" * 70)
    
    db = get_db()
    llm_service = get_mock_llm_service()
    processor = BatchProcessor(llm_service)
    
    # Create test user with 2 slots
    user_id = db.create_user("Inspect User", "inspect@example.com")
    
    # Create slot 1 - ELA
    slot1_id = db.create_class_slot(
        user_id=user_id,
        slot_number=1,
        subject="ELA",
        grade="3",
        homeroom="3-1",
        proficiency_levels='["Entering"]',
        primary_teacher_file="input/Lang Lesson Plans 9_15_25-9_19_25.docx"
    )
    db.update_class_slot(slot1_id, primary_teacher_name="Lang")
    
    # Create slot 2 - Math
    slot2_id = db.create_class_slot(
        user_id=user_id,
        slot_number=2,
        subject="Math",
        grade="3",
        homeroom="3-1",
        proficiency_levels='["Emerging"]',
        primary_teacher_file="input/9_15-9_19 Davies Lesson Plans.docx"
    )
    db.update_class_slot(slot2_id, primary_teacher_name="Davies")
    
    # Process week
    print("\nProcessing week with 2 slots...")
    result = await processor.process_user_week(
        user_id=user_id,
        week_of="9/15-9/19",
        provider="mock"
    )
    
    print(f"\nResult: {result['success']}")
    print(f"Consolidated: {result.get('consolidated')}")
    print(f"Total slots: {result.get('total_slots')}")
    print(f"Output file: {result.get('output_file')}")
    
    # Open and inspect the DOCX
    if result.get('output_file'):
        output_path = Path(result['output_file'])
        if output_path.exists():
            print(f"\n" + "=" * 70)
            print("INSPECTING DOCX CONTENT")
            print("=" * 70)
            
            doc = Document(output_path)
            
            # Check total number of tables
            print(f"\nTotal tables in document: {len(doc.tables)}")
            print(f"Expected for 2 slots: 4 tables (2 metadata + 2 daily plans)")
            
            if len(doc.tables) >= 4:
                print("✓ Document has multiple slot tables!")
            else:
                print("✗ Document may not have separate tables per slot")
            
            # Check first metadata table (Table 0)
            print("\n--- FIRST SLOT METADATA TABLE ---")
            metadata_table = doc.tables[0]
            row = metadata_table.rows[0]
            for i, cell in enumerate(row.cells):
                print(f"Cell {i}: {cell.text[:100]}")
            
            # Check daily plans table (Table 1)
            print("\n--- DAILY PLANS TABLE (Monday column) ---")
            daily_table = doc.tables[1]
            
            # Row labels
            row_labels = [
                "Unit/Lesson",
                "Objective",
                "Anticipatory Set",
                "Tailored Instruction",
                "Misconceptions",
                "Assessment",
                "Homework"
            ]
            
            # Monday is column 1 (column 0 is row labels)
            for i, label in enumerate(row_labels, 1):
                cell_text = daily_table.rows[i].cells[1].text
                print(f"\n{label}:")
                # Show first 200 chars
                preview = cell_text[:200].replace('\n', ' | ')
                print(f"  {preview}...")
                if "Slot 1:" in cell_text and "Slot 2:" in cell_text:
                    print(f"  ✓ Contains both slots")
                elif "Slot 1:" in cell_text or "Slot 2:" in cell_text:
                    print(f"  ⚠ Contains only one slot")
                else:
                    print(f"  ✗ No slot headers found")
        else:
            print(f"\nERROR: Output file not found at {output_path}")
    
    # Cleanup
    db.delete_user(user_id)
    
    print("\n" + "=" * 70)
    print("INSPECTION COMPLETE")
    print("=" * 70)


if __name__ == "__main__":
    asyncio.run(main())
