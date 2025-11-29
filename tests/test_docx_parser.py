"""
Tests for DOCX parser functionality.
"""

import pytest
import sys
from pathlib import Path
from docx import Document

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_parser import DOCXParser, parse_docx


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file for testing."""
    doc = Document()
    
    # Add title
    doc.add_heading('Weekly Lesson Plan', 0)
    doc.add_paragraph('Week of: 10/07-10/11')
    
    # Add Math section
    doc.add_heading('Math', 1)
    doc.add_paragraph('Objectives: Students will be able to solve linear equations.')
    doc.add_paragraph('Activities: Practice problems, group work, exit ticket.')
    doc.add_paragraph('Assessment: Exit ticket on solving equations.')
    doc.add_paragraph('Materials: Whiteboard, worksheets, calculators.')
    
    # Add Science section
    doc.add_heading('Science', 1)
    doc.add_paragraph('Objectives: Students will understand the water cycle.')
    doc.add_paragraph('Activities: Watch video, create diagram, lab experiment.')
    doc.add_paragraph('Assessment: Diagram completion and lab report.')
    
    # Save to temp file
    file_path = tmp_path / "sample_lesson.docx"
    doc.save(str(file_path))
    
    return str(file_path)


def test_parser_initialization(sample_docx):
    """Test parser initialization."""
    parser = DOCXParser(sample_docx)
    
    assert parser.file_path.exists()
    assert parser.doc is not None
    assert len(parser.paragraphs) > 0


def test_get_full_text(sample_docx):
    """Test getting full text from document."""
    parser = DOCXParser(sample_docx)
    full_text = parser.get_full_text()
    
    assert 'Weekly Lesson Plan' in full_text
    assert 'Math' in full_text
    assert 'Science' in full_text


def test_find_subject_sections(sample_docx):
    """Test finding subject sections."""
    parser = DOCXParser(sample_docx)
    sections = parser.find_subject_sections()
    
    assert 'math' in sections
    assert 'science' in sections
    assert len(sections['math']) > 0
    assert len(sections['science']) > 0


def test_extract_by_heading(sample_docx):
    """Test extracting content by heading."""
    parser = DOCXParser(sample_docx)
    math_content = parser.extract_by_heading('Math')
    
    assert len(math_content) > 0
    assert any('linear equations' in para for para in math_content)


def test_extract_subject_content(sample_docx):
    """Test extracting structured subject content."""
    parser = DOCXParser(sample_docx)
    content = parser.extract_subject_content('Math')
    
    assert content['subject'] == 'Math'
    assert 'full_text' in content
    assert 'objectives' in content
    assert 'activities' in content
    assert 'assessments' in content
    assert 'materials' in content


def test_extract_objectives(sample_docx):
    """Test extracting objectives."""
    parser = DOCXParser(sample_docx)
    content = parser.extract_subject_content('Math')
    
    assert len(content['objectives']) > 0
    assert any('linear equations' in obj for obj in content['objectives'])


def test_extract_activities(sample_docx):
    """Test extracting activities."""
    parser = DOCXParser(sample_docx)
    content = parser.extract_subject_content('Science')
    
    assert len(content['activities']) > 0
    assert any('video' in act or 'diagram' in act for act in content['activities'])


def test_extract_assessments(sample_docx):
    """Test extracting assessments."""
    parser = DOCXParser(sample_docx)
    content = parser.extract_subject_content('Math')
    
    assert len(content['assessments']) > 0
    assert any('Exit ticket' in assess for assess in content['assessments'])


def test_extract_materials(sample_docx):
    """Test extracting materials."""
    parser = DOCXParser(sample_docx)
    content = parser.extract_subject_content('Math')
    
    assert len(content['materials']) > 0
    assert any('worksheets' in mat or 'calculators' in mat for mat in content['materials'])


def test_list_available_subjects(sample_docx):
    """Test listing available subjects."""
    parser = DOCXParser(sample_docx)
    subjects = parser.list_available_subjects()
    
    assert 'math' in subjects
    assert 'science' in subjects


def test_extract_week_info(sample_docx):
    """Test extracting week information."""
    parser = DOCXParser(sample_docx)
    week_info = parser.extract_week_info()
    
    assert week_info is not None
    assert '10/07' in week_info or '10/11' in week_info


def test_get_metadata(sample_docx):
    """Test getting document metadata."""
    parser = DOCXParser(sample_docx)
    metadata = parser.get_metadata()
    
    assert 'paragraph_count' in metadata
    assert 'table_count' in metadata
    assert 'available_subjects' in metadata
    assert metadata['paragraph_count'] > 0


def test_parse_docx_convenience_function(sample_docx):
    """Test the convenience function."""
    result = parse_docx(sample_docx, subject='Math')
    
    assert result['subject'] == 'Math'
    assert 'full_text' in result
    
    # Test without subject
    result_all = parse_docx(sample_docx)
    assert 'full_text' in result_all
    assert 'metadata' in result_all
    assert 'available_subjects' in result_all


def test_file_not_found():
    """Test handling of non-existent file."""
    with pytest.raises(FileNotFoundError):
        DOCXParser("nonexistent_file.docx")


def test_table_extraction(tmp_path):
    """Test extracting tables from document."""
    doc = Document()
    
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Day"
    table.cell(0, 1).text = "Activity"
    table.cell(1, 0).text = "Monday"
    table.cell(1, 1).text = "Introduction"
    table.cell(2, 0).text = "Tuesday"
    table.cell(2, 1).text = "Practice"
    
    file_path = tmp_path / "table_test.docx"
    doc.save(str(file_path))
    
    parser = DOCXParser(str(file_path))
    
    assert len(parser.tables) == 1
    assert parser.tables[0][0][0] == "Day"
    assert parser.tables[0][1][1] == "Introduction"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
