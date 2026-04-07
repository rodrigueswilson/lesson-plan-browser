"""Unit tests for ELA per-lesson detailed plan table detection and parsing."""

from __future__ import annotations

import json
import sqlite3
import sys
from pathlib import Path

import pytest
from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[1]
_SCRAPER_DIR = str(_REPO_ROOT / "tools" / "scraper")
if _SCRAPER_DIR not in sys.path:
    sys.path.insert(0, _SCRAPER_DIR)
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from tools.scraper.ela_lesson_plan_table import (  # noqa: E402
    is_ela_lesson_plan_table,
    parse_ela_lesson_plan_table,
)


def _p(text: str, bold: bool = False) -> dict:
    style = {"bold": True} if bold else None
    return {
        "type": "paragraph",
        "text": text,
        "is_bullet": False,
        "ilvl": 0,
        "runs": [{"type": "run", "text": text, "style": style}],
    }


def _cell(*paragraphs: dict) -> dict:
    return {"type": "table-cell", "grid_span": None, "v_merge": None, "content": list(paragraphs)}


def _row(*cells: dict) -> dict:
    return {"type": "table-row", "cells": list(cells)}


def _table(*rows: dict) -> dict:
    return {"type": "table", "rows": list(rows), "depth": 0}


def _htmlize(content: list) -> str:
    parts = []
    for e in content:
        if e.get("type") == "paragraph":
            inner = ""
            for run in e.get("runs") or []:
                t = run.get("text") or ""
                if (run.get("style") or {}).get("bold"):
                    t = f"<b>{t}</b>"
                inner += t
            if inner:
                parts.append(f"<p>{inner}</p>")
    return "".join(parts)


def test_is_lesson_plan_detects_typical_grid() -> None:
    tbl = _table(
        _row(_cell(_p("Lesson 1: The U.S. Constitution")), _cell(_p(""))),
        _row(_cell(_p("Learning Intention")), _cell(_p("Success Criteria"))),
        _row(_cell(_p("I am learning to read.")), _cell(_p("I can read."))),
        _row(_cell(_p("NJSLS Standards")), _cell(_p(""))),
        _row(_cell(_p("Priority Standards: RI.CR.2.1")), _cell(_p(""))),
    )
    assert is_ela_lesson_plan_table(tbl) is True


def test_is_lesson_plan_detects_grid_with_blank_row_after_title() -> None:
    """Grade 2 Unit 8 tab export: spacer row between Lesson N title and LI|SC headers."""
    tbl = _table(
        _row(_cell(_p("Lesson 1: The U.S. Constitution")), _cell(_p(""))),
        _row(_cell(_p("")), _cell(_p(""))),
        _row(_cell(_p("Learning Intention")), _cell(_p("Success Criteria"))),
        _row(_cell(_p("I am learning to define key words.")), _cell(_p("I can define X."))),
        _row(_cell(_p("NJSLS Standards")), _cell(_p(""))),
        _row(_cell(_p("Priority Standards: RI.CR.2.1")), _cell(_p(""))),
    )
    assert is_ela_lesson_plan_table(tbl) is True
    out = parse_ela_lesson_plan_table(tbl, _htmlize)
    assert out is not None
    assert out["lesson_number"] == 1
    assert "define key words" in out["learning_intention_html"]
    assert "I can define" in out["success_criteria_html"]


def test_is_lesson_plan_rejects_summary_shape() -> None:
    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Learning Intentions and Success Criteria")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
        _row(_cell(_p("1")), _cell(_p("a")), _cell(_p("b")), _cell(_p("c"))),
    )
    assert is_ela_lesson_plan_table(tbl) is False


def test_parse_lesson1_minimal_bands() -> None:
    proc_cell = _cell(
        _p("Anticipatory Set: Unit Introduction"),
        _p("Open the lesson."),
        _p("Learning Procedures:"),
        _p("Teach the skill."),
        _p("Engagement with the Content"),
        _p("Students read."),
        _p("Daily Instructional Task:"),
        _p("Partner discussion."),
    )
    tbl = _table(
        _row(_cell(_p("Lesson 1: The U.S. Constitution")), _cell(_p(""))),
        _row(_cell(_p("Learning Intention")), _cell(_p("Success Criteria"))),
        _row(_cell(_p("I am learning to define terms.")), _cell(_p("I can define X."))),
        _row(_cell(_p("NJSLS Standards")), _cell(_p(""))),
        _row(_cell(_p("Priority Standards: RI.CR.2.1")), _cell(_p(""))),
        _row(_cell(_p("Key Instructional Practices")), _cell(_p(""))),
        _row(
            _cell(_p("Key Questions Posed During the Lesson: What?")),
            _cell(_p("Instructional Routines and Assessments: Preview SL.PE.2.1")),
        ),
        _row(
            _cell(_p("Vocabulary: word1")),
            _cell(_p("Instructional Resources: Doc A")),
        ),
        _row(proc_cell, _cell(_p(""))),
        _row(
            _cell(_p("Differentiation: scaffold.")),
            _cell(_p("Addressing Misconceptions: clarify.")),
        ),
    )
    out = parse_ela_lesson_plan_table(tbl, _htmlize)
    assert out is not None
    assert out["lesson_number"] == 1
    assert "Constitution" in out["lesson_title"]
    assert "define terms" in out["learning_intention_html"]
    assert "I can define" in out["success_criteria_html"]
    assert "RI.CR.2.1" in out["njsls_standards_html"]
    assert "What?" in out["key_questions_html"]
    assert "Preview SL.PE.2.1" in out["instructional_routines_assessments_html"]
    assert "word1" in out["vocabulary_cell_html"]
    assert "Doc A" in out["instructional_resources_cell_html"]
    assert "Unit Introduction" in (out.get("anticipatory_set_html") or "")
    assert "Teach the skill" in (out.get("learning_procedures_html") or "")
    assert "Students read" in (out.get("engagement_with_content_html") or "")
    assert "Partner discussion" in (out.get("daily_instructional_task_html") or "")
    assert "scaffold" in out["differentiation_html"]
    assert "clarify" in out["addressing_misconceptions_html"]
    assert out["schema_version"] == 1


def test_parse_priority_standards_without_njsls_label_row() -> None:
    """Empty title row then merged Priority Standards cell (Grade 3 Unit 8 lesson 1 pattern)."""
    tbl = _table(
        _row(_cell(_p("Lesson 1: Sample Title")), _cell(_p(""))),
        _row(_cell(_p("Learning Intention")), _cell(_p("Success Criteria"))),
        _row(_cell(_p("LI here.")), _cell(_p("SC here."))),
        _row(_cell(_p("")), _cell(_p(""))),
        _row(_cell(_p("Priority Standards: RI.ZZ.9.9")), _cell(_p(""))),
        _row(_cell(_p("Key Instructional Practices")), _cell(_p(""))),
        _row(
            _cell(_p("Key Questions: Q?")),
            _cell(_p("Instructional Routines: R1")),
        ),
        _row(
            _cell(_p("Vocabulary: v1")),
            _cell(_p("Instructional Resources: res1")),
        ),
    )
    out = parse_ela_lesson_plan_table(tbl, _htmlize)
    assert out is not None
    assert "RI.ZZ.9.9" in (out.get("njsls_standards_html") or "")
    assert "Q?" in out["key_questions_html"]


def test_parse_njsls_priority_standards_label_row_merges_following_body() -> None:
    tbl = _table(
        _row(_cell(_p("Lesson 6: Freedom")), _cell(_p(""))),
        _row(_cell(_p("Learning Intention")), _cell(_p("Success Criteria"))),
        _row(_cell(_p("LI body")), _cell(_p("SC body"))),
        _row(_cell(_p("NJSLS Priority Standards")), _cell(_p(""))),
        _row(_cell(_p("RI.CR.3.1")), _cell(_p(""))),
        _row(_cell(_p("Key Instructional Practices")), _cell(_p(""))),
        _row(_cell(_p("Key Questions: Q?")), _cell(_p("Instructional Routines and Assessments: R?"))),
    )
    out = parse_ela_lesson_plan_table(tbl, _htmlize)
    assert out is not None
    assert "RI.CR.3.1" in out["njsls_standards_html"]
    assert "NJSLS Priority Standards" not in out["njsls_standards_html"]


def test_parser_normalizes_heading_prefix_and_star_bullets() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    from tools.scraper.ela_lesson_plan_table import _normalize_docx_export_html

    raw = (
        "<p>"
        "Anticipatory Set: Unit Introduction<br>"
        "* First item<br>"
        "* Second item"
        "</p>"
    )
    got = _normalize_docx_export_html(raw)
    assert "<b>Anticipatory Set:</b> Unit Introduction" in got
    assert "<ul>" in got and "</ul>" in got
    assert "<li>First item</li>" in got
    assert "<li>Second item</li>" in got


def test_parser_normalizes_heading_prefix_and_linebreak_bullets() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    from tools.scraper.ela_lesson_plan_table import _normalize_docx_export_html

    raw = (
        "<p>"
        "Anticipatory Set: Unit Introduction<br>"
        "Introduce essential question.<br>"
        "Invite partner discussion.<br>"
        "Learning Procedures: Plan<br>"
        "Model vocabulary.<br>"
        "Read aloud pages."
        "</p>"
    )
    got = _normalize_docx_export_html(raw)
    assert "<b>Anticipatory Set:</b> Unit Introduction" in got
    assert "<b>Learning Procedures:</b> Plan" in got
    assert got.count("<ul>") >= 2
    assert "<li>Introduce essential question.</li>" in got
    assert "<li>Invite partner discussion.</li>" in got
    assert "<li>Model vocabulary.</li>" in got


def test_parser_splits_read_aloud_heading_and_standards_codes() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    from tools.scraper.ela_lesson_plan_table import _normalize_docx_export_html

    raw = (
        "<p>"
        "Read Aloud and Text Dependent Questions RI.IT.2.3, RI.CR.2.1, RI.CI.2.2, "
        "SL.PE.2.1, 6.1.2.HistoryCC.3, 6.1.2.HistoryUP.1, 6.1.2.CivicsPI.6:"
        "</p>"
    )
    got = _normalize_docx_export_html(raw)
    assert "<b>Read Aloud and Text Dependent Questions:</b>" in got
    assert "<li>RI.IT.2.3, RI.CR.2.1, RI.CI.2.2, SL.PE.2.1, 6.1.2.HistoryCC.3, 6.1.2.HistoryUP.1, 6.1.2.CivicsPI.6:</li>" in got
    # Standards line should remain normal text inside li, not bolded.
    assert "<li><b>RI.IT.2.3" not in got


def test_parser_splits_daily_instructional_task_heading_and_standards_codes() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    from tools.scraper.ela_lesson_plan_table import _normalize_docx_export_html

    raw = (
        "<p>"
        "<b>Daily Instructional Task:</b> Partner Discussions RI.CR.2.1, RI.IT.2.3, "
        "SL.PE.2.1A-C, SL.II.2.2, SL.ES.2.3, 6.1.2.HistoryUP.1:"
        "</p>"
    )
    got = _normalize_docx_export_html(raw)
    assert "<b>Daily Instructional Task:</b>" in got
    assert (
        "<li>Partner Discussions RI.CR.2.1, RI.IT.2.3, SL.PE.2.1A-C, "
        "SL.II.2.2, SL.ES.2.3, 6.1.2.HistoryUP.1:</li>"
    ) in got
    assert "<li><b>Partner Discussions" not in got


def test_collects_all_cells_links_for_resource_persistence() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    import table_extractor as te

    plan = {
        "instructional_resources_cell_html": (
            '<p><a href="https://docs.google.com/document/d/AAA/edit">Resource Doc</a></p>'
        ),
        "key_questions_html": (
            '<p>See <a href="https://example.org/q1">Question Guide</a></p>'
        ),
        "learning_procedures_html": (
            '<p>Warm-up video: <a href="https://watch.screencastify.com/v/VID">Video</a></p>'
        ),
    }
    links = te._collect_links_from_ela_structured_plan(plan)
    urls = {x["url"] for x in links}
    assert "https://docs.google.com/document/d/AAA/edit" in urls
    assert "https://example.org/q1" in urls
    assert "https://watch.screencastify.com/v/VID" in urls

    lesson_data = {"_resources": []}
    te._merge_links_into_resources(lesson_data, links)
    assert len(lesson_data["_resources"]) == 3
    doc_res = next(r for r in lesson_data["_resources"] if "docs.google.com/document/d/AAA" in r["url"])
    assert doc_res.get("google_id") == "AAA"
    # Dedupe behavior: re-merging same links should not duplicate.
    te._merge_links_into_resources(lesson_data, links)
    assert len(lesson_data["_resources"]) == 3


def test_collects_links_from_structured_json_string() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    import table_extractor as te

    raw = json.dumps(
        {
            "instructional_resources_cell_html": (
                '<p><a href="https://docs.google.com/document/d/ABC/edit">Doc A</a></p>'
            ),
            "key_questions_html": (
                '<p><a href="https://example.org/guide">Guide</a></p>'
            ),
        }
    )
    links = te._collect_links_from_ela_structured_json(raw)
    urls = {x["url"] for x in links}
    assert "https://docs.google.com/document/d/ABC/edit" in urls
    assert "https://example.org/guide" in urls


def test_collects_nested_anchor_inner_href_safety_net() -> None:
    sys.path.insert(0, _SCRAPER_DIR)
    import table_extractor as te

    plan = {
        "instructional_resources_cell_html": (
            '<p><a href="https://docs.google.com/document/d/OUTER/edit">'
            '<a href="https://docs.google.com/document/d/INNER/edit">Nested</a>'
            "</a></p>"
        )
    }
    links = te._collect_links_from_ela_structured_plan(plan)
    urls = {x["url"] for x in links}
    assert "https://docs.google.com/document/d/OUTER/edit" in urls
    assert "https://docs.google.com/document/d/INNER/edit" in urls


@pytest.mark.skipif(not (_REPO_ROOT / "data" / "curriculum.db").is_file(), reason="data/curriculum.db not present")
def test_ela_ingest_persists_ela_lesson_plan_structured(tmp_path: Path) -> None:
    db_copy = tmp_path / "curriculum_plan_test.db"
    db_copy.write_bytes((_REPO_ROOT / "data" / "curriculum.db").read_bytes())

    docx_path = tmp_path / "ela_plan_unit.docx"
    doc = Document()
    t0 = doc.add_table(rows=3, cols=4)
    t0.rows[0].cells[0].text = "Summary of Key Learning"
    for i, h in enumerate(
        [
            "Lesson",
            "Learning Intentions and Success Criteria",
            "Daily Instructional Task",
            "Content and Learning Strategies",
        ]
    ):
        t0.rows[1].cells[i].text = h
    r2 = t0.rows[2].cells
    r2[0].text = "1"
    r2[1].text = "Learning Intention:\nX.\nSuccess Criteria:\nY."
    r2[2].text = "Task"
    r2[3].text = "SL.1.1"

    t1 = doc.add_table(rows=10, cols=2)
    t1.rows[0].cells[0].text = "Lesson 1: Plan Ingest Check"
    t1.rows[1].cells[0].text = "Learning Intention"
    t1.rows[1].cells[1].text = "Success Criteria"
    t1.rows[2].cells[0].text = "LI body"
    t1.rows[2].cells[1].text = "SC body"
    t1.rows[3].cells[0].text = "NJSLS Standards"
    t1.rows[4].cells[0].text = "Priority Standards: RI.ZZ.1.1"
    t1.rows[5].cells[0].text = "Key Instructional Practices"
    t1.rows[6].cells[0].text = "Key Questions Posed During the Lesson: Q1"
    t1.rows[6].cells[1].text = "Instructional Routines and Assessments: R1"
    t1.rows[7].cells[0].text = "Vocabulary: v1"
    t1.rows[7].cells[1].text = "Instructional Resources: res1"
    c8 = t1.rows[8].cells[0]
    c8.text = "Anticipatory Set: start"
    c8.add_paragraph("Daily Instructional Task:")
    c8.add_paragraph("finish")
    t1.rows[9].cells[0].text = "Differentiation: diff"
    t1.rows[9].cells[1].text = "Addressing Misconceptions: misc"

    doc.save(docx_path)

    unit_id = "ELA_PLAN_TEST_UNIT"
    conn = sqlite3.connect(db_copy)
    conn.execute(
        """INSERT OR REPLACE INTO units
        (id, grade, subject, unit_number, title, description, source)
        VALUES (?, 2, 'ELA', 8, 'ELA plan test', '', 'test')""",
        (unit_id,),
    )
    conn.execute(
        "INSERT OR IGNORE INTO units_intro (unit_id, essential_questions, enduring_understandings) VALUES (?, '[]', '[]')",
        (unit_id,),
    )
    conn.commit()
    conn.close()

    sys.path.insert(0, _SCRAPER_DIR)
    from table_extractor import RecursiveTableParser

    parser = RecursiveTableParser(db_path=str(db_copy))
    n = parser.ingest_to_curriculum(
        str(docx_path),
        unit_id,
        subject="ELA",
        source_url="https://docs.google.com/document/d/test_plan_table/edit",
        parser_version="table_extractor@test_plan",
    )
    assert n >= 1

    conn = sqlite3.connect(db_copy)
    row = conn.execute(
        "SELECT ela_lesson_plan_structured FROM lessons WHERE unit_id = ? AND lesson_number = 1",
        (unit_id,),
    ).fetchone()
    conn.close()
    assert row is not None and row[0]
    data = json.loads(row[0])
    assert data.get("schema_version") == 1
    assert "LI body" in data.get("learning_intention_html", "")
    assert "RI.ZZ.1.1" in data.get("njsls_standards_html", "")
    assert "finish" in data.get("daily_instructional_task_html", "")
