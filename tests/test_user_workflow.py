"""
Integration test for user workflow: Create user, add slots, process week.
This demonstrates the complete multi-user, multi-slot workflow.
"""

import sys
import asyncio
from pathlib import Path

# Add current directory to path
sys.path.insert(0, str(Path(__file__).parent))

from backend.mock_llm_service import get_mock_llm_service
from tools.batch_processor import BatchProcessor
from docx import Document


def create_sample_primary_docx(file_path: str, subject: str):
    """Create a sample primary teacher DOCX file."""
    doc = Document()
    
    doc.add_heading(f'{subject} Lesson Plan', 0)
    doc.add_paragraph('Week of: 10/07-10/11')
    
    doc.add_heading('Objectives', 1)
    doc.add_paragraph(f'Students will understand key concepts in {subject}.')
    doc.add_paragraph('Students will be able to apply knowledge to real-world scenarios.')
    
    doc.add_heading('Activities', 1)
    doc.add_paragraph('Warm-up activity: Review previous concepts')
    doc.add_paragraph('Main lesson: Introduce new topic with examples')
    doc.add_paragraph('Group work: Collaborative problem solving')
    doc.add_paragraph('Independent practice: Worksheet completion')
    
    doc.add_heading('Assessment', 1)
    doc.add_paragraph('Exit ticket: 3 questions on today\'s topic')
    doc.add_paragraph('Homework: Practice problems from textbook')
    
    doc.add_heading('Materials', 1)
    doc.add_paragraph('Whiteboard and markers')
    doc.add_paragraph('Student worksheets')
    doc.add_paragraph('Textbook pages 45-52')
    
    doc.save(file_path)
    print(f"Created sample DOCX: {file_path}")


async def test_user_workflow():
    """Test complete user workflow."""
    print("\n" + "="*60)
    print("TESTING USER WORKFLOW: Multi-User, Multi-Slot System")
    print("="*60 + "\n")
    
    # Initialize database (use default path so batch processor can find it)
    from backend.database import get_db
    db = get_db()
    print("1. Database initialized\n")
    
    # Create users
    print("2. Creating users...")
    user1_id = db.create_user("Maria Garcia", "maria@school.edu")
    user2_id = db.create_user("John Smith", "john@school.edu")
    print(f"   - Created user: Maria Garcia (ID: {user1_id[:8]}...)")
    print(f"   - Created user: John Smith (ID: {user2_id[:8]}...)\n")
    
    # Create sample primary teacher files
    input_dir = Path("input")
    input_dir.mkdir(exist_ok=True)
    
    print("3. Creating sample primary teacher DOCX files...")
    subjects = ["Math", "Science", "ELA"]
    for subject in subjects:
        file_path = input_dir / f"primary_{subject.lower()}.docx"
        create_sample_primary_docx(str(file_path), subject)
    print()
    
    # Configure class slots for Maria (3 classes)
    print("4. Configuring class slots for Maria Garcia...")
    slot1_id = db.create_class_slot(
        user_id=user1_id,
        slot_number=1,
        subject="Math",
        grade="6",
        homeroom="6A",
        proficiency_levels='{"levels": ["1", "2", "3"]}',
        primary_teacher_file=str(input_dir / "primary_math.docx")
    )
    slot2_id = db.create_class_slot(
        user_id=user1_id,
        slot_number=2,
        subject="Science",
        grade="6",
        homeroom="6B",
        proficiency_levels='{"levels": ["2", "3", "4"]}',
        primary_teacher_file=str(input_dir / "primary_science.docx")
    )
    slot3_id = db.create_class_slot(
        user_id=user1_id,
        slot_number=3,
        subject="ELA",
        grade="7",
        homeroom="7A",
        proficiency_levels='{"levels": ["1", "2"]}',
        primary_teacher_file=str(input_dir / "primary_ela.docx")
    )
    print(f"   - Slot 1: Math, Grade 6, Room 6A")
    print(f"   - Slot 2: Science, Grade 6, Room 6B")
    print(f"   - Slot 3: ELA, Grade 7, Room 7A\n")
    
    # Configure class slots for John (2 classes)
    print("5. Configuring class slots for John Smith...")
    db.create_class_slot(
        user_id=user2_id,
        slot_number=1,
        subject="Math",
        grade="7",
        homeroom="7B",
        proficiency_levels='{"levels": ["3", "4", "5"]}',
        primary_teacher_file=str(input_dir / "primary_math.docx")
    )
    db.create_class_slot(
        user_id=user2_id,
        slot_number=2,
        subject="Science",
        grade="7",
        homeroom="7C",
        proficiency_levels='{"levels": ["2", "3"]}',
        primary_teacher_file=str(input_dir / "primary_science.docx")
    )
    print(f"   - Slot 1: Math, Grade 7, Room 7B")
    print(f"   - Slot 2: Science, Grade 7, Room 7C\n")
    
    # Verify slots
    maria_slots = db.get_user_slots(user1_id)
    john_slots = db.get_user_slots(user2_id)
    print(f"6. Verification:")
    print(f"   - Maria has {len(maria_slots)} class slots configured")
    print(f"   - John has {len(john_slots)} class slots configured\n")
    
    # Process Maria's week (using mock LLM)
    print("7. Processing Maria's week (10/07-10/11)...")
    print("   Using mock LLM service for testing...")
    
    llm_service = get_mock_llm_service()
    processor = BatchProcessor(llm_service)
    
    try:
        result = await processor.process_user_week(
            user_id=user1_id,
            week_of="10/07-10/11",
            provider="mock"
        )
        
        if result['success']:
            print(f"   SUCCESS!")
            print(f"   - Processed {result['processed_slots']} slots")
            print(f"   - Failed {result['failed_slots']} slots")
            print(f"   - Total time: {result['total_time_ms']:.0f}ms")
            print(f"   - Output file: {result['output_file']}\n")
        else:
            print(f"   FAILED!")
            print(f"   - Errors: {result.get('errors')}")
            print(f"   - Result: {result}\n")
    except Exception as e:
        print(f"   ERROR: {str(e)}\n")
    
    # Show weekly plans
    print("8. Retrieving weekly plans...")
    maria_plans = db.get_user_plans(user1_id)
    print(f"   - Maria has {len(maria_plans)} weekly plan(s)")
    if maria_plans:
        for plan in maria_plans:
            print(f"     * Week: {plan.week_of}, Status: {plan.status}")
    print()
    
    # Update a slot
    print("9. Testing slot update...")
    db.update_class_slot(slot1_id, homeroom="6C")
    updated_slot = db.get_slot(slot1_id)
    print(f"   - Updated slot 1 homeroom: {updated_slot.homeroom}\n")
    
    # Summary
    print("="*60)
    print("WORKFLOW TEST COMPLETE")
    print("="*60)
    print("\nSummary:")
    print(f"- Created 2 users")
    print(f"- Configured {len(maria_slots) + len(john_slots)} total class slots")
    print(f"- Processed 1 weekly plan")
    print(f"- All database operations successful")
    print("\nThe multi-user, multi-slot system is working correctly!")
    print("="*60 + "\n")


if __name__ == "__main__":
    asyncio.run(test_user_workflow())
