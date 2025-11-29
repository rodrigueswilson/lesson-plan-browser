"""
Test slot-aware content extraction to prevent cross-slot contamination.

Validates that extract_subject_content_for_slot() only extracts from
the specified slot's tables, not from other slots with the same subject.
"""

import pytest
from pathlib import Path
from docx import Document
from tools.docx_parser import DOCXParser
import io
from PIL import Image as PILImage


def create_test_doc_with_duplicate_subjects(tmp_path):
    """Create a document with 2 Math slots with different content."""
    doc_path = tmp_path / "test_duplicate_subjects.docx"
    doc = Document()
    
    # Slot 1: Math (Grade 3)
    meta1 = doc.add_table(rows=3, cols=2)
    meta1.rows[0].cells[0].text = "Name: Teacher A"
    meta1.rows[1].cells[0].text = "Subject: Math"
    meta1.rows[1].cells[1].text = "Grade 3"
    
    daily1 = doc.add_table(rows=3, cols=6)
    daily1.rows[0].cells[0].text = "Day"
    daily1.rows[0].cells[1].text = "Monday"
    daily1.rows[0].cells[2].text = "Tuesday"
    daily1.rows[1].cells[0].text = "Objective"
    daily1.rows[1].cells[1].text = "SLOT1_MONDAY_OBJECTIVE"
    daily1.rows[1].cells[2].text = "SLOT1_TUESDAY_OBJECTIVE"
    daily1.rows[2].cells[0].text = "Anticipatory Set"
    daily1.rows[2].cells[1].text = "SLOT1_MONDAY_ANTICIPATORY"
    daily1.rows[2].cells[2].text = "SLOT1_TUESDAY_ANTICIPATORY"
    
    # Slot 2: Math (Grade 4) - DIFFERENT CONTENT
    meta2 = doc.add_table(rows=3, cols=2)
    meta2.rows[0].cells[0].text = "Name: Teacher B"
    meta2.rows[1].cells[0].text = "Subject: Math"
    meta2.rows[1].cells[1].text = "Grade 4"
    
    daily2 = doc.add_table(rows=3, cols=6)
    daily2.rows[0].cells[0].text = "Day"
    daily2.rows[0].cells[1].text = "Monday"
    daily2.rows[0].cells[2].text = "Tuesday"
    daily2.rows[1].cells[0].text = "Objective"
    daily2.rows[1].cells[1].text = "SLOT2_MONDAY_OBJECTIVE"
    daily2.rows[1].cells[2].text = "SLOT2_TUESDAY_OBJECTIVE"
    daily2.rows[2].cells[0].text = "Anticipatory Set"
    daily2.rows[2].cells[1].text = "SLOT2_MONDAY_ANTICIPATORY"
    daily2.rows[2].cells[2].text = "SLOT2_TUESDAY_ANTICIPATORY"
    
    # Signature table
    sig = doc.add_table(rows=2, cols=2)
    sig.rows[0].cells[0].text = "Signature"
    
    doc.save(str(doc_path))
    return doc_path


class TestSlotAwareContentExtraction:
    """Test slot-aware content extraction prevents cross-slot contamination."""
    
    def test_extract_content_for_slot_isolates_slots(self, tmp_path):
        """Test that each slot only gets its own content, not other slots'."""
        doc_path = create_test_doc_with_duplicate_subjects(tmp_path)
        parser = DOCXParser(str(doc_path))
        
        # Extract content for Slot 1
        slot1_content = parser.extract_subject_content_for_slot(1, "Math")
        
        # Should have Slot 1's content
        assert slot1_content['found'] is True
        assert slot1_content['slot_number'] == 1
        assert slot1_content['table_idx'] == 1  # Slot 1's daily table
        
        full_text1 = slot1_content['full_text']
        assert "SLOT1_MONDAY_OBJECTIVE" in full_text1
        assert "SLOT1_TUESDAY_OBJECTIVE" in full_text1
        assert "SLOT1_MONDAY_ANTICIPATORY" in full_text1
        
        # Should NOT have Slot 2's content
        assert "SLOT2_MONDAY_OBJECTIVE" not in full_text1
        assert "SLOT2_TUESDAY_OBJECTIVE" not in full_text1
        assert "SLOT2_MONDAY_ANTICIPATORY" not in full_text1
        
        # Extract content for Slot 2
        slot2_content = parser.extract_subject_content_for_slot(2, "Math")
        
        # Should have Slot 2's content
        assert slot2_content['found'] is True
        assert slot2_content['slot_number'] == 2
        assert slot2_content['table_idx'] == 3  # Slot 2's daily table
        
        full_text2 = slot2_content['full_text']
        assert "SLOT2_MONDAY_OBJECTIVE" in full_text2
        assert "SLOT2_TUESDAY_OBJECTIVE" in full_text2
        assert "SLOT2_MONDAY_ANTICIPATORY" in full_text2
        
        # Should NOT have Slot 1's content
        assert "SLOT1_MONDAY_OBJECTIVE" not in full_text2
        assert "SLOT1_TUESDAY_OBJECTIVE" not in full_text2
        assert "SLOT1_MONDAY_ANTICIPATORY" not in full_text2
    
    def test_extract_content_for_slot_table_content_dict(self, tmp_path):
        """Test that table_content dict only contains slot's data."""
        doc_path = create_test_doc_with_duplicate_subjects(tmp_path)
        parser = DOCXParser(str(doc_path))
        
        # Extract for Slot 1
        slot1_content = parser.extract_subject_content_for_slot(1, "Math")
        table_content1 = slot1_content['table_content']
        
        # Check Monday content (case-sensitive key)
        assert 'Monday' in table_content1
        monday_content = table_content1['Monday']
        
        # Should have Slot 1's Monday content
        found_slot1_objective = False
        found_slot1_anticipatory = False
        for label, text in monday_content.items():
            if "SLOT1_MONDAY_OBJECTIVE" in text:
                found_slot1_objective = True
            if "SLOT1_MONDAY_ANTICIPATORY" in text:
                found_slot1_anticipatory = True
            # Should NOT have Slot 2's content
            assert "SLOT2_MONDAY_OBJECTIVE" not in text
            assert "SLOT2_MONDAY_ANTICIPATORY" not in text
        
        assert found_slot1_objective, "Slot 1 Monday objective not found"
        assert found_slot1_anticipatory, "Slot 1 Monday anticipatory not found"
    
    def test_extract_images_for_slot_with_real_images(self, tmp_path):
        """Test that image extraction only gets slot's images."""
        doc_path = tmp_path / "test_slot_images.docx"
        doc = Document()
        
        # Slot 1 with RED image
        meta1 = doc.add_table(rows=2, cols=2)
        meta1.rows[0].cells[0].text = "Name: Teacher 1"
        
        daily1 = doc.add_table(rows=2, cols=3)
        daily1.rows[0].cells[1].text = "Monday"
        daily1.rows[1].cells[0].text = "Objective"
        
        # Add RED image to Slot 1
        img1 = PILImage.new('RGB', (2, 2), color='red')
        img1_bytes = io.BytesIO()
        img1.save(img1_bytes, format='PNG')
        img1_bytes.seek(0)
        run1 = daily1.rows[1].cells[1].paragraphs[0].add_run()
        run1.add_picture(img1_bytes, width=1)
        
        # Slot 2 with BLUE image
        meta2 = doc.add_table(rows=2, cols=2)
        meta2.rows[0].cells[0].text = "Name: Teacher 2"
        
        daily2 = doc.add_table(rows=2, cols=3)
        daily2.rows[0].cells[1].text = "Monday"
        daily2.rows[1].cells[0].text = "Objective"
        
        # Add BLUE image to Slot 2
        img2 = PILImage.new('RGB', (2, 2), color='blue')
        img2_bytes = io.BytesIO()
        img2.save(img2_bytes, format='PNG')
        img2_bytes.seek(0)
        run2 = daily2.rows[1].cells[1].paragraphs[0].add_run()
        run2.add_picture(img2_bytes, width=1)
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        # Test extraction
        parser = DOCXParser(str(doc_path))
        
        # Slot 1 should have 1 image from table 1
        slot1_images = parser.extract_images_for_slot(1)
        assert len(slot1_images) == 1, f"Slot 1 should have 1 image, got {len(slot1_images)}"
        assert slot1_images[0]['table_idx'] == 1
        
        # Slot 2 should have 1 image from table 3
        slot2_images = parser.extract_images_for_slot(2)
        assert len(slot2_images) == 1, f"Slot 2 should have 1 image, got {len(slot2_images)}"
        assert slot2_images[0]['table_idx'] == 3
        
        # Images should be different (different data)
        assert slot1_images[0]['data'] != slot2_images[0]['data']
        
        # No overlap
        slot1_table_indices = {img['table_idx'] for img in slot1_images}
        slot2_table_indices = {img['table_idx'] for img in slot2_images}
        assert len(slot1_table_indices & slot2_table_indices) == 0


class TestSlotAwareIntegration:
    """Integration tests for complete slot-aware extraction."""
    
    def test_all_extractions_use_same_tables(self, tmp_path):
        """Test that content, hyperlinks, and images all use same table indices."""
        doc_path = tmp_path / "test_integration.docx"
        doc = Document()
        
        # Create 2 slots
        for slot_num in range(1, 3):
            # Metadata
            meta = doc.add_table(rows=2, cols=2)
            meta.rows[0].cells[0].text = f"Name: Teacher {slot_num}"
            
            # Daily plans with content
            daily = doc.add_table(rows=2, cols=3)
            daily.rows[0].cells[1].text = "Monday"
            daily.rows[1].cells[0].text = "Objective"
            daily.rows[1].cells[1].text = f"SLOT{slot_num}_CONTENT"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        
        parser = DOCXParser(str(doc_path))
        
        # Extract everything for Slot 1
        content1 = parser.extract_subject_content_for_slot(1, "Test")
        hyperlinks1 = parser.extract_hyperlinks_for_slot(1)
        images1 = parser.extract_images_for_slot(1)
        
        # All should reference tables 0-1
        assert content1['table_idx'] == 1
        for link in hyperlinks1:
            assert link['table_idx'] in [0, 1]
        for img in images1:
            assert img['table_idx'] in [0, 1]
        
        # Extract everything for Slot 2
        content2 = parser.extract_subject_content_for_slot(2, "Test")
        hyperlinks2 = parser.extract_hyperlinks_for_slot(2)
        images2 = parser.extract_images_for_slot(2)
        
        # All should reference tables 2-3
        assert content2['table_idx'] == 3
        for link in hyperlinks2:
            assert link['table_idx'] in [2, 3]
        for img in images2:
            assert img['table_idx'] in [2, 3]
        
        # Content should be different
        assert "SLOT1_CONTENT" in content1['full_text']
        assert "SLOT2_CONTENT" in content2['full_text']
        assert "SLOT2_CONTENT" not in content1['full_text']
        assert "SLOT1_CONTENT" not in content2['full_text']
    
    def test_extract_content_with_teacher_name_disambiguation(self, tmp_path):
        """Test extract_subject_content_for_slot with teacher_name parameter."""
        # Create doc with same teacher, same subject, different slots (Daniela's bug scenario)
        doc_path = tmp_path / "test_teacher_name_extraction.docx"
        doc = Document()
        
        # Slot 1: Science Grade 6 (Grande)
        meta1 = doc.add_table(rows=3, cols=2)
        meta1.rows[0].cells[0].text = "Name: Mariela Grande"
        meta1.rows[1].cells[0].text = "Subject: Science"
        meta1.rows[1].cells[1].text = "Grade: 6"
        
        daily1 = doc.add_table(rows=3, cols=6)
        daily1.rows[0].cells[0].text = "Day"
        daily1.rows[0].cells[1].text = "Monday"
        daily1.rows[1].cells[0].text = "Objective"
        daily1.rows[1].cells[1].text = "GRADE6_SCIENCE_MONDAY"
        daily1.rows[2].cells[0].text = "Anticipatory Set"
        daily1.rows[2].cells[1].text = "GRADE6_SCIENCE_ANTICIPATORY"
        
        # Slot 2: Science Grade 2 (Grande)
        meta2 = doc.add_table(rows=3, cols=2)
        meta2.rows[0].cells[0].text = "Name: Mariela Grande"
        meta2.rows[1].cells[0].text = "Subject: Science"
        meta2.rows[1].cells[1].text = "Grade: 2"
        
        daily2 = doc.add_table(rows=3, cols=6)
        daily2.rows[0].cells[0].text = "Day"
        daily2.rows[0].cells[1].text = "Monday"
        daily2.rows[1].cells[0].text = "Objective"
        daily2.rows[1].cells[1].text = "GRADE2_SCIENCE_MONDAY"
        daily2.rows[2].cells[0].text = "Anticipatory Set"
        daily2.rows[2].cells[1].text = "GRADE2_SCIENCE_ANTICIPATORY"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        parser = DOCXParser(str(doc_path))
        
        # Extract Slot 1 content with teacher_name
        # Since both slots have same teacher, it should use slot_number as primary
        slot1_content = parser.extract_subject_content_for_slot(
            slot_number=1,
            subject="Science",
            teacher_name="Mariela Grande"
        )
        
        assert slot1_content['found'] is True
        assert slot1_content['slot_number'] == 1
        assert "GRADE6_SCIENCE_MONDAY" in slot1_content['full_text']
        assert "GRADE6_SCIENCE_ANTICIPATORY" in slot1_content['full_text']
        assert "GRADE2_SCIENCE_MONDAY" not in slot1_content['full_text']
        
        # Extract Slot 2 content with teacher_name
        slot2_content = parser.extract_subject_content_for_slot(
            slot_number=2,
            subject="Science",
            teacher_name="Mariela Grande"
        )
        
        assert slot2_content['found'] is True
        assert slot2_content['slot_number'] == 2
        assert "GRADE2_SCIENCE_MONDAY" in slot2_content['full_text']
        assert "GRADE2_SCIENCE_ANTICIPATORY" in slot2_content['full_text']
        assert "GRADE6_SCIENCE_MONDAY" not in slot2_content['full_text']
    
    def test_extract_content_with_teacher_name_different_teachers(self, tmp_path):
        """Test teacher_name disambiguation when different teachers teach same subject."""
        doc_path = tmp_path / "test_different_teachers_same_subject.docx"
        doc = Document()
        
        # Slot 1: Science (Morais)
        meta1 = doc.add_table(rows=3, cols=2)
        meta1.rows[0].cells[0].text = "Name: Catarina Morais"
        meta1.rows[1].cells[0].text = "Subject: Science"
        
        daily1 = doc.add_table(rows=3, cols=6)
        daily1.rows[0].cells[0].text = "Day"
        daily1.rows[0].cells[1].text = "Monday"
        daily1.rows[1].cells[0].text = "Objective"
        daily1.rows[1].cells[1].text = "MORAIS_SCIENCE_MONDAY"
        
        # Slot 2: Science (Grande)
        meta2 = doc.add_table(rows=3, cols=2)
        meta2.rows[0].cells[0].text = "Name: Mariela Grande"
        meta2.rows[1].cells[0].text = "Subject: Science"
        
        daily2 = doc.add_table(rows=3, cols=6)
        daily2.rows[0].cells[0].text = "Day"
        daily2.rows[0].cells[1].text = "Monday"
        daily2.rows[1].cells[0].text = "Objective"
        daily2.rows[1].cells[1].text = "GRANDE_SCIENCE_MONDAY"
        
        # Signature
        sig = doc.add_table(rows=2, cols=2)
        sig.rows[0].cells[0].text = "Signature"
        
        doc.save(str(doc_path))
        parser = DOCXParser(str(doc_path))
        
        # Extract with Grande teacher_name - should get Slot 2
        grande_content = parser.extract_subject_content_for_slot(
            slot_number=1,  # Requested slot (but should be overridden)
            subject="Science",
            teacher_name="Mariela Grande"
        )
        
        assert grande_content['found'] is True
        assert grande_content['slot_number'] == 2  # Should find Slot 2 (Grande)
        assert "GRANDE_SCIENCE_MONDAY" in grande_content['full_text']
        assert "MORAIS_SCIENCE_MONDAY" not in grande_content['full_text']
        
        # Extract with Morais teacher_name - should get Slot 1
        morais_content = parser.extract_subject_content_for_slot(
            slot_number=1,
            subject="Science",
            teacher_name="Catarina Morais"
        )
        
        assert morais_content['found'] is True
        assert morais_content['slot_number'] == 1  # Should find Slot 1 (Morais)
        assert "MORAIS_SCIENCE_MONDAY" in morais_content['full_text']
        assert "GRANDE_SCIENCE_MONDAY" not in morais_content['full_text']
    
    def test_extract_content_without_teacher_name_backward_compat(self, tmp_path):
        """Test backward compatibility: extract_subject_content_for_slot without teacher_name."""
        doc_path = create_test_doc_with_duplicate_subjects(tmp_path)
        parser = DOCXParser(str(doc_path))
        
        # Should work without teacher_name parameter (backward compatibility)
        slot1_content = parser.extract_subject_content_for_slot(
            slot_number=1,
            subject="Math"
            # teacher_name not provided
        )
        
        assert slot1_content['found'] is True
        assert slot1_content['slot_number'] == 1
        assert "SLOT1_MONDAY_OBJECTIVE" in slot1_content['full_text']
    
    def test_extract_content_teacher_name_not_found_fallback(self, tmp_path):
        """Test fallback when teacher_name doesn't match any slot."""
        doc_path = create_test_doc_with_duplicate_subjects(tmp_path)
        parser = DOCXParser(str(doc_path))
        
        # Teacher name doesn't match - should fall back to slot_number
        slot1_content = parser.extract_subject_content_for_slot(
            slot_number=1,
            subject="Math",
            teacher_name="Unknown Teacher"  # Not in document
        )
        
        # Should still extract from Slot 1 (using slot_number)
        assert slot1_content['found'] is True
        assert slot1_content['slot_number'] == 1
        assert "SLOT1_MONDAY_OBJECTIVE" in slot1_content['full_text']