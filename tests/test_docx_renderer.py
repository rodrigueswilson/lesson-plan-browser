"""
Tests for DOCX Renderer.

Tests:
1. Basic rendering
2. Metadata population
3. Daily plan population
4. Markdown formatting
5. Missing optional fields
6. Template preservation
"""

import sys
import os
from pathlib import Path
import json
import tempfile

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.docx_renderer import DOCXRenderer
from tools.markdown_to_docx import MarkdownToDocx
from docx import Document


def test_basic_rendering():
    """Test basic DOCX rendering."""
    print("\nTest: Basic Rendering")
    
    # Load fixture
    fixture_path = "tests/fixtures/valid_lesson_minimal.json"
    with open(fixture_path, 'r', encoding='utf-8') as f:
        json_data = json.load(f)
    
    # Render
    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_basic_rendering.docx"
    
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)
    
    assert success, "Rendering should succeed"
    assert Path(output_path).exists(), "Output file should exist"
    
    # Verify content
    doc = Document(output_path)
    assert len(doc.tables) == 3, "Should have 3 tables"
    
    print("  PASS: Basic rendering works")
    return True


def test_metadata_population():
    """Test metadata table population."""
    print("\nTest: Metadata Population")
    
    # Create minimal JSON with metadata
    json_data = {
        "metadata": {
            "week_of": "10/6-10/10",
            "grade": "7",
            "subject": "Social Studies",
            "homeroom": "302",
            "teacher_name": "Ms. Smith"
        },
        "days": {
            "monday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test"
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"}
            },
            "tuesday": {
                "unit_lesson": "Test",
                "objective": {"content_objective": "Test", "student_goal": "Test", "wida_objective": "Test"},
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"}
            },
            "wednesday": {
                "unit_lesson": "Test",
                "objective": {"content_objective": "Test", "student_goal": "Test", "wida_objective": "Test"},
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"}
            },
            "thursday": {
                "unit_lesson": "Test",
                "objective": {"content_objective": "Test", "student_goal": "Test", "wida_objective": "Test"},
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"}
            },
            "friday": {
                "unit_lesson": "Test",
                "objective": {"content_objective": "Test", "student_goal": "Test", "wida_objective": "Test"},
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"}
            }
        }
    }
    
    # Render
    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_metadata.docx"
    
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)
    
    assert success, "Rendering should succeed"
    
    # Verify metadata
    doc = Document(output_path)
    metadata_table = doc.tables[0]
    row = metadata_table.rows[0]
    
    assert "Ms. Smith" in row.cells[0].text, "Should contain teacher name"
    assert "7" in row.cells[1].text, "Should contain grade"
    assert "302" in row.cells[2].text, "Should contain homeroom"
    assert "Social Studies" in row.cells[3].text, "Should contain subject"
    assert "10/6-10/10" in row.cells[4].text, "Should contain week"
    
    print("  PASS: Metadata population works")
    return True


def test_daily_plan_population():
    """Test daily plan population."""
    print("\nTest: Daily Plan Population")
    
    # Load fixture
    fixture_path = "tests/fixtures/valid_lesson_minimal.json"
    with open(fixture_path, 'r', encoding='utf-8') as f:
        json_data = json.load(f)
    
    # Render
    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_daily_plans.docx"
    
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)
    
    assert success, "Rendering should succeed"
    
    # Verify daily plans
    doc = Document(output_path)
    daily_table = doc.tables[1]
    
    # Check Monday column (col 1)
    monday_unit = daily_table.rows[1].cells[1].text
    assert "Unit One Lesson Seven" in monday_unit, "Should contain Monday unit"
    
    monday_objective = daily_table.rows[2].cells[1].text
    assert "Content:" in monday_objective, "Should contain content objective"
    assert "Student Goal:" in monday_objective, "Should contain student goal"
    assert "WIDA/Bilingual:" in monday_objective, "Should contain WIDA objective"
    
    print("  PASS: Daily plan population works")
    return True


def test_markdown_formatting():
    """Test markdown formatting in cells."""
    print("\nTest: Markdown Formatting")
    
    # Create JSON with markdown
    json_data = {
        "metadata": {
            "week_of": "10/6-10/10",
            "grade": "7",
            "subject": "Test"
        },
        "days": {
            day: {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test with **bold** and *italic*",
                    "student_goal": "Test",
                    "wida_objective": "Test"
                },
                "anticipatory_set": {
                    "original_content": "Test with **bold text**"
                },
                "tailored_instruction": {
                    "original_content": "Test",
                    "ell_support": [
                        {
                            "strategy_name": "Test Strategy",
                            "implementation": "Test with *italic*",
                            "proficiency_levels": "1-5"
                        }
                    ]
                },
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"}
            }
            for day in ["monday", "tuesday", "wednesday", "thursday", "friday"]
        }
    }
    
    # Render
    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_markdown.docx"
    
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)
    
    assert success, "Rendering should succeed"
    
    # Verify formatting (basic check - full verification would require inspecting runs)
    doc = Document(output_path)
    daily_table = doc.tables[1]
    
    # Check that content was added
    monday_objective = daily_table.rows[2].cells[1].text
    assert "bold" in monday_objective, "Should contain bold text"
    assert "italic" in monday_objective, "Should contain italic text"
    
    print("  PASS: Markdown formatting works")
    return True


def test_missing_optional_fields():
    """Test handling of missing optional fields."""
    print("\nTest: Missing Optional Fields")
    
    # Create minimal JSON (no optional fields)
    json_data = {
        "metadata": {
            "week_of": "10/6-10/10",
            "grade": "7",
            "subject": "Test"
        },
        "days": {
            day: {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test"
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": ""}
            }
            for day in ["monday", "tuesday", "wednesday", "thursday", "friday"]
        }
    }
    
    # Render
    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_missing_fields.docx"
    
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)
    
    assert success, "Rendering should succeed even with missing optional fields"
    
    print("  PASS: Missing optional fields handled correctly")
    return True


def test_template_preservation():
    """Test that template formatting is preserved."""
    print("\nTest: Template Preservation")
    
    # Load original template
    template_path = "input/Lesson Plan Template SY'25-26.docx"
    original_doc = Document(template_path)
    
    # Get original table count
    original_table_count = len(original_doc.tables)
    
    # Render
    fixture_path = "tests/fixtures/valid_lesson_minimal.json"
    with open(fixture_path, 'r', encoding='utf-8') as f:
        json_data = json.load(f)
    
    output_path = "output/test_template_preservation.docx"
    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)
    
    assert success, "Rendering should succeed"
    
    # Verify structure preserved
    rendered_doc = Document(output_path)
    assert len(rendered_doc.tables) == original_table_count, "Should preserve table count"
    
    # Verify table structure
    daily_table = rendered_doc.tables[1]
    assert len(daily_table.rows) == 8, "Should have 8 rows in daily plans table"
    assert len(daily_table.columns) == 6, "Should have 6 columns in daily plans table"
    
    print("  PASS: Template structure preserved")
    return True


def test_markdown_to_docx_utilities():
    """Test markdown to DOCX utility functions."""
    print("\nTest: Markdown to DOCX Utilities")
    
    # Test clean_markdown
    text_with_markdown = "This is **bold** and *italic* text"
    clean_text = MarkdownToDocx.clean_markdown(text_with_markdown)
    assert "**" not in clean_text, "Should remove bold markers"
    assert "*" not in clean_text, "Should remove italic markers"
    assert "bold" in clean_text, "Should keep text content"
    
    # Test bullet pattern
    bullet_text = "- This is a bullet"
    match = MarkdownToDocx.BULLET_PATTERN.match(bullet_text)
    assert match is not None, "Should match bullet pattern"
    assert match.group(1) == "This is a bullet", "Should extract bullet text"
    
    print("  PASS: Markdown utilities work correctly")
    return True


def test_vocabulary_cognate_awareness_block():
    """Test that vocabulary/cognate awareness block is rendered from structured data."""
    print("\nTest: Vocabulary / Cognate Awareness Block")

    # Minimal JSON with vocabulary_cognates and a cognate_awareness strategy
    json_data = {
        "metadata": {
            "week_of": "10/6-10/10",
            "grade": "7",
            "subject": "Social Studies",
        },
        "days": {
            "monday": {
                "unit_lesson": "Unit One Lesson Seven",
                "objective": {
                    "content_objective": "Students will be able to explain key concepts.",
                    "student_goal": "I will explain key concepts.",
                    "wida_objective": "Students will use language to explain key concepts (ELD-SS.6-8.Explain.Reading/Writing) by using cognate awareness.",
                },
                "anticipatory_set": {"original_content": "Hook"},
                "vocabulary_cognates": [
                    {"english": "map", "portuguese": "mapa"},
                    {"english": "capital", "portuguese": "capital"},
                    {"english": "coast", "portuguese": "costa"},
                    {"english": "population", "portuguese": "população"},
                    {"english": "legend", "portuguese": "legenda"},
                    {"english": "scale", "portuguese": "escala"},
                ],
                "tailored_instruction": {
                    "original_content": "Core instruction.",
                    "co_teaching_model": {
                        "model_name": "Station Teaching",
                        "rationale": "Test",
                        "wida_context": "Test",
                        "phase_plan": [],
                        "implementation_notes": [],
                        "warnings": [],
                    },
                    "ell_support": [
                        {
                            "strategy_id": "cognate_awareness",
                            "strategy_name": "Cognate Awareness",
                            "implementation": "Use the 6 cognate pairs at the bilingual station.",
                            "proficiency_levels": "Levels 1-4",
                        }
                    ],
                    "materials": ["Word cards"],
                },
                "misconceptions": {"original_content": "None"},
                "assessment": {"primary_assessment": "Exit ticket"},
                "homework": {"original_content": "Read"},
            },
            # Minimal stubs for other days
            "tuesday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
            "wednesday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
            "thursday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
            "friday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
        },
    }

    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_vocabulary_cognate_awareness.docx"

    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)

    assert success, "Rendering should succeed"

    doc = Document(output_path)
    daily_table = doc.tables[1]

    # Monday column (col 1) Tailored Instruction row (index 4 per renderer.INSTRUCTION_ROW)
    instruction_cell_text = daily_table.rows[4].cells[1].text
    assert "Vocabulary / Cognate Awareness:" in instruction_cell_text
    # Check a couple of representative pairs for correct formatting
    assert "**map**" in instruction_cell_text or "map" in instruction_cell_text
    assert "mapa" in instruction_cell_text
    assert "scale" in instruction_cell_text
    assert "escala" in instruction_cell_text

    print("  PASS: Vocabulary / Cognate Awareness block rendered from structured data")
    return True


def test_sentence_frames_grouped_by_level_block():
    """Test that sentence frames are rendered in a grouped-by-level block from structured data."""
    print("\nTest: Sentence Frames Grouped-By-Level Block")

    json_data = {
        "metadata": {
            "week_of": "10/6-10/10",
            "grade": "7",
            "subject": "Social Studies",
        },
        "days": {
            "monday": {
                "unit_lesson": "Unit One Lesson Seven",
                "objective": {
                    "content_objective": "Students will be able to explain key concepts.",
                    "student_goal": "I will explain key concepts.",
                    "wida_objective": "Students will use language to explain key concepts (ELD-SS.6-8.Explain.Reading/Writing) by using sentence frames.",
                },
                "anticipatory_set": {"original_content": "Hook"},
                "sentence_frames": [
                    {  # Levels 1-2
                        "proficiency_level": "levels_1_2",
                        "english": "This is ___",
                        "portuguese": "Isto é ___",
                        "language_function": "identify",
                        "frame_type": "frame",
                    },
                    {
                        "proficiency_level": "levels_1_2",
                        "english": "I see ___",
                        "portuguese": "Eu vejo ___",
                        "language_function": "describe",
                        "frame_type": "frame",
                    },
                    {
                        "proficiency_level": "levels_1_2",
                        "english": "It has ___",
                        "portuguese": "Tem ___",
                        "language_function": "describe",
                        "frame_type": "frame",
                    },
                    {  # Levels 3-4
                        "proficiency_level": "levels_3_4",
                        "english": "First ___, then ___",
                        "portuguese": "Primeiro ___, depois ___",
                        "language_function": "sequence",
                        "frame_type": "frame",
                    },
                    {
                        "proficiency_level": "levels_3_4",
                        "english": "This shows ___ because ___",
                        "portuguese": "Isso mostra ___ porque ___",
                        "language_function": "explain",
                        "frame_type": "frame",
                    },
                    {
                        "proficiency_level": "levels_3_4",
                        "english": "I think ___ because ___",
                        "portuguese": "Eu acho ___ porque ___",
                        "language_function": "justify",
                        "frame_type": "frame",
                    },
                    {  # Levels 5-6
                        "proficiency_level": "levels_5_6",
                        "english": "Evidence suggests that ___",
                        "portuguese": "As evidências sugerem que ___",
                        "language_function": "argue",
                        "frame_type": "stem",
                    },
                    {
                        "proficiency_level": "levels_5_6",
                        "english": "How does ___ relate to ___?",
                        "portuguese": "Como ___ se relaciona com ___?",
                        "language_function": "analyze",
                        "frame_type": "open_question",
                    },
                ],
                "tailored_instruction": {
                    "original_content": "Core instruction.",
                    "co_teaching_model": {
                        "model_name": "Station Teaching",
                        "rationale": "Test",
                        "wida_context": "Test",
                        "phase_plan": [],
                        "implementation_notes": [],
                        "warnings": [],
                    },
                    "ell_support": [],
                    "materials": ["Word cards"],
                },
                "misconceptions": {"original_content": "None"},
                "assessment": {"primary_assessment": "Exit ticket"},
                "homework": {"original_content": "Read"},
            },
            # Minimal stubs for other days
            "tuesday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
            "wednesday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
            "thursday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
            "friday": {
                "unit_lesson": "Test",
                "objective": {
                    "content_objective": "Test",
                    "student_goal": "Test",
                    "wida_objective": "Test",
                },
                "anticipatory_set": {"original_content": "Test"},
                "tailored_instruction": {"original_content": "Test"},
                "misconceptions": {"original_content": "Test"},
                "assessment": {"primary_assessment": "Test"},
                "homework": {"original_content": "Test"},
            },
        },
    }

    template_path = "input/Lesson Plan Template SY'25-26.docx"
    output_path = "output/test_sentence_frames_grouped.docx"

    renderer = DOCXRenderer(template_path)
    success = renderer.render(json_data, output_path)

    assert success, "Rendering should succeed"

    doc = Document(output_path)
    daily_table = doc.tables[1]

    # Tailored Instruction row (index 4 per renderer.INSTRUCTION_ROW)
    instruction_cell_text = daily_table.rows[4].cells[1].text
    assert "Sentence Frames / Stems / Questions:" in instruction_cell_text
    assert "Levels 1-2" in instruction_cell_text
    assert "Levels 3-4" in instruction_cell_text
    assert "Levels 5-6" in instruction_cell_text
    assert "This is" in instruction_cell_text and "identify" in instruction_cell_text
    assert "Isto" in instruction_cell_text or "Isto".lower() in instruction_cell_text.lower()

    print("  PASS: Sentence frames block rendered and grouped by level from structured data")
    return True


def run_all_tests():
    """Run all tests."""
    print("=" * 60)
    print("DOCX Renderer Tests")
    print("=" * 60)
    
    tests = [
        test_basic_rendering,
        test_metadata_population,
        test_daily_plan_population,
        test_markdown_formatting,
        test_missing_optional_fields,
        test_template_preservation,
        test_markdown_to_docx_utilities,
        test_vocabulary_cognate_awareness_block,
        test_sentence_frames_grouped_by_level_block,
    ]
    
    passed = 0
    failed = 0
    
    for test in tests:
        try:
            if test():
                passed += 1
        except AssertionError as e:
            print(f"  FAIL: {e}")
            failed += 1
        except Exception as e:
            print(f"  ERROR: {e}")
            import traceback
            traceback.print_exc()
            failed += 1
    
    print("\n" + "=" * 60)
    print(f"Results: {passed}/{len(tests)} passed")
    if failed > 0:
        print(f"         {failed}/{len(tests)} failed")
    print("=" * 60)
    
    return failed == 0


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
