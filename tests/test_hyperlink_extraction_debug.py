"""Debug hyperlink extraction with detailed error reporting."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')
files_to_test = [
    folder / 'Morais 10_20_25 - 10_24_25.docx',
    folder / 'Mrs. Grande Science 10_20- 10_24.docx'
]

for file_path in files_to_test:
    print("=" * 80)
    print(f"📄 {file_path.name}")
    print("=" * 80)
    
    try:
        doc = Document(str(file_path))
        print(f"✅ Document opened")
        print(f"   Paragraphs: {len(doc.paragraphs)}")
        print(f"   Tables: {len(doc.tables)}")
        
        # Check paragraphs for hyperlinks
        para_links = 0
        for para in doc.paragraphs:
            try:
                hyperlinks = para._element.xpath('.//w:hyperlink')
                if hyperlinks:
                    para_links += len(hyperlinks)
                    print(f"\n   Found {len(hyperlinks)} hyperlinks in paragraph:")
                    for hl in hyperlinks:
                        r_id = hl.get(qn('r:id'))
                        if r_id and r_id in para.part.rels:
                            url = para.part.rels[r_id].target_ref
                            text = ''.join(node.text for node in hl.xpath('.//w:t') if node.text)
                            print(f"      • {text[:40]} → {url[:50]}")
            except Exception as e:
                print(f"   ⚠️  Error in paragraph: {e}")
        
        print(f"\n   Total paragraph hyperlinks: {para_links}")
        
        # Check tables for hyperlinks
        table_links = 0
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        try:
                            hyperlinks = para._element.xpath('.//w:hyperlink')
                            if hyperlinks:
                                table_links += len(hyperlinks)
                                row_label = row.cells[0].text.strip()[:20] if row.cells else ""
                                print(f"\n   Found {len(hyperlinks)} hyperlinks in table {table_idx}, row {row_idx} ({row_label}), cell {cell_idx}:")
                                for hl in hyperlinks:
                                    r_id = hl.get(qn('r:id'))
                                    if r_id and r_id in para.part.rels:
                                        url = para.part.rels[r_id].target_ref
                                        text = ''.join(node.text for node in hl.xpath('.//w:t') if node.text)
                                        print(f"      • {text[:40]} → {url[:50]}")
                        except Exception as e:
                            print(f"   ⚠️  Error in table cell: {e}")
        
        print(f"\n   Total table hyperlinks: {table_links}")
        print(f"\n   📊 TOTAL HYPERLINKS: {para_links + table_links}")
        
    except Exception as e:
        print(f"❌ Error opening document: {e}")
        import traceback
        traceback.print_exc()
    
    print()
