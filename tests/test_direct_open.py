"""Test opening files directly with python-docx."""
from pathlib import Path
from docx import Document

folder = Path(r'F:\rodri\Documents\OneDrive\AS\Daniela LP\25 W43')
morais_file = folder / 'Morais 10_20_25 - 10_24_25.docx'

print(f"Testing: {morais_file}")
print(f"Exists: {morais_file.exists()}")
print(f"Absolute path: {morais_file.absolute()}")

try:
    # Try opening directly
    doc = Document(str(morais_file.absolute()))
    print(f"✅ Opened successfully!")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    # Check for hyperlinks
    hyperlink_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if hasattr(run, '_element'):
                for child in run._element:
                    if child.tag.endswith('hyperlink'):
                        hyperlink_count += 1
    
    print(f"Hyperlinks in paragraphs: {hyperlink_count}")
    
except Exception as e:
    print(f"❌ Error: {e}")
    import traceback
    traceback.print_exc()
