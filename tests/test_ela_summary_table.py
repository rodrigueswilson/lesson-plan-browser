"""Unit tests for ELA Summary of Key Learning table detection and parsing."""

from __future__ import annotations

import json
import sqlite3
import sys
from pathlib import Path

import pytest
from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[1]
_SCRAPER_DIR = str(_REPO_ROOT / "tools" / "scraper")
if _SCRAPER_DIR not in sys.path:
    sys.path.insert(0, _SCRAPER_DIR)
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from tools.scraper.ela_summary_table import (  # noqa: E402
    is_summary_key_learning_table,
    parse_summary_rows,
)
from tools.scraper.subject_config import SubjectConfig  # noqa: E402


def _p(text: str, bold: bool = False) -> dict:
    style = {"bold": True} if bold else None
    return {"type": "paragraph", "text": text, "is_bullet": False, "ilvl": 0, "runs": [{"type": "run", "text": text, "style": style}]}


def _cell(*paragraphs: dict) -> dict:
    return {"type": "table-cell", "grid_span": None, "v_merge": None, "content": list(paragraphs)}


def _row(*cells: dict) -> dict:
    return {"type": "table-row", "cells": list(cells)}


def _table(*rows: dict) -> dict:
    return {"type": "table", "rows": list(rows), "depth": 0}


def _htmlize(content: list) -> str:
    parts = []
    for e in content:
        if e.get("type") == "paragraph":
            inner = ""
            for run in e.get("runs") or []:
                t = run.get("text") or ""
                if (run.get("style") or {}).get("bold"):
                    t = f"<b>{t}</b>"
                inner += t
            if inner:
                parts.append(f"<p>{inner}</p>")
    return "".join(parts)


def _pat():
    cfg = SubjectConfig.get_config("ELA")
    return cfg["patterns"].get("standards_summary_mentions") or cfg["patterns"]["standards"]


def test_is_summary_detects_valid_table() -> None:
    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Learning Intentions and Success Criteria")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
        _row(
            _cell(_p("1")),
            _cell(
                _p(
                    "Learning Intention:\nDefine terms.\nSuccess Criteria:\nI can explain.",
                )
            ),
            _cell(_p("Oral Task", bold=True), _p("Partner discussion steps.")),
            _cell(_p("Preview SL.PE.2.1A-C and RI.CR.2.1")),
        ),
    )
    assert is_summary_key_learning_table(tbl) is True


def test_is_summary_rejects_bad_header_row() -> None:
    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Wrong header")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
        _row(_cell(_p("1")), _cell(_p("a")), _cell(_p("b")), _cell(_p("c"))),
    )
    assert is_summary_key_learning_table(tbl) is False


def test_is_summary_rejects_too_few_rows() -> None:
    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Learning Intentions and Success Criteria")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
    )
    assert is_summary_key_learning_table(tbl) is False


def test_parse_golden_grade2_summary_lessons_1_and_2() -> None:
    """Fixture text from Grade 2 Unit 8 Summary of Key Learning (operator paste)."""
    b1 = """Learning Intention:
I am learning to define key words in a text.
I am learning that there was a reason the Constitution was written.

Success Criteria:
I can define: Constitution, protest, representatives
I can work with a partner to explain why states wanted the Virginia or New Jersey Plan."""
    c1_intro = (
        "Students will engage in a structured partner discussion to answer the question: "
        '"Why would smaller states want the New Jersey Plan while larger states would want the Virginia Plan?" '
    )
    c1_rest = (
        "Give each pair a copy of the partner discussion rubric.\n"
        "Each student takes turns being the speaker and the listener."
    )
    d1 = """Unit Introduction: Essential Question
Preview SL.PE.2.1A-C
Vocabulary L .VL.2..2, 
Read Aloud and Text Dependent Questions RI.CR.2.1, RI.IT.2.3, L .VL.2..2,SL.PE.2.1A-C, SL.II.2.2, 6.1.2.HistoryUP.1:
Daily Instructional Task: Partner Discussions -Oral Explanatory Task RI.CR.2.1,  RI.IT.2.3, SL.PE.2.1A-C, SL.II.2.2, SL.ES.2.3, 6.1.2.HistoryUP.1:"""

    b2 = """Learning Intention:
We are learning to define key terms in texts.
We are learning about the constitution.

Success Criteria:
I can explain the term, compromise and cite evidence from the text to support my definition.
I can explain what the US Constitution is and why it is important."""
    c2_body = (
        "Write an informational paragraph explaining what the U.S. Constitution is and why it is important.\n"
        "In your writing, be sure to:\n"
        "Introduce your topic by explaining what the U.S. Constitution is."
    )
    d2 = """Review RI.IT.2.3, RI.CI.2.2, SL.PE.2.1A-C
Vocabulary L.VL.2..2
Read Aloud and Text Dependent Questions RI.IT.2.3, RI.CR.2.1, RI.CI.2.2, SL.PE.2.1A-C, 6.1.2.HistoryCC.3, 6.1.2.HistoryUP.1, 6.1.2.CivicsPI.6:
Collaborative Discussion RI.IT.2.3, RI.CR.2.1, RI.CI.2.2, SL.PE.2.1A-C, 6.1.2.HistoryCC.3,6.1.2.HistoryUP.1, 6.1.2.CivicsPI.6:
Daily Instructional Task: Explanatory Writing W.IW.2.2A-C, 6.1.2.HistoryCC.3, 6.1.2.HistoryUP.1, 6.1.2.CivicsPI.6"""

    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Learning Intentions and Success Criteria")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
        _row(
            _cell(_p("1")),
            _cell(_p(b1)),
            _cell(_p("Oral Explanatory Task", bold=True), _p(c1_intro), _p(c1_rest)),
            _cell(_p(d1)),
        ),
        _row(
            _cell(_p("2")),
            _cell(_p(b2)),
            _cell(_p("Explanatory Writing", bold=True), _p(c2_body)),
            _cell(_p(d2)),
        ),
    )
    out = parse_summary_rows(tbl, _htmlize, _pat())
    assert 1 in out and 2 in out
    r1 = out[1]
    assert "define key words" in r1["learning_intention"]
    assert "Constitution, protest" in r1["success_criteria"]
    assert r1["daily_task_title"] == "Oral Explanatory Task"
    assert "partner discussion" in r1["daily_task_body"].lower()
    m1 = set(r1["standards_mentions"])
    assert "SL.PE.2.1A-C" in m1
    assert "RI.CR.2.1" in m1
    assert "6.1.2.HistoryUP.1" in m1
    assert "SL.ES.2.3" in m1
    assert "L.VL.2.2" in m1

    r2 = out[2]
    assert r2["daily_task_title"] == "Explanatory Writing"
    m2 = set(r2["standards_mentions"])
    assert "W.IW.2.2A-C" in m2
    assert "6.1.2.HistoryCC.3" in m2
    assert "6.1.2.CivicsPI.6" in m2


def test_column_c_fallback_uses_first_line_when_not_bold() -> None:
    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Learning Intentions and Success Criteria")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
        _row(
            _cell(_p("1")),
            _cell(_p("Learning Intention:\nX.\nSuccess Criteria:\nY.")),
            _cell(_p("Plain Task Title"), _p("Body line two.")),
            _cell(_p("RI.CR.2.1")),
        ),
    )
    out = parse_summary_rows(tbl, _htmlize, _pat())
    assert out[1]["daily_task_title"] == "Plain Task Title"
    assert "Body line two." in out[1]["daily_task_body"]


def test_parse_summary_rows_splits_columns() -> None:
    tbl = _table(
        _row(_cell(_p("Summary of Key Learning"))),
        _row(
            _cell(_p("Lesson")),
            _cell(_p("Learning Intentions and Success Criteria")),
            _cell(_p("Daily Instructional Task")),
            _cell(_p("Content and Learning Strategies")),
        ),
        _row(
            _cell(_p("1")),
            _cell(_p("Learning Intention:\nOne line.\nSuccess Criteria:\nTwo line.")),
            _cell(_p("Explanatory Writing", bold=True), _p("Write a paragraph.")),
            _cell(_p("Standards RI.CR.3.1 and 6.1.2.HistoryUP.1")),
        ),
    )
    out = parse_summary_rows(tbl, _htmlize, _pat())
    assert 1 in out
    row = out[1]
    assert row["learning_intention"] == "One line."
    assert row["success_criteria"] == "Two line."
    assert row["daily_task_title"] == "Explanatory Writing"
    assert "Write a paragraph." in row["daily_task_body"]
    assert "RI.CR.3.1" in row["standards_mentions"]


@pytest.mark.skipif(not (_REPO_ROOT / "data" / "curriculum.db").is_file(), reason="data/curriculum.db not present")
def test_ela_ingest_persists_ela_key_learning_summary(tmp_path: Path) -> None:
    db_copy = tmp_path / "curriculum_test.db"
    db_copy.write_bytes((_REPO_ROOT / "data" / "curriculum.db").read_bytes())

    docx_path = tmp_path / "ela_summary_unit.docx"
    doc = Document()
    t = doc.add_table(rows=3, cols=4)
    t.rows[0].cells[0].text = "Summary of Key Learning"
    for i, h in enumerate(
        [
            "Lesson",
            "Learning Intentions and Success Criteria",
            "Daily Instructional Task",
            "Content and Learning Strategies",
        ]
    ):
        t.rows[1].cells[i].text = h
    t.rows[2].cells[0].text = "1"
    t.rows[2].cells[1].text = "Learning Intention:\nDo something.\nSuccess Criteria:\nI can do it."
    c2 = t.rows[2].cells[2]
    c2.text = ""
    p = c2.paragraphs[0]
    r = p.add_run("Visual Task")
    r.bold = True
    p2 = c2.add_paragraph("Draw a chart.")
    t.rows[2].cells[3].text = "Use SL.PE.2.1A-C"
    doc.add_paragraph("Lesson 1: Sample Title for Ingest")

    doc.save(docx_path)

    unit_id = "ELA_SUMMARY_TEST_UNIT"
    conn = sqlite3.connect(db_copy)
    conn.execute(
        """INSERT OR REPLACE INTO units
        (id, grade, subject, unit_number, title, description, source)
        VALUES (?, 3, 'ELA', 9, 'ELA summary test', '', 'test')""",
        (unit_id,),
    )
    conn.execute(
        "INSERT OR IGNORE INTO units_intro (unit_id, essential_questions, enduring_understandings) VALUES (?, '[]', '[]')",
        (unit_id,),
    )
    conn.commit()
    conn.close()

    sys.path.insert(0, _SCRAPER_DIR)
    from table_extractor import RecursiveTableParser

    parser = RecursiveTableParser(db_path=str(db_copy))
    n = parser.ingest_to_curriculum(
        str(docx_path),
        unit_id,
        subject="ELA",
        source_url="https://docs.google.com/document/d/test_summary_table/edit",
        parser_version="table_extractor@test_summary",
    )
    assert n >= 1

    conn = sqlite3.connect(db_copy)
    row = conn.execute(
        "SELECT ela_key_learning_summary FROM lessons WHERE unit_id = ? AND lesson_number = 1",
        (unit_id,),
    ).fetchone()
    conn.close()
    assert row is not None and row[0]
    data = json.loads(row[0])
    assert data.get("schema_version") == 1
    assert "Do something." in data.get("learning_intention", "")
    assert data.get("daily_task_title") == "Visual Task"
    assert any("SL.PE.2.1A-C" in m for m in data.get("standards_mentions", []))
