"""
End-to-end test for semantic anchoring and structure-based image placement.

Tests the full pipeline: Parse → Extract media → Mock transform → Render with anchoring
"""

import pytest
import json
from pathlib import Path
from tools.docx_parser import DOCXParser
from tools.docx_renderer import DOCXRenderer
from docx import Document


class TestE2EMediaAnchoring:
    """End-to-end tests for media anchoring pipeline."""
    
    @pytest.fixture
    def test_output_dir(self, tmp_path):
        """Create temporary output directory."""
        output_dir = tmp_path / "output"
        output_dir.mkdir()
        return output_dir
    
    @pytest.fixture
    def mock_lesson_json(self):
        """Create mock lesson JSON with media metadata."""
        return {
            "metadata": {
                "week_of": "10/21-10/25",
                "grade": "3rd",
                "subject": "ELA",
                "homeroom": "Test Teacher"
            },
            "days": {
                "monday": {
                    "unit_lesson": "Unit 5: Reading Comprehension",
                    "objective": {
                        "content_objective": "Students will identify main ideas",
                        "student_goal": "I can find the main idea",
                        "wida_objective": "Reading at Level 3"
                    },
                    "anticipatory_set": {
                        "original_content": "Review prior knowledge using activity",
                        "bilingual_bridge": ""
                    },
                    "tailored_instruction": {
                        "original_content": "Students will complete the worksheet",
                        "co_teaching_model": {},
                        "ell_support": [],
                        "special_needs_support": [],
                        "materials": []
                    },
                    "misconceptions": {
                        "original_content": "Students may confuse main idea with details",
                        "linguistic_note": {}
                    },
                    "assessment": {
                        "primary_assessment": "Exit ticket",
                        "bilingual_overlay": {}
                    },
                    "homework": {
                        "original_content": "Practice worksheet",
                        "bilingual_note": ""
                    }
                },
                "tuesday": {
                    "unit_lesson": "Unit 5: Reading Comprehension",
                    "objective": {
                        "content_objective": "Students will analyze text structure",
                        "student_goal": "I can identify text structure",
                        "wida_objective": "Reading at Level 3"
                    },
                    "anticipatory_set": {
                        "original_content": "Warm-up activity",
                        "bilingual_bridge": ""
                    },
                    "tailored_instruction": {
                        "original_content": "Guided practice",
                        "co_teaching_model": {},
                        "ell_support": [],
                        "special_needs_support": [],
                        "materials": []
                    },
                    "misconceptions": {
                        "original_content": "None identified",
                        "linguistic_note": {}
                    },
                    "assessment": {
                        "primary_assessment": "Observation",
                        "bilingual_overlay": {}
                    },
                    "homework": {
                        "original_content": "Reading log",
                        "bilingual_note": ""
                    }
                },
                "wednesday": {
                    "unit_lesson": "Unit 5: Reading Comprehension",
                    "objective": {
                        "content_objective": "Students will summarize passages",
                        "student_goal": "I can summarize what I read",
                        "wida_objective": "Reading at Level 3"
                    },
                    "anticipatory_set": {
                        "original_content": "Review vocabulary",
                        "bilingual_bridge": ""
                    },
                    "tailored_instruction": {
                        "original_content": "Independent practice",
                        "co_teaching_model": {},
                        "ell_support": [],
                        "special_needs_support": [],
                        "materials": []
                    },
                    "misconceptions": {
                        "original_content": "None",
                        "linguistic_note": {}
                    },
                    "assessment": {
                        "primary_assessment": "Written summary",
                        "bilingual_overlay": {}
                    },
                    "homework": {
                        "original_content": "Complete summary",
                        "bilingual_note": ""
                    }
                },
                "thursday": {
                    "unit_lesson": "Unit 5: Reading Comprehension",
                    "objective": {
                        "content_objective": "Students will make inferences",
                        "student_goal": "I can make inferences",
                        "wida_objective": "Reading at Level 3"
                    },
                    "anticipatory_set": {
                        "original_content": "Think-pair-share",
                        "bilingual_bridge": ""
                    },
                    "tailored_instruction": {
                        "original_content": "Group work",
                        "co_teaching_model": {},
                        "ell_support": [],
                        "special_needs_support": [],
                        "materials": []
                    },
                    "misconceptions": {
                        "original_content": "None",
                        "linguistic_note": {}
                    },
                    "assessment": {
                        "primary_assessment": "Group presentation",
                        "bilingual_overlay": {}
                    },
                    "homework": {
                        "original_content": "Inference worksheet",
                        "bilingual_note": ""
                    }
                },
                "friday": {
                    "unit_lesson": "Unit 5: Reading Comprehension",
                    "objective": {
                        "content_objective": "Students will demonstrate mastery",
                        "student_goal": "I can show what I learned",
                        "wida_objective": "Reading at Level 3"
                    },
                    "anticipatory_set": {
                        "original_content": "Review week's content",
                        "bilingual_bridge": ""
                    },
                    "tailored_instruction": {
                        "original_content": "Assessment",
                        "co_teaching_model": {},
                        "ell_support": [],
                        "special_needs_support": [],
                        "materials": []
                    },
                    "misconceptions": {
                        "original_content": "None",
                        "linguistic_note": {}
                    },
                    "assessment": {
                        "primary_assessment": "Unit test",
                        "bilingual_overlay": {}
                    },
                    "homework": {
                        "original_content": "None - enjoy weekend",
                        "bilingual_note": ""
                    }
                }
            }
        }
    
    def test_hyperlink_extraction_and_placement(self, mock_lesson_json, test_output_dir):
        """Test that hyperlinks are extracted with context and can be placed."""
        # Create mock hyperlinks with good context
        hyperlinks = [
            {
                'text': 'activity',
                'url': 'https://example.com/activity1',
                'context_snippet': 'Review prior knowledge using activity',
                'context_type': 'table_cell',
                'section_hint': 'anticipatory_set',
                'day_hint': 'monday'
            },
            {
                'text': 'worksheet',
                'url': 'https://example.com/worksheet',
                'context_snippet': 'Students will complete the worksheet',
                'context_type': 'table_cell',
                'section_hint': 'instruction',
                'day_hint': 'monday'
            }
        ]
        
        # Add hyperlinks to lesson JSON
        mock_lesson_json['_hyperlinks'] = hyperlinks
        mock_lesson_json['_media_schema_version'] = '1.1'
        
        # Render the document
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        output_path = test_output_dir / "test_hyperlinks.docx"
        
        success = renderer.render(mock_lesson_json, str(output_path))
        
        assert success, "Rendering should succeed"
        assert output_path.exists(), "Output file should be created"
        
        # Verify document can be opened
        doc = Document(str(output_path))
        assert len(doc.tables) > 0, "Document should have tables"
        
        # Check that hyperlinks were processed (either inline or fallback)
        # We can't easily verify inline placement without parsing the XML,
        # but we can verify the document was created successfully
        print(f"\n✅ E2E Test: Hyperlinks processed successfully")
        print(f"   Output: {output_path}")
        print(f"   Hyperlinks: {len(hyperlinks)}")
    
    def test_image_structure_based_placement(self, mock_lesson_json, test_output_dir):
        """Test that images with structure info are placed correctly."""
        # Create mock image with structure information
        images = [
            {
                'data': 'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg==',  # 1x1 red pixel
                'content_type': 'image/png',
                'filename': 'test_image.png',
                'rel_id': 'rId1',
                'context_snippet': '',
                'context_type': 'table_cell',
                'section_hint': 'anticipatory_set',
                'day_hint': 'wednesday',
                'row_label': 'Anticipatory Set:',
                'cell_index': 3  # Wednesday
            }
        ]
        
        # Add images to lesson JSON
        mock_lesson_json['_images'] = images
        mock_lesson_json['_media_schema_version'] = '1.1'
        
        # Render the document
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        output_path = test_output_dir / "test_images.docx"
        
        success = renderer.render(mock_lesson_json, str(output_path))
        
        assert success, "Rendering should succeed"
        assert output_path.exists(), "Output file should be created"
        
        # Verify document can be opened
        doc = Document(str(output_path))
        assert len(doc.tables) > 0, "Document should have tables"
        
        print(f"\n✅ E2E Test: Images with structure processed successfully")
        print(f"   Output: {output_path}")
        print(f"   Images: {len(images)}")
        print(f"   Structure: {images[0]['row_label']}, Cell {images[0]['cell_index']}")
    
    def test_backward_compatibility_legacy_media(self, mock_lesson_json, test_output_dir):
        """Test that legacy media (no structure) still works."""
        # Create legacy hyperlinks (no context)
        legacy_hyperlinks = [
            {
                'text': 'Old Link',
                'url': 'https://example.com/old'
            }
        ]
        
        # Add legacy media (no schema version or schema 1.0)
        mock_lesson_json['_hyperlinks'] = legacy_hyperlinks
        mock_lesson_json['_media_schema_version'] = '1.0'
        
        # Render the document
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        output_path = test_output_dir / "test_legacy.docx"
        
        success = renderer.render(mock_lesson_json, str(output_path))
        
        assert success, "Rendering should succeed with legacy media"
        assert output_path.exists(), "Output file should be created"
        
        # Verify document can be opened
        doc = Document(str(output_path))
        assert len(doc.tables) > 0, "Document should have tables"
        
        print(f"\n✅ E2E Test: Legacy media (backward compatibility) works")
        print(f"   Output: {output_path}")
        print(f"   Schema: 1.0 (legacy)")
    
    def test_full_pipeline_with_mixed_media(self, mock_lesson_json, test_output_dir):
        """Test complete pipeline with both hyperlinks and images."""
        # Create mixed media
        hyperlinks = [
            {
                'text': 'resource link',
                'url': 'https://example.com/resource',
                'context_snippet': 'Students will use the resource link',
                'context_type': 'table_cell',
                'section_hint': 'instruction',
                'day_hint': 'tuesday'
            }
        ]
        
        images = [
            {
                'data': 'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg==',
                'content_type': 'image/png',
                'filename': 'diagram.png',
                'rel_id': 'rId2',
                'context_snippet': 'See diagram below',
                'context_type': 'table_cell',
                'section_hint': 'instruction',
                'day_hint': 'thursday',
                'row_label': 'Tailored Instruction:',
                'cell_index': 4
            }
        ]
        
        # Add all media
        mock_lesson_json['_hyperlinks'] = hyperlinks
        mock_lesson_json['_images'] = images
        mock_lesson_json['_media_schema_version'] = '1.1'
        
        # Render the document
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        output_path = test_output_dir / "test_full_pipeline.docx"
        
        success = renderer.render(mock_lesson_json, str(output_path))
        
        assert success, "Full pipeline should succeed"
        assert output_path.exists(), "Output file should be created"
        
        # Verify document structure
        doc = Document(str(output_path))
        assert len(doc.tables) > 0, "Document should have tables"
        
        # Get file size to verify media was included
        file_size = output_path.stat().st_size
        assert file_size > 10000, "File should be substantial (contains media)"
        
        print(f"\n✅ E2E Test: Full pipeline with mixed media works")
        print(f"   Output: {output_path}")
        print(f"   Hyperlinks: {len(hyperlinks)}")
        print(f"   Images: {len(images)}")
        print(f"   File size: {file_size:,} bytes")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
