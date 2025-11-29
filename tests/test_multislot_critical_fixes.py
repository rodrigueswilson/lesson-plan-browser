"""
Test suite for critical multi-slot fixes.

Tests two critical issues that were fixed before production:
1. First slot clearing when i > 0 (skipped empty slots)
2. Image slot filtering to prevent cross-contamination
"""

import pytest
from pathlib import Path
from docx import Document
from tools.docx_renderer import DOCXRenderer


# Use the actual template file for testing
TEMPLATE_PATH = Path("d:/LP/input/Lesson Plan Template SY'25-26.docx")


class TestCriticalFixes:
    """Test critical fixes for production readiness."""
    
    def test_first_slot_clears_when_slot_0_empty(self):
        """
        CRITICAL FIX 1: Verify first written slot clears cell even when i > 0.
        
        Scenario: Slot 0 has no content, Slot 1 is first to write.
        Bug: append_mode=(i > 0) would be True, leaving template content.
        Fix: append_mode=written_any tracks actual writes, not loop index.
        """
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        renderer = DOCXRenderer(str(TEMPLATE_PATH))
        
        # Pre-populate cell with "old content" (simulates template content)
        table.rows[renderer.UNIT_LESSON_ROW].cells[1].text = "OLD TEMPLATE CONTENT"
        
        slots = [
            {
                'slot_number': 0,
                'subject': 'Empty Slot',
                # No unit_lesson - this slot will be skipped
            },
            {
                'slot_number': 1,
                'subject': 'ELA',
                'unit_lesson': 'Unit 1: Reading'
            },
            {
                'slot_number': 2,
                'subject': 'Math',
                'unit_lesson': 'LESSON 9: AREA'
            },
        ]
        
        renderer._fill_multi_slot_day(table, 1, slots, day_name='Monday')
        
        # Check Unit/Lesson row
        unit_cell = table.rows[renderer.UNIT_LESSON_ROW].cells[1]
        unit_text = unit_cell.text
        
        # CRITICAL: Old template content should be CLEARED
        assert 'OLD TEMPLATE CONTENT' not in unit_text, \
            "First written slot (i=1) should clear cell, not append to template content"
        
        # Verify correct content is present
        assert 'Slot 1: ELA' in unit_text
        assert 'Unit 1: Reading' in unit_text
        assert 'Slot 2: Math' in unit_text
        assert 'LESSON 9: AREA' in unit_text
    
    def test_first_slot_clears_when_multiple_empty_slots(self):
        """
        Verify first written slot clears even when multiple empty slots precede it.
        """
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        renderer = DOCXRenderer(str(TEMPLATE_PATH))
        
        # Pre-populate cell
        table.rows[renderer.OBJECTIVE_ROW].cells[1].text = "OLD OBJECTIVE CONTENT"
        
        slots = [
            {'slot_number': 0, 'subject': 'Empty 1'},  # No content
            {'slot_number': 1, 'subject': 'Empty 2'},  # No content
            {'slot_number': 2, 'subject': 'Empty 3'},  # No content
            {
                'slot_number': 3,
                'subject': 'Science',
                'objective': {'content_objective': 'Students will observe.'}
            },
        ]
        
        renderer._fill_multi_slot_day(table, 1, slots, day_name='Tuesday')
        
        # Check Objective row
        obj_cell = table.rows[renderer.OBJECTIVE_ROW].cells[1]
        obj_text = obj_cell.text
        
        # CRITICAL: Old content should be cleared
        assert 'OLD OBJECTIVE CONTENT' not in obj_text, \
            "First written slot (i=3) should clear cell"
        
        # Verify correct content
        assert 'Slot 3: Science' in obj_text
        assert 'Students will observe' in obj_text
    
    def test_image_slot_filtering_prevents_cross_contamination(self):
        """
        CRITICAL FIX 2: Verify images are filtered by _source_slot.
        
        Bug: Images weren't checked for _source_slot, allowing cross-contamination.
        Fix: Added same slot/subject filtering as hyperlinks.
        """
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        renderer = DOCXRenderer(str(TEMPLATE_PATH))
        
        # Create images for different slots
        images = [
            {
                'data': 'base64_ela_image',
                'caption': 'ELA Image',
                'filename': 'ela_image.jpg',
                '_source_slot': 1,
                '_source_subject': 'ELA',
                'section_hint': 'anticipatory_set',
                'day_hint': 'Monday'
            },
            {
                'data': 'base64_math_image',
                'caption': 'Math Image',
                'filename': 'math_image.jpg',
                '_source_slot': 2,
                '_source_subject': 'Math',
                'section_hint': 'anticipatory_set',
                'day_hint': 'Monday'
            },
        ]
        
        # Render Slot 1 (ELA) - should only see ELA image
        renderer._fill_cell(
            table, renderer.ANTICIPATORY_SET_ROW, 1,
            "ELA anticipatory content",
            pending_images=images,
            current_slot_number=1,
            current_subject='ELA',
            day_name='Monday',
            section_name='anticipatory_set'
        )
        
        # CRITICAL: Math image should still be in list (not grabbed by ELA slot)
        assert len(images) >= 1, "Math image should not be removed by ELA slot"
        
        # Verify Math image is still present
        math_images = [img for img in images if img.get('_source_slot') == 2]
        assert len(math_images) == 1, "Math image should still be in pending list"
        assert math_images[0]['filename'] == 'math_image.jpg'
    
    def test_image_subject_filtering(self):
        """
        Verify images are also filtered by _source_subject for extra safety.
        """
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        renderer = DOCXRenderer(str(TEMPLATE_PATH))
        
        images = [
            {
                'data': 'img1',
                'caption': 'Science Image',
                'filename': 'science.jpg',
                '_source_slot': 1,
                '_source_subject': 'Science',
                'section_hint': 'instruction',
                'day_hint': 'Wednesday'
            },
            {
                'data': 'img2',
                'caption': 'History Image',
                'filename': 'history.jpg',
                '_source_slot': 1,  # Same slot
                '_source_subject': 'History',  # Different subject
                'section_hint': 'instruction',
                'day_hint': 'Wednesday'
            },
        ]
        
        # Render with Science subject
        renderer._fill_cell(
            table, renderer.INSTRUCTION_ROW, 1,
            "Science instruction content",
            pending_images=images,
            current_slot_number=1,
            current_subject='Science',
            day_name='Wednesday',
            section_name='instruction'
        )
        
        # History image should not be grabbed even though slot matches
        history_images = [img for img in images if img.get('_source_subject') == 'History']
        assert len(history_images) == 1, "History image should not be removed by Science subject"
    
    def test_combined_slot_and_subject_filtering(self):
        """
        Verify both slot and subject filtering work together.
        """
        doc = Document()
        table = doc.add_table(rows=8, cols=2)
        renderer = DOCXRenderer(str(TEMPLATE_PATH))
        
        hyperlinks = [
            {'text': 'ELA Link', 'url': 'http://ela.com', '_source_slot': 1, '_source_subject': 'ELA'},
            {'text': 'Math Link', 'url': 'http://math.com', '_source_slot': 2, '_source_subject': 'Math'},
        ]
        
        images = [
            {
                'data': 'img1', 'caption': 'ELA Image', 'filename': 'ela.jpg',
                '_source_slot': 1, '_source_subject': 'ELA',
                'section_hint': 'assessment', 'day_hint': 'Thursday'
            },
            {
                'data': 'img2', 'caption': 'Math Image', 'filename': 'math.jpg',
                '_source_slot': 2, '_source_subject': 'Math',
                'section_hint': 'assessment', 'day_hint': 'Thursday'
            },
        ]
        
        # Render Slot 1 (ELA)
        renderer._fill_cell(
            table, renderer.ASSESSMENT_ROW, 1,
            "ELA Link assessment content",
            pending_hyperlinks=hyperlinks,
            pending_images=images,
            current_slot_number=1,
            current_subject='ELA',
            day_name='Thursday',
            section_name='assessment'
        )
        
        # Verify Math items remain
        math_links = [h for h in hyperlinks if h.get('_source_slot') == 2]
        math_imgs = [i for i in images if i.get('_source_slot') == 2]
        
        assert len(math_links) == 1, "Math hyperlink should not be removed"
        assert len(math_imgs) == 1, "Math image should not be removed"


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
