"""
Tests for DOCX style normalization utilities (normalize_styles_via_file, etc.).
Session 8: combined-original styles.
"""
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from lxml import etree

from tools.docx_utils import diagnose_style_conflicts, normalize_styles_via_file


def test_normalize_styles_via_file_returns_valid_stream():
    """normalize_styles_via_file returns a non-empty BytesIO and result opens as a doc with styles."""
    master = Document()
    master.add_paragraph("Master")
    target = Document()
    target.add_paragraph("Target")

    stream = normalize_styles_via_file(master, target)
    assert stream is not None
    assert isinstance(stream, BytesIO)
    assert len(stream.getvalue()) > 0

    stream.seek(0)
    doc = Document(stream)
    assert len(doc.styles) >= 1


def test_diagnose_style_conflicts_returns_dict():
    """diagnose_style_conflicts returns a dict with expected keys."""
    master = Document()
    target = Document()
    diagnosis = diagnose_style_conflicts(master, target)
    assert isinstance(diagnosis, dict)
    assert "style_counts" in diagnosis
    assert "conflicting_styles" in diagnosis
    assert diagnosis["style_counts"]["master"] >= 0
    assert diagnosis["style_counts"]["target"] >= 0


def test_normalize_styles_via_file_repairs_styles_without_names(tmp_path):
    """Malformed template table styles get names before being copied to output."""
    master = Document()
    master.add_paragraph("Master")
    target = Document()
    target.add_paragraph("Target")

    master_path = tmp_path / "master.docx"
    target_path = tmp_path / "target.docx"
    master.save(master_path)
    target.save(target_path)

    _add_unnamed_table_style(master_path, "a1")

    repaired_master = Document(master_path)
    target_doc = Document(target_path)
    stream = normalize_styles_via_file(repaired_master, target_doc)

    assert stream is not None
    stream.seek(0)
    styles_xml = zipfile.ZipFile(stream).read("word/styles.xml")
    root = etree.fromstring(styles_xml)
    style = _style_by_id(root, "a1")

    assert style is not None
    assert style.find(_w_tag("name")).get(_w_tag("val")) == "a1"


def test_normalize_styles_via_file_uniquifies_reused_hyperlink_relationships(
    tmp_path,
):
    """Repeated hyperlink rIds are expanded before Word has to repair them."""
    master = Document()
    master.add_paragraph("Master")
    target = Document()
    target.add_paragraph("Target")

    target_path = tmp_path / "target.docx"
    target.save(target_path)
    _add_reused_hyperlinks(target_path)

    stream = normalize_styles_via_file(master, Document(target_path))

    assert stream is not None
    stream.seek(0)
    doc_xml = zipfile.ZipFile(stream).read("word/document.xml")
    rels_xml = zipfile.ZipFile(stream).read("word/_rels/document.xml.rels")
    doc_root = etree.fromstring(doc_xml)
    rels_root = etree.fromstring(rels_xml)
    rel_ids = [
        hyperlink.get(_r_tag("id"))
        for hyperlink in doc_root.findall(f".//{_w_tag('hyperlink')}")
    ]
    hyperlink_relationships = [
        rel
        for rel in rels_root.findall(_rel_tag("Relationship"))
        if rel.get("Type")
        == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    ]

    assert len(rel_ids) == 2
    assert len(set(rel_ids)) == 2
    assert len(hyperlink_relationships) == 2


def _add_unnamed_table_style(docx_path: Path, style_id: str) -> None:
    original = docx_path.read_bytes()
    with zipfile.ZipFile(BytesIO(original), "r") as source_zip:
        styles_xml = source_zip.read("word/styles.xml")
        styles_root = etree.fromstring(styles_xml)

        existing = _style_by_id(styles_root, style_id)
        if existing is not None:
            styles_root.remove(existing)

        style = etree.Element(_w_tag("style"))
        style.set(_w_tag("type"), "table")
        style.set(_w_tag("styleId"), style_id)

        based_on = etree.Element(_w_tag("basedOn"))
        based_on.set(_w_tag("val"), "TableNormal")
        style.append(based_on)

        tbl_pr = etree.Element(_w_tag("tblPr"))
        style.append(tbl_pr)
        styles_root.append(style)

        repaired_styles = etree.tostring(
            styles_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

        output = BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as output_zip:
            for item in source_zip.infolist():
                data = (
                    repaired_styles
                    if item.filename == "word/styles.xml"
                    else source_zip.read(item.filename)
                )
                output_zip.writestr(item, data)

    docx_path.write_bytes(output.getvalue())


def _add_reused_hyperlinks(docx_path: Path) -> None:
    original = docx_path.read_bytes()
    with zipfile.ZipFile(BytesIO(original), "r") as source_zip:
        document_root = etree.fromstring(source_zip.read("word/document.xml"))
        rels_root = etree.fromstring(source_zip.read("word/_rels/document.xml.rels"))

        body = document_root.find(_w_tag("body"))
        sect_pr = body.find(_w_tag("sectPr"))
        insert_at = list(body).index(sect_pr) if sect_pr is not None else len(body)
        for index in range(2):
            body.insert(insert_at + index, _hyperlink_paragraph("rId50", f"Link {index}"))

        relationship = etree.Element(_rel_tag("Relationship"))
        relationship.set("Id", "rId50")
        relationship.set(
            "Type",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        )
        relationship.set("Target", "https://example.com")
        relationship.set("TargetMode", "External")
        rels_root.append(relationship)

        updated_document = etree.tostring(
            document_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )
        updated_rels = etree.tostring(
            rels_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

        output = BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as output_zip:
            for item in source_zip.infolist():
                if item.filename == "word/document.xml":
                    data = updated_document
                elif item.filename == "word/_rels/document.xml.rels":
                    data = updated_rels
                else:
                    data = source_zip.read(item.filename)
                output_zip.writestr(item, data)

    docx_path.write_bytes(output.getvalue())


def _hyperlink_paragraph(rel_id: str, text: str):
    paragraph = etree.Element(_w_tag("p"))
    hyperlink = etree.SubElement(paragraph, _w_tag("hyperlink"))
    hyperlink.set(_r_tag("id"), rel_id)
    run = etree.SubElement(hyperlink, _w_tag("r"))
    text_element = etree.SubElement(run, _w_tag("t"))
    text_element.text = text
    return paragraph


def _style_by_id(root, style_id: str):
    for style in root.findall(_w_tag("style")):
        if style.get(_w_tag("styleId")) == style_id:
            return style
    return None


def _w_tag(local_name: str) -> str:
    return f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{local_name}"


def _r_tag(local_name: str) -> str:
    return f"{{http://schemas.openxmlformats.org/officeDocument/2006/relationships}}{local_name}"


def _rel_tag(local_name: str) -> str:
    return f"{{http://schemas.openxmlformats.org/package/2006/relationships}}{local_name}"
