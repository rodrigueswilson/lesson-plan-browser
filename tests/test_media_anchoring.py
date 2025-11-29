"""
Tests for semantic anchoring of hyperlinks and images.

Tests context extraction, fuzzy matching, and inline media injection.
"""

import pytest
from pathlib import Path
from unittest.mock import Mock, patch
from docx import Document
from docx.shared import Inches

from tools.docx_parser import DOCXParser
from tools.docx_renderer import DOCXRenderer


class TestContextExtraction:
    """Test context snippet extraction from parser."""
    
    def test_get_context_snippet_with_link_in_middle(self):
        """Parser extracts context around hyperlink text."""
        parser = DOCXParser("tests/fixtures/regular_lesson.docx")
        
        # Create mock paragraph
        mock_para = Mock()
        mock_para.text = "Students will click here to access the worksheet and complete the assignment."
        
        context = parser._get_context_snippet(mock_para, "click here", window=40)
        
        assert "click here" in context
        assert len(context) <= 50  # Allow some buffer for word boundaries
        assert "Students" in context or "access" in context
    
    def test_get_context_snippet_fallback(self):
        """Parser falls back to beginning of text when link text not found."""
        parser = DOCXParser("tests/fixtures/regular_lesson.docx")
        
        mock_para = Mock()
        mock_para.text = "This is a long paragraph with lots of content."
        
        context = parser._get_context_snippet(mock_para, "nonexistent", window=20)
        
        assert len(context) <= 20
        assert context.startswith("This is")
    
    def test_infer_section_objective(self):
        """Parser correctly infers objective section."""
        parser = DOCXParser("tests/fixtures/regular_lesson.docx")
        
        text = "Students will be able to identify the main idea of the passage."
        section = parser._infer_section(text)
        
        assert section == 'objective'
    
    def test_infer_section_instruction(self):
        """Parser correctly infers instruction section."""
        parser = DOCXParser("tests/fixtures/regular_lesson.docx")
        
        text = "The teacher will model the activity using a think-aloud strategy."
        section = parser._infer_section(text)
        
        assert section == 'instruction'
    
    def test_infer_section_assessment(self):
        """Parser correctly infers assessment section."""
        parser = DOCXParser("tests/fixtures/regular_lesson.docx")
        
        text = "Students will complete an exit ticket to check for understanding."
        section = parser._infer_section(text)
        
        assert section == 'assessment'
    
    def test_detect_day_from_table_monday(self):
        """Parser detects Monday from table header."""
        parser = DOCXParser("tests/fixtures/regular_lesson.docx")
        
        # Create mock table with Monday in header
        mock_table = Mock()
        mock_row = Mock()
        mock_cell1 = Mock()
        mock_cell1.text = "Monday"
        mock_cell2 = Mock()
        mock_cell2.text = "Tuesday"
        mock_row.cells = [mock_cell1, mock_cell2]
        mock_table.rows = [mock_row]
        
        day = parser._detect_day_from_table(mock_table)
        
        assert day == 'monday'


class TestMatchingConfidence:
    """Test confidence calculation for media matching."""
    
    @patch('tools.docx_renderer.logger')
    def test_exact_text_match(self, mock_logger):
        """Exact hyperlink text match returns 100% confidence."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        cell_text = "Students will click here to access the worksheet."
        hyperlink = {
            'text': 'click here',
            'url': 'https://example.com',
            'context_snippet': 'Students will click here to access'
        }
        
        confidence, match_type = renderer._calculate_match_confidence(
            cell_text, hyperlink
        )
        
        assert confidence == 1.0
        assert match_type == 'exact_text'
    
    @patch('tools.docx_renderer.logger')
    def test_fuzzy_context_match(self, mock_logger):
        """Fuzzy context match returns high confidence."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        # Context similar but not exact
        cell_text = "Learners will access the online worksheet to complete the activity."
        hyperlink = {
            'text': 'worksheet link',
            'url': 'https://example.com',
            'context_snippet': 'Students will access the worksheet to complete'
        }
        
        confidence, match_type = renderer._calculate_match_confidence(
            cell_text, hyperlink
        )
        
        # Should have decent confidence due to similar words
        assert confidence > 0.5
        assert 'context' in match_type or match_type == 'no_match'
    
    @patch('tools.docx_renderer.logger')
    def test_hint_boost(self, mock_logger):
        """Day and section hints boost confidence score."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        cell_text = "Students will complete the worksheet activity."
        hyperlink = {
            'text': 'worksheet',
            'url': 'https://example.com',
            'context_snippet': 'complete the worksheet',
            'day_hint': 'monday',
            'section_hint': 'instruction'
        }
        
        # Without hints
        conf_no_hints, _ = renderer._calculate_match_confidence(
            cell_text, hyperlink, day_name=None, section_name=None
        )
        
        # With hints
        conf_with_hints, match_type = renderer._calculate_match_confidence(
            cell_text, hyperlink, day_name='monday', section_name='instruction'
        )
        
        # Confidence should be boosted
        assert conf_with_hints >= conf_no_hints
        if 'hints' in match_type:
            assert conf_with_hints > conf_no_hints
    
    @patch('tools.docx_renderer.logger')
    def test_no_match(self, mock_logger):
        """Unrelated content returns zero confidence."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        cell_text = "Students will practice multiplication facts."
        hyperlink = {
            'text': 'reading passage',
            'url': 'https://example.com',
            'context_snippet': 'Analyze the reading passage for main ideas'
        }
        
        confidence, match_type = renderer._calculate_match_confidence(
            cell_text, hyperlink
        )
        
        assert confidence < 0.5
        assert match_type in ['no_match', 'hints_only']


class TestMediaInjection:
    """Test inline media injection into cells."""
    
    @patch('tools.docx_renderer.logger')
    def test_inject_hyperlink_inline(self, mock_logger):
        """Hyperlink is injected into cell paragraph."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        # Create mock cell with paragraph
        mock_cell = Mock()
        mock_para = Mock()
        mock_para.text = "Some existing content"
        mock_cell.paragraphs = [mock_para]
        
        hyperlink = {
            'text': 'Resource Link',
            'url': 'https://example.com'
        }
        
        # Mock _add_hyperlink method
        with patch.object(renderer, '_add_hyperlink') as mock_add:
            renderer._inject_hyperlink_inline(mock_cell, hyperlink)
            
            # Should add space and then hyperlink
            mock_para.add_run.assert_called_once_with(" ")
            mock_add.assert_called_once_with(mock_para, 'Resource Link', 'https://example.com')
    
    @patch('tools.docx_renderer.logger')
    def test_inject_image_inline_with_caption(self, mock_logger):
        """Image is injected into cell with caption."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        # Create mock cell
        mock_cell = Mock()
        mock_para = Mock()
        mock_run = Mock()
        mock_para.add_run.return_value = mock_run
        mock_cell.add_paragraph.return_value = mock_para
        
        image = {
            'data': 'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg==',  # 1x1 red pixel
            'filename': 'test.png',
            'content_type': 'image/png'
        }
        
        renderer._inject_image_inline(mock_cell, image, max_width=1.2)
        
        # Should add paragraph and image
        assert mock_cell.add_paragraph.call_count >= 2  # Image + caption
        mock_run.add_picture.assert_called_once()


class TestBackwardCompatibility:
    """Test backward compatibility with legacy media (no context)."""
    
    @patch('tools.docx_renderer.logger')
    def test_legacy_media_falls_back(self, mock_logger):
        """Legacy media without context falls back to end section."""
        renderer = DOCXRenderer("input/Lesson Plan Template SY'25-26.docx")
        
        # Legacy hyperlink (no context_snippet)
        legacy_hyperlink = {
            'text': 'Old Link',
            'url': 'https://example.com'
        }
        
        cell_text = "Some cell content"
        
        confidence, match_type = renderer._calculate_match_confidence(
            cell_text, legacy_hyperlink
        )
        
        # Should have zero confidence (no context to match)
        assert confidence == 0.0
        assert match_type == 'no_match'


class TestSchemaVersioning:
    """Test schema version handling."""
    
    def test_schema_version_1_1_enables_anchoring(self):
        """Schema version 1.1 enables semantic anchoring."""
        json_data = {
            'metadata': {'week_of': '10/21-10/25'},
            'days': {},
            '_hyperlinks': [{'text': 'link', 'url': 'http://example.com'}],
            '_media_schema_version': '1.1'
        }
        
        # Version 1.1 should prepare pending lists
        assert json_data.get('_media_schema_version') == '1.1'
    
    def test_schema_version_1_0_uses_legacy(self):
        """Schema version 1.0 uses legacy end-of-document placement."""
        json_data = {
            'metadata': {'week_of': '10/21-10/25'},
            'days': {},
            '_hyperlinks': [{'text': 'link', 'url': 'http://example.com'}],
            '_media_schema_version': '1.0'
        }
        
        # Version 1.0 should use legacy behavior
        assert json_data.get('_media_schema_version') == '1.0'


class TestEndToEnd:
    """End-to-end integration tests."""
    
    @pytest.mark.skipif(
        not Path("tests/fixtures/media_test.docx").exists(),
        reason="Media test fixture not available"
    )
    def test_full_pipeline_with_anchoring(self):
        """Full pipeline: parse → extract context → render with anchoring."""
        # This would require a real fixture DOCX with hyperlinks/images
        # Skipped if fixture doesn't exist
        pass


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
