"""
Validation script for python-docx table width API.
Tests different approaches to set equal column widths.
"""

from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path
import sys

def test_column_width_approach():
    """Test setting column widths using table.columns API."""
    
    # Create a simple test document
    doc = Document()
    
    # Add a table with 3 columns and 2 rows
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Add some content
    table.cell(0, 0).text = "Column 1"
    table.cell(0, 1).text = "Column 2"
    table.cell(0, 2).text = "Column 3"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    table.cell(1, 2).text = "Data 3"
    
    # Approach 1: Set width on each column
    print("Testing Approach 1: table.columns[i].width")
    try:
        target_width = Inches(2.0)
        for i, column in enumerate(table.columns):
            column.width = target_width
            print(f"  Column {i}: width set to {target_width}")
        
        # Verify
        for i, column in enumerate(table.columns):
            print(f"  Column {i}: width = {column.width}")
        
        # Save test file
        output_path = Path('output/test_table_width_approach1.docx')
        output_path.parent.mkdir(exist_ok=True)
        doc.save(str(output_path))
        print(f"  Saved to: {output_path}")
        print("  SUCCESS: Approach 1 works!")
        return True
        
    except Exception as e:
        print(f"  FAILED: {type(e).__name__}: {e}")
        return False

def test_cell_width_approach():
    """Test setting width on individual cells."""
    
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Add content
    table.cell(0, 0).text = "Column 1"
    table.cell(0, 1).text = "Column 2"
    table.cell(0, 2).text = "Column 3"
    
    print("\nTesting Approach 2: cell.width")
    try:
        target_width = Inches(2.0)
        
        # Set width on first row cells
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = target_width
            print(f"  Cell {i}: width set to {target_width}")
        
        # Verify
        for i, cell in enumerate(table.rows[0].cells):
            print(f"  Cell {i}: width = {cell.width}")
        
        # Save test file
        output_path = Path('output/test_table_width_approach2.docx')
        output_path.parent.mkdir(exist_ok=True)
        doc.save(str(output_path))
        print(f"  Saved to: {output_path}")
        print("  SUCCESS: Approach 2 works!")
        return True
        
    except Exception as e:
        print(f"  FAILED: {type(e).__name__}: {e}")
        return False

def test_with_existing_template():
    """Test with the actual lesson plan template."""
    
    template_path = Path('input/Lesson Plan Template SY\'25-26.docx')
    
    if not template_path.exists():
        print(f"\nSkipping template test: {template_path} not found")
        return None
    
    print(f"\nTesting with actual template: {template_path}")
    
    try:
        doc = Document(str(template_path))
        
        if not doc.tables:
            print("  No tables found in template")
            return None
        
        print(f"  Found {len(doc.tables)} table(s)")
        
        # Test on first table
        table = doc.tables[0]
        print(f"  Table has {len(table.columns)} columns")
        
        # Calculate equal width
        total_width = Inches(6.5)  # Standard page width minus margins
        col_width = int(total_width / len(table.columns))  # Must be int!
        
        print(f"  Setting each column to {col_width}")
        
        for i, column in enumerate(table.columns):
            original_width = column.width
            column.width = col_width
            new_width = column.width
            print(f"    Column {i}: {original_width} -> {new_width}")
        
        # Save
        output_path = Path('output/test_template_equal_widths.docx')
        output_path.parent.mkdir(exist_ok=True)
        doc.save(str(output_path))
        print(f"  Saved to: {output_path}")
        print("  SUCCESS: Template test works!")
        return True
        
    except Exception as e:
        print(f"  FAILED: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    print("=" * 60)
    print("PYTHON-DOCX TABLE WIDTH VALIDATION")
    print("=" * 60)
    
    results = []
    
    # Test approach 1
    results.append(("Approach 1 (column.width)", test_column_width_approach()))
    
    # Test approach 2
    results.append(("Approach 2 (cell.width)", test_cell_width_approach()))
    
    # Test with template
    template_result = test_with_existing_template()
    if template_result is not None:
        results.append(("Template test", template_result))
    
    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    
    for name, result in results:
        status = "PASS" if result else "FAIL"
        print(f"{status}: {name}")
    
    # Exit code
    if all(r for r in results if r is not None):
        print("\nAll tests PASSED!")
        sys.exit(0)
    else:
        print("\nSome tests FAILED!")
        sys.exit(1)
